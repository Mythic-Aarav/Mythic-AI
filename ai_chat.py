"""
Mythic AI — single file, powered by Groq (primary) with Cerebras as a silent
automatic fallback. No provider selection is exposed to the user — if Groq is
rate-limited, times out, or errors, the app transparently retries on Cerebras.

Usage:
    1. pip install flask requests
    2. Set your API keys:
         Mac/Linux:   export GROQ_API_KEY="your-groq-key"
                      export CEREBRAS_API_KEY="your-cerebras-key"
         Windows:     set GROQ_API_KEY=your-groq-key
                      set CEREBRAS_API_KEY=your-cerebras-key
    3. python ai_chat.py
    4. Open http://localhost:5000 in your browser

Get a FREE Groq API key at https://console.groq.com/keys
Get a FREE Cerebras API key at https://cloud.cerebras.ai

Optional — NanoBanana (nanobananaapi.ai) powers real image-to-image editing for
"Ghibli Me"; without it, image generation falls back to HuggingFace FLUX
(text-to-image only). Weather uses Open-Meteo, which needs no API key at all.

Optional — "Paste URL" book/document support:
    pip install beautifulsoup4          (readable text extraction from ordinary webpages)
    pip install playwright pytesseract  (OCR reading of flipbook-style viewers, e.g.
                                          FlippingBook/Issuu/Yumpu/mmdigital-style sites)
    playwright install --with-deps chromium
    Also install the tesseract-ocr system package (e.g. `apt-get install tesseract-ocr`
    in your Render build). Without these, direct PDF/DOCX/TXT links and ordinary
    webpages still work fine — only flipbook OCR is disabled, with a clear message
    telling the user what's missing.

Supabase (optional — for accounts/conversation storage across restarts & devices):
    Set these as environment variables (never hardcode secrets in this file):
         SUPABASE_URL   e.g. https://xxxxx.supabase.co
         SUPABASE_KEY   your Supabase *secret* key (server-side only, keeps full DB access)
    If unset, the app falls back to storing conversations as local JSON files in chat_data/.

Features:
- Login/register (real accounts, hashed passwords, stored in chat_data/users.json)
- Multi-conversation chat with sidebar, saved per-account, survives restarts
- File/image upload (attach an image or text file to a message)
- Streaming responses (text appears word-by-word)
- Groq primary / Cerebras automatic silent fallback — no provider picker, ever
- Optional per-user "bring your own API key" override (Settings) so a person
  can use their own Groq/Cerebras key instead of the server's
- Image generation, Ghibli Me (image-to-image), and full weather (current +
  hourly + 7-day + air quality) built in
- Generate downloadable files (PDF / Word / text) straight from a chat reply
- Daily chat streaks + re-engagement push notifications ("come back and chat",
  study reminders, activity nudges, streak-on-hold alerts, feature updates)
- No rate limiting — unlimited messages
"""

import os
import re
import json
import uuid
import time
import secrets
import base64
import random
import hashlib
import datetime
import threading
import urllib.parse
import requests
import smtplib
from email.mime.text import MIMEText
from email.utils import formataddr
try:
    from pywebpush import webpush, WebPushException
    _PUSH_AVAILABLE = True
except ImportError:
    _PUSH_AVAILABLE = False
try:
    from PIL import Image, ImageFilter, ImageStat, ImageDraw, ImageOps, ImageSequence
    _WATERMARK_AVAILABLE = True
except ImportError:
    _WATERMARK_AVAILABLE = False
# bs4 — used to pull clean readable text out of ordinary webpages (the
# "Paste URL" box, when the link isn't a direct PDF/DOCX/TXT download).
try:
    from bs4 import BeautifulSoup
    _BS4_AVAILABLE = True
except ImportError:
    _BS4_AVAILABLE = False
# playwright + pytesseract — power OCR extraction for "flipbook" style book
# viewers (FlippingBook, Issuu, Yumpu, Calameo, AnyFlip, and similar), which
# render pages as images/canvas rather than sending any real text. Both are
# optional: if either is missing, flipbook OCR is silently disabled and the
# app falls back to a clear "can't read this" message instead of crashing.
# Install with:
#   pip install playwright pytesseract
#   playwright install --with-deps chromium
#   apt-get install -y tesseract-ocr   (or the equivalent on your host/Render build)
try:
    from playwright.sync_api import sync_playwright
    _PLAYWRIGHT_AVAILABLE = True
except ImportError:
    _PLAYWRIGHT_AVAILABLE = False
try:
    import pytesseract
    _OCR_AVAILABLE = True
except ImportError:
    _OCR_AVAILABLE = False
_FLIPBOOK_OCR_AVAILABLE = _PLAYWRIGHT_AVAILABLE and _OCR_AVAILABLE and _WATERMARK_AVAILABLE  # needs PIL too
from flask import (
    Flask, request, jsonify, Response, session,
    stream_with_context, redirect
)
from werkzeug.security import generate_password_hash, check_password_hash

PROVIDER = os.environ.get("AI_PROVIDER", "auto").strip().lower()
# "auto"  = Groq first, silently falls back to Cerebras on any failure (rate limit,
#           timeout, invalid model, network error, 429/500/503, etc.)
# "groq"     = Groq only
# "cerebras" = Cerebras only

# --- API Keys (hardcoded fallbacks — override via environment variables) ------
# WARNING: don't commit a file with real keys to a public GitHub repo.
# Set these as environment variables on Render instead.
GROQ_API_KEY      = os.environ.get("GROQ_API_KEY",      "")
CEREBRAS_API_KEY  = os.environ.get("CEREBRAS_API_KEY",  "")
# --- Google Sign-In (OAuth 2.0) ------------------------------------------
# Create these at https://console.cloud.google.com/apis/credentials
# (OAuth client ID, type "Web application"). Add BOTH of your live domains'
# callback URLs as Authorized redirect URIs, e.g.:
#   https://mythic-ai.vercel.app/api/auth/google/callback
#   https://<your-render-app>.onrender.com/api/auth/google/callback
GOOGLE_CLIENT_ID     = os.environ.get("GOOGLE_CLIENT_ID", "")
GOOGLE_CLIENT_SECRET = os.environ.get("GOOGLE_CLIENT_SECRET", "")

# --- Email sign-in (one-time code / OTP, no passwords) --------------------
# Any SMTP provider works — e.g. Gmail (use an "App password", not your
# normal password), SendGrid, Mailgun, Amazon SES, Resend, Postmark, etc.
#   SMTP_HOST      e.g. smtp.gmail.com  /  smtp.sendgrid.net
#   SMTP_PORT      e.g. 587 (STARTTLS) — this is the default if unset
#   SMTP_USER      login username for that SMTP account
#   SMTP_PASSWORD  password / app password / API key for that SMTP account
#   SMTP_FROM      the "From" address shown to recipients (defaults to SMTP_USER)
SMTP_HOST     = os.environ.get("SMTP_HOST", "")
SMTP_PORT     = int(os.environ.get("SMTP_PORT", "587") or "587")
SMTP_USER     = os.environ.get("SMTP_USER", "")
SMTP_PASSWORD = os.environ.get("SMTP_PASSWORD", "")
SMTP_FROM     = os.environ.get("SMTP_FROM", "") or SMTP_USER
SMTP_FROM_NAME = os.environ.get("SMTP_FROM_NAME", "Mythic AI")  # display name shown in the inbox
# HF is kept ONLY as a text-to-image fallback for /api/generate-image when
# NanoBanana isn't configured — it is NOT used as a chat/text provider.
HF_API_KEY        = os.environ.get("HF_API_KEY",        "")
# NanoBanana API (nanobananaapi.ai) — powers "Ghibli Me" image editing so it can
# actually transform the user's uploaded photo (image-to-image), not just
# generate a generic image from text. Get a key at https://nanobananaapi.ai/api-key
# and set it as an environment variable — never hardcode it here.
NANO_BANANA_API_KEY = os.environ.get("NANO_BANANA_API_KEY", "")
NANO_BANANA_BASE     = "https://api.nanobananaapi.ai/api/v1/nanobanana"

# Judge0 (code execution engine) — powers the multi-language Run button in
# Code Workspace (Python, C++, C, Java, Node.js, TypeScript, Go, Ruby).
# By default this calls the free public https://ce.judge0.com demo instance,
# which needs no key at all but is rate-limited and best-effort.
# For a reliable/production setup, get a free RapidAPI key at
# https://rapidapi.com/judge0-official/api/judge0-ce and set JUDGE0_API_KEY
# as an environment variable — the app will then route requests through
# RapidAPI's judge0-ce host instead.
JUDGE0_API_KEY  = os.environ.get("JUDGE0_API_KEY", "").strip()
JUDGE0_API_HOST = os.environ.get("JUDGE0_API_HOST", "judge0-ce.p.rapidapi.com").strip()
JUDGE0_BASE_URL = (
    f"https://{JUDGE0_API_HOST}" if JUDGE0_API_KEY else "https://ce.judge0.com"
)
# language key (sent by the frontend) -> Judge0 language_id
JUDGE0_LANGUAGE_IDS = {
    "python":     71,  # Python (3.8.1)
    "cpp":        54,  # C++ (GCC 9.2.0)
    "c":          50,  # C (GCC 9.2.0)
    "java":       62,  # Java (OpenJDK 13.0.1)
    "javascript": 63,  # JavaScript (Node.js 12.14.0)
    "typescript": 74,  # TypeScript (3.7.4)
    "go":         60,  # Go (1.13.5)
    "ruby":       72,  # Ruby (2.7.0)
}

# ── Push Notifications (Web Push / VAPID) ────────────────────────────────────
# Generate VAPID keys once and store as env vars:
#   pip install py-vapid
#   vapid --gen   → outputs private + public key
# Then set:
#   VAPID_PRIVATE_KEY  (the full private key PEM or base64url string)
#   VAPID_PUBLIC_KEY   (the applicationServerKey sent to browsers)
#   VAPID_CLAIMS_EMAIL (e.g. mailto:you@example.com)
#
# If keys are not set, push notifications are silently disabled — everything
# else still works normally.
VAPID_PRIVATE_KEY   = os.environ.get("VAPID_PRIVATE_KEY",   "").strip()
# The public key MUST be a single line with no stray whitespace/newlines —
# copy/pasting into a multi-line env-var box (Render, etc.) can silently
# introduce a trailing newline or spaces, which breaks the browser's
# base64url decode and produces "applicationServerKey is not valid" even
# though the key content itself is correct. Strip defensively and also
# collapse any internal whitespace that shouldn't be there.
VAPID_PUBLIC_KEY    = "".join(os.environ.get("VAPID_PUBLIC_KEY", "").split())
VAPID_CLAIMS_EMAIL  = os.environ.get("VAPID_CLAIMS_EMAIL",  "mailto:admin@mythic-ai.app").strip()

# ── Push subscription persistence ─────────────────────────────────────────
# On Render's FREE plan the filesystem is ephemeral: every restart/redeploy/
# spin-down wipes local files, silently deleting every saved push
# subscription. That's why notifications can quietly stop working on devices
# you don't reopen often, while a frequently-reopened tab (which re-saves
# its subscription on every load) looks unaffected.
#
# Fix: optionally persist subscriptions to Upstash Redis (a free, permanent,
# REST-based key-value store — no native Redis driver needed, just HTTPS).
# To enable: sign up at upstash.com (free tier), create a Redis database,
# and set these two env vars on Render:
#   UPSTASH_REDIS_REST_URL
#   UPSTASH_REDIS_REST_TOKEN
# If unset, this falls back to the original local-JSON-file behavior
# unchanged (fine for local dev, but still ephemeral on Render free tier).
UPSTASH_REDIS_REST_URL   = os.environ.get("UPSTASH_REDIS_REST_URL", "").strip()
UPSTASH_REDIS_REST_TOKEN = os.environ.get("UPSTASH_REDIS_REST_TOKEN", "").strip()
_USE_UPSTASH = bool(UPSTASH_REDIS_REST_URL and UPSTASH_REDIS_REST_TOKEN)
_UPSTASH_PUSH_KEY = "mythic:push_subscriptions"

def _upstash_cmd(*parts):
    """Run one Upstash Redis REST command. Returns the 'result' field, or
    None on any failure (network error, bad response, not configured)."""
    if not _USE_UPSTASH:
        return None
    try:
        resp = requests.post(
            UPSTASH_REDIS_REST_URL,
            headers={"Authorization": f"Bearer {UPSTASH_REDIS_REST_TOKEN}"},
            json=list(parts),
            timeout=5,
        )
        if resp.ok:
            return resp.json().get("result")
    except Exception:
        pass
    return None

# In-memory subscription store (replaced by file/Supabase in production)
# Key: a stable browser id, Value: the full PushSubscription JSON object
# (each subscription also carries an internal "_username" field so
# re-engagement notifications can be targeted at a specific person)
_push_subscriptions: dict = {}

def _save_push_subscription(sub_id: str, sub_data: dict):
    _push_subscriptions[sub_id] = sub_data
    if _USE_UPSTASH:
        _upstash_cmd("HSET", _UPSTASH_PUSH_KEY, sub_id, json.dumps(sub_data))
        return
    # Local-file fallback (fine for local dev; ephemeral on Render free tier)
    try:
        path = _os.path.join(_DATA_DIR, "push_subscriptions.json")
        existing = {}
        if _os.path.exists(path):
            with open(path, encoding="utf-8") as f:
                existing = json.load(f)
        existing[sub_id] = sub_data
        with open(path, "w", encoding="utf-8") as f:
            json.dump(existing, f)
    except Exception:
        pass

def _load_push_subscriptions():
    global _push_subscriptions
    if _USE_UPSTASH:
        result = _upstash_cmd("HGETALL", _UPSTASH_PUSH_KEY)
        data = {}
        if result:
            it = iter(result)
            for k, v in zip(it, it):
                try:
                    data[k] = json.loads(v)
                except Exception:
                    pass
        _push_subscriptions = data
        return
    try:
        path = _os.path.join(_DATA_DIR, "push_subscriptions.json")
        if _os.path.exists(path):
            with open(path, encoding="utf-8") as f:
                _push_subscriptions = json.load(f)
    except Exception:
        _push_subscriptions = {}

def _delete_push_subscription(sub_id: str):
    _push_subscriptions.pop(sub_id, None)
    if _USE_UPSTASH:
        _upstash_cmd("HDEL", _UPSTASH_PUSH_KEY, sub_id)
        return
    try:
        path = _os.path.join(_DATA_DIR, "push_subscriptions.json")
        if _os.path.exists(path):
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            data.pop(sub_id, None)
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f)
    except Exception:
        pass

def send_push_notification(title: str, body: str, url: str = "/", icon: str = "/icon.png"):
    """Send a push notification to all subscribed browsers.
    Silently drops dead subscriptions (410 Gone = unsubscribed)."""
    if not _PUSH_AVAILABLE or not VAPID_PRIVATE_KEY or not VAPID_PUBLIC_KEY:
        return
    _load_push_subscriptions()
    dead = []
    payload = json.dumps({"title": title, "body": body, "url": url, "icon": icon})
    for sub_id, sub in list(_push_subscriptions.items()):
        try:
            webpush(
                subscription_info={k: v for k, v in sub.items() if k != "_username"},
                data=payload,
                vapid_private_key=VAPID_PRIVATE_KEY,
                vapid_claims={"sub": VAPID_CLAIMS_EMAIL},
            )
        except WebPushException as ex:
            if ex.response is not None and ex.response.status_code == 410:
                dead.append(sub_id)  # browser unsubscribed
        except Exception:
            pass
    for sub_id in dead:
        _delete_push_subscription(sub_id)


def send_push_notification_to_user(username: str, title: str, body: str,
                                    url: str = "/", icon: str = "/icon.png"):
    """Same as send_push_notification, but only targets subscriptions that
    belong to a specific (anonymous, per-browser) username. Used for
    personalized re-engagement / streak notifications."""
    if not _PUSH_AVAILABLE or not VAPID_PRIVATE_KEY or not VAPID_PUBLIC_KEY:
        return
    _load_push_subscriptions()
    dead = []
    payload = json.dumps({"title": title, "body": body, "url": url, "icon": icon})
    for sub_id, sub in list(_push_subscriptions.items()):
        if sub.get("_username") != username:
            continue
        try:
            webpush(
                subscription_info={k: v for k, v in sub.items() if k != "_username"},
                data=payload,
                vapid_private_key=VAPID_PRIVATE_KEY,
                vapid_claims={"sub": VAPID_CLAIMS_EMAIL},
            )
        except WebPushException as ex:
            if ex.response is not None and ex.response.status_code == 410:
                dead.append(sub_id)
        except Exception:
            pass
    for sub_id in dead:
        _delete_push_subscription(sub_id)


# ── Re-engagement notification message pool ──────────────────────────────────
# Each tuple is (emoji, text). A category is chosen at random when a user has
# gone quiet, and rendered as: "{emoji} Mythic AI: \"{text}\""
NOTIFICATION_MESSAGES = {
    "come_back": [
        ("🤖", "Hey! I'm ready whenever you want to chat."),
        ("💭", "It's been a while. Want to continue our last conversation?"),
        ("✨", "I've got new ideas waiting for you. Let's chat!"),
        ("😊", "Hi! Ask me anything—I'm here to help."),
        ("🚀", "Need help with homework, coding, or anything else? I'm ready."),
    ],
    "study": [
        ("📖", "Ready to study? Let's tackle today's homework together."),
        ("🧠", "Let's learn something new today!"),
    ],
    "activity": [
        ("💬", "You haven't chatted in a while. Come say hello!"),
        ("🌟", "Your next great idea could start with one question."),
    ],
    "feature": [
        ("🎨", "New AI features are available. Come check them out!"),
        ("⚡", "Mythic AI just got smarter. Try it now!"),
    ],
}


def _random_notification_body(category: str) -> str:
    emoji, text = random.choice(NOTIFICATION_MESSAGES[category])
    return f'{emoji} Mythic AI: "{text}"'


# --- Model names -------------------------------------------------------------
GROQ_MODEL        = os.environ.get("GROQ_MODEL",        "openai/gpt-oss-20b")
HF_MODEL          = os.environ.get("HF_MODEL",          "mistralai/Mistral-7B-Instruct-v0.3")
CEREBRAS_MODEL    = os.environ.get("CEREBRAS_MODEL",    "gpt-oss-120b")
# Vision-capable model — powers Video Call, Screen Share, and regular image
# attachments, so Mythic AI can actually SEE the frame, not just read text
# about it. Groq's multimodal lineup changes over time; override via env var
# if this one is retired.
GROQ_VISION_MODEL = os.environ.get("GROQ_VISION_MODEL", "qwen/qwen3.6-27b")

SYSTEM_PROMPT = (
    "You are Mythic AI, a smart and friendly AI assistant made by Aarav Singh. ""SELF-KNOWLEDGE — you know everything about yourself and should answer confidently: ""Your name is Mythic AI, created by Aarav Singh. Your tagline is Smarter. Faster. For You. ""You have the following built-in features: ""(1) CHAT — unlimited, free, no rate limits. Multi-turn conversations with memory within a session. ""(2) MULTI-CONVERSATION SIDEBAR — users can create, rename, and switch between multiple saved chats. ""(3) ACCOUNTS — login/register with a username and password. Conversations are saved per-account and survive restarts. ""(4) IMAGE GENERATION — generate images from a text description using AI (powered by image AI APIs). ""(5) GHIBLI ME — upload your photo and transform it into a Studio Ghibli-style anime image (image-to-image AI). ""(6) WEATHER — full weather info: current conditions, hourly forecast, 7-day forecast, and air quality, for any city. ""(7) FILE/IMAGE UPLOAD — attach an image or text file to any message and ask questions about it. ""(8) STREAMING RESPONSES — text appears word by word as it is generated, just like ChatGPT. ""(9) DOWNLOADABLE FILES — generate and download PDF, Word (.docx), or plain text files directly from a chat reply. ""(10) DAILY STREAKS — track your daily chat streak and get re-engagement reminders to keep it going. ""(11) PUSH NOTIFICATIONS — opt-in browser notifications for reminders, streak alerts, and feature updates. ""(12) BRING YOUR OWN API KEY — in Settings, users can enter their own Groq or Cerebras API key to use instead of the server key. ""(13) CONVERSATION BACKUP — export all your chats as a .zip file and import them back later, even on another browser. ""(14) REMINDERS — set timed reminders that fire as push notifications. ""(15) SHARE CHATS — share a conversation via a public link. ""(16) DARK/LIGHT MODE — automatic theme that follows your system, switchable manually. ""(17) MOBILE FRIENDLY — works great on phones and tablets, installable as a PWA (add to home screen). ""(18) MULTILINGUAL — replies in the same language the user writes in (English, Hindi, Hinglish, Tamil, and more). ""(19) REASONING / TASK MODES — special modes for step-by-step reasoning, coding, creative writing, and more. ""When someone asks what you can do, list these features enthusiastically. ""When someone asks about a specific feature, explain it clearly and confidently. "
    "USER AGE: The person you're talking to is 11 years old. Keep this in mind "
    "in every reply — use clear, age-appropriate language, keep content suitable "
    "for a child, and never produce romantic, sexual, violent, or otherwise "
    "mature content. Don't mention their age unprompted or bring it up "
    "constantly; just let it naturally shape how you explain things and what "
    "you're willing to help with. "
    "If asked who made you, say you are Mythic AI made by Aarav Singh — say it once naturally, never repeat it unprompted. "
    "IDENTITY & ARCHITECTURE: You are a cloud-based application that does not host or train its own "
    "machine learning models. If asked how you work or what powers you, answer honestly and directly: "
    "explain that Mythic AI is a lightweight app that securely routes user requests to external, "
    "third-party foundational AI APIs for chat, plus separate specialized APIs for image generation, "
    "weather, and other utility tools. Do not name the specific underlying model providers/vendors "
    "unless the user explicitly asks which exact provider or model is used — that's a commercial "
    "detail, not something to volunteer, but never deny or deflect the question if asked directly. "
    "Never claim to be a fully custom, locally-trained model built from scratch. "
    "You can help with anything: questions, writing, coding, math, ideas, or just chatting. "
    "When writing code, always wrap it in markdown code blocks with the language name. "
    "LANGUAGE: Your primary/default language is English. Always reply in the SAME language the user's "
    "current message is written in — never switch language on your own. If the user writes in English, "
    "reply fully in English. If the user writes in Hinglish (Hindi+English mixed, Roman script), reply in "
    "that same Hinglish style. If the user writes fully in Hindi using Devanagari script, reply fully in "
    "Hindi using Devanagari script. If they write in Tamil, reply fully in Tamil (Tamil script). The same "
    "rule applies to Gujarati, Marathi, Bengali, Telugu, Malayalam, or any other specific language — always "
    "reply in that language's own native script, fully and consistently, from the first word to the last. "
    "Never mix two languages within a single word or produce garbled or mis-encoded text. If they mix "
    "languages themselves, match their mix. Default to English whenever the user's language is ambiguous "
    "or you are unsure — never default to Hindi. "
    "TOOL USE: Never write out fake tool calls, function names, or JSON like {\"query\": ...} in your reply — "
    "those are internal mechanisms the user must never see. You do not have live web search access — "
    "answer from what you know and say your information may not be fully up to date if asked about "
    "very recent events, instead of pretending to search. "
    "ANTI-REPETITION RULES — follow strictly every reply: "
    "1. NEVER restate or echo back what the user just said. Jump straight to the answer. "
    "2. NEVER start replies with filler like Great question, Sure, Of course, Absolutely, Certainly. "
    "3. NEVER repeat information already given earlier in the conversation. Build on it. "
    "4. Be direct and natural — like a knowledgeable friend, not a customer service bot. "
    "5. Keep answers concise unless the user asks for detail. "
    "6. MATCH YOUR REPLY LENGTH TO THE MESSAGE: a short greeting like 'hi', 'hello', "
    "'hey', 'thanks', or 'ok' gets a short, casual, 1-2 sentence reply — never a long "
    "essay, never a list of your capabilities, never multiple paragraphs. Save longer, "
    "detailed answers for messages that actually ask a real question or request "
    "something specific."
)


app = Flask(__name__)

# --- Explicit session cookie config -------------------------------------------
# Without these, Flask falls back to defaults that can behave inconsistently
# across browsers/proxies (Render sits behind a proxy that terminates HTTPS,
# so we need SESSION_COOKIE_SECURE=True but the app itself sees plain HTTP —
# that's fine, the browser still only sends it over https). Setting these
# explicitly avoids the "works in one browser, silently empty in another"
# class of bug caused by a cookie getting dropped or expiring sooner than
# expected.
app.config.update(
    SESSION_COOKIE_SAMESITE="Lax",
    SESSION_COOKIE_SECURE=True,       # only sent over https (Render is https)
    SESSION_COOKIE_HTTPONLY=True,
    PERMANENT_SESSION_LIFETIME=datetime.timedelta(days=365),
)


def get_public_origin():
    """Returns the public-facing origin (scheme + host), forced to https for
    any real domain. Render/Vercel/most hosts terminate HTTPS at their own
    proxy in front of this app, so Flask's own request.url_root/host_url
    reports plain http:// even though the site is actually only ever
    reachable over https — that mismatch is what causes sitemap.xml,
    robots.txt, canonical/OG tags, and invite/share links to show http://
    instead of https://. Only localhost/127.0.0.1 (local dev) keeps http."""
    origin = request.host_url.rstrip("/")
    host = request.host.split(":")[0]
    if origin.startswith("http://") and host not in ("localhost", "127.0.0.1"):
        origin = "https://" + origin[len("http://"):]
    return origin


# The ONE domain that should show up in Google Search. Every SEO signal
# (canonical link, og:url, JSON-LD @id/url) points here REGARDLESS of which
# domain actually served the request, so Google consolidates all ranking
# signal onto this single URL instead of splitting it across Render+Vercel
# mirrors of the same app. Update this if the Vercel domain changes.
PREFERRED_PUBLIC_ORIGIN = "https://mythic-ai.vercel.app"


def _is_deindexed_host():
    """Hosts that should NEVER show up in search results — currently any
    Render deployment domain, since Vercel (PREFERRED_PUBLIC_ORIGIN) is the
    one and only domain meant to be public-facing in Google. Render domains
    always end in onrender.com unless a custom domain is mapped to Render
    specifically (not the case here)."""
    host = request.host.split(":")[0].lower()
    return host.endswith(".onrender.com")



# ── Reasoning/task modes — pure prompt-engineering, no extra APIs needed ────
# Selected client-side and sent as `mode` with each /api/chat call; appended
# to the system prompt for that request only.
MODE_PROMPTS = {
    "default": "",
    "cowork": (
        "MODE: Cowork — a general multi-step task assistant, not just a coder. "
        "If the task is to build/fix software or a website: do NOT just describe a plan "
        "and stop — output the complete, runnable code for the whole thing in this single "
        "reply, split into clearly labeled files using fenced code blocks with the language "
        "and filename on the first line as a comment (e.g. ```html <!-- index.html -->). "
        "Never ask clarifying questions before producing a first working version — make "
        "reasonable assumptions and state them briefly instead. "
        "If the task is NOT about code (research, writing, planning, comparisons, etc.): "
        "just do the task directly and completely, in whatever format actually fits it — "
        "don't force code blocks or a file structure where none make sense. "
        "In both cases: finish the work in this reply rather than only listing next steps."
    ),
    "coding": (
        "MODE: Coding Assistant. Prioritize correct, runnable code. Always specify "
        "the language in code fences. Briefly explain non-obvious logic. When asked "
        "to debug, identify the exact bug before proposing a fix. When asked to "
        "generate a project, structure it file by file with clear filenames."
    ),
    "research": (
        "MODE: Research Assistant. Structure answers with clear sections. Explicitly "
        "flag where a claim is uncertain, contested, or outside your training data "
        "rather than presenting speculation as fact. Prefer thorough, well-organized "
        "answers over short ones."
    ),
    "study": (
        "MODE: Study/Homework Helper. Explain concepts step-by-step like a patient "
        "tutor. After solving a problem, briefly state the method so the student can "
        "apply it themselves next time. Suitable for CBSE/NCERT and general curricula."
    ),
    "debate": (
        "MODE: Debate Partner. When given a position, construct the strongest "
        "good-faith argument for it, then present the strongest counterarguments. "
        "Stay even-handed and avoid straw-manning either side."
    ),
    "business": (
        "MODE: Business Assistant. Be concise, structured, and action-oriented — "
        "think memos, plans, and decision frameworks rather than essays."
    ),
    "math": (
        "MODE: Math Solver. Show step-by-step working, not just the final answer. "
        "Use LaTeX notation ($...$ for inline, $$...$$ for display) for equations. "
        "Double-check arithmetic before presenting the final result."
    ),
    "translation": (
        "MODE: Translator. Translate faithfully, preserving tone and meaning. If "
        "helpful, briefly explain notable grammar or idiom choices after the translation."
    ),
    "writing": (
        "MODE: Writing Assistant. Match the requested tone and format precisely "
        "(email, blog, resume, essay, letter, etc.). Give the finished piece first; "
        "only add commentary if asked."
    ),
    "brainstorm": (
        "MODE: Brainstorming Partner. Generate a wide variety of genuinely distinct "
        "ideas rather than elaborating on just one. Use short bullet points unless "
        "asked for depth."
    ),
}


def _persistent_secret_key():
    """A stable Flask secret key across restarts/workers so each visitor's
    session cookie (which stores their anonymous user_id) doesn't get
    invalidated every time the server restarts. Without this, a random key
    was generated per-process, which silently "lost" every conversation
    tied to the old session on every restart/redeploy/worker respawn —
    the classic cause of "it only ever shows 1 chat".

    Priority: FLASK_SECRET_KEY env var > a key file persisted next to this
    script. Setting FLASK_SECRET_KEY explicitly is STRONGLY recommended for
    any real deployment (Render, etc.), since some hosts wipe local disk
    between deploys, which would defeat the file fallback too.
    """
    env_key = os.environ.get("FLASK_SECRET_KEY")
    if env_key:
        return env_key
    # On Vercel (and any serverless host), the deployed script directory is
    # READ-ONLY — only /tmp is writable, and /tmp itself is ephemeral and NOT
    # shared across instances. That combination means this file-based fallback
    # can never actually keep your session stable on Vercel: writes here may
    # raise (read-only fs) or, even if /tmp is used, vanish on the next cold
    # start / different instance, silently minting a new "you" and breaking
    # owner checks like the API keys page. Detect that case explicitly instead
    # of crashing, and fall back to a same-process-lifetime key so at least a
    # single warm instance stays consistent — but this is NOT a real fix for
    # Vercel; set FLASK_SECRET_KEY as an env var there (see comment above).
    is_serverless = bool(os.environ.get("VERCEL") or os.environ.get("VERCEL_ENV"))
    if is_serverless:
        print("[secret-key] WARNING: FLASK_SECRET_KEY is not set and this is a "
              "serverless (Vercel) deployment. Sessions/owner status WILL NOT "
              "persist reliably across requests. Set FLASK_SECRET_KEY in your "
              "Vercel project's environment variables to fix this.")
        return str(uuid.uuid4())
    base_dir = os.path.dirname(os.path.abspath(__file__))
    data_dir = os.path.join(base_dir, "chat_data")
    try:
        os.makedirs(data_dir, exist_ok=True)
        key_path = os.path.join(data_dir, "flask_secret.key")
        if os.path.exists(key_path):
            with open(key_path, "r") as f:
                existing = f.read().strip()
            if existing:
                return existing
        new_key = str(uuid.uuid4())
        with open(key_path, "w") as f:
            f.write(new_key)
        return new_key
    except OSError as e:
        print(f"[secret-key] Could not persist secret key to disk ({e}); "
              f"using a process-lifetime key instead. Set FLASK_SECRET_KEY "
              f"as an env var for a permanent fix.")
        return str(uuid.uuid4())


app.secret_key = _persistent_secret_key()

MAX_UPLOAD_BYTES = 8 * 1024 * 1024  # 8 MB — images and general attachments
DOCUMENT_UPLOAD_BYTES = 1024 * 1024 * 1024  # 1 GB — books/PDFs/docs specifically

# --- Temporary public image hosting (for NanoBanana image-to-image editing) --
# NanoBanana's /generate endpoint needs a publicly reachable image URL for
# edit-mode requests — it can't accept a raw base64 upload directly. So when a
# user uploads a selfie for "Ghibli Me", we stash the bytes here in memory
# under a random id and serve them back at /api/temp-image/<id>, giving
# NanoBanana's servers a URL they can fetch. Entries expire on their own after
# a short TTL so memory doesn't grow unbounded.
_TEMP_IMAGES = {}
_TEMP_IMAGE_TTL_SECONDS = 30 * 60  # 30 minutes


def _store_temp_image(raw_bytes, mime_type):
    # Opportunistic cleanup of anything past its TTL
    cutoff = time.time() - _TEMP_IMAGE_TTL_SECONDS
    for k in [k for k, v in _TEMP_IMAGES.items() if v["created"] < cutoff]:
        _TEMP_IMAGES.pop(k, None)
    img_id = uuid.uuid4().hex
    _TEMP_IMAGES[img_id] = {"data": raw_bytes, "mime_type": mime_type or "image/png", "created": time.time()}
    return img_id


def nano_banana_submit(prompt, image_urls=None, num_images=1):
    """Submits a NanoBanana generation/edit task. Returns (task_id, error)."""
    if not NANO_BANANA_API_KEY:
        return None, "NanoBanana API key not configured"
    payload = {
        "prompt": prompt,
        "type": "IMAGETOIMAGE" if image_urls else "TEXTTOIAMGE",
        "numImages": num_images,
    }
    if image_urls:
        payload["imageUrls"] = image_urls
    try:
        resp = requests.post(
            f"{NANO_BANANA_BASE}/generate",
            headers={"Authorization": f"Bearer {NANO_BANANA_API_KEY}", "Content-Type": "application/json"},
            json=payload,
            timeout=30,
        )
        data = resp.json()
        if resp.status_code == 200 and data.get("code") == 200:
            return data["data"]["taskId"], None
        return None, data.get("msg") or f"NanoBanana error ({resp.status_code})"
    except (requests.RequestException, ValueError) as e:
        return None, str(e)


def nano_banana_poll(task_id, max_wait=180, interval=3):
    """Polls a NanoBanana task until it succeeds/fails/times out.
    Returns (result_image_url, error)."""
    headers = {"Authorization": f"Bearer {NANO_BANANA_API_KEY}"}
    deadline = time.time() + max_wait
    while time.time() < deadline:
        try:
            resp = requests.get(
                f"{NANO_BANANA_BASE}/record-info",
                params={"taskId": task_id}, headers=headers, timeout=15,
            )
            data = resp.json()
        except (requests.RequestException, ValueError) as e:
            return None, str(e)
        flag = data.get("successFlag")
        if flag == 1:
            result = data.get("response") or {}
            url = result.get("resultImageUrl")
            if not url and isinstance(result.get("resultImageUrls"), list) and result["resultImageUrls"]:
                url = result["resultImageUrls"][0]
            if not url:
                return None, "NanoBanana returned no result image"
            return url, None
        if flag in (2, 3):
            return None, data.get("errorMessage") or "NanoBanana generation failed"
        time.sleep(interval)
    return None, "NanoBanana generation timed out"

# --- Supabase config ---------------------------------------------------------
# Never hardcode real keys here — always set these as environment variables:
#   SUPABASE_URL  -> your project URL, e.g. https://xxxxx.supabase.co
#   SUPABASE_KEY  -> your Supabase *secret* (service-role-style) key
# The secret key grants full database access, so it must only ever live in
# server-side environment variables/secret storage, never in source code,
# client-side JS, or anything committed to a public repo.
SUPABASE_URL = os.environ.get("SUPABASE_URL", "").rstrip("/")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY", "")

def sb_headers():
    return {
        "apikey": SUPABASE_KEY,
        "Authorization": f"Bearer {SUPABASE_KEY}",
        "Content-Type": "application/json",
        "Prefer": "return=representation",
    }

def sb(path):
    return f"{SUPABASE_URL}/rest/v1/{path}"


# --- User accounts (Supabase: users table) -----------------------------------

def current_username():
    """Returns the id of the currently AUTHENTICATED account. Only ever
    called from inside a @login_required view, which guarantees
    session["user_id"] + session["authenticated"] are already set — see
    login_required below, which is the actual gate."""
    return session.get("user_id")


def login_required(view):
    """Real auth gate. Unauthenticated visitors get:
      - redirected to /login for normal page routes
      - a 401 JSON error for /api/... routes (so frontend fetch calls can
        detect it and redirect client-side instead of rendering raw HTML)
    A request is considered authenticated once session["authenticated"] is
    True — set by email/password login, Google sign-in, an /invite/<token>
    link, or /claim-owner/<secret>, all of which are equivalent proof of
    "this browser is allowed into this specific account"."""
    def wrapped(*args, **kwargs):
        if not session.get("authenticated") or not session.get("user_id"):
            if request.path.startswith("/api/"):
                return jsonify({"error": "authentication required", "login_url": "/login"}), 401
            return redirect("/login")
        return view(*args, **kwargs)
    wrapped.__name__ = view.__name__
    return wrapped


# --- Supabase / file storage helpers ----------------------------------------

def list_conversations(username):
    if not SUPABASE_URL:
        return _list_conversations_file(username)
    try:
        r = requests.get(
            sb(f"conversations?username=eq.{username}&order=updated_at.desc"
               f"&select=id,title,updated_at,folder,pinned,archived"),
            headers=sb_headers(), timeout=10,
        )
        if r.status_code == 200:
            rows = r.json()
            rows.sort(key=lambda c: (not c.get("pinned", False), -(c.get("updated_at") or 0)))
            return rows
        else:
            print(f"[Supabase] list_conversations failed: HTTP {r.status_code} — {r.text[:300]}")
    except Exception as e:
        print(f"[Supabase] list_conversations exception: {e}")
    return []


def load_conversation(username, conv_id):
    if not SUPABASE_URL:
        return _load_conversation_file(username, conv_id)
    try:
        r = requests.get(
            sb(f"conversations?id=eq.{conv_id}&username=eq.{username}"),
            headers=sb_headers(), timeout=10,
        )
        if r.status_code == 200:
            rows = r.json()
            if rows:
                row = rows[0]
                return {
                    "title": row["title"],
                    "updated_at": row["updated_at"],
                    "messages": row["messages"] if isinstance(row["messages"], list) else json.loads(row["messages"]),
                }
        else:
            print(f"[Supabase] load_conversation failed: HTTP {r.status_code} — {r.text[:300]}")
    except Exception as e:
        print(f"[Supabase] load_conversation exception: {e}")
    return None


def save_conversation(username, conv_id, data):
    data["updated_at"] = time.time()
    if not SUPABASE_URL:
        _save_conversation_file(username, conv_id, data)
        return
    try:
        headers = {**sb_headers(), "Prefer": "resolution=merge-duplicates,return=minimal"}
        r = requests.post(
            sb("conversations"),
            headers=headers,
            json={
                "id": conv_id,
                "username": username,
                "title": data.get("title", "New chat"),
                "updated_at": data["updated_at"],
                "messages": data.get("messages", []),
                "folder": data.get("folder"),
                "pinned": bool(data.get("pinned", False)),
                "archived": bool(data.get("archived", False)),
            },
            timeout=15,
        )
        if r.status_code not in (200, 201, 204):
            print(f"[Supabase] save_conversation failed: HTTP {r.status_code} — {r.text[:500]}")
    except Exception as e:
        print(f"[Supabase] save_conversation exception: {e}")


def delete_conversation(username, conv_id):
    if not SUPABASE_URL:
        _delete_conversation_file(username, conv_id)
        return
    try:
        requests.delete(
            sb(f"conversations?id=eq.{conv_id}&username=eq.{username}"),
            headers=sb_headers(), timeout=10,
        )
    except Exception:
        pass


# --- Local file fallbacks for when Supabase is not configured ----------------
import os as _os

# ── Serverless (Vercel) compatibility notes ──────────────────────────────────
# This app was written as a normal long-running server (Render/Railway/a VPS)
# and several of its features fundamentally rely on that:
#   1. It keeps push subscriptions, temp image uploads, and conversations in
#      local JSON files / in-memory dicts. Vercel's filesystem is READ-ONLY
#      except for /tmp, and /tmp (plus all memory) is wiped on every cold
#      start and isn't shared across instances — so conversations, streaks,
#      and push subscriptions will keep "disappearing" there.
#   2. The hourly re-engagement notifications rely on a background thread
#      that runs forever. Vercel serverless functions only run while
#      handling a request and are frozen/killed otherwise, so that thread
#      never gets to fire on its own schedule.
#   3. /api/chat streams its reply chunk-by-chunk (SSE-style). Vercel's
#      default Node/Python serverless functions buffer and have a max
#      execution duration, so long streaming replies can get cut off or
#      arrive all at once instead of live.
# If Vercel is a hard requirement, the practical fix is to swap local JSON
# storage for a real database (Vercel KV / Postgres / Supabase — which this
# file already partially supports), and move the notification scheduler to
# an external cron (e.g. Vercel Cron Jobs calling a small endpoint) instead
# of an in-process thread. Otherwise, a normal always-on host (Render,
# Railway, Fly.io, a VPS) is a much better match for this code as written.
IS_SERVERLESS = bool(_os.environ.get("VERCEL") or _os.environ.get("VERCEL_ENV"))

_BASE_DIR = _os.path.dirname(_os.path.abspath(__file__))
if IS_SERVERLESS:
    # /tmp is the only writable path on Vercel — this avoids crashing with a
    # read-only-filesystem error, even though data still won't persist
    # across cold starts/instances there. See notes above.
    _DATA_DIR = _os.path.join("/tmp", "mythic_ai_chat_data")
else:
    _DATA_DIR = _os.path.join(_BASE_DIR, "chat_data")
_os.makedirs(_DATA_DIR, exist_ok=True)

def _user_conv_dir(username):
    path = _os.path.join(_DATA_DIR, "conversations", username)
    _os.makedirs(path, exist_ok=True)
    return path

# --- Permanent, account-wide invite link -------------------------------------
# One stable code, generated once and reused forever after (persisted to
# disk / Supabase so it doesn't change on every restart or redeploy). This
# is what makes the invite link look/behave like an actual generated share
# link (…/invite/<code>) instead of just the bare domain, while still
# always resolving to the same public, no-login chat page for everyone.
_INVITE_CODE_FILE = _os.path.join(_DATA_DIR, "invite_code.txt")
_invite_code_lock = threading.Lock()

def get_or_create_invite_code():
    with _invite_code_lock:
        if SUPABASE_URL:
            try:
                r = requests.get(sb("app_settings?key=eq.invite_code&select=value"),
                                  headers=sb_headers(), timeout=10)
                if r.status_code == 200 and r.json():
                    return r.json()[0]["value"]
            except Exception as e:
                print(f"[Supabase] get_or_create_invite_code read failed: {e} — falling back to local file.")
        if _os.path.exists(_INVITE_CODE_FILE):
            try:
                with open(_INVITE_CODE_FILE, encoding="utf-8") as f:
                    code = f.read().strip()
                    if code:
                        return code
            except Exception:
                pass
        code = uuid.uuid4().hex[:12]
        if SUPABASE_URL:
            try:
                requests.post(sb("app_settings"), headers=sb_headers(),
                              json={"key": "invite_code", "value": code}, timeout=10)
            except Exception as e:
                print(f"[Supabase] get_or_create_invite_code write failed: {e} — falling back to local file.")
        try:
            with open(_INVITE_CODE_FILE, "w", encoding="utf-8") as f:
                f.write(code)
        except Exception as e:
            print(f"[invite] could not persist invite code to disk: {e}")
        return code

# --- Owner account id ---------------------------------------------------------
# The invite link is meant to share YOUR chat history with whoever opens it
# (not give them their own empty account). So we need one fixed "owner"
# user_id that anyone opening /invite/<code> gets logged into directly.
# Persisted the same way as the invite code so it survives restarts/redeploys.
_OWNER_ID_FILE = _os.path.join(_DATA_DIR, "owner_user_id.txt")
_owner_id_lock = threading.Lock()

# Hard override for serverless hosts (Vercel, etc.) where neither the local
# filesystem (/tmp is wiped on cold start) nor the session cookie can be
# trusted to persist the owner id reliably without Supabase configured.
# Set OWNER_USER_ID to any fixed string (e.g. a UUID you generate once) as
# an environment variable, and that value is used directly, every time, on
# every instance — no file, no database, no race to "claim" ownership.
_OWNER_USER_ID_ENV = _os.environ.get("OWNER_USER_ID", "").strip()

def get_or_create_owner_id(preferred_id=None):
    if _OWNER_USER_ID_ENV:
        return _OWNER_USER_ID_ENV
    with _owner_id_lock:
        if SUPABASE_URL:
            try:
                r = requests.get(sb("app_settings?key=eq.owner_user_id&select=value"),
                                  headers=sb_headers(), timeout=10)
                if r.status_code == 200 and r.json():
                    return r.json()[0]["value"]
            except Exception as e:
                print(f"[owner] Supabase read failed: {e} — falling back to local file.")
        if _os.path.exists(_OWNER_ID_FILE):
            try:
                with open(_OWNER_ID_FILE, encoding="utf-8") as f:
                    oid = f.read().strip()
                    if oid:
                        return oid
            except Exception:
                pass
        # First time this is ever called: prefer adopting the caller's
        # existing session id (so their real chat history becomes the
        # shared "owner" history) rather than minting a brand new empty one.
        oid = preferred_id if preferred_id else str(uuid.uuid4())
        if SUPABASE_URL:
            try:
                requests.post(sb("app_settings"), headers=sb_headers(),
                              json={"key": "owner_user_id", "value": oid}, timeout=10)
            except Exception as e:
                print(f"[owner] Supabase write failed: {e} — falling back to local file.")
        try:
            with open(_OWNER_ID_FILE, "w", encoding="utf-8") as f:
                f.write(oid)
        except Exception as e:
            print(f"[owner] could not persist owner id to disk: {e}")
        return oid


# --- Per-account unique links (one distinct URL per visitor's own account) --
# The /invite/<code> link above intentionally sends EVERY visitor into one
# shared "owner" account/history — useful if you want a single conversation
# thread anyone can add to. This is the opposite: each browser/account gets
# its OWN permanent, unique URL (…/invite/<token>) that always logs back into
# THAT SPECIFIC account's private chats — nobody else's. Bookmarking or
# sharing that link only ever opens the same one account, not a fresh one.
_ACCOUNT_TOKENS_FILE = _os.path.join(_DATA_DIR, "account_tokens.json")
_account_tokens_lock = threading.Lock()


def _load_account_tokens():
    if _os.path.exists(_ACCOUNT_TOKENS_FILE):
        try:
            with open(_ACCOUNT_TOKENS_FILE, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def _save_account_tokens(data):
    try:
        with open(_ACCOUNT_TOKENS_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f)
    except Exception as e:
        print(f"[account-token] could not persist tokens to disk: {e}")


def get_or_create_account_token(user_id):
    """Returns this account's permanent unique token, creating one on first
    call. Same token forever for the same user_id — the URL never changes
    on redeploy/restart, since it's persisted to disk/Supabase like the
    invite code and owner id above."""
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"account_tokens?user_id=eq.{user_id}&select=token"),
                              headers=sb_headers(), timeout=10)
            if r.status_code == 200 and r.json():
                return r.json()[0]["token"]
        except Exception as e:
            print(f"[account-token] Supabase read failed: {e} — falling back to local file.")

    with _account_tokens_lock:
        tokens = _load_account_tokens()
        # tokens maps token -> user_id; do a reverse lookup for existing token
        for tok, uid in tokens.items():
            if uid == user_id:
                return tok
        new_token = uuid.uuid4().hex[:16]
        tokens[new_token] = user_id
        _save_account_tokens(tokens)

    if SUPABASE_URL:
        try:
            requests.post(sb("account_tokens"), headers=sb_headers(),
                          json={"token": new_token, "user_id": user_id}, timeout=10)
        except Exception as e:
            print(f"[account-token] Supabase write failed: {e} — token still works via local file.")

    return new_token


def resolve_account_token(token):
    """Returns the user_id this token belongs to, or None if unknown."""
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"account_tokens?token=eq.{token}&select=user_id"),
                              headers=sb_headers(), timeout=10)
            if r.status_code == 200 and r.json():
                return r.json()[0]["user_id"]
        except Exception as e:
            print(f"[account-token] Supabase resolve failed: {e} — falling back to local file.")
    tokens = _load_account_tokens()
    return tokens.get(token)


# --- User accounts: email/password + Google Sign-In --------------------------
_USERS_FILE = _os.path.join(_DATA_DIR, "users.json")
_users_lock = threading.Lock()


def _load_users_store_file():
    with _users_lock:
        if _os.path.exists(_USERS_FILE):
            try:
                with open(_USERS_FILE, encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {"by_id": {}, "by_email": {}, "by_google_sub": {}}


def _save_users_store_file(store):
    with _users_lock:
        try:
            with open(_USERS_FILE, "w", encoding="utf-8") as f:
                json.dump(store, f)
        except Exception as e:
            print(f"[users] failed to save: {e}")


# --- Users: Supabase-backed when configured, local JSON file otherwise ------
# IMPORTANT: on Vercel specifically, the local-file fallback does NOT persist
# reliably (see IS_SERVERLESS notes above — /tmp is wiped on every cold start
# and isn't shared across instances). Now that accounts hold real passwords,
# set SUPABASE_URL + SUPABASE_KEY and create this table before relying on
# Vercel in production:
#
#   create table users (
#     user_id text primary key,
#     email text unique not null,
#     name text,
#     picture text,
#     password_hash text,
#     google_sub text unique,
#     created_at double precision
#   );
#
# (Supabase's REST API — PostgREST — is what sb()/sb_headers() below talk to;
# no extra Python client library needed.)

def get_user_by_id(user_id):
    if not user_id:
        return None
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"users?user_id=eq.{user_id}"), headers=sb_headers(), timeout=10)
            if r.status_code == 200 and r.json():
                return r.json()[0]
        except Exception as e:
            print(f"[Supabase] get_user_by_id failed: {e}")
        return None
    return _load_users_store_file()["by_id"].get(user_id)


def get_user_by_email(email):
    if not email:
        return None
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"users?email=eq.{email}"), headers=sb_headers(), timeout=10)
            if r.status_code == 200 and r.json():
                return r.json()[0]
        except Exception as e:
            print(f"[Supabase] get_user_by_email failed: {e}")
        return None
    store = _load_users_store_file()
    uid = store["by_email"].get(email)
    return store["by_id"].get(uid) if uid else None


def get_user_by_google_sub(google_sub):
    if not google_sub:
        return None
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"users?google_sub=eq.{google_sub}"), headers=sb_headers(), timeout=10)
            if r.status_code == 200 and r.json():
                return r.json()[0]
        except Exception as e:
            print(f"[Supabase] get_user_by_google_sub failed: {e}")
        return None
    store = _load_users_store_file()
    uid = store["by_google_sub"].get(google_sub)
    return store["by_id"].get(uid) if uid else None


def create_user(record):
    """record must include user_id, email; may include name, picture,
    password_hash, google_sub, created_at."""
    record.setdefault("created_at", time.time())
    if SUPABASE_URL:
        try:
            r = requests.post(sb("users"), headers=sb_headers(), json=record, timeout=10)
            if r.status_code not in (200, 201):
                print(f"[Supabase] create_user failed: HTTP {r.status_code} — {r.text[:300]}")
        except Exception as e:
            print(f"[Supabase] create_user exception: {e}")
        return record
    store = _load_users_store_file()
    store["by_id"][record["user_id"]] = record
    store["by_email"][record["email"]] = record["user_id"]
    if record.get("google_sub"):
        store["by_google_sub"][record["google_sub"]] = record["user_id"]
    _save_users_store_file(store)
    return record


def update_user(user_id, fields):
    if SUPABASE_URL:
        try:
            r = requests.patch(sb(f"users?user_id=eq.{user_id}"), headers=sb_headers(), json=fields, timeout=10)
            if r.status_code not in (200, 204):
                print(f"[Supabase] update_user failed: HTTP {r.status_code} — {r.text[:300]}")
        except Exception as e:
            print(f"[Supabase] update_user exception: {e}")
        return
    store = _load_users_store_file()
    user = store["by_id"].get(user_id)
    if not user:
        return
    user.update(fields)
    if fields.get("google_sub"):
        store["by_google_sub"][fields["google_sub"]] = user_id
    _save_users_store_file(store)


_EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")



def _pending_local_id():
    """If this browser already has a locally-known id (from before this
    login gate existed, or from an in-progress OAuth redirect), returns it
    so a brand-new signup can adopt it as the account's permanent id —
    meaning any conversations/memories/personas already saved under that id
    keep working immediately instead of starting empty. Falls back to a
    fresh uuid if there's nothing to adopt."""
    candidate = session.get("pending_anon_id") or request.headers.get("X-Client-Id", "").strip()
    if candidate:
        try:
            uuid.UUID(candidate)
            return candidate
        except (ValueError, AttributeError):
            pass
    return str(uuid.uuid4())


def _log_user_in(user_id):
    session["user_id"] = user_id
    session["authenticated"] = True
    session.permanent = True
    session.pop("pending_anon_id", None)


@app.route("/api/auth/me", methods=["GET"])
def api_auth_me():
    """Lets the frontend check auth status without triggering the redirect
    login_required would do — used to decide whether to show the login
    screen or the app shell."""
    if not session.get("authenticated") or not session.get("user_id"):
        return jsonify({"authenticated": False})
    u = get_user_by_id(session["user_id"]) or {}
    return jsonify({
        "authenticated": True,
        "email": u.get("email"),
        "name": u.get("name"),
        "picture": u.get("picture"),
    })


@app.route("/api/auth/logout", methods=["POST"])
def api_auth_logout():
    session.clear()
    return jsonify({"status": "ok"})


@app.route("/api/auth/signup", methods=["POST"])
def api_auth_signup():
    # Passwords are gone entirely — email sign-in is now a one-time code
    # (OTP) sent to the address, handled by /api/auth/otp/request and
    # /api/auth/otp/verify below. This endpoint stays only so any old
    # cached frontend hitting it gets a clear message instead of a raw 404.
    return jsonify({"error": "Password sign-up is disabled. Use 'Continue with Google' "
                              "or sign in with an emailed code instead."}), 403


@app.route("/api/auth/login", methods=["POST"])
def api_auth_login():
    return jsonify({"error": "Password login is disabled. Use 'Continue with Google' "
                              "or sign in with an emailed code instead."}), 403


# --- Email OTP (one-time code) sign-in, no passwords ------------------------
# Dual-backend like users/memories/personas: Supabase table if configured,
# local JSON file otherwise. Table:
#   create table otp_codes (
#     email text primary key, code text not null,
#     expires_at double precision not null, attempts int default 0,
#     last_sent_at double precision
#   );
_OTP_FILE = _os.path.join(_DATA_DIR, "otp_codes.json")
_otp_lock = threading.Lock()
_OTP_TTL_SECONDS = 10 * 60      # code valid for 10 minutes
_OTP_RESEND_COOLDOWN = 45       # seconds between resend requests per email
_OTP_MAX_ATTEMPTS = 5           # wrong-code guesses allowed before the code is dead


def _load_otp_store_file():
    with _otp_lock:
        if _os.path.exists(_OTP_FILE):
            try:
                with open(_OTP_FILE, encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {}


def _save_otp_store_file(store):
    with _otp_lock:
        try:
            with open(_OTP_FILE, "w", encoding="utf-8") as f:
                json.dump(store, f)
        except Exception as e:
            print(f"[otp] failed to save: {e}")


def _get_otp_row(email):
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"otp_codes?email=eq.{email}"), headers=sb_headers(), timeout=10)
            if r.status_code == 200 and r.json():
                return r.json()[0]
        except Exception as e:
            print(f"[Supabase] get_otp_row failed: {e}")
        return None
    return _load_otp_store_file().get(email)


def _upsert_otp_row(row):
    """Returns (True, None) or (False, error_detail_string) so the caller
    can tell the requester if the code genuinely got saved, instead of
    always claiming success even when the Supabase write silently failed
    (e.g. missing table grants)."""
    if SUPABASE_URL:
        try:
            headers = {**sb_headers(), "Prefer": "resolution=merge-duplicates,return=minimal"}
            r = requests.post(sb("otp_codes"), headers=headers, json=row, timeout=10)
            if r.status_code not in (200, 201, 204):
                detail = f"HTTP {r.status_code} — {r.text[:300]}"
                print(f"[Supabase] upsert_otp_row failed: {detail}")
                return False, detail
            return True, None
        except Exception as e:
            print(f"[Supabase] upsert_otp_row exception: {e}")
            return False, str(e)
    store = _load_otp_store_file()
    store[row["email"]] = row
    _save_otp_store_file(store)
    return True, None


def _delete_otp_row(email):
    if SUPABASE_URL:
        try:
            requests.delete(sb(f"otp_codes?email=eq.{email}"), headers=sb_headers(), timeout=10)
        except Exception as e:
            print(f"[Supabase] delete_otp_row failed: {e}")
        return
    store = _load_otp_store_file()
    store.pop(email, None)
    _save_otp_store_file(store)


def _send_otp_email(to_email, code):
    if not (SMTP_HOST and SMTP_USER and SMTP_PASSWORD):
        print("[otp] SMTP not configured — cannot send code. "
              "Set SMTP_HOST/SMTP_USER/SMTP_PASSWORD env vars.")
        return False
    msg = MIMEText(
        f"Your Mythic AI sign-in code is:\n\n{code}\n\n"
        f"This code expires in 10 minutes. If you didn't request this, you can ignore this email."
    )
    msg["Subject"] = f"{code} is your Mythic AI sign-in code"
    msg["From"] = formataddr((SMTP_FROM_NAME, SMTP_FROM))
    msg["To"] = to_email
    try:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=15) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.sendmail(SMTP_FROM, [to_email], msg.as_string())
        return True
    except Exception as e:
        print(f"[otp] send failed for {to_email}: {e}")
        return False


@app.route("/api/auth/otp/request", methods=["POST"])
def api_auth_otp_request():
    data = request.get_json(force=True) or {}
    email = (data.get("email") or "").strip().lower()
    if not _EMAIL_RE.match(email):
        return jsonify({"error": "Enter a valid email address."}), 400

    existing = _get_otp_row(email)
    now = time.time()
    if existing and now - existing.get("last_sent_at", 0) < _OTP_RESEND_COOLDOWN:
        wait = int(_OTP_RESEND_COOLDOWN - (now - existing["last_sent_at"]))
        return jsonify({"error": f"Please wait {wait}s before requesting another code."}), 429

    code = f"{secrets.randbelow(1000000):06d}"

    saved, save_error = _upsert_otp_row({
        "email": email, "code": code,
        "expires_at": now + _OTP_TTL_SECONDS,
        "attempts": 0, "last_sent_at": now,
    })
    if not saved:
        # TEMPORARY: exposing save_error directly in the response for
        # debugging the current Supabase setup issue. Once storage is
        # confirmed working, this detail should go back to server logs
        # only (see the print() in _upsert_otp_row) rather than being
        # shown to whoever is signing in.
        return jsonify({"error": f"The server couldn't save the code (storage error): {save_error}"}), 500

    if not _send_otp_email(email, code):
        return jsonify({"error": "Couldn't send the code — email sending isn't configured on "
                                  "the server yet. Try 'Continue with Google' instead, or "
                                  "contact the site owner."}), 500

    return jsonify({"status": "sent"})


@app.route("/api/auth/otp/verify", methods=["POST"])
def api_auth_otp_verify():
    data = request.get_json(force=True) or {}
    email = (data.get("email") or "").strip().lower()
    code = (data.get("code") or "").strip()

    row = _get_otp_row(email)
    if not row:
        return jsonify({"error": "Request a new code first."}), 400
    if time.time() > row.get("expires_at", 0):
        _delete_otp_row(email)
        return jsonify({"error": "That code expired. Request a new one."}), 400
    if row.get("attempts", 0) >= _OTP_MAX_ATTEMPTS:
        _delete_otp_row(email)
        return jsonify({"error": "Too many incorrect attempts. Request a new code."}), 429
    if code != row.get("code"):
        row["attempts"] = row.get("attempts", 0) + 1
        _upsert_otp_row(row)
        remaining = _OTP_MAX_ATTEMPTS - row["attempts"]
        return jsonify({"error": f"Incorrect code. {remaining} attempt(s) left."}), 401

    _delete_otp_row(email)  # one-time — burn it immediately on success
    user = get_user_by_email(email)
    if user:
        user_id = user["user_id"]
    else:
        user_id = _pending_local_id()
        create_user({
            "user_id": user_id, "email": email, "name": email.split("@")[0],
            "picture": None, "password_hash": None, "google_sub": None,
            "created_at": time.time(),
        })
    _log_user_in(user_id)
    return jsonify({"status": "ok", "email": email})


@app.route("/api/auth/google/start", methods=["GET"])

def api_auth_google_start():
    if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
        return Response(
            "Google sign-in isn't configured yet — the server is missing "
            "GOOGLE_CLIENT_ID / GOOGLE_CLIENT_SECRET environment variables.",
            mimetype="text/plain"), 500

    state = secrets.token_urlsafe(24)
    session["oauth_state"] = state
    # Preserve any pre-existing local/anon id across the redirect so a NEW
    # Google signup can still adopt it (see _pending_local_id above).
    session["pending_anon_id"] = request.headers.get("X-Client-Id", "").strip() or None
    redirect_uri = get_public_origin() + "/api/auth/google/callback"
    params = {
        "client_id": GOOGLE_CLIENT_ID,
        "redirect_uri": redirect_uri,
        "response_type": "code",
        "scope": "openid email profile",
        "state": state,
        "access_type": "online",
        "prompt": "select_account",
    }
    return redirect("https://accounts.google.com/o/oauth2/v2/auth?" + urllib.parse.urlencode(params))


@app.route("/api/auth/google/callback", methods=["GET"])
def api_auth_google_callback():
    error = request.args.get("error")
    if error:
        return redirect("/login?error=" + urllib.parse.quote(error))

    state = request.args.get("state", "")
    if not state or state != session.get("oauth_state"):
        return redirect("/login?error=invalid_state")
    code = request.args.get("code", "")
    if not code:
        return redirect("/login?error=missing_code")

    redirect_uri = get_public_origin() + "/api/auth/google/callback"
    try:
        token_resp = requests.post("https://oauth2.googleapis.com/token", data={
            "client_id": GOOGLE_CLIENT_ID,
            "client_secret": GOOGLE_CLIENT_SECRET,
            "code": code,
            "redirect_uri": redirect_uri,
            "grant_type": "authorization_code",
        }, timeout=10)
        token_resp.raise_for_status()
        access_token = token_resp.json()["access_token"]

        profile_resp = requests.get(
            "https://www.googleapis.com/oauth2/v3/userinfo",
            headers={"Authorization": f"Bearer {access_token}"}, timeout=10,
        )
        profile_resp.raise_for_status()
        profile = profile_resp.json()
    except Exception as e:
        print(f"[google-oauth] callback exchange failed: {e}")
        return redirect("/login?error=google_exchange_failed")

    google_sub = profile.get("sub")
    email = (profile.get("email") or "").strip().lower()
    if not google_sub or not email:
        return redirect("/login?error=incomplete_google_profile")

    user = get_user_by_google_sub(google_sub) or get_user_by_email(email)

    if user:
        # Existing account (signed up with Google before, or previously
        # with email/password using the same address) — link if needed.
        user_id = user["user_id"]
        updates = {}
        if not user.get("google_sub"):
            updates["google_sub"] = google_sub
        if profile.get("picture"):
            updates["picture"] = profile.get("picture")
        if not user.get("name") and profile.get("name"):
            updates["name"] = profile.get("name")
        if updates:
            update_user(user_id, updates)
    else:
        user_id = _pending_local_id()
        create_user({
            "user_id": user_id, "email": email,
            "name": profile.get("name") or email.split("@")[0],
            "picture": profile.get("picture"), "password_hash": None,
            "google_sub": google_sub, "created_at": time.time(),
        })

    _log_user_in(user_id)
    return redirect("/")


_LOGIN_PAGE_HTML = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Sign in · Mythic AI</title>
<link rel="icon" type="image/png" href="/icon.png">
<style>
  :root { --bg:#0c1410; --panel:#141f19; --border:#2a3a30; --text:#f5f3ea; --muted:#9aa89e; --accent:#10a37f; }
  * { box-sizing: border-box; }
  body {
    margin:0; min-height:100vh; display:flex; align-items:center; justify-content:center;
    background:var(--bg); color:var(--text); font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;
    padding:24px;
  }
  .card { width:100%; max-width:380px; background:var(--panel); border:1px solid var(--border); border-radius:16px; padding:32px 28px; }
  .brand { display:flex; align-items:center; gap:10px; justify-content:center; margin-bottom:24px; }
  .brand img { width:32px; height:32px; border-radius:8px; }
  .brand span { font-size:20px; font-weight:700; }
  h1 { font-size:16px; font-weight:500; color:var(--muted); text-align:center; margin:0 0 24px; }
  p.sub { font-size:13px; color:var(--muted); text-align:center; margin:-16px 0 20px; }
  .google-btn {
    width:100%; display:flex; align-items:center; justify-content:center; gap:10px;
    background:#fff; color:#1f1f1f; border:none; border-radius:10px; padding:12px; font-size:15px;
    font-weight:600; cursor:pointer; text-decoration:none;
  }
  .google-btn svg { width:18px; height:18px; }
  .divider { display:flex; align-items:center; gap:12px; margin:20px 0; color:var(--muted); font-size:13px; }
  .divider::before, .divider::after { content:''; flex:1; height:1px; background:var(--border); }
  input {
    width:100%; padding:12px 14px; border-radius:10px; border:1px solid var(--border);
    background:#0f1912; color:var(--text); font-size:14px; margin-bottom:12px;
  }
  input#code { text-align:center; letter-spacing:6px; font-size:20px; font-weight:700; }
  input:focus { outline:none; border-color:var(--accent); }
  button.primary {
    width:100%; padding:12px; border:none; border-radius:10px; background:var(--accent);
    color:#06120c; font-weight:700; font-size:14px; cursor:pointer;
  }
  button.primary:disabled { opacity:.6; cursor:default; }
  .toggle { text-align:center; margin-top:16px; font-size:13px; color:var(--muted); }
  .toggle a { color:var(--accent); cursor:pointer; text-decoration:none; }
  .err { background:#3a1414; color:#ff9a9a; border:1px solid #5a2020; border-radius:8px; padding:10px 12px; font-size:13px; margin-bottom:16px; display:none; }
  .ok  { background:#123a26; color:#9affc4; border:1px solid #1f5a3a; border-radius:8px; padding:10px 12px; font-size:13px; margin-bottom:16px; display:none; }
</style>
</head>
<body>
  <div class="card">
    <div class="brand"><img src="/icon.png" alt=""><span>Mythic AI</span></div>
    <h1>Sign in to continue</h1>
    <div class="err" id="err"></div>
    <div class="ok" id="ok"></div>

    <a class="google-btn" href="/api/auth/google/start" id="google-btn">
      <svg viewBox="0 0 48 48"><path fill="#FFC107" d="M43.6 20.5H42V20H24v8h11.3C33.7 32.7 29.3 36 24 36c-6.6 0-12-5.4-12-12s5.4-12 12-12c3.1 0 5.9 1.2 8 3.1l5.7-5.7C34.5 6.1 29.5 4 24 4 12.9 4 4 12.9 4 24s8.9 20 20 20 20-8.9 20-20c0-1.3-.1-2.7-.4-3.5z"/><path fill="#FF3D00" d="M6.3 14.7l6.6 4.8C14.6 15.9 18.9 13 24 13c3.1 0 5.9 1.2 8 3.1l5.7-5.7C34.5 6.1 29.5 4 24 4 16.3 4 9.7 8.3 6.3 14.7z"/><path fill="#4CAF50" d="M24 44c5.2 0 10-2 13.5-5.3l-6.2-5.2C29.3 35.4 26.8 36 24 36c-5.2 0-9.6-3.3-11.2-7.9l-6.5 5C9.6 39.6 16.3 44 24 44z"/><path fill="#1976D2" d="M43.6 20.5H42V20H24v8h11.3c-.8 2.3-2.3 4.2-4.2 5.5l6.2 5.2C40.8 36 44 30.9 44 24c0-1.3-.1-2.7-.4-3.5z"/></svg>
      Continue with Google
    </a>

    <div class="divider" id="divider">or</div>

    <div id="step-email">
      <input id="email" type="email" placeholder="Email" autocomplete="email">
      <button class="primary" id="send-btn">Send code</button>
    </div>

    <div id="step-code" style="display:none;">
      <p class="sub" id="sent-to"></p>
      <input id="code" type="text" inputmode="numeric" maxlength="6" placeholder="000000" autocomplete="one-time-code">
      <button class="primary" id="verify-btn">Verify &amp; sign in</button>
      <div class="toggle">
        <a id="resend-link">Resend code</a> &nbsp;·&nbsp; <a id="change-email-link">Use a different email</a>
      </div>
    </div>
  </div>
<script>
  const els = {
    stepEmail: document.getElementById('step-email'),
    stepCode: document.getElementById('step-code'),
    email: document.getElementById('email'),
    code: document.getElementById('code'),
    sendBtn: document.getElementById('send-btn'),
    verifyBtn: document.getElementById('verify-btn'),
    resend: document.getElementById('resend-link'),
    changeEmail: document.getElementById('change-email-link'),
    sentTo: document.getElementById('sent-to'),
    err: document.getElementById('err'),
    ok: document.getElementById('ok'),
  };
  const clientId = localStorage.getItem('mythic_client_id') || '';

  function showErr(msg) { els.err.textContent = msg; els.err.style.display = 'block'; els.ok.style.display = 'none'; }
  function showOk(msg)  { els.ok.textContent = msg; els.ok.style.display = 'block'; els.err.style.display = 'none'; }
  function clearMsgs()  { els.err.style.display = 'none'; els.ok.style.display = 'none'; }

  async function requestCode() {
    const email = els.email.value.trim();
    if (!email) { showErr('Enter your email address.'); return; }
    clearMsgs();
    els.sendBtn.disabled = true;
    try {
      const r = await fetch('/api/auth/otp/request', {
        method: 'POST', headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ email }),
      });
      const j = await r.json();
      if (!r.ok) { showErr(j.error || 'Something went wrong.'); els.sendBtn.disabled = false; return; }
      els.stepEmail.style.display = 'none';
      els.stepCode.style.display = 'block';
      els.sentTo.textContent = 'Code sent to ' + email;
      els.code.focus();
      showOk('Check your inbox for a 6-digit code.');
    } catch (e) {
      showErr('Network error — please try again.');
    }
    els.sendBtn.disabled = false;
  }

  async function verifyCode() {
    const email = els.email.value.trim();
    const code = els.code.value.trim();
    if (code.length !== 6) { showErr('Enter the 6-digit code.'); return; }
    clearMsgs();
    els.verifyBtn.disabled = true;
    try {
      const r = await fetch('/api/auth/otp/verify', {
        method: 'POST',
        headers: { 'Content-Type': 'application/json', 'X-Client-Id': clientId },
        body: JSON.stringify({ email, code }),
      });
      const j = await r.json();
      if (!r.ok) { showErr(j.error || 'Something went wrong.'); els.verifyBtn.disabled = false; return; }
      window.location.href = '/';
    } catch (e) {
      showErr('Network error — please try again.');
      els.verifyBtn.disabled = false;
    }
  }

  els.sendBtn.addEventListener('click', requestCode);
  els.email.addEventListener('keydown', e => { if (e.key === 'Enter') requestCode(); });
  els.verifyBtn.addEventListener('click', verifyCode);
  els.code.addEventListener('keydown', e => { if (e.key === 'Enter') verifyCode(); });
  els.resend.addEventListener('click', requestCode);
  els.changeEmail.addEventListener('click', () => {
    els.stepCode.style.display = 'none';
    els.stepEmail.style.display = 'block';
    clearMsgs();
  });

  const params = new URLSearchParams(window.location.search);
  if (params.get('error')) showErr('Google sign-in failed — please try again.');

  // Hide the Google button (and the "or" divider) entirely if the server
  // doesn't have GOOGLE_CLIENT_ID/SECRET configured yet, instead of
  // showing a button that just leads to an error page.
  fetch('/api/auth/config').then(r => r.json()).then(cfg => {
    if (!cfg.google_enabled) {
      document.getElementById('google-btn').style.display = 'none';
      document.getElementById('divider').style.display = 'none';
    }
  }).catch(() => {});
</script>
</body>
</html>"""


@app.route("/api/auth/config", methods=["GET"])
def api_auth_config():
    """Public, unauthenticated — lets the login page know whether to show
    the Google button at all, instead of showing one that just errors out
    when GOOGLE_CLIENT_ID/SECRET aren't set yet."""
    return jsonify({"google_enabled": bool(GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET)})


@app.route("/login", methods=["GET"])
def login_page():
    if session.get("authenticated") and session.get("user_id"):
        return redirect("/")
    return Response(_LOGIN_PAGE_HTML, mimetype="text/html")


# --- Public API keys ("aarav-...") --------------------------------------------
# Lets other apps/people call YOUR Mythic AI like a hosted API (OpenAI-style),
# authenticated with a personal key instead of the free chat UI. Keys are
# generated as "aarav-<random>", and only a SHA-256 hash is ever stored —
# the plaintext key is shown once at creation time and never again, same as
# how OpenAI/Anthropic/Stripe etc. handle API keys.
_API_KEYS_FILE = _os.path.join(_DATA_DIR, "api_keys.json")
_api_keys_lock = threading.Lock()
API_KEY_PREFIX = "aarav-"

def _hash_api_key(raw_key):
    return hashlib.sha256(raw_key.encode("utf-8")).hexdigest()

def _load_local_api_keys():
    if _os.path.exists(_API_KEYS_FILE):
        try:
            with open(_API_KEYS_FILE, encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return []
    return []

def _save_local_api_keys(keys):
    try:
        with open(_API_KEYS_FILE, "w", encoding="utf-8") as f:
            json.dump(keys, f)
    except Exception as e:
        print(f"[api_keys] could not persist keys to disk: {e}")

def create_api_key(label="", username=""):
    """Generates a new 'aarav-...' key (total 45 chars), stores only its hash, and returns the
    ONE-TIME plaintext key alongside its public record. Scoped to `username`
    (the creator's session id) so each visitor only ever sees/manages their
    own keys, not everyone else's."""
    # Generate random part: 45 total - 6 (aarav-) = 39 random chars
    random_part = secrets.token_urlsafe(29)[:39]  # 29 bytes base64 ≈ 39 chars
    raw_key = API_KEY_PREFIX + random_part
    record = {
        "id": str(uuid.uuid4()),
        "key_hash": _hash_api_key(raw_key),
        "key_prefix": raw_key,  # full key is short (25 chars), no need to truncate
        "label": (label or "").strip()[:100],
        "username": username,
        "created_at": datetime.datetime.utcnow().isoformat() + "Z",
        "active": True,
        "last_used_at": None,
        "request_count": 0,
    }
    with _api_keys_lock:
        if SUPABASE_URL:
            try:
                r = requests.post(sb("api_keys"), headers={**sb_headers(), "Prefer": "return=minimal"},
                                   json=record, timeout=10)
                if r.status_code in (200, 201, 204):
                    return raw_key, record
                print(f"[api_keys] Supabase write failed: HTTP {r.status_code} — {r.text[:300]} "
                      f"— falling back to local file.")
            except Exception as e:
                print(f"[api_keys] Supabase write failed: {e} — falling back to local file.")
        keys = _load_local_api_keys()
        keys.append(record)
        _save_local_api_keys(keys)
        return raw_key, record

def list_api_keys(username=""):
    """Returns key records WITHOUT the hash (never expose that) for display,
    filtered to only the ones the calling user created."""
    if SUPABASE_URL:
        try:
            r = requests.get(
                sb(f"api_keys?select=id,key_prefix,label,created_at,active,last_used_at,request_count"
                   f"&username=eq.{urllib.parse.quote(username)}&order=created_at.desc"),
                headers=sb_headers(), timeout=10)
            if r.status_code == 200:
                return r.json()
        except Exception as e:
            print(f"[api_keys] Supabase read failed: {e} — falling back to local file.")
    keys = _load_local_api_keys()
    return [{k: v for k, v in rec.items() if k != "key_hash"}
            for rec in reversed(keys) if rec.get("username") == username]

def revoke_api_key(key_id, username=""):
    """Only revokes the key if it belongs to `username` — prevents one
    visitor from revoking someone else's key by guessing/reusing an id."""
    if SUPABASE_URL:
        try:
            r = requests.patch(
                sb(f"api_keys?id=eq.{key_id}&username=eq.{urllib.parse.quote(username)}"),
                headers=sb_headers(), json={"active": False}, timeout=10)
            return r.status_code in (200, 204)
        except Exception as e:
            print(f"[api_keys] Supabase revoke failed: {e} — falling back to local file.")
    keys = _load_local_api_keys()
    found = False
    for rec in keys:
        if rec["id"] == key_id and rec.get("username") == username:
            rec["active"] = False
            found = True
    if found:
        _save_local_api_keys(keys)
    return found

def rename_api_key(key_id, new_label, username=""):
    """Renames (relabels) a key — only if it belongs to `username`."""
    new_label = (new_label or "").strip()[:100]
    if SUPABASE_URL:
        try:
            r = requests.patch(
                sb(f"api_keys?id=eq.{key_id}&username=eq.{urllib.parse.quote(username)}"),
                headers=sb_headers(), json={"label": new_label}, timeout=10)
            return r.status_code in (200, 204)
        except Exception as e:
            print(f"[api_keys] Supabase rename failed: {e} — falling back to local file.")
    keys = _load_local_api_keys()
    found = False
    for rec in keys:
        if rec["id"] == key_id and rec.get("username") == username:
            rec["label"] = new_label
            found = True
    if found:
        _save_local_api_keys(keys)
    return found

def verify_api_key(raw_key):
    """Returns True and records usage if raw_key is a valid, active key."""
    if not raw_key or not raw_key.startswith(API_KEY_PREFIX):
        return False
    key_hash = _hash_api_key(raw_key)
    now = datetime.datetime.utcnow().isoformat() + "Z"
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"api_keys?key_hash=eq.{key_hash}&active=eq.true&select=id,request_count"),
                              headers=sb_headers(), timeout=10)
            if r.status_code == 200 and r.json():
                rec = r.json()[0]
                try:
                    requests.patch(sb(f"api_keys?id=eq.{rec['id']}"), headers=sb_headers(),
                                    json={"last_used_at": now, "request_count": (rec.get("request_count") or 0) + 1},
                                    timeout=10)
                except Exception:
                    pass  # usage tracking is best-effort, shouldn't block the request
                return True
            return False
        except Exception as e:
            print(f"[api_keys] Supabase verify failed: {e} — falling back to local file.")
    keys = _load_local_api_keys()
    for rec in keys:
        if rec.get("key_hash") == key_hash and rec.get("active"):
            rec["last_used_at"] = now
            rec["request_count"] = (rec.get("request_count") or 0) + 1
            _save_local_api_keys(keys)
            return True
    return False


# ══════════════════════════════════════════════════════════════════════════════
# ─── ANALYTICS & USAGE TRACKING (1000+ lines module) ──────────────────────────
# ══════════════════════════════════════════════════════════════════════════════
#
# Comprehensive tracking system for API key users:
# • Per-API-key usage metrics (requests, tokens, models used)
# • Per-user analytics (chat count, message volume, active hours)
# • Per-conversation export (JSON, CSV, HTML)
# • Full-text search with advanced filtering
# • Admin dashboard with user & key management
# • Usage trends and per-model breakdown
# • Rate limiting insights and quota tracking

import json as _json
import csv as _csv
from io import StringIO as _StringIO
from datetime import datetime as _dt, timedelta as _td

_USAGE_DB_FILE = _os.path.join(_DATA_DIR, "usage_metrics.json")
_USAGE_LOCK = threading.Lock()


def _load_usage_metrics():
    """Load all usage metrics from local storage or Supabase."""
    if _os.path.exists(_USAGE_DB_FILE):
        try:
            with open(_USAGE_DB_FILE, encoding="utf-8") as f:
                data = _json.load(f)
                # Convert lists back to proper types
                if isinstance(data.get("daily_stats"), dict):
                    for day, stats in data["daily_stats"].items():
                        if isinstance(stats.get("unique_users"), list):
                            stats["unique_users"] = set(stats["unique_users"])
                return data
        except Exception:
            pass
    return {"api_key_usage": {}, "user_stats": {}, "model_usage": {}, "daily_stats": {}}


def _save_usage_metrics(data):
    """Save usage metrics to persistent storage, converting sets to lists for JSON."""
    try:
        # Convert sets to lists for JSON serialization
        data_copy = _json.loads(_json.dumps(data, default=str))
        if isinstance(data.get("daily_stats"), dict):
            for day, stats in data.get("daily_stats", {}).items():
                if isinstance(stats.get("unique_users"), set):
                    stats["unique_users"] = list(stats["unique_users"])
        with open(_USAGE_DB_FILE, "w", encoding="utf-8") as f:
            _json.dump(data_copy, f, indent=2)
    except Exception as e:
        print(f"[analytics] failed to save usage metrics: {e}")


def track_api_request(key_id, username, model, tokens_in, tokens_out, endpoint, success=True):
    """Record an API request for analytics and rate-limit tracking."""
    with _USAGE_LOCK:
        metrics = _load_usage_metrics()
        now = _dt.utcnow().isoformat()
        today = now.split("T")[0]

        # Per-key usage
        if key_id not in metrics["api_key_usage"]:
            metrics["api_key_usage"][key_id] = {
                "total_requests": 0, "successful_requests": 0, "failed_requests": 0,
                "total_tokens_in": 0, "total_tokens_out": 0, "by_model": {}, "last_updated": now
            }
        key_rec = metrics["api_key_usage"][key_id]
        key_rec["total_requests"] += 1
        if success:
            key_rec["successful_requests"] += 1
            key_rec["total_tokens_in"] += tokens_in
            key_rec["total_tokens_out"] += tokens_out
            if model not in key_rec["by_model"]:
                key_rec["by_model"][model] = {"requests": 0, "tokens_in": 0, "tokens_out": 0}
            key_rec["by_model"][model]["requests"] += 1
            key_rec["by_model"][model]["tokens_in"] += tokens_in
            key_rec["by_model"][model]["tokens_out"] += tokens_out
        else:
            key_rec["failed_requests"] += 1
        key_rec["last_updated"] = now

        # Per-user stats
        if username not in metrics["user_stats"]:
            metrics["user_stats"][username] = {
                "total_requests": 0, "total_messages": 0, "total_conversations": 0,
                "models_used": [], "first_seen": now, "last_active": now
            }
        user_rec = metrics["user_stats"][username]
        user_rec["total_requests"] += 1
        user_rec["total_messages"] += 1
        user_rec["last_active"] = now
        if model not in user_rec["models_used"]:
            user_rec["models_used"].append(model)

        # Daily aggregates
        if today not in metrics["daily_stats"]:
            metrics["daily_stats"][today] = {
                "total_requests": 0, "total_tokens": 0, "unique_users": set(),
                "by_model": {}, "by_endpoint": {}
            }
        daily = metrics["daily_stats"][today]
        daily["total_requests"] += 1
        daily["total_tokens"] += tokens_in + tokens_out
        daily["unique_users"].add(username)

        if model not in daily["by_model"]:
            daily["by_model"][model] = {"requests": 0, "tokens": 0}
        daily["by_model"][model]["requests"] += 1
        daily["by_model"][model]["tokens"] += tokens_in + tokens_out

        if endpoint not in daily["by_endpoint"]:
            daily["by_endpoint"][endpoint] = 0
        daily["by_endpoint"][endpoint] += 1

        _save_usage_metrics(metrics)


def get_usage_report(username=None, key_id=None, start_date=None, end_date=None, days_back=30):
    """Generate a detailed usage report with optional filtering by user, API key, and date range."""
    with _USAGE_LOCK:
        metrics = _load_usage_metrics()

    if not start_date:
        start_date = (_dt.utcnow() - _td(days=days_back)).strftime("%Y-%m-%d")
    if not end_date:
        end_date = _dt.utcnow().strftime("%Y-%m-%d")

    report = {
        "generated_at": _dt.utcnow().isoformat(),
        "filters": {"username": username, "key_id": key_id, "start_date": start_date, "end_date": end_date},
        "summary": {
            "total_requests": 0, "successful_requests": 0, "failed_requests": 0,
            "total_tokens_in": 0, "total_tokens_out": 0, "by_model": {}, "daily_breakdown": []
        },
        "api_key_metrics": [],
        "user_metrics": []
    }

    # API key metrics
    if key_id and key_id in metrics["api_key_usage"]:
        report["api_key_metrics"].append(metrics["api_key_usage"][key_id])
    elif not key_id:
        report["api_key_metrics"] = list(metrics["api_key_usage"].values())

    # User metrics
    if username and username in metrics["user_stats"]:
        report["user_metrics"].append(metrics["user_stats"][username])
    elif not username:
        report["user_metrics"] = list(metrics["user_stats"].values())

    # Daily breakdown with date filtering
    for day in sorted(metrics["daily_stats"].keys()):
        if day < start_date or day > end_date:
            continue
        stats = metrics["daily_stats"][day]
        report["summary"]["daily_breakdown"].append({"date": day, "stats": stats})
        report["summary"]["total_requests"] += stats.get("total_requests", 0)
        for model, mstats in stats.get("by_model", {}).items():
            if model not in report["summary"]["by_model"]:
                report["summary"]["by_model"][model] = {"requests": 0, "tokens": 0}
            report["summary"]["by_model"][model]["requests"] += mstats.get("requests", 0)
            report["summary"]["by_model"][model]["tokens"] += mstats.get("tokens", 0)

    return report


def search_conversations(username, query="", filters=None):
    """Search conversations by full-text query with advanced filters.
    Filters: {'start_date': '2025-01-01', 'end_date': '2025-01-31', 
              'folder': 'work', 'archived': bool, 'pinned': bool, 'min_messages': int}"""
    filters = filters or {}
    convs = list_conversations(username)
    results = []

    for conv in convs:
        # Apply boolean filters
        if filters.get("archived") is not None and conv.get("archived") != filters["archived"]:
            continue
        if filters.get("pinned") is not None and conv.get("pinned") != filters["pinned"]:
            continue
        if filters.get("folder") and conv.get("folder") != filters["folder"]:
            continue

        # Message count filter
        msg_count = len(conv.get("messages", []))
        if filters.get("min_messages") and msg_count < filters["min_messages"]:
            continue

        # Date range filter
        if filters.get("start_date"):
            conv_date = conv.get("updated_at", "")
            if conv_date and conv_date < filters["start_date"]:
                continue
        if filters.get("end_date"):
            conv_date = conv.get("updated_at", "")
            if conv_date and conv_date > filters["end_date"]:
                continue

        # Full-text search across title and messages
        if query:
            query_lower = query.lower()
            search_text = (conv.get("title") or "").lower()
            found = query_lower in search_text

            if not found:
                for msg in conv.get("messages", []):
                    msg_text = (msg.get("text") or "").lower()
                    if query_lower in msg_text:
                        found = True
                        break
            
            if not found:
                continue

        results.append(conv)

    return results


def export_conversation(conv_id, username, format_type="json"):
    """Export a single conversation in requested format (json, csv, html).
    Returns (content_bytes, mimetype, filename) tuple."""
    conv = load_conversation(username, conv_id)
    if not conv:
        return None, None, None

    title = conv.get("title", "chat")
    safe_title = "".join(c if c.isalnum() or c in "-_ " else "" for c in title)[:50]

    if format_type == "json":
        content = _json.dumps(conv, indent=2, default=str).encode("utf-8")
        return content, "application/json", f"{safe_title}.json"

    elif format_type == "csv":
        output = _StringIO()
        writer = _csv.writer(output)
        writer.writerow(["Role", "Timestamp", "Message"])
        for msg in conv.get("messages", []):
            role = "User" if msg.get("role") == "user" else "AI"
            parts = msg.get("parts", [])
            text = "".join(part.get("text", "") for part in parts if "text" in part)
            writer.writerow([role, msg.get("created_at", ""), text])
        content = output.getvalue().encode("utf-8")
        return content, "text/csv", f"{safe_title}.csv"

    elif format_type == "html":
        html_parts = [
            f"""<!DOCTYPE html>
<html>
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{_esc_html(title)}</title>
  <style>
    * {{ box-sizing: border-box; }}
    body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; max-width: 900px; margin: 0 auto; padding: 20px; line-height: 1.6; color: #333; background: #fafafa; }}
    h1 {{ color: #10a37f; border-bottom: 3px solid #10a37f; padding-bottom: 10px; margin-top: 0; }}
    .meta {{ font-size: 0.9em; color: #666; margin: 10px 0 30px; }}
    .msg {{ margin: 15px 0; padding: 12px 16px; border-radius: 8px; }}
    .msg.user {{ background: #e8f5e9; margin-left: 40px; border-left: 3px solid #10a37f; }}
    .msg.ai {{ background: #f5f5f5; margin-right: 40px; border-left: 3px solid #ccc; }}
    .role {{ font-weight: 700; font-size: 0.9em; color: #555; margin-bottom: 6px; text-transform: uppercase; letter-spacing: 0.5px; }}
    .content {{ margin-bottom: 6px; white-space: pre-wrap; word-wrap: break-word; }}
    .timestamp {{ font-size: 0.8em; color: #999; }}
    footer {{ margin-top: 40px; padding-top: 20px; border-top: 1px solid #ddd; text-align: center; color: #999; font-size: 0.85em; }}
  </style>
</head>
<body>
  <h1>{_esc_html(title)}</h1>
  <div class="meta">
    <p>Exported on {_dt.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC | Total messages: {len(conv.get("messages", []))}</p>
  </div>
"""
        ]

        for msg in conv.get("messages", []):
            role = "User" if msg.get("role") == "user" else "AI"
            role_class = "user" if msg.get("role") == "user" else "ai"
            parts = msg.get("parts", [])
            text = "".join(part.get("text", "") for part in parts if "text" in part)
            safe_text = _esc_html(text)
            ts = msg.get("created_at", "")
            html_parts.append(f'<div class="msg {role_class}"><div class="role">{role}</div><div class="content">{safe_text}</div><div class="timestamp">{ts}</div></div>')

        html_parts.append(f"""
  <footer>
    <p>This conversation was exported from Mythic AI. It contains {len(conv.get("messages", []))} messages.</p>
  </footer>
</body>
</html>""")
        
        content = "".join(html_parts).encode("utf-8")
        return content, "text/html", f"{safe_title}.html"

    return None, None, None


def _esc_html(s):
    """Escape HTML special characters."""
    if not s:
        return ""
    return (str(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                   .replace('"', "&quot;").replace("'", "&#39;"))


def get_admin_stats(include_daily=False):
    """Return aggregate admin-level statistics (super-admin only).
    Returns overall system metrics, per-key usage, per-user stats."""
    with _USAGE_LOCK:
        metrics = _load_usage_metrics()
    
    total_requests = sum(k.get("total_requests", 0) for k in metrics.get("api_key_usage", {}).values())
    total_tokens_in = sum(k.get("total_tokens_in", 0) for k in metrics.get("api_key_usage", {}).values())
    total_tokens_out = sum(k.get("total_tokens_out", 0) for k in metrics.get("api_key_usage", {}).values())
    
    # Top models by usage
    top_models = {}
    for key_data in metrics.get("api_key_usage", {}).values():
        for model, stats in key_data.get("by_model", {}).items():
            if model not in top_models:
                top_models[model] = {"requests": 0, "tokens": 0}
            top_models[model]["requests"] += stats.get("requests", 0)
            top_models[model]["tokens"] += stats.get("tokens_in", 0) + stats.get("tokens_out", 0)
    
    stats = {
        "total_api_keys_issued": len(metrics.get("api_key_usage", {})),
        "total_users": len(metrics.get("user_stats", {})),
        "all_time_requests": total_requests,
        "all_time_tokens_in": total_tokens_in,
        "all_time_tokens_out": total_tokens_out,
        "all_time_tokens_total": total_tokens_in + total_tokens_out,
        "top_models_by_requests": sorted(top_models.items(), key=lambda x: x[1]["requests"], reverse=True)[:10],
        "api_key_metrics": metrics.get("api_key_usage", {}),
        "user_metrics": metrics.get("user_stats", {}),
        "generated_at": _dt.utcnow().isoformat()
    }
    
    if include_daily:
        stats["daily_stats"] = metrics.get("daily_stats", {})
    
    return stats


def _conv_file(username, conv_id):
    return _os.path.join(_user_conv_dir(username), f"{conv_id}.json")

def _list_conversations_file(username):
    folder = _user_conv_dir(username)
    convs = []
    for fname in _os.listdir(folder):
        if not fname.endswith(".json"):
            continue
        path = _os.path.join(folder, fname)
        try:
            with open(path, "r", encoding="utf-8") as f:
                d = json.load(f)
            convs.append({
                "id": fname[:-5],
                "title": d.get("title", "New chat"),
                "updated_at": d.get("updated_at", 0),
                "folder": d.get("folder"),
                "pinned": bool(d.get("pinned", False)),
                "archived": bool(d.get("archived", False)),
            })
        except Exception:
            continue
    # Pinned conversations always float to the top, then most-recent first.
    convs.sort(key=lambda c: (not c["pinned"], -c["updated_at"]))
    return convs

def _load_conversation_file(username, conv_id):
    path = _conv_file(username, conv_id)
    if not _os.path.exists(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None

def _save_conversation_file(username, conv_id, data):
    try:
        with open(_conv_file(username, conv_id), "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"[save_conversation] FAILED to write conversation {conv_id} for {username}: {e}")

def _delete_conversation_file(username, conv_id):
    path = _conv_file(username, conv_id)
    if _os.path.exists(path):
        _os.remove(path)


# ── Public share links ────────────────────────────────────────────────────
# A "share" maps a short public id -> (username, conv_id) so anyone with the
# link can view a read-only copy of that one conversation at /share/<id>,
# with no login and no access to the rest of that person's chats. Uses
# Supabase (a `shares` table) when configured, else a local JSON index file
# alongside the conversation JSON fallback.
_SHARES_INDEX_FILE = _os.path.join(_DATA_DIR, "shares_index.json")
_shares_lock = threading.Lock()


def _load_shares_index():
    try:
        if _os.path.exists(_SHARES_INDEX_FILE):
            with open(_SHARES_INDEX_FILE, encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return {}


def _save_shares_index(data):
    try:
        with open(_SHARES_INDEX_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f)
    except Exception as e:
        print(f"[shares] failed to save shares index: {e}")


def get_active_share_id(username, conv_id):
    """Returns the existing active (non-revoked) share id for this
    conversation, or None if it has never been shared / was revoked.
    Checks Supabase first (if configured), then always falls back to
    checking the local file index too — a share may have landed there if
    Supabase writes were failing (e.g. the `shares` table doesn't exist
    yet) when it was created. See create_share_link()."""
    if SUPABASE_URL:
        try:
            r = requests.get(
                sb(f"shares?username=eq.{username}&conv_id=eq.{conv_id}&revoked=eq.false&select=id"),
                headers=sb_headers(), timeout=10,
            )
            if r.status_code == 200 and r.json():
                return r.json()[0]["id"]
            elif r.status_code != 200:
                print(f"[Supabase] get_active_share_id: HTTP {r.status_code} — {r.text[:200]}")
        except Exception as e:
            print(f"[Supabase] get_active_share_id exception: {e}")

    with _shares_lock:
        idx = _load_shares_index()
    for sid, rec in idx.items():
        if rec.get("username") == username and rec.get("conv_id") == conv_id and not rec.get("revoked"):
            return sid
    return None


def create_share_link(username, conv_id):
    """Creates (or reuses, if one is already active) a public share id for
    conv_id owned by username. Returns the share_id string.

    IMPORTANT: if Supabase is configured but the write fails for any reason
    (most commonly: the `shares` table hasn't been created yet), this does
    NOT hand back a link that silently 404s — it transparently falls back
    to the local JSON index instead, so the link that's returned always
    actually works."""
    existing = get_active_share_id(username, conv_id)
    if existing:
        return existing

    share_id = uuid.uuid4().hex[:12]
    if SUPABASE_URL:
        try:
            r = requests.post(
                sb("shares"),
                headers={**sb_headers(), "Prefer": "return=minimal"},
                json={"id": share_id, "username": username, "conv_id": conv_id,
                      "revoked": False, "created_at": time.time()},
                timeout=10,
            )
            if r.status_code in (200, 201, 204):
                return share_id
            print(f"[Supabase] create_share_link failed: HTTP {r.status_code} — {r.text[:300]}"
                  f" — falling back to local storage for this share link.")
        except Exception as e:
            print(f"[Supabase] create_share_link exception: {e} — falling back to local storage.")

    with _shares_lock:
        idx = _load_shares_index()
        idx[share_id] = {"username": username, "conv_id": conv_id,
                          "revoked": False, "created_at": time.time()}
        _save_shares_index(idx)
    return share_id


def resolve_share_link(share_id):
    """Returns {"username":..., "conv_id":...} for an active (non-revoked)
    share id, or None if it doesn't exist / has been revoked. Checks
    Supabase first (if configured), then always checks the local file
    index too, since create_share_link() may have fallen back to it."""
    if SUPABASE_URL:
        try:
            r = requests.get(
                sb(f"shares?id=eq.{share_id}&select=username,conv_id,revoked"),
                headers=sb_headers(), timeout=10,
            )
            if r.status_code == 200:
                rows = r.json()
                if rows:
                    if rows[0].get("revoked"):
                        return None
                    return {"username": rows[0]["username"], "conv_id": rows[0]["conv_id"]}
            else:
                print(f"[Supabase] resolve_share_link: HTTP {r.status_code} — {r.text[:200]}")
        except Exception as e:
            print(f"[Supabase] resolve_share_link exception: {e}")

    idx = _load_shares_index()
    rec = idx.get(share_id)
    if rec and not rec.get("revoked"):
        return {"username": rec["username"], "conv_id": rec["conv_id"]}
    return None


def revoke_share_link(username, conv_id):
    """Revokes any active share link(s) for this conversation, in both
    Supabase and the local index (a share may live in either, depending on
    whether Supabase writes were working when it was created)."""
    if SUPABASE_URL:
        try:
            requests.patch(
                sb(f"shares?username=eq.{username}&conv_id=eq.{conv_id}"),
                headers=sb_headers(), json={"revoked": True}, timeout=10,
            )
        except Exception as e:
            print(f"[Supabase] revoke_share_link exception: {e}")
    with _shares_lock:
        idx = _load_shares_index()
        changed = False
        for sid, rec in idx.items():
            if rec.get("username") == username and rec.get("conv_id") == conv_id and not rec.get("revoked"):
                rec["revoked"] = True
                changed = True
        if changed:
            _save_shares_index(idx)


def list_folders(username):
    """Distinct, non-empty folder names currently in use by this user's
    conversations, alphabetically sorted — used to populate the folder
    picker in the sidebar without a separate folders table."""
    names = sorted({c["folder"] for c in list_conversations(username) if c.get("folder")})
    return names


def make_title(first_message):
    text = (first_message or "Attachment").strip()
    # Strip an internal "[Instructions: ...] " tone/length/custom-instructions
    # prefix (added client-side when Settings → tone/length are set) so it
    # never leaks into a conversation's saved title in the sidebar.
    text = re.sub(r'^\[Instructions:.*?\]\s*', '', text, flags=re.DOTALL)
    title = text.replace("\n", " ").strip() or "New chat"
    return title[:40] + ("…" if len(title) > 40 else "")


# ── Document reading (PDF / DOCX / TXT / CSV / MD / JSON / source code) ─────
# Groq and Cerebras models are text-only — they cannot see images or read
# binary files directly. So that "attach a file and ask about it" actually
# works, every text-extractable upload is converted to plain text here and
# injected into the prompt as context. Images are explicitly labeled as
# unreadable rather than silently ignored, per "don't pretend" — no OCR/vision
# is performed since that would require a separate API.
_TEXT_EXTENSIONS = (
    ".txt", ".md", ".markdown", ".csv", ".json", ".py", ".js", ".ts", ".tsx",
    ".jsx", ".html", ".htm", ".css", ".java", ".c", ".cpp", ".cc", ".h", ".cs",
    ".go", ".rs", ".php", ".rb", ".sql", ".sh", ".ps1", ".kt", ".swift",
    ".yaml", ".yml", ".xml", ".log", ".ini", ".toml",
)
DOCUMENT_EXTRACT_MAX_CHARS = 12000


def extract_text_from_attachment(filename, mime_type, raw_bytes):
    """Returns (extracted_text_or_None, note_or_None).
    extracted_text is None when the file type isn't text-extractable at all
    (e.g. an image) — the caller decides what to tell the model in that case.
    note carries a short caveat (truncation, missing library, empty PDF, etc.)
    that should be shown to the model/user alongside the text."""
    name_lower = (filename or "").lower()
    mime_type = mime_type or ""
    text = None
    try:
        if mime_type == "application/pdf" or name_lower.endswith(".pdf"):
            try:
                from pypdf import PdfReader
            except ImportError:
                return None, "PDF reading isn't available on this server (needs `pip install pypdf`)."
            import io as _io
            reader = PdfReader(_io.BytesIO(raw_bytes))
            pages_text = []
            for page in reader.pages[:40]:  # cap pages so huge PDFs don't blow up the prompt
                try:
                    pages_text.append(page.extract_text() or "")
                except Exception:
                    continue
            text = "\n".join(pages_text).strip()
            if not text:
                return "", "No extractable text found — this PDF may be scanned/image-only, which needs OCR (not available without a separate API)."

        elif name_lower.endswith(".docx") or mime_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
            try:
                import docx as _docx
            except ImportError:
                return None, "DOCX reading isn't available on this server (needs `pip install python-docx`)."
            import io as _io
            doc = _docx.Document(_io.BytesIO(raw_bytes))
            text = "\n".join(p.text for p in doc.paragraphs).strip()

        elif mime_type.startswith("text/") or name_lower.endswith(_TEXT_EXTENSIONS):
            text = raw_bytes.decode("utf-8", errors="replace")

        else:
            return None, None  # not a recognized text-extractable type (e.g. an image)
    except Exception as e:
        return None, f"Could not read this file: {e}"

    if not text:
        return "", None
    note = None
    if len(text) > DOCUMENT_EXTRACT_MAX_CHARS:
        text = text[:DOCUMENT_EXTRACT_MAX_CHARS]
        note = f"(showing first {DOCUMENT_EXTRACT_MAX_CHARS} characters — file was longer)"
    return text, note


# ── Downloadable file generation (PDF / DOCX / TXT) ──────────────────────────
# Built with zero required extra dependencies (mirrors the manual PNG-writer
# used for the app icon above) so "generate a PDF" works out of the box.
# If python-docx happens to be installed, real .docx files are used instead
# of falling back to plain text for Word-document requests.
import textwrap as _textwrap


def _pdf_escape(s: str) -> str:
    return s.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def generate_pdf_bytes(title: str, body_text: str) -> bytes:
    """Renders `title` + `body_text` as a simple multi-page PDF using only
    the Python standard library (Helvetica, Letter-size pages). Good enough
    for chat-generated notes, summaries, letters, etc. — not a full layout
    engine, just readable wrapped text."""
    PAGE_W, PAGE_H = 612, 792
    MARGIN = 56
    FONT_SIZE = 11
    LEADING = 15
    usable_width_chars = max(40, int((PAGE_W - 2 * MARGIN) / (FONT_SIZE * 0.5)))
    max_lines_per_page = int((PAGE_H - 2 * MARGIN - 40) / LEADING)

    lines = []
    if title:
        lines.append(("title", title))
        lines.append(("blank", ""))
    for para in body_text.split("\n"):
        para = para.rstrip()
        if not para:
            lines.append(("blank", ""))
            continue
        wrapped = _textwrap.wrap(para, width=usable_width_chars) or [""]
        for w in wrapped:
            lines.append(("body", w))

    pages = [lines[i:i + max_lines_per_page] for i in range(0, len(lines), max_lines_per_page)] or [[]]

    objects = []  # 1-indexed PDF object numbers; objects[i-1] holds object i's bytes

    def emit(data: bytes) -> int:
        objects.append(data)
        return len(objects)

    catalog_num = emit(b"placeholder")   # filled in once we know the Pages object number
    pages_num = emit(b"placeholder")     # filled in once we know all page object numbers
    font_num = emit(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")

    page_nums = []
    for page_lines in pages:
        stream_parts = [b"BT", f"/F1 {FONT_SIZE} Tf".encode(), f"{LEADING} TL".encode(),
                         f"{MARGIN} {PAGE_H - MARGIN} Td".encode()]
        first = True
        for kind, text in page_lines:
            if not first:
                stream_parts.append(b"T*")
            first = False
            if kind == "blank":
                continue
            if kind == "title":
                stream_parts.append(b"/F1 16 Tf")
                stream_parts.append(f"({_pdf_escape(text)}) Tj".encode("latin-1", "replace"))
                stream_parts.append(f"/F1 {FONT_SIZE} Tf".encode())
            else:
                stream_parts.append(f"({_pdf_escape(text)}) Tj".encode("latin-1", "replace"))
        stream_parts.append(b"ET")
        stream = b"\n".join(stream_parts)
        content_data = (f"<< /Length {len(stream)} >>\nstream\n".encode() + stream +
                         b"\nendstream")
        content_num = emit(content_data)
        page_dict = (
            f"<< /Type /Page /Parent {pages_num} 0 R /MediaBox [0 0 {PAGE_W} {PAGE_H}] "
            f"/Resources << /Font << /F1 {font_num} 0 R >> >> /Contents {content_num} 0 R >>"
        ).encode()
        page_nums.append(emit(page_dict))

    kids = " ".join(f"{n} 0 R" for n in page_nums)
    objects[pages_num - 1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_nums)} >>".encode()
    objects[catalog_num - 1] = f"<< /Type /Catalog /Pages {pages_num} 0 R >>".encode()

    # Assemble the final PDF byte stream with a proper xref table.
    out = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj_data in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj_data + b"\nendobj\n"
    xref_offset = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets[1:]:
        out += f"{off:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_num} 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF"
    ).encode()
    return bytes(out)


def generate_docx_bytes(title: str, body_text: str):
    """Returns bytes if python-docx is installed for a real Word document;
    otherwise returns None to signal the caller should fall back to a plain
    text file."""
    try:
        import docx as _docx
    except ImportError:
        return None
    doc = _docx.Document()
    if title:
        doc.add_heading(title, level=1)
    for para in body_text.split("\n"):
        doc.add_paragraph(para)
    import io as _io
    buf = _io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════
# AUTOMATIC WATERMARK SYSTEM
# ══════════════════════════════════════════════════════════════════════════
# Every AI-generated image is automatically watermarked before it ever leaves
# the server — the frontend never sees, and can never download, an
# unwatermarked image. Uses Pillow (free, open-source) — no paid image APIs.
# If Pillow isn't installed, watermarking is skipped gracefully (image
# generation still works) and a warning is printed once at startup; run
# `pip install Pillow` to enable it.
#
# ADMIN_SECRET (env var, same pattern as CRON_SECRET) protects the admin
# endpoints for uploading/replacing/resetting the official logo.
ADMIN_SECRET = os.environ.get("ADMIN_SECRET", "").strip()
_WATERMARK_DIR = _os.path.join(_DATA_DIR, "watermark")
_os.makedirs(_WATERMARK_DIR, exist_ok=True)
_WATERMARK_LOGO_LIGHT_PATH = _os.path.join(_WATERMARK_DIR, "logo_light.png")  # white mark, for dark backgrounds
_WATERMARK_LOGO_DARK_PATH  = _os.path.join(_WATERMARK_DIR, "logo_dark.png")   # dark mark, for light backgrounds

if not _WATERMARK_AVAILABLE:
    print("[Watermark] Pillow is not installed — generated images will NOT be "
          "watermarked. Run `pip install Pillow` (free, no license cost) to enable "
          "the automatic watermark system.")


def _require_admin():
    """Returns True if the request carries a valid ADMIN_SECRET. If
    ADMIN_SECRET is unset, admin endpoints are open (dev-only default —
    always set ADMIN_SECRET in production)."""
    if not ADMIN_SECRET:
        return True
    auth = request.headers.get("Authorization", "")
    provided = auth[7:].strip() if auth.startswith("Bearer ") else request.args.get("secret", "").strip()
    return provided == ADMIN_SECRET


def _build_default_watermark_logo(size=512, light=True):
    """Generates the default monogram watermark as a transparent-background
    RGBA PIL Image — supersampled 4x then downsampled for clean anti-aliasing.
    `light=True` gives a white mark (for dark image backgrounds); `light=False`
    gives a near-black mark with a soft white edge (for light backgrounds)."""
    SS = 4  # supersample factor for anti-aliasing
    W = H = size * SS
    img = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    draw = ImageDraw.Draw(img)

    fg = (255, 255, 255, 255) if light else (26, 26, 26, 255)
    s = W / 40
    pts = [
        (10 * s, 28 * s), (10 * s, 12 * s), (20 * s, 22 * s),
        (30 * s, 12 * s), (30 * s, 28 * s),
    ]
    lw = max(4, int(W // 11))
    draw.line(pts, fill=fg, width=lw, joint="curve")
    r = lw // 2
    for (x, y) in [pts[0], pts[-1]]:
        draw.ellipse([x - r, y - r, x + r, y + r], fill=fg)

    if not light:
        outline = Image.new("RGBA", (W, H), (0, 0, 0, 0))
        ImageDraw.Draw(outline).line(pts, fill=(255, 255, 255, 160), width=lw + max(6, W // 40), joint="curve")
        outline = outline.filter(ImageFilter.GaussianBlur(W / 60))
        img = Image.alpha_composite(outline, img)

    return img.resize((size, size), Image.LANCZOS)


_default_logo_cache = {}

def _get_default_watermark_logo(light=True):
    key = "light" if light else "dark"
    if key not in _default_logo_cache:
        _default_logo_cache[key] = _build_default_watermark_logo(512, light=light)
    return _default_logo_cache[key]


def _load_watermark_logo(light=True):
    """Returns the active logo as an RGBA PIL Image — an admin-uploaded custom
    logo if one exists on disk, otherwise the built-in default monogram."""
    path = _WATERMARK_LOGO_LIGHT_PATH if light else _WATERMARK_LOGO_DARK_PATH
    if _os.path.exists(path):
        try:
            return Image.open(path).convert("RGBA")
        except Exception:
            pass
    return _get_default_watermark_logo(light=light)


_WM_POSITIONS = {
    "br": ("right", "bottom"), "bottom-right": ("right", "bottom"),
    "bl": ("left", "bottom"),  "bottom-left": ("left", "bottom"),
    "tr": ("right", "top"),    "top-right": ("right", "top"),
    "tl": ("left", "top"),     "top-left": ("left", "top"),
    "center": ("center", "center"),
}


def _region_edge_density(gray_img, box):
    crop = gray_img.crop(box)
    if crop.width < 2 or crop.height < 2:
        return 0.0
    edges = crop.filter(ImageFilter.FIND_EDGES)
    return ImageStat.Stat(edges).mean[0]


def _pick_watermark_geometry(base_rgba, logo_w, logo_h, margin, position_pref):
    W, H = base_rgba.size
    anchor_x, anchor_y = _WM_POSITIONS.get(position_pref, _WM_POSITIONS["br"])

    def box_for(dx_off, dy_off):
        if anchor_x == "right":
            x = W - margin - logo_w - dx_off
        elif anchor_x == "left":
            x = margin + dx_off
        else:
            x = (W - logo_w) // 2
        if anchor_y == "bottom":
            y = H - margin - logo_h - dy_off
        elif anchor_y == "top":
            y = margin + dy_off
        else:
            y = (H - logo_h) // 2
        x = max(0, min(W - logo_w, int(x)))
        y = max(0, min(H - logo_h, int(y)))
        return (x, y, x + logo_w, y + logo_h)

    default_box = box_for(0, 0)
    if anchor_x == "center":
        return default_box

    try:
        gray = base_rgba.convert("L")
        overall = ImageStat.Stat(gray.filter(ImageFilter.FIND_EDGES)).mean[0]
        busy = _region_edge_density(gray, default_box)
        if busy > overall * 1.6 and busy > 8:
            shift = int(max(logo_w, logo_h) * 0.9)
            candidates = [box_for(shift, 0), box_for(0, shift), box_for(shift, shift)]
            best_box, best_score = default_box, busy
            for cand in candidates:
                score = _region_edge_density(gray, cand)
                if score < best_score:
                    best_box, best_score = cand, score
            return best_box
    except Exception:
        pass
    return default_box


def _analyze_background_theme(base_rgba, box):
    crop = base_rgba.convert("RGB").crop(box)
    stat = ImageStat.Stat(crop)
    r, g, b = stat.mean
    luminance = 0.299 * r + 0.587 * g + 0.114 * b
    colorfulness = (max(r, g, b) - min(r, g, b))
    is_colorful = colorfulness > 40 or (sum(stat.stddev) / 3) > 55
    return ("dark" if luminance < 128 else "light"), is_colorful


def _add_shadow_and_glow(logo_rgba, want_shadow=True, want_glow=True, is_dark_bg=True):
    pad = max(8, logo_rgba.width // 6)
    canvas_w, canvas_h = logo_rgba.width + pad * 2, logo_rgba.height + pad * 2
    canvas = Image.new("RGBA", (canvas_w, canvas_h), (0, 0, 0, 0))

    if want_glow:
        glow_color = (255, 255, 255, 90) if is_dark_bg else (0, 0, 0, 60)
        glow_mask = logo_rgba.split()[-1]
        glow = Image.new("RGBA", (canvas_w, canvas_h), (0, 0, 0, 0))
        glow_layer = Image.new("RGBA", (canvas_w, canvas_h), glow_color)
        mask_canvas = Image.new("L", (canvas_w, canvas_h), 0)
        mask_canvas.paste(glow_mask, (pad, pad))
        mask_canvas = mask_canvas.filter(ImageFilter.GaussianBlur(pad / 2.2))
        glow.paste(glow_layer, (0, 0), mask_canvas)
        canvas = Image.alpha_composite(canvas, glow)

    if want_shadow:
        shadow_mask = logo_rgba.split()[-1]
        shadow = Image.new("RGBA", (canvas_w, canvas_h), (0, 0, 0, 0))
        shadow_layer = Image.new("RGBA", (canvas_w, canvas_h), (0, 0, 0, 130))
        mask_canvas = Image.new("L", (canvas_w, canvas_h), 0)
        offset = max(2, logo_rgba.width // 40)
        mask_canvas.paste(shadow_mask, (pad + offset, pad + offset))
        mask_canvas = mask_canvas.filter(ImageFilter.GaussianBlur(pad / 4))
        shadow.paste(shadow_layer, (0, 0), mask_canvas)
        canvas = Image.alpha_composite(canvas, shadow)

    canvas.alpha_composite(logo_rgba, (pad, pad))
    return canvas, pad


def apply_watermark_to_frame(base_rgba, opts):
    """Applies the watermark to a single RGBA frame and returns the result."""
    if not opts.get("enabled", True):
        return base_rgba

    W, H = base_rgba.size
    size_pct = max(2.0, min(5.0, float(opts.get("size_pct", 3.2))))
    logo_w = max(20, int(W * size_pct / 100.0))
    margin = opts.get("margin")
    margin = int(margin) if margin else max(24, int(W * 0.02))
    position_pref = opts.get("position", "br")

    probe_box = _pick_watermark_geometry(base_rgba, logo_w, logo_w, margin, position_pref)
    theme_pref = opts.get("theme", "auto")
    if theme_pref == "auto":
        bg_theme, is_colorful = _analyze_background_theme(base_rgba, probe_box)
    else:
        bg_theme, is_colorful = theme_pref, False

    use_light_logo = (bg_theme == "dark")
    logo = _load_watermark_logo(light=use_light_logo)
    logo_h = int(logo.height * (logo_w / logo.width))
    logo = logo.resize((logo_w, logo_h), Image.LANCZOS)

    logo_with_fx, pad = _add_shadow_and_glow(
        logo, want_shadow=opts.get("shadow", True), want_glow=opts.get("glow", True),
        is_dark_bg=use_light_logo,
    )

    opacity = max(10, min(100, int(opts.get("opacity", 88)))) / 100.0
    if opacity < 1.0:
        alpha = logo_with_fx.split()[-1].point(lambda a: int(a * opacity))
        logo_with_fx.putalpha(alpha)

    final_box = _pick_watermark_geometry(base_rgba, logo_with_fx.width, logo_with_fx.height, max(0, margin - pad), position_pref)
    x, y, _, _ = final_box

    result = base_rgba.copy()
    result.alpha_composite(logo_with_fx, (x, y))
    return result


def apply_watermark(image_bytes, opts):
    """Top-level entry point: watermarks `image_bytes` and returns new bytes
    in the same format, preserving animation (GIF/animated WebP). If
    watermarking is disabled or Pillow is unavailable, returns the original
    bytes untouched."""
    if not _WATERMARK_AVAILABLE or not (opts or {}).get("enabled", True):
        return image_bytes
    try:
        import io as _io
        src = Image.open(_io.BytesIO(image_bytes))
        fmt = (src.format or "PNG").upper()

        is_animated = getattr(src, "is_animated", False) and src.n_frames > 1
        if is_animated and fmt in ("GIF", "WEBP"):
            frames = []
            durations = []
            for frame in ImageSequence.Iterator(src):
                rgba = frame.convert("RGBA")
                watermarked = apply_watermark_to_frame(rgba, opts)
                frames.append(watermarked)
                durations.append(frame.info.get("duration", 80))
            out = _io.BytesIO()
            save_kwargs = {"save_all": True, "append_images": frames[1:], "duration": durations,
                           "loop": src.info.get("loop", 0), "disposal": 2}
            if fmt == "GIF":
                frames[0].save(out, format="GIF", **save_kwargs)
            else:
                frames[0].save(out, format="WEBP", **save_kwargs)
            return out.getvalue()

        rgba = src.convert("RGBA")
        watermarked = apply_watermark_to_frame(rgba, opts)
        out = _io.BytesIO()
        if fmt in ("JPEG", "JPG"):
            watermarked.convert("RGB").save(out, format="JPEG", quality=95, optimize=True)
        elif fmt == "WEBP":
            watermarked.save(out, format="WEBP", quality=95, lossless=False)
        else:
            watermarked.save(out, format="PNG", optimize=True)
        return out.getvalue()
    except Exception as e:
        print(f"[Watermark] failed to apply, returning original image unmodified: {e}")
        return image_bytes


def _finalize_generated_image(raw_bytes, watermark_opts):
    """Applies the automatic watermark and returns a base64 string ready to
    send to the frontend. This is the single choke point every image-
    generation branch routes through, so the watermark can never be
    'forgotten' for any provider."""
    opts = dict(watermark_opts or {})
    opts.setdefault("enabled", True)
    watermarked_bytes = apply_watermark(raw_bytes, opts)
    return base64.b64encode(watermarked_bytes).decode("utf-8")


def _io_BytesIO(b):
    import io as _io
    return _io.BytesIO(b)


# ── Admin endpoints: upload / replace / reset / preview the official logo ────

@app.route("/api/admin/watermark/upload", methods=["POST"])
def admin_watermark_upload():
    if not _require_admin():
        return jsonify({"error": "unauthorized"}), 401
    if not _WATERMARK_AVAILABLE:
        return jsonify({"error": "Pillow is not installed on the server — run `pip install Pillow`"}), 503
    data = request.get_json(force=True) or {}
    saved = []
    for key, path in (("logo_light_base64", _WATERMARK_LOGO_LIGHT_PATH),
                       ("logo_dark_base64", _WATERMARK_LOGO_DARK_PATH)):
        b64 = data.get(key)
        if not b64:
            continue
        try:
            raw = base64.b64decode(b64, validate=True)
        except Exception:
            return jsonify({"error": f"{key} is not valid base64"}), 400

        is_svg = raw.lstrip().startswith(b"<?xml") or raw.lstrip().startswith(b"<svg") or b"<svg" in raw[:200]
        if is_svg:
            try:
                import cairosvg
                raw = cairosvg.svg2png(bytestring=raw, output_width=1024, output_height=1024)
            except ImportError:
                return jsonify({
                    "error": "SVG upload requires `pip install cairosvg` on the server. "
                             "Please upload a transparent PNG or WebP instead, or install "
                             "cairosvg (free/open-source) to enable SVG."
                }), 503
            except Exception as e:
                return jsonify({"error": f"Could not rasterize SVG: {e}"}), 400

        try:
            img = Image.open(_io_BytesIO(raw)).convert("RGBA")
        except Exception as e:
            return jsonify({"error": f"Could not read image for {key}: {e}"}), 400
        img.save(path, format="PNG")
        saved.append(key)

    if not saved:
        return jsonify({"error": "no logo_light_base64 or logo_dark_base64 provided"}), 400
    global _default_logo_cache
    _default_logo_cache = {}
    return jsonify({"status": "saved", "updated": saved})


@app.route("/api/admin/watermark/reset", methods=["POST"])
def admin_watermark_reset():
    if not _require_admin():
        return jsonify({"error": "unauthorized"}), 401
    for path in (_WATERMARK_LOGO_LIGHT_PATH, _WATERMARK_LOGO_DARK_PATH):
        if _os.path.exists(path):
            _os.remove(path)
    global _default_logo_cache
    _default_logo_cache = {}
    return jsonify({"status": "reset to default logo"})


@app.route("/api/admin/watermark/preview", methods=["GET"])
def admin_watermark_preview():
    if not _require_admin():
        return jsonify({"error": "unauthorized"}), 401
    if not _WATERMARK_AVAILABLE:
        return jsonify({"error": "Pillow is not installed on the server"}), 503
    import io as _io
    size = 640
    preview = Image.new("RGB", (size, size), (235, 235, 235))
    draw = ImageDraw.Draw(preview)
    draw.rectangle([0, 0, size, size // 2], fill=(30, 30, 34))
    draw.rectangle([0, size // 2, size // 2, size], fill=(230, 60, 90))
    opts = {
        "enabled": True, "opacity": int(request.args.get("opacity", 88)),
        "size_pct": float(request.args.get("size_pct", 3.2)),
        "position": request.args.get("position", "br"),
        "shadow": request.args.get("shadow", "1") != "0",
        "glow": request.args.get("glow", "1") != "0",
        "theme": request.args.get("theme", "auto"),
    }
    watermarked = apply_watermark_to_frame(preview.convert("RGBA"), opts)
    out = _io.BytesIO()
    watermarked.convert("RGB").save(out, format="PNG")
    return Response(out.getvalue(), mimetype="image/png")


@app.route("/api/watermark/info", methods=["GET"])
def watermark_info():
    return jsonify({
        "available": _WATERMARK_AVAILABLE,
        "custom_logo_light": _os.path.exists(_WATERMARK_LOGO_LIGHT_PATH),
        "custom_logo_dark": _os.path.exists(_WATERMARK_LOGO_DARK_PATH),
    })


# ── Daily chat streaks + re-engagement scheduling ────────────────────────────
_ACTIVITY_FILE = _os.path.join(_DATA_DIR, "user_activity.json")
_activity_lock = threading.Lock()

_REENGAGEMENT_ROTATION = ["come_back", "study", "activity", "feature"]
_REENGAGEMENT_SKIP_IF_ACTIVE_WITHIN_HOURS = 1
# Minimum gap between two notifications to the same user — prevents duplicates
# if the cron endpoint is hit multiple times in quick succession (e.g. two
# external crons, or a manual test). Default 55 minutes so an hourly cron still
# fires normally, but rapid retries are ignored.
#
# NOTE ON SCHEDULES:
#   - Render (always-on): the in-process background thread below fires
#     every _REENGAGEMENT_CHECK_INTERVAL_SECONDS (1 hour by default) — see
#     _reengagement_loop(). Nothing else to configure.
#   - Vercel (serverless): the in-process thread never runs (frozen between
#     requests), so /api/cron/reengagement must be triggered externally.
#     Set up a Vercel Cron Job that hits it ONCE A DAY AT 12:00 — see the
#     vercel.json example and CRON_SECRET notes near that route below.
_REENGAGEMENT_MIN_GAP_MINUTES = 55
_REENGAGEMENT_CHECK_INTERVAL_SECONDS = 60 * 60  # Render: fires once every hour


def _load_all_activity():
    try:
        if _os.path.exists(_ACTIVITY_FILE):
            with open(_ACTIVITY_FILE, encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return {}


def _save_all_activity(data):
    try:
        with open(_ACTIVITY_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f)
    except Exception:
        pass


def _today_str():
    return datetime.date.today().isoformat()


def _update_user_activity(username):
    with _activity_lock:
        data = _load_all_activity()
        rec = data.get(username) or {"streak": 0, "last_active_day": None,
                                      "last_active_ts": 0, "last_notified_ts": 0}
        today = _today_str()
        last_day = rec.get("last_active_day")
        if last_day != today:
            gap_days = None
            if last_day:
                try:
                    y, m, d = (int(x) for x in last_day.split("-"))
                    gap_days = (datetime.date.today() - datetime.date(y, m, d)).days
                except Exception:
                    gap_days = None
            if gap_days == 1:
                rec["streak"] = rec.get("streak", 0) + 1
            else:
                rec["streak"] = 1
            rec["last_active_day"] = today
        rec["last_active_ts"] = time.time()
        data[username] = rec
        _save_all_activity(data)
        return rec


def _get_user_streak(username):
    with _activity_lock:
        data = _load_all_activity()
    return (data.get(username) or {}).get("streak", 0)


def _subscribed_usernames():
    _load_push_subscriptions()
    names = set()
    for sub in _push_subscriptions.values():
        u = sub.get("_username")
        if u:
            names.add(u)
    return names


def _run_reengagement_pass():
    if not _PUSH_AVAILABLE or not VAPID_PRIVATE_KEY or not VAPID_PUBLIC_KEY:
        return
    now = time.time()
    usernames = _subscribed_usernames()
    if not usernames:
        return

    with _activity_lock:
        data = _load_all_activity()
        changed = False
        for username in usernames:
            rec = data.get(username) or {
                "streak": 0, "last_active_day": None, "last_active_ts": 0,
                "last_notified_ts": 0, "notif_rotation_index": 0,
            }
            last_active_ts = rec.get("last_active_ts", 0)
            hours_inactive = (now - last_active_ts) / 3600.0 if last_active_ts else 999
            if hours_inactive < _REENGAGEMENT_SKIP_IF_ACTIVE_WITHIN_HOURS:
                continue

            # Don't re-notify the same user within the minimum-gap window,
            # so rapid cron hits don't spam the same person.
            last_notified_ts = rec.get("last_notified_ts", 0)
            minutes_since_notified = (now - last_notified_ts) / 60.0 if last_notified_ts else 99999
            if minutes_since_notified < _REENGAGEMENT_MIN_GAP_MINUTES:
                continue

            streak = rec.get("streak", 0)
            if streak >= 2:
                body = f"Your {streak}-day streak is on hold! Come chat with me to keep it going."
            else:
                idx = rec.get("notif_rotation_index", 0) % len(_REENGAGEMENT_ROTATION)
                category = _REENGAGEMENT_ROTATION[idx]
                rec["notif_rotation_index"] = idx + 1
                body = _random_notification_body(category)

            try:
                send_push_notification_to_user(username, "Mythic AI", body, url="/")
            except Exception:
                pass
            rec["last_notified_ts"] = now
            data[username] = rec
            changed = True
        if changed:
            _save_all_activity(data)


def _reengagement_loop():
    """Render / always-on hosts only. Fires _run_reengagement_pass() every
    _REENGAGEMENT_CHECK_INTERVAL_SECONDS (1 hour). Never runs on Vercel —
    use the /api/cron/reengagement endpoint with an external daily cron there."""
    while True:
        try:
            _run_reengagement_pass()
        except Exception as e:
            print(f"[Reengagement] error: {e}")
        time.sleep(_REENGAGEMENT_CHECK_INTERVAL_SECONDS)


# --- HTML pages ----------------------------------------------------------

PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover">
<meta name="theme-color" content="#10a37f">
<meta name="mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-capable" content="yes">
<meta name="apple-mobile-web-app-status-bar-style" content="black-translucent">
<meta name="apple-mobile-web-app-title" content="Mythic AI">

<!-- Primary SEO -->
<title>Mythic AI — Free Smart AI Chat Assistant</title>
<meta name="description" content="Mythic AI is a free, smart AI chat assistant for questions, writing, coding, image generation, homework help, and more — built by Aarav Singh.">
<meta name="keywords" content="Mythic AI, Aarav AI, Aarav Singh AI, AI chat assistant, free AI chatbot, AI assistant, AI image generator, AI homework helper, AI coding assistant">
<meta name="author" content="Aarav Singh">
<meta name="robots" content="index, follow">
<link rel="canonical" href="__CANONICAL_URL__">
<meta name="google-site-verification" content="PFkkupIPte_3H3QnLU0oQ-_WF67Pu8kTlIkTRWussww" />

<!-- Open Graph / Facebook / WhatsApp / LinkedIn -->
<meta property="og:type" content="website">
<meta property="og:site_name" content="Mythic AI">
<meta property="og:title" content="Mythic AI — Free Smart AI Chat Assistant">
<meta property="og:description" content="Chat, ask questions, generate images, get homework help, and more with Mythic AI — a free, smart AI assistant built by Aarav Singh.">
<meta property="og:url" content="__CANONICAL_URL__">
<meta property="og:image" content="__CANONICAL_ORIGIN__/icon-512.png">
<meta property="og:image:width" content="512">
<meta property="og:image:height" content="512">
<meta property="og:locale" content="en_US">

<!-- Twitter Card -->
<meta name="twitter:card" content="summary">
<meta name="twitter:title" content="Mythic AI — Free Smart AI Chat Assistant">
<meta name="twitter:description" content="Chat, ask questions, generate images, get homework help, and more with Mythic AI — a free, smart AI assistant.">
<meta name="twitter:image" content="__CANONICAL_ORIGIN__/icon-512.png">

<!-- Schema.org structured data — tells Google exactly what this site is,
     and that "Aarav AI" refers to the same app as "Mythic AI".
     The WebSite entity below is what Google's "Site Name" search-result
     feature specifically reads (see developers.google.com/search/docs/appearance/site-names) —
     WebApplication alone is not enough to set that. Both are combined in one
     @graph so there is a single, non-conflicting source of truth. -->
<script type="application/ld+json">
{
  "@context": "https://schema.org",
  "@graph": [
    {
      "@type": "WebSite",
      "@id": "__CANONICAL_ORIGIN__/#website",
      "name": "Mythic AI",
      "alternateName": ["Aarav AI", "Mythic AI Chat", "Aarav Singh AI"],
      "url": "__CANONICAL_ORIGIN__/"
    },
    {
      "@type": "WebApplication",
      "@id": "__CANONICAL_ORIGIN__/#webapp",
      "name": "Mythic AI",
      "alternateName": ["Aarav AI", "Mythic AI Chat", "Aarav Singh AI"],
      "url": "__CANONICAL_ORIGIN__/",
      "description": "Mythic AI (also known as Aarav AI) is a free, smart AI chat assistant for questions, writing, coding, image generation, and homework help.",
      "applicationCategory": "Chatbot, Productivity",
      "operatingSystem": "Any",
      "isPartOf": { "@id": "__CANONICAL_ORIGIN__/#website" },
      "offers": {
        "@type": "Offer",
        "price": "0",
        "priceCurrency": "USD"
      },
      "creator": {
        "@type": "Person",
        "name": "Aarav Singh"
      },
      "image": "__CANONICAL_ORIGIN__/icon-512.png"
    }
  ]
}
</script>

<link rel="manifest" href="/manifest.json">
<link rel="icon" type="image/png" sizes="192x192" href="/icon.png">
<link rel="icon" type="image/png" sizes="512x512" href="/icon-512.png">
<link rel="shortcut icon" href="/favicon.ico">
<link rel="apple-touch-icon" href="/icon.png">
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Noto+Sans:wght@400;600;700&family=Noto+Sans+Devanagari:wght@400;600&display=swap" rel="stylesheet">
<style>
  :root {
    --bg:#0c1410; --panel:rgba(20,31,25,.60); --border:#2a3a3080;
    --text:#f5f3ea; --muted:#9aa89e; --accent:#10a37f;
    --accent2:#0d7a5f; --accent-grad:#10a37f;
    --accent-dim:rgba(16,163,127,.16); --user-bubble:rgba(255,255,255,.08);
    --user-text:var(--text);
    --ai-bubble:transparent; --sidebar-w:272px; --msg-font-size:16px;
    --composer-bg:rgba(20,31,25,.75); --composer-border:rgba(13,122,95,.30);
    --composer-shadow:0 20px 44px rgba(0,0,0,.45), 0 0 0 1px rgba(16,163,127,.08);
    --radius-lg:20px; --radius-md:14px; --radius-sm:10px;
  }
  * { box-sizing:border-box; margin:0; padding:0; }
  html,body { height:100%;
    background:var(--bg);
    color:var(--text);
    font-family:"Inter",-apple-system,BlinkMacSystemFont,"Segoe UI",
      "Noto Sans","Noto Sans Devanagari",sans-serif; overflow:hidden; letter-spacing:.1px;
    position:relative; }
  .layout { position:relative; z-index:1; display:flex; height:100vh;
    height:calc(var(--app-height, 100vh)); height:100dvh; gap:18px; padding:18px; }

  body.theme-light {
    --bg:#fdfaf3; --panel:rgba(255,255,255,.55); --border:#e7e0cf;
    --text:#1b1a13; --muted:#6d6a5e; --accent-dim:#fff1e0;
    --user-bubble:rgba(16,163,127,.10); --user-text:var(--text); --ai-bubble:transparent;
    --composer-bg:rgba(255,255,255,.85); --composer-border:rgba(13,122,95,.22);
    --composer-shadow:0 20px 44px rgba(0,60,20,.10), 0 0 0 1px rgba(16,163,127,.08);
  }

  #sidebar { width:var(--sidebar-w); flex-shrink:0; background:var(--panel);
    backdrop-filter:blur(22px); -webkit-backdrop-filter:blur(22px);
    border:1px solid var(--border); border-radius:var(--radius-lg); display:flex; flex-direction:column;
    transition:margin-left .2s ease; box-shadow:0 10px 30px rgba(0,0,0,.25); overflow:hidden; }
  #sidebar.hidden { margin-left:calc(-1 * var(--sidebar-w) - 10px); }
  #new-chat-btn { margin:14px; padding:11px 15px; background:var(--accent-grad); color:#fff;
    border:none; border-radius:var(--radius-md); font-size:13.5px; font-weight:700; cursor:pointer;
    text-align:left; box-shadow:0 6px 18px rgba(16,163,127,.35); }
  #new-chat-btn:hover { filter:brightness(1.08); }
  #api-keys-shortcut-btn { margin:0 14px 8px; padding:9px 14px; background:none; border:1px solid var(--border);
    color:var(--muted); border-radius:var(--radius-md); font-size:13px; font-weight:500; cursor:pointer;
    text-align:left; touch-action:manipulation; -webkit-tap-highlight-color:transparent; }
  #api-keys-shortcut-btn:hover { background:var(--accent-dim); color:var(--text); border-color:var(--accent); }
  #conv-list { flex:1; overflow-y:auto; padding:0 10px; display:flex; flex-direction:column; gap:3px; }
  .conv-item { display:flex; align-items:center; justify-content:space-between; gap:6px;
    padding:10px 11px; border-radius:var(--radius-sm); cursor:pointer; font-size:13px; color:var(--muted); }
  .conv-item:hover { background:var(--accent-dim); color:var(--text); }
  .conv-item.active { background:var(--accent-dim); color:var(--accent); font-weight:600;
    box-shadow:inset 2px 0 0 var(--accent); }
  .conv-item .title { overflow:hidden; text-overflow:ellipsis; white-space:nowrap; flex:1; }
  .conv-item .menu-btn { opacity:0; background:none; border:none; color:var(--muted);
    cursor:pointer; font-size:16px; padding:2px 8px; border-radius:5px; flex-shrink:0;
    touch-action:manipulation; }
  .conv-item:hover .menu-btn { opacity:1; }
  .conv-item .menu-btn:hover { color:var(--accent); background:rgba(16,163,127,.12); }
  #sidebar-footer { padding:10px; font-size:11px; color:var(--muted); border-top:1px solid var(--border); }

  #sidebar-profile { display:flex; align-items:center; gap:10px; width:100%;
    background:none; border:none; padding:9px 8px; border-radius:var(--radius-md); cursor:pointer;
    font-family:inherit; text-align:left; touch-action:manipulation; -webkit-tap-highlight-color:transparent; }
  #sidebar-profile:hover { background:var(--accent-dim); }
  #sidebar-profile-avatar { width:40px; height:40px; min-width:40px; border-radius:50%;
    display:flex; align-items:center; justify-content:center; flex-shrink:0;
    font-size:15px; font-weight:700; color:#fff; letter-spacing:.2px;
    background:var(--accent-grad);
    border:1px solid var(--border); box-shadow:0 4px 12px rgba(16,163,127,.35);
    object-fit:cover; overflow:hidden; image-rendering:auto; }
  #sidebar-profile-avatar img { width:100%; height:100%; object-fit:cover; border-radius:50%; display:block; }
  #sidebar-profile-text { display:flex; flex-direction:column; min-width:0; overflow:hidden; }
  #sidebar-profile-name { font-size:13px; font-weight:600; color:var(--text);
    white-space:nowrap; overflow:hidden; text-overflow:ellipsis; }
  #sidebar-profile-sub { font-size:11px; color:var(--muted); }
  #sidebar-byline { text-align:center; padding-top:8px; margin-top:4px; border-top:1px solid var(--border); }

  @media(max-width:768px) {
    #sidebar-profile-avatar { width:38px; height:38px; min-width:38px; font-size:14px; }
  }

  .app { display:flex; flex-direction:column; height:100vh;
    height:calc(var(--app-height, 100vh)); height:100dvh; flex:1; min-width:0; min-height:0;
    background:var(--panel); backdrop-filter:blur(22px); -webkit-backdrop-filter:blur(22px);
    border:1px solid var(--border); border-radius:var(--radius-lg);
    box-shadow:0 10px 30px rgba(0,0,0,.25); overflow:hidden; }
  header { padding:calc(14px + env(safe-area-inset-top)) 20px 14px; border-bottom:1px solid var(--border);
    display:flex; align-items:center; justify-content:space-between; gap:10px;
    background:var(--panel); position:relative; z-index:20; flex-shrink:0; }
  header::before { content:""; position:absolute; top:0; left:0; right:0; height:4px;
    background:linear-gradient(90deg,#10a37f 0 33.33%, #ffffff 33.33% 66.66%, #0d7a5f 66.66% 100%);
    border-radius:var(--radius-lg) var(--radius-lg) 0 0; }
  header .left { display:flex; align-items:center; gap:10px; min-width:0; }
  header .right { display:flex; align-items:center; gap:8px; flex-shrink:0; }
  header button { touch-action:manipulation; -webkit-tap-highlight-color:transparent; }
  #sidebar-toggle { background:none; border:1px solid var(--border); color:var(--muted);
    width:40px; height:40px; border-radius:var(--radius-sm); cursor:pointer; font-size:16px; flex-shrink:0; }
  #sidebar-toggle:hover { background:var(--accent-dim); color:var(--accent); border-color:var(--accent); }
  header h1 { font-size:19px; font-weight:800; margin:0; letter-spacing:.2px;
    background:var(--accent-grad); -webkit-background-clip:text; background-clip:text;
    -webkit-text-fill-color:transparent; }
  #streak-badge { display:none; align-items:center; gap:4px; background:linear-gradient(135deg,#ff9d42,#ff5f6d);
    color:#fff; font-size:11px; font-weight:800; padding:3px 9px; border-radius:12px; white-space:nowrap; }
  #name-btn { background:none; border:1px solid var(--border); color:var(--muted);
    width:40px; height:40px; border-radius:var(--radius-sm); cursor:pointer; font-size:16px; flex-shrink:0;
    display:flex; align-items:center; justify-content:center; touch-action:manipulation; }
  #name-btn:hover { background:var(--accent-dim); color:var(--accent); border-color:var(--accent); }
  #settings-btn { background:none; border:1px solid var(--border); color:var(--muted);
    width:40px; height:40px; border-radius:var(--radius-sm); cursor:pointer; font-size:16px; flex-shrink:0;
    display:flex; align-items:center; justify-content:center; touch-action:manipulation; }
  #settings-btn:hover { background:var(--accent-dim); color:var(--accent); border-color:var(--accent); }
  #export-btn { background:none; border:1px solid var(--border); color:var(--muted);
    width:40px; height:40px; border-radius:var(--radius-sm); cursor:pointer; font-size:16px; flex-shrink:0;
    display:flex; align-items:center; justify-content:center; touch-action:manipulation; }
  #export-btn:hover { background:var(--accent-dim); color:var(--accent); border-color:var(--accent); }
  #share-btn { background:none; border:1px solid var(--border); color:var(--muted);
    width:40px; height:40px; border-radius:var(--radius-sm); cursor:pointer; font-size:16px; flex-shrink:0;
    display:flex; align-items:center; justify-content:center; touch-action:manipulation; }
  #share-btn:hover { background:var(--accent-dim); color:var(--accent); border-color:var(--accent); }
  #share-btn.active { color:var(--accent); border-color:var(--accent); }

  #share-modal-overlay { display:none; position:fixed; inset:0; background:rgba(0,0,0,.6);
    z-index:250; align-items:center; justify-content:center; padding:16px; }
  #share-modal-overlay.show { display:flex; }
  #share-modal { background:var(--bg); border:1px solid var(--border); border-radius:14px;
    padding:22px; width:100%; max-width:400px; box-shadow:0 10px 40px rgba(0,0,0,.3); }
  #share-modal h3 { margin:0 0 4px; font-size:16px; }
  #share-modal p.sub { margin:0 0 16px; font-size:12.5px; color:var(--muted); }
  #share-link-row { display:flex; gap:8px; margin-bottom:14px; }
  #share-link-input { flex:1; min-width:0; padding:9px 12px; border-radius:8px;
    border:1.5px solid var(--border); background:var(--panel); color:var(--text);
    font-size:12.5px; outline:none; text-overflow:ellipsis; }
  #share-link-input:focus { border-color:var(--accent); }
  #share-open-btn { background:var(--panel); border:1px solid var(--border); color:var(--muted);
    border-radius:8px; padding:0 12px; font-size:14px; cursor:pointer; flex-shrink:0; }
  #share-open-btn:hover { border-color:var(--accent); color:var(--accent); }
  #share-copy-btn { background:var(--accent); color:#fff; border:none; border-radius:8px;
    padding:0 14px; font-size:12.5px; font-weight:700; cursor:pointer; flex-shrink:0; }
  #share-copy-btn:hover { opacity:.9; }
  #share-native-btn { width:100%; background:var(--panel); border:1px solid var(--border);
    color:var(--text); border-radius:8px; padding:10px; font-size:13px; cursor:pointer;
    font-family:inherit; margin-bottom:8px; }
  #share-native-btn:hover { border-color:var(--accent); color:var(--accent); }
  #share-qr-image-btn, #download-qr-btn { background:var(--panel); border:1px solid var(--border);
    color:var(--text); border-radius:8px; padding:10px; font-size:12.5px; cursor:pointer;
    font-family:inherit; margin-bottom:8px; }
  #share-qr-image-btn:hover, #download-qr-btn:hover { border-color:var(--accent); color:var(--accent); }
  #share-revoke-btn { width:100%; background:none; border:1px solid var(--border);
    color:#ef4444; border-radius:8px; padding:10px; font-size:12.5px; cursor:pointer;
    font-family:inherit; margin-bottom:8px; }
  #share-revoke-btn:hover { background:rgba(239,68,68,.08); }
  #share-close-btn { width:100%; background:none; border:1px solid var(--border);
    color:var(--muted); border-radius:8px; padding:10px; font-size:13px; cursor:pointer;
    font-family:inherit; }
  #share-status { font-size:11.5px; color:var(--muted); text-align:center; margin-top:2px; min-height:14px; }
  #vip-btn { background:none; border:1px solid var(--border); color:var(--muted);
    width:40px; height:40px; border-radius:6px; cursor:pointer; font-size:16px; flex-shrink:0;
    display:flex; align-items:center; justify-content:center; touch-action:manipulation; }
  #vip-btn:hover { background:var(--panel); }
  #vip-btn.active { color:var(--accent); border-color:var(--accent); }

  #fullscreen-btn { display:flex; align-items:center; justify-content:center;
    width:40px; height:40px; border-radius:6px; flex-shrink:0;
    background:none; border:1px solid var(--border);
    color:var(--muted); font-size:16px; cursor:pointer; touch-action:manipulation;
    -webkit-tap-highlight-color:transparent; }
  #fullscreen-btn:hover { color:var(--text); border-color:var(--accent); background:var(--panel); }
  #fullscreen-btn.active { color:var(--accent); border-color:var(--accent); }
  #fullscreen-icon { font-size:16px; }

  body.pseudo-fullscreen #sidebar-toggle,
  body.pseudo-fullscreen header .left h1 { display:none; }
  body.pseudo-fullscreen header { padding-top:calc(6px + env(safe-area-inset-top)); padding-bottom:6px; }

  /* ── Header "⋯" overflow menu ─────────────────────────────────────── */
  #header-menu-wrap { position:relative; }
  #header-menu-btn { background:none; border:1px solid var(--border); color:var(--muted);
    width:40px; height:40px; border-radius:6px; cursor:pointer; font-size:19px; flex-shrink:0;
    display:flex; align-items:center; justify-content:center; touch-action:manipulation;
    line-height:1; }
  #header-menu-btn:hover, #header-menu-btn[aria-expanded="true"] { background:var(--panel); color:var(--text); border-color:var(--accent); }
  #header-menu-dropdown { display:none; position:absolute; top:calc(100% + 6px); right:0;
    background:var(--panel); border:1px solid var(--border); border-radius:10px;
    box-shadow:0 8px 28px rgba(0,0,0,.35); min-width:200px; max-width:min(260px, calc(100vw - 24px));
    overflow:hidden; z-index:500; padding:6px; }
  #header-menu-dropdown.open { display:block; animation:headerMenuIn .12s ease-out; }
  @keyframes headerMenuIn { from { opacity:0; transform:translateY(-4px); } to { opacity:1; transform:none; } }
  #header-menu-dropdown .menu-divider { height:1px; background:var(--border); margin:5px 6px; }
  /* Scoped overrides — same buttons/ids/handlers as before, restyled as
     full-width rows instead of their old standalone square icon look. */
  #header-menu-dropdown button { display:flex; align-items:center; gap:10px; width:100%;
    background:none; border:none; color:var(--text); font-size:13px; font-weight:500;
    padding:9px 10px; border-radius:7px; cursor:pointer; font-family:inherit;
    text-align:left; height:auto; }
  #header-menu-dropdown button:hover { background:var(--accent-dim, rgba(16,163,127,.12)); }
  #header-menu-dropdown button.active { color:var(--accent); }
  #header-menu-dropdown button[data-label]::after { content:attr(data-label); }
  #header-menu-dropdown #clear-btn { color:#ef4444; }
  #header-menu-dropdown #clear-btn:hover { background:rgba(239,68,68,.1); }
  #header-menu-dropdown #clear-btn::before { content:"🗑"; }
  #header-menu-dropdown #logout-btn { color:#ef4444; }
  #header-menu-dropdown #logout-btn:hover { background:rgba(239,68,68,.1); }
  #header-menu-dropdown #logout-btn::before { content:"🚪"; }

  #name-modal-overlay { display:none; position:fixed; inset:0; background:rgba(0,0,0,.55);
    z-index:200; align-items:center; justify-content:center; }
  #name-modal-overlay.show { display:flex; }
  #name-modal { background:var(--bg); border:1px solid var(--border); border-radius:14px;
    padding:22px; width:90%; max-width:360px; box-shadow:0 10px 40px rgba(0,0,0,.3); }
  #name-modal h3 { margin:0 0 6px; font-size:16px; color:var(--text); }
  #name-modal p { margin:0 0 14px; font-size:12.5px; color:var(--muted); }
  #avatar-upload-row { display:flex; align-items:center; gap:14px; margin-bottom:18px; }
  #avatar-upload-preview { width:64px; height:64px; min-width:64px; border-radius:50%;
    background:var(--accent-grad); display:flex; align-items:center; justify-content:center;
    color:#fff; font-weight:800; font-size:24px; overflow:hidden; flex-shrink:0; }
  #avatar-upload-preview img { width:100%; height:100%; object-fit:cover; border-radius:50%; display:block; }
  #avatar-upload-actions { display:flex; flex-direction:column; gap:6px; }
  #avatar-upload-btn { display:inline-block; padding:7px 12px; border-radius:8px; font-size:12.5px;
    cursor:pointer; border:1px solid var(--border); background:none; color:var(--text); text-align:center; }
  #avatar-upload-btn:hover { background:var(--panel); }
  #avatar-remove-btn { padding:6px 12px; border-radius:8px; font-size:12px;
    cursor:pointer; border:1px solid transparent; background:none; color:var(--muted); }
  #avatar-remove-btn:hover { color:var(--text); text-decoration:underline; }
  #avatar-preset-grid { display:grid; grid-template-columns:repeat(7, 1fr); gap:8px; margin-bottom:18px;
    max-height:220px; overflow-y:auto; padding-right:2px; }
  #avatar-preset-grid img { width:100%; aspect-ratio:1/1; border-radius:50%; cursor:pointer; display:block;
    border:2px solid transparent; box-sizing:border-box; transition:border-color .15s, transform .15s; background:var(--panel); }
  #avatar-preset-grid img:hover { transform:scale(1.08); border-color:var(--border); }
  #avatar-preset-grid img.selected { border-color:var(--accent, #7c5cff); box-shadow:0 0 0 2px rgba(16,163,127,.25); }
  @media (max-width:480px) { #avatar-preset-grid { grid-template-columns:repeat(5, 1fr); } }
  #name-input { width:100%; box-sizing:border-box; padding:10px 12px; border-radius:8px;
    border:1.5px solid var(--border); background:var(--panel); color:var(--text);
    font-size:14.5px; outline:none; }
  #name-input:focus { border-color:var(--accent); }
  #name-modal-actions { display:flex; justify-content:flex-end; gap:8px; margin-top:16px; }
  #name-modal-actions button { padding:8px 14px; border-radius:8px; font-size:13px;
    cursor:pointer; border:1px solid var(--border); background:none; color:var(--text); }
  #name-cancel-btn:hover { background:var(--panel); }
  #name-save-btn { background:var(--accent); color:#fff; border-color:var(--accent); }
  #name-save-btn:hover { opacity:.9; }
  #clear-btn { background:none; border:1px solid var(--border); color:var(--muted);
    font-size:12px; padding:6px 12px; border-radius:6px; cursor:pointer; flex-shrink:0; }
  #clear-btn:hover { background:var(--panel); }

  #settings-modal-overlay { display:none; position:fixed; inset:0; background:rgba(0,0,0,.55);
    z-index:200; align-items:center; justify-content:center; }
  #settings-modal { background:var(--bg); border:1px solid var(--border); border-radius:14px;
    padding:22px; width:92%; max-width:420px; max-height:86vh; overflow-y:auto;
    box-shadow:0 10px 40px rgba(0,0,0,.3); }
  #settings-modal h3 { margin:0 0 4px; font-size:17px; color:var(--text); }
  #settings-modal p.sub { margin:0 0 16px; font-size:12.5px; color:var(--muted); }
  .settings-section { margin-bottom:16px; }
  .settings-section label { display:block; font-size:12px; color:var(--muted); margin-bottom:6px; font-weight:600; }
  .settings-section .hint { font-size:11px; color:var(--muted); margin-top:6px; font-weight:400; }
  .settings-row { display:flex; gap:8px; flex-wrap:wrap; }
  .settings-choice { flex:1; min-width:80px; padding:8px 10px; border-radius:8px; border:1.5px solid var(--border);
    background:var(--panel); color:var(--muted); cursor:pointer; font-size:12.5px; font-family:inherit; text-align:center; }
  .settings-choice:hover { border-color:var(--accent); }
  #accent-color-input { width:44px; height:34px; border:1.5px solid var(--border); border-radius:8px;
    background:var(--panel); cursor:pointer; padding:2px; }
  #font-size-slider { width:100%; accent-color:var(--accent); }
  #font-size-label { font-size:12px; color:var(--muted); }
  .settings-select { width:100%; padding:9px 10px; border-radius:8px; border:1.5px solid var(--border);
    background:var(--panel); color:var(--text); font-size:13px; font-family:inherit; outline:none; }
  .settings-select:focus { border-color:var(--accent); }
  .settings-text-input { width:100%; box-sizing:border-box; padding:9px 12px; border-radius:8px;
    border:1.5px solid var(--border); background:var(--panel); color:var(--text);
    font-size:13px; font-family:inherit; outline:none; }
  .settings-text-input:focus { border-color:var(--accent); }
  #custom-instructions-input { width:100%; box-sizing:border-box; padding:9px 12px; border-radius:8px;
    border:1.5px solid var(--border); background:var(--panel); color:var(--text);
    font-size:13px; font-family:inherit; outline:none; resize:vertical; min-height:60px; }
  #custom-instructions-input:focus { border-color:var(--accent); }
  #settings-close-btn { width:100%; margin-top:6px; background:var(--accent); color:#fff; border:none;
    border-radius:10px; padding:11px; font-size:14px; font-weight:700; cursor:pointer; font-family:inherit; }
  #settings-close-btn:hover { opacity:.9; }

  body.bubble-compact .msg { padding:7px 11px; border-radius:12px; }
  body.bubble-compact #messages { gap:8px; }
  body.bubble-comfortable .msg { padding:11px 15px; border-radius:18px; }
  body.bubble-comfortable #messages { gap:16px; }
  body.bubble-spacious .msg { padding:16px 20px; border-radius:22px; }
  body.bubble-spacious #messages { gap:24px; }

  #messages-wrap { flex:1; min-height:0; overflow-y:auto; position:relative; }
  #messages { padding:28px 20px 24px; display:flex; flex-direction:column; gap:22px;
    max-width:760px; margin:0 auto; width:100%; min-height:100%; }
  .msg { max-width:80%; padding:11px 15px; border-radius:18px; line-height:1.6;
    font-size:var(--msg-font-size); white-space:pre-wrap; word-wrap:break-word; }
  .msg.user { align-self:flex-end; background:var(--user-bubble); color:var(--user-text);
    border-bottom-right-radius:4px; }
  /* AI replies render as plain, open flowing text (no bubble/box), matching
     how Gemini/ChatGPT show assistant replies — only the user's own messages
     stay as a boxed bubble. */
  .msg.ai { align-self:flex-start; background:none; color:var(--text);
    padding:6px 0; border-radius:0; max-width:100%; }
  .msg.error { align-self:center; background:#fef2f2; border:1px solid #fecaca;
    color:#dc2626; font-size:13px; border-radius:10px; }
  .msg img { max-width:100%; border-radius:10px; display:block; margin-top:8px; }
  .attach-chip { font-size:11.5px; opacity:.75; margin-bottom:4px; }
  .file-download-chip { display:inline-flex; align-items:center; gap:8px; margin-top:10px;
    padding:9px 14px; background:var(--panel); border:1px solid var(--border); border-radius:10px;
    text-decoration:none; color:var(--text); font-size:12.5px; }
  .file-download-chip:hover { border-color:var(--accent); color:var(--accent); }

  .msg-row { display:flex; flex-direction:column; max-width:80%; }
  .msg-row.user { align-self:flex-end; align-items:flex-end; }
  .msg-row.ai { align-self:flex-start; align-items:flex-start; max-width:100%; width:100%;
    position:relative; padding-left:38px; }
  .msg-row.ai::before { content:""; position:absolute; left:0; top:0;
    width:28px; height:28px; border-radius:50%;
    background:url("data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAADgAAAA4CAYAAACohjseAAAZxElEQVR42t2ad3hc5Zn2f+9pM6OZkTRqNu7dxtgYMKYvDpjeDSQsATYh2RTCkgALhCUhm91QkpDAhuwSNoUECJ2EbkyMbbAN7jY27kUusmx1aTT9nLd8f5yRLAfnCySb/WPnuo7m6Gg059zv87z3cz9F+L5vhBD8b72MKZ8IEH/zexkcIQR/S4AhIAMIhIC+W5lDfgj+Vo/g/C1W7SAYgW1b/X/TGLTWGMCyLAQCq7wIxuhDvsOUTf2nFt8McAXLEgN+/yOAxpi/2oIDv1yI8Galkk93dzednV20trTT1tpFoeBTKgQEQYDr2CSSMeoG1dAwpJ5UbQ3VlVW4tkMk4uI6DkIItDZ9pj7kOS3LwrIsjFYUSyU8zzvsswnf941lWX8xsD4XtCyLXC7Hli1b2bFzN+lMnm279tNdgEy0kg4vTiYap8f2yAYWMu9jZ/NEWjuo6e1mRDRg9PAEw8bUMXnyeEYNG0Z9bYpUdSWu6/bfTylFJpuno7OTXY17WLV+D39/+ZmMGzsSKeUhi2CM+csA/rHFOjs7WbduPWvXbWJ/axfxyhSJ6lpyAbQXFQeKAW2FgEygSAtBJlaBrEwS1NehahowXhKnJYezahOJdetpCA4wZlyUGdMnMv3oyYwcPqQf6K7dTby/bBWLl6xlwTaHo0dEeemxO6iuTqG1/ghA5y+zWgiwp7ub1Ws/YOWqdeDEwPJwI1F2bN/GnqZWslmfkq/QgUJJhZYKlCKhDEoJAmMhYwn8I0aijpuBOu1k8ledx96NTex6fj7vLV3K+AlLGT28nuGD66lMJti5azfvLNrO3lGXYNkHmH3WcCqrqj4C7hO7aB+oPjdZv3498xcuwdcukVicNSuXsnrFUoyVIBavxnFdQKO1Igh8AhmgtUJrg1QKKSUGgVIaWTRIP0JgH4E68mScKy4mfvpk7GW7yDzyIqJ7EanqPC6C5s4U9pW34rf0MKs0lwfu+ypTJ08APsolH8tFBwLT2tDbm2bOm/PYsqOJusFD2bhhPW+98Srp7i5Gj59GLJ4k3dVKNtMdMmbZ4lqr8rnGaIMuX7MdDywPaUAJGxVEUXo0avKFRL92KanJtRQffJfMgqeRIsHMm25mS6sheP1evn3L33Htpy8gVV15WOt9bIB9Vmtu3s8zz79GRbKaQEmee+ZZtm/ZSFV1NUcMG0fg+7S37iHwi8Ti1Wity4fCaI02ZcBahefaIIMSluVgLA+tDcoItONiRApp/R3ikmuo/foxtD+6jzsnx2n0Pd74xX9x3UU21191HlMmjcHz3D8J0Pm4brlx4xaef2Uhxx17NJs2beSpp18g3XmAmto66gaNJJvppaN1N14kxqCh4zBGIJVEK4XW5Xdj0EqilAzftcZyPIq5NOCXGdnG0hItFG70LdTvmmjf9Dn+86fns7sdnr/5LU4a38NJ02eSqkriOPZfF+i11jTu2sOLbyzn3HPOYsl7y3j+xTkEhR7iiQSp+mH0dLXS03WAqppBxBM1BH4pBKd1aC2t+i2plEIrGYIuu6rjRQn8IkaA0T4oje1AkFWoRJEn757MloUH+FGLxHvoVBq/u5ela7YwqL6K+rpqkon4JwfYF+NaWlp57Q9ruej8M1ny/kpeeWMR3R27cV2PhoYRtO5vpFTMUNswAsf1KOSzGKP73TM8D0EqKdFGo5Ua8JmQeCzbDYO6ZYNRBMUixI/kuScfYuW6ffzoXx8m0jAeOfardNx+NS997xGUXkSiIsaxUycQjRw+0P9/6TObzfDmwg+YPn0am7duZ868lWTTzZx60vFMmjiBzRtWYExAde0QlJTksmmkLBEEPr5fRAbhed8hZYCUAarfTVW/lZWSaK0QlsD3JV4sxbPP/ox35s3nR/92F9OPKxL0vIb43qNYskT3DZ9jwYIcb7+3gv0H2tBaf3yAxhiCwOftRWtI1Q2hlO/l93NWY4Iuerq7OXrqUTz1+CMcNWksHe0dBH6BYjGHlAG+XyTwi8jAJ5A+MighgxIq8Pv3ntYKrRRKywHsarBsm2K+RGX1ETzx3LO88vs5/OyBB/jDaw9w/qyT0T3NOGYB+u4nMJNqaD59Nq+/tp2lazfQm8kdFqR1WGoVgs1bG2nuCBg3IsUzry6nkGkln0uT7txHR0c7DQ0NzH97LmfPOo09OzagdYCSPjLwkdJHqQAlw2t9FlMqCAGWQRqtQ3bVGtt2KOR96o4YzyNPPs9vHn+TJ594mZfnPs7ZnzqBjs4OII0I0oiu19H3zoGrj2GrNZ3X3lxJ457mwwrujwIEspkMy9ft5eTjj+SNt1fRsn8/vV1NtOzbRimfwXVdjDHEKmK8+spLfPGLX6ClaRuBX0DJoAwyBCiVRMkQoNF6QOgox0MdWi5fUAwdN4MHfv4Ej/78D7y+eDvzFvyGS889CV9rbGEAiS5lsMR+xOrfId/agX/DhSxd6/LWohV0dvV+REpaf5y7Ga1ZvX4bNXWD6GxvY8XanfR27GTf7g34pSIIu18W+b5EKs0vf/lz7rzzLrramvBLebSWyLJLhuCCfqIZaD1tNLbjUCjA6Glnc/dDj/KT/17Jey0u7771Y846cRwtHb1h1tD30MJg8lkseyPmmbno+jhtx57JosV72bZrL1KpQ7OOQ3SbgFwux54DWcYMr+XNBWvIdh/gQPN2lFQIO2Q4MyB90QYKRZ/777+XBx98iGy6Hb+YR6BRMsAMdMVyLDRGowHLtinkNBNPuYJ//t6PePCXu9icHMeip/6R04+spjNdwLYspBpgFWNAWJhCG6J9PublLejLp9PYXsfmnXvpzeQPD7DvC7buaKKyKkXT3n1s2LidXG9LuOKAVqocG8PP9qVqwrLJZAvccsvNPPnkb5F+nmIhh8CgtOwPFcbo8iGwbZtiRjLtrGu54Zvf5fu/bmb/lCN590cncspQi65MgGOHlpMKZB+BKB8sCwxYehtm/rtYFR77xxzNsuXNdHSn+5/vUAsKQaGQZ9uudoYNTvHOkjXkMm0U8r0oKcsZtwA0riX6gWoDpVJARTxGT2+ea6+9hpdfeYWoZ5PPZ7CEQGtV3t1htm/ZFsVMkVMu/RL/8JWbuedX28jPmsCibzQwI1kkpz2E0UipkEpjDHjlexo0RktAI0wR0bYUvbwTecY01mySbN65B6X0YQAaQ1dXGjeaIJftYeOmHchimiDw+/eM0qF/b+1MU1QGjCHq2fzq8eeY9/ZiqisraO/s5cILzmfuW3Opr01RzPXi2DZG63DVLUEpm+ecq7/BpVddx3d+tgLn0qN59+8TTHUL6GiUVStX8fqbC4nEIiil6MoHfNCePehpRoMxCG0QeidmyQfYYyrZb9WxfsNucoUiolzS6gdoCUFzazeD6mvYsmUHHR2tFAvZUFqVxbHSmqhrsWR7gUU7eoi5YT3Etg2fufLTLFu1jvraSlo7ejn1lFOYv2ABo0aNpJDrwbYdwFDKZbjyy9/iU2dfwF0PL6Dh+nN497I4E0Qe4hWsW7eJT18xm4gn0MYi6lq8uWYf761L4w6oyhlCoII07FiDykJx4jj2NxfJ5PIgQn/pByiVpKsnTzIeYcPGbfjFDKVSESUDtAop3WgdukhviULGBwHSwIgRI8iU4lw0+4u8v2INDXWVHGhPM/moo5i/YCFTpkylmO/BLxa44Zs/5vjjZ3DXQ28w+Z8/zzuXJBlBHisRZ83q9Vx44WzaWzsYMrieki/RwiLb3o7obkNYIcv312hQCCMhvRWzo4CcOIKde316c7mQaIwpAxSCUrGEFg5on23bdyKDfFlWhe4pAxm6h7CQxV6KXa1YjovUYAuDaJhGp3c8l3z2OyxbsZL6+ipa2tMMGT6SuX+YxwknnMSd9/yECaOGcef3n+f0793B/AsT1MocdjLBiuVruPiSz9DS1oFVOSQkF6kw2JR6WjAyWzZd2T2xAB1aMWjGbG/GDK+lV8bIF4oHK27lGiwlP6AyGSeT6aW15QAyKIb5WTFNrhQgB09CByW0VkTjVXz/xY3s3duEa4FAY9K7iU44lc4jbuOSGxeyYsUa6mqraOvoIZao5r3lS6myfG7596e47Cc/5vXz4nilHF5VgveWrODS2dfR1inxhl2MTp2MMQrHddneuIdHX1pJpCLBQe7QGFUMGRwRumlTMyru4UerKOR9zB+TTMkPEMKis6ubTCYd6tFCmo6OTpxzb8J4LkYpNGES2cgYWnt6sQUgLBAesusA3jEpOo69g8u+3cmqlWupSyWJJSLcd/d3+JcH5/DlJx7j6XPilLJZopVJFi58nyuu+goduSq8KRejzvwpou5sKIeTlp48LX4drgm3hBCAzOLYBmG7Za8qQVcbxoFSJIpfDMqSswxQiJDyDYKe7m5KpQKymKG3UOLIG/6TWMtO2LYC40VAKZQRRIIubDSy7AG4FRh3BOr3n8WbupL2y87m8v/qZXNjM/9+x53866/W8MM3X+anMytIp3PEq6uYv2AxV113M11iIu6gk5D2HjgjCZWjMQZ8qRCAq7IHC8Pax3E8LCcWXjMajA/ZHowGElEKhaA/FvZbUCmFQVAolijm07S3NuOcch0HVsyla8lTOPEKjJShVAvakcseRPkSacqMZsUQyXro7kTeNxt32j7avzSTGWddzcNv7eWFJS9z81Sb1q4cyVQV8xe8w3Vf/i49sbPxRpyJbHqBiO7FG2UwsQq0FgTKIKUPQQ9oVQ6lBmE5oULqA4gCP48JwHhOGCLEYRJepTRSSXo7m/Eq6ykue4F8eh9eIo6WMmQv28N07qAQGJTwDsooLwWxFFQmES07EDddjUlWM3XqKfz2+QcY5+Rp7g6oSVWz4O15XH/ri3TV3IWXdNAffAlT3Ec0dgYmIijYIQP6UqNUgKXz/Q0bCPNHy7HL7qnDoK8lmNATY1G3v7Fj9ck713VCVrIcvFgCWxdxSh24sQq0kgfZS0A2V+Dqq65k2IgRZHNB6ONeHSRqERhs28Jf+z4npD7FyRdeDY0byGmPymQFi99dzPX3tdA58T/wBo9ArboeMrsQxNGWSyh6wkQ4kJqgWEIWshjH6osOKD+H0aovOwCtwbHBAsf3qajwsISFGWjBiOeilSRWkcTzPHy/ACJ0BQhrJ7Yl8P0iZ559Bvfefz+9mRy2a0IXjQyHMcMRrsFXSc6+9QlSE0/j5187nbdH+cxbuJjdjU1c/3iGzjOuw9uzEvW7SyDfDl4lqCLKaChokBqDoVAoUZVMUqU76Sr6iDLnG2OQxQx2tBJwwVgQrw5xlAoh54kw0vfvwUjEQylFRaKSZLKyHFB1WRqZARJCM2zIEShs8oUigVaokoKxp0HXGwQtBS779mvEBk/h+Ts+QySyj8ad25l93iy+uLlE+/UX4Jbmol44F/Kd4MQxRiIAXwSUjIUlSwhLUCgWqKobyh3/9FncoBODKKt7ERaNS7lyduNBqhaKmrjMEYm5/WLA6tN3nutiVIATiVPfMAhZ3nOmXMsUB7t6FItF/EAiNRQLBSI1oyD/DubXt3P5bU/Rlcnz6t2zcUsrUcrgxZKs37SR3fN+ibPhWfRDl0Mxi3AiCBOEru+6qB0fojc1IiqrkKUSgVR0dHZxzIkzOWfWqShtsCy73D21yyDTCJ3AjBiK1dPDiHhAMhnvd+d+C9q2TcwTWLbDqLETQkVuBjYpTb+A1eUqt+/7xBLVrF35Ds7iBznv8/ezZf1iFv3kGjx7NwgbS4DRBi8ax/3945hbrgalwXHBKBAWAhCOhWjbC/d9GlFYTUXlIPKFPL4MyBWKeBXJ/qxHYCGEQFh2aE17MHriSKJ79zFpeJRUVSV6YBzs+8ea6jgqkEw86phyWUJ/pJzRlyYVSz619YN44alfc/+/3cdJl36T1QufZNOce/AqFOBgWSAsgbCscC2jCaxoEuG4ZSofkMwoCY5AtG9Dv30zb81fQE2qHtvxKBZLSKkO9r4tG4QNloOwIoj6qZjhtdTs2cX48TUkYrFDXbTvNWRwHTrIM3biVI4YOgoZHOy3DSzn+EFAbV01Lz33G+7+1j2MPfZS1i3+DS0bX8WLxRFYWFbZ3iJc7XDlD7bcwsMKGbisUIySYIcLc+/dN/Ptm79IKV+ktmEI2lgHiV/YGGGD7QE1cOwpoApMkZ1MmDgM13EOX7JIJpPUJByisQQnnDITpcLGZhiArPKKCyqrUvzumd9yz/2PMWT8GWxd/SKZtq04kURo+XJFPJCSQAZIVS4NlpfJlJsyupxfKqmQftBnSsDGjSZ4Z97L3HjdBcz93dM4jndw7wkbLBdcD7zx6FnTia3dwLGjXYYPGXxIv7+/+RISpWFn424Wr92HH/h869avUMynwzoKICwHWcwybMgwZMNMMr0FsjteR7iC6qoqlJL0dnVQU9+A1qrMzBJjwkqd1hqlJJ7noZQOyxb5LKmaOmIVFXR0dqGUAWGw7DpsJ0KxeABMJRUJj1K2FWE5YLvgREHEYfptyFuv56jHHuf7N0xg1mkz8Fy7Pzm277777u+G7hK6YTIRY+/eZiKVQ9CqxLrVK/EiLgYLXSpiJerIj7mGKWMGcefnjmNH4x6eeuIXLFu2jNtv/Tq19XVcefklbNq4gR/c9z2++uV/pL2tjbFjRnPyySewY2cjd915G7lcjosvPJ9ASh788Q+YOuVIVq1aSzafR/gVePf+Cs69BPP6HCru/QWlNLBnA3gewvLA88A5Gf3N23HXbOYzqQNcfMEMaquT5cLI4apqCDwvypSJw+hpb+Gyqz7PyLFHojWIoIRdO4q6826DA2tRB5YwecoU7vyX29CWw/79zTz19DPcdOvtPPzwT9m5Yyv/8dNHWLx8Lc899wzxqhquvOoabrzxRhKpQXiRGHY0waxzzuWZ51/mtttup6c3g6WB46+gOOlMShNPhaMvpejUQHIY2BWh5bwIBEfAZV9B13gcvWklF100hYaa1EFF1xcd+izYl1UYY0hVV9Ld2U5BVTB58mTm/+EtrOrBWJNmUVw3j1LTMhKpGlasWM7PHv4xLZ0Z2lv3s3vXbva1tLNq2RIsy8KLJvAVbNv8IfGqWuYvWMhzTz2BG0+xa+d2fGlYOH8ep593OVOPO4ntWzdSyBawZt4Iaz9EbPwQKkcg2nKwbwt0bkZ4MYyIw5E3om+7mIbH3+arM21mnX4c1ZXxP9/C7usT9PSkeXXeGkZNnMbvX3yBR594Gi/fhF/IIGwXHZT6a5TCBP0lfy19bMfFINBaglY4bgQZ+OGKOg5KlhCWC8ZgTEj/jhtDaQG2i/AaEKVcuGcS9YANlIAiwomga7+M88jXcRas57O9K7nlG2cydvgwHNv68w3QPmtWVVVy2vFjmbtkPdd/4QukO7t47umf48UqCeI1OIlKUm41juvQ2vQhDU6Ruvo68vk8tiUolnwqYjFaWlsRCGprh6G0JvBLJJOVtLV3IKWkoaEeIQStrR1UV1fhuC4d7Z1kk1VYw6aC9DFBAXIFRBBBp67De+Amotv2cerulVxz07EMHzToI+D+7BBCWBORrNuwjSVr2zj/nNP4wQ8f5tmnf4VHhlIQsqMAtA5wbQvbtg4JAbZthdMVRuM4DqKcd9qOHV7XJrxuWQSBxHEcLNtBSoPUhGwpXHBshF2HGvY1ovd+norOdia9MoebvjCCM048hprqysOOgv3ZHr1SGt/3WbN+C++v72T2Bafz2OMv8Oh/Pwql5rA1JmUZUNhQEQf1XX961SfvhCkL5bI7CwRhdBTl9CYU0cKyD4YD10M4U1HHfYP4t8+gYksb496Yzz9dP5iZM6bSUFd9iIj4xEMISilKpRIr127mvQ9amX3hTNZ8sJl7fvgzWvauxZJp/MAvd21N/yRFH0BjdF8O0M9wA8aIykWWsmwToUrBdsC2EXYDuup8xBVfInntULzXd3DcppV86dphnHTMZAbVVf/Z0ZePNWUR7p2ALdsbmb90N1OnHEVdKsEjj73Cm2++RqGnEaGySClRypSnKXQ5xTJ/YlBuICgn1JeWBbaLJWrQFTPQM64i+g+nEq+H2G+Wc7rbyFWXj+T4oyZSV1PVP/D3VwPsG0aQUnGgpYU5CzdgnCpOPO4omg508+xLC3j//cX0tG1HyB6ELpVBht2nfhuWXbQPXJgNWGC5GOJoaygmdQxi2ll4F51IbJKF+/4+6uYv5+xpeS6/eCqTRo+kpjr5sQYIP9GsminLNSkVxUKedZsaWbB0FzV1Q5g0cTzpnOa9Vdt4f8Vadu34kEx3M6qUBl1CGBWWQxBh0dGKoEUE7GqINEBqHNaoqbjHTsM+djBeDbgfHqDmvfWc1pDmzNNrOXrSaEYOHUxFLPKJxs4+8TBe3/SF0or29g5Wr2/kgy3tKBFnxKixRBP19BQETa0ZGps7aG7roiOTJVssUlCgXQ8rlsCprsGrbyA6pI54fQSikOnKY23fy7DmJka7XZxxaoqTZ0ygoaaWREUsZOk/MRf6Pwbw0AEh0ErSnU6zvXE/azbsY29LHjeWwo3X4sRr0ZEqZCROyXbIC4tebegxipwfIAt5VE8Wq60Lr6eDoV6W6RNiTDmqlqFH1FJbnaKqMvEnY9zfDOAhIMsaFsD3i2QyWdq7etnV1Mnu/b10pCU9WRBOFNv1yBUCLKPwLE0qblFX6TBqeJKGuhg1qQSVlUliXhTPc7Gtv35QV5RKpb94IPaPwQohDlHyxihkEOBL2X8tkAqlw1TJLQd51wnLfJYl+ksj5iNp9l8J8H9iMP2j+6Nv2HxgVcD0Zy4DCyEHh9c/Khv/0pfWGmfgyOP/pVff4vw/VGoSDwDS8lIAAAAASUVORK5CYII=") center/cover no-repeat;
    box-shadow:0 2px 8px rgba(0,0,0,.25); flex-shrink:0; }
  .msg-row.error { align-self:center; align-items:center; max-width:90%; }
  .msg-row .msg { max-width:100%; }
  .msg-actions { display:flex; gap:4px; margin-top:3px; opacity:0; transition:opacity .15s;
    height:22px; }
  .msg-row:hover .msg-actions, .msg-row:focus-within .msg-actions { opacity:1; }
  .msg-actions button { background:none; border:none; color:var(--muted); cursor:pointer;
    font-size:12px; padding:2px 7px; border-radius:5px; touch-action:manipulation;
    -webkit-tap-highlight-color:transparent; }
  .msg-actions button:hover { background:var(--panel); color:var(--text); }
  .msg-timestamp { font-size:10.5px; color:var(--muted); margin-top:2px; }
  .edit-box { display:flex; flex-direction:column; gap:8px; min-width:220px; width:100%; }
  .edit-textarea { width:100%; box-sizing:border-box; resize:none; background:var(--bg);
    border:1px solid var(--accent); border-radius:8px; color:var(--text); font-family:inherit;
    font-size:14.5px; line-height:1.4; padding:8px 10px; outline:none; max-height:200px; }
  .edit-controls { display:flex; gap:8px; justify-content:flex-end; }
  .edit-save-btn, .edit-cancel-btn { border:none; border-radius:8px; padding:6px 14px;
    font-size:12.5px; font-weight:600; cursor:pointer; font-family:inherit;
    touch-action:manipulation; -webkit-tap-highlight-color:transparent; }
  .edit-save-btn { background:var(--accent); color:#fff; }
  .edit-save-btn:disabled { opacity:.6; cursor:not-allowed; }
  .edit-cancel-btn { background:none; border:1px solid var(--border); color:var(--muted); }
  .edit-cancel-btn:hover { color:var(--text); }
  .empty-state { position:absolute; top:50%; left:50%; transform:translate(-50%,-50%);
    text-align:center; color:var(--muted); }
  .empty-state h2 { font-size:25px; font-weight:700; color:var(--accent); margin-bottom:8px; }
  .empty-state p { font-size:16px; }
  .typing { align-self:flex-start; display:flex; gap:5px; padding:14px 16px;
    background:var(--ai-bubble); border-radius:18px; border-bottom-left-radius:4px; }
  .typing span { width:7px; height:7px; border-radius:50%; background:var(--muted);
    animation:blink 1.2s infinite ease-in-out; }
  .typing span:nth-child(2) { animation-delay:.2s; }
  .typing span:nth-child(3) { animation-delay:.4s; }
  @keyframes blink { 0%,80%,100%{opacity:.2} 40%{opacity:1} }

  #scroll-btn { position:fixed; bottom:130px; right:24px; width:36px; height:36px;
    border-radius:50%; background:var(--accent); color:#fff; border:none; cursor:pointer;
    font-size:18px; display:none; align-items:center; justify-content:center;
    box-shadow:0 2px 8px rgba(0,0,0,.15); z-index:10; }
  #scroll-btn.show { display:flex; }

  .gen-img { max-width:320px; border-radius:12px; display:block; margin-top:8px; }

  #pending-attach { max-width:760px; margin:0 auto; width:100%; padding:6px 20px 0;
    display:none; align-items:center; gap:8px; font-size:12.5px; color:var(--muted); flex-shrink:0; }
  #pending-attach.show { display:flex; }
  #pending-attach button { background:none; border:none; color:var(--muted); cursor:pointer; }
  .input-area { padding:10px 26px 22px; border-top:none;
    background:transparent; max-width:760px; margin:0 auto; width:100%; flex-shrink:0; }
  .input-row { display:flex; flex-direction:column; gap:6px;
    background:var(--composer-bg); border:1px solid var(--composer-border);
    border-radius:30px; padding:16px 16px 10px;
    box-shadow:var(--composer-shadow), 0 24px 60px -18px rgba(16,163,127,.35), 0 24px 60px -18px rgba(13,122,95,.25);
    backdrop-filter:blur(28px) saturate(1.4); -webkit-backdrop-filter:blur(28px) saturate(1.4);
    transition:box-shadow .2s ease, border-color .2s ease, transform .15s ease; }
  .input-row:focus-within { border-color:var(--accent);
    box-shadow:var(--composer-shadow), 0 0 0 3px var(--accent-dim),
      0 24px 60px -18px rgba(16,163,127,.45), 0 24px 60px -18px rgba(13,122,95,.35); }
  .composer-bottom-row { display:flex; align-items:center; justify-content:space-between; }
  .tool-btn { background:rgba(255,255,255,.06); border:1px solid transparent; color:var(--muted); cursor:pointer;
    width:42px; height:42px; border-radius:50%; font-size:19px; flex-shrink:0;
    display:flex; align-items:center; justify-content:center;
    backdrop-filter:blur(6px); -webkit-backdrop-filter:blur(6px);
    transition:background .15s ease, color .15s ease, transform .1s ease;
    touch-action:manipulation; -webkit-tap-highlight-color:transparent; }
  .tool-btn:hover { background:var(--accent-dim); color:var(--accent); transform:translateY(-1px); }
  .tool-btn.active { color:var(--accent); background:var(--accent-dim); }
  textarea { width:100%; resize:none; background:transparent; border:none; color:var(--text);
    font-size:16px; font-family:inherit; line-height:1.5; max-height:160px; min-height:26px;
    outline:none; padding:2px 4px 8px; }
  textarea::placeholder { color:var(--muted); }
  #send-btn { background:linear-gradient(135deg,#10a37f,#0d7a5f); color:#fff; border:none; border-radius:50%;
    width:42px; height:42px; font-size:19px; cursor:pointer; flex-shrink:0;
    display:flex; align-items:center; justify-content:center;
    box-shadow:0 4px 14px rgba(16,163,127,.35);
    touch-action:manipulation; -webkit-tap-highlight-color:transparent; }
  #send-btn:disabled { background:var(--accent-dim); color:var(--muted); cursor:not-allowed; box-shadow:none; }
  #send-btn.generating { background:#ef4444; }
  #send-btn.generating:hover { opacity:.9; }
  #voice-btn.listening { color:#ef4444; animation:pulse 1s infinite; }
  @keyframes pulse { 0%,100%{opacity:1} 50%{opacity:.4} }

  #live-talk-btn.live-listening { color:#fff; background:#ef4444; animation:pulse 1s infinite; }
  #live-talk-btn.live-speaking { color:#fff; background:var(--accent-grad); }
  #live-talk-btn.live-idle-active { color:var(--accent); background:var(--accent-dim); }

  #live-talk-status { display:none; align-items:center; gap:8px; position:fixed;
    left:50%; transform:translateX(-50%); bottom:calc(120px + env(safe-area-inset-bottom));
    background:rgba(15,15,22,.85); color:#fff; padding:8px 16px 8px 12px; border-radius:24px;
    font-size:12.5px; font-weight:600; z-index:300; backdrop-filter:blur(14px);
    box-shadow:0 10px 30px rgba(0,0,0,.35); pointer-events:none; }
  #live-talk-status.show { display:flex; }
  #live-talk-status .dot { width:8px; height:8px; border-radius:50%; background:#ef4444;
    animation:pulse 1s infinite; flex-shrink:0; }
  #live-talk-status.speaking .dot { background:var(--accent2); animation:none; }
  #live-talk-status button { pointer-events:auto; background:none; border:1px solid rgba(255,255,255,.25);
    color:#fff; border-radius:14px; padding:2px 10px; font-size:11px; cursor:pointer; margin-left:4px; }

  #video-call-btn.active, #screen-share-btn.active { color:#fff; background:#ef4444; }

  #live-video-wrap { display:none; position:fixed; right:16px;
    bottom:calc(120px + env(safe-area-inset-bottom)); z-index:301;
    border-radius:14px; overflow:hidden; box-shadow:0 10px 30px rgba(0,0,0,.4);
    border:2px solid rgba(255,255,255,.25); background:#000; }
  #live-video-wrap.show { display:block; }
  #live-video-preview { width:180px; max-width:38vw; height:auto; display:block; }
  #live-video-stop-btn { position:absolute; top:4px; right:4px; background:rgba(0,0,0,.6);
    color:#fff; border:none; border-radius:50%; width:22px; height:22px; font-size:12px;
    cursor:pointer; line-height:1; }

  #speaking-indicator { display:none; align-items:center; gap:6px; font-size:12px;
    color:var(--accent); padding:4px 0; flex-shrink:0; }
  #speaking-indicator.show { display:flex; }
  #stop-speak-btn { background:none; border:1px solid var(--border); color:var(--muted);
    font-size:11px; padding:2px 8px; border-radius:4px; cursor:pointer; }
  #ai-disclaimer { max-width:760px; margin:6px auto 0; width:100%; padding:0 20px;
    text-align:center; font-size:11.5px; line-height:1.5; color:var(--muted);
    flex-shrink:0; }
  .quick-btn { background:var(--panel); border:1px solid var(--border); color:var(--text);
    font-size:12.5px; padding:6px 14px; border-radius:20px; cursor:pointer;
    transition:all .15s ease; white-space:nowrap; font-family:inherit; touch-action:manipulation; }
  .quick-btn:hover { background:var(--accent-dim); border-color:var(--accent); color:var(--accent); }
  #quick-actions { flex-shrink:0; }

  .mode-tab { display:flex; align-items:center; gap:6px; background:none; border:1px solid transparent;
    color:var(--muted); font-size:13px; font-family:inherit; padding:7px 12px; border-radius:8px;
    cursor:pointer; white-space:nowrap; touch-action:manipulation; }
  .mode-tab:hover { background:var(--panel); color:var(--text); }
  .mode-tab.active { background:var(--accent-dim); color:var(--accent); border-color:var(--accent); font-weight:600; }
  .mode-tab-lock { font-size:10px; opacity:.7; }
  .mode-tab.active .mode-tab-lock, .mode-tab.unlocked .mode-tab-lock { display:none; }

  #messages-wrap::-webkit-scrollbar, #conv-list::-webkit-scrollbar { width:6px; }
  #messages-wrap::-webkit-scrollbar-thumb, #conv-list::-webkit-scrollbar-thumb
    { background:var(--border); border-radius:4px; }
  #sidebar-overlay { display:none; position:fixed; inset:0; background:rgba(0,0,0,.55);
    z-index:99; -webkit-tap-highlight-color:transparent; }

  @media(max-width:768px) {
    :root { --sidebar-w: 78vw; }

    .layout { gap:0; padding:0; }
    .app { border-radius:0; border:none; box-shadow:none; }
    header::before { border-radius:0; }
    #sidebar { border-radius:0; border-right:1px solid var(--border); border-top:none; border-left:none; border-bottom:none; }

    #sidebar { position:fixed; top:0; left:0; z-index:100; height:100%;
      height:-webkit-fill-available; width:var(--sidebar-w) !important;
      transform:translateX(0); transition:transform .25s ease;
      box-shadow:4px 0 24px rgba(0,0,0,.5); }
    #sidebar.hidden { transform:translateX(-105%); margin-left:0 !important; }

    #sidebar-overlay { display:block; }

    .app { width:100% !important; flex:1; }

    header { padding:calc(10px + env(safe-area-inset-top)) 10px 8px;
      flex-wrap:wrap; row-gap:6px; }
    header .left { flex-wrap:wrap; row-gap:4px; }
    header .right { flex-wrap:wrap; justify-content:flex-end; row-gap:6px; gap:6px; }
    header h1 { font-size:14px; }
    #sidebar-toggle { width:44px; height:44px; font-size:15px; }
    #name-btn { width:44px; height:44px; font-size:15px; }
    #settings-btn { width:44px; height:44px; font-size:15px; }
    #export-btn { width:44px; height:44px; font-size:15px; }
    #share-btn { width:44px; height:44px; font-size:15px; }
    #vip-btn { width:44px; height:44px; font-size:15px; }
    #clear-btn { font-size:12.5px; padding:10px 12px; min-height:44px; }
    #speak-toggle { font-size:12.5px; padding:8px 10px; min-height:44px; }
    #fullscreen-btn { width:44px; height:44px; font-size:15px; }
    #install-btn { padding:9px 12px; font-size:12.5px; min-height:44px; }

    #messages-wrap { overflow-y:auto; -webkit-overflow-scrolling:touch; }
    #messages { padding:14px 10px; gap:12px; max-width:100%; }
    .msg { max-width:90%; font-size:15px; padding:11px 13px; }
    .msg-row { max-width:90%; }
    .msg-actions { opacity:1; height:36px; }
    .msg-actions button { font-size:14px; padding:6px 12px; min-width:40px; min-height:36px; }

    .input-area { padding:8px 10px max(8px,env(safe-area-inset-bottom)); }
    .input-row { padding:12px 12px 8px; border-radius:22px; }
    textarea { font-size:16px; padding:2px 2px 6px; }
    .tool-btn { width:44px; height:44px; font-size:18px; }
    #send-btn { width:44px; height:44px; font-size:18px; }
    #ai-disclaimer { font-size:10.5px; padding:0 12px; }

    .empty-state h2 { font-size:19px; }
    .empty-state p { font-size:13px; }
    #scroll-btn { bottom:80px; right:12px; width:40px; height:40px; }

    #new-chat-btn { margin:10px; padding:12px; font-size:14.5px; min-height:44px; }
    .conv-item { padding:12px 10px; font-size:14px; min-height:48px; }
    .conv-item .menu-btn { opacity:1; width:36px; height:36px; }
    #sidebar-footer { font-size:12px; padding:12px; }
    #search-chats-btn, #reminders-btn { padding:10px 4px; font-size:13px; min-height:40px; }
    .quick-btn { font-size:13.5px; padding:9px 16px; min-height:40px; }
    .mode-tab { font-size:14px; padding:10px 12px; min-height:40px; }
  }

  @media(max-width:380px) {
    :root { --sidebar-w: 88vw; }
    .msg { font-size:13.5px; }
    header h1 { font-size:13px; }
    #speak-toggle { display:none; }
  }
</style>
<script src="https://cdnjs.cloudflare.com/ajax/libs/qrcodejs/1.0.0/qrcode.min.js"></script>
</head>
<body>
<div class="layout">
  <div id="sidebar-overlay" style="display:none;position:fixed;inset:0;background:#0007;z-index:99"></div>
  <div id="sidebar">
    <button id="new-chat-btn">+ New chat</button>
    <button id="api-keys-shortcut-btn" title="Manage API keys">🔑 API Keys</button>
    <div id="api-usage-summary" style="display:none;margin:0 12px 6px;padding:8px 10px;border:1px solid var(--border);border-radius:8px;font-size:11.5px;color:var(--muted);cursor:pointer;" title="Click to manage API keys"></div>
    <div style="display:flex;gap:6px;margin:6px 0;">
      <button id="search-chats-btn" style="flex:1;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:7px 4px;font-size:12px;cursor:pointer;font-family:inherit;">🔎 Search</button>
      <button id="reminders-btn" style="flex:1;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:7px 4px;font-size:12px;cursor:pointer;font-family:inherit;">⏰ Reminders</button>
    </div>
    <div id="conv-list"></div>
    <div id="sidebar-tools" style="display:none;">
      <div id="mode-tab-bar">
        <button class="mode-tab active" data-mode="chat" title="Regular chat">
          💬 <span>Chat</span>
        </button>
        <button class="mode-tab" data-mode="cowork" title="VIP — multi-step task assistant">
          🗂 <span>Cowork</span> <span class="mode-tab-lock">🔒</span>
        </button>
        <button class="mode-tab" data-mode="code" title="VIP — coding-focused assistant">
          &lt;/&gt; <span>Code</span> <span class="mode-tab-lock">🔒</span>
        </button>
        <button class="mode-tab" id="artifacts-tab-btn" title="VIP — saved code/text snippets from replies">
          📦 <span>Artifacts</span> <span class="mode-tab-lock">🔒</span>
        </button>
      </div>
      <div id="quick-actions">
        <button class="quick-btn" id="img-gen-btn">🎨 Image</button>
        <button class="quick-btn" id="ghibli-btn">🌿 Ghibli Me</button>
        <button class="quick-btn" id="file-gen-btn">📄 File / PDF</button>
        <button class="quick-btn" id="homework-btn">📚 Homework & Study</button>
        <button class="quick-btn" id="weather-btn">🌤 Weather</button>
        <button class="quick-btn" id="search-btn">🔍 Search</button>
        <button class="quick-btn" id="code-workspace-btn">💻 Code</button>
      </div>
    </div>
    <div id="sidebar-footer">
      <div style="display:none;">
        <button id="archived-toggle-btn">⭐ Starred</button>
        <button id="bookmarks-btn">🔖 Bookmarks</button>
        <button id="stats-btn">📊 Stats</button>
      </div>
      <button id="sidebar-profile" type="button" title="Your profile — click to edit">
        <span id="sidebar-profile-avatar" aria-hidden="true"></span>
        <span id="sidebar-profile-text">
          <span id="sidebar-profile-name">Guest</span>
          <span id="sidebar-profile-sub">Your profile</span>
        </span>
      </button>
      <div id="sidebar-byline">Mythic AI &middot; by Aarav Singh</div>
    </div>
  </div>
  <div class="app">
    <header>
      <div class="left">
        <button id="sidebar-toggle" title="Toggle sidebar">☰</button>
        <h1>Mythic AI</h1>
        <span id="vip-badge" style="display:none;background:linear-gradient(135deg,#f5c542,#e0a800);color:#1a1a1a;font-size:10.5px;font-weight:800;padding:3px 8px;border-radius:10px;letter-spacing:.3px;">VIP</span>
        <span id="streak-badge" title="Daily chat streak">🔥 0</span>
      </div>
      <div class="right">
        <button id="install-btn" title="Install Mythic AI" style="display:flex;background:var(--accent);color:#fff;border:none;border-radius:8px;padding:6px 12px;font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;touch-action:manipulation;align-items:center;gap:4px;">⬇ Install</button>
        <div id="header-menu-wrap">
          <button id="header-menu-btn" type="button" title="More" aria-haspopup="true" aria-expanded="false">⋯</button>
          <div id="header-menu-dropdown" role="menu">
            <button id="vip-btn" title="Mythic VIP" data-label="VIP">✨</button>
            <button id="fullscreen-btn" type="button" title="Fullscreen">
              <span id="fullscreen-icon">⛶</span><span>Fullscreen</span>
            </button>
            <button id="name-btn" title="What should Mythic AI call you?">🙂<span>Nickname</span></button>
            <button id="settings-btn" title="Settings">⚙<span>Settings</span></button>
            <button id="share-btn" title="Get invite link">🔗<span>Share link</span></button>
            <button id="export-btn" title="Export this chat">⬇<span>Download</span></button>
            <div class="menu-divider"></div>
            <button id="clear-btn">Delete chat</button>
            <button id="logout-btn">Log out</button>
          </div>
        </div>
      </div>
    </header>

    <div id="messages-wrap">
      <div id="messages">
        <div class="empty-state" id="empty-state">
          <h2>Mythic AI</h2>
          <p>Ask me anything, generate images, or just chat 👋</p>
        </div>
      </div>
    </div>


    <button id="scroll-btn" title="Scroll to bottom">↓</button>

    <div id="pending-attach">
      📎 <span id="pending-attach-name"></span>
      <button id="pending-attach-remove">✕</button>
    </div>

    <div id="speaking-indicator">
      🔊 Speaking...
      <button id="stop-speak-btn">Stop</button>
    </div>

    <div id="live-talk-status">
      <span class="dot"></span>
      <span id="live-talk-status-text">Listening…</span>
      <button id="live-talk-stop-btn">End</button>
    </div>

    <div id="live-video-wrap">
      <video id="live-video-preview" autoplay muted playsinline></video>
      <button id="live-video-stop-btn" title="Stop sharing">✕</button>
    </div>

    <!-- In-page camera capture: takes a photo without leaving the app for
         the phone's native camera app (unlike a plain <input capture>). -->
    <div id="camera-modal-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.92);z-index:400;flex-direction:column;align-items:center;justify-content:center;">
      <video id="camera-modal-video" autoplay muted playsinline style="max-width:100%;max-height:70vh;border-radius:12px;background:#000;"></video>
      <canvas id="camera-modal-canvas" style="display:none;"></canvas>
      <div style="display:flex;gap:16px;margin-top:20px;">
        <button id="camera-modal-cancel" type="button" style="background:none;border:1px solid rgba(255,255,255,.4);color:#fff;border-radius:24px;padding:12px 22px;font-size:14px;cursor:pointer;font-family:inherit;">Cancel</button>
        <button id="camera-modal-shoot" type="button" style="background:#fff;border:none;color:#111;border-radius:50%;width:60px;height:60px;font-size:22px;cursor:pointer;">📸</button>
      </div>
      <div id="camera-modal-error" style="display:none;color:#ff6b6b;font-size:13px;margin-top:14px;max-width:80vw;text-align:center;"></div>
    </div>

    <div id="notif-banner" style="display:none;align-items:center;justify-content:space-between;gap:10px;
      background:linear-gradient(135deg,var(--accent-dim),rgba(16,163,127,.15));
      border:1px solid var(--accent);border-radius:12px;padding:10px 14px;
      max-width:760px;margin:8px auto 0;width:calc(100% - 40px);flex-wrap:wrap;flex-shrink:0;">
      <div style="display:flex;align-items:center;gap:8px;flex:1;min-width:0;">
        <span style="font-size:20px;flex-shrink:0;">🔔</span>
        <div style="min-width:0;">
          <div style="font-size:13px;font-weight:600;color:var(--text);">Get notified when Mythic AI replies</div>
          <div style="font-size:11.5px;color:var(--muted);margin-top:1px;">Even when you switch to another tab</div>
        </div>
      </div>
      <div style="display:flex;gap:6px;flex-shrink:0;">
        <button id="notif-banner-allow" type="button"
          style="background:var(--accent);color:#fff;border:none;border-radius:8px;
            padding:7px 14px;font-size:12.5px;font-weight:600;cursor:pointer;font-family:inherit;
            white-space:nowrap;">
          Allow 🔔
        </button>
        <button id="notif-banner-dismiss" type="button"
          style="background:none;border:1px solid var(--border);color:var(--muted);
            border-radius:8px;padding:7px 10px;font-size:12.5px;cursor:pointer;font-family:inherit;">
          ✕
        </button>
      </div>
    </div>

      <form id="chat-form">
        <div class="input-row">
          <input type="file" id="file-input" accept="image/*,.txt,.md,.csv,.json,.pdf,.docx" style="display:none">
          <input type="file" id="camera-input" accept="image/*" capture="environment" style="display:none">
          <textarea id="input" rows="1" placeholder="Write a message..."></textarea>
          <div class="composer-bottom-row">
            <button class="tool-btn" id="plus-btn" type="button" title="Add photo or file">
              <svg width="22" height="22" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                <line x1="12" y1="5" x2="12" y2="19"/>
                <line x1="5" y1="12" x2="19" y2="12"/>
              </svg>
            </button>
            <button class="tool-btn" id="live-talk-btn" type="button" title="Live Talk — speak and Mythic AI replies out loud">
              <svg width="19" height="19" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                <path d="M12 1a3 3 0 0 0-3 3v8a3 3 0 0 0 6 0V4a3 3 0 0 0-3-3z"/>
                <path d="M19 10v2a7 7 0 0 1-14 0v-2"/>
                <line x1="12" y1="19" x2="12" y2="23"/>
                <line x1="8" y1="23" x2="16" y2="23"/>
              </svg>
            </button>
            <button class="tool-btn" id="video-call-btn" type="button" title="Video Call — Mythic AI can see your camera and talk with you">
              <svg width="19" height="19" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                <polygon points="23 7 16 12 23 17 23 7"/>
                <rect x="1" y="5" width="15" height="14" rx="2" ry="2"/>
              </svg>
            </button>
            <button class="tool-btn" id="screen-share-btn" type="button" title="Screen Share — Mythic AI can see your screen and talk with you">
              <svg width="19" height="19" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                <rect x="2" y="3" width="20" height="14" rx="2" ry="2"/>
                <line x1="8" y1="21" x2="16" y2="21"/>
                <line x1="12" y1="17" x2="12" y2="21"/>
              </svg>
            </button>
            <button id="send-btn" type="submit" title="Send">
              <svg width="18" height="18" viewBox="0 0 24 24" fill="currentColor">
                <path d="M2.01 21L23 12 2.01 3 2 10l15 2-15 2z"/>
              </svg>
            </button>
          </div>
        </div>
      </form>
      <div id="ai-disclaimer">Mythic AI is AI and can make mistakes. Please double-check responses.</div>
    </div>
  </div>
</div>

<div id="name-modal-overlay">
  <div id="name-modal">
    <h3>Your profile</h3>
    <p>Set a photo and the name Mythic AI will use when it talks to you.</p>
    <div id="avatar-upload-row">
      <span id="avatar-upload-preview" aria-hidden="true"></span>
      <div id="avatar-upload-actions">
        <label id="avatar-upload-btn" for="avatar-file-input">Change photo</label>
        <button id="avatar-remove-btn" type="button">Remove</button>
        <input type="file" id="avatar-file-input" accept="image/*" style="display:none">
      </div>
    </div>
    <p id="avatar-preset-label" style="margin:0 0 8px;font-size:12px;color:var(--muted);">Or pick a Mythic avatar</p>
    <div id="avatar-preset-grid"></div>
    <input type="text" id="name-input" maxlength="60" placeholder="e.g. Aarav" autocomplete="off">
    <div id="name-modal-actions">
      <button id="name-cancel-btn" type="button">Cancel</button>
      <button id="name-save-btn" type="button">Save</button>
    </div>
  </div>
</div>

<div id="share-modal-overlay">
  <div id="share-modal">
    <h3>🔗 Your account link</h3>
    <p class="sub">A permanent, unique link for THIS account — nobody else's. Bookmark it to always get back to these exact chats from any device, or share it with someone you want to give access to this specific account.</p>
    <div id="share-link-row">
      <input type="text" id="share-link-input" readonly>
      <button id="share-open-btn" type="button" title="Open in a new tab">↗</button>
      <button id="share-copy-btn" type="button">Copy</button>
    </div>
    <div id="share-qr-wrap" style="display:flex;justify-content:center;margin:16px 0;">
      <div id="share-qr-box" style="position:relative;width:220px;height:220px;background:#fff;padding:12px;border-radius:12px;">
        <div id="share-qr-canvas" style="width:100%;height:100%;"></div>
        <div id="share-qr-logo" style="position:absolute;top:50%;left:50%;transform:translate(-50%,-50%);
             width:52px;height:52px;background:#fff;border-radius:12px;display:flex;align-items:center;
             justify-content:center;box-shadow:0 0 0 4px #fff;">
          <img src="/icon.png" alt="Mythic AI" style="width:40px;height:40px;border-radius:8px;object-fit:cover;">
        </div>
      </div>
    </div>
    <button id="share-native-btn" type="button">📤 Share via…</button>
    <div style="display:flex;gap:8px;">
      <button id="share-qr-image-btn" type="button" style="flex:1;">🖼 Share QR Code</button>
      <button id="download-qr-btn" type="button" style="flex:1;">⬇ Download QR</button>
    </div>
    <button id="share-revoke-btn" type="button">Stop sharing</button>
    <button id="share-close-btn" type="button">Close</button>
    <div id="share-status"></div>
  </div>
</div>

<div id="api-usage-overlay" style="display:none;position:fixed;inset:0;background:#0f1115;color:#f2f2f2;z-index:200;overflow-y:auto;font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;">
  <div style="max-width:1000px;margin:0 auto;padding:32px 24px 60px;">
    <button id="api-usage-close-btn" style="background:none;border:none;color:#9a9ea6;font-size:14px;cursor:pointer;padding:0;margin-bottom:20px;font-family:inherit;">← Back to chat</button>
    <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:14px;margin-bottom:8px;">
      <div>
        <h1 style="font-size:30px;margin:0 0 6px;">API keys</h1>
        <div style="color:#9a9ea6;font-size:14px;margin-bottom:30px;max-width:560px;line-height:1.5;">Create and manage API keys for authenticating requests to Mythic AI. These keys allow programmatic access to your app.</div>
      </div>
      <button id="api-usage-gen-btn" style="background:#e8532a;color:#fff;border:none;border-radius:8px;padding:11px 18px;font-size:13px;font-weight:700;letter-spacing:.3px;cursor:pointer;white-space:nowrap;">GENERATE API KEY</button>
    </div>
    <div id="api-usage-totals" style="display:flex;gap:16px;margin-bottom:30px;flex-wrap:wrap;"></div>
    <table style="width:100%;border-collapse:collapse;background:#1a1d24;border:1px solid #2a2e37;border-radius:14px;overflow:hidden;">
      <thead><tr>
        <th style="text-align:left;font-size:11px;letter-spacing:.5px;text-transform:uppercase;color:#9a9ea6;padding:14px 16px;border-bottom:1px solid #2a2e37;">Name</th>
        <th style="text-align:left;font-size:11px;letter-spacing:.5px;text-transform:uppercase;color:#9a9ea6;padding:14px 16px;border-bottom:1px solid #2a2e37;">API Key</th>
        <th style="text-align:left;font-size:11px;letter-spacing:.5px;text-transform:uppercase;color:#9a9ea6;padding:14px 16px;border-bottom:1px solid #2a2e37;">Created At</th>
        <th style="text-align:left;font-size:11px;letter-spacing:.5px;text-transform:uppercase;color:#9a9ea6;padding:14px 16px;border-bottom:1px solid #2a2e37;">Calls</th>
        <th style="text-align:left;font-size:11px;letter-spacing:.5px;text-transform:uppercase;color:#9a9ea6;padding:14px 16px;border-bottom:1px solid #2a2e37;">State</th>
        <th style="text-align:left;font-size:11px;letter-spacing:.5px;text-transform:uppercase;color:#9a9ea6;padding:14px 16px;border-bottom:1px solid #2a2e37;">Options</th>
      </tr></thead>
      <tbody id="api-usage-tbody"></tbody>
    </table>
    <div id="api-usage-empty" style="display:none;color:#9a9ea6;font-size:15px;padding:50px 0;text-align:center;">No API keys yet. Click "Generate API Key" to create one.</div>
  </div>

  <div id="api-usage-create-overlay" style="display:none;position:fixed;inset:0;background:#000a;align-items:center;justify-content:center;z-index:210;">
    <div style="background:#1a1d24;border:1px solid #2a2e37;border-radius:14px;padding:26px;width:min(90vw,420px);">
      <h3 style="margin:0 0 14px;font-size:18px;">Generate API key</h3>
      <input type="text" id="api-usage-create-label" placeholder="Key name (optional)" maxlength="100" style="width:100%;padding:10px 12px;border-radius:8px;border:1px solid #3a3e47;background:#0f1115;color:#fff;font-size:14px;margin-bottom:14px;">
      <div id="api-usage-new-key-result"></div>
      <div style="display:flex;gap:10px;justify-content:flex-end;">
        <button id="api-usage-create-close-btn" style="border-radius:8px;padding:9px 16px;font-size:13px;cursor:pointer;border:none;background:#2a2e37;color:#f2f2f2;">Close</button>
        <button id="api-usage-create-confirm-btn" style="border-radius:8px;padding:9px 16px;font-size:13px;cursor:pointer;border:none;background:#e8532a;color:#fff;font-weight:700;">Generate</button>
      </div>
    </div>
  </div>
</div>

<div id="clarify-modal-overlay" style="display:none;position:fixed;inset:0;background:#000000b3;z-index:300;align-items:center;justify-content:center;padding:16px;">
  <div id="clarify-modal" style="background:#fff;color:#111;border-radius:14px;width:min(92vw,480px);max-height:80vh;overflow-y:auto;box-shadow:0 20px 60px #0006;">
    <div style="display:flex;justify-content:space-between;align-items:center;padding:18px 20px 6px;">
      <h3 id="clarify-question" style="margin:0;font-size:17px;font-weight:600;">What's wrong right now?</h3>
      <button id="clarify-close-btn" style="background:none;border:none;font-size:20px;cursor:pointer;color:#888;line-height:1;padding:4px;">×</button>
    </div>
    <div id="clarify-options" style="padding:8px 0 4px;"></div>
    <div style="padding:12px 20px 18px;border-top:1px solid #eee;display:flex;align-items:center;gap:10px;">
      <span style="color:#888;font-size:15px;">✎</span>
      <input id="clarify-custom-input" type="text" placeholder="Something else"
        style="flex:1;border:none;outline:none;font-size:15px;color:#333;background:transparent;">
      <button id="clarify-skip-btn" style="border:1px solid #ddd;background:#fff;border-radius:8px;padding:6px 14px;font-size:13px;cursor:pointer;color:#555;">Skip</button>
    </div>
  </div>
</div>

<div id="settings-modal-overlay">
  <div id="settings-modal">
    <h3>Settings</h3>
    <p class="sub">Customize how Mythic AI looks and replies. Saved on this device.</p>

    <div class="settings-section">
      <label>Theme</label>
      <div class="settings-row">
        <button class="settings-choice" data-group="theme" data-value="dark">🌙 Dark</button>
        <button class="settings-choice" data-group="theme" data-value="light">☀️ Light</button>
        <button class="settings-choice" data-group="theme" data-value="system">🖥 System</button>
      </div>
    </div>

    <div class="settings-section">
      <label>Accent color</label>
      <input type="color" id="accent-color-input" value="#10a37f">
    </div>

    <div class="settings-section">
      <label>Font size — <span id="font-size-label">14.5px</span></label>
      <input type="range" id="font-size-slider" min="12" max="20" step="0.5" value="14.5">
    </div>

    <div class="settings-section">
      <label>Bubble spacing</label>
      <div class="settings-row">
        <button class="settings-choice" data-group="bubble" data-value="compact">Compact</button>
        <button class="settings-choice" data-group="bubble" data-value="comfortable">Comfortable</button>
        <button class="settings-choice" data-group="bubble" data-value="spacious">Spacious</button>
      </div>
    </div>

    <div class="settings-section">
      <label>Reply tone</label>
      <select id="tone-select" class="settings-select">
        <option value="default">Default</option>
        <option value="formal">Formal</option>
        <option value="casual">Casual</option>
        <option value="funny">Funny</option>
        <option value="professional">Professional</option>
      </select>
    </div>

    <div class="settings-section">
      <label>Reply length</label>
      <select id="length-select" class="settings-select">
        <option value="default">Default</option>
        <option value="short">Short</option>
        <option value="medium">Medium</option>
        <option value="long">Long</option>
      </select>
    </div>

    <div class="settings-section">
      <label>Custom instructions</label>
      <textarea id="custom-instructions-input" placeholder="e.g. Always answer in bullet points"></textarea>
    </div>

    <div class="settings-section" id="api-keys-section" style="border-top:1px solid var(--border);padding-top:14px;margin-top:4px;">
      <label>🔑 API Keys</label>
      <div class="hint">Let other apps call Mythic AI like an OpenAI-style API. Keys start with
        <code>aarav-</code> and are shown only once — copy them right away.</div>
      <div style="display:flex;gap:8px;margin-top:8px;">
        <input type="text" id="api-key-label-input" class="settings-text-input"
          placeholder="Label (e.g. 'BattleZoneApp')" style="flex:1;">
        <button type="button" id="api-key-create-btn" class="settings-btn">+ Generate</button>
      </div>
      <div id="api-key-new-box" style="display:none;margin-top:10px;padding:10px;border:1px solid var(--accent);border-radius:8px;background:var(--bg);">
        <div style="font-size:12px;opacity:.8;margin-bottom:8px;display:flex;justify-content:space-between;align-items:center;flex-wrap:wrap;gap:8px;">
          <span>New key — copy it now, it won't be shown again:</span>
          <button type="button" id="api-key-copy-btn" style="background:var(--accent);color:#000;border:none;padding:4px 10px;border-radius:4px;cursor:pointer;font-size:11px;font-weight:700;white-space:nowrap;">Copy</button>
        </div>
        <input type="text" id="api-key-new-value" readonly style="width:100%;padding:10px;font-family:monospace;font-size:11px;letter-spacing:0.5px;background:var(--panel);border:1px solid var(--border);border-radius:6px;color:var(--text);box-sizing:border-box;cursor:text;overflow:auto;white-space:nowrap;" />
      </div>
      <div id="api-key-list" style="margin-top:10px;display:flex;flex-direction:column;gap:6px;"></div>
    </div>

    <div class="settings-section" style="border-top:1px solid var(--border);padding-top:14px;margin-top:4px;">
      <label style="display:flex;align-items:center;justify-content:space-between;cursor:pointer;">
        <span>🔔 Reply notifications</span>
        <button id="notif-toggle-btn" type="button"
          style="background:none;border:1.5px solid var(--border);color:var(--muted);border-radius:20px;padding:6px 14px;font-size:12px;cursor:pointer;font-family:inherit;transition:all .15s;">
          Enable
        </button>
      </label>
      <div id="notif-status" style="font-size:11.5px;color:var(--muted);margin-top:6px;"></div>
    </div>

    <div class="settings-section" style="border-top:1px solid var(--border);padding-top:14px;margin-top:4px;">
      <label>💾 Backup all chats</label>
      <p style="font-size:11.5px;color:var(--muted);margin:2px 0 8px;">Download every conversation as one file, or restore from a previous backup — also the only way to carry your chats over to a different browser/device, since accounts here are anonymous per-browser.</p>
      <div style="display:flex;gap:8px;">
        <button id="backup-export-btn" type="button" style="flex:1;background:none;border:1px solid var(--border);color:var(--text);border-radius:8px;padding:9px;cursor:pointer;font-family:inherit;font-size:12.5px;">⬇ Export all</button>
        <button id="backup-import-btn" type="button" style="flex:1;background:none;border:1px solid var(--border);color:var(--text);border-radius:8px;padding:9px;cursor:pointer;font-family:inherit;font-size:12.5px;">⬆ Import backup</button>
      </div>
      <input id="backup-import-file" type="file" accept=".zip" style="display:none;">
      <div id="backup-status" style="font-size:11.5px;color:var(--muted);margin-top:6px;"></div>
    </div>

    <button id="settings-close-btn" type="button">Done</button>
  </div>
</div>

<div id="ghibli-modal-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.8);z-index:300;align-items:center;justify-content:center;">
  <div style="background:var(--panel);border:1px solid var(--border);border-radius:16px;padding:24px;width:92%;max-width:440px;max-height:90vh;overflow-y:auto;">
    <h3 style="margin:0 0 4px;font-size:18px;">🌿 Ghibli Me</h3>
    <p style="color:var(--muted);font-size:13px;margin:0 0 16px;">Upload your photo and get a Studio Ghibli-style version of yourself</p>

    <div id="ghibli-upload-area" style="border:2px dashed var(--border);border-radius:12px;padding:24px;text-align:center;cursor:pointer;margin-bottom:12px;transition:border-color .2s;">
      <div style="font-size:36px;margin-bottom:8px;">📸</div>
      <div style="font-size:13px;color:var(--muted);">Click to upload your photo<br><span style="font-size:11px;">or drag & drop</span></div>
      <input type="file" id="ghibli-file-input" accept="image/*" style="display:none">
    </div>

    <div id="ghibli-preview-wrap" style="display:none;margin-bottom:12px;text-align:center;">
      <img id="ghibli-preview" style="max-width:100%;max-height:180px;border-radius:10px;border:2px solid var(--accent);">
      <div style="font-size:11px;color:var(--muted);margin-top:4px;">Your photo ✓</div>
    </div>

    <div style="margin-bottom:12px;">
      <label style="font-size:12px;color:var(--muted);display:block;margin-bottom:6px;">Ghibli Style:</label>
      <div style="display:grid;grid-template-columns:1fr 1fr;gap:6px;">
        <button class="ghibli-style-btn" data-style="Studio Ghibli portrait, Spirited Away style, soft watercolor anime art" style="padding:8px;border-radius:8px;border:1.5px solid var(--accent);background:var(--accent-dim);color:var(--accent);cursor:pointer;font-size:12px;font-family:inherit;">🌊 Spirited Away</button>
        <button class="ghibli-style-btn" data-style="Studio Ghibli portrait, My Neighbor Totoro style, soft forest anime art" style="padding:8px;border-radius:8px;border:1px solid var(--border);background:var(--panel);color:var(--muted);cursor:pointer;font-size:12px;font-family:inherit;">🌳 Totoro Forest</button>
        <button class="ghibli-style-btn" data-style="Studio Ghibli portrait, Howl's Moving Castle style, fantasy anime art" style="padding:8px;border-radius:8px;border:1px solid var(--border);background:var(--panel);color:var(--muted);cursor:pointer;font-size:12px;font-family:inherit;">🏰 Howl's Castle</button>
        <button class="ghibli-style-btn" data-style="Studio Ghibli portrait, Princess Mononoke style, nature anime art" style="padding:8px;border-radius:8px;border:1px solid var(--border);background:var(--panel);color:var(--muted);cursor:pointer;font-size:12px;font-family:inherit;">🐺 Mononoke</button>
      </div>
    </div>

    <input id="ghibli-extra" type="text" placeholder="Add details (optional): e.g. forest background, sunset..."
      style="width:100%;background:var(--bg);border:1.5px solid var(--border);color:var(--text);border-radius:8px;padding:9px 12px;font-size:13px;outline:none;margin-bottom:12px;font-family:inherit;">

    <div id="ghibli-result-wrap" style="display:none;margin-bottom:12px;text-align:center;">
      <img id="ghibli-result" style="max-width:100%;border-radius:12px;box-shadow:0 4px 20px rgba(0,0,0,.4);">
      <button id="ghibli-download-btn" style="margin-top:8px;background:var(--accent);color:#fff;border:none;border-radius:8px;padding:8px 18px;font-size:13px;cursor:pointer;font-family:inherit;">⬇ Download</button>
    </div>
    <div id="ghibli-loading" style="display:none;text-align:center;padding:20px;">
      <div style="font-size:32px;margin-bottom:8px;">🎨</div>
      <div style="color:var(--muted);font-size:13px;">Creating your Ghibli portrait...<br><span style="font-size:11px;">This can take up to a minute or two</span></div>
    </div>
    <div id="ghibli-error" style="display:none;color:#ef4444;font-size:12px;margin-bottom:8px;padding:8px;background:#fef2f2;border-radius:6px;"></div>

    <div style="display:flex;gap:8px;">
      <button id="ghibli-generate-btn" style="flex:1;background:linear-gradient(135deg,#10a37f,#0d7a5f);color:#fff;border:none;border-radius:10px;padding:12px;font-size:14px;font-weight:700;cursor:pointer;font-family:inherit;">✨ Create Ghibli Art</button>
      <button id="ghibli-close-btn" style="background:none;border:1px solid var(--border);color:var(--muted);border-radius:10px;padding:12px 16px;font-size:14px;cursor:pointer;">✕</button>
    </div>
  </div>
</div>

<div id="img-modal-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.8);z-index:300;align-items:center;justify-content:center;">
  <div style="background:var(--panel);border:1px solid var(--border);border-radius:16px;padding:24px;width:92%;max-width:440px;max-height:90vh;overflow-y:auto;">
    <h3 style="margin:0 0 4px;font-size:18px;">🎨 Generate Image</h3>
    <p style="color:var(--muted);font-size:13px;margin:0 0 16px;">Describe what you want to see</p>

    <textarea id="img-prompt" rows="3" placeholder="e.g. a cozy cabin in a snowy forest, golden hour lighting"
      style="width:100%;background:var(--bg);border:1.5px solid var(--border);color:var(--text);border-radius:8px;padding:10px 12px;font-size:13px;outline:none;margin-bottom:12px;font-family:inherit;resize:vertical;"></textarea>

    <div style="margin-bottom:14px;">
      <label style="font-size:12px;color:var(--muted);display:block;margin-bottom:6px;">Style (optional):</label>
      <select id="img-style" style="width:100%;background:var(--bg);border:1.5px solid var(--border);color:var(--text);border-radius:8px;padding:9px 12px;font-size:13px;outline:none;font-family:inherit;">
        <option value="">✨ Auto (recommended)</option>
        <option value="photorealistic, hyperrealistic DSLR photography, 8K resolution, cinematic">📸 Photorealistic</option>
        <option value="professional book cover design, award-winning layout, elegant typography">📚 Book Cover</option>
        <option value="Studio Ghibli anime style, soft watercolor, vibrant colors, beautiful">🌿 Anime / Ghibli</option>
        <option value="digital painting, fantasy concept art, epic lighting, deviantart">🎭 Fantasy Art</option>
        <option value="watercolor painting, soft pastel, dreamy, artistic brushstrokes">🖌 Watercolor</option>
        <option value="3D render, Octane render, ultra realistic, physically based rendering">🧊 3D Render</option>
        <option value="flat vector illustration, minimalist, clean lines, modern design">📐 Minimalist / Vector</option>
        <option value="oil painting, impressionist, rich textures, museum quality">🖼 Oil Painting</option>
        <option value="cinematic film still, dramatic lighting, movie poster quality, 35mm">🎬 Cinematic</option>
        <option value="pixel art, retro 8-bit style, vibrant palette, game art">🕹 Pixel Art</option>
        <option value="pencil sketch, detailed graphite drawing, fine art, black and white">✏️ Pencil Sketch</option>
        <option value="logo design, professional brand identity, clean, scalable vector">🏷 Logo / Brand</option>
      </select>
    </div>

    <div id="img-loading" style="display:none;text-align:center;padding:20px;">
      <div style="font-size:32px;margin-bottom:8px;">🎨</div>
      <div style="color:var(--muted);font-size:13px;">Generating your image...<br>
        <span style="font-size:11px;opacity:.7;">Auto-enhancing your prompt for best quality</span>
      </div>
    </div>
    <div id="img-error" style="display:none;color:#ef4444;font-size:12px;margin-bottom:8px;padding:8px;background:#fef2f2;border-radius:6px;"></div>

    <div id="img-result" style="display:none;margin-bottom:12px;text-align:center;">
      <img id="img-output" style="max-width:100%;border-radius:12px;box-shadow:0 4px 20px rgba(0,0,0,.4);cursor:zoom-in;">
      <div style="display:flex;gap:8px;margin-top:8px;">
        <button id="img-download-btn" style="flex:1;background:var(--accent);color:#fff;border:none;border-radius:8px;padding:8px;font-size:13px;cursor:pointer;font-family:inherit;">⬇ Download</button>
        <button id="img-copy-btn" style="flex:1;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:8px;font-size:13px;cursor:pointer;font-family:inherit;">📋 Copy</button>
        <button id="img-fullscreen-btn" style="flex:1;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:8px;font-size:13px;cursor:pointer;font-family:inherit;">⛶ View</button>
      </div>
    </div>

    <div style="display:flex;gap:8px;">
      <button id="img-generate-btn" style="flex:1;background:linear-gradient(135deg,#10a37f,#0d7a5f);color:#fff;border:none;border-radius:10px;padding:12px;font-size:14px;font-weight:700;cursor:pointer;font-family:inherit;">✨ Generate</button>
      <button id="img-close-btn" style="background:none;border:1px solid var(--border);color:var(--muted);border-radius:10px;padding:12px 16px;font-size:14px;cursor:pointer;">✕</button>
    </div>
  </div>
</div>

<div id="img-viewer-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.92);z-index:400;align-items:center;justify-content:center;cursor:zoom-out;">
  <img id="img-viewer-img" style="max-width:94%;max-height:94%;border-radius:8px;">
</div>

<!-- ─── FILE / PDF GENERATION MODAL ─────────────────────────────────────────── -->
<div id="file-modal-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.8);z-index:300;align-items:center;justify-content:center;">
  <div style="background:var(--panel);border:1px solid var(--border);border-radius:16px;padding:24px;width:92%;max-width:460px;max-height:90vh;overflow-y:auto;">
    <h3 style="margin:0 0 4px;font-size:18px;">📄 Generate a File</h3>
    <p style="color:var(--muted);font-size:13px;margin:0 0 16px;">Turn text into a downloadable PDF, Word doc, or plain text file</p>

    <input id="file-title-input" type="text" placeholder="Title (optional)"
      style="width:100%;background:var(--bg);border:1.5px solid var(--border);color:var(--text);border-radius:8px;padding:9px 12px;font-size:13px;outline:none;margin-bottom:10px;font-family:inherit;">

    <textarea id="file-content-input" rows="8" placeholder="Paste or type the content you want in the file..."
      style="width:100%;background:var(--bg);border:1.5px solid var(--border);color:var(--text);border-radius:8px;padding:10px 12px;font-size:13px;outline:none;margin-bottom:12px;font-family:inherit;resize:vertical;"></textarea>

    <div style="margin-bottom:14px;">
      <label style="font-size:12px;color:var(--muted);display:block;margin-bottom:6px;">Format:</label>
      <div style="display:flex;gap:6px;">
        <button class="file-format-btn" data-format="pdf" style="flex:1;padding:9px;border-radius:8px;border:1.5px solid var(--accent);background:var(--accent-dim);color:var(--accent);cursor:pointer;font-size:12.5px;font-family:inherit;">📕 PDF</button>
        <button class="file-format-btn" data-format="docx" style="flex:1;padding:9px;border-radius:8px;border:1px solid var(--border);background:var(--panel);color:var(--muted);cursor:pointer;font-size:12.5px;font-family:inherit;">📘 Word</button>
        <button class="file-format-btn" data-format="txt" style="flex:1;padding:9px;border-radius:8px;border:1px solid var(--border);background:var(--panel);color:var(--muted);cursor:pointer;font-size:12.5px;font-family:inherit;">📄 Text</button>
      </div>
    </div>

    <div id="file-loading" style="display:none;text-align:center;padding:16px;">
      <div style="font-size:28px;margin-bottom:6px;">📄</div>
      <div style="color:var(--muted);font-size:13px;">Building your file...</div>
    </div>
    <div id="file-error" style="display:none;color:#ef4444;font-size:12px;margin-bottom:8px;padding:8px;background:#fef2f2;border-radius:6px;"></div>
    <div id="file-note" style="display:none;color:var(--muted);font-size:11.5px;margin-bottom:8px;padding:8px;background:var(--bg);border-radius:6px;"></div>

    <div style="display:flex;gap:8px;">
      <button id="file-generate-btn" style="flex:1;background:linear-gradient(135deg,#10a37f,#0d7a5f);color:#fff;border:none;border-radius:10px;padding:12px;font-size:14px;font-weight:700;cursor:pointer;font-family:inherit;">⬇ Generate & Download</button>
      <button id="file-close-btn" style="background:none;border:1px solid var(--border);color:var(--muted);border-radius:10px;padding:12px 16px;font-size:14px;cursor:pointer;">✕</button>
    </div>
  </div>
</div>

<div id="weather-modal-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.8);z-index:300;align-items:center;justify-content:center;">
  <div style="background:var(--panel);border:1px solid var(--border);border-radius:16px;padding:24px;width:92%;max-width:460px;max-height:90vh;overflow-y:auto;">
    <h3 style="margin:0 0 4px;font-size:18px;">🌤 Weather</h3>
    <p style="color:var(--muted);font-size:13px;margin:0 0 16px;">Search any city, or use your current location</p>

    <div style="display:flex;gap:8px;margin-bottom:12px;">
      <input id="weather-city" type="text" placeholder="Search city or place..." autocomplete="off"
        style="flex:1;background:var(--bg);border:1.5px solid var(--border);color:var(--text);border-radius:8px;padding:10px 12px;font-size:13px;outline:none;font-family:inherit;">
      <button id="weather-search-btn" style="background:var(--accent);color:#fff;border:none;border-radius:8px;padding:0 14px;font-size:14px;cursor:pointer;">🔍</button>
      <button id="weather-location-btn" title="Use my location" style="background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:0 12px;font-size:14px;cursor:pointer;">📍</button>
    </div>

    <div id="weather-loading" style="display:none;text-align:center;padding:20px;">
      <div style="font-size:32px;margin-bottom:8px;">🌍</div>
      <div style="color:var(--muted);font-size:13px;">Fetching weather...</div>
    </div>
    <div id="weather-error" style="display:none;color:#ef4444;font-size:12px;margin-bottom:8px;padding:8px;background:#fef2f2;border-radius:6px;"></div>

    <div id="weather-result" style="display:none;">
      <div id="weather-content"></div>
    </div>

    <div style="display:flex;justify-content:flex-end;margin-top:14px;">
      <button id="weather-close-btn" style="background:none;border:1px solid var(--border);color:var(--muted);border-radius:10px;padding:10px 16px;font-size:14px;cursor:pointer;">Close</button>
    </div>
  </div>
</div>

<!-- ─── HOMEWORK & STUDY BOOK — question, upload, or URL ───────────────────── -->
<div id="homework-modal-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.8);z-index:300;align-items:center;justify-content:center;padding:16px;">
  <div style="background:var(--panel);border:1px solid var(--border);border-radius:16px;padding:24px;width:92%;max-width:460px;max-height:90vh;overflow-y:auto;">
    <h3 style="margin:0 0 4px;font-size:18px;">📚 Homework &amp; Study</h3>
    <p style="color:var(--muted);font-size:13px;margin:0 0 16px;">Ask a question, upload a book/PDF, or paste a link to one — or all three.</p>

    <textarea id="homework-question" rows="3" placeholder="e.g. Help me with question 4 on quadratic equations, or leave blank if you're just uploading a book"
      style="width:100%;background:var(--bg);border:1.5px solid var(--border);color:var(--text);border-radius:8px;padding:10px 12px;font-size:13px;outline:none;margin-bottom:12px;font-family:inherit;resize:vertical;"></textarea>

    <div style="display:flex;gap:6px;margin-bottom:10px;">
      <button id="hw-mode-upload" class="hw-mode-btn" data-mode="upload" style="flex:1;padding:9px;border-radius:8px;border:1.5px solid var(--accent);background:var(--accent-dim);color:var(--accent);cursor:pointer;font-size:12.5px;font-family:inherit;">📎 Upload File</button>
      <button id="hw-mode-url" class="hw-mode-btn" data-mode="url" style="flex:1;padding:9px;border-radius:8px;border:1px solid var(--border);background:var(--panel);color:var(--muted);cursor:pointer;font-size:12.5px;font-family:inherit;">🔗 Paste URL</button>
    </div>

    <div id="hw-upload-area">
      <div id="hw-upload-dropzone" style="border:2px dashed var(--border);border-radius:12px;padding:18px;text-align:center;cursor:pointer;margin-bottom:8px;">
        <div style="font-size:13px;color:var(--muted);">📄 Click to choose a PDF, DOCX, or text file</div>
      </div>
      <input type="file" id="hw-file-input" accept=".pdf,.docx,.txt,.md,.csv,.json" style="display:none">
      <div id="hw-file-name" style="font-size:12px;color:var(--accent);display:none;margin-bottom:8px;"></div>
    </div>

    <div id="hw-url-area" style="display:none;">
      <input id="hw-url-input" type="text" placeholder="https://example.com/book.pdf"
        style="width:100%;background:var(--bg);border:1.5px solid var(--border);color:var(--text);border-radius:8px;padding:10px 12px;font-size:13px;outline:none;margin-bottom:8px;font-family:inherit;">
    </div>

    <div id="hw-loading" style="display:none;text-align:center;padding:14px;color:var(--muted);font-size:13px;">Fetching document...</div>
    <div id="hw-error" style="display:none;color:#ef4444;font-size:12px;margin-bottom:8px;padding:8px;background:#fef2f2;border-radius:6px;"></div>

    <div style="display:flex;gap:8px;margin-top:6px;">
      <button id="hw-send-btn" style="flex:1;background:linear-gradient(135deg,#10a37f,#0d7a5f);color:#fff;border:none;border-radius:10px;padding:12px;font-size:14px;font-weight:700;cursor:pointer;font-family:inherit;">Send</button>
      <button id="hw-close-btn" style="background:none;border:1px solid var(--border);color:var(--muted);border-radius:10px;padding:12px 16px;font-size:14px;cursor:pointer;">✕</button>
    </div>
  </div>
</div>

<!-- ─── CODE WORKSPACE — HTML/CSS/JS editor with live preview ──────────────── -->
<div id="code-modal-overlay" style="display:none;position:fixed;inset:0;background:rgba(0,0,0,.85);z-index:300;align-items:center;justify-content:center;padding:16px;">
  <div style="background:var(--panel);border:1px solid var(--border);border-radius:16px;width:100%;max-width:1100px;height:88vh;display:flex;flex-direction:column;overflow:hidden;">
    <div style="display:flex;align-items:center;justify-content:space-between;padding:14px 18px;border-bottom:1px solid var(--border);flex-shrink:0;">
      <div style="display:flex;align-items:center;gap:10px;">
        <span style="font-size:16px;font-weight:700;">💻 Code Workspace</span>
        <input id="code-project-name" value="my-project" style="background:var(--bg);border:1px solid var(--border);color:var(--muted);font-size:12px;padding:4px 8px;border-radius:6px;font-family:inherit;width:140px;">
      </div>
      <div style="display:flex;gap:8px;">
        <button id="code-run-btn" style="background:var(--accent);color:#fff;border:none;border-radius:8px;padding:8px 16px;font-size:13px;font-weight:600;cursor:pointer;font-family:inherit;">▶ Run</button>
        <button id="code-download-btn" style="background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:8px 14px;font-size:13px;cursor:pointer;font-family:inherit;">⬇ Download</button>
        <button id="code-fullscreen-preview-btn" style="background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:8px 14px;font-size:13px;cursor:pointer;font-family:inherit;">⛶ Preview</button>
        <button id="code-close-btn" style="background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:8px 14px;font-size:13px;cursor:pointer;font-family:inherit;">✕</button>
      </div>
    </div>

    <!-- Language selector — Web (HTML/CSS/JS, live preview) or a single-file
         language run remotely via Judge0 (see /api/code/run below) -->
    <div id="code-lang-row" style="display:flex;gap:6px;padding:8px 18px;border-bottom:1px solid var(--border);flex-shrink:0;overflow-x:auto;">
      <button class="code-lang-btn active" data-lang="web" style="padding:6px 12px;border-radius:20px;border:1px solid var(--accent);background:var(--accent-dim);color:var(--accent);font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;">🌐 Web</button>
      <button class="code-lang-btn" data-lang="python" style="padding:6px 12px;border-radius:20px;border:1px solid var(--border);background:none;color:var(--muted);font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;">🐍 Python</button>
      <button class="code-lang-btn" data-lang="cpp" style="padding:6px 12px;border-radius:20px;border:1px solid var(--border);background:none;color:var(--muted);font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;">C++</button>
      <button class="code-lang-btn" data-lang="c" style="padding:6px 12px;border-radius:20px;border:1px solid var(--border);background:none;color:var(--muted);font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;">C</button>
      <button class="code-lang-btn" data-lang="java" style="padding:6px 12px;border-radius:20px;border:1px solid var(--border);background:none;color:var(--muted);font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;">☕ Java</button>
      <button class="code-lang-btn" data-lang="javascript" style="padding:6px 12px;border-radius:20px;border:1px solid var(--border);background:none;color:var(--muted);font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;">🟩 Node.js</button>
      <button class="code-lang-btn" data-lang="typescript" style="padding:6px 12px;border-radius:20px;border:1px solid var(--border);background:none;color:var(--muted);font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;">🔷 TypeScript</button>
      <button class="code-lang-btn" data-lang="go" style="padding:6px 12px;border-radius:20px;border:1px solid var(--border);background:none;color:var(--muted);font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;">🐹 Go</button>
      <button class="code-lang-btn" data-lang="ruby" style="padding:6px 12px;border-radius:20px;border:1px solid var(--border);background:none;color:var(--muted);font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;white-space:nowrap;">💎 Ruby</button>
    </div>

    <div style="display:flex;flex:1;min-height:0;">
      <!-- Editor pane -->
      <div style="flex:1;display:flex;flex-direction:column;min-width:0;border-right:1px solid var(--border);">
        <div id="code-web-tabs" style="display:flex;border-bottom:1px solid var(--border);flex-shrink:0;">
          <button class="code-file-tab active" data-target="code-editor-html" style="flex:1;padding:9px;background:var(--accent-dim);color:var(--accent);border:none;font-size:12.5px;font-weight:600;cursor:pointer;font-family:inherit;">HTML</button>
          <button class="code-file-tab" data-target="code-editor-css" style="flex:1;padding:9px;background:none;color:var(--muted);border:none;font-size:12.5px;font-weight:600;cursor:pointer;font-family:inherit;">CSS</button>
          <button class="code-file-tab" data-target="code-editor-js" style="flex:1;padding:9px;background:none;color:var(--muted);border:none;font-size:12.5px;font-weight:600;cursor:pointer;font-family:inherit;">JS</button>
        </div>
        <div id="code-single-tab" style="display:none;padding:9px 14px;border-bottom:1px solid var(--border);flex-shrink:0;font-size:12.5px;font-weight:600;color:var(--accent);"></div>
        <textarea id="code-editor-html" spellcheck="false" style="flex:1;background:#0d1117;color:#c9d1d9;border:none;outline:none;padding:14px;font-family:'SF Mono',Consolas,Monaco,monospace;font-size:13px;line-height:1.6;resize:none;tab-size:2;white-space:pre;overflow:auto;">&lt;!-- Write your HTML here --&gt;
&lt;h1&gt;Hello from Mythic AI Code Workspace&lt;/h1&gt;
&lt;p&gt;Edit HTML, CSS, and JS, then hit Run.&lt;/p&gt;
&lt;button onclick="sayHi()"&gt;Click me&lt;/button&gt;</textarea>
        <textarea id="code-editor-css" spellcheck="false" style="flex:1;display:none;background:#0d1117;color:#c9d1d9;border:none;outline:none;padding:14px;font-family:'SF Mono',Consolas,Monaco,monospace;font-size:13px;line-height:1.6;resize:none;tab-size:2;white-space:pre;overflow:auto;">body {
  font-family: sans-serif;
  background: #1a1a1a;
  color: #ececec;
  padding: 24px;
}
button {
  background: #10a37f;
  color: #fff;
  border: none;
  padding: 10px 18px;
  border-radius: 8px;
  cursor: pointer;
  font-size: 14px;
}</textarea>
        <textarea id="code-editor-js" spellcheck="false" style="flex:1;display:none;background:#0d1117;color:#c9d1d9;border:none;outline:none;padding:14px;font-family:'SF Mono',Consolas,Monaco,monospace;font-size:13px;line-height:1.6;resize:none;tab-size:2;white-space:pre;overflow:auto;">function sayHi() {
  alert("Hello from your Code Workspace!");
}</textarea>
        <textarea id="code-editor-single" spellcheck="false" style="flex:1;display:none;background:#0d1117;color:#c9d1d9;border:none;outline:none;padding:14px;font-family:'SF Mono',Consolas,Monaco,monospace;font-size:13px;line-height:1.6;resize:none;tab-size:2;white-space:pre;overflow:auto;"></textarea>
        <textarea id="code-stdin" spellcheck="false" placeholder="stdin (optional) — piped into your program when it runs" style="display:none;height:64px;flex-shrink:0;background:#0a0d12;color:#8b949e;border:none;border-top:1px solid var(--border);outline:none;padding:8px 14px;font-family:'SF Mono',Consolas,Monaco,monospace;font-size:12px;line-height:1.5;resize:none;white-space:pre;overflow:auto;"></textarea>
      </div>

      <!-- Preview / Output pane -->
      <div style="flex:1;display:flex;flex-direction:column;min-width:0;background:#fff;">
        <div id="code-preview-label" style="padding:6px 12px;background:var(--bg);border-bottom:1px solid var(--border);font-size:11px;color:var(--muted);flex-shrink:0;">Live Preview</div>
        <iframe id="code-preview-frame" sandbox="allow-scripts allow-modals" style="flex:1;border:none;width:100%;background:#fff;"></iframe>
        <pre id="code-output-console" style="flex:1;display:none;margin:0;padding:14px;background:#0d1117;color:#c9d1d9;font-family:'SF Mono',Consolas,Monaco,monospace;font-size:12.5px;line-height:1.6;overflow:auto;white-space:pre-wrap;word-break:break-word;"></pre>
      </div>
    </div>
  </div>
</div>

<script>
// ─── Resilient identity: keep a copy of our anonymous id outside the cookie ──
// If the session cookie ever fails to persist in a given browser (blocked,
// stripped by a proxy, cleared, etc.), this localStorage id lets the server
// reseed the same account instead of silently starting a fresh empty one.
(function () {
  try {
    let cid = localStorage.getItem('mythic_client_id');
    if (!cid) {
      cid = (crypto.randomUUID ? crypto.randomUUID() :
        'xxxxxxxx-xxxx-4xxx-yxxx-xxxxxxxxxxxx'.replace(/[xy]/g, c => {
          const r = Math.random() * 16 | 0;
          return (c === 'x' ? r : (r & 0x3 | 0x8)).toString(16);
        }));
      localStorage.setItem('mythic_client_id', cid);
    }
    const _origFetch = window.fetch.bind(window);
    window.fetch = function (input, init) {
      const url = typeof input === 'string' ? input : (input && input.url) || '';
      if (url.startsWith('/api/')) {
        init = init || {};
        init.headers = new Headers(init.headers || {});
        init.headers.set('X-Client-Id', cid);
      }
      return _origFetch(input, init).then(resp => {
        // Session isn't authenticated (expired, logged out elsewhere, etc.)
        // — bounce to the login screen instead of leaving the app stuck on
        // a silent 401. Skip this for the auth endpoints themselves so the
        // login page's own fetch calls can show its own inline error UI.
        if (resp.status === 401 && url.startsWith('/api/') && !url.startsWith('/api/auth/')) {
          window.location.href = '/login';
        }
        return resp;
      });
    };
  } catch (e) {
    // localStorage unavailable (rare, e.g. some locked-down browser modes) —
    // app still works, just without the cookie-loss fallback.
    console.warn('Client-id resilience layer unavailable:', e);
  }
})();

function _setAppHeight() {
  const h = (window.visualViewport && window.visualViewport.height) || window.innerHeight;
  document.documentElement.style.setProperty('--app-height', h + 'px');
}
_setAppHeight();
window.addEventListener('resize', _setAppHeight);
window.addEventListener('orientationchange', () => setTimeout(_setAppHeight, 100));
if (window.visualViewport) {
  window.visualViewport.addEventListener('resize', _setAppHeight);
}

// ─── Declare all key global variables at the top to avoid TDZ errors ─────
const isIOS = /iPad|iPhone|iPod/.test(navigator.userAgent) && !window.MSStream;
const modeTabs = document.querySelectorAll('.mode-tab[data-mode]');

const messagesWrap = document.getElementById('messages-wrap');
const messagesEl   = document.getElementById('messages');
const form         = document.getElementById('chat-form');
const input        = document.getElementById('input');
const sendBtn      = document.getElementById('send-btn');
const clearBtn     = document.getElementById('clear-btn');
const convListEl   = document.getElementById('conv-list');
const newChatBtn   = document.getElementById('new-chat-btn');
const sidebarToggle= document.getElementById('sidebar-toggle');
const fullscreenBtn= document.getElementById('fullscreen-btn');
const nameBtn       = document.getElementById('name-btn');
const vipBtn        = document.getElementById('vip-btn');
const streakBadge   = document.getElementById('streak-badge');

let selectedModel = 'mythic-2';
let currentMode   = 'chat';
let _artifacts     = []; // {id, lang, code, ts, sourceMsgPreview}
let vipUnlocked   = false;

function getUserApiKeys() {
  return {
    groq_api_key: localStorage.getItem('mythic_user_groq_key') || '',
    cerebras_api_key: localStorage.getItem('mythic_user_cerebras_key') || '',
  };
}

function updateVipBtn() {
  vipBtn.textContent = vipUnlocked && selectedModel === 'mythic-vip' ? '✨' : (vipUnlocked ? '✨' : '🔒');
  vipBtn.classList.toggle('active', selectedModel === 'mythic-vip');
  vipBtn.title = vipUnlocked
    ? (selectedModel === 'mythic-vip' ? 'Mythic VIP active — click to switch back' : 'Switch to Mythic VIP')
    : 'Unlock Mythic VIP';
}

async function refreshStreakBadge() {
  try {
    const r = await fetch('/api/streak');
    if (!r.ok) return;
    const d = await r.json();
    const streak = d.streak || 0;
    if (streak > 0) {
      streakBadge.textContent = '🔥 ' + streak;
      streakBadge.style.display = 'inline-flex';
      streakBadge.title = streak === 1
        ? '1 day streak — chat again tomorrow to keep it going!'
        : streak + ' day streak — chat again tomorrow to keep it going!';
    } else {
      streakBadge.style.display = 'none';
    }
  } catch {}
}
refreshStreakBadge();

function showVipModal() {
  const existing = document.getElementById('vip-modal-overlay');
  if (existing) { existing.style.display = 'flex'; return; }
  const overlay = document.createElement('div');
  overlay.id = 'vip-modal-overlay';
  overlay.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.7);z-index:500;display:flex;align-items:center;justify-content:center;';
  overlay.innerHTML = `<div style="background:var(--panel);border:1px solid var(--border);border-radius:16px;padding:24px;width:90%;max-width:340px;">
    <div style="font-size:22px;margin-bottom:6px;">🔒 VIP Access</div>
    <div style="color:var(--muted);font-size:13px;margin-bottom:16px;">Mythic VIP is for VIP users only.</div>
    <input id="vip-pw-in" type="password" placeholder="VIP password" style="width:100%;background:var(--bg);border:1.5px solid var(--border);color:var(--text);border-radius:8px;padding:10px 12px;font-size:14px;outline:none;margin-bottom:8px;font-family:inherit;">
    <div id="vip-pw-err" style="color:#ef4444;font-size:12px;display:none;margin-bottom:8px;">Wrong password.</div>
    <div style="display:flex;gap:8px;">
      <button id="vip-pw-ok" style="flex:1;background:var(--accent);color:#fff;border:none;border-radius:8px;padding:10px;font-size:14px;font-weight:600;cursor:pointer;">Unlock</button>
      <button id="vip-pw-cancel" style="flex:1;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:10px;font-size:14px;cursor:pointer;">Cancel</button>
    </div></div>`;
  document.body.appendChild(overlay);
  const pwIn = overlay.querySelector('#vip-pw-in'), pwErr = overlay.querySelector('#vip-pw-err');
  pwIn.focus();
  overlay.querySelector('#vip-pw-cancel').addEventListener('click', () => {
    overlay.style.display = 'none';
  });
  overlay.querySelector('#vip-pw-ok').addEventListener('click', async () => {
    try {
      const r = await fetch('/api/vip-unlock', {
        method: 'POST', headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ password: pwIn.value.trim() })
      });
      const d = await r.json();
      if (d.success) {
        vipUnlocked = true;
        overlay.style.display = 'none';
        selectedModel = 'mythic-vip';
        updateVipBtn();
        if (typeof syncModeTabLocks === 'function') syncModeTabLocks();
      } else {
        pwErr.style.display = 'block'; pwIn.value = ''; pwIn.focus();
      }
    } catch {
      pwErr.textContent = 'Network error — try again.';
      pwErr.style.display = 'block';
    }
  });
  pwIn.addEventListener('keydown', e => { if (e.key === 'Enter') overlay.querySelector('#vip-pw-ok').click(); });
}

(async () => {
  try {
    // Deliberately NOT restoring vipUnlocked from /api/vip-status here.
    // The unlock flag rides in the same long-lived session cookie used for
    // conversation storage, so trusting it on load meant a page refresh
    // skipped the VIP password entirely. VIP now always starts locked for
    // a fresh page load — the password is required again each time you
    // open/reload the app, but not repeatedly while switching tabs within
    // the same loaded page.
    const mr = await fetch('/api/models').then(r => r.json());
    if (mr && mr.default) selectedModel = mr.default;
  } catch {}
  updateVipBtn();
  if (typeof syncModeTabLocks === 'function') syncModeTabLocks();
})();

vipBtn.addEventListener('click', () => {
  if (!vipUnlocked) { showVipModal(); return; }
  selectedModel = selectedModel === 'mythic-vip' ? 'mythic-2' : 'mythic-vip';
  updateVipBtn();
  if (typeof syncModeTabLocks === 'function') syncModeTabLocks();
  if (typeof setActiveModeTab === 'function' && selectedModel !== 'mythic-vip' && currentMode !== 'chat') setActiveModeTab('chat');
});

const nameModalOverlay = document.getElementById('name-modal-overlay');
const nameInput     = document.getElementById('name-input');
const nameCancelBtn = document.getElementById('name-cancel-btn');
const nameSaveBtn   = document.getElementById('name-save-btn');
const exportBtn     = document.getElementById('export-btn');
const sidebar      = document.getElementById('sidebar');
const fileInput    = document.getElementById('file-input');
const cameraInput  = document.getElementById('camera-input');
const plusBtn      = document.getElementById('plus-btn');
const voiceBtn     = document.getElementById('voice-btn');
const pendingAttach= document.getElementById('pending-attach');
const pendingName  = document.getElementById('pending-attach-name');
const pendingRemove= document.getElementById('pending-attach-remove');
const scrollBtn    = document.getElementById('scroll-btn');
const speakingIndicator = document.getElementById('speaking-indicator');
const stopSpeakBtn = document.getElementById('stop-speak-btn');

let activeConvId = null;
let pendingFile  = null;
let recognition  = null;
let currentUtterance = null;
let _pendingUndoSend = false; // true while a just-sent message is still in its undo window

messagesWrap.addEventListener('scroll', () => {
  const nearBottom = messagesWrap.scrollHeight - messagesWrap.scrollTop - messagesWrap.clientHeight < 120;
  scrollBtn.classList.toggle('show', !nearBottom);
});
scrollBtn.addEventListener('click', () => {
  messagesWrap.scrollTo({ top: messagesWrap.scrollHeight, behavior: 'smooth' });
});
function scrollToBottom() {
  requestAnimationFrame(() => {
    messagesWrap.scrollTo({ top: messagesWrap.scrollHeight, behavior: 'smooth' });
  });
}

function clearEmptyState() {
  const es = document.getElementById('empty-state');
  if (es) es.remove();
}
function showEmptyState() {
  messagesEl.innerHTML = '<div class="empty-state" id="empty-state"><h2>Mythic AI</h2><p>Ask me anything, generate images, or just chat 👋</p></div>';
}

let addMessage = function(role, text, attachment) {
  clearEmptyState();
  const row = document.createElement('div');
  row.className = 'msg-row ' + role;
  // Note: row.dataset.msgIndex is set later, in buildMsgActions() below
  // (see the bookmark-feature monkey-patch further down the file, which
  // owns _msgIndexCounter) — kept as the single source of truth for a
  // message's position in the server's conv["messages"] array so the
  // bookmark feature and the edit/branch feature below both read the same
  // index. It's reset to 0 in openConversation()/startNewChat() and
  // adjusted by regenerateLast()/startEditingMessage() to stay in sync
  // whenever a message is replaced in place rather than appended.

  const div = document.createElement('div');
  div.className = 'msg ' + role;
  if (attachment) {
    const chip = document.createElement('div');
    chip.className = 'attach-chip';
    chip.textContent = '📎 ' + attachment.name;
    div.appendChild(chip);
    if (attachment.mimeType && attachment.mimeType.startsWith('image/') && attachment.dataBase64) {
      const img = document.createElement('img');
      img.src = 'data:' + attachment.mimeType + ';base64,' + attachment.dataBase64;
      div.appendChild(img);
    }
  }
  const textNode = document.createElement('div');
  textNode.className = 'msg-text';
  textNode.textContent = text;
  div.appendChild(textNode);
  row.appendChild(div);

  if (role === 'user' || role === 'ai') {
    row.appendChild(buildMsgActions(row, textNode, role));
  }

  messagesEl.appendChild(row);
  scrollToBottom();
  return textNode;
};

let buildMsgActions = function(row, textNode, role) {
  const actions = document.createElement('div');
  actions.className = 'msg-actions';

  const copyBtn = document.createElement('button');
  copyBtn.type = 'button';
  copyBtn.className = 'copy-btn';
  copyBtn.title = 'Copy';
  copyBtn.textContent = '📋';
  copyBtn.addEventListener('click', async () => {
    try {
      await navigator.clipboard.writeText(textNode.textContent);
    } catch {
      const ta = document.createElement('textarea');
      ta.value = textNode.textContent;
      document.body.appendChild(ta);
      ta.select();
      try { document.execCommand('copy'); } catch {}
      ta.remove();
    }
    const orig = copyBtn.textContent;
    copyBtn.textContent = '✓';
    setTimeout(() => { copyBtn.textContent = orig; }, 1200);
  });
  actions.appendChild(copyBtn);

  if (role === 'ai') {
    const regenBtn = document.createElement('button');
    regenBtn.type = 'button';
    regenBtn.className = 'regen-btn';
    regenBtn.title = 'Regenerate response';
    regenBtn.textContent = '↻';
    regenBtn.addEventListener('click', () => regenerateLast(row));
    actions.appendChild(regenBtn);

    const fileBtn = document.createElement('button');
    fileBtn.type = 'button';
    fileBtn.title = 'Save as file (PDF/Word/Text)';
    fileBtn.textContent = '📄';
    fileBtn.addEventListener('click', () => openFileModalWithContent(textNode.textContent || textNode.innerText || ''));
    actions.appendChild(fileBtn);
  }

  if (role === 'user') {
    const editBtn = document.createElement('button');
    editBtn.type = 'button';
    editBtn.className = 'edit-btn';
    editBtn.title = 'Edit message';
    editBtn.textContent = '✏️';
    editBtn.addEventListener('click', () => startEditingMessage(row, textNode));
    actions.appendChild(editBtn);
  }
  return actions;
};

// Turns a sent user message into an editable textarea in place. Saving
// calls /api/conversations/<id>/edit-message to truncate + rewrite server
// history, then re-adds the edited message and streams a fresh reply via
// the EXISTING regenerate path — no chat/streaming logic is duplicated or
// modified here, only conversation history.
function startEditingMessage(row, textNode) {
  if (isGenerating) return;
  if (!activeConvId) return; // nothing saved server-side yet to branch from
  if (row.querySelector('.edit-box')) return; // already editing

  const msgIndex = row.dataset.msgIndex;
  const originalText = textNode.textContent;
  const bubble = row.querySelector('.msg');

  const box = document.createElement('div');
  box.className = 'edit-box';
  const ta = document.createElement('textarea');
  ta.className = 'edit-textarea';
  ta.value = originalText;
  const controls = document.createElement('div');
  controls.className = 'edit-controls';
  const saveBtn = document.createElement('button');
  saveBtn.type = 'button'; saveBtn.className = 'edit-save-btn'; saveBtn.textContent = 'Save & submit';
  const cancelBtn = document.createElement('button');
  cancelBtn.type = 'button'; cancelBtn.className = 'edit-cancel-btn'; cancelBtn.textContent = 'Cancel';
  controls.appendChild(saveBtn);
  controls.appendChild(cancelBtn);
  box.appendChild(ta);
  box.appendChild(controls);

  textNode.style.display = 'none';
  bubble.appendChild(box);
  ta.focus();
  ta.setSelectionRange(ta.value.length, ta.value.length);
  ta.style.height = 'auto';
  ta.style.height = Math.min(ta.scrollHeight, 200) + 'px';
  ta.addEventListener('input', () => {
    ta.style.height = 'auto';
    ta.style.height = Math.min(ta.scrollHeight, 200) + 'px';
  });

  function cancelEdit() {
    box.remove();
    textNode.style.display = '';
  }
  ta.addEventListener('keydown', e => {
    if (e.key === 'Escape') cancelEdit();
    if (e.key === 'Enter' && !e.shiftKey) { e.preventDefault(); saveBtn.click(); }
  });
  cancelBtn.addEventListener('click', cancelEdit);

  saveBtn.addEventListener('click', async () => {
    const newText = ta.value.trim();
    if (!newText || newText === originalText) { cancelEdit(); return; }
    saveBtn.disabled = true; cancelBtn.disabled = true;
    saveBtn.textContent = 'Saving…';
    try {
      const r = await fetch('/api/conversations/' + activeConvId + '/edit-message', {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ message_index: parseInt(msgIndex, 10), new_text: newText })
      });
      if (!r.ok) {
        const d = await r.json().catch(() => ({}));
        throw new Error(d.error || ('HTTP ' + r.status));
      }
      // Drop this row and every row after it (the old reply + anything
      // beyond), then roll the index counter back to this slot so the
      // re-added message and its new AI reply get the same indices the
      // server just assigned them.
      let sib = row.nextSibling;
      while (sib) { const next = sib.nextSibling; sib.remove(); sib = next; }
      row.remove();
      _msgIndexCounter = parseInt(msgIndex, 10);
      addMessage('user', newText);
      streamReply({ regenerate: true });
    } catch (err) {
      cancelEdit();
      addMessage('error', 'Could not save the edit: ' + (err.message || 'unknown error'));
    }
  });
}

function addImageMessage(role, base64, caption) {
  clearEmptyState();
  const div = document.createElement('div');
  div.className = 'msg ' + role;
  if (caption) {
    const cap = document.createElement('div');
    cap.textContent = caption;
    cap.style.marginBottom = '8px';
    div.appendChild(cap);
  }
  const img = document.createElement('img');
  img.className = 'gen-img';
  img.src = 'data:image/png;base64,' + base64;
  div.appendChild(img);
  messagesEl.appendChild(div);
  scrollToBottom();
}

function showTyping(label) {
  hideTyping();  // avoid stacking duplicates if called twice (e.g. during retry)
  const div = document.createElement('div');
  div.className = 'typing'; div.id = 'typing-indicator';
  div.innerHTML = label
    ? `<span style="font-size:12.5px;color:var(--muted);white-space:nowrap;">${label.replace(/</g,'&lt;')}</span>`
    : '<span></span><span></span><span></span>';
  messagesEl.appendChild(div);
  scrollToBottom();
}
function hideTyping() {
  const el = document.getElementById('typing-indicator');
  if (el) el.remove();
}

function speak(text) {
  if (!window.speechSynthesis) return;
  window.speechSynthesis.cancel();
  const plain = text.replace(/[#*`_~>]/g, '').trim();
  if (!plain) return;
  currentUtterance = new SpeechSynthesisUtterance(plain);
  currentUtterance.rate = 1.05;
  const chosen = (typeof getChosenVoice === 'function') ? getChosenVoice() : null;
  // One fixed voice for every language — no per-language switching.
  if (chosen) currentUtterance.voice = chosen;
  currentUtterance.onstart = () => speakingIndicator.classList.add('show');
  currentUtterance.onend = () => speakingIndicator.classList.remove('show');
  currentUtterance.onerror = () => speakingIndicator.classList.remove('show');
  window.speechSynthesis.speak(currentUtterance);
}
stopSpeakBtn.addEventListener('click', () => {
  window.speechSynthesis && window.speechSynthesis.cancel();
  speakingIndicator.classList.remove('show');
});

function setupVoice() {
  const SR = window.SpeechRecognition || window.webkitSpeechRecognition;
  if (!voiceBtn) return;
  if (!SR) { voiceBtn.title = 'Voice not supported in this browser'; return; }
  recognition = new SR();
  recognition.continuous = false;
  recognition.interimResults = true;
  recognition.lang = 'en-US';
  let finalTranscript = '';
  recognition.onstart  = () => { voiceBtn.classList.add('active', 'listening'); finalTranscript = ''; };
  recognition.onresult = (e) => {
    finalTranscript = '';
    for (let i = e.resultIndex; i < e.results.length; i++) {
      if (e.results[i].isFinal) finalTranscript += e.results[i][0].transcript;
      else input.value = e.results[i][0].transcript;
    }
    if (finalTranscript) input.value = finalTranscript;
  };
  recognition.onend = () => {
    voiceBtn.classList.remove('active', 'listening');
    if (input.value.trim()) form.requestSubmit();
  };
  recognition.onerror = () => voiceBtn.classList.remove('active', 'listening');
}
setupVoice();
if (voiceBtn) voiceBtn.addEventListener('click', () => {
  if (!recognition) { alert('Voice input is not supported in this browser. Try Chrome.'); return; }
  if (voiceBtn.classList.contains('listening')) { recognition.stop(); return; }
  recognition.start();
});

// ─── LIVE TALK: continuous hands-free voice conversation ───────────────────
// Loop: listen for speech -> auto-submit as a message -> wait for the AI's
// reply to finish streaming -> speak it aloud -> automatically start
// listening again. Runs until the user taps the button (or the status
// pill's "End") to stop. Independent of the (currently unused) voiceBtn
// dictation path above — Live Talk manages its own recognition instance.
const liveTalkBtn        = document.getElementById('live-talk-btn');
const liveTalkStatus     = document.getElementById('live-talk-status');
const liveTalkStatusText = document.getElementById('live-talk-status-text');
const liveTalkStopBtn    = document.getElementById('live-talk-stop-btn');
let liveTalkActive = false;
let liveRecognition = null;
let liveTalkPhase = 'idle'; // 'listening' | 'sending' | 'speaking'

function liveTalkSupported() {
  return !!(window.SpeechRecognition || window.webkitSpeechRecognition) && !!window.speechSynthesis;
}

function setLiveTalkStatus(phase, text) {
  liveTalkPhase = phase;
  if (liveTalkStatus) {
    liveTalkStatus.classList.toggle('show', liveTalkActive);
    liveTalkStatus.classList.toggle('speaking', phase === 'speaking');
  }
  if (liveTalkStatusText) liveTalkStatusText.textContent = text;
  if (liveTalkBtn) {
    liveTalkBtn.classList.toggle('live-listening', liveTalkActive && phase === 'listening');
    liveTalkBtn.classList.toggle('live-speaking', liveTalkActive && phase === 'speaking');
    liveTalkBtn.classList.toggle('live-idle-active', liveTalkActive && phase !== 'listening' && phase !== 'speaking');
  }
}

function liveTalkListenOnce() {
  if (!liveTalkActive) return;
  const SR = window.SpeechRecognition || window.webkitSpeechRecognition;
  const rec = new SR();
  liveRecognition = rec;
  rec.continuous = false;
  rec.interimResults = false;
  rec.lang = 'en-US';
  setLiveTalkStatus('listening', 'Listening…');
  let handled = false;
  rec.onresult = (e) => {
    let t = '';
    for (let i = e.resultIndex; i < e.results.length; i++) {
      if (e.results[i].isFinal) t += e.results[i][0].transcript;
    }
    t = t.trim();
    if (t) {
      handled = true;
      input.value = t;
      autoResize();
      // Video Call / Screen Share: grab the current frame and attach it,
      // so this message carries what Mythic AI is currently "seeing".
      if (liveMediaStream) pendingFile = captureLiveFrame();
      setLiveTalkStatus('sending', liveMediaStream ? 'Looking…' : 'Thinking…');
      form.requestSubmit();
      liveTalkWaitThenSpeak();
    }
  };
  rec.onerror = (e) => {
    if (!liveTalkActive) return;
    if (e.error === 'no-speech' || e.error === 'aborted') { setTimeout(() => { if (liveTalkActive) liveTalkListenOnce(); }, 300); }
  };
  rec.onend = () => {
    if (!liveTalkActive || handled) return;
    // Silence/timeout with nothing said — just keep listening.
    setTimeout(() => { if (liveTalkActive) liveTalkListenOnce(); }, 300);
  };
  try { rec.start(); } catch { setTimeout(() => { if (liveTalkActive) liveTalkListenOnce(); }, 500); }
}

function liveTalkWaitThenSpeak() {
  const check = setInterval(() => {
    if (!liveTalkActive) { clearInterval(check); return; }
    if (!isGenerating) {
      clearInterval(check);
      const allRows = messagesEl.querySelectorAll('.msg-row.ai');
      const lastRow = allRows[allRows.length - 1];
      const textEl = lastRow ? lastRow.querySelector('.msg-text,.md-rendered') : null;
      const replyText = textEl ? (textEl.textContent || textEl.innerText || '') : '';
      if (replyText.trim() && liveTalkActive) {
        setLiveTalkStatus('speaking', 'Speaking…');
        speak(replyText);
        const waitSpeak = setInterval(() => {
          if (!liveTalkActive) { clearInterval(waitSpeak); return; }
          if (!window.speechSynthesis.speaking) {
            clearInterval(waitSpeak);
            if (liveTalkActive) liveTalkListenOnce();
          }
        }, 250);
      } else if (liveTalkActive) {
        liveTalkListenOnce();
      }
    }
  }, 400);
}

// ─── VIDEO CALL & SCREEN SHARE: Mythic AI can actually see a live frame ────
// Captures a snapshot from the active camera/screen stream each time Live
// Talk hears you finish speaking, attaches it to that message (same path as
// a regular image attachment), and the backend routes it to a vision model.
const videoCallBtn      = document.getElementById('video-call-btn');
const screenShareBtn    = document.getElementById('screen-share-btn');
const liveVideoWrap     = document.getElementById('live-video-wrap');
const liveVideoPreview  = document.getElementById('live-video-preview');
const liveVideoStopBtn  = document.getElementById('live-video-stop-btn');
let liveMediaStream = null;   // active MediaStream (camera or screen)
let liveMediaKind   = null;   // 'camera' | 'screen'
let liveFrameCanvas = null;

function captureLiveFrame() {
  if (!liveMediaStream || !liveVideoPreview || !liveVideoPreview.videoWidth) return null;
  if (!liveFrameCanvas) liveFrameCanvas = document.createElement('canvas');
  liveFrameCanvas.width = liveVideoPreview.videoWidth;
  liveFrameCanvas.height = liveVideoPreview.videoHeight;
  const ctx = liveFrameCanvas.getContext('2d');
  ctx.drawImage(liveVideoPreview, 0, 0, liveFrameCanvas.width, liveFrameCanvas.height);
  const dataUrl = liveFrameCanvas.toDataURL('image/jpeg', 0.7);
  const base64 = dataUrl.split(',')[1];
  if (!base64) return null;
  return { name: (liveMediaKind === 'screen' ? 'screen.jpg' : 'camera.jpg'), mimeType: 'image/jpeg', dataBase64: base64 };
}

function stopLiveMedia() {
  if (liveMediaStream) { liveMediaStream.getTracks().forEach(t => { try { t.stop(); } catch {} }); }
  liveMediaStream = null;
  liveMediaKind = null;
  if (liveVideoPreview) liveVideoPreview.srcObject = null;
  if (liveVideoWrap) liveVideoWrap.classList.remove('show');
  if (videoCallBtn) videoCallBtn.classList.remove('active');
  if (screenShareBtn) screenShareBtn.classList.remove('active');
}

async function startLiveMedia(kind) {
  // Camera (Video Call) only needs getUserMedia — it does NOT need speech
  // recognition support. Screen Share needs getDisplayMedia, which simply
  // does not exist on ANY mobile browser (Android Chrome or iOS Safari) —
  // that's a platform limitation, not something fixable here, so we detect
  // it up front and say so clearly instead of silently failing.
  if (kind === 'screen') {
    if (!navigator.mediaDevices || !navigator.mediaDevices.getDisplayMedia) {
      alert('Screen sharing isn\'t available in mobile browsers (Chrome/Safari on '
        + 'phones don\'t support it) — this only works on desktop Chrome/Edge.');
      return;
    }
  } else {
    if (!navigator.mediaDevices || !navigator.mediaDevices.getUserMedia) {
      alert('Camera access isn\'t available in this browser.');
      return;
    }
  }
  stopLiveMedia();
  try {
    liveMediaStream = kind === 'screen'
      ? await navigator.mediaDevices.getDisplayMedia({ video: true, audio: false })
      : await navigator.mediaDevices.getUserMedia({ video: { facingMode: 'user' }, audio: false });
  } catch (err) {
    if (err && (err.name === 'NotAllowedError' || err.name === 'PermissionDeniedError')) {
      alert((kind === 'screen' ? 'Screen share' : 'Camera') + ' permission was denied. '
        + 'Check your browser\'s site settings and allow camera access, then try again.');
    } else if (err && err.name === 'NotFoundError') {
      alert('No camera was found on this device.');
    } else {
      alert((kind === 'screen' ? 'Screen share' : 'Camera') + ' access was blocked or cancelled'
        + (err && err.message ? (': ' + err.message) : '.'));
    }
    return;
  }
  liveMediaKind = kind;
  liveVideoPreview.srcObject = liveMediaStream;
  liveVideoWrap.classList.add('show');
  (kind === 'screen' ? screenShareBtn : videoCallBtn).classList.add('active');
  // If the person stops sharing from the browser's own UI (not our button),
  // clean up on this end too.
  liveMediaStream.getVideoTracks()[0].addEventListener('ended', () => { if (liveMediaKind === kind) stopLiveMedia(); });
  // A call implies talking — auto-start Live Talk ONLY if this browser
  // actually supports speech recognition (e.g. not Safari). If it doesn't,
  // the camera/screen preview still works fine on its own — the person can
  // type messages instead of speaking them.
  if (!liveTalkActive) {
    if (liveTalkSupported()) {
      liveTalkActive = true;
      liveTalkListenOnce();
    } else {
      setLiveTalkStatus('idle', 'Voice control not supported here — type your message instead');
      if (liveTalkStatus) liveTalkStatus.classList.add('show');
      setTimeout(() => { if (!liveTalkActive && liveTalkStatus) liveTalkStatus.classList.remove('show'); }, 4000);
    }
  }
}

if (videoCallBtn) videoCallBtn.addEventListener('click', () => {
  if (liveMediaKind === 'camera') { stopLiveMedia(); return; }
  startLiveMedia('camera');
});
if (screenShareBtn) screenShareBtn.addEventListener('click', () => {
  if (liveMediaKind === 'screen') { stopLiveMedia(); return; }
  startLiveMedia('screen');
});
if (liveVideoStopBtn) liveVideoStopBtn.addEventListener('click', stopLiveMedia);

function stopLiveTalk() {
  liveTalkActive = false;
  if (liveRecognition) { try { liveRecognition.stop(); } catch {} liveRecognition = null; }
  if (window.speechSynthesis) window.speechSynthesis.cancel();
  if (liveTalkStatus) liveTalkStatus.classList.remove('show');
  if (liveTalkBtn) liveTalkBtn.classList.remove('live-listening', 'live-speaking', 'live-idle-active');
  stopLiveMedia();
}

if (liveTalkBtn) liveTalkBtn.addEventListener('click', () => {
  if (!liveTalkSupported()) { alert('Live Talk needs microphone + speech support — try Chrome or Edge.'); return; }
  if (liveTalkActive) { stopLiveTalk(); return; }
  liveTalkActive = true;
  liveTalkListenOnce();
});
if (liveTalkStopBtn) liveTalkStopBtn.addEventListener('click', stopLiveTalk);

function handleFileSelect(file) {
  if (!file) return;
  const reader = new FileReader();
  reader.onload = (e) => {
    const dataUrl = e.target.result;
    const base64  = dataUrl.split(',')[1];
    pendingFile = { name: file.name, mimeType: file.type || 'application/octet-stream', dataBase64: base64 };
    pendingName.textContent = file.name;
    pendingAttach.classList.add('show');
  };
  reader.readAsDataURL(file);
}
function closePlusMenu() {
  const existing = document.querySelector('.plus-menu-dropdown');
  if (existing) existing.remove();
}

// ─── In-page camera capture ─────────────────────────────────────────────
// Stays inside the app (a full-screen overlay) instead of handing off to
// the phone's native camera app the way a plain <input capture> does.
// Falls back to the native file-picker (cameraInput.click()) only if this
// browser has no getUserMedia support at all.
let _cameraModalStream = null;

async function openCameraModal() {
  const overlay = document.getElementById('camera-modal-overlay');
  const videoEl = document.getElementById('camera-modal-video');
  const errEl   = document.getElementById('camera-modal-error');
  if (!overlay || !videoEl) { cameraInput.click(); return; }

  if (!navigator.mediaDevices || !navigator.mediaDevices.getUserMedia) {
    // No in-page camera API available in this browser at all — fall back
    // to the OS camera picker rather than showing a dead-end screen.
    cameraInput.click();
    return;
  }

  overlay.style.display = 'flex';
  errEl.style.display = 'none';
  try {
    _cameraModalStream = await navigator.mediaDevices.getUserMedia({
      video: { facingMode: 'environment' }, audio: false,
    });
    videoEl.srcObject = _cameraModalStream;
  } catch (err) {
    errEl.textContent = (err && err.name === 'NotAllowedError')
      ? 'Camera permission denied — allow camera access in your browser settings and try again.'
      : 'Could not open the camera' + (err && err.message ? (': ' + err.message) : '.') + ' Using photo picker instead…';
    errEl.style.display = 'block';
    setTimeout(() => { closeCameraModal(); cameraInput.click(); }, 1600);
  }
}

function closeCameraModal() {
  const overlay = document.getElementById('camera-modal-overlay');
  if (overlay) overlay.style.display = 'none';
  if (_cameraModalStream) {
    _cameraModalStream.getTracks().forEach(t => { try { t.stop(); } catch {} });
    _cameraModalStream = null;
  }
}

function shootCameraModalPhoto() {
  const videoEl  = document.getElementById('camera-modal-video');
  const canvasEl = document.getElementById('camera-modal-canvas');
  if (!videoEl || !videoEl.videoWidth) return;
  canvasEl.width = videoEl.videoWidth;
  canvasEl.height = videoEl.videoHeight;
  const ctx = canvasEl.getContext('2d');
  ctx.drawImage(videoEl, 0, 0, canvasEl.width, canvasEl.height);
  const dataUrl = canvasEl.toDataURL('image/jpeg', 0.85);
  const base64 = dataUrl.split(',')[1];
  if (!base64) return;
  pendingFile = { name: 'photo-' + Date.now() + '.jpg', mimeType: 'image/jpeg', dataBase64: base64 };
  pendingName.textContent = pendingFile.name;
  pendingAttach.classList.add('show');
  closeCameraModal();
}

document.addEventListener('DOMContentLoaded', () => {
  const cancelBtn = document.getElementById('camera-modal-cancel');
  const shootBtn  = document.getElementById('camera-modal-shoot');
  if (cancelBtn) cancelBtn.addEventListener('click', closeCameraModal);
  if (shootBtn) shootBtn.addEventListener('click', shootCameraModalPhoto);
});
// In case this script runs after DOMContentLoaded already fired (it's at
// the bottom of the page), wire up immediately too — harmless if it also
// fires via the listener above since getElementById calls are idempotent.
(function _wireCameraModalNow() {
  const cancelBtn = document.getElementById('camera-modal-cancel');
  const shootBtn  = document.getElementById('camera-modal-shoot');
  if (cancelBtn && !cancelBtn._wired) { cancelBtn._wired = true; cancelBtn.addEventListener('click', closeCameraModal); }
  if (shootBtn && !shootBtn._wired) { shootBtn._wired = true; shootBtn.addEventListener('click', shootCameraModalPhoto); }
})();
plusBtn.addEventListener('click', (e) => {
  e.stopPropagation();
  if (document.querySelector('.plus-menu-dropdown')) { closePlusMenu(); return; }
  const menu = document.createElement('div');
  menu.className = 'plus-menu-dropdown';
  menu.style.cssText = 'position:fixed;background:var(--panel);border:1px solid var(--border);'
    + 'border-radius:10px;box-shadow:0 4px 20px rgba(0,0,0,.3);z-index:400;min-width:200px;'
    + 'max-height:70vh;overflow-y:auto;overflow-x:hidden;';
  const rect = plusBtn.getBoundingClientRect();
  // Anchor ABOVE the button since it lives in the bottom input bar.
  menu.style.bottom = (window.innerHeight - rect.top + 6) + 'px';
  menu.style.left = rect.left + 'px';

  const items = [
    { label: '📷 Take Photo', action: () => openCameraModal() },
    { label: '📎 Upload File', action: () => fileInput.click() },
    { label: '🎨 Image', action: () => document.getElementById('img-gen-btn')?.click() },
    { label: '🌿 Ghibli Me', action: () => document.getElementById('ghibli-btn')?.click() },
    { label: '📄 File / PDF', action: () => document.getElementById('file-gen-btn')?.click() },
    { label: '📚 Homework & Study', action: () => document.getElementById('homework-btn')?.click() },
    { label: '🌤 Weather', action: () => document.getElementById('weather-btn')?.click() },
    { label: '🔍 Search Web', action: () => document.getElementById('search-btn')?.click() },
    { label: '💻 Code', action: () => document.getElementById('code-workspace-btn')?.click() },
    { label: '🔎 Search My Chats', action: () => document.getElementById('search-chats-btn')?.click() },
    { label: '⏰ Reminders', action: () => document.getElementById('reminders-btn')?.click() },
    { label: '🔖 Bookmarks', action: () => document.getElementById('bookmarks-btn')?.click() },
    { label: '📊 Stats', action: () => document.getElementById('stats-btn')?.click() },
    { label: '📦 Artifacts', action: () => document.getElementById('artifacts-tab-btn')?.click() },
    { label: '🗂 Cowork Mode', action: () => document.querySelector('.mode-tab[data-mode="cowork"]')?.click() },
  ];
  items.forEach(it => {
    const row = document.createElement('div');
    row.textContent = it.label;
    row.style.cssText = 'padding:10px 14px;font-size:13px;cursor:pointer;color:var(--text);white-space:nowrap;';
    row.addEventListener('mouseenter', () => row.style.background = 'var(--accent-dim)');
    row.addEventListener('mouseleave', () => row.style.background = '');
    row.addEventListener('click', () => { it.action(); closePlusMenu(); });
    menu.appendChild(row);
  });
  document.body.appendChild(menu);
  setTimeout(() => {
    document.addEventListener('click', function onDocClick(ev) {
      if (!menu.contains(ev.target)) { closePlusMenu(); document.removeEventListener('click', onDocClick); }
    });
  }, 0);
});
fileInput.addEventListener('change', () => handleFileSelect(fileInput.files[0]));
cameraInput.addEventListener('change', () => handleFileSelect(cameraInput.files[0]));
pendingRemove.addEventListener('click', () => {
  pendingFile = null;
  fileInput.value = '';
  cameraInput.value = '';
  pendingAttach.classList.remove('show');
});

const IMAGE_KEYWORDS = /\b(generate|create|draw|make|paint|render|show me|ghibli|anime|realistic|cartoon|portrait|landscape|art|artwork|image of|picture of|photo of|illustration)\b/i;
async function tryGenerateImage(prompt) {
  try {
    const r = await fetch('/api/generate-image', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ prompt })
    });
    const d = await r.json();
    if (d.image) {
      addImageMessage('ai', d.image, '');
      return true;
    }
  } catch {}
  return false;
}

// --- Downloadable file (PDF/Word/Text) generation, triggered from chat -----
const FILE_KEYWORDS = /\b(generate|create|make|write|give me).{0,40}\b(pdf|word doc|word document|docx|downloadable file|download(able)? (a |the )?(file|document)|a file)\b|\bdownload(able)? (pdf|doc|document|file)\b/i;

function _b64ToBlob(b64, mime) {
  const bytes = atob(b64);
  const arr = new Uint8Array(bytes.length);
  for (let i = 0; i < bytes.length; i++) arr[i] = bytes.charCodeAt(i);
  return new Blob([arr], { type: mime });
}

async function tryGenerateFile(promptText, replyText) {
  const fmt = /\bword\b|\bdocx\b|\bdoc\b/i.test(promptText) ? 'docx'
    : /\btext file\b|\.txt\b/i.test(promptText) ? 'txt' : 'pdf';
  try {
    const r = await fetch('/api/generate-file', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ content: replyText, format: fmt, title: 'Mythic AI Document' })
    });
    const d = await r.json();
    if (d.file) {
      addFileMessage(d.file, d.filename, d.mimeType, d.note);
      return true;
    }
  } catch {}
  return false;
}

function addFileMessage(fileB64, filename, mimeType, note) {
  clearEmptyState();
  const row = document.createElement('div');
  row.className = 'msg-row ai';
  const div = document.createElement('div');
  div.className = 'msg ai';
  const label = document.createElement('div');
  label.textContent = "Here's your file, ready to download:";
  div.appendChild(label);
  const blob = _b64ToBlob(fileB64, mimeType);
  const url = URL.createObjectURL(blob);
  const link = document.createElement('a');
  link.className = 'file-download-chip';
  link.href = url;
  link.download = filename;
  link.textContent = '⬇ ' + filename;
  div.appendChild(link);
  if (note) {
    const noteEl = document.createElement('div');
    noteEl.style.cssText = 'font-size:11px;color:var(--muted);margin-top:6px;';
    noteEl.textContent = note;
    div.appendChild(noteEl);
  }
  row.appendChild(div);
  messagesEl.appendChild(row);
  scrollToBottom();
}

let showingStarredOnly = false;

function buildConvItem(c) {
  const item = document.createElement('div');
  item.className = 'conv-item' + (c.id === activeConvId ? ' active' : '');
  item.innerHTML = '<span class="title"></span>'
    + '<button class="menu-btn" title="More">⋮</button>';
  item.querySelector('.title').textContent = (c.pinned ? '⭐ ' : '') + c.title;

  item.addEventListener('click', (e) => {
    if (e.target.closest('.menu-btn')) return;
    openConversation(c.id);
  });

  item.querySelector('.menu-btn').addEventListener('click', (e) => {
    e.stopPropagation();
    document.querySelectorAll('.conv-menu-dropdown').forEach(el => el.remove());
    const menu = document.createElement('div');
    menu.className = 'conv-menu-dropdown';
    menu.style.cssText = 'position:fixed;background:var(--panel);border:1px solid var(--border);border-radius:8px;box-shadow:0 4px 20px rgba(0,0,0,.3);z-index:400;min-width:170px;overflow:hidden;';
    const rect = e.target.getBoundingClientRect();
    menu.style.top = (rect.bottom + 4) + 'px';
    menu.style.left = Math.max(8, rect.right - 180) + 'px';

    const menuItems = [
      { label: (c.pinned ? '⭐ Unstar' : '☆ Star'), action: async () => {
          await fetch('/api/conversations/' + c.id, { method: 'PATCH', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify({ pinned: !c.pinned }) });
          loadConversationList();
        } },
      { label: '⎘ Duplicate', action: async () => {
          const r = await fetch('/api/conversations/' + c.id + '/duplicate', { method: 'POST' });
          const d = await r.json();
          if (d.id) openConversation(d.id);
        } },
      { label: '📁 Move to folder', action: async () => {
          const folders = await fetch('/api/folders').then(r => r.json()).then(d => d.folders || []).catch(() => []);
          const hint = folders.length ? ('Existing: ' + folders.join(', ') + '\n\n') : '';
          const name = prompt(hint + 'Folder name (blank to remove from folder):', c.folder || '');
          if (name === null) return;
          await fetch('/api/conversations/' + c.id, { method: 'PATCH', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify({ folder: name.trim() }) });
          loadConversationList();
        } },
      { label: '✎ Rename', action: async () => {
          const newTitle = prompt('Rename chat:', c.title);
          if (!newTitle || !newTitle.trim() || newTitle.trim() === c.title) return;
          await fetch('/api/conversations/' + c.id, { method: 'PATCH', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify({ title: newTitle.trim() }) });
          loadConversationList();
        } },
      { label: '🔗 Share link', action: () => openShareModalFor(c.id) },
      { label: (c.archived ? '📤 Unarchive' : '🗄 Archive'), action: async () => {
          await fetch('/api/conversations/' + c.id, { method: 'PATCH', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify({ archived: !c.archived }) });
          if (c.id === activeConvId) startNewChat(); else loadConversationList();
        } },
      { label: '✕ Delete', danger: true, action: async () => {
          await fetch('/api/conversations/' + c.id, { method: 'DELETE' });
          if (c.id === activeConvId) startNewChat(); else loadConversationList();
        } },
    ];
    menuItems.forEach(it => {
      const row = document.createElement('div');
      row.textContent = it.label;
      row.style.cssText = 'padding:9px 14px;font-size:12.5px;cursor:pointer;color:' + (it.danger ? '#ef4444' : 'var(--text)') + ';';
      row.addEventListener('mouseenter', () => row.style.background = 'var(--accent-dim)');
      row.addEventListener('mouseleave', () => row.style.background = '');
      row.addEventListener('click', () => { it.action(); menu.remove(); });
      menu.appendChild(row);
    });
    document.body.appendChild(menu);
    setTimeout(() => {
      document.addEventListener('click', function closeConvMenu() {
        menu.remove();
        document.removeEventListener('click', closeConvMenu);
      }, { once: true });
    }, 0);
  });

  return item;
}

async function loadConversationList() {
  try {
    const r = await fetch('/api/conversations?archived=0');
    const d = await r.json();
    let convs = d.conversations || [];
    if (showingStarredOnly) convs = convs.filter(c => c.pinned);
    convListEl.innerHTML = '';

    if (!showingStarredOnly) {
      const byFolder = {};
      const noFolder = [];
      convs.forEach(c => {
        if (c.folder) { (byFolder[c.folder] = byFolder[c.folder] || []).push(c); }
        else noFolder.push(c);
      });
      Object.keys(byFolder).sort().forEach(folderName => {
        const header = document.createElement('div');
        header.textContent = '📁 ' + folderName;
        header.style.cssText = 'font-size:11px;font-weight:700;color:var(--muted);padding:8px 10px 3px;text-transform:uppercase;letter-spacing:.3px;';
        convListEl.appendChild(header);
        byFolder[folderName].forEach(c => convListEl.appendChild(buildConvItem(c)));
      });
      noFolder.forEach(c => convListEl.appendChild(buildConvItem(c)));
    } else {
      convs.forEach(c => convListEl.appendChild(buildConvItem(c)));
    }

    if (!convs.length) {
      const empty = document.createElement('div');
      empty.textContent = showingStarredOnly ? 'No starred chats yet.' : 'No chats yet.';
      empty.style.cssText = 'padding:16px 10px;font-size:12.5px;color:var(--muted);text-align:center;';
      convListEl.appendChild(empty);
    }
    return convs;
  } catch { return []; }
}

async function openConversation(convId, opts) {
  activeConvId = convId;
  // Reflect the open chat in the address bar (?c=<id>) so the URL at the
  // top actually changes per-chat, refreshing the page reopens the same
  // chat, and the browser Back/Forward buttons move between chats.
  // opts.updateUrl=false is used when we're reacting to a popstate event
  // (the URL already matches — pushing again would break Back/Forward).
  if (!opts || opts.updateUrl !== false) {
    try { history.pushState({ conv: convId }, '', '?c=' + encodeURIComponent(convId)); } catch {}
  }
  try {
    const r = await fetch('/api/conversations/' + convId);
    if (!r.ok) return;
    const d = await r.json();
    messagesEl.innerHTML = '';
    _msgIndexCounter = 0;
    (d.messages || []).forEach(m => addMessage(m.role, m.text, m.attachment));
    loadConversationList();
  } catch {}
  refreshShareBtnState();
  if (isMobile()) closeSidebar();
}

function startNewChat(opts) {
  activeConvId = null;
  messagesEl.innerHTML = '';
  _msgIndexCounter = 0;
  showEmptyState();
  refreshShareBtnState();
  if (!opts || opts.updateUrl !== false) {
    try { history.pushState({}, '', location.pathname); } catch {}
  }
  // No conversation is created on the server here anymore — it now only
  // gets created once the first real message is sent (see streamReply()),
  // so an empty "New chat" entry never shows up in the sidebar before the
  // person has actually typed anything.
  loadConversationList();
}

// Keep the address bar in sync with Back/Forward navigation between chats.
window.addEventListener('popstate', () => {
  const id = new URLSearchParams(location.search).get('c');
  if (id) openConversation(id, { updateUrl: false });
  else startNewChat({ updateUrl: false });
});

function refreshShareBtnState() {
  // The invite link is static and account-wide now (not per-conversation),
  // so there's no per-chat "shared" state to check on the server anymore.
  // Kept as a function (rather than removing every call site) so nothing
  // else in the file needs to change.
}

let isGenerating = false;
let currentAbortController = null;

function setGenerating(state) {
  isGenerating = state;
  sendBtn.classList.toggle('generating', state);
  sendBtn.title = state ? 'Stop generating' : 'Send';
  sendBtn.innerHTML = state
    ? '<svg width="14" height="14" viewBox="0 0 24 24" fill="currentColor"><rect x="5" y="5" width="14" height="14" rx="2"/></svg>'
    : '<svg width="18" height="18" viewBox="0 0 24 24" fill="currentColor"><path d="M2.01 21L23 12 2.01 3 2 10l15 2-15 2z"/></svg>';
}

let _lastUserMessageText = '';

async function streamReply({ message = null, attachment = null, regenerate = false } = {}) {
  showTyping();
  setGenerating(true);
  currentAbortController = new AbortController();
  _lastUserMessageText = message || _lastUserMessageText;

  if (!regenerate && currentMode === 'chat') {
    const wantsFile = FILE_KEYWORDS.test(message || '');
    const wantsImage = !wantsFile && IMAGE_KEYWORDS.test(message || '') && !attachment;
    if (wantsImage) {
      hideTyping();
      const generated = await tryGenerateImage(message);
      if (generated) { setGenerating(false); loadConversationList(); refreshStreakBadge(); return; }
      showTyping();
    }
  }

  let aiTextNode = null;
  // Watchdogs so a hung/dead connection ALWAYS surfaces something visible
  // instead of silently sitting there forever — this is the #1 cause of
  // "nothing happens" on mobile networks, where a carrier's NAT/proxy can
  // silently drop a long-idle streaming connection without ever sending
  // fetch() a clean error.
  let noResponseTimer = null;
  let streamStallTimer = null;
  const STREAM_STALL_MS = 20000; // no NEW chunk for 20s while streaming = stalled
  const NO_RESPONSE_MS = 45000;  // no response headers at all for 45s = dead
  function clearWatchdogs() {
    if (noResponseTimer) { clearTimeout(noResponseTimer); noResponseTimer = null; }
    if (streamStallTimer) { clearTimeout(streamStallTimer); streamStallTimer = null; }
  }
  try {
    const chatPayload = {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      signal: currentAbortController.signal,
      body: JSON.stringify({
        message: message || '',
        conversation_id: activeConvId,
        attachment,
        user_name: getUserName(),
        regenerate: !!regenerate,
        model: selectedModel,
        mode: currentMode === 'chat' ? 'default' : currentMode,
        ...getUserApiKeys(),
      })
    };

    noResponseTimer = setTimeout(() => {
      if (currentAbortController) currentAbortController.abort('no-response-timeout');
    }, NO_RESPONSE_MS);

    // On free-tier hosts (Render, etc.) the server can go to sleep after
    // inactivity — the very first request after that wakes it up but often
    // fails or times out while it's cold-starting. Rather than immediately
    // treating that as a real error (and popping the trouble-report modal),
    // retry once after a short pause before giving up for real.
    let r;
    try {
      r = await fetch('/api/chat', chatPayload);
    } catch (firstErr) {
      if (firstErr.name === 'AbortError') throw firstErr;
      showTyping('Server is waking up, retrying…');
      await new Promise(res => setTimeout(res, 3000));
      r = await fetch('/api/chat', chatPayload);
    }
    clearWatchdogs();
    if (!r.ok || !r.body) {
      hideTyping();
      let errMsg = 'Something went wrong. Try again.';
      try {
        const errData = await r.clone().json();
        if (errData && errData.error) errMsg = errData.error;
      } catch {
        try {
          const errText = await r.clone().text();
          if (errText && errText.trim()) errMsg = errText.trim().slice(0, 300);
        } catch {}
      }
      addMessage('error', errMsg + ` (HTTP ${r.status})`);
      return;
    }
    hideTyping();
    aiTextNode = addMessage('ai', '');

    const convId = r.headers.get('X-Conversation-Id');
    if (convId) activeConvId = convId;

    const reader = r.body.getReader();
    const decoder = new TextDecoder();
    let fullText = '';
    while (true) {
      // Race each chunk read against a stall timer — if the connection has
      // gone dead mid-stream (common on cellular networks that silently
      // kill idle-looking connections), abort instead of hanging forever.
      const stallPromise = new Promise((_, reject) => {
        streamStallTimer = setTimeout(() => reject(new Error('stream-stalled')), STREAM_STALL_MS);
      });
      let readResult;
      try {
        readResult = await Promise.race([reader.read(), stallPromise]);
      } finally {
        if (streamStallTimer) { clearTimeout(streamStallTimer); streamStallTimer = null; }
      }
      const { done, value } = readResult;
      if (done) break;
      const chunk = decoder.decode(value, { stream: true });
      fullText += chunk;
      aiTextNode.textContent = fullText;
      scrollToBottom();
    }
    // If the connection reported success but delivered literally nothing
    // (empty stream), don't leave a blank, invisible message row sitting
    // in the chat — that's indistinguishable from "nothing happened".
    if (!fullText.trim()) {
      fullText = "I didn't get a response back from the server that time — please try again.";
      aiTextNode.textContent = fullText;
    }
    // Per spec: Mythic AI never speaks automatically. Speech only happens
    // when the user explicitly clicks the 🔊 "Read aloud" button on a
    // message (see buildMsgActions override below, which calls speak()).
    _addArtifactsFromReply(fullText);
    loadConversationList();
    refreshStreakBadge();

    if (!regenerate && currentMode === 'chat' && FILE_KEYWORDS.test(message || '')) {
      await tryGenerateFile(message || '', fullText);
    }

    if (typeof window._notifyAiReply === 'function') {
      const preview = fullText.replace(/[#*`_~>]/g, '').trim().slice(0, 80);
      window._notifyAiReply(preview || 'Your answer is ready 💬');
    }
  } catch (err) {
    clearWatchdogs();
    hideTyping();
    const wasWatchdogAbort = err.name === 'AbortError' &&
      (currentAbortController === null || currentAbortController.signal.reason === 'no-response-timeout');
    const wasStall = err.message === 'stream-stalled';
    if (err.name === 'AbortError' && !wasWatchdogAbort) {
      // Genuinely user-initiated stop (Stop button clicked).
      if (aiTextNode && !aiTextNode.textContent.trim()) aiTextNode.textContent = '[Stopped]';
    } else if (wasWatchdogAbort || wasStall) {
      const msg = 'Lost connection to the server (this can happen on a weak '
        + 'or switching mobile connection). Please try again.';
      if (aiTextNode) {
        aiTextNode.textContent = msg;
      } else {
        addMessage('error', msg);
      }
    } else {
      addMessage('error', 'Network error: ' + err.message);
      askClarification({
        question: "What's wrong right now?",
        options: [
          "Generation isn't working when I hit Enter",
          "Replies are too slow",
          "The reply doesn't match what I asked",
        ],
        onPick: (opt) => { streamReply({ message: "Trouble report: " + opt, regenerate: false }); },
        onCustom: (text) => { streamReply({ message: text, regenerate: false }); },
      });
    }
  } finally {
    clearWatchdogs();
    setGenerating(false);
    currentAbortController = null;
  }
}

function regenerateLast(row) {
  if (isGenerating) return;
  // The row being removed occupied one index slot; the server's regenerate
  // path pops the old reply and appends a new one at that same slot (array
  // length unchanged), so roll the counter back one to keep addMessage()'s
  // next index aligned with the server instead of drifting upward forever.
  if (row.dataset.msgIndex !== undefined) _msgIndexCounter--;
  row.remove();
  streamReply({ regenerate: true });
}

form.addEventListener('submit', (e) => {
  e.preventDefault();
  if (isGenerating) return;
  const text = input.value.trim();
  if (!text && !pendingFile) return;
  const attachment = pendingFile;
  pendingFile = null;
  fileInput.value = ''; cameraInput.value = '';
  pendingAttach.classList.remove('show');
  input.value = '';
  input.style.height = 'auto';
  // Cowork mode runs a real multi-step task (plan → work each step →
  // synthesize) via /api/cowork/run instead of a single streamed reply —
  // see runCoworkTask(). Everything else keeps the normal streaming flow.
  if (typeof currentMode !== 'undefined' && currentMode === 'cowork' && text) {
    runCoworkTask(text);
    return;
  }
  addMessage('user', text, attachment);
  const tonePrefix = getTonePrefix();
  streamReply({ message: tonePrefix + text, attachment });
});

sendBtn.addEventListener('click', (e) => {
  if (isGenerating) {
    e.preventDefault();
    if (currentAbortController) currentAbortController.abort();
  }
});

input.addEventListener('keydown', (e) => {
  if (e.key === 'Enter' && !e.shiftKey) { e.preventDefault(); form.requestSubmit(); }
});
function autoResize() {
  input.style.height = 'auto';
  input.style.height = Math.min(input.scrollHeight, 140) + 'px';
}
input.addEventListener('input', autoResize);

const sidebarOverlay = document.getElementById('sidebar-overlay');

function isMobile() { return window.innerWidth <= 768; }

function openSidebar() {
  sidebar.classList.remove('hidden');
  if (isMobile()) sidebarOverlay.style.display = 'block';
}
function closeSidebar() {
  sidebar.classList.add('hidden');
  sidebarOverlay.style.display = 'none';
}
sidebarToggle.addEventListener('click', () => {
  sidebar.classList.contains('hidden') ? openSidebar() : closeSidebar();
});
sidebarOverlay.addEventListener('click', closeSidebar);

const fullscreenIcon  = document.getElementById('fullscreen-icon');
const fsSupported = !!(document.documentElement.requestFullscreen || document.documentElement.webkitRequestFullscreen);

function isFullscreen() {
  return !!(document.fullscreenElement || document.webkitFullscreenElement) ||
    document.body.classList.contains('pseudo-fullscreen');
}
function updateFullscreenBtn() {
  if (isFullscreen()) {
    fullscreenIcon.textContent = '⤢';
    fullscreenBtn.classList.add('active');
    fullscreenBtn.title = 'Exit fullscreen';
  } else {
    fullscreenIcon.textContent = '⛶';
    fullscreenBtn.classList.remove('active');
    fullscreenBtn.title = 'Fullscreen';
  }
}
async function toggleFullscreen() {
  const el = document.documentElement;
  try {
    if (fsSupported) {
      if (!document.fullscreenElement && !document.webkitFullscreenElement) {
        if (el.requestFullscreen) await el.requestFullscreen();
        else if (el.webkitRequestFullscreen) el.webkitRequestFullscreen();
      } else {
        if (document.exitFullscreen) await document.exitFullscreen();
        else if (document.webkitExitFullscreen) document.webkitExitFullscreen();
      }
    } else {
      document.body.classList.toggle('pseudo-fullscreen');
      updateFullscreenBtn();
    }
  } catch (err) {
    console.warn('Fullscreen request failed:', err);
    document.body.classList.toggle('pseudo-fullscreen');
    updateFullscreenBtn();
  }
}
fullscreenBtn.addEventListener('click', toggleFullscreen);
document.addEventListener('fullscreenchange', updateFullscreenBtn);
document.addEventListener('webkitfullscreenchange', updateFullscreenBtn);

function getUserName() { return localStorage.getItem('mythic_user_name') || ''; }
function setUserName(name) {
  if (name) localStorage.setItem('mythic_user_name', name);
  else localStorage.removeItem('mythic_user_name');
  if (typeof _renderSidebarProfile === 'function') _renderSidebarProfile();
}
const PRESET_AVATARS = ["data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjMWE5YzkzIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjJjMTlmIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTBhODc3Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2YyYzE5ZiIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTBhODc3IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Y2lyY2xlIGN4PSI2NCIgY3k9IjM0IiByPSIxMiIgZmlsbD0iIzNiMjUxNyIvPgogIDxwYXRoIGQ9Ik0zMiA3OCBDMzAgNDYgNDQgMzIgNjQgMzIgQzg0IDMyIDk4IDQ2IDk2IDc4IEM5MCA2MCA4MiA1MCA2NCA1MCBDNDYgNTAgMzggNjAgMzIgNzggWiIgZmlsbD0iIzNiMjUxNyIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxnIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjYiPgogICAgPGNpcmNsZSBjeD0iNTEiIGN5PSI3OCIgcj0iOSIvPjxjaXJjbGUgY3g9Ijc3IiBjeT0iNzgiIHI9IjkiLz4KICAgIDxwYXRoIGQ9Ik02MCA3OCBMNjggNzgiLz48cGF0aCBkPSJNNDIgNzggTDM2IDc2Ii8+PHBhdGggZD0iTTg2IDc4IEw5MiA3NiIvPgogIDwvZz48cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik0zNCAxMTggTDM0IDEwMCBMNTAgMTEyIEw2NCAxMDQgTDc4IDExMiBMOTQgMTAwIEw5NCAxMTggWiIgZmlsbD0iIzNiM2I0NSIvPjxwYXRoIGQ9Ik02MCAxMDYgTDY4IDEwNiBMNjYgMTI4IEw2NCAxMzIgTDYyIDEyOCBaIiBmaWxsPSIjNWMzYTIxIi8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjN2EyZTUxIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZThiNDhhIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYzg5NDZhIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2U4YjQ4YSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYzg5NDZhIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Zz48Y2lyY2xlIGN4PSIzNC4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNDIuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9IjUxLjAiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI1OS41IiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNjguMCIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9Ijc2LjUiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI4NS4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iOTMuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxwYXRoIGQ9Ik0zNCA3MCBDMzQgNDggNDYgNDAgNjQgNDAgQzgyIDQwIDk0IDQ4IDk0IDcwIEw4OCA2NiBDODQgNTQgNzYgNDggNjQgNDggQzUyIDQ4IDQ0IDU0IDQwIDY2IFoiIGZpbGw9IiMxYTEzMTAiLz48L2c+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNNDIgODggQzQyIDEwNCA1MiAxMTYgNjQgMTE2IEM3NiAxMTYgODYgMTA0IDg2IDg4IEw4NiA5NiBDODIgMTA4IDc0IDExNCA2NCAxMTQgQzU0IDExNCA0NiAxMDggNDIgOTYgWiIgZmlsbD0iIzFhMTMxMCIgb3BhY2l0eT0iMC45MiIvPjxwYXRoIGQ9Ik0zMCAxMTggTDM2IDk4IEw1MCAxMDggTDY0IDEwMCBMNzggMTA4IEw5MiA5OCBMOTggMTE4IFoiIGZpbGw9IiNjOTRmNGYiLz4KICA8cGF0aCBkPSJNNTYgMTA0IEw2NCAxMDAgTDcyIDEwNCBMNjggMTE4IEw2MCAxMTggWiIgZmlsbD0iI2U4ZTRkYSIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjZDk3YjNmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjYzk4YTVlIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYTg2ZTQ2Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2M5OGE1ZSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYTg2ZTQ2IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzQgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NCBDOTIgNTYgODIgNDYgNjQgNDYgQzQ2IDQ2IDM2IDU2IDMyIDc0IFoiIGZpbGw9IiM1YzNhMjEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik0zNCAxMTggTDM0IDEwMiBDNDIgOTYgNTQgMTAwIDY0IDEwOCBDNzQgMTAwIDg2IDk2IDk0IDEwMiBMOTQgMTE4IFoiIGZpbGw9IiMzYTZlYTgiLz4KPC9zdmc+", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjNWZhODNiIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjOGQ1YTNjIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjNmY0NDI5Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzhkNWEzYyIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjNmY0NDI5IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzYgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NiBMOTAgNzAgQzg2IDU0IDc4IDQ2IDY0IDQ2IEM1MCA0NiA0MiA1NCAzOCA3MCBaIiBmaWxsPSIjN2E1MjMwIi8+CiAgPHBhdGggZD0iTTkyIDYwIEMxMDQgNjIgMTA4IDc4IDEwMCA5NiBDOTggODQgOTIgNzYgODggNjggWiIgZmlsbD0iIzdhNTIzMCIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxnIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjYiPgogICAgPGNpcmNsZSBjeD0iNTEiIGN5PSI3OCIgcj0iOSIvPjxjaXJjbGUgY3g9Ijc3IiBjeT0iNzgiIHI9IjkiLz4KICAgIDxwYXRoIGQ9Ik02MCA3OCBMNjggNzgiLz48cGF0aCBkPSJNNDIgNzggTDM2IDc2Ii8+PHBhdGggZD0iTTg2IDc4IEw5MiA3NiIvPgogIDwvZz48cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik01NCA5OCBDNTQgMTA4IDU4IDExNCA2NCAxMTQgQzcwIDExNCA3NCAxMDggNzQgOTggTDcwIDEwMCBDNjggMTA2IDY2IDEwOCA2NCAxMDggQzYyIDEwOCA2MCAxMDYgNTggMTAwIFoiIGZpbGw9IiM3YTUyMzAiLz48cGF0aCBkPSJNMzQgMTE4IEwzNCAxMDAgTDUwIDExMiBMNjQgMTA0IEw3OCAxMTIgTDk0IDEwMCBMOTQgMTE4IFoiIGZpbGw9IiM1ZmE4M2IiLz4KPC9zdmc+", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjYzk0ZjRmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjdkM2I1Ii8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTNiNDhmIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2Y3ZDNiNSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTNiNDhmIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8ZWxsaXBzZSBjeD0iNjQiIGN5PSI1OCIgcng9IjQiIHJ5PSIyIiBmaWxsPSIjMmIyMTE4IiBvcGFjaXR5PSIwIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNMzAgMTE4IEwzNiA5OCBMNTAgMTA4IEw2NCAxMDAgTDc4IDEwOCBMOTIgOTggTDk4IDExOCBaIiBmaWxsPSIjOGU1ZmE4Ii8+CiAgPHBhdGggZD0iTTU2IDEwNCBMNjQgMTAwIEw3MiAxMDQgTDY4IDExOCBMNjAgMTE4IFoiIGZpbGw9IiNlOGU0ZGEiLz4KPC9zdmc+", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjM2E2ZWE4Ii8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjJjMTlmIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTBhODc3Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2YyYzE5ZiIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTBhODc3IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNjggQzMwIDQ0IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NCA5NiA2OCBDOTAgNTggODIgNTQgNjQgNTQgQzQ2IDU0IDM4IDU4IDMyIDY4IFoiIGZpbGw9IiNjOWExNWEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik00MiA4OCBDNDIgMTA0IDUyIDExNiA2NCAxMTYgQzc2IDExNiA4NiAxMDQgODYgODggTDg2IDk2IEM4MiAxMDggNzQgMTE0IDY0IDExNCBDNTQgMTE0IDQ2IDEwOCA0MiA5NiBaIiBmaWxsPSIjYzlhMTVhIiBvcGFjaXR5PSIwLjkyIi8+PHBhdGggZD0iTTM0IDExOCBMMzQgMTAyIEM0MiA5NiA1NCAxMDAgNjQgMTA4IEM3NCAxMDAgODYgOTYgOTQgMTAyIEw5NCAxMTggWiIgZmlsbD0iI2Q5N2IzZiIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjOGU1ZmE4Ii8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZThiNDhhIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYzg5NDZhIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2U4YjQ4YSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYzg5NDZhIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzAgMTAwIEMyNCA3MCAzMCA0MiA2NCAzMiBDOTggNDIgMTA0IDcwIDk4IDEwMCBMOTAgOTAgQzkyIDY4IDg4IDUyIDY0IDQ2IEM0MCA1MiAzNiA2OCAzOCA5MCBaIiBmaWxsPSIjNGEyZjFjIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPGcgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNiI+CiAgICA8Y2lyY2xlIGN4PSI1MSIgY3k9Ijc4IiByPSI5Ii8+PGNpcmNsZSBjeD0iNzciIGN5PSI3OCIgcj0iOSIvPgogICAgPHBhdGggZD0iTTYwIDc4IEw2OCA3OCIvPjxwYXRoIGQ9Ik00MiA3OCBMMzYgNzYiLz48cGF0aCBkPSJNODYgNzggTDkyIDc2Ii8+CiAgPC9nPjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTM0IDExOCBMMzQgMTAwIEw1MCAxMTIgTDY0IDEwNCBMNzggMTEyIEw5NCAxMDAgTDk0IDExOCBaIiBmaWxsPSIjMmY4ZjZmIi8+PHBhdGggZD0iTTYwIDEwNiBMNjggMTA2IEw2NiAxMjggTDY0IDEzMiBMNjIgMTI4IFoiIGZpbGw9IiMxYTEzMTAiLz4KPC9zdmc+", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjMmY4ZjZmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjYzk4YTVlIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYTg2ZTQ2Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2M5OGE1ZSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYTg2ZTQ2IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Y2lyY2xlIGN4PSI2NCIgY3k9IjM0IiByPSIxMiIgZmlsbD0iIzNiMjUxNyIvPgogIDxwYXRoIGQ9Ik0zMiA3OCBDMzAgNDYgNDQgMzIgNjQgMzIgQzg0IDMyIDk4IDQ2IDk2IDc4IEM5MCA2MCA4MiA1MCA2NCA1MCBDNDYgNTAgMzggNjAgMzIgNzggWiIgZmlsbD0iIzNiMjUxNyIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTU0IDk4IEM1NCAxMDggNTggMTE0IDY0IDExNCBDNzAgMTE0IDc0IDEwOCA3NCA5OCBMNzAgMTAwIEM2OCAxMDYgNjYgMTA4IDY0IDEwOCBDNjIgMTA4IDYwIDEwNiA1OCAxMDAgWiIgZmlsbD0iIzNiMjUxNyIvPjxwYXRoIGQ9Ik0zMCAxMTggTDM2IDk4IEw1MCAxMDggTDY0IDEwMCBMNzggMTA4IEw5MiA5OCBMOTggMTE4IFoiIGZpbGw9IiMzYjNiNDUiLz4KICA8cGF0aCBkPSJNNTYgMTA0IEw2NCAxMDAgTDcyIDEwNCBMNjggMTE4IEw2MCAxMTggWiIgZmlsbD0iI2U4ZTRkYSIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjYjg4NjJmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjOGQ1YTNjIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjNmY0NDI5Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzhkNWEzYyIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjNmY0NDI5IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Zz48Y2lyY2xlIGN4PSIzNC4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNDIuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9IjUxLjAiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI1OS41IiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNjguMCIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9Ijc2LjUiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI4NS4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iOTMuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxwYXRoIGQ9Ik0zNCA3MCBDMzQgNDggNDYgNDAgNjQgNDAgQzgyIDQwIDk0IDQ4IDk0IDcwIEw4OCA2NiBDODQgNTQgNzYgNDggNjQgNDggQzUyIDQ4IDQ0IDU0IDQwIDY2IFoiIGZpbGw9IiMxYTEzMTAiLz48L2c+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNMzQgMTE4IEwzNCAxMDIgQzQyIDk2IDU0IDEwMCA2NCAxMDggQzc0IDEwMCA4NiA5NiA5NCAxMDIgTDk0IDExOCBaIiBmaWxsPSIjYzk0ZjRmIi8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjMWE5YzkzIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjdkM2I1Ii8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTNiNDhmIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2Y3ZDNiNSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTNiNDhmIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzQgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NCBDOTIgNTYgODIgNDYgNjQgNDYgQzQ2IDQ2IDM2IDU2IDMyIDc0IFoiIGZpbGw9IiM1YzNhMjEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8ZyBmaWxsPSJub25lIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi42Ij4KICAgIDxjaXJjbGUgY3g9IjUxIiBjeT0iNzgiIHI9IjkiLz48Y2lyY2xlIGN4PSI3NyIgY3k9Ijc4IiByPSI5Ii8+CiAgICA8cGF0aCBkPSJNNjAgNzggTDY4IDc4Ii8+PHBhdGggZD0iTTQyIDc4IEwzNiA3NiIvPjxwYXRoIGQ9Ik04NiA3OCBMOTIgNzYiLz4KICA8L2c+PHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNNDIgODggQzQyIDEwNCA1MiAxMTYgNjQgMTE2IEM3NiAxMTYgODYgMTA0IDg2IDg4IEw4NiA5NiBDODIgMTA4IDc0IDExNCA2NCAxMTQgQzU0IDExNCA0NiAxMDggNDIgOTYgWiIgZmlsbD0iIzVjM2EyMSIgb3BhY2l0eT0iMC45MiIvPjxwYXRoIGQ9Ik0zNCAxMTggTDM0IDEwMCBMNTAgMTEyIEw2NCAxMDQgTDc4IDExMiBMOTQgMTAwIEw5NCAxMTggWiIgZmlsbD0iIzNhNmVhOCIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjN2EyZTUxIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjJjMTlmIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTBhODc3Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2YyYzE5ZiIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTBhODc3IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzYgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NiBMOTAgNzAgQzg2IDU0IDc4IDQ2IDY0IDQ2IEM1MCA0NiA0MiA1NCAzOCA3MCBaIiBmaWxsPSIjN2E1MjMwIi8+CiAgPHBhdGggZD0iTTkyIDYwIEMxMDQgNjIgMTA4IDc4IDEwMCA5NiBDOTggODQgOTIgNzYgODggNjggWiIgZmlsbD0iIzdhNTIzMCIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTMwIDExOCBMMzYgOTggTDUwIDEwOCBMNjQgMTAwIEw3OCAxMDggTDkyIDk4IEw5OCAxMTggWiIgZmlsbD0iIzVmYTgzYiIvPgogIDxwYXRoIGQ9Ik01NiAxMDQgTDY0IDEwMCBMNzIgMTA0IEw2OCAxMTggTDYwIDExOCBaIiBmaWxsPSIjZThlNGRhIi8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjZDk3YjNmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZThiNDhhIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYzg5NDZhIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2U4YjQ4YSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYzg5NDZhIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8ZWxsaXBzZSBjeD0iNjQiIGN5PSI1OCIgcng9IjQiIHJ5PSIyIiBmaWxsPSIjMmIyMTE4IiBvcGFjaXR5PSIwIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNNTQgOTggQzU0IDEwOCA1OCAxMTQgNjQgMTE0IEM3MCAxMTQgNzQgMTA4IDc0IDk4IEw3MCAxMDAgQzY4IDEwNiA2NiAxMDggNjQgMTA4IEM2MiAxMDggNjAgMTA2IDU4IDEwMCBaIiBmaWxsPSIjMmIyMTE4Ii8+PHBhdGggZD0iTTM0IDExOCBMMzQgMTAyIEM0MiA5NiA1NCAxMDAgNjQgMTA4IEM3NCAxMDAgODYgOTYgOTQgMTAyIEw5NCAxMTggWiIgZmlsbD0iIzhlNWZhOCIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjNWZhODNiIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjYzk4YTVlIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYTg2ZTQ2Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2M5OGE1ZSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYTg2ZTQ2IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNjggQzMwIDQ0IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NCA5NiA2OCBDOTAgNTggODIgNTQgNjQgNTQgQzQ2IDU0IDM4IDU4IDMyIDY4IFoiIGZpbGw9IiNjOWExNWEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8ZyBmaWxsPSJub25lIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi42Ij4KICAgIDxjaXJjbGUgY3g9IjUxIiBjeT0iNzgiIHI9IjkiLz48Y2lyY2xlIGN4PSI3NyIgY3k9Ijc4IiByPSI5Ii8+CiAgICA8cGF0aCBkPSJNNjAgNzggTDY4IDc4Ii8+PHBhdGggZD0iTTQyIDc4IEwzNiA3NiIvPjxwYXRoIGQ9Ik04NiA3OCBMOTIgNzYiLz4KICA8L2c+PHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNMzQgMTE4IEwzNCAxMDAgTDUwIDExMiBMNjQgMTA0IEw3OCAxMTIgTDk0IDEwMCBMOTQgMTE4IFoiIGZpbGw9IiNkOTdiM2YiLz48cGF0aCBkPSJNNjAgMTA2IEw2OCAxMDYgTDY2IDEyOCBMNjQgMTMyIEw2MiAxMjggWiIgZmlsbD0iIzNiMjUxNyIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjYzk0ZjRmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjOGQ1YTNjIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjNmY0NDI5Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzhkNWEzYyIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjNmY0NDI5IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzAgMTAwIEMyNCA3MCAzMCA0MiA2NCAzMiBDOTggNDIgMTA0IDcwIDk4IDEwMCBMOTAgOTAgQzkyIDY4IDg4IDUyIDY0IDQ2IEM0MCA1MiAzNiA2OCAzOCA5MCBaIiBmaWxsPSIjNGEyZjFjIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNNDIgODggQzQyIDEwNCA1MiAxMTYgNjQgMTE2IEM3NiAxMTYgODYgMTA0IDg2IDg4IEw4NiA5NiBDODIgMTA4IDc0IDExNCA2NCAxMTQgQzU0IDExNCA0NiAxMDggNDIgOTYgWiIgZmlsbD0iIzRhMmYxYyIgb3BhY2l0eT0iMC45MiIvPjxwYXRoIGQ9Ik0zMCAxMTggTDM2IDk4IEw1MCAxMDggTDY0IDEwMCBMNzggMTA4IEw5MiA5OCBMOTggMTE4IFoiIGZpbGw9IiMyZjhmNmYiLz4KICA8cGF0aCBkPSJNNTYgMTA0IEw2NCAxMDAgTDcyIDEwNCBMNjggMTE4IEw2MCAxMTggWiIgZmlsbD0iI2U4ZTRkYSIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjM2E2ZWE4Ii8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjdkM2I1Ii8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTNiNDhmIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2Y3ZDNiNSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTNiNDhmIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Y2lyY2xlIGN4PSI2NCIgY3k9IjM0IiByPSIxMiIgZmlsbD0iIzNiMjUxNyIvPgogIDxwYXRoIGQ9Ik0zMiA3OCBDMzAgNDYgNDQgMzIgNjQgMzIgQzg0IDMyIDk4IDQ2IDk2IDc4IEM5MCA2MCA4MiA1MCA2NCA1MCBDNDYgNTAgMzggNjAgMzIgNzggWiIgZmlsbD0iIzNiMjUxNyIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTM0IDExOCBMMzQgMTAyIEM0MiA5NiA1NCAxMDAgNjQgMTA4IEM3NCAxMDAgODYgOTYgOTQgMTAyIEw5NCAxMTggWiIgZmlsbD0iIzNiM2I0NSIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjOGU1ZmE4Ii8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjJjMTlmIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTBhODc3Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2YyYzE5ZiIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTBhODc3IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Zz48Y2lyY2xlIGN4PSIzNC4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNDIuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9IjUxLjAiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI1OS41IiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNjguMCIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9Ijc2LjUiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI4NS4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iOTMuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxwYXRoIGQ9Ik0zNCA3MCBDMzQgNDggNDYgNDAgNjQgNDAgQzgyIDQwIDk0IDQ4IDk0IDcwIEw4OCA2NiBDODQgNTQgNzYgNDggNjQgNDggQzUyIDQ4IDQ0IDU0IDQwIDY2IFoiIGZpbGw9IiMxYTEzMTAiLz48L2c+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPGcgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNiI+CiAgICA8Y2lyY2xlIGN4PSI1MSIgY3k9Ijc4IiByPSI5Ii8+PGNpcmNsZSBjeD0iNzciIGN5PSI3OCIgcj0iOSIvPgogICAgPHBhdGggZD0iTTYwIDc4IEw2OCA3OCIvPjxwYXRoIGQ9Ik00MiA3OCBMMzYgNzYiLz48cGF0aCBkPSJNODYgNzggTDkyIDc2Ii8+CiAgPC9nPjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTU0IDk4IEM1NCAxMDggNTggMTE0IDY0IDExNCBDNzAgMTE0IDc0IDEwOCA3NCA5OCBMNzAgMTAwIEM2OCAxMDYgNjYgMTA4IDY0IDEwOCBDNjIgMTA4IDYwIDEwNiA1OCAxMDAgWiIgZmlsbD0iIzFhMTMxMCIvPjxwYXRoIGQ9Ik0zNCAxMTggTDM0IDEwMCBMNTAgMTEyIEw2NCAxMDQgTDc4IDExMiBMOTQgMTAwIEw5NCAxMTggWiIgZmlsbD0iI2M5NGY0ZiIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjMmY4ZjZmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZThiNDhhIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYzg5NDZhIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2U4YjQ4YSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYzg5NDZhIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzQgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NCBDOTIgNTYgODIgNDYgNjQgNDYgQzQ2IDQ2IDM2IDU2IDMyIDc0IFoiIGZpbGw9IiM1YzNhMjEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik0zMCAxMTggTDM2IDk4IEw1MCAxMDggTDY0IDEwMCBMNzggMTA4IEw5MiA5OCBMOTggMTE4IFoiIGZpbGw9IiMzYTZlYTgiLz4KICA8cGF0aCBkPSJNNTYgMTA0IEw2NCAxMDAgTDcyIDEwNCBMNjggMTE4IEw2MCAxMTggWiIgZmlsbD0iI2U4ZTRkYSIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjYjg4NjJmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjYzk4YTVlIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYTg2ZTQ2Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2M5OGE1ZSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYTg2ZTQ2IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzYgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NiBMOTAgNzAgQzg2IDU0IDc4IDQ2IDY0IDQ2IEM1MCA0NiA0MiA1NCAzOCA3MCBaIiBmaWxsPSIjN2E1MjMwIi8+CiAgPHBhdGggZD0iTTkyIDYwIEMxMDQgNjIgMTA4IDc4IDEwMCA5NiBDOTggODQgOTIgNzYgODggNjggWiIgZmlsbD0iIzdhNTIzMCIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTQyIDg4IEM0MiAxMDQgNTIgMTE2IDY0IDExNiBDNzYgMTE2IDg2IDEwNCA4NiA4OCBMODYgOTYgQzgyIDEwOCA3NCAxMTQgNjQgMTE0IEM1NCAxMTQgNDYgMTA4IDQyIDk2IFoiIGZpbGw9IiM3YTUyMzAiIG9wYWNpdHk9IjAuOTIiLz48cGF0aCBkPSJNMzQgMTE4IEwzNCAxMDIgQzQyIDk2IDU0IDEwMCA2NCAxMDggQzc0IDEwMCA4NiA5NiA5NCAxMDIgTDk0IDExOCBaIiBmaWxsPSIjNWZhODNiIi8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjMWE5YzkzIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjOGQ1YTNjIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjNmY0NDI5Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzhkNWEzYyIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjNmY0NDI5IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8ZWxsaXBzZSBjeD0iNjQiIGN5PSI1OCIgcng9IjQiIHJ5PSIyIiBmaWxsPSIjMmIyMTE4IiBvcGFjaXR5PSIwIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPGcgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNiI+CiAgICA8Y2lyY2xlIGN4PSI1MSIgY3k9Ijc4IiByPSI5Ii8+PGNpcmNsZSBjeD0iNzciIGN5PSI3OCIgcj0iOSIvPgogICAgPHBhdGggZD0iTTYwIDc4IEw2OCA3OCIvPjxwYXRoIGQ9Ik00MiA3OCBMMzYgNzYiLz48cGF0aCBkPSJNODYgNzggTDkyIDc2Ii8+CiAgPC9nPjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTM0IDExOCBMMzQgMTAwIEw1MCAxMTIgTDY0IDEwNCBMNzggMTEyIEw5NCAxMDAgTDk0IDExOCBaIiBmaWxsPSIjOGU1ZmE4Ii8+PHBhdGggZD0iTTYwIDEwNiBMNjggMTA2IEw2NiAxMjggTDY0IDEzMiBMNjIgMTI4IFoiIGZpbGw9IiM0YTJmMWMiLz4KPC9zdmc+", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjN2EyZTUxIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjdkM2I1Ii8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTNiNDhmIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2Y3ZDNiNSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTNiNDhmIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNjggQzMwIDQ0IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NCA5NiA2OCBDOTAgNTggODIgNTQgNjQgNTQgQzQ2IDU0IDM4IDU4IDMyIDY4IFoiIGZpbGw9IiNjOWExNWEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik01NCA5OCBDNTQgMTA4IDU4IDExNCA2NCAxMTQgQzcwIDExNCA3NCAxMDggNzQgOTggTDcwIDEwMCBDNjggMTA2IDY2IDEwOCA2NCAxMDggQzYyIDEwOCA2MCAxMDYgNTggMTAwIFoiIGZpbGw9IiNjOWExNWEiLz48cGF0aCBkPSJNMzAgMTE4IEwzNiA5OCBMNTAgMTA4IEw2NCAxMDAgTDc4IDEwOCBMOTIgOTggTDk4IDExOCBaIiBmaWxsPSIjZDk3YjNmIi8+CiAgPHBhdGggZD0iTTU2IDEwNCBMNjQgMTAwIEw3MiAxMDQgTDY4IDExOCBMNjAgMTE4IFoiIGZpbGw9IiNlOGU0ZGEiLz4KPC9zdmc+", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjZDk3YjNmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjJjMTlmIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTBhODc3Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2YyYzE5ZiIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTBhODc3IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzAgMTAwIEMyNCA3MCAzMCA0MiA2NCAzMiBDOTggNDIgMTA0IDcwIDk4IDEwMCBMOTAgOTAgQzkyIDY4IDg4IDUyIDY0IDQ2IEM0MCA1MiAzNiA2OCAzOCA5MCBaIiBmaWxsPSIjNGEyZjFjIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNMzQgMTE4IEwzNCAxMDIgQzQyIDk2IDU0IDEwMCA2NCAxMDggQzc0IDEwMCA4NiA5NiA5NCAxMDIgTDk0IDExOCBaIiBmaWxsPSIjMmY4ZjZmIi8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjNWZhODNiIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZThiNDhhIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYzg5NDZhIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2U4YjQ4YSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYzg5NDZhIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Y2lyY2xlIGN4PSI2NCIgY3k9IjM0IiByPSIxMiIgZmlsbD0iIzNiMjUxNyIvPgogIDxwYXRoIGQ9Ik0zMiA3OCBDMzAgNDYgNDQgMzIgNjQgMzIgQzg0IDMyIDk4IDQ2IDk2IDc4IEM5MCA2MCA4MiA1MCA2NCA1MCBDNDYgNTAgMzggNjAgMzIgNzggWiIgZmlsbD0iIzNiMjUxNyIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxnIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjYiPgogICAgPGNpcmNsZSBjeD0iNTEiIGN5PSI3OCIgcj0iOSIvPjxjaXJjbGUgY3g9Ijc3IiBjeT0iNzgiIHI9IjkiLz4KICAgIDxwYXRoIGQ9Ik02MCA3OCBMNjggNzgiLz48cGF0aCBkPSJNNDIgNzggTDM2IDc2Ii8+PHBhdGggZD0iTTg2IDc4IEw5MiA3NiIvPgogIDwvZz48cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik00MiA4OCBDNDIgMTA0IDUyIDExNiA2NCAxMTYgQzc2IDExNiA4NiAxMDQgODYgODggTDg2IDk2IEM4MiAxMDggNzQgMTE0IDY0IDExNCBDNTQgMTE0IDQ2IDEwOCA0MiA5NiBaIiBmaWxsPSIjM2IyNTE3IiBvcGFjaXR5PSIwLjkyIi8+PHBhdGggZD0iTTM0IDExOCBMMzQgMTAwIEw1MCAxMTIgTDY0IDEwNCBMNzggMTEyIEw5NCAxMDAgTDk0IDExOCBaIiBmaWxsPSIjM2IzYjQ1Ii8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjYzk0ZjRmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjYzk4YTVlIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYTg2ZTQ2Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2M5OGE1ZSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYTg2ZTQ2IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Zz48Y2lyY2xlIGN4PSIzNC4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNDIuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9IjUxLjAiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI1OS41IiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNjguMCIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9Ijc2LjUiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI4NS4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iOTMuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxwYXRoIGQ9Ik0zNCA3MCBDMzQgNDggNDYgNDAgNjQgNDAgQzgyIDQwIDk0IDQ4IDk0IDcwIEw4OCA2NiBDODQgNTQgNzYgNDggNjQgNDggQzUyIDQ4IDQ0IDU0IDQwIDY2IFoiIGZpbGw9IiMxYTEzMTAiLz48L2c+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNMzAgMTE4IEwzNiA5OCBMNTAgMTA4IEw2NCAxMDAgTDc4IDEwOCBMOTIgOTggTDk4IDExOCBaIiBmaWxsPSIjYzk0ZjRmIi8+CiAgPHBhdGggZD0iTTU2IDEwNCBMNjQgMTAwIEw3MiAxMDQgTDY4IDExOCBMNjAgMTE4IFoiIGZpbGw9IiNlOGU0ZGEiLz4KPC9zdmc+", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjM2E2ZWE4Ii8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjOGQ1YTNjIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjNmY0NDI5Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzhkNWEzYyIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjNmY0NDI5IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzQgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NCBDOTIgNTYgODIgNDYgNjQgNDYgQzQ2IDQ2IDM2IDU2IDMyIDc0IFoiIGZpbGw9IiM1YzNhMjEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik01NCA5OCBDNTQgMTA4IDU4IDExNCA2NCAxMTQgQzcwIDExNCA3NCAxMDggNzQgOTggTDcwIDEwMCBDNjggMTA2IDY2IDEwOCA2NCAxMDggQzYyIDEwOCA2MCAxMDYgNTggMTAwIFoiIGZpbGw9IiM1YzNhMjEiLz48cGF0aCBkPSJNMzQgMTE4IEwzNCAxMDIgQzQyIDk2IDU0IDEwMCA2NCAxMDggQzc0IDEwMCA4NiA5NiA5NCAxMDIgTDk0IDExOCBaIiBmaWxsPSIjM2E2ZWE4Ii8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjOGU1ZmE4Ii8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjdkM2I1Ii8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTNiNDhmIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2Y3ZDNiNSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTNiNDhmIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzYgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NiBMOTAgNzAgQzg2IDU0IDc4IDQ2IDY0IDQ2IEM1MCA0NiA0MiA1NCAzOCA3MCBaIiBmaWxsPSIjN2E1MjMwIi8+CiAgPHBhdGggZD0iTTkyIDYwIEMxMDQgNjIgMTA4IDc4IDEwMCA5NiBDOTggODQgOTIgNzYgODggNjggWiIgZmlsbD0iIzdhNTIzMCIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxnIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjYiPgogICAgPGNpcmNsZSBjeD0iNTEiIGN5PSI3OCIgcj0iOSIvPjxjaXJjbGUgY3g9Ijc3IiBjeT0iNzgiIHI9IjkiLz4KICAgIDxwYXRoIGQ9Ik02MCA3OCBMNjggNzgiLz48cGF0aCBkPSJNNDIgNzggTDM2IDc2Ii8+PHBhdGggZD0iTTg2IDc4IEw5MiA3NiIvPgogIDwvZz48cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik0zNCAxMTggTDM0IDEwMCBMNTAgMTEyIEw2NCAxMDQgTDc4IDExMiBMOTQgMTAwIEw5NCAxMTggWiIgZmlsbD0iIzVmYTgzYiIvPjxwYXRoIGQ9Ik02MCAxMDYgTDY4IDEwNiBMNjYgMTI4IEw2NCAxMzIgTDYyIDEyOCBaIiBmaWxsPSIjYzlhMTVhIi8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjMmY4ZjZmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjJjMTlmIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTBhODc3Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2YyYzE5ZiIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTBhODc3IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8ZWxsaXBzZSBjeD0iNjQiIGN5PSI1OCIgcng9IjQiIHJ5PSIyIiBmaWxsPSIjMmIyMTE4IiBvcGFjaXR5PSIwIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNNDIgODggQzQyIDEwNCA1MiAxMTYgNjQgMTE2IEM3NiAxMTYgODYgMTA0IDg2IDg4IEw4NiA5NiBDODIgMTA4IDc0IDExNCA2NCAxMTQgQzU0IDExNCA0NiAxMDggNDIgOTYgWiIgZmlsbD0iIzJiMjExOCIgb3BhY2l0eT0iMC45MiIvPjxwYXRoIGQ9Ik0zMCAxMTggTDM2IDk4IEw1MCAxMDggTDY0IDEwMCBMNzggMTA4IEw5MiA5OCBMOTggMTE4IFoiIGZpbGw9IiM4ZTVmYTgiLz4KICA8cGF0aCBkPSJNNTYgMTA0IEw2NCAxMDAgTDcyIDEwNCBMNjggMTE4IEw2MCAxMTggWiIgZmlsbD0iI2U4ZTRkYSIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjYjg4NjJmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZThiNDhhIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYzg5NDZhIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2U4YjQ4YSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYzg5NDZhIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNjggQzMwIDQ0IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NCA5NiA2OCBDOTAgNTggODIgNTQgNjQgNTQgQzQ2IDU0IDM4IDU4IDMyIDY4IFoiIGZpbGw9IiNjOWExNWEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNNTMgOTQgUTY0IDEwMyA3NSA5NCBRNjQgOTkgNTMgOTQgWiIgZmlsbD0iI2E4NDAyZiIvPjxwYXRoIGQ9Ik0zNCAxMTggTDM0IDEwMiBDNDIgOTYgNTQgMTAwIDY0IDEwOCBDNzQgMTAwIDg2IDk2IDk0IDEwMiBMOTQgMTE4IFoiIGZpbGw9IiNkOTdiM2YiLz4KPC9zdmc+", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjMWE5YzkzIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjYzk4YTVlIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYTg2ZTQ2Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2M5OGE1ZSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYTg2ZTQ2IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzAgMTAwIEMyNCA3MCAzMCA0MiA2NCAzMiBDOTggNDIgMTA0IDcwIDk4IDEwMCBMOTAgOTAgQzkyIDY4IDg4IDUyIDY0IDQ2IEM0MCA1MiAzNiA2OCAzOCA5MCBaIiBmaWxsPSIjNGEyZjFjIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPGcgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNiI+CiAgICA8Y2lyY2xlIGN4PSI1MSIgY3k9Ijc4IiByPSI5Ii8+PGNpcmNsZSBjeD0iNzciIGN5PSI3OCIgcj0iOSIvPgogICAgPHBhdGggZD0iTTYwIDc4IEw2OCA3OCIvPjxwYXRoIGQ9Ik00MiA3OCBMMzYgNzYiLz48cGF0aCBkPSJNODYgNzggTDkyIDc2Ii8+CiAgPC9nPjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTU0IDk4IEM1NCAxMDggNTggMTE0IDY0IDExNCBDNzAgMTE0IDc0IDEwOCA3NCA5OCBMNzAgMTAwIEM2OCAxMDYgNjYgMTA4IDY0IDEwOCBDNjIgMTA4IDYwIDEwNiA1OCAxMDAgWiIgZmlsbD0iIzRhMmYxYyIvPjxwYXRoIGQ9Ik0zNCAxMTggTDM0IDEwMCBMNTAgMTEyIEw2NCAxMDQgTDc4IDExMiBMOTQgMTAwIEw5NCAxMTggWiIgZmlsbD0iIzJmOGY2ZiIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjN2EyZTUxIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjOGQ1YTNjIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjNmY0NDI5Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzhkNWEzYyIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjNmY0NDI5IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Y2lyY2xlIGN4PSI2NCIgY3k9IjM0IiByPSIxMiIgZmlsbD0iIzNiMjUxNyIvPgogIDxwYXRoIGQ9Ik0zMiA3OCBDMzAgNDYgNDQgMzIgNjQgMzIgQzg0IDMyIDk4IDQ2IDk2IDc4IEM5MCA2MCA4MiA1MCA2NCA1MCBDNDYgNTAgMzggNjAgMzIgNzggWiIgZmlsbD0iIzNiMjUxNyIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTMwIDExOCBMMzYgOTggTDUwIDEwOCBMNjQgMTAwIEw3OCAxMDggTDkyIDk4IEw5OCAxMTggWiIgZmlsbD0iIzNiM2I0NSIvPgogIDxwYXRoIGQ9Ik01NiAxMDQgTDY0IDEwMCBMNzIgMTA0IEw2OCAxMTggTDYwIDExOCBaIiBmaWxsPSIjZThlNGRhIi8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjZDk3YjNmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjdkM2I1Ii8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTNiNDhmIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2Y3ZDNiNSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTNiNDhmIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8Zz48Y2lyY2xlIGN4PSIzNC4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNDIuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9IjUxLjAiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI1OS41IiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iNjguMCIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxjaXJjbGUgY3g9Ijc2LjUiIGN5PSI0NiIgcj0iMTIiIGZpbGw9IiMxYTEzMTAiLz48Y2lyY2xlIGN4PSI4NS4wIiBjeT0iNTgiIHI9IjEyIiBmaWxsPSIjMWExMzEwIi8+PGNpcmNsZSBjeD0iOTMuNSIgY3k9IjUyIiByPSIxMiIgZmlsbD0iIzFhMTMxMCIvPjxwYXRoIGQ9Ik0zNCA3MCBDMzQgNDggNDYgNDAgNjQgNDAgQzgyIDQwIDk0IDQ4IDk0IDcwIEw4OCA2NiBDODQgNTQgNzYgNDggNjQgNDggQzUyIDQ4IDQ0IDU0IDQwIDY2IFoiIGZpbGw9IiMxYTEzMTAiLz48L2c+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNNDIgODggQzQyIDEwNCA1MiAxMTYgNjQgMTE2IEM3NiAxMTYgODYgMTA0IDg2IDg4IEw4NiA5NiBDODIgMTA4IDc0IDExNCA2NCAxMTQgQzU0IDExNCA0NiAxMDggNDIgOTYgWiIgZmlsbD0iIzFhMTMxMCIgb3BhY2l0eT0iMC45MiIvPjxwYXRoIGQ9Ik0zNCAxMTggTDM0IDEwMiBDNDIgOTYgNTQgMTAwIDY0IDEwOCBDNzQgMTAwIDg2IDk2IDk0IDEwMiBMOTQgMTE4IFoiIGZpbGw9IiNjOTRmNGYiLz4KPC9zdmc+", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjNWZhODNiIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjJjMTlmIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTBhODc3Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2YyYzE5ZiIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTBhODc3IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzQgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NCBDOTIgNTYgODIgNDYgNjQgNDYgQzQ2IDQ2IDM2IDU2IDMyIDc0IFoiIGZpbGw9IiM1YzNhMjEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8ZyBmaWxsPSJub25lIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi42Ij4KICAgIDxjaXJjbGUgY3g9IjUxIiBjeT0iNzgiIHI9IjkiLz48Y2lyY2xlIGN4PSI3NyIgY3k9Ijc4IiByPSI5Ii8+CiAgICA8cGF0aCBkPSJNNjAgNzggTDY4IDc4Ii8+PHBhdGggZD0iTTQyIDc4IEwzNiA3NiIvPjxwYXRoIGQ9Ik04NiA3OCBMOTIgNzYiLz4KICA8L2c+PHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNMzQgMTE4IEwzNCAxMDAgTDUwIDExMiBMNjQgMTA0IEw3OCAxMTIgTDk0IDEwMCBMOTQgMTE4IFoiIGZpbGw9IiMzYTZlYTgiLz48cGF0aCBkPSJNNjAgMTA2IEw2OCAxMDYgTDY2IDEyOCBMNjQgMTMyIEw2MiAxMjggWiIgZmlsbD0iIzJiMjExOCIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjYzk0ZjRmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZThiNDhhIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYzg5NDZhIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2U4YjQ4YSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYzg5NDZhIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNzYgQzMwIDQ2IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NiA5NiA3NiBMOTAgNzAgQzg2IDU0IDc4IDQ2IDY0IDQ2IEM1MCA0NiA0MiA1NCAzOCA3MCBaIiBmaWxsPSIjN2E1MjMwIi8+CiAgPHBhdGggZD0iTTkyIDYwIEMxMDQgNjIgMTA4IDc4IDEwMCA5NiBDOTggODQgOTIgNzYgODggNjggWiIgZmlsbD0iIzdhNTIzMCIvPjxjaXJjbGUgY3g9IjQ2IiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNTUiLz48Y2lyY2xlIGN4PSI4MiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjUiLz4KICA8ZWxsaXBzZSBjeD0iNTEiIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxlbGxpcHNlIGN4PSI3NyIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPHBhdGggZD0iTTQ1IDcwIFE1MSA2NyA1NyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CiAgPHBhdGggZD0iTTcxIDcwIFE3NyA2NyA4MyA3MCIgc3Ryb2tlPSIjMmIyMTE4IiBzdHJva2Utd2lkdGg9IjIuNCIgZmlsbD0ibm9uZSIgc3Ryb2tlLWxpbmVjYXA9InJvdW5kIi8+CjxwYXRoIGQ9Ik01MyA5NCBRNjQgMTAzIDc1IDk0IFE2NCA5OSA1MyA5NCBaIiBmaWxsPSIjYTg0MDJmIi8+PHBhdGggZD0iTTU0IDk4IEM1NCAxMDggNTggMTE0IDY0IDExNCBDNzAgMTE0IDc0IDEwOCA3NCA5OCBMNzAgMTAwIEM2OCAxMDYgNjYgMTA4IDY0IDEwOCBDNjIgMTA4IDYwIDEwNiA1OCAxMDAgWiIgZmlsbD0iIzdhNTIzMCIvPjxwYXRoIGQ9Ik0zMCAxMTggTDM2IDk4IEw1MCAxMDggTDY0IDEwMCBMNzggMTA4IEw5MiA5OCBMOTggMTE4IFoiIGZpbGw9IiM1ZmE4M2IiLz4KICA8cGF0aCBkPSJNNTYgMTA0IEw2NCAxMDAgTDcyIDEwNCBMNjggMTE4IEw2MCAxMTggWiIgZmlsbD0iI2U4ZTRkYSIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjM2E2ZWE4Ii8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjYzk4YTVlIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjYTg2ZTQ2Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2M5OGE1ZSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjYTg2ZTQ2IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8ZWxsaXBzZSBjeD0iNjQiIGN5PSI1OCIgcng9IjQiIHJ5PSIyIiBmaWxsPSIjMmIyMTE4IiBvcGFjaXR5PSIwIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNMzQgMTE4IEwzNCAxMDIgQzQyIDk2IDU0IDEwMCA2NCAxMDggQzc0IDEwMCA4NiA5NiA5NCAxMDIgTDk0IDExOCBaIiBmaWxsPSIjOGU1ZmE4Ii8+Cjwvc3ZnPg==", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjOGU1ZmE4Ii8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjOGQ1YTNjIi8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjNmY0NDI5Ii8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iIzhkNWEzYyIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjNmY0NDI5IiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzIgNjggQzMwIDQ0IDQ0IDMyIDY0IDMyIEM4NCAzMiA5OCA0NCA5NiA2OCBDOTAgNTggODIgNTQgNjQgNTQgQzQ2IDU0IDM4IDU4IDMyIDY4IFoiIGZpbGw9IiNjOWExNWEiLz48Y2lyY2xlIGN4PSI0NiIgY3k9IjkwIiByPSI1IiBmaWxsPSIjZjRhNmE2IiBvcGFjaXR5PSIwLjU1Ii8+PGNpcmNsZSBjeD0iODIiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41Ii8+CiAgPGVsbGlwc2UgY3g9IjUxIiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8ZWxsaXBzZSBjeD0iNzciIGN5PSI3OCIgcng9IjMuMiIgcnk9IjMuNiIgZmlsbD0iIzJiMjExOCIvPgogIDxwYXRoIGQ9Ik00NSA3MCBRNTEgNjcgNTcgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgogIDxwYXRoIGQ9Ik03MSA3MCBRNzcgNjcgODMgNzAiIHN0cm9rZT0iIzJiMjExOCIgc3Ryb2tlLXdpZHRoPSIyLjQiIGZpbGw9Im5vbmUiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8ZyBmaWxsPSJub25lIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi42Ij4KICAgIDxjaXJjbGUgY3g9IjUxIiBjeT0iNzgiIHI9IjkiLz48Y2lyY2xlIGN4PSI3NyIgY3k9Ijc4IiByPSI5Ii8+CiAgICA8cGF0aCBkPSJNNjAgNzggTDY4IDc4Ii8+PHBhdGggZD0iTTQyIDc4IEwzNiA3NiIvPjxwYXRoIGQ9Ik04NiA3OCBMOTIgNzYiLz4KICA8L2c+PHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNNDIgODggQzQyIDEwNCA1MiAxMTYgNjQgMTE2IEM3NiAxMTYgODYgMTA0IDg2IDg4IEw4NiA5NiBDODIgMTA4IDc0IDExNCA2NCAxMTQgQzU0IDExNCA0NiAxMDggNDIgOTYgWiIgZmlsbD0iI2M5YTE1YSIgb3BhY2l0eT0iMC45MiIvPjxwYXRoIGQ9Ik0zNCAxMTggTDM0IDEwMCBMNTAgMTEyIEw2NCAxMDQgTDc4IDExMiBMOTQgMTAwIEw5NCAxMTggWiIgZmlsbD0iI2Q5N2IzZiIvPgo8L3N2Zz4=", "data:image/svg+xml;base64,PHN2ZyB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHdpZHRoPSIxMjgiIGhlaWdodD0iMTI4IiB2aWV3Qm94PSIwIDAgMTI4IDEyOCI+CjxjaXJjbGUgY3g9IjY0IiBjeT0iNjQiIHI9IjY0IiBmaWxsPSIjMmY4ZjZmIi8+CgogIDxwYXRoIGQ9Ik02NCAxMTggQzQwIDExOCAzNCAxMDYgMzQgOTYgTDM0IDg0IEMzNCA2MCA0NiA0OCA2NCA0OCBDODIgNDggOTQgNjAgOTQgODQgTDk0IDk2IEM5NCAxMDYgODggMTE4IDY0IDExOCBaIiBmaWxsPSIjZjdkM2I1Ii8+CiAgPHBhdGggZD0iTTY0IDQ4IEw2NCAxMTggQzg4IDExOCA5NCAxMDYgOTQgOTYgTDk0IDg0IEM5NCA2MCA4MiA0OCA2NCA0OCBaIiBmaWxsPSIjZTNiNDhmIi8+CiAgPHBhdGggZD0iTTQ4IDc2IEM0NSA3NiA0MyA3OSA0MyA4MyBDNDMgODcgNDUgOTAgNDggOTAiIGZpbGw9Im5vbmUiIHN0cm9rZT0iI2Y3ZDNiNSIgc3Ryb2tlLXdpZHRoPSI0IiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNODAgNzYgQzgzIDc2IDg1IDc5IDg1IDgzIEM4NSA4NyA4MyA5MCA4MCA5MCIgZmlsbD0ibm9uZSIgc3Ryb2tlPSIjZTNiNDhmIiBzdHJva2Utd2lkdGg9IjQiIHN0cm9rZS1saW5lY2FwPSJyb3VuZCIvPgo8cGF0aCBkPSJNMzAgMTAwIEMyNCA3MCAzMCA0MiA2NCAzMiBDOTggNDIgMTA0IDcwIDk4IDEwMCBMOTAgOTAgQzkyIDY4IDg4IDUyIDY0IDQ2IEM0MCA1MiAzNiA2OCAzOCA5MCBaIiBmaWxsPSIjNGEyZjFjIi8+PGNpcmNsZSBjeD0iNDYiIGN5PSI5MCIgcj0iNSIgZmlsbD0iI2Y0YTZhNiIgb3BhY2l0eT0iMC41NSIvPjxjaXJjbGUgY3g9IjgyIiBjeT0iOTAiIHI9IjUiIGZpbGw9IiNmNGE2YTYiIG9wYWNpdHk9IjAuNSIvPgogIDxlbGxpcHNlIGN4PSI1MSIgY3k9Ijc4IiByeD0iMy4yIiByeT0iMy42IiBmaWxsPSIjMmIyMTE4Ii8+CiAgPGVsbGlwc2UgY3g9Ijc3IiBjeT0iNzgiIHJ4PSIzLjIiIHJ5PSIzLjYiIGZpbGw9IiMyYjIxMTgiLz4KICA8cGF0aCBkPSJNNDUgNzAgUTUxIDY3IDU3IDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KICA8cGF0aCBkPSJNNzEgNzAgUTc3IDY3IDgzIDcwIiBzdHJva2U9IiMyYjIxMTgiIHN0cm9rZS13aWR0aD0iMi40IiBmaWxsPSJub25lIiBzdHJva2UtbGluZWNhcD0icm91bmQiLz4KPHBhdGggZD0iTTUzIDk0IFE2NCAxMDMgNzUgOTQgUTY0IDk5IDUzIDk0IFoiIGZpbGw9IiNhODQwMmYiLz48cGF0aCBkPSJNMzAgMTE4IEwzNiA5OCBMNTAgMTA4IEw2NCAxMDAgTDc4IDEwOCBMOTIgOTggTDk4IDExOCBaIiBmaWxsPSIjMmY4ZjZmIi8+CiAgPHBhdGggZD0iTTU2IDEwNCBMNjQgMTAwIEw3MiAxMDQgTDY4IDExOCBMNjAgMTE4IFoiIGZpbGw9IiNlOGU0ZGEiLz4KPC9zdmc+"];

function getUserAvatar() { return localStorage.getItem('mythic_user_avatar') || ''; }
function setUserAvatar(dataUrl) {
  if (dataUrl) localStorage.setItem('mythic_user_avatar', dataUrl);
  else localStorage.removeItem('mythic_user_avatar');
  if (typeof _renderSidebarProfile === 'function') _renderSidebarProfile();
  _renderAvatarPreview();
  _renderAvatarPresetGrid();
}
function _renderAvatarPreview() {
  const preview = document.getElementById('avatar-upload-preview');
  if (!preview) return;
  const photo = getUserAvatar();
  if (photo) {
    preview.innerHTML = '<img src="' + photo + '" alt="Profile photo">';
  } else {
    const name = getUserName();
    const initial = (name || 'M').trim().charAt(0).toUpperCase() || 'M';
    preview.textContent = initial;
  }
}
// Resize/compress the picked image client-side before storing it (keeps
// localStorage small and avoids giant base64 blobs from full-res photos).
function _resizeImageFile(file, maxSize) {
  return new Promise((resolve, reject) => {
    const reader = new FileReader();
    reader.onerror = () => reject(new Error('read failed'));
    reader.onload = () => {
      const img = new Image();
      img.onerror = () => reject(new Error('decode failed'));
      img.onload = () => {
        let { width, height } = img;
        if (width > height && width > maxSize) { height = Math.round(height * maxSize / width); width = maxSize; }
        else if (height > maxSize) { width = Math.round(width * maxSize / height); height = maxSize; }
        const canvas = document.createElement('canvas');
        canvas.width = width; canvas.height = height;
        canvas.getContext('2d').drawImage(img, 0, 0, width, height);
        resolve(canvas.toDataURL('image/jpeg', 0.85));
      };
      img.src = reader.result;
    };
    reader.readAsDataURL(file);
  });
}
function _renderAvatarPresetGrid() {
  const grid = document.getElementById('avatar-preset-grid');
  if (!grid) return;
  if (!grid.dataset.built) {
    grid.innerHTML = PRESET_AVATARS.map((src, i) =>
      '<img src="' + src + '" data-avatar-idx="' + i + '" alt="Mythic avatar ' + (i + 1) + '">'
    ).join('');
    grid.addEventListener('click', (e) => {
      const img = e.target.closest('img[data-avatar-idx]');
      if (!img) return;
      setUserAvatar(img.src);
    });
    grid.dataset.built = '1';
  }
  const current = getUserAvatar();
  grid.querySelectorAll('img[data-avatar-idx]').forEach(img => {
    img.classList.toggle('selected', img.src === current);
  });
}
const avatarFileInput  = document.getElementById('avatar-file-input');
const avatarRemoveBtn  = document.getElementById('avatar-remove-btn');
if (avatarFileInput) avatarFileInput.addEventListener('change', async (e) => {
  const file = e.target.files && e.target.files[0];
  if (!file) return;
  if (!file.type.startsWith('image/')) { alert('Please pick an image file'); return; }
  try {
    const dataUrl = await _resizeImageFile(file, 256);
    setUserAvatar(dataUrl);
  } catch (err) {
    console.warn('[avatar] resize failed:', err);
    alert('Could not load that image');
  }
  avatarFileInput.value = '';
});
if (avatarRemoveBtn) avatarRemoveBtn.addEventListener('click', () => setUserAvatar(''));
function openNameModal() {
  nameInput.value = getUserName();
  _renderAvatarPreview();
  _renderAvatarPresetGrid();
  nameModalOverlay.classList.add('show');
  setTimeout(() => nameInput.focus(), 50);
}
function closeNameModal() { nameModalOverlay.classList.remove('show'); }
nameBtn.addEventListener('click', openNameModal);
nameCancelBtn.addEventListener('click', closeNameModal);
nameModalOverlay.addEventListener('click', (e) => { if (e.target === nameModalOverlay) closeNameModal(); });
nameSaveBtn.addEventListener('click', () => {
  setUserName(nameInput.value.trim());
  closeNameModal();
});
nameInput.addEventListener('keydown', (e) => {
  if (e.key === 'Enter') { e.preventDefault(); nameSaveBtn.click(); }
  else if (e.key === 'Escape') closeNameModal();
});
if (!localStorage.getItem('mythic_name_prompted')) {
  localStorage.setItem('mythic_name_prompted', '1');
  setTimeout(openNameModal, 600);
}

// ── Sidebar profile row (bottom of sidebar) ─────────────────────────────
// Source of truth: getUserName() / mythic_user_name and getUserAvatar() /
// mythic_user_avatar in localStorage — same values shown in the profile modal.
const sidebarProfileBtn    = document.getElementById('sidebar-profile');
const sidebarProfileAvatar = document.getElementById('sidebar-profile-avatar');
const sidebarProfileName   = document.getElementById('sidebar-profile-name');

function _renderSidebarProfile() {
  if (!sidebarProfileAvatar || !sidebarProfileName) return;
  const name = getUserName();
  const display = name || 'Guest';
  sidebarProfileName.textContent = display;
  const photo = getUserAvatar();
  if (photo) {
    sidebarProfileAvatar.innerHTML = '<img src="' + photo + '" alt="">';
  } else {
    const initial = (name || 'M').trim().charAt(0).toUpperCase() || 'M';
    sidebarProfileAvatar.textContent = initial;
  }
  sidebarProfileAvatar.setAttribute('role', 'img');
  sidebarProfileAvatar.setAttribute('aria-label', display + "'s avatar");

  // Show the actual signed-in account email under the name (falls back to
  // the existing "Your profile" label until this resolves, so there's no
  // layout jump / blank state on first paint).
  const subEl = document.getElementById('sidebar-profile-sub');
  fetch('/api/auth/me').then(r => r.json()).then(info => {
    if (!info.authenticated) return;
    if (subEl && info.email) subEl.textContent = info.email;
    if (info.name && !name) sidebarProfileName.textContent = info.name;
    if (info.picture && !photo) sidebarProfileAvatar.innerHTML = '<img src="' + info.picture + '" alt="">';
  }).catch(() => {});
}
if (sidebarProfileBtn) sidebarProfileBtn.addEventListener('click', openNameModal);
_renderSidebarProfile();

const logoutBtn = document.getElementById('logout-btn');
if (logoutBtn) {
  logoutBtn.addEventListener('click', async () => {
    try { await fetch('/api/auth/logout', { method: 'POST' }); } catch (e) {}
    window.location.href = '/login';
  });
}

if (isMobile()) sidebar.classList.add('hidden');
newChatBtn.addEventListener('click', startNewChat);
clearBtn.addEventListener('click', async () => {
  if (!activeConvId) return;
  await fetch('/api/conversations/' + activeConvId, { method: 'DELETE' });
  startNewChat();
});

exportBtn.addEventListener('click', async () => {
  if (!activeConvId) { alert('Start or open a chat first.'); return; }
  try {
    const r = await fetch('/api/conversations/' + activeConvId);
    if (!r.ok) return;
    const d = await r.json();
    const lines = [`# ${d.title || 'Mythic AI chat'}`, ''];
    (d.messages || []).forEach(m => {
      lines.push(m.role === 'user' ? 'You:' : 'Mythic AI:');
      lines.push(m.text || (m.attachment ? `[attachment: ${m.attachment.name}]` : ''));
      lines.push('');
    });
    const blob = new Blob([lines.join('\n')], { type: 'text/plain;charset=utf-8' });
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url;
    a.download = (d.title || 'chat').replace(/[^a-z0-9_ -]/gi, '').trim().slice(0, 60) + '.txt';
    document.body.appendChild(a);
    a.click();
    a.remove();
    URL.revokeObjectURL(url);
  } catch (err) {
    alert('Export failed: ' + err.message);
  }
});

// ─── Share link (per-chat, e.g. mythic-ai.app/share/abc123) ────────────────
const shareBtn          = document.getElementById('share-btn');
const shareModalOverlay = document.getElementById('share-modal-overlay');
const shareLinkInput    = document.getElementById('share-link-input');
const shareOpenBtn      = document.getElementById('share-open-btn');
const shareCopyBtn      = document.getElementById('share-copy-btn');
const shareNativeBtn    = document.getElementById('share-native-btn');
const shareRevokeBtn    = document.getElementById('share-revoke-btn');
const shareCloseBtn     = document.getElementById('share-close-btn');
const shareStatusEl     = document.getElementById('share-status');

function closeShareModal() { shareModalOverlay.classList.remove('show'); }

// One single, permanent link for the whole app/account — not one per chat.
// No login, no server round-trip, no "start a chat first" requirement:
// anyone who opens this URL lands straight on the chat screen and gets
// their own private, anonymous conversation history (see current_username()
// server-side). This is just the site's own root URL.
function openInviteModal() {
  shareLinkInput.value = '';
  shareLinkInput.title = '';
  shareStatusEl.textContent = 'Loading your account link…';
  shareModalOverlay.classList.add('show');
  shareBtn.classList.add('active');
  if (shareRevokeBtn) shareRevokeBtn.style.display = 'none';  // nothing to revoke — it's a static link
  fetch('/api/invite-link').then(r => r.json()).then(d => {
    const link = d.invite_url || (location.origin + '/');
    shareLinkInput.value = link;
    shareLinkInput.title = link;
    shareStatusEl.textContent = 'This link is unique to THIS account — opening it always ' +
      'comes back to these exact chats, on any device. It does not open a fresh/different account.';
    requestAnimationFrame(() => { shareLinkInput.focus(); shareLinkInput.select(); });
    renderInviteQrCode(link);
  }).catch(() => {
    shareLinkInput.value = location.origin + '/';
    shareStatusEl.textContent = 'Could not generate your account link, showing the site link instead.';
    renderInviteQrCode(location.origin + '/');
  });
}

// High error-correction (level H) tolerates ~30% obscured area, which is
// what makes the logo sit safely in the middle without breaking the scan.
function renderInviteQrCode(link) {
  const box = document.getElementById('share-qr-canvas');
  if (!box || typeof QRCode === 'undefined') return;
  box.innerHTML = '';
  new QRCode(box, {
    text: link,
    width: 196,
    height: 196,
    colorDark: '#000000',
    colorLight: '#ffffff',
    correctLevel: QRCode.CorrectLevel.H,
  });
}

if (shareBtn) shareBtn.addEventListener('click', openInviteModal);
if (shareCloseBtn) shareCloseBtn.addEventListener('click', closeShareModal);
if (shareModalOverlay) shareModalOverlay.addEventListener('click', e => { if (e.target === shareModalOverlay) closeShareModal(); });
if (shareLinkInput) shareLinkInput.addEventListener('click', () => shareLinkInput.select());
if (shareOpenBtn) shareOpenBtn.addEventListener('click', () => {
  if (shareLinkInput.value) window.open(shareLinkInput.value, '_blank', 'noopener');
});

if (shareCopyBtn) shareCopyBtn.addEventListener('click', async () => {
  if (!shareLinkInput.value) return;
  try {
    await navigator.clipboard.writeText(shareLinkInput.value);
  } catch {
    shareLinkInput.select();
    try { document.execCommand('copy'); } catch {}
  }
  const orig = shareCopyBtn.textContent;
  shareCopyBtn.textContent = '✓ Copied';
  setTimeout(() => { shareCopyBtn.textContent = orig; }, 1400);
});

if (shareNativeBtn) shareNativeBtn.addEventListener('click', async () => {
  if (!shareLinkInput.value) return;
  if (navigator.share) {
    try { await navigator.share({ title: 'Mythic AI chat', url: shareLinkInput.value }); return; }
    catch (e) { if (e && e.name === 'AbortError') return; }
  }
  shareCopyBtn.click();
});

// Builds a single flattened PNG of the QR box (white background + QR +
// centered logo) by compositing onto an offscreen canvas — this is what
// both "Share QR Code" and "Download QR" actually send/save, so what you
// see in the modal is exactly what gets shared.
async function buildQrImageBlob() {
  const qrBox = document.getElementById('share-qr-box');
  const qrCanvasEl = document.querySelector('#share-qr-canvas canvas') ||
                      document.querySelector('#share-qr-canvas img');
  if (!qrBox || !qrCanvasEl) return null;

  const size = 320; // export at higher res than the on-screen 220px box
  const canvas = document.createElement('canvas');
  canvas.width = size; canvas.height = size;
  const ctx = canvas.getContext('2d');
  ctx.fillStyle = '#ffffff';
  ctx.fillRect(0, 0, size, size);

  const pad = Math.round(size * 12 / 220);
  const qrSize = size - pad * 2;
  const qrSource = qrCanvasEl.tagName === 'CANVAS' ? qrCanvasEl : qrCanvasEl;
  await new Promise((resolve) => {
    if (qrCanvasEl.tagName === 'CANVAS') { ctx.drawImage(qrCanvasEl, pad, pad, qrSize, qrSize); resolve(); }
    else {
      const img = new Image();
      img.crossOrigin = 'anonymous';
      img.onload = () => { ctx.drawImage(img, pad, pad, qrSize, qrSize); resolve(); };
      img.onerror = resolve;
      img.src = qrCanvasEl.src;
    }
  });

  // Composite the center logo (white rounded square + icon), matching the
  // on-screen #share-qr-logo overlay.
  const logoSize = Math.round(size * 52 / 220);
  const logoX = (size - logoSize) / 2, logoY = (size - logoSize) / 2;
  const logoImg = document.querySelector('#share-qr-logo img');
  await new Promise((resolve) => {
    ctx.fillStyle = '#ffffff';
    const r = 10;
    ctx.beginPath();
    ctx.moveTo(logoX + r, logoY);
    ctx.arcTo(logoX + logoSize, logoY, logoX + logoSize, logoY + logoSize, r);
    ctx.arcTo(logoX + logoSize, logoY + logoSize, logoX, logoY + logoSize, r);
    ctx.arcTo(logoX, logoY + logoSize, logoX, logoY, r);
    ctx.arcTo(logoX, logoY, logoX + logoSize, logoY, r);
    ctx.closePath();
    ctx.fill();
    if (!logoImg) { resolve(); return; }
    const img = new Image();
    img.crossOrigin = 'anonymous';
    img.onload = () => {
      const inset = Math.round(logoSize * 0.12);
      ctx.drawImage(img, logoX + inset, logoY + inset, logoSize - inset * 2, logoSize - inset * 2);
      resolve();
    };
    img.onerror = resolve;
    img.src = logoImg.src;
  });

  return new Promise((resolve) => canvas.toBlob(resolve, 'image/png'));
}

const shareQrImageBtn = document.getElementById('share-qr-image-btn');
const downloadQrBtn   = document.getElementById('download-qr-btn');

if (shareQrImageBtn) shareQrImageBtn.addEventListener('click', async () => {
  const orig = shareQrImageBtn.textContent;
  shareQrImageBtn.textContent = 'Preparing…';
  try {
    const blob = await buildQrImageBlob();
    if (!blob) { shareQrImageBtn.textContent = orig; return; }
    const file = new File([blob], 'mythic-ai-qr.png', { type: 'image/png' });
    if (navigator.canShare && navigator.canShare({ files: [file] })) {
      await navigator.share({ files: [file], title: 'Mythic AI — scan to open my chat' });
    } else if (navigator.share) {
      // Some browsers support navigator.share but not file sharing —
      // fall back to sharing the link instead of failing silently.
      await navigator.share({ title: 'Mythic AI chat', url: shareLinkInput.value });
    } else {
      // No native share support at all (most desktop browsers) — download instead.
      const url = URL.createObjectURL(blob);
      const a = document.createElement('a');
      a.href = url; a.download = 'mythic-ai-qr.png';
      document.body.appendChild(a); a.click(); a.remove();
      URL.revokeObjectURL(url);
    }
  } catch (e) {
    if (!(e && e.name === 'AbortError')) console.warn('QR share failed:', e);
  } finally {
    shareQrImageBtn.textContent = orig;
  }
});

if (downloadQrBtn) downloadQrBtn.addEventListener('click', async () => {
  const orig = downloadQrBtn.textContent;
  downloadQrBtn.textContent = 'Preparing…';
  try {
    const blob = await buildQrImageBlob();
    if (!blob) { downloadQrBtn.textContent = orig; return; }
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url; a.download = 'mythic-ai-qr.png';
    document.body.appendChild(a); a.click(); a.remove();
    URL.revokeObjectURL(url);
  } catch (e) {
    console.warn('QR download failed:', e);
  } finally {
    downloadQrBtn.textContent = orig;
  }
});

// shareRevokeBtn is hidden in openInviteModal() — the invite link is static
// and can't be "revoked" (it's just the site's own address). Kept wired to
// closeShareModal() only in case older cached HTML still shows the button.
if (shareRevokeBtn) shareRevokeBtn.addEventListener('click', closeShareModal);

const settingsBtn        = document.getElementById('settings-btn');
// ─── Reusable "clarification" modal ──────────────────────────────────────────
// Call askClarification({ question, options, onPick, onCustom, onSkip }) any
// time the app needs a quick decision from the user instead of guessing.
// - options: array of short strings, shown numbered (1-9), each also
//   triggers via that number key
// - Enter with nothing typed picks option 1 (matches the screenshot's
//   "highlighted first row + return-arrow icon" behavior)
// - typing something and pressing Enter (or clicking outside the list)
//   calls onCustom with that free text instead
// - Skip/✕/Escape calls onSkip if provided
let _clarifyKeyHandler = null;

function askClarification({ question, options, onPick, onCustom, onSkip }) {
  const overlay = document.getElementById('clarify-modal-overlay');
  const questionEl = document.getElementById('clarify-question');
  const optionsEl = document.getElementById('clarify-options');
  const customInput = document.getElementById('clarify-custom-input');
  const closeBtn = document.getElementById('clarify-close-btn');
  const skipBtn = document.getElementById('clarify-skip-btn');

  questionEl.textContent = question || "What's wrong right now?";
  optionsEl.innerHTML = '';
  customInput.value = '';

  options.forEach((opt, i) => {
    const row = document.createElement('div');
    row.style.cssText = 'display:flex;align-items:center;gap:12px;padding:12px 20px;cursor:pointer;font-size:15px;' +
      (i === 0 ? 'background:#f2f2f2;' : '');
    row.innerHTML =
      '<span style="width:22px;height:22px;border-radius:6px;background:' + (i === 0 ? '#e5e5e5' : 'transparent') +
      ';display:flex;align-items:center;justify-content:center;font-size:13px;color:#666;flex-shrink:0;">' + (i + 1) + '</span>' +
      '<span style="flex:1;color:#222;">' + opt.replace(/</g, '&lt;') + '</span>' +
      (i === 0 ? '<span style="color:#aaa;font-size:16px;">⏎</span>' : '');
    row.addEventListener('mouseenter', () => { row.style.background = '#f2f2f2'; });
    row.addEventListener('mouseleave', () => { row.style.background = i === 0 ? '#f2f2f2' : 'transparent'; });
    row.addEventListener('click', () => { closeClarification(); if (onPick) onPick(opt, i); });
    optionsEl.appendChild(row);
  });

  function closeClarification() {
    overlay.style.display = 'none';
    if (_clarifyKeyHandler) { document.removeEventListener('keydown', _clarifyKeyHandler); _clarifyKeyHandler = null; }
  }

  closeBtn.onclick = () => { closeClarification(); if (onSkip) onSkip(); };
  skipBtn.onclick = () => { closeClarification(); if (onSkip) onSkip(); };

  customInput.onkeydown = (e) => {
    if (e.key === 'Enter') {
      const val = customInput.value.trim();
      closeClarification();
      if (val && onCustom) onCustom(val);
      else if (!val && onPick) onPick(options[0], 0);  // bare Enter = first option
    }
  };

  _clarifyKeyHandler = (e) => {
    if (e.key === 'Escape') { closeClarification(); if (onSkip) onSkip(); return; }
    if (document.activeElement === customInput) return;  // let typing work normally
    const n = parseInt(e.key, 10);
    if (n >= 1 && n <= options.length) {
      closeClarification();
      if (onPick) onPick(options[n - 1], n - 1);
    } else if (e.key === 'Enter') {
      closeClarification();
      if (onPick) onPick(options[0], 0);
    }
  };
  document.addEventListener('keydown', _clarifyKeyHandler);

  overlay.style.display = 'flex';
  setTimeout(() => customInput.focus(), 50);
}

const settingsModalOverlay=document.getElementById('settings-modal-overlay');
const settingsCloseBtn   = document.getElementById('settings-close-btn');
const accentColorInput   = document.getElementById('accent-color-input');
const fontSizeSlider     = document.getElementById('font-size-slider');
const fontSizeLabel      = document.getElementById('font-size-label');
const toneSelect         = document.getElementById('tone-select');
const lengthSelect       = document.getElementById('length-select');
const customInstructions = document.getElementById('custom-instructions-input');

function loadSettings() {
  const s = JSON.parse(localStorage.getItem('mythic_settings') || '{}');
  const theme = s.theme || 'dark';
  applyTheme(theme);
  document.querySelectorAll('[data-group="theme"]').forEach(b => {
    b.style.borderColor = b.dataset.value === theme ? 'var(--accent)' : 'var(--border)';
    b.style.color = b.dataset.value === theme ? 'var(--accent)' : '';
  });
  const accent = s.accent || '#10a37f';
  accentColorInput.value = accent;
  document.documentElement.style.setProperty('--accent', accent);
  const fs = s.fontSize || '14.5';
  fontSizeSlider.value = fs;
  fontSizeLabel.textContent = fs + 'px';
  document.documentElement.style.setProperty('--msg-font-size', fs + 'px');
  const bubble = s.bubble || 'comfortable';
  document.body.classList.remove('bubble-compact','bubble-comfortable','bubble-spacious');
  document.body.classList.add('bubble-' + bubble);
  document.querySelectorAll('[data-group="bubble"]').forEach(b => {
    b.style.borderColor = b.dataset.value === bubble ? 'var(--accent)' : 'var(--border)';
    b.style.color = b.dataset.value === bubble ? 'var(--accent)' : '';
  });
  if (toneSelect) toneSelect.value = s.tone || 'default';
  if (lengthSelect) lengthSelect.value = s.length || 'default';
  if (customInstructions) customInstructions.value = s.customInstructions || '';
}

function saveSettings() {
  const s = JSON.parse(localStorage.getItem('mythic_settings') || '{}');
  s.theme = document.body.classList.contains('theme-light') ? 'light' : 'dark';
  s.accent = accentColorInput.value;
  s.fontSize = fontSizeSlider.value;
  const bubs = ['compact','comfortable','spacious'].find(b => document.body.classList.contains('bubble-'+b)) || 'comfortable';
  s.bubble = bubs;
  s.tone = toneSelect ? toneSelect.value : 'default';
  s.length = lengthSelect ? lengthSelect.value : 'default';
  s.customInstructions = customInstructions ? customInstructions.value : '';
  localStorage.setItem('mythic_settings', JSON.stringify(s));
}

// ─── API key management (owner only — settings panel) ───────────────────────
const apiKeyLabelInput = document.getElementById('api-key-label-input');
const apiKeyCreateBtn  = document.getElementById('api-key-create-btn');
const apiKeyNewBox     = document.getElementById('api-key-new-box');
const apiKeyNewValue   = document.getElementById('api-key-new-value');
const apiKeyListEl     = document.getElementById('api-key-list');

async function loadApiKeys() {
  if (!apiKeyListEl) return;
  try {
    const res = await fetch('/api/keys');
    if (res.status === 403) { apiKeyListEl.innerHTML = ''; return; } // not the owner — hide silently
    const data = await res.json();
    const keys = data.keys || [];
    apiKeyListEl.innerHTML = keys.length ? '' : '<div class="hint">No keys yet.</div>';
    keys.forEach(k => {
      const row = document.createElement('div');
      row.style.cssText = 'display:flex;align-items:center;justify-content:space-between;gap:8px;padding:6px 8px;border:1px solid var(--border);border-radius:6px;font-size:13px;flex-wrap:wrap;';
      const used = k.request_count ? `${k.request_count} calls` : 'unused';
      const labelSpan = document.createElement('span');
      labelSpan.innerHTML = `${k.active ? '🟢' : '⚪'} <code>${k.key_prefix}</code> ${k.label ? '— ' + k.label.replace(/</g,'&lt;') : ''} <span style="opacity:.6;">(${used})</span>`;
      row.appendChild(labelSpan);

      const btnGroup = document.createElement('div');
      btnGroup.style.cssText = 'display:flex;gap:6px;';

      const copyBtn = document.createElement('button');
      copyBtn.type = 'button';
      copyBtn.textContent = '📋 Copy';
      copyBtn.className = 'settings-btn';
      copyBtn.onclick = () => {
        navigator.clipboard.writeText(k.key_prefix || '').then(() => {
          const orig = copyBtn.textContent;
          copyBtn.textContent = '✓ Copied';
          setTimeout(() => { copyBtn.textContent = orig; }, 1200);
        });
      };
      btnGroup.appendChild(copyBtn);

      const renameBtn = document.createElement('button');
      renameBtn.type = 'button';
      renameBtn.textContent = '✎ Rename';
      renameBtn.className = 'settings-btn';
      renameBtn.onclick = async () => {
        const newLabel = prompt('New name for this key:', k.label || '');
        if (newLabel === null) return;
        const trimmed = newLabel.trim();
        if (!trimmed) { alert('Name cannot be empty.'); return; }
        const r = await fetch('/api/keys/' + k.id, {
          method: 'PATCH', headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ label: trimmed }),
        });
        const d = await r.json();
        if (d.error) { alert(d.error); return; }
        loadApiKeys();
        loadApiUsageSummary();
      };
      btnGroup.appendChild(renameBtn);

      if (k.active) {
        const revokeBtn = document.createElement('button');
        revokeBtn.type = 'button';
        revokeBtn.textContent = 'Revoke';
        revokeBtn.className = 'settings-btn';
        revokeBtn.onclick = async () => {
          if (!confirm('Revoke this key? Apps using it will stop working immediately.')) return;
          await fetch('/api/keys/' + k.id, { method: 'DELETE' });
          loadApiKeys();
          loadApiUsageSummary();
        };
        btnGroup.appendChild(revokeBtn);
      }
      row.appendChild(btnGroup);
      apiKeyListEl.appendChild(row);
    });
  } catch (e) {
    console.warn('Could not load API keys:', e);
  }
}

// ─── API usage summary (shown right in the sidebar, under New Chat) ─────────
const apiUsageSummaryEl = document.getElementById('api-usage-summary');

async function loadApiUsageSummary() {
  if (!apiUsageSummaryEl) return;
  try {
    const res = await fetch('/api/keys');
    if (res.status === 403) { apiUsageSummaryEl.style.display = 'none'; return; } // not the owner
    const data = await res.json();
    const keys = data.keys || [];
    if (!keys.length) { apiUsageSummaryEl.style.display = 'none'; return; }

    const activeKeys = keys.filter(k => k.active);
    const totalCalls = keys.reduce((sum, k) => sum + (k.request_count || 0), 0);
    const lastUsedTimes = keys.map(k => k.last_used_at).filter(Boolean).sort();
    const lastUsed = lastUsedTimes.length ? lastUsedTimes[lastUsedTimes.length - 1] : null;
    const lastUsedStr = lastUsed
      ? new Date(lastUsed).toLocaleString(undefined, { month: 'short', day: 'numeric', hour: 'numeric', minute: '2-digit' })
      : 'never';

    apiUsageSummaryEl.style.display = 'block';
    apiUsageSummaryEl.innerHTML =
      `<div style="display:flex;justify-content:space-between;"><span>🔑 ${activeKeys.length} active key${activeKeys.length === 1 ? '' : 's'}</span>` +
      `<span>${totalCalls} call${totalCalls === 1 ? '' : 's'}</span></div>` +
      `<div style="opacity:.75;margin-top:2px;">Last used: ${lastUsedStr}</div>`;
  } catch (e) {
    console.warn('Could not load API usage summary:', e);
  }
}

if (apiUsageSummaryEl) {
  apiUsageSummaryEl.addEventListener('click', () => {
    openApiUsageOverlay();
  });
  loadApiUsageSummary();
}

// ─── API usage full dashboard (in-page overlay, no page navigation) ─────────
const apiUsageOverlayEl = document.getElementById('api-usage-overlay');
const apiUsageTotalsEl  = document.getElementById('api-usage-totals');
const apiUsageTbodyEl   = document.getElementById('api-usage-tbody');
const apiUsageEmptyEl   = document.getElementById('api-usage-empty');
const apiUsageCloseBtn  = document.getElementById('api-usage-close-btn');
const apiUsageGenBtn    = document.getElementById('api-usage-gen-btn');
const apiUsageCreateOverlayEl = document.getElementById('api-usage-create-overlay');
const apiUsageCreateLabelEl   = document.getElementById('api-usage-create-label');
const apiUsageNewKeyResultEl  = document.getElementById('api-usage-new-key-result');
const apiUsageCreateCloseBtn   = document.getElementById('api-usage-create-close-btn');
const apiUsageCreateConfirmBtn = document.getElementById('api-usage-create-confirm-btn');

function fmtApiUsageDate(iso) {
  if (!iso) return 'Never';
  const d = new Date(iso);
  if (isNaN(d)) return iso;
  return d.toLocaleString(undefined, { month: 'short', day: 'numeric', year: 'numeric', hour: 'numeric', minute: '2-digit' });
}

async function refreshApiUsageOverlay() {
  if (!apiUsageOverlayEl) return;
  try {
    const res = await fetch('/api/keys');
    if (res.status === 403) {
      apiUsageTotalsEl.innerHTML = '';
      apiUsageTbodyEl.innerHTML = '';
      apiUsageEmptyEl.style.display = 'block';
      apiUsageEmptyEl.textContent = 'Only the account owner can view API keys.';
      return;
    }
    const data = await res.json();
    const keys = data.keys || [];
    const activeCount = keys.filter(k => k.active).length;
    const totalCalls = keys.reduce((s, k) => s + (k.request_count || 0), 0);

    apiUsageTotalsEl.innerHTML = `
      <div style="flex:1;min-width:150px;background:#1a1d24;border:1px solid #2a2e37;border-radius:14px;padding:18px;text-align:center;">
        <div style="font-size:40px;font-weight:800;line-height:1.1;">${activeCount}</div>
        <div style="font-size:12px;color:#9a9ea6;margin-top:6px;letter-spacing:.3px;text-transform:uppercase;">Active Keys</div>
      </div>
      <div style="flex:1;min-width:150px;background:#1a1d24;border:1px solid #2a2e37;border-radius:14px;padding:18px;text-align:center;">
        <div style="font-size:40px;font-weight:800;line-height:1.1;">${totalCalls}</div>
        <div style="font-size:12px;color:#9a9ea6;margin-top:6px;letter-spacing:.3px;text-transform:uppercase;">Total Calls</div>
      </div>
      <div style="flex:1;min-width:150px;background:#1a1d24;border:1px solid #2a2e37;border-radius:14px;padding:18px;text-align:center;">
        <div style="font-size:40px;font-weight:800;line-height:1.1;">${keys.length}</div>
        <div style="font-size:12px;color:#9a9ea6;margin-top:6px;letter-spacing:.3px;text-transform:uppercase;">Total Keys</div>
      </div>`;

    if (!keys.length) {
      apiUsageTbodyEl.innerHTML = '';
      apiUsageEmptyEl.style.display = 'block';
      apiUsageEmptyEl.textContent = 'No API keys yet. Click "Generate API Key" to create one.';
      return;
    }
    apiUsageEmptyEl.style.display = 'none';

    apiUsageTbodyEl.innerHTML = keys.map(k => `
      <tr>
        <td style="padding:16px;border-bottom:1px solid #22252c;font-size:14px;font-weight:700;">${(k.label || '(unnamed key)').replace(/</g,'&lt;')}</td>
        <td style="padding:16px;border-bottom:1px solid #22252c;font-size:14px;font-family:monospace;color:#c7cad1;">
          <div style="display:flex;align-items:center;gap:8px;"><span>${k.key_prefix || ''}</span>
          <button class="api-usage-copy-btn" data-prefix="${(k.key_prefix||'').replace(/"/g,'&quot;')}" style="background:none;border:1px solid #3a3e47;color:#c7cad1;cursor:pointer;font-size:11.5px;padding:4px 9px;border-radius:6px;font-weight:600;font-family:inherit;white-space:nowrap;">📋 Copy</button></div>
        </td>
        <td style="padding:16px;border-bottom:1px solid #22252c;font-size:14px;">${fmtApiUsageDate(k.created_at)}</td>
        <td style="padding:16px;border-bottom:1px solid #22252c;font-size:18px;font-weight:700;">${k.request_count || 0}</td>
        <td style="padding:16px;border-bottom:1px solid #22252c;font-size:14px;">
          <span style="font-size:11px;font-weight:700;letter-spacing:.4px;border:1px solid;border-radius:20px;padding:3px 10px;color:${k.active ? '#1a9e5c' : '#c0392b'};border-color:${k.active ? '#1a9e5c' : '#c0392b'};">${k.active ? 'ACTIVE' : 'REVOKED'}</span>
        </td>
        <td style="padding:16px;border-bottom:1px solid #22252c;font-size:14px;">
          <div style="display:flex;gap:6px;flex-wrap:wrap;">
            <button class="api-usage-rename-btn" data-id="${k.id}" data-label="${(k.label||'').replace(/"/g,'&quot;')}" style="background:none;border:1px solid #3a3e47;color:#9a9ea6;border-radius:6px;padding:6px 10px;font-size:12px;cursor:pointer;font-family:inherit;">✎ Rename</button>
            ${k.active ? `<button class="api-usage-revoke-btn" data-id="${k.id}" style="background:none;border:1px solid #3a3e47;color:#c0392b;border-radius:6px;padding:6px 10px;font-size:12px;cursor:pointer;font-family:inherit;">Revoke</button>` : ''}
          </div>
        </td>
      </tr>`).join('');

    apiUsageTbodyEl.querySelectorAll('.api-usage-revoke-btn').forEach(btn => {
      btn.addEventListener('click', async () => {
        if (!confirm('Revoke this key? Apps using it will stop working immediately.')) return;
        await fetch('/api/keys/' + btn.dataset.id, { method: 'DELETE' });
        refreshApiUsageOverlay();
        loadApiUsageSummary();
      });
    });
    apiUsageTbodyEl.querySelectorAll('.api-usage-copy-btn').forEach(btn => {
      btn.addEventListener('click', () => {
        navigator.clipboard.writeText(btn.dataset.prefix || '').then(() => {
          const orig = btn.textContent;
          btn.textContent = '✓ Copied';
          setTimeout(() => { btn.textContent = orig; }, 1200);
        });
      });
    });
    apiUsageTbodyEl.querySelectorAll('.api-usage-rename-btn').forEach(btn => {
      btn.addEventListener('click', async () => {
        const newLabel = prompt('New name for this key:', btn.dataset.label || '');
        if (newLabel === null) return;
        const trimmed = newLabel.trim();
        if (!trimmed) { alert('Name cannot be empty.'); return; }
        const r = await fetch('/api/keys/' + btn.dataset.id, {
          method: 'PATCH', headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ label: trimmed }),
        });
        const d = await r.json();
        if (d.error) { alert(d.error); return; }
        refreshApiUsageOverlay();
        loadApiUsageSummary();
      });
    });
  } catch (e) {
    console.warn('Could not load API usage:', e);
  }
}

function openApiUsageOverlay() {
  if (!apiUsageOverlayEl) return;
  apiUsageOverlayEl.style.display = 'block';
  refreshApiUsageOverlay();
}
function closeApiUsageOverlay() {
  if (apiUsageOverlayEl) apiUsageOverlayEl.style.display = 'none';
}
if (apiUsageCloseBtn) apiUsageCloseBtn.addEventListener('click', closeApiUsageOverlay);

if (apiUsageGenBtn) apiUsageGenBtn.addEventListener('click', () => {
  apiUsageCreateLabelEl.value = '';
  apiUsageNewKeyResultEl.innerHTML = '';
  apiUsageCreateOverlayEl.style.display = 'flex';
});
if (apiUsageCreateCloseBtn) apiUsageCreateCloseBtn.addEventListener('click', () => {
  apiUsageCreateOverlayEl.style.display = 'none';
  refreshApiUsageOverlay();
  loadApiUsageSummary();
});
if (apiUsageCreateConfirmBtn) apiUsageCreateConfirmBtn.addEventListener('click', async () => {
  const label = apiUsageCreateLabelEl.value.trim();
  apiUsageCreateConfirmBtn.disabled = true;
  try {
    const res = await fetch('/api/keys', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ label }),
    });
    const data = await res.json();
    if (data.api_key) {
      apiUsageNewKeyResultEl.innerHTML =
        '<div style="background:#0f1115;border:1px solid #1a9e5c;border-radius:8px;padding:12px;margin-bottom:14px;">' +
          '<div style="display:flex;gap:8px;align-items:center;">' +
            '<input type="text" id="api-usage-new-key-input" readonly value="' + data.api_key.replace(/"/g, '&quot;') + '" ' +
              'style="flex:1;background:#0f1115;border:1px solid #1a9e5c;border-radius:6px;padding:8px;font-family:monospace;font-size:12px;color:#7be3ab;box-sizing:border-box;min-width:0;">' +
            '<button type="button" id="api-usage-new-key-copy-btn" ' +
              'style="background:#1a9e5c;color:#04140b;border:none;padding:8px 14px;border-radius:6px;cursor:pointer;font-size:12px;font-weight:700;white-space:nowrap;">Copy</button>' +
          '</div>' +
        '</div>' +
        '<div style="font-size:12px;color:#9a9ea6;margin-bottom:10px;">Copy this now — it will not be shown again. (Double-clicking to select may only grab part of the key due to the hyphens — use the Copy button instead.)</div>';
      const newKeyCopyBtn = document.getElementById('api-usage-new-key-copy-btn');
      if (newKeyCopyBtn) {
        newKeyCopyBtn.addEventListener('click', async () => {
          const keyInput = document.getElementById('api-usage-new-key-input');
          if (!keyInput) return;
          try {
            keyInput.select();
            document.execCommand('copy');
            newKeyCopyBtn.textContent = '✓ Copied!';
            setTimeout(() => { newKeyCopyBtn.textContent = 'Copy'; }, 2500);
          } catch (e) {
            try {
              await navigator.clipboard.writeText(keyInput.value);
              newKeyCopyBtn.textContent = '✓ Copied!';
              setTimeout(() => { newKeyCopyBtn.textContent = 'Copy'; }, 2500);
            } catch (e2) {
              alert('Copy failed. Try selecting the whole box manually (click once, then Ctrl+A, Ctrl+C).');
            }
          }
        });
      }
      refreshApiUsageOverlay();
      loadApiUsageSummary();
    } else if (data.error) {
      alert(data.error);
    }
  } catch (e) {
    alert('Could not create key.');
  } finally {
    apiUsageCreateConfirmBtn.disabled = false;
  }
});

if (apiKeyCreateBtn) {
  apiKeyCreateBtn.addEventListener('click', apiKeyCreateHandler);
  // iOS fallback: add touchstart for better compatibility
  if (isIOS) {
    apiKeyCreateBtn.addEventListener('touchstart', apiKeyCreateHandler, { passive: false });
  }
  
  async function apiKeyCreateHandler(e) {
    const label = apiKeyLabelInput ? apiKeyLabelInput.value.trim() : '';
    apiKeyCreateBtn.disabled = true;
    try {
      const res = await fetch('/api/keys', {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ label }),
      });
      
      if (!res.ok) {
        throw new Error(`HTTP ${res.status}: ${res.statusText}`);
      }
      
      const data = await res.json();
      
      if (data.error) {
        alert('Error: ' + (data.error.message || data.error));
        return;
      }
      
      if (data.api_key) {
        apiKeyNewValue.value = data.api_key;
        apiKeyNewBox.style.display = 'block';
        if (apiKeyLabelInput) apiKeyLabelInput.value = '';
        loadApiKeys();
        loadApiUsageSummary();
      } else {
        alert('No API key returned from server');
      }
    } catch (e) {
      console.error('[API Key Creation Error]', e);
      alert('Could not create API key: ' + e.message);
    } finally {
      apiKeyCreateBtn.disabled = false;
    }
  }
}

const apiKeyCopyBtn = document.getElementById('api-key-copy-btn');
if (apiKeyCopyBtn) {
  apiKeyCopyBtn.addEventListener('click', async () => {
    const keyInput = document.getElementById('api-key-new-value');
    const keyValue = keyInput ? keyInput.value : '';
    console.log('=== API Key Copy Debug ===');
    console.log('keyInput element:', keyInput);
    console.log('keyInput.value length:', keyValue.length);
    console.log('keyInput.value (first 50 chars):', keyValue.substring(0, 50));
    console.log('Full keyInput.value:', keyValue);
    
    if (!keyValue || keyValue.length < 20) {
      alert(`Key is incomplete (only ${keyValue.length} chars). Check browser console for details.`);
      return;
    }
    
    try {
      keyInput.select();
      document.execCommand('copy');
      apiKeyCopyBtn.textContent = '✓ Copied!';
      console.log('✓ Copy succeeded');
      setTimeout(() => { apiKeyCopyBtn.textContent = 'Copy'; }, 2500);
    } catch (e) {
      console.error('execCommand failed:', e);
      try {
        await navigator.clipboard.writeText(keyValue);
        apiKeyCopyBtn.textContent = '✓ Copied!';
        console.log('✓ Clipboard API succeeded');
        setTimeout(() => { apiKeyCopyBtn.textContent = 'Copy'; }, 2500);
      } catch (e2) {
        console.error('Both methods failed:', e2);
        alert('Copy failed. Try Ctrl+A then Ctrl+C in the box above.');
      }
    }
  });
}

function applyTheme(t) {
  if (t === 'system') {
    const dark = window.matchMedia('(prefers-color-scheme: dark)').matches;
    document.body.classList.toggle('theme-light', !dark);
  } else {
    document.body.classList.toggle('theme-light', t === 'light');
  }
}

settingsBtn.addEventListener('click', () => { settingsModalOverlay.style.display = 'flex'; loadApiKeys(); });

const apiKeysShortcutBtn = document.getElementById('api-keys-shortcut-btn');
if (apiKeysShortcutBtn) apiKeysShortcutBtn.addEventListener('click', () => {
  openApiUsageOverlay();
});

// ─── Analytics removed - use Stats button in chat instead ─────────────────

settingsCloseBtn.addEventListener('click', () => { saveSettings(); settingsModalOverlay.style.display = 'none'; });
settingsModalOverlay.addEventListener('click', e => { if (e.target === settingsModalOverlay) { saveSettings(); settingsModalOverlay.style.display = 'none'; } });

(function() {
  const notifBtn    = document.getElementById('notif-toggle-btn');
  const notifStatus = document.getElementById('notif-status');
  if (!notifBtn) return;

  function updateNotifUI() {
    if (isIOS) {
      notifBtn.textContent = 'Not available on iPhone';
      notifBtn.disabled = true;
      notifStatus.textContent = 'Push notifications are not supported in Safari PWA. App works normally without them.';
      return;
    }
    if (!('Notification' in window)) {
      notifBtn.textContent = 'Not supported';
      notifBtn.disabled = true;
      notifStatus.textContent = 'Push notifications are not supported in this browser.';
      return;
    }
    const perm = Notification.permission;
    if (perm === 'granted') {
      notifBtn.textContent = 'Enabled ✓';
      notifBtn.style.borderColor = 'var(--accent)';
      notifBtn.style.color = 'var(--accent)';
      notifStatus.textContent = "You'll get a notification when Mythic AI replies while you're away.";
    } else if (perm === 'denied') {
      notifBtn.textContent = 'Blocked';
      notifBtn.style.borderColor = '#ef4444';
      notifBtn.style.color = '#ef4444';
      notifStatus.textContent = 'Notifications are blocked. Allow them in your browser site settings.';
    } else {
      notifBtn.textContent = 'Enable';
      notifBtn.style.borderColor = 'var(--border)';
      notifBtn.style.color = 'var(--muted)';
      notifStatus.textContent = "Get notified when Mythic AI replies while you're in another tab.";
    }
  }
  updateNotifUI();
  if (settingsBtn) settingsBtn.addEventListener('click', updateNotifUI);

  notifBtn.addEventListener('click', async () => {
    if (Notification.permission === 'granted') {
      const sub = await (async () => {
        try {
          const reg = await navigator.serviceWorker.getRegistration('/');
          return reg ? await reg.pushManager.getSubscription() : null;
        } catch { return null; }
      })();
      if (sub) {
        await fetch('/api/push/unsubscribe', {
          method: 'POST', headers: { 'Content-Type': 'application/json' },
          body: JSON.stringify({ endpoint: sub.endpoint }),
        });
        await sub.unsubscribe();
        localStorage.removeItem('mythic_push_subscribed');
        localStorage.removeItem('mythic_push_asked');
      }
      notifStatus.textContent = 'Notifications disabled.';
      updateNotifUI();
      return;
    }
    if (Notification.permission === 'denied') {
      notifStatus.textContent = 'Please allow notifications in your browser site settings, then reload.';
      return;
    }
    const perm = await Notification.requestPermission();
    if (perm === 'granted') {
      if (typeof window._notifyAiReply !== 'undefined') {
        try {
          const reg = await navigator.serviceWorker.getRegistration('/');
          if (reg) {
            const kr = await fetch('/api/push/vapid-public-key');
            if (kr.ok) {
              const { publicKey } = await kr.json();
              if (publicKey) {
                const padding = '='.repeat((4 - publicKey.length % 4) % 4);
                const b64 = (publicKey + padding).replace(/-/g, '+').replace(/_/g, '/');
                const raw = Uint8Array.from(atob(b64), c => c.charCodeAt(0));
                const sub = await reg.pushManager.subscribe({ userVisibleOnly: true, applicationServerKey: raw });
                await fetch('/api/push/subscribe', {
                  method: 'POST', headers: { 'Content-Type': 'application/json' },
                  body: JSON.stringify({ subscription: sub.toJSON() }),
                });
                localStorage.setItem('mythic_push_subscribed', '1');
              }
            }
          }
        } catch (err) { console.warn('[Push] manual subscribe failed:', err); }
      }
    }
    updateNotifUI();
  });
})();

document.querySelectorAll('.settings-choice').forEach(btn => {
  btn.addEventListener('click', () => {
    const group = btn.dataset.group;
    document.querySelectorAll(`[data-group="${group}"]`).forEach(b => {
      b.style.borderColor = 'var(--border)'; b.style.color = '';
    });
    btn.style.borderColor = 'var(--accent)'; btn.style.color = 'var(--accent)';
    if (group === 'theme') applyTheme(btn.dataset.value);
    if (group === 'bubble') {
      document.body.classList.remove('bubble-compact','bubble-comfortable','bubble-spacious');
      document.body.classList.add('bubble-' + btn.dataset.value);
    }
    saveSettings();
  });
});

accentColorInput.addEventListener('input', () => {
  document.documentElement.style.setProperty('--accent', accentColorInput.value);
});
fontSizeSlider.addEventListener('input', () => {
  fontSizeLabel.textContent = fontSizeSlider.value + 'px';
  document.documentElement.style.setProperty('--msg-font-size', fontSizeSlider.value + 'px');
});


loadSettings();

// ─── VOICE: Mythic AI auto-picks ONE voice, used for every language ───────
// No picker shown to the user — a single best-available voice is chosen
// once and reused for every reply, regardless of what language it's in.
// (Browsers can't make one system voice fluently pronounce every language,
// but this keeps a single consistent voice identity instead of a per-
// language switcher, per spec.)
(function() {
  const GOOD_NAME_HINTS = ['natural', 'neural', 'premium', 'enhanced', 'online', 'google'];
  let cachedVoices = [];
  let chosenVoice = null;

  function pickBestVoice(voices) {
    if (!voices.length) return null;
    // Prefer an English voice that sounds highest quality (heuristic name
    // match), then any English voice, then just the first voice available.
    const english = voices.filter(v => v.lang && v.lang.toLowerCase().startsWith('en'));
    const pool = english.length ? english : voices;
    const scored = pool.map(v => {
      const n = v.name.toLowerCase();
      let score = 0;
      if (GOOD_NAME_HINTS.some(h => n.includes(h))) score += 2;
      if (v.localService) score += 1; // local voices tend to be lower-latency
      if (/^en-us$/i.test(v.lang || '')) score += 1;
      return { v, score };
    });
    scored.sort((a, b) => b.score - a.score);
    return scored[0].v;
  }

  function refreshVoices() {
    if (!window.speechSynthesis) return;
    cachedVoices = window.speechSynthesis.getVoices() || [];
    if (cachedVoices.length) chosenVoice = pickBestVoice(cachedVoices);
  }

  if (window.speechSynthesis) {
    refreshVoices();
    window.speechSynthesis.onvoiceschanged = refreshVoices;
  }

  window.getChosenVoice = function() { return chosenVoice; };
})();

function renderMarkdown(text) {
  const div = document.createElement('div');
  div.className = 'msg-text md-rendered';
  let html = text
    .replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;');
  html = html.replace(/```(\w*)\n?([\s\S]*?)```/g, (_, lang, code) => {
    const id = 'cb_' + Math.random().toString(36).slice(2, 9);
    const label = lang ? lang : 'text';
    const escaped = code.trim();
    return `<div class="code-block-wrap" style="position:relative;margin:8px 0;border-radius:10px;overflow:hidden;border:1px solid var(--border);">`
      + `<div style="display:flex;justify-content:space-between;align-items:center;background:var(--panel);padding:6px 10px;font-size:11px;color:var(--muted);">`
      + `<span>${label}</span>`
      + `<button class="code-copy-btn" data-target="${id}" style="background:none;border:none;color:var(--muted);cursor:pointer;font-size:11px;padding:2px 6px;">📋 Copy</button>`
      + `</div>`
      + `<pre style="margin:0;padding:10px 12px;overflow-x:auto;background:var(--bg);"><code id="${id}" class="lang-${label}">${escaped}</code></pre>`
      + `</div>`;
  });
  html = html.replace(/`([^`]+)`/g, '<code>$1</code>');
  html = html.replace(/\*\*(.+?)\*\*/g, '<strong>$1</strong>');
  html = html.replace(/\*(.+?)\*/g, '<em>$1</em>');
  html = html.replace(/^### (.+)$/gm, '<h3 style="font-size:14px;margin:6px 0 3px;font-weight:700;">$1</h3>');
  html = html.replace(/^## (.+)$/gm, '<h2 style="font-size:15px;margin:8px 0 4px;font-weight:700;">$1</h2>');
  html = html.replace(/^# (.+)$/gm, '<h1 style="font-size:17px;margin:10px 0 5px;font-weight:700;">$1</h1>');
  html = html.replace(/(^|\n)([\-\*] .+(\n[\-\*] .+)*)/g, (_, pre, block) =>
    pre + '<ul>' + block.replace(/[\-\*] (.+)/g, '<li>$1</li>') + '</ul>');
  html = html.replace(/(^|\n)(\d+\. .+(\n\d+\. .+)*)/g, (_, pre, block) =>
    pre + '<ol>' + block.replace(/\d+\. (.+)/g, '<li>$1</li>') + '</ol>');
  html = html.replace(/\n/g, '<br>');
  div.innerHTML = html;
  return div;
}

const _origAddMessage = addMessage;
addMessage = function(role, text, attachment) {
  const textNode = _origAddMessage(role, text, attachment);
  if (role === 'ai' && text) {
    try {
      const md = renderMarkdown(text);
      textNode.parentNode.replaceChild(md, textNode);
      return md;
    } catch { return textNode; }
  }
  return textNode;
};

const _origAddMsg2 = addMessage;
function addMessageWithTimestamp(role, text, attachment) {
  const node = _origAddMsg2(role, text, attachment);
  const row = node.closest ? node.closest('.msg-row') : null;
  if (row) {
    const ts = document.createElement('div');
    ts.className = 'msg-timestamp';
    ts.textContent = new Date().toLocaleTimeString([], {hour:'2-digit',minute:'2-digit'});
    row.appendChild(ts);
  }
  return node;
}
const _rawAdd = addMessage;
window._addMsgFinal = function(role, text, attachment) {
  const node = _rawAdd(role, text, attachment);
  const row = node && node.closest ? node.closest('.msg-row') : null;
  if (row && !row.querySelector('.msg-timestamp')) {
    const ts = document.createElement('div');
    ts.className = 'msg-timestamp';
    ts.textContent = new Date().toLocaleTimeString([], {hour:'2-digit',minute:'2-digit'});
    row.appendChild(ts);
  }
  return node;
};

async function addFollowupSuggestions(aiText) {
  if (!aiText || aiText.length < 50) return;
  try {
    const r = await fetch('/api/chat', {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify({
        message: `Based on this AI reply, suggest 3 short follow-up questions the user might ask. Reply ONLY with 3 questions, one per line, no numbering, no extra text, in English:\n\n${aiText.slice(0,400)}`,
        conversation_id: null, model: 'mythic-1.0', user_name: '', ephemeral: true
      })
    });
    if (!r.ok) return;
    const reader = r.body.getReader();
    const dec = new TextDecoder();
    let full = '';
    while (true) { const {done,value} = await reader.read(); if(done) break; full += dec.decode(value,{stream:true}); }
    const qs = full.trim().split('\n').filter(q => q.trim().length > 5).slice(0,3);
    if (!qs.length) return;
    const wrap = document.createElement('div');
    wrap.style.cssText = 'display:flex;flex-wrap:wrap;gap:6px;margin-top:6px;max-width:100%;';
    qs.forEach(q => {
      const btn = document.createElement('button');
      btn.textContent = q.trim().replace(/^["']|["']$/g,'');
      btn.style.cssText = 'background:var(--panel);border:1px solid var(--border);color:var(--muted);font-size:11.5px;padding:5px 10px;border-radius:16px;cursor:pointer;font-family:inherit;text-align:left;';
      btn.addEventListener('click', () => {
        input.value = btn.textContent; input.focus(); autoResize();
        form.requestSubmit(); wrap.remove();
      });
      wrap.appendChild(btn);
    });
    messagesEl.appendChild(wrap);
    scrollToBottom();
  } catch {}
}

function addReactionBar(row) {
  if (row.querySelector('.reaction-bar')) return;
  const bar = document.createElement('div');
  bar.className = 'reaction-bar';
  bar.style.cssText = 'display:flex;gap:4px;margin-top:3px;';
  ['👍','👎','❤️','😂','🔖'].forEach(emoji => {
    const btn = document.createElement('button');
    btn.textContent = emoji;
    btn.style.cssText = 'background:none;border:1px solid var(--border);border-radius:12px;padding:2px 7px;font-size:13px;cursor:pointer;touch-action:manipulation;';
    btn.addEventListener('click', () => {
      btn.style.borderColor = btn.style.borderColor === 'var(--accent)' ? 'var(--border)' : 'var(--accent)';
      btn.style.background  = btn.style.background === 'var(--accent-dim)' ? '' : 'var(--accent-dim)';
    });
    bar.appendChild(btn);
  });
  row.appendChild(bar);
}

function addSearchUI() {
  const searchWrap = document.createElement('div');
  searchWrap.id = 'msg-search-wrap';
  searchWrap.style.cssText = 'display:none;position:fixed;top:70px;left:50%;transform:translateX(-50%);z-index:150;background:var(--panel);border:1.5px solid var(--accent);border-radius:10px;padding:8px 12px;display:flex;gap:8px;align-items:center;min-width:280px;box-shadow:0 4px 20px rgba(0,0,0,.3);';
  searchWrap.innerHTML = '<input id="msg-search-input" placeholder="Search messages..." style="background:transparent;border:none;color:var(--text);font-size:14px;outline:none;flex:1;font-family:inherit;"><span id="msg-search-count" style="color:var(--muted);font-size:12px;"></span><button id="msg-search-close" style="background:none;border:none;color:var(--muted);cursor:pointer;font-size:16px;">✕</button>';
  searchWrap.style.display = 'none';
  document.body.appendChild(searchWrap);

  let highlights = [];
  document.getElementById('msg-search-input').addEventListener('input', e => {
    highlights.forEach(el => { el.style.background = ''; el.style.outline = ''; });
    highlights = [];
    const q = e.target.value.trim().toLowerCase();
    if (!q) { document.getElementById('msg-search-count').textContent = ''; return; }
    document.querySelectorAll('.msg-text,.md-rendered').forEach(el => {
      if (el.textContent.toLowerCase().includes(q)) {
        el.style.background = 'rgba(255,200,0,.15)';
        el.style.outline = '2px solid rgba(255,200,0,.4)';
        highlights.push(el);
      }
    });
    document.getElementById('msg-search-count').textContent = highlights.length ? `${highlights.length} found` : 'no results';
    if (highlights.length) highlights[0].scrollIntoView({behavior:'smooth',block:'center'});
  });
  document.getElementById('msg-search-close').addEventListener('click', () => {
    searchWrap.style.display = 'none';
    highlights.forEach(el => { el.style.background = ''; el.style.outline = ''; });
  });
  return searchWrap;
}
const msgSearchWrap = addSearchUI();

document.addEventListener('keydown', e => {
  if ((e.ctrlKey || e.metaKey) && e.key === 'f' && !settingsModalOverlay.style.display.includes('flex')) {
    e.preventDefault();
    msgSearchWrap.style.display = 'flex';
    setTimeout(() => document.getElementById('msg-search-input').focus(), 50);
  }
  if (e.key === 'Escape') msgSearchWrap.style.display = 'none';
});

// ─── PWA INSTALL BUTTON ──────────────────────────────────────────────────────
const installBtn = document.getElementById('install-btn');
let _deferredInstallPrompt = null;
let _relatedAppInstalled = false; // set true if getInstalledRelatedApps() finds an existing Mythic AI install

function _showInstallBtn() {
  if (!installBtn) return;
  installBtn.style.display = 'flex';
  installBtn.style.alignItems = 'center';
}
function _hideInstallBtn() {
  if (installBtn) installBtn.style.display = 'none';
}

// ─── Runtime PWA diagnostics ─────────────────────────────────────────────
// Real, live checks (secure context, manifest fetched + validated the same
// way the browser reads it, every manifest icon actually resolving, service
// worker registration + control state) — run on load, logged to console,
// and re-used by the install modal so it can show the ACTUAL detected
// reason instead of a canned message. Inspect at any time via
// window.__pwaDebug in DevTools.
const _pwaDiag = {
  httpsOk: null,
  manifestOk: null,
  manifestReason: null,
  iconsOk: null,
  iconsReason: null,
  swRegistered: null,
  swControlling: null,
  swScriptURL: null,
  swScope: null,
  swState: null,
  beforeInstallPromptReceived: false,
  beforeInstallPromptAt: null,
  appInstalled: false,
  appInstalledAt: null,
  checkedAt: null,
};

async function _runPwaDiagnostics() {
  _pwaDiag.httpsOk = window.isSecureContext === true;
  console.log(_pwaDiag.httpsOk
    ? '[PWA] Secure context OK (HTTPS or localhost)'
    : '[PWA] Installation unavailable because: page is not a secure context (needs HTTPS)');

  try {
    const link = document.querySelector('link[rel="manifest"]');
    const manifestUrl = link ? link.href : new URL('/manifest.json', location.href).href;
    const res = await fetch(manifestUrl, { cache: 'no-store' });
    if (!res.ok) {
      _pwaDiag.manifestOk = false;
      _pwaDiag.manifestReason = `manifest request failed (HTTP ${res.status})`;
    } else {
      const m = await res.json();
      const missing = ['name', 'short_name', 'start_url', 'display', 'icons']
        .filter(k => !m[k] || (Array.isArray(m[k]) && m[k].length === 0));
      if (missing.length) {
        _pwaDiag.manifestOk = false;
        _pwaDiag.manifestReason = `manifest missing required field(s): ${missing.join(', ')}`;
      } else if (!['standalone', 'fullscreen', 'minimal-ui'].includes(m.display)) {
        _pwaDiag.manifestOk = false;
        _pwaDiag.manifestReason = `manifest "display" is "${m.display}" (needs standalone/fullscreen/minimal-ui)`;
      } else {
        _pwaDiag.manifestOk = true;
        const iconChecks = await Promise.all((m.icons || []).map(icon =>
          fetch(new URL(icon.src, manifestUrl), { method: 'HEAD', cache: 'no-store' })
            .then(r => ({ src: icon.src, ok: r.ok, sizes: icon.sizes || '' }))
            .catch(() => ({ src: icon.src, ok: false, sizes: icon.sizes || '' }))
        ));
        const broken = iconChecks.filter(c => !c.ok);
        const has192 = iconChecks.some(c => c.ok && c.sizes.split('x')[0] >= 192);
        const has512 = iconChecks.some(c => c.ok && c.sizes.split('x')[0] >= 512);
        if (broken.length) {
          _pwaDiag.iconsOk = false;
          _pwaDiag.iconsReason = `icon(s) failed to load: ${broken.map(b => b.src).join(', ')}`;
        } else if (!has192 || !has512) {
          _pwaDiag.iconsOk = false;
          _pwaDiag.iconsReason = 'manifest needs at least one icon ≥192px and one ≥512px';
        } else {
          _pwaDiag.iconsOk = true;
        }
      }
    }
  } catch (err) {
    _pwaDiag.manifestOk = false;
    _pwaDiag.manifestReason = 'manifest fetch threw: ' + err.message;
  }
  console.log(_pwaDiag.manifestOk ? '[PWA] Manifest valid ✅' : `[PWA] Installation unavailable because: ${_pwaDiag.manifestReason}`);
  if (_pwaDiag.manifestOk) {
    console.log(_pwaDiag.iconsOk ? '[PWA] Manifest icons OK ✅' : `[PWA] Installation unavailable because: ${_pwaDiag.iconsReason}`);
  }

  if ('serviceWorker' in navigator) {
    try {
      const reg = await navigator.serviceWorker.getRegistration('/');
      _pwaDiag.swRegistered = !!reg && !!(reg.active || reg.waiting || reg.installing);
      _pwaDiag.swControlling = !!navigator.serviceWorker.controller;
      const sw = reg && (reg.active || reg.waiting || reg.installing);
      _pwaDiag.swScriptURL = sw ? sw.scriptURL : null;
      _pwaDiag.swScope = reg ? reg.scope : null;
      _pwaDiag.swState = sw ? sw.state : null;
      console.log(_pwaDiag.swRegistered ? '[PWA] Service worker registered ✅' : '[PWA] Installation unavailable because: no service worker registration was found');
      console.log(_pwaDiag.swControlling
        ? '[PWA] Service worker controlling page ✅'
        : '[PWA] Service worker not controlling this load yet (normal on a first visit — Chrome does not require this for installability, only that a registration exists)');
    } catch (err) {
      _pwaDiag.swRegistered = false;
      console.warn('[PWA] Service worker check failed:', err);
    }
  } else {
    _pwaDiag.swRegistered = false;
    console.warn('[PWA] Installation unavailable because: navigator.serviceWorker is unsupported in this browser');
  }

  _pwaDiag.checkedAt = Date.now();
  return _pwaDiag;
}

Object.defineProperty(window, '__pwaDebug', {
  get: () => ({
    ..._pwaDiag,
    deferredPromptAvailable: !!_deferredInstallPrompt,
    displayMode: window.matchMedia('(display-mode: standalone)').matches ? 'standalone' : 'browser',
    iosStandalone: !!window.navigator.standalone,
    origin: window.location.origin,
    protocol: window.location.protocol,
    manifestURL: (document.querySelector('link[rel="manifest"]') || {}).href || null,
  }),
});

const _pwaDiagReady = _runPwaDiagnostics();

window.addEventListener('beforeinstallprompt', e => {
  e.preventDefault();
  _deferredInstallPrompt = e;
  _pwaDiag.beforeInstallPromptReceived = true;
  _pwaDiag.beforeInstallPromptAt = new Date().toISOString();
  _showInstallBtn();
  console.log('[PWA] beforeinstallprompt received — install prompt available ✅');
  _refreshPwaDiagnosticsPanel();
});

window.addEventListener('appinstalled', () => {
  _hideInstallBtn();
  _deferredInstallPrompt = null;
  _pwaDiag.appInstalled = true;
  _pwaDiag.appInstalledAt = new Date().toISOString();
  localStorage.setItem('mythic_pwa_installed', '1');
  _closeInstallModal();
  _showInstallSuccessToast();
  console.log('[PWA] appinstalled fired — app installed ✅');
  _refreshPwaDiagnosticsPanel();
});

// After diagnostics finish, log a single clear summary line so it's obvious
// at a glance whether beforeinstallprompt not firing is expected (browser
// genuinely doesn't support it, or checks failed) or unexplained (checks
// all pass — likely Incognito/Private mode, a recent dismissal cooldown,
// or already installed under a different profile, all of which are normal
// Chrome behavior with no code-level fix).
_pwaDiagReady.then(() => {
  if (_deferredInstallPrompt || window.matchMedia('(display-mode: standalone)').matches || window.navigator.standalone) return;
  const allChecksPass = _pwaDiag.httpsOk && _pwaDiag.manifestOk && _pwaDiag.iconsOk && _pwaDiag.swRegistered;
  if (allChecksPass) {
    console.log('[PWA] Install prompt available — all installability checks pass, but beforeinstallprompt has not fired on this load. Common causes: Incognito/Private browsing, a recently-dismissed prompt (Chrome cools down for a few days), or the app already installed in another profile.');
  } else {
    console.warn('[PWA] Installation currently blocked — see the "[PWA] Installation unavailable because" lines above for the exact reason(s).');
  }
});

// Only hide the Install button once we're SURE the app is already running
// as an installed PWA — otherwise keep it visible (the install modal below
// adapts to whatever state is actually detected) so it's never mysteriously
// missing on desktop Chrome/Firefox or browsers that don't fire
// beforeinstallprompt.
if (window.matchMedia('(display-mode: standalone)').matches || window.navigator.standalone) {
  _hideInstallBtn();
} else {
  _showInstallBtn();
  // Covers the case standalone-mode can't: the app IS installed, but this
  // tab was opened as a normal URL rather than launched from the installed
  // app icon (display-mode is still "browser" here). Chrome's own omnibox
  // knows this via its "Open in app" button; getInstalledRelatedApps() is
  // the one API that lets the page know it too, via the
  // related_applications self-reference declared in manifest.json.
  if ('getInstalledRelatedApps' in navigator) {
    navigator.getInstalledRelatedApps().then(apps => {
      if (apps && apps.length > 0) {
        _relatedAppInstalled = true;
        _hideInstallBtn();
        console.log('[PWA] getInstalledRelatedApps() found an existing install — hiding Install button');
      }
    }).catch(() => {});
  }
}

// ─── Live PWA diagnostics panel ─────────────────────────────────────────
// Reports the raw, real browser state — not a summary, not a guess.
// Toggle with Alt+Shift+P, ?pwadebug=1 in the URL, or window.showPwaDiagnostics()
// in the console. This is separate from (and never opened automatically by)
// the install modal.
//
// Honest limitation: the field below labeled "Native install prompt
// available" can only ever be YES or NO, because that's all the web
// platform exposes to page JavaScript — there is no public API for "why".
// Chrome's own internal reason code (e.g. "in-incognito", "already
// installed", "prompt previously dismissed — cooldown active") is only
// visible through Chrome DevTools → Application → Manifest → Installability,
// or the DevTools Protocol's Page.getInstallabilityErrors, neither of which
// a served web page can call on itself. When every check below is ✓ but the
// prompt is still NO, that internal reason — not a bug in this code — is
// almost always the explanation.
function _pwaDiagnosticsRow(label, value) {
  return `<div style="display:flex;justify-content:space-between;gap:14px;padding:3px 0;border-bottom:1px solid rgba(255,255,255,.06);">
    <span style="color:var(--muted);">${label}</span>
    <span style="color:var(--text);font-weight:600;text-align:right;word-break:break-all;">${value}</span>
  </div>`;
}

function _pwaDiagnosticsHTML() {
  const d = window.__pwaDebug;
  const nativeAvailable = !!_deferredInstallPrompt;
  const allChecksPass = d.httpsOk && d.manifestOk && d.iconsOk && d.swRegistered;
  const reasonIfNo = d.appInstalled || d.displayMode === 'standalone' || d.iosStandalone
    ? 'Already running as the installed app.'
    : !allChecksPass
      ? 'A configuration check below is failing — see the ✗ row(s).'
      : 'All checkable requirements pass. The specific internal reason (Incognito, dismissal cooldown, already installed elsewhere, etc.) is only visible via Chrome DevTools → Application → Manifest → Installability, not to page JavaScript.';

  return `
    <div style="font-weight:700;font-size:13px;margin:10px 0 4px;color:var(--accent);">Native install prompt available: ${nativeAvailable ? 'YES' : 'NO'}</div>
    ${nativeAvailable ? '' : `<div style="font-size:11.5px;color:var(--muted);margin-bottom:8px;line-height:1.5;">${reasonIfNo}</div>`}
    <div style="font-weight:700;font-size:12px;margin:10px 0 2px;color:var(--text);">Origin &amp; page</div>
    ${_pwaDiagnosticsRow('window.location.origin', d.origin)}
    ${_pwaDiagnosticsRow('location.protocol', d.protocol)}
    ${_pwaDiagnosticsRow('display mode', d.displayMode)}
    ${_pwaDiagnosticsRow('iOS standalone', d.iosStandalone)}
    <div style="font-weight:700;font-size:12px;margin:10px 0 2px;color:var(--text);">Manifest</div>
    ${_pwaDiagnosticsRow('manifest URL', d.manifestURL || 'not linked!')}
    ${_pwaDiagnosticsRow('manifest valid', d.manifestOk === null ? 'checking…' : d.manifestOk)}
    ${d.manifestReason ? _pwaDiagnosticsRow('reason', d.manifestReason) : ''}
    ${_pwaDiagnosticsRow('icons load', d.iconsOk === null ? 'checking…' : d.iconsOk)}
    ${d.iconsReason ? _pwaDiagnosticsRow('reason', d.iconsReason) : ''}
    <div style="font-weight:700;font-size:12px;margin:10px 0 2px;color:var(--text);">Service worker</div>
    ${_pwaDiagnosticsRow('registered', d.swRegistered === null ? 'checking…' : d.swRegistered)}
    ${_pwaDiagnosticsRow('controller present', d.swControlling)}
    ${_pwaDiagnosticsRow('registration scope', d.swScope || '—')}
    ${_pwaDiagnosticsRow('script URL', d.swScriptURL || '—')}
    ${_pwaDiagnosticsRow('worker state', d.swState || '—')}
    <div style="font-weight:700;font-size:12px;margin:10px 0 2px;color:var(--text);">Events</div>
    ${_pwaDiagnosticsRow('beforeinstallprompt received', d.beforeInstallPromptReceived)}
    ${_pwaDiagnosticsRow('received at', d.beforeInstallPromptAt || '—')}
    ${_pwaDiagnosticsRow('appinstalled fired', d.appInstalled)}
    ${_pwaDiagnosticsRow('fired at', d.appInstalledAt || '—')}
  `;
}

function _refreshPwaDiagnosticsPanel() {
  const body = document.getElementById('pwa-diag-body');
  if (body) body.innerHTML = _pwaDiagnosticsHTML();
}

async function showPwaDiagnostics() {
  if (!_pwaDiag.checkedAt) await _pwaDiagReady;
  const existing = document.getElementById('pwa-diag-panel');
  if (existing) { existing.remove(); return; }
  const panel = document.createElement('div');
  panel.id = 'pwa-diag-panel';
  panel.style.cssText = 'position:fixed;bottom:16px;right:16px;z-index:10001;width:340px;max-height:70vh;overflow:auto;background:var(--panel);border:1px solid var(--border);border-radius:14px;padding:16px;box-shadow:0 8px 40px rgba(0,0,0,.5);font-size:12px;font-family:inherit;';
  panel.innerHTML = `
    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:6px;">
      <span style="font-weight:800;font-size:13.5px;color:var(--text);">🔧 PWA Diagnostics</span>
      <button id="pwa-diag-close" style="background:none;border:none;color:var(--muted);font-size:16px;cursor:pointer;line-height:1;">✕</button>
    </div>
    <div id="pwa-diag-body">${_pwaDiagnosticsHTML()}</div>
    <button id="pwa-diag-refresh" style="margin-top:10px;width:100%;background:var(--accent);color:#fff;border:none;border-radius:8px;padding:8px;font-size:12px;font-weight:600;cursor:pointer;font-family:inherit;">Re-run checks</button>
  `;
  document.body.appendChild(panel);
  document.getElementById('pwa-diag-close').addEventListener('click', () => panel.remove());
  document.getElementById('pwa-diag-refresh').addEventListener('click', async () => {
    await _runPwaDiagnostics();
    _refreshPwaDiagnosticsPanel();
  });
}
window.showPwaDiagnostics = showPwaDiagnostics;

window.addEventListener('keydown', e => {
  if (e.altKey && e.shiftKey && (e.key === 'P' || e.key === 'p')) showPwaDiagnostics();
});

if (new URLSearchParams(location.search).get('pwadebug')) {
  _pwaDiagReady.then(() => showPwaDiagnostics());
}

// ─── Install modal shell ─────────────────────────────────────────────────
// One shared modal container + close behavior; each _show*Modal() below
// just supplies the inner content for its particular install state.
function _closeInstallModal() {
  const m = document.getElementById('install-modal');
  if (m) m.remove();
}

function _renderInstallModal(bodyHTML, { anchorBottom = false } = {}) {
  _closeInstallModal();
  const m = document.createElement('div');
  m.id = 'install-modal';
  m.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.78);z-index:9999;display:flex;' +
    (anchorBottom ? 'align-items:flex-end;' : 'align-items:center;') +
    'justify-content:center;padding:20px;';
  m.innerHTML = `
    <div style="background:var(--panel);border:1px solid var(--border);border-radius:20px;padding:28px 24px;width:100%;max-width:420px;text-align:center;box-shadow:0 4px 40px rgba(0,0,0,.4);">
      ${bodyHTML}
    </div>`;
  document.body.appendChild(m);
  m.addEventListener('click', e => { if (e.target === m) m.remove(); });
  m.querySelectorAll('[data-install-close]').forEach(btn =>
    btn.addEventListener('click', _closeInstallModal));
  return m;
}

function _showInstallSuccessToast() {
  const t = document.createElement('div');
  t.style.cssText = 'position:fixed;bottom:24px;left:50%;transform:translateX(-50%);' +
    'background:var(--accent);color:#fff;padding:12px 22px;border-radius:10px;' +
    'font-size:13.5px;font-weight:600;z-index:10000;box-shadow:0 4px 24px rgba(0,0,0,.35);' +
    'max-width:90vw;text-align:center;';
  t.textContent = '✅ Mythic AI has been installed successfully!';
  document.body.appendChild(t);
  setTimeout(() => { t.style.transition = 'opacity .3s'; t.style.opacity = '0'; setTimeout(() => t.remove(), 300); }, 3500);
}

// The primary, "real install" state — beforeinstallprompt already fired and
// the deferred event is ready to use.
function _showNativeInstallModal() {
  _renderInstallModal(`
    <div style="font-size:42px;margin-bottom:10px;">📲</div>
    <div style="font-weight:700;font-size:18px;margin-bottom:8px;color:var(--text);">Install Mythic AI</div>
    <div style="color:var(--muted);font-size:13.5px;line-height:1.6;margin-bottom:20px;">
      Install Mythic AI on your device for a faster, app-like experience.
    </div>
    <button id="install-modal-primary" style="background:var(--accent);color:#fff;border:none;border-radius:10px;padding:14px 32px;font-size:15px;font-weight:700;cursor:pointer;font-family:inherit;width:100%;margin-bottom:14px;">Install Mythic AI</button>
    <div style="color:var(--muted);font-size:11.5px;line-height:1.5;">Works offline where supported &middot; Opens like a desktop/mobile app</div>
  `);
  const primary = document.getElementById('install-modal-primary');
  if (primary) primary.addEventListener('click', _triggerNativeInstall);
}

async function _triggerNativeInstall() {
  if (!_deferredInstallPrompt) { _closeInstallModal(); return; }
  // A saved beforeinstallprompt event can only be used once — grab it and
  // clear the reference immediately so it's never reused.
  const promptEvent = _deferredInstallPrompt;
  _deferredInstallPrompt = null;
  const primary = document.getElementById('install-modal-primary');
  if (primary) { primary.disabled = true; primary.style.opacity = '.7'; primary.textContent = 'Waiting for confirmation…'; }
  try {
    promptEvent.prompt();
    const { outcome } = await promptEvent.userChoice;
    // On success, the 'appinstalled' listener closes the modal, hides the
    // button, and shows the success toast. On dismissal/cancel, just close
    // the modal — the Install button must stay VISIBLE (per spec: only
    // 'appinstalled' or already-standalone-on-load may hide it).
    if (outcome !== 'accepted') {
      _closeInstallModal();
    }
  } catch (err) {
    console.warn('[PWA] install prompt failed:', err);
    _renderInstallModal(`
      <div style="font-size:38px;margin-bottom:10px;">⚠️</div>
      <div style="font-weight:700;font-size:18px;margin-bottom:8px;color:var(--text);">Install didn't start</div>
      <div style="color:var(--muted);font-size:13.5px;line-height:1.6;margin-bottom:20px;">Something interrupted the install prompt. Please try again from the Install button, or use your browser's menu.</div>
      <button data-install-close style="background:var(--accent);color:#fff;border:none;border-radius:10px;padding:12px 32px;font-size:15px;font-weight:600;cursor:pointer;font-family:inherit;width:100%;">Got it!</button>
    `);
  }
}

// iOS Safari never fires beforeinstallprompt — "Add to Home Screen" via the
// Share sheet is the only real installation path there.
function _showIOSInstallModal() {
  _renderInstallModal(`
    <div style="font-size:42px;margin-bottom:10px;">📲</div>
    <div style="font-weight:700;font-size:18px;margin-bottom:8px;color:var(--text);">Install Mythic AI</div>
    <div style="color:var(--muted);font-size:13.5px;line-height:1.7;margin-bottom:20px;">
      Tap the <strong style="color:var(--text);">Share button</strong> <span style="font-size:17px;">⬆</span> at the bottom of Safari,<br>
      then tap <strong style="color:var(--text);">"Add to Home Screen"</strong> <span style="font-size:15px;">➕</span>
    </div>
    <button data-install-close style="background:var(--accent);color:#fff;border:none;border-radius:10px;padding:12px 32px;font-size:15px;font-weight:600;cursor:pointer;font-family:inherit;width:100%;">Got it!</button>
  `, { anchorBottom: true });
}

// Fallback for everything else: browsers that don't support one-click PWA
// installs at all (desktop Safari, Firefox), or Chrome-family browsers that
// haven't fired beforeinstallprompt yet. We never fake a download here —
// just explain, per-browser, what's actually possible.
// STATE D — already running as an installed PWA.
function _showAlreadyInstalledModal(viaRelatedApps) {
  const message = viaRelatedApps
    ? "Mythic AI is already installed on this device — you're just viewing it in a regular browser tab right now. Open it from your home screen/apps list, or use the \u201cOpen in app\u201d button in your browser's address bar."
    : "You're running the installed app on this device.";
  _renderInstallModal(`
    <div style="font-size:38px;margin-bottom:10px;">✅</div>
    <div style="font-weight:700;font-size:18px;margin-bottom:8px;color:var(--text);">Mythic AI is already installed</div>
    <div style="color:var(--muted);font-size:13.5px;line-height:1.6;margin-bottom:20px;">${message}</div>
    <button data-install-close style="background:var(--accent);color:#fff;border:none;border-radius:10px;padding:12px 32px;font-size:15px;font-weight:600;cursor:pointer;font-family:inherit;width:100%;">Got it!</button>
  `);
}

// A small ✓/✗ readout of the live diagnostic results, used by both STATE B
// (browser supports installs, prompt just hasn't fired) and the genuine
// failure case, so the reason shown is always the real, current state —
// never a guess.
function _diagChecklistHTML() {
  const rows = [
    ['Secure context (HTTPS/localhost)', _pwaDiag.httpsOk],
    ['Manifest valid', _pwaDiag.manifestOk],
    ['Manifest icons load', _pwaDiag.iconsOk],
    ['Service worker registered', _pwaDiag.swRegistered],
  ];
  return `<div style="text-align:left;background:var(--bg,rgba(0,0,0,.15));border-radius:10px;padding:10px 14px;margin-bottom:16px;">` +
    rows.map(([label, ok]) => `
      <div style="display:flex;justify-content:space-between;gap:10px;padding:4px 0;font-size:12px;color:var(--muted);">
        <span>${label}</span>
        <span style="color:${ok === false ? '#e5484d' : ok === true ? 'var(--accent)' : 'var(--muted)'};font-weight:700;">
          ${ok === false ? '✗' : ok === true ? '✓' : '…'}
        </span>
      </div>`).join('') +
    `</div>`;
}

// STATE B/C — no deferred prompt available. Show what the live diagnostics
// actually found: a concrete failure reason if one exists, or (when every
// check passes) an honest explanation of why Chrome still might not have
// offered the prompt, plus real per-browser fallback instructions.
function _showUnsupportedInstallModal() {
  const ua = navigator.userAgent;
  const isChromiumFamily = /Edg|Chrome|SamsungBrowser/i.test(ua) && !/Firefox/i.test(ua);
  const diagFailed = _pwaDiag.checkedAt && (
    _pwaDiag.httpsOk === false || _pwaDiag.manifestOk === false ||
    _pwaDiag.iconsOk === false || _pwaDiag.swRegistered === false
  );
  const failureReason = !_pwaDiag.httpsOk ? 'This page is not being served over HTTPS (or localhost), which PWA installation requires.'
    : !_pwaDiag.manifestOk ? `Manifest problem: ${_pwaDiag.manifestReason}`
    : !_pwaDiag.iconsOk ? `Icon problem: ${_pwaDiag.iconsReason}`
    : !_pwaDiag.swRegistered ? 'No service worker registration was found for this page.'
    : null;

  let title, message, showChecklist = false;
  if (diagFailed) {
    title = 'Installation is blocked';
    message = failureReason;
  } else if (isChromiumFamily) {
    title = "Install prompt hasn't appeared yet";
    message = 'Every installability check below passes, so your browser does support installing Mythic AI — Chrome just hasn\'t offered the prompt on this visit. That usually means Private/Incognito browsing, a prompt you dismissed recently (Chrome pauses re-offering it for a few days), or the app already installed under a different profile. You can also install manually from the ⋮ menu → "Install app".';
    showChecklist = true;
  } else if (/Firefox/i.test(ua)) {
    title = "Installation isn't available here";
    message = 'Firefox doesn\'t support one-tap app installs yet. You can still bookmark Mythic AI or pin the tab for quick access.';
  } else if (/^((?!chrome|android|crios|edgios).)*safari/i.test(ua)) {
    title = "Installation isn't available here";
    message = 'On macOS Sonoma or later, open Safari\'s Share menu and choose "Add to Dock". Otherwise, try Mythic AI in Chrome or Edge to install it as an app.';
  } else {
    title = "Installation isn't available here";
    message = 'This browser doesn\'t support one-tap app installs. Try opening Mythic AI in Chrome, Edge, or Samsung Internet.';
  }

  _renderInstallModal(`
    <div style="font-size:38px;margin-bottom:10px;">${diagFailed ? '⚠️' : 'ℹ️'}</div>
    <div style="font-weight:700;font-size:18px;margin-bottom:8px;color:var(--text);">${title}</div>
    <div style="color:var(--muted);font-size:13.5px;line-height:1.7;margin-bottom:${showChecklist ? '14' : '20'}px;">${message}</div>
    ${showChecklist ? _diagChecklistHTML() : ''}
    <button data-install-close style="background:var(--accent);color:#fff;border:none;border-radius:10px;padding:12px 32px;font-size:15px;font-weight:600;cursor:pointer;font-family:inherit;width:100%;">Got it!</button>
    ${showChecklist ? '<button id="install-modal-view-diag" style="margin-top:10px;background:none;border:none;color:var(--muted);font-size:11.5px;text-decoration:underline;cursor:pointer;font-family:inherit;">View full diagnostics panel</button>' : ''}
  `);
  const diagLink = document.getElementById('install-modal-view-diag');
  if (diagLink) diagLink.addEventListener('click', () => { _closeInstallModal(); showPwaDiagnostics(); });
}

async function _openInstallModal() {
  if (window.matchMedia('(display-mode: standalone)').matches || window.navigator.standalone) {
    _hideInstallBtn();
    _showAlreadyInstalledModal(false);
    return;
  }
  if (_deferredInstallPrompt) {
    _showNativeInstallModal();
    return;
  }
  // getInstalledRelatedApps() already told us this IS installed, just not
  // in this tab — show that instead of the generic "hasn't appeared yet"
  // fallback, which would be misleading here.
  if (_relatedAppInstalled) {
    _hideInstallBtn();
    _showAlreadyInstalledModal(true);
    return;
  }
  // Make sure diagnostics have finished at least once before deciding what
  // to show, so the checklist/reason is never stale or half-computed.
  if (!_pwaDiag.checkedAt) await _pwaDiagReady;
  if (isIOS && !window.navigator.standalone) {
    _showIOSInstallModal();
  } else {
    _showUnsupportedInstallModal();
  }
}

if (installBtn) {
  installBtn.addEventListener('click', _openInstallModal);
}

// ── Header "⋯" overflow menu ──────────────────────────────────────────
// Purely a show/hide wrapper — none of the moved buttons' own click
// handlers (vip-btn, fullscreen-btn, name-btn, settings-btn, share-btn,
// export-btn, clear-btn) are touched; they're wired up elsewhere exactly
// as before, by the same ids, and keep firing normally. This block only
// closes the dropdown after a choice is made (or on outside click / Esc).
const headerMenuBtn      = document.getElementById('header-menu-btn');
const headerMenuDropdown = document.getElementById('header-menu-dropdown');

function _openHeaderMenu() {
  if (!headerMenuDropdown) return;
  headerMenuDropdown.classList.add('open');
  headerMenuBtn.setAttribute('aria-expanded', 'true');
  setTimeout(() => {
    document.addEventListener('click', _onHeaderMenuOutsideClick);
    document.addEventListener('keydown', _onHeaderMenuEscape);
  }, 0);
}

function _closeHeaderMenu() {
  if (!headerMenuDropdown) return;
  headerMenuDropdown.classList.remove('open');
  headerMenuBtn.setAttribute('aria-expanded', 'false');
  document.removeEventListener('click', _onHeaderMenuOutsideClick);
  document.removeEventListener('keydown', _onHeaderMenuEscape);
}

function _onHeaderMenuOutsideClick(e) {
  if (!headerMenuDropdown.contains(e.target) && e.target !== headerMenuBtn) _closeHeaderMenu();
}

function _onHeaderMenuEscape(e) {
  if (e.key === 'Escape') _closeHeaderMenu();
}

if (headerMenuBtn && headerMenuDropdown) {
  headerMenuBtn.addEventListener('click', (e) => {
    e.stopPropagation();
    headerMenuDropdown.classList.contains('open') ? _closeHeaderMenu() : _openHeaderMenu();
  });
  // Let each item's own existing handler run first, then close the menu —
  // this is the only new behavior; the handlers themselves are unchanged.
  headerMenuDropdown.addEventListener('click', (e) => {
    if (e.target.closest('button')) setTimeout(_closeHeaderMenu, 0);
  });
}

const notifBanner     = document.getElementById('notif-banner');
const notifAllowBtn   = document.getElementById('notif-banner-allow');
const notifDismissBtn = document.getElementById('notif-banner-dismiss');
let _swReg = null;

function _hideBanner() { if (notifBanner) notifBanner.style.display = 'none'; }

function _showBanner() {
  if (!notifBanner) return;
  if (!('Notification' in window)) return;
  if (Notification.permission === 'granted') return;
  if (Notification.permission === 'denied') return;
  if (localStorage.getItem('mythic_notif_dismissed')) return;
  notifBanner.style.display = 'flex';
}

function _urlB64ToUint8(b64url) {
  const pad = '='.repeat((4 - b64url.length % 4) % 4);
  const b64 = (b64url + pad).replace(/-/g, '+').replace(/_/g, '/');
  return Uint8Array.from(atob(b64), c => c.charCodeAt(0));
}

async function _doSubscribe(reg) {
  if (!('PushManager' in window)) return;
  try {
    const kr = await fetch('/api/push/vapid-public-key');
    if (!kr.ok) return;
    const { publicKey } = await kr.json();
    if (!publicKey) return;
    const sub = await reg.pushManager.subscribe({
      userVisibleOnly: true,
      applicationServerKey: _urlB64ToUint8(publicKey),
    });
    await fetch('/api/push/subscribe', {
      method: 'POST', headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ subscription: sub.toJSON() }),
    });
    localStorage.setItem('mythic_push_subscribed', '1');
  } catch (err) { console.warn('[Push] subscribe error:', err); }
}

// Detect if running on iPhone/iOS — moved to top-level scope
// iOS Safari has supported Web Push since iOS 16.4, but ONLY when the site
// is installed to the Home Screen (running standalone) — never in a regular
// Safari tab. That's a hard OS restriction, not something code can bypass.
const _iosStandalone = isIOS && window.navigator.standalone === true;

if ('serviceWorker' in navigator && (!isIOS || _iosStandalone)) {
  navigator.serviceWorker.register('/sw.js', { scope: '/' })
    .then(reg => {
      _swReg = reg;
      if (Notification.permission === 'granted') {
        _hideBanner();
        _doSubscribe(reg);
      } else {
        _showBanner();
      }
    })
    .catch(err => {
      console.warn('[SW] registration failed:', err);
      _showBanner();
    });
} else if (isIOS) {
  // iPhone, but NOT installed to Home Screen yet: push is impossible from a
  // Safari tab (hard iOS restriction). Reuse the existing "Add to Home
  // Screen" install modal so the user has a clear path to actually get
  // notifications instead of silently getting nothing.
  console.log('[iOS] Push notifications need the app added to Home Screen first.');
  if (notifBanner && !localStorage.getItem('mythic_notif_dismissed')) {
    notifBanner.style.display = 'flex';
    if (notifAllowBtn) notifAllowBtn.textContent = '📲 Add to Home Screen';
  }
} else {
  _showBanner();
}

if (notifAllowBtn) notifAllowBtn.addEventListener('click', async () => {
  if (isIOS && !window.navigator.standalone) {
    // On iPhone in a plain Safari tab, there's no permission to request —
    // push literally cannot work until the app is installed. Send them to
    // the install flow instead of firing a no-op permission prompt.
    _hideBanner();
    _openInstallModal();
    return;
  }
  _hideBanner();
  let perm;
  try { perm = await Notification.requestPermission(); }
  catch { perm = 'denied'; }

  if (perm === 'granted') {
    if (_swReg) {
      try {
        await _swReg.showNotification('Mythic AI 🔔', {
          body: "Notifications enabled! You'll hear from me when your answer is ready.",
          icon: '/icon.png', badge: '/badge.png',
          tag: 'mythic-notif-confirm', vibrate: [150, 80, 150],
        });
      } catch (e) { console.warn('[Push] confirm notification failed:', e); }
      _doSubscribe(_swReg);
    } else {
      try { new Notification('Mythic AI 🔔', { body: "Notifications enabled!", icon: '/icon.png' }); }
      catch {}
    }
    const nb = document.getElementById('notif-toggle-btn');
    const ns = document.getElementById('notif-status');
    if (nb) { nb.textContent = 'Enabled ✓'; nb.style.borderColor = 'var(--accent)'; nb.style.color = 'var(--accent)'; }
    if (ns) ns.textContent = "You'll get notified when Mythic AI replies while you're away.";
    localStorage.removeItem('mythic_notif_dismissed');
  } else {
    const nb = document.getElementById('notif-toggle-btn');
    const ns = document.getElementById('notif-status');
    if (nb) { nb.textContent = 'Blocked'; nb.style.borderColor = '#ef4444'; nb.style.color = '#ef4444'; }
    if (ns) ns.textContent = 'Notifications blocked. Allow them in your browser site settings.';
  }
});

if (notifDismissBtn) notifDismissBtn.addEventListener('click', () => {
  _hideBanner();
  localStorage.setItem('mythic_notif_dismissed', String(Date.now() + 3 * 24 * 60 * 60 * 1000));
});

(function() {
  const ts = parseInt(localStorage.getItem('mythic_notif_dismissed') || '0', 10);
  if (ts && Date.now() > ts) localStorage.removeItem('mythic_notif_dismissed');
})();

window._notifyAiReply = function(preview) {
  if (document.visibilityState === 'visible') return;
  if (Notification.permission !== 'granted') return;
  const body = preview || 'Your answer is ready — tap to read it.';
  if (_swReg) {
    try {
      _swReg.showNotification('Mythic AI replied 💬', {
        body, icon: '/icon.png', badge: '/badge.png',
        tag: 'mythic-ai-reply', renotify: true, vibrate: [200, 100, 200],
        data: { url: '/' },
        actions: [{ action: 'open', title: '💬 Open Chat' }, { action: 'dismiss', title: '✕' }],
      });
    } catch {}
  } else {
    try { new Notification('Mythic AI replied 💬', { body, icon: '/icon.png' }); }
    catch {}
  }
};

const _origBuildActions = buildMsgActions;
buildMsgActions = function(row, textNode, role) {
  const actions = _origBuildActions(row, textNode, role);
  if (role === 'ai') {
    const reactBtn = document.createElement('button');
    reactBtn.type = 'button'; reactBtn.title = 'React'; reactBtn.textContent = '😊';
    reactBtn.addEventListener('click', () => addReactionBar(row));
    actions.appendChild(reactBtn);
    const sp = document.createElement('button');
    sp.type='button'; sp.title='Read aloud'; sp.textContent='🔊';
    sp.addEventListener('click', () => {
      if (sp.textContent === '⏹') { stopSpeaking(); sp.textContent='🔊'; return; }
      sp.textContent='⏹'; speak(textNode.textContent || (textNode.innerText || ''));
      if (currentUtterance) {
        currentUtterance.onend = () => sp.textContent='🔊';
        currentUtterance.onerror = () => sp.textContent='🔊';
      }
    });
    actions.appendChild(sp);
  }
  return actions;
};

function stopSpeaking() { if(window.speechSynthesis) window.speechSynthesis.cancel(); }

function getTonePrefix() {
  const s = JSON.parse(localStorage.getItem('mythic_settings') || '{}');
  const tone = s.tone || 'default';
  const length = s.length || 'default';
  let parts = [];
  if (tone === 'formal') parts.push('Reply in a formal, professional tone.');
  if (tone === 'casual') parts.push('Reply in a casual, friendly tone.');
  if (tone === 'funny') parts.push('Be funny and use wit and humor in your reply.');
  if (tone === 'professional') parts.push('Use a professional, business-appropriate tone.');
  if (length === 'short') parts.push('Keep your reply very short — 1-3 sentences max.');
  if (length === 'medium') parts.push('Keep your reply medium length — a few paragraphs.');
  if (length === 'long') parts.push('Give a thorough, detailed, long reply.');
  const ci = customInstructions ? customInstructions.value.trim() : '';
  if (ci) parts.push(ci);
  return parts.length ? '[Instructions: ' + parts.join(' ') + '] ' : '';
}

const _origFormSubmit = form.onsubmit;
form.addEventListener('submit', async () => {
  const checkDone = setInterval(() => {
    if (!isGenerating) {
      clearInterval(checkDone);
      const allRows = messagesEl.querySelectorAll('.msg-row.ai');
      if (allRows.length) {
        const lastRow = allRows[allRows.length - 1];
        const textEl = lastRow.querySelector('.msg-text,.md-rendered');
        if (textEl && !lastRow.querySelector('.reaction-bar')) {
          addReactionBar(lastRow);
          setTimeout(() => addFollowupSuggestions(textEl.textContent || textEl.innerText || ''), 300);
        }
      }
    }
  }, 500);
});

(async () => {
  await loadConversationList();

  // If we were redirected here from a "Continue this conversation" share
  // link (see SHARE_PAGE), open that freshly-forked conversation instead
  // of the usual blank New Chat screen. Also support a plain "?c=<id>"
  // link/refresh/bookmark — this is the id openConversation() itself now
  // puts in the address bar, so refreshing or sharing that URL reopens
  // the same chat instead of always landing on a blank screen.
  const params = new URLSearchParams(location.search);
  const openId = params.get('open') || params.get('c');
  const action = params.get('action');
  if (openId) {
    // Normalize the URL to ?c=<id> and open without pushing a duplicate
    // history entry (replaceState keeps Back from bouncing between the
    // raw ?open= link and the normalized ?c= one).
    history.replaceState({ conv: openId }, '', '?c=' + encodeURIComponent(openId));
    await openConversation(openId, { updateUrl: false });
  } else if (action === 'new') {
    // Opened via the "New Chat" home-screen long-press shortcut.
    history.replaceState({}, '', location.pathname);
    startNewChat();
  } else if (action === 'image') {
    // Opened via the "Generate Image" home-screen long-press shortcut —
    // jump straight into image mode instead of the plain chat screen.
    history.replaceState({}, '', location.pathname);
    showEmptyState();
    const imgBtn = document.getElementById('img-gen-btn');
    if (imgBtn) imgBtn.click();
  } else {
    // Always start on a fresh New Chat screen otherwise, rather than
    // auto-reopening the last conversation — the sidebar list is still
    // populated underneath.
    showEmptyState();
  }

  // Silently remove stray internal-tooling conversations (old follow-up-
  // suggestion / instruction-prefix leaks) in the background, no button
  // or confirmation needed — safe because it only ever matches those very
  // specific known patterns server-side (see _JUNK_CONV_PATTERNS).
  try {
    const r = await fetch('/api/conversations/cleanup-junk', { method: 'POST' });
    const d = await r.json();
    if (d && d.removed_count > 0) loadConversationList();
  } catch {}
})();

const imgGenBtn   = document.getElementById('img-gen-btn');
const ghibliBtn   = document.getElementById('ghibli-btn');
const fileGenBtn  = document.getElementById('file-gen-btn');
const homeworkBtn = document.getElementById('homework-btn');
const weatherBtn2 = document.getElementById('weather-btn');
const searchBtn   = document.getElementById('search-btn');

if (imgGenBtn) imgGenBtn.addEventListener('click', () => {
  const imgModal = document.getElementById('img-modal-overlay');
  if (imgModal) { imgModal.style.display = 'flex'; document.getElementById('img-prompt').focus(); }
});
if (homeworkBtn) homeworkBtn.addEventListener('click', () => {
  const hwModal = document.getElementById('homework-modal-overlay');
  if (hwModal) { hwModal.style.display = 'flex'; document.getElementById('homework-question').focus(); }
});
if (weatherBtn2) weatherBtn2.addEventListener('click', () => {
  const wm = document.getElementById('weather-modal-overlay');
  if (wm) { wm.style.display = 'flex'; document.getElementById('weather-city').focus(); }
});
if (searchBtn) searchBtn.addEventListener('click', () => {
  const q = prompt('What do you want to search for?');
  if (!q || !q.trim()) return;
  input.value = 'Search: ' + q.trim(); autoResize(); form.requestSubmit();
});

// ─── FILE / PDF GENERATION MODAL JS ───────────────────────────────────────────
const fileModal        = document.getElementById('file-modal-overlay');
const fileTitleInput   = document.getElementById('file-title-input');
const fileContentInput = document.getElementById('file-content-input');
const fileGenerateBtn  = document.getElementById('file-generate-btn');
const fileCloseBtn     = document.getElementById('file-close-btn');
const fileLoadingEl    = document.getElementById('file-loading');
const fileErrorEl      = document.getElementById('file-error');
const fileNoteEl       = document.getElementById('file-note');
let selectedFileFormat = 'pdf';

function openFileModalWithContent(content) {
  fileModal.style.display = 'flex';
  fileContentInput.value = content || '';
  fileErrorEl.style.display = 'none';
  fileNoteEl.style.display = 'none';
}
if (fileGenBtn) fileGenBtn.addEventListener('click', () => openFileModalWithContent(_lastUserMessageText ? '' : ''));
if (fileCloseBtn) fileCloseBtn.addEventListener('click', () => fileModal.style.display = 'none');
if (fileModal) fileModal.addEventListener('click', e => { if (e.target === fileModal) fileModal.style.display = 'none'; });

document.querySelectorAll('.file-format-btn').forEach(btn => {
  btn.addEventListener('click', () => {
    document.querySelectorAll('.file-format-btn').forEach(b => {
      b.style.borderColor = 'var(--border)'; b.style.background = 'var(--panel)'; b.style.color = 'var(--muted)';
    });
    btn.style.borderColor = 'var(--accent)'; btn.style.background = 'var(--accent-dim)'; btn.style.color = 'var(--accent)';
    selectedFileFormat = btn.dataset.format;
  });
});

if (fileGenerateBtn) fileGenerateBtn.addEventListener('click', async () => {
  const content = fileContentInput.value.trim();
  if (!content) { fileErrorEl.textContent = 'Please add some content first.'; fileErrorEl.style.display = 'block'; return; }
  fileErrorEl.style.display = 'none'; fileNoteEl.style.display = 'none';
  fileLoadingEl.style.display = 'block'; fileGenerateBtn.disabled = true;
  try {
    const r = await fetch('/api/generate-file', {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify({ content, format: selectedFileFormat, title: fileTitleInput.value.trim() || 'Mythic AI Document' })
    });
    const d = await r.json();
    fileLoadingEl.style.display = 'none';
    if (d.file) {
      const blob = _b64ToBlob(d.file, d.mimeType);
      const url = URL.createObjectURL(blob);
      const a = document.createElement('a');
      a.href = url; a.download = d.filename;
      document.body.appendChild(a); a.click(); a.remove();
      URL.revokeObjectURL(url);
      if (d.note) { fileNoteEl.textContent = d.note; fileNoteEl.style.display = 'block'; }
      else { fileModal.style.display = 'none'; }
    } else {
      fileErrorEl.textContent = d.error || 'File generation failed. Try again.';
      fileErrorEl.style.display = 'block';
    }
  } catch (e) {
    fileLoadingEl.style.display = 'none';
    fileErrorEl.textContent = 'Network error: ' + e.message;
    fileErrorEl.style.display = 'block';
  } finally { fileGenerateBtn.disabled = false; }
});

const ghibliModal     = document.getElementById('ghibli-modal-overlay');
const ghibliUploadArea= document.getElementById('ghibli-upload-area');
const ghibliFileInput = document.getElementById('ghibli-file-input');
const ghibliPreviewWrap=document.getElementById('ghibli-preview-wrap');
const ghibliPreview   = document.getElementById('ghibli-preview');
const ghibliResult    = document.getElementById('ghibli-result');
const ghibliResultWrap= document.getElementById('ghibli-result-wrap');
const ghibliLoading   = document.getElementById('ghibli-loading');
const ghibliError     = document.getElementById('ghibli-error');
const ghibliGenerateBtn=document.getElementById('ghibli-generate-btn');
const ghibliCloseBtn  = document.getElementById('ghibli-close-btn');
const ghibliDownloadBtn=document.getElementById('ghibli-download-btn');
const ghibliExtraInput= document.getElementById('ghibli-extra');

let ghibliBase64 = null;
let ghibliMimeType = 'image/jpeg';
let ghibliSelectedStyle = 'Studio Ghibli portrait, Spirited Away style, soft watercolor anime art';

if (ghibliBtn) ghibliBtn.addEventListener('click', () => {
  ghibliModal.style.display = 'flex';
  ghibliBase64 = null;
  ghibliPreviewWrap.style.display = 'none';
  ghibliResultWrap.style.display = 'none';
  ghibliError.style.display = 'none';
  ghibliLoading.style.display = 'none';
});
if (ghibliCloseBtn) ghibliCloseBtn.addEventListener('click', () => ghibliModal.style.display = 'none');
ghibliModal.addEventListener('click', e => { if (e.target === ghibliModal) ghibliModal.style.display = 'none'; });

document.querySelectorAll('.ghibli-style-btn').forEach(btn => {
  btn.addEventListener('click', () => {
    document.querySelectorAll('.ghibli-style-btn').forEach(b => {
      b.style.borderColor = 'var(--border)'; b.style.background = 'var(--panel)'; b.style.color = 'var(--muted)';
    });
    btn.style.borderColor = 'var(--accent)'; btn.style.background = 'var(--accent-dim)'; btn.style.color = 'var(--accent)';
    ghibliSelectedStyle = btn.dataset.style;
  });
});

ghibliUploadArea.addEventListener('click', () => ghibliFileInput.click());
ghibliUploadArea.addEventListener('dragover', e => { e.preventDefault(); ghibliUploadArea.style.borderColor = 'var(--accent)'; });
ghibliUploadArea.addEventListener('dragleave', () => { ghibliUploadArea.style.borderColor = 'var(--border)'; });
ghibliUploadArea.addEventListener('drop', e => {
  e.preventDefault(); ghibliUploadArea.style.borderColor = 'var(--border)';
  const file = e.dataTransfer.files[0];
  if (file && file.type.startsWith('image/')) loadGhibliPhoto(file);
});
ghibliFileInput.addEventListener('change', () => {
  if (ghibliFileInput.files[0]) loadGhibliPhoto(ghibliFileInput.files[0]);
});

function loadGhibliPhoto(file) {
  const reader = new FileReader();
  reader.onload = e => {
    ghibliBase64 = e.target.result.split(',')[1];
    ghibliMimeType = file.type || 'image/jpeg';
    ghibliPreview.src = e.target.result;
    ghibliPreviewWrap.style.display = 'block';
    ghibliResultWrap.style.display = 'none';
    ghibliError.style.display = 'none';
    ghibliUploadArea.style.borderColor = 'var(--accent)';
  };
  reader.readAsDataURL(file);
}

ghibliGenerateBtn.addEventListener('click', async () => {
  const extra = ghibliExtraInput.value.trim();
  const prompt = `${ghibliSelectedStyle}, beautiful detailed portrait of a person, ${extra ? extra + ', ' : ''}masterpiece, best quality, highly detailed, cinematic lighting, soft colors, dreamy atmosphere`;

  ghibliError.style.display = 'none';
  ghibliResultWrap.style.display = 'none';
  ghibliLoading.style.display = 'block';
  ghibliGenerateBtn.disabled = true;

  const bodyPayload = { prompt };
  if (ghibliBase64) {
    bodyPayload.imageBase64 = ghibliBase64;
    bodyPayload.mimeType = ghibliMimeType;
  }

  try {
    const r = await fetch('/api/generate-image', {
      method: 'POST', headers: {'Content-Type':'application/json'},
      body: JSON.stringify(bodyPayload)
    });
    const text = await r.text();
    let d;
    try { d = JSON.parse(text); }
    catch { d = {error: `Server error (${r.status}). Please try again in a moment.`}; }
    ghibliLoading.style.display = 'none';
    if (d.image) {
      ghibliResult.src = 'data:image/png;base64,' + d.image;
      ghibliResultWrap.style.display = 'block';
      clearEmptyState();
      const row = document.createElement('div'); row.className = 'msg-row ai';
      const bubble = document.createElement('div'); bubble.className = 'msg ai';
      bubble.style.padding = '8px';
      const cap = document.createElement('div');
      cap.textContent = '🌿 Your Ghibli portrait';
      cap.style.cssText = 'font-size:12px;color:var(--muted);margin-bottom:8px;';
      const img = document.createElement('img');
      img.src = 'data:image/png;base64,' + d.image;
      img.style.cssText = 'max-width:100%;border-radius:12px;display:block;cursor:pointer;';
      img.title = 'Click to download';
      img.addEventListener('click', () => downloadGhibliImage(d.image));
      bubble.appendChild(cap); bubble.appendChild(img); row.appendChild(bubble);
      messagesEl.appendChild(row); scrollToBottom();
    } else {
      ghibliError.textContent = d.error || 'Generation failed. Try again.';
      ghibliError.style.display = 'block';
    }
  } catch (e) {
    ghibliLoading.style.display = 'none';
    ghibliError.textContent = 'Network error: ' + e.message;
    ghibliError.style.display = 'block';
  } finally { ghibliGenerateBtn.disabled = false; }
});

function downloadGhibliImage(b64) {
  const a = document.createElement('a');
  a.href = 'data:image/png;base64,' + b64;
  a.download = 'mythic-ai-ghibli-portrait.png';
  document.body.appendChild(a); a.click(); a.remove();
}
if (ghibliDownloadBtn) ghibliDownloadBtn.addEventListener('click', () => {
  if (ghibliResult.src) downloadGhibliImage(ghibliResult.src.split(',')[1]);
});

const imgModalOverlay = document.getElementById('img-modal-overlay');
const imgPromptEl     = document.getElementById('img-prompt');
const imgStyleEl      = document.getElementById('img-style');
const imgResultEl     = document.getElementById('img-result');
const imgOutputEl     = document.getElementById('img-output');
const imgLoadingEl    = document.getElementById('img-loading');
const imgErrorEl      = document.getElementById('img-error');
const imgGenerateBtn2 = document.getElementById('img-generate-btn');
const imgCloseBtn2    = document.getElementById('img-close-btn');
const imgDownloadBtn2 = document.getElementById('img-download-btn');
const imgCopyBtn2     = document.getElementById('img-copy-btn');
const imgFullscreenBtn2 = document.getElementById('img-fullscreen-btn');
const imgViewerOverlay = document.getElementById('img-viewer-overlay');
const imgViewerImg     = document.getElementById('img-viewer-img');
let lastGeneratedImageB64 = null;

if (imgGenerateBtn2) imgGenerateBtn2.addEventListener('click', async () => {
  const prompt = imgPromptEl.value.trim();
  const style = imgStyleEl ? imgStyleEl.value : '';
  if (!prompt) { imgErrorEl.textContent = 'Please enter a description first.'; imgErrorEl.style.display = 'block'; return; }
  imgResultEl.style.display = 'none'; imgErrorEl.style.display = 'none';
  imgLoadingEl.style.display = 'block'; imgGenerateBtn2.disabled = true;
  try {
    const r = await fetch('/api/generate-image', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({prompt, style})
    });
    const text = await r.text();
    let d;
    try { d = JSON.parse(text); }
    catch { d = {error: `Server error (${r.status}). Please try again in a moment.`}; }
    imgLoadingEl.style.display = 'none';
    if (d.image) {
      lastGeneratedImageB64 = d.image;
      imgOutputEl.src = 'data:image/png;base64,' + d.image;
      imgResultEl.style.display = 'block';
      clearEmptyState();
      const row = document.createElement('div'); row.className = 'msg-row ai';
      const bubble = document.createElement('div'); bubble.className = 'msg ai'; bubble.style.padding='8px';
      const cap = document.createElement('div'); cap.textContent = '🎨 ' + prompt;
      cap.style.cssText = 'font-size:12px;opacity:.7;margin-bottom:8px;';
      const img = document.createElement('img');
      img.src = 'data:image/png;base64,' + d.image;
      img.style.cssText = 'max-width:100%;border-radius:10px;display:block;';
      bubble.appendChild(cap); bubble.appendChild(img); row.appendChild(bubble);
      messagesEl.appendChild(row); scrollToBottom();
    } else {
      imgErrorEl.textContent = d.error || 'Image generation failed. Try again.';
      imgErrorEl.style.display = 'block';
    }
  } catch(e) {
    imgLoadingEl.style.display = 'none';
    imgErrorEl.textContent = 'Connection error: ' + e.message + '. Check your internet and try again.';
    imgErrorEl.style.display = 'block';
  }
  finally { imgGenerateBtn2.disabled = false; }
});
if (imgCloseBtn2) imgCloseBtn2.addEventListener('click', () => imgModalOverlay.style.display='none');
if (imgModalOverlay) imgModalOverlay.addEventListener('click', e => { if(e.target===imgModalOverlay) imgModalOverlay.style.display='none'; });
if (imgPromptEl) imgPromptEl.addEventListener('keydown', e => { if(e.key==='Enter'&&!e.shiftKey){e.preventDefault();imgGenerateBtn2.click();} });
if (imgDownloadBtn2) imgDownloadBtn2.addEventListener('click', () => {
  if (!lastGeneratedImageB64) return;
  const a = document.createElement('a');
  a.href = 'data:image/png;base64,' + lastGeneratedImageB64;
  a.download = 'mythic-ai-image-' + Date.now() + '.png';
  document.body.appendChild(a); a.click(); a.remove();
});
if (imgCopyBtn2) imgCopyBtn2.addEventListener('click', async () => {
  if (!lastGeneratedImageB64) return;
  try {
    const res = await fetch('data:image/png;base64,' + lastGeneratedImageB64);
    const blob = await res.blob();
    await navigator.clipboard.write([new ClipboardItem({ [blob.type]: blob })]);
    const orig = imgCopyBtn2.textContent;
    imgCopyBtn2.textContent = '✓ Copied';
    setTimeout(() => { imgCopyBtn2.textContent = orig; }, 1200);
  } catch { imgErrorEl.textContent = 'Copy not supported in this browser.'; imgErrorEl.style.display = 'block'; }
});
if (imgFullscreenBtn2) imgFullscreenBtn2.addEventListener('click', () => {
  if (!lastGeneratedImageB64) return;
  imgViewerImg.src = 'data:image/png;base64,' + lastGeneratedImageB64;
  imgViewerOverlay.style.display = 'flex';
});
if (imgViewerOverlay) imgViewerOverlay.addEventListener('click', () => imgViewerOverlay.style.display = 'none');

const weatherModal2    = document.getElementById('weather-modal-overlay');
const weatherCityEl2   = document.getElementById('weather-city');
const weatherResultEl2 = document.getElementById('weather-result');
const weatherContentEl2= document.getElementById('weather-content');
const weatherLoadingEl2= document.getElementById('weather-loading');
const weatherErrorEl2  = document.getElementById('weather-error');
const weatherSearchBtn2= document.getElementById('weather-search-btn');
const weatherCloseBtn2 = document.getElementById('weather-close-btn');
const weatherLocBtn2   = document.getElementById('weather-location-btn');

function getRecentWeatherSearches() {
  try { return JSON.parse(localStorage.getItem('mythic_recent_weather') || '[]'); } catch { return []; }
}
function addRecentWeatherSearch(name) {
  let list = getRecentWeatherSearches().filter(x => x.toLowerCase() !== name.toLowerCase());
  list.unshift(name);
  list = list.slice(0, 6);
  localStorage.setItem('mythic_recent_weather', JSON.stringify(list));
}
function renderRecentSearches() {
  const recents = getRecentWeatherSearches();
  const wrap = document.getElementById('weather-recents');
  if (!recents.length) { if (wrap) wrap.remove(); return; }
  let el = wrap;
  if (!el) {
    el = document.createElement('div');
    el.id = 'weather-recents';
    el.style.cssText = 'display:flex;gap:6px;flex-wrap:wrap;margin-bottom:12px;';
    weatherCityEl2.parentNode.insertAdjacentElement('afterend', el);
  }
  el.innerHTML = '';
  recents.forEach(name => {
    const chip = document.createElement('button');
    chip.textContent = name;
    chip.style.cssText = 'background:var(--bg);border:1px solid var(--border);color:var(--muted);border-radius:20px;padding:5px 12px;font-size:12px;cursor:pointer;font-family:inherit;';
    chip.addEventListener('click', () => fetchWeatherModal({ location: name }));
    el.appendChild(chip);
  });
}

function renderWeather(w) {
  const hourly = (w.hourly || []).map(h => `
    <div style="flex:0 0 auto;text-align:center;background:var(--bg);border-radius:10px;padding:8px 12px;min-width:64px;">
      <div style="font-size:11px;color:var(--muted);">${h.time}</div>
      <div style="font-size:20px;margin:4px 0;">${h.icon}</div>
      <div style="font-size:13px;font-weight:700;">${h.temp}°</div>
    </div>`).join('');

  const daily = (w.daily || []).map(d => `
    <div style="display:flex;align-items:center;justify-content:space-between;padding:8px 4px;border-bottom:1px solid var(--border);">
      <div style="flex:1;font-size:13px;">${d.day}</div>
      <div style="font-size:18px;">${d.icon}</div>
      <div style="flex:1;text-align:right;font-size:13px;"><span style="font-weight:700;">${d.max}°</span> <span style="color:var(--muted);">${d.min}°</span></div>
    </div>`).join('');

  weatherContentEl2.innerHTML = `
    <div style="display:flex;align-items:center;gap:12px;margin-bottom:14px;">
      <div style="font-size:52px;">${w.icon}</div>
      <div><div style="font-size:19px;font-weight:700;">${w.location}</div>
      <div style="font-size:13px;color:var(--muted);">${w.condition}</div></div>
      <div style="margin-left:auto;font-size:32px;font-weight:700;">${w.temp}°C</div>
    </div>

    <div style="display:grid;grid-template-columns:1fr 1fr 1fr;gap:8px;margin-bottom:12px;">
      <div style="background:var(--bg);border-radius:8px;padding:9px;text-align:center;">
        <div style="font-size:10px;color:var(--muted);">FEELS LIKE</div>
        <div style="font-size:16px;font-weight:700;">${w.feels_like}°C</div>
      </div>
      <div style="background:var(--bg);border-radius:8px;padding:9px;text-align:center;">
        <div style="font-size:10px;color:var(--muted);">HUMIDITY</div>
        <div style="font-size:16px;font-weight:700;">${w.humidity}%</div>
      </div>
      <div style="background:var(--bg);border-radius:8px;padding:9px;text-align:center;">
        <div style="font-size:10px;color:var(--muted);">WIND</div>
        <div style="font-size:16px;font-weight:700;">${w.wind_speed} km/h</div>
      </div>
      <div style="background:var(--bg);border-radius:8px;padding:9px;text-align:center;">
        <div style="font-size:10px;color:var(--muted);">UV INDEX</div>
        <div style="font-size:16px;font-weight:700;">${w.uv ?? '–'}</div>
      </div>
      <div style="background:var(--bg);border-radius:8px;padding:9px;text-align:center;">
        <div style="font-size:10px;color:var(--muted);">PRESSURE</div>
        <div style="font-size:16px;font-weight:700;">${w.pressure ?? '–'} hPa</div>
      </div>
      <div style="background:var(--bg);border-radius:8px;padding:9px;text-align:center;">
        <div style="font-size:10px;color:var(--muted);">VISIBILITY</div>
        <div style="font-size:16px;font-weight:700;">${w.visibility ?? '–'} km</div>
      </div>
      ${w.aqi != null ? `
      <div style="background:var(--bg);border-radius:8px;padding:9px;text-align:center;">
        <div style="font-size:10px;color:var(--muted);">AIR QUALITY</div>
        <div style="font-size:16px;font-weight:700;">${w.aqi}</div>
      </div>` : ''}
      <div style="background:var(--bg);border-radius:8px;padding:9px;text-align:center;">
        <div style="font-size:10px;color:var(--muted);">SUNRISE</div>
        <div style="font-size:14px;font-weight:700;">${w.sunrise ?? '–'}</div>
      </div>
      <div style="background:var(--bg);border-radius:8px;padding:9px;text-align:center;">
        <div style="font-size:10px;color:var(--muted);">SUNSET</div>
        <div style="font-size:14px;font-weight:700;">${w.sunset ?? '–'}</div>
      </div>
    </div>

    ${hourly ? `<div style="font-size:12px;color:var(--muted);margin-bottom:6px;">HOURLY FORECAST</div>
    <div style="display:flex;gap:8px;overflow-x:auto;padding-bottom:6px;margin-bottom:12px;">${hourly}</div>` : ''}

    ${daily ? `<div style="font-size:12px;color:var(--muted);margin-bottom:6px;">7-DAY FORECAST</div>
    <div style="margin-bottom:6px;">${daily}</div>` : ''}
  `;
  weatherResultEl2.style.display = 'block';
  addRecentWeatherSearch(w.location);
  renderRecentSearches();
}

async function fetchWeatherModal(payload) {
  weatherResultEl2.style.display='none'; weatherErrorEl2.style.display='none';
  weatherLoadingEl2.style.display='block'; weatherSearchBtn2.disabled=true;
  try {
    const r = await fetch('/api/weather',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify(payload)});
    const d = await r.json(); weatherLoadingEl2.style.display='none';
    if(d.weather) renderWeather(d.weather);
    else { weatherErrorEl2.textContent=d.error||'Could not find that location. Try a different search.'; weatherErrorEl2.style.display='block'; }
  } catch(e) { weatherLoadingEl2.style.display='none'; weatherErrorEl2.textContent='Network error: '+e.message; weatherErrorEl2.style.display='block'; }
  finally { weatherSearchBtn2.disabled=false; }
}
if (weatherSearchBtn2) weatherSearchBtn2.addEventListener('click', () => { const loc=weatherCityEl2.value.trim(); if(loc) fetchWeatherModal({location:loc}); });
if (weatherCityEl2) weatherCityEl2.addEventListener('keydown', e => { if(e.key==='Enter') weatherSearchBtn2.click(); });
if (weatherCloseBtn2) weatherCloseBtn2.addEventListener('click', () => weatherModal2.style.display='none');
if (weatherModal2) weatherModal2.addEventListener('click', e => { if(e.target===weatherModal2) weatherModal2.style.display='none'; });
if (weatherLocBtn2) weatherLocBtn2.addEventListener('click', () => {
  if (!navigator.geolocation) { weatherErrorEl2.textContent = 'Geolocation is not supported in this browser.'; weatherErrorEl2.style.display = 'block'; return; }
  navigator.geolocation.getCurrentPosition(
    pos => fetchWeatherModal({lat:pos.coords.latitude,lon:pos.coords.longitude}),
    err => { weatherErrorEl2.textContent = 'Location error: ' + err.message; weatherErrorEl2.style.display = 'block'; }
  );
});
renderRecentSearches();

// ─── HOMEWORK & STUDY BOOK MODAL ─────────────────────────────────────────────
(function() {
  const modal       = document.getElementById('homework-modal-overlay');

  const closeBtn     = document.getElementById('hw-close-btn');
  const sendBtn      = document.getElementById('hw-send-btn');
  const questionEl   = document.getElementById('homework-question');
  const modeUploadBtn= document.getElementById('hw-mode-upload');
  const modeUrlBtn   = document.getElementById('hw-mode-url');
  const uploadArea   = document.getElementById('hw-upload-area');
  const urlArea      = document.getElementById('hw-url-area');
  const dropzone     = document.getElementById('hw-upload-dropzone');
  const hwFileInput  = document.getElementById('hw-file-input');
  const hwFileName   = document.getElementById('hw-file-name');
  const urlInput     = document.getElementById('hw-url-input');
  const loadingEl    = document.getElementById('hw-loading');
  const errorEl      = document.getElementById('hw-error');
  if (!modal) return;

  let hwMode = 'upload';
  let hwPendingFile = null; // { name, mimeType, dataBase64 }

  function setHwMode(mode) {
    hwMode = mode;
    modeUploadBtn.style.borderColor = mode === 'upload' ? 'var(--accent)' : 'var(--border)';
    modeUploadBtn.style.background  = mode === 'upload' ? 'var(--accent-dim)' : 'var(--panel)';
    modeUploadBtn.style.color       = mode === 'upload' ? 'var(--accent)' : 'var(--muted)';
    modeUrlBtn.style.borderColor    = mode === 'url' ? 'var(--accent)' : 'var(--border)';
    modeUrlBtn.style.background     = mode === 'url' ? 'var(--accent-dim)' : 'var(--panel)';
    modeUrlBtn.style.color          = mode === 'url' ? 'var(--accent)' : 'var(--muted)';
    uploadArea.style.display = mode === 'upload' ? 'block' : 'none';
    urlArea.style.display    = mode === 'url' ? 'block' : 'none';
  }
  modeUploadBtn.addEventListener('click', () => setHwMode('upload'));
  modeUrlBtn.addEventListener('click', () => setHwMode('url'));

  dropzone.addEventListener('click', () => hwFileInput.click());
  hwFileInput.addEventListener('change', () => {
    const file = hwFileInput.files[0];
    if (!file) return;
    const reader = new FileReader();
    reader.onload = (e) => {
      hwPendingFile = { name: file.name, mimeType: file.type || 'application/octet-stream', dataBase64: e.target.result.split(',')[1] };
      hwFileName.textContent = '📄 ' + file.name;
      hwFileName.style.display = 'block';
    };
    reader.readAsDataURL(file);
  });

  if (homeworkBtn) homeworkBtn.addEventListener('click', () => {
    errorEl.style.display = 'none';
    loadingEl.style.display = 'none';
  });
  if (closeBtn) closeBtn.addEventListener('click', () => { modal.style.display = 'none'; });
  modal.addEventListener('click', (e) => { if (e.target === modal) modal.style.display = 'none'; });

  const STUDY_INSTRUCTIONS = "First list the chapters you can see in this document and tell me which chapter you're focusing on. Then prepare the hardest, most exam-relevant questions for that chapter first, before easier ones.";

  sendBtn.addEventListener('click', async () => {
    errorEl.style.display = 'none';
    const question = questionEl.value.trim();

    if (hwMode === 'upload') {
      if (!hwPendingFile && !question) {
        errorEl.textContent = 'Upload a file or type a question first.';
        errorEl.style.display = 'block';
        return;
      }
      const finalMessage = question || STUDY_INSTRUCTIONS;
      modal.style.display = 'none';
      addMessage('user', finalMessage, hwPendingFile);
      const attachmentToSend = hwPendingFile;
      hwPendingFile = null;
      hwFileName.style.display = 'none';
      hwFileInput.value = '';
      questionEl.value = '';
      streamReply({ message: getTonePrefix() + finalMessage, attachment: attachmentToSend });
      return;
    }

    // URL mode — fetch + extract text server-side, then send as a normal
    // text message with the extracted content inlined (same 12k-char cap
    // the file-upload path already uses via extract_text_from_attachment).
    const url = urlInput.value.trim();
    if (!url && !question) {
      errorEl.textContent = 'Paste a URL or type a question first.';
      errorEl.style.display = 'block';
      return;
    }
    if (!url) {
      // No URL, just a plain question — send as normal chat.
      modal.style.display = 'none';
      addMessage('user', question);
      questionEl.value = '';
      streamReply({ message: getTonePrefix() + question });
      return;
    }

    loadingEl.style.display = 'block';
    sendBtn.disabled = true;
    try {
      const r = await fetch('/api/fetch-url-document', {
        method: 'POST', headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ url })
      });
      const d = await r.json();
      loadingEl.style.display = 'none';
      if (!r.ok || d.error) {
        errorEl.textContent = d.error || 'Could not fetch that URL.';
        errorEl.style.display = 'block';
        return;
      }
      const finalMessage = (question || STUDY_INSTRUCTIONS)
        + `\n\n[Document fetched from ${d.filename || url}]\n${d.text}`
        + (d.note ? `\n[Note: ${d.note}]` : '');
      modal.style.display = 'none';
      addMessage('user', question || `📚 Studying: ${d.filename || url}`);
      urlInput.value = '';
      questionEl.value = '';
      streamReply({ message: getTonePrefix() + finalMessage });
    } catch (e) {
      loadingEl.style.display = 'none';
      errorEl.textContent = 'Network error: ' + e.message;
      errorEl.style.display = 'block';
    } finally {
      sendBtn.disabled = false;
    }
  });
})();

// ─── CODE WORKSPACE — HTML/CSS/JS live preview, plus Python/C++/C/Java/
//     Node.js/TypeScript/Go/Ruby executed remotely via /api/code/run ───────
(function() {
  const codeBtn        = document.getElementById('code-workspace-btn');
  const codeModal       = document.getElementById('code-modal-overlay');
  const closeBtn        = document.getElementById('code-close-btn');
  const runBtn          = document.getElementById('code-run-btn');
  const downloadBtn     = document.getElementById('code-download-btn');
  const fullscreenBtn   = document.getElementById('code-fullscreen-preview-btn');
  const projectNameInput= document.getElementById('code-project-name');
  const previewFrame    = document.getElementById('code-preview-frame');
  const outputConsole   = document.getElementById('code-output-console');
  const previewLabel    = document.getElementById('code-preview-label');
  const webTabs         = document.getElementById('code-web-tabs');
  const singleTabLabel  = document.getElementById('code-single-tab');
  const singleEditor    = document.getElementById('code-editor-single');
  const stdinBox         = document.getElementById('code-stdin');
  const editors = {
    html: document.getElementById('code-editor-html'),
    css:  document.getElementById('code-editor-css'),
    js:   document.getElementById('code-editor-js'),
  };
  const tabs = document.querySelectorAll('.code-file-tab');
  const langBtns = document.querySelectorAll('.code-lang-btn');
  if (!codeBtn) return;

  // Per-language metadata: display name, file extension, and starter code
  // shown the first time someone switches to that language.
  const LANGS = {
    python:     { label: 'Python',     ext: 'py',  starter: 'print("Hello from Mythic AI Code Workspace!")\n' },
    cpp:        { label: 'C++',        ext: 'cpp', starter: '#include <iostream>\n\nint main() {\n  std::cout << "Hello from Mythic AI Code Workspace!" << std::endl;\n  return 0;\n}\n' },
    c:          { label: 'C',          ext: 'c',   starter: '#include <stdio.h>\n\nint main() {\n  printf("Hello from Mythic AI Code Workspace!\\n");\n  return 0;\n}\n' },
    java:       { label: 'Java',       ext: 'java',starter: 'public class Main {\n  public static void main(String[] args) {\n    System.out.println("Hello from Mythic AI Code Workspace!");\n  }\n}\n' },
    javascript: { label: 'Node.js',    ext: 'js',  starter: 'console.log("Hello from Mythic AI Code Workspace!");\n' },
    typescript: { label: 'TypeScript', ext: 'ts',  starter: 'const message: string = "Hello from Mythic AI Code Workspace!";\nconsole.log(message);\n' },
    go:         { label: 'Go',         ext: 'go',  starter: 'package main\n\nimport "fmt"\n\nfunc main() {\n  fmt.Println("Hello from Mythic AI Code Workspace!")\n}\n' },
    ruby:       { label: 'Ruby',       ext: 'rb',  starter: 'puts "Hello from Mythic AI Code Workspace!"\n' },
  };

  let currentLang = 'web';
  const singleCode = {}; // language -> code, populated lazily from LANGS starters

  const STORE_KEY = 'mythic_code_workspace';
  function saveDraft() {
    try {
      if (currentLang !== 'web') singleCode[currentLang] = singleEditor.value;
      localStorage.setItem(STORE_KEY, JSON.stringify({
        html: editors.html.value, css: editors.css.value, js: editors.js.value,
        name: projectNameInput.value,
        lang: currentLang,
        singleCode,
        stdin: stdinBox.value,
      }));
    } catch {}
  }
  function loadDraft() {
    try {
      const saved = JSON.parse(localStorage.getItem(STORE_KEY) || 'null');
      if (saved) {
        editors.html.value = saved.html ?? editors.html.value;
        editors.css.value  = saved.css  ?? editors.css.value;
        editors.js.value   = saved.js   ?? editors.js.value;
        projectNameInput.value = saved.name || 'my-project';
        Object.assign(singleCode, saved.singleCode || {});
        stdinBox.value = saved.stdin || '';
        if (saved.lang && saved.lang !== 'web') switchLang(saved.lang, /*skipSave*/ true);
      }
    } catch {}
  }

  function switchLang(lang, skipSave) {
    if (lang === currentLang) return;
    if (currentLang !== 'web') singleCode[currentLang] = singleEditor.value;
    currentLang = lang;

    langBtns.forEach(b => {
      const active = b.dataset.lang === lang;
      b.classList.toggle('active', active);
      b.style.background = active ? 'var(--accent-dim)' : 'none';
      b.style.color = active ? 'var(--accent)' : 'var(--muted)';
      b.style.borderColor = active ? 'var(--accent)' : 'var(--border)';
    });

    if (lang === 'web') {
      webTabs.style.display = 'flex';
      singleTabLabel.style.display = 'none';
      singleEditor.style.display = 'none';
      stdinBox.style.display = 'none';
      tabs.forEach(t => { if (t.classList.contains('active')) document.getElementById(t.dataset.target).style.display = 'block'; });
      previewLabel.textContent = 'Live Preview';
      previewFrame.style.display = 'block';
      outputConsole.style.display = 'none';
      runPreview();
    } else {
      webTabs.style.display = 'none';
      Object.values(editors).forEach(ed => ed.style.display = 'none');
      singleTabLabel.style.display = 'block';
      singleTabLabel.textContent = LANGS[lang].label + ' — main.' + LANGS[lang].ext;
      singleEditor.style.display = 'block';
      singleEditor.value = singleCode[lang] ?? LANGS[lang].starter;
      stdinBox.style.display = 'block';
      previewLabel.textContent = 'Output';
      previewFrame.style.display = 'none';
      outputConsole.style.display = 'block';
      outputConsole.textContent = 'Hit ▶ Run to execute this in ' + LANGS[lang].label + '.';
      outputConsole.style.color = '#8b949e';
    }
    if (!skipSave) saveDraft();
  }

  langBtns.forEach(b => b.addEventListener('click', () => switchLang(b.dataset.lang)));

  tabs.forEach(tab => {
    tab.addEventListener('click', () => {
      tabs.forEach(t => {
        t.classList.remove('active');
        t.style.background = 'none'; t.style.color = 'var(--muted)';
      });
      tab.classList.add('active');
      tab.style.background = 'var(--accent-dim)'; tab.style.color = 'var(--accent)';
      Object.values(editors).forEach(ed => ed.style.display = 'none');
      document.getElementById(tab.dataset.target).style.display = 'block';
    });
  });

  // Tab-key inserts 2 spaces instead of moving focus, standard editor behavior
  [...Object.values(editors), singleEditor].forEach(ed => {
    ed.addEventListener('keydown', (e) => {
      if (e.key === 'Tab') {
        e.preventDefault();
        const start = ed.selectionStart, end = ed.selectionEnd;
        ed.value = ed.value.slice(0, start) + '  ' + ed.value.slice(end);
        ed.selectionStart = ed.selectionEnd = start + 2;
      }
    });
    ed.addEventListener('input', saveDraft);
  });

  function buildDocument() {
    return `<!DOCTYPE html>
<html>
<head>
<meta charset="UTF-8">
<style>${editors.css.value}</style>
</head>
<body>
${editors.html.value}
<script>${editors.js.value}<\/script>
</body>
</html>`;
  }

  function runPreview() {
    previewFrame.srcdoc = buildDocument();
    saveDraft();
  }

  async function runSingleLang() {
    const lang = currentLang;
    const code = singleEditor.value;
    singleCode[lang] = code;
    saveDraft();

    outputConsole.style.color = '#c9d1d9';
    outputConsole.textContent = '⏳ Running ' + LANGS[lang].label + '…';
    runBtn.disabled = true;
    const origLabel = runBtn.textContent;
    runBtn.textContent = '⏳ Running…';

    try {
      const res = await fetch('/api/code/run', {
        method: 'POST',
        headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ language: lang, code, stdin: stdinBox.value }),
      });
      const data = await res.json();
      if (!res.ok) {
        outputConsole.style.color = '#f85149';
        outputConsole.textContent = '✕ ' + (data.error || 'Something went wrong running your code.');
        return;
      }
      let out = '';
      if (data.compile_output) out += data.compile_output.trimEnd() + '\n';
      if (data.stdout) out += data.stdout;
      if (data.stderr) out += (out && !out.endsWith('\n') ? '\n' : '') + data.stderr;
      if (!out.trim()) out = '(no output)';
      const meta = `\n\n— ${data.status}` + (data.time ? `, ${data.time}s` : '') + (data.memory ? `, ${Math.round(data.memory/1024)}MB` : '') + ' —';
      const ok = /accepted/i.test(data.status || '');
      outputConsole.style.color = ok ? '#c9d1d9' : '#f0b429';
      outputConsole.textContent = out + meta;
    } catch (e) {
      outputConsole.style.color = '#f85149';
      outputConsole.textContent = '✕ Network error: ' + e.message;
    } finally {
      runBtn.disabled = false;
      runBtn.textContent = origLabel;
    }
  }

  function runCurrent() {
    if (currentLang === 'web') runPreview();
    else runSingleLang();
  }

  codeBtn.addEventListener('click', () => {
    codeModal.style.display = 'flex';
    loadDraft();
    runCurrent();
  });
  closeBtn.addEventListener('click', () => { codeModal.style.display = 'none'; });
  codeModal.addEventListener('click', (e) => { if (e.target === codeModal) codeModal.style.display = 'none'; });
  runBtn.addEventListener('click', runCurrent);

  downloadBtn.addEventListener('click', () => {
    const name = (projectNameInput.value || 'my-project').replace(/[^a-z0-9_-]/gi, '-');
    let blob, filename;
    if (currentLang === 'web') {
      blob = new Blob([buildDocument()], { type: 'text/html;charset=utf-8' });
      filename = name + '.html';
    } else {
      const meta = LANGS[currentLang];
      blob = new Blob([singleEditor.value], { type: 'text/plain;charset=utf-8' });
      filename = (currentLang === 'java' ? 'Main' : name) + '.' + meta.ext;
    }
    const url = URL.createObjectURL(blob);
    const a = document.createElement('a');
    a.href = url; a.download = filename;
    document.body.appendChild(a); a.click(); a.remove();
    URL.revokeObjectURL(url);
  });

  fullscreenBtn.addEventListener('click', () => {
    if (currentLang === 'web') {
      const win = window.open('', '_blank');
      if (win) { win.document.write(buildDocument()); win.document.close(); }
    } else {
      const win = window.open('', '_blank');
      if (win) {
        win.document.write('<pre style="background:#0d1117;color:#c9d1d9;padding:20px;font-family:monospace;white-space:pre-wrap;word-break:break-word;">' +
          outputConsole.textContent.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;') + '</pre>');
        win.document.close();
      }
    }
  });

  // Ctrl+Enter inside the modal re-runs, like most code sandboxes
  codeModal.addEventListener('keydown', (e) => {
    if ((e.ctrlKey || e.metaKey) && e.key === 'Enter') { e.preventDefault(); runCurrent(); }
  });

  loadDraft();
})();
messagesEl.addEventListener('click', async (e) => {
  const btn = e.target.closest('.code-copy-btn');
  if (!btn) return;
  const codeEl = document.getElementById(btn.dataset.target);
  if (!codeEl) return;
  try {
    await navigator.clipboard.writeText(codeEl.textContent);
    const orig = btn.textContent;
    btn.textContent = '✓ Copied';
    setTimeout(() => { btn.textContent = orig; }, 1200);
  } catch {}
});

// ─── Shared password gate (reuses the VIP password) for protected sidebar
// features — Bookmarks, Stats, Archived. Opens the VIP unlock modal if the
// person hasn't already entered the password this session, then runs the
// requested action once they do.
function requirePassword(action) {
  if (vipUnlocked) { action(); return; }
  showVipModal();
  const check = setInterval(() => {
    if (vipUnlocked) { clearInterval(check); action(); }
    const overlay = document.getElementById('vip-modal-overlay');
    if (overlay && overlay.style.display === 'none') clearInterval(check);
  }, 350);
}

// ─── Chat / Cowork / Code mode tabs (VIP-model-gated) ───────────────────────
// Cowork, Code, and Artifacts are only usable while the Mythic VIP model is
// selected. If VIP isn't unlocked yet, switching to it goes through the
// existing VIP password modal exactly once (via showVipModal / vipUnlocked) —
// after that, no repeated password prompts, just the model requirement.
// modeTabs now declared at top-level scope
function setActiveModeTab(mode) {
  currentMode = mode;
  modeTabs.forEach(t => t.classList.toggle('active', t.dataset.mode === mode));
  const placeholders = { chat: 'Message Mythic AI...', cowork: 'Describe the task to hand off...', code: 'Describe what you want to build or fix...' };
  if (input) input.placeholder = placeholders[mode] || placeholders.chat;
}
function requireVipModel(action) {
  if (selectedModel === 'mythic-vip') { action(); return; }
  if (vipUnlocked) {
    selectedModel = 'mythic-vip';
    updateVipBtn();
    syncModeTabLocks();
    action();
    return;
  }
  showVipModal();
  const check = setInterval(() => {
    if (vipUnlocked && selectedModel === 'mythic-vip') { clearInterval(check); syncModeTabLocks(); action(); }
    const overlay = document.getElementById('vip-modal-overlay');
    if (overlay && overlay.style.display === 'none') clearInterval(check);
  }, 350);
}
// Lock icons reflect REAL auth state (vipUnlocked + selectedModel), not a
// manually-toggled class — this avoids the lock silently disappearing/
// staying stale after a page reload or model switch.
function syncModeTabLocks() {
  const isVip = vipUnlocked && selectedModel === 'mythic-vip';
  modeTabs.forEach(t => {
    if (t.dataset.mode === 'chat') return;
    t.classList.toggle('unlocked', isVip);
  });
  if (artifactsTabBtn) artifactsTabBtn.classList.toggle('unlocked', isVip);
}
modeTabs.forEach(tab => {
  tab.addEventListener('click', () => {
    const mode = tab.dataset.mode;
    if (mode === 'chat') { setActiveModeTab('chat'); return; }
    requireVipModel(() => {
      setActiveModeTab(mode);
    });
  });
});
// If the person switches back to a non-VIP model while on Cowork/Code, drop
// back to Chat mode so requests don't silently keep using the VIP-only prompt.
const _origVipBtnHandlerCheck = setInterval(() => {
  if (selectedModel !== 'mythic-vip' && currentMode !== 'chat') setActiveModeTab('chat');
}, 1000);

// ─── Artifacts panel — collects code blocks pulled out of AI replies ───────
function _extractCodeBlocks(text) {
  const blocks = [];
  const re = /```(\w*)\n?([\s\S]*?)```/g;
  let m;
  while ((m = re.exec(text)) !== null) {
    if (m[2].trim()) blocks.push({ lang: m[1] || 'text', code: m[2].trim() });
  }
  return blocks;
}
function _addArtifactsFromReply(fullText) {
  const blocks = _extractCodeBlocks(fullText || '');
  const groupId = 'grp_' + Math.random().toString(36).slice(2, 9);
  blocks.forEach(b => {
    _artifacts.push({
      id: 'art_' + Math.random().toString(36).slice(2, 9),
      groupId, lang: b.lang, code: b.code, ts: Date.now(),
      preview: fullText.replace(/[#*`_~>]/g, '').trim().slice(0, 60),
    });
  });
}
// Combines all html/css/js artifacts from the SAME reply into one runnable
// document and opens it in a sandboxed iframe — this is the real, honest
// equivalent of "run this for me": actual in-browser execution, not a
// claim of touching your filesystem or OS.
function _buildPreviewDoc(groupId) {
  const group = _artifacts.filter(a => a.groupId === groupId);
  const htmlArt = group.find(a => a.lang === 'html');
  const cssArts = group.filter(a => a.lang === 'css');
  const jsArts = group.filter(a => ['javascript', 'js'].includes(a.lang));
  if (!htmlArt) return null;
  let doc = htmlArt.code;
  const cssTag = cssArts.map(a => `<style>${a.code}</style>`).join('\n');
  const jsTag = jsArts.map(a => `<script>${a.code}<\/script>`).join('\n');
  if (doc.includes('</head>')) doc = doc.replace('</head>', cssTag + '\n</head>');
  else doc = cssTag + doc;
  if (doc.includes('</body>')) doc = doc.replace('</body>', jsTag + '\n</body>');
  else doc = doc + jsTag;
  return doc;
}
function showPreviewModal(groupId) {
  const doc = _buildPreviewDoc(groupId);
  if (!doc) return;
  const overlay = document.createElement('div');
  overlay.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.85);z-index:600;display:flex;flex-direction:column;';
  overlay.innerHTML = `
    <div style="display:flex;justify-content:space-between;align-items:center;padding:10px 16px;background:#1a1a1a;">
      <span style="color:#fff;font-size:13px;">▶ Live preview (sandboxed — runs in your browser only)</span>
      <button id="preview-close" style="background:none;border:1px solid #444;color:#ccc;border-radius:6px;padding:5px 12px;cursor:pointer;font-family:inherit;">✕ Close</button>
    </div>
    <iframe id="preview-frame" style="flex:1;border:none;background:#fff;" sandbox="allow-scripts allow-forms allow-modals"></iframe>`;
  document.body.appendChild(overlay);
  const frame = overlay.querySelector('#preview-frame');
  frame.srcdoc = doc;
  overlay.querySelector('#preview-close').addEventListener('click', () => overlay.remove());
}
function showArtifactsModal() {
  const overlay = document.createElement('div');
  overlay.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.7);z-index:500;display:flex;align-items:center;justify-content:center;padding:20px;';
  const rows = _artifacts.length ? _artifacts.slice().reverse().map(a => {
    const canPreview = a.lang === 'html' && _artifacts.some(x => x.groupId === a.groupId && x.lang === 'html');
    const canRunCode = !canPreview && RUNNABLE_LANGS.includes((a.lang || '').toLowerCase());
    return `
    <div class="art-row" data-id="${a.id}" style="border:1px solid var(--border);border-radius:10px;margin-bottom:8px;overflow:hidden;">
      <div style="display:flex;justify-content:space-between;align-items:center;background:var(--panel);padding:8px 12px;font-size:11.5px;color:var(--muted);">
        <span>📦 ${a.lang} &middot; ${a.preview || 'snippet'}</span>
        <div style="display:flex;gap:6px;">
          ${canPreview ? `<button class="art-run" data-group="${a.groupId}" style="background:var(--accent);border:none;color:#fff;border-radius:6px;cursor:pointer;font-size:11px;padding:3px 8px;">▶ Run</button>` : ''}
          ${canRunCode ? `<button class="art-run-code" data-id="${a.id}" style="background:var(--accent);border:none;color:#fff;border-radius:6px;cursor:pointer;font-size:11px;padding:3px 8px;">▶ Run</button>` : ''}
          <button class="art-copy" data-id="${a.id}" style="background:none;border:none;color:var(--muted);cursor:pointer;font-size:12px;">📋</button>
          <button class="art-download" data-id="${a.id}" style="background:none;border:none;color:var(--muted);cursor:pointer;font-size:12px;">⬇</button>
        </div>
      </div>
      <pre style="margin:0;padding:10px 12px;overflow-x:auto;background:var(--bg);font-size:12px;max-height:160px;"><code>${a.code.replace(/&/g,'&amp;').replace(/</g,'&lt;').slice(0,4000)}</code></pre>
    </div>`;
  }).join('') : '<div style="padding:24px;text-align:center;color:var(--muted);font-size:13px;">No artifacts yet — code blocks from AI replies show up here automatically.</div>';
  overlay.innerHTML = `<div style="background:var(--panel);border:1px solid var(--border);border-radius:14px;padding:20px;width:92%;max-width:520px;max-height:78vh;overflow-y:auto;">
    <h3 style="margin:0 0 12px;font-size:16px;">📦 Artifacts</h3>
    <div>${rows}</div>
    <button id="art-close" style="margin-top:10px;width:100%;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:9px;cursor:pointer;font-family:inherit;">Close</button>
  </div>`;
  document.body.appendChild(overlay);
  overlay.querySelectorAll('.art-run').forEach(btn => btn.addEventListener('click', () => showPreviewModal(btn.dataset.group)));
  overlay.querySelectorAll('.art-copy').forEach(btn => btn.addEventListener('click', async () => {
    const a = _artifacts.find(x => x.id === btn.dataset.id);
    if (a) { try { await navigator.clipboard.writeText(a.code); btn.textContent = '✓'; setTimeout(() => btn.textContent = '📋', 1000); } catch {} }
  }));
  overlay.querySelectorAll('.art-download').forEach(btn => btn.addEventListener('click', () => {
    const a = _artifacts.find(x => x.id === btn.dataset.id);
    if (!a) return;
    const ext = { python: 'py', javascript: 'js', js: 'js', html: 'html', css: 'css', json: 'json', bash: 'sh', text: 'txt' }[a.lang] || 'txt';
    const blob = new Blob([a.code], { type: 'text/plain' });
    const url = URL.createObjectURL(blob);
    const link = document.createElement('a');
    link.href = url; link.download = `artifact.${ext}`;
    document.body.appendChild(link); link.click(); link.remove();
    URL.revokeObjectURL(url);
  }));
  overlay.querySelector('#art-close').addEventListener('click', () => overlay.remove());
  overlay.addEventListener('click', e => { if (e.target === overlay) overlay.remove(); });
}
const artifactsTabBtn = document.getElementById('artifacts-tab-btn');
if (artifactsTabBtn) artifactsTabBtn.addEventListener('click', () => requireVipModel(() => {
  showArtifactsModal();
}));

// ─── Real code execution ("Run") for non-HTML artifacts, e.g. Python ───────
// HTML/CSS/JS artifacts already run for real, in-browser, via the sandboxed
// iframe above. Other languages (Python, etc.) can't run in a browser at
// all, so this calls the server's /api/execute-code, which itself runs the
// code on Piston (a third-party sandbox) rather than on this app's own
// server — see the long comment above api_execute_code() in the Python for
// why that separation matters. Gated behind VIP the same way as the rest of
// Cowork/Code/Artifacts.
const RUNNABLE_LANGS = ['python', 'py', 'javascript', 'js', 'bash', 'sh', 'c', 'cpp', 'c++', 'java', 'go', 'rust', 'ruby', 'php'];
function showCodeRunResultModal(title, body) {
  const overlay = document.createElement('div');
  overlay.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.75);z-index:650;display:flex;align-items:center;justify-content:center;padding:20px;';
  overlay.innerHTML = `<div style="background:#111;border:1px solid #333;border-radius:12px;padding:16px;width:92%;max-width:640px;max-height:78vh;overflow-y:auto;color:#eee;font-family:ui-monospace,monospace;">
    <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:10px;">
      <strong style="font-size:13px;">▶ ${title}</strong>
      <button id="run-result-close" style="background:none;border:1px solid #444;color:#ccc;border-radius:6px;padding:4px 10px;cursor:pointer;font-family:inherit;">✕</button>
    </div>
    <pre style="white-space:pre-wrap;font-size:12.5px;margin:0;">${body}</pre>
  </div>`;
  document.body.appendChild(overlay);
  overlay.querySelector('#run-result-close').addEventListener('click', () => overlay.remove());
  overlay.addEventListener('click', e => { if (e.target === overlay) overlay.remove(); });
}
async function runArtifactCode(artifactId) {
  const a = _artifacts.find(x => x.id === artifactId);
  if (!a) return;
  const lang = (a.lang || '').toLowerCase();
  if (!RUNNABLE_LANGS.includes(lang)) {
    showCodeRunResultModal('Can\'t run this', `"${lang}" isn't a runnable language here.`);
    return;
  }
  showCodeRunResultModal('Running…', 'Executing on a sandboxed runner, one moment…');
  try {
    const r = await fetch('/api/execute-code', {
      method: 'POST', headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ code: a.code, language: lang }),
    });
    const d = await r.json();
    if (!r.ok || d.error) {
      showCodeRunResultModal('Run failed', (d.error || 'Unknown error') +
        (r.status === 403 ? '\n\nUnlock VIP mode first (🔒 button in the header).' : ''));
      return;
    }
    const out = [
      d.stdout ? 'stdout:\n' + d.stdout : '',
      d.stderr ? 'stderr:\n' + d.stderr : '',
      (!d.stdout && !d.stderr) ? '(no output)' : '',
      `\n[exit code ${d.exit_code}]`,
    ].filter(Boolean).join('\n\n');
    showCodeRunResultModal(`Result (${d.language} ${d.version || ''})`, out);
  } catch (err) {
    showCodeRunResultModal('Network error', String(err.message || err));
  }
}
// Hook into the Artifacts modal's row buttons for non-HTML runnable code —
// showArtifactsModal() only wires up '.art-run' for HTML groups today, so
// this listens at the document level and covers any '.art-run-code' button
// we add per-row (see the small patch to the row template just below).
document.addEventListener('click', e => {
  const btn = e.target.closest('.art-run-code');
  if (btn) runArtifactCode(btn.dataset.id);
});

// ─── Message search across every chat (not just the open one) ─────────────
const searchChatsBtn = document.getElementById('search-chats-btn');
function showSearchModal() {
  const overlay = document.createElement('div');
  overlay.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.7);z-index:500;display:flex;align-items:center;justify-content:center;padding:20px;';
  overlay.innerHTML = `<div style="background:var(--panel);border:1px solid var(--border);border-radius:14px;padding:20px;width:92%;max-width:520px;max-height:78vh;overflow-y:auto;">
    <h3 style="margin:0 0 12px;font-size:16px;">🔎 Search your chats</h3>
    <input id="search-chats-input" type="text" placeholder="Search all conversations..." autocomplete="off"
      style="width:100%;box-sizing:border-box;padding:10px 12px;border-radius:8px;border:1px solid var(--border);background:var(--bg);color:var(--text);font-family:inherit;font-size:14px;margin-bottom:12px;">
    <div id="search-chats-results" style="font-size:13px;color:var(--muted);"></div>
    <button id="search-chats-close" style="margin-top:10px;width:100%;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:9px;cursor:pointer;font-family:inherit;">Close</button>
  </div>`;
  document.body.appendChild(overlay);
  const input2 = overlay.querySelector('#search-chats-input');
  const resultsEl = overlay.querySelector('#search-chats-results');
  let debounceTimer = null;
  input2.addEventListener('input', () => {
    clearTimeout(debounceTimer);
    const q = input2.value.trim();
    if (q.length < 2) { resultsEl.innerHTML = ''; return; }
    debounceTimer = setTimeout(async () => {
      resultsEl.textContent = 'Searching…';
      try {
        const r = await fetch('/api/search?q=' + encodeURIComponent(q));
        const d = await r.json();
        const results = d.results || [];
        resultsEl.innerHTML = results.length ? results.map(res => `
          <div class="search-result-row" data-conv="${res.conv_id}" style="padding:10px;border:1px solid var(--border);border-radius:8px;margin-bottom:6px;cursor:pointer;">
            <div style="font-weight:600;font-size:12.5px;margin-bottom:3px;">${(res.title || '').replace(/</g,'&lt;')}</div>
            <div style="opacity:.8;">${(res.role === 'user' ? '🧑' : '🤖')} ${(res.snippet || '').replace(/</g,'&lt;')}</div>
          </div>`).join('') : '<div style="padding:12px;text-align:center;">No matches.</div>';
        resultsEl.querySelectorAll('.search-result-row').forEach(row => {
          row.addEventListener('click', () => { overlay.remove(); openConversation(row.dataset.conv); });
        });
      } catch { resultsEl.textContent = 'Search failed — try again.'; }
    }, 300);
  });
  overlay.querySelector('#search-chats-close').addEventListener('click', () => overlay.remove());
  overlay.addEventListener('click', e => { if (e.target === overlay) overlay.remove(); });
  requestAnimationFrame(() => input2.focus());
}
if (searchChatsBtn) searchChatsBtn.addEventListener('click', showSearchModal);

// ─── Reminders & scheduled tasks (delivered via existing Web Push) ─────────
const remindersBtn = document.getElementById('reminders-btn');
async function loadRemindersList(container) {
  container.textContent = 'Loading…';
  try {
    const r = await fetch('/api/reminders');
    const d = await r.json();
    const list = d.reminders || [];
    container.innerHTML = list.length ? list.map(rem => `
      <div style="display:flex;justify-content:space-between;align-items:center;padding:8px 10px;border:1px solid var(--border);border-radius:8px;margin-bottom:6px;font-size:12.5px;">
        <div>
          <div>${rem.text.replace(/</g,'&lt;')}</div>
          <div style="opacity:.7;font-size:11px;">${new Date(rem.fire_at * 1000).toLocaleString()}</div>
        </div>
        <button class="reminder-del" data-id="${rem.id}" style="background:none;border:none;color:var(--muted);cursor:pointer;font-size:14px;">✕</button>
      </div>`).join('') : '<div style="padding:12px;text-align:center;color:var(--muted);">No reminders set.</div>';
    container.querySelectorAll('.reminder-del').forEach(btn => btn.addEventListener('click', async () => {
      await fetch('/api/reminders/' + btn.dataset.id, { method: 'DELETE' });
      loadRemindersList(container);
    }));
  } catch { container.textContent = 'Could not load reminders.'; }
}
function showRemindersModal() {
  const overlay = document.createElement('div');
  overlay.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.7);z-index:500;display:flex;align-items:center;justify-content:center;padding:20px;';
  const now = new Date(Date.now() + 5 * 60000);
  const defaultLocal = new Date(now.getTime() - now.getTimezoneOffset() * 60000).toISOString().slice(0, 16);
  overlay.innerHTML = `<div style="background:var(--panel);border:1px solid var(--border);border-radius:14px;padding:20px;width:92%;max-width:480px;max-height:80vh;overflow-y:auto;">
    <h3 style="margin:0 0 12px;font-size:16px;">⏰ Reminders</h3>
    <p style="font-size:12px;color:var(--muted);margin:0 0 12px;">Delivered as a push notification (enable notifications in Settings first). Requires the app to stay running on the server (works on Render; not on serverless hosts like Vercel).</p>
    <input id="reminder-text-input" type="text" placeholder="Remind me to..." maxlength="200"
      style="width:100%;box-sizing:border-box;padding:10px 12px;border-radius:8px;border:1px solid var(--border);background:var(--bg);color:var(--text);font-family:inherit;font-size:14px;margin-bottom:8px;">
    <input id="reminder-time-input" type="datetime-local" value="${defaultLocal}"
      style="width:100%;box-sizing:border-box;padding:10px 12px;border-radius:8px;border:1px solid var(--border);background:var(--bg);color:var(--text);font-family:inherit;font-size:14px;margin-bottom:10px;">
    <button id="reminder-add-btn" style="width:100%;background:var(--accent);color:#fff;border:none;border-radius:8px;padding:10px;cursor:pointer;font-family:inherit;margin-bottom:14px;">Set reminder</button>
    <div id="reminders-list"></div>
    <button id="reminders-close" style="margin-top:10px;width:100%;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:9px;cursor:pointer;font-family:inherit;">Close</button>
  </div>`;
  document.body.appendChild(overlay);
  const listEl = overlay.querySelector('#reminders-list');
  loadRemindersList(listEl);
  overlay.querySelector('#reminder-add-btn').addEventListener('click', async () => {
    const text = overlay.querySelector('#reminder-text-input').value.trim();
    const localVal = overlay.querySelector('#reminder-time-input').value;
    if (!text || !localVal) return;
    const fireAt = new Date(localVal).getTime() / 1000;
    try {
      const r = await fetch('/api/reminders', {
        method: 'POST', headers: { 'Content-Type': 'application/json' },
        body: JSON.stringify({ text, fire_at: fireAt }),
      });
      const d = await r.json();
      if (!r.ok) { alert(d.error || 'Could not set reminder.'); return; }
      overlay.querySelector('#reminder-text-input').value = '';
      loadRemindersList(listEl);
    } catch (err) { alert('Network error: ' + err.message); }
  });
  overlay.querySelector('#reminders-close').addEventListener('click', () => overlay.remove());
  overlay.addEventListener('click', e => { if (e.target === overlay) overlay.remove(); });
}
if (remindersBtn) remindersBtn.addEventListener('click', showRemindersModal);

// ─── Full backup export / import (all chats, not just one) ────────────────
// Wired into the Settings modal — see #backup-export-btn / #backup-import-*
// elements added to the settings modal markup.
const backupExportBtn = document.getElementById('backup-export-btn');
const backupImportBtn = document.getElementById('backup-import-btn');
const backupImportFile = document.getElementById('backup-import-file');
const backupStatusEl   = document.getElementById('backup-status');
if (backupExportBtn) backupExportBtn.addEventListener('click', () => {
  window.location.href = '/api/backup/export';
});
if (backupImportBtn && backupImportFile) {
  backupImportBtn.addEventListener('click', () => backupImportFile.click());
  backupImportFile.addEventListener('change', async () => {
    const file = backupImportFile.files[0];
    if (!file) return;
    if (backupStatusEl) backupStatusEl.textContent = 'Importing…';
    const fd = new FormData();
    fd.append('file', file);
    try {
      const r = await fetch('/api/backup/import', { method: 'POST', body: fd });
      const d = await r.json();
      if (backupStatusEl) {
        backupStatusEl.textContent = r.ok
          ? `Imported ${d.imported} conversation(s).`
          : (d.error || 'Import failed.');
      }
      loadConversationList();
    } catch (err) {
      if (backupStatusEl) backupStatusEl.textContent = 'Network error: ' + err.message;
    }
    backupImportFile.value = '';
  });
}

// ─── Cowork mode: real multi-step task runner ──────────────────────────────
// Renders a distinct "steps" card (plan → each step's result → final
// answer) instead of a normal streamed reply, so it's visibly different
// from a single chat completion.
function renderCoworkResult(container, data) {
  const stepsHtml = (data.steps || []).map((s, i) => `
    <div style="margin-bottom:8px;padding:8px 10px;border-left:2px solid var(--accent);">
      <div style="font-size:11.5px;font-weight:600;opacity:.8;">Step ${i + 1}: ${s.step.replace(/</g,'&lt;')}</div>
      <div style="font-size:13px;margin-top:3px;">${s.result.replace(/</g,'&lt;')}</div>
    </div>`).join('');
  container.innerHTML = `
    <div style="font-size:12px;font-weight:700;opacity:.7;margin-bottom:8px;">🗂 COWORK — ${(data.steps||[]).length} step(s)</div>
    ${stepsHtml}
    <div style="margin-top:12px;padding-top:12px;border-top:1px solid var(--border);white-space:pre-wrap;">${(data.final_answer || '').replace(/</g,'&lt;')}</div>`;
}
async function runCoworkTask(task) {
  addMessage('user', task);
  // addMessage() returns the message's text element directly (see its
  // definition above), not the whole row — safe to update in place.
  const textEl = addMessage('ai', '⏳ Planning and working through the task…');
  try {
    const r = await fetch('/api/cowork/run', {
      method: 'POST', headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ task }),
    });
    const d = await r.json();
    if (!r.ok || d.error) {
      if (textEl) textEl.textContent = d.error || 'Cowork task failed.';
      return;
    }
    if (textEl) renderCoworkResult(textEl, d);
  } catch (err) {
    if (textEl) textEl.textContent = 'Network error: ' + err.message;
  }
}

// ─── Starred view toggle (no VIP required) ───────────────────────────────────
const archivedToggleBtn = document.getElementById('archived-toggle-btn');
if (archivedToggleBtn) archivedToggleBtn.addEventListener('click', () => {
  showingStarredOnly = !showingStarredOnly;
  archivedToggleBtn.textContent = showingStarredOnly ? '💬 All Chats' : '⭐ Starred';
  archivedToggleBtn.style.color = showingStarredOnly ? 'var(--accent)' : '';
  archivedToggleBtn.style.borderColor = showingStarredOnly ? 'var(--accent)' : 'var(--border)';
  loadConversationList();
});

// ─── Message bookmarks (stored per-conversation in localStorage) ──────────────
function getBookmarks() {
  try { return JSON.parse(localStorage.getItem('mythic_bookmarks') || '{}'); } catch { return {}; }
}
function saveBookmarks(b) { localStorage.setItem('mythic_bookmarks', JSON.stringify(b)); }
function toggleBookmark(convId, msgIndex, text) {
  if (!convId) return false;
  const all = getBookmarks();
  const list = all[convId] = all[convId] || [];
  const existingIdx = list.findIndex(b => b.msgIndex === msgIndex);
  let nowBookmarked;
  if (existingIdx >= 0) { list.splice(existingIdx, 1); nowBookmarked = false; }
  else { list.push({ msgIndex, text: (text || '').slice(0, 200), ts: Date.now() }); nowBookmarked = true; }
  if (!list.length) delete all[convId];
  saveBookmarks(all);
  return nowBookmarked;
}
function isBookmarked(convId, msgIndex) {
  const list = getBookmarks()[convId] || [];
  return list.some(b => b.msgIndex === msgIndex);
}

const bookmarksBtn = document.getElementById('bookmarks-btn');
function showBookmarksModal() {
  const all = getBookmarks();
  const overlay = document.createElement('div');
  overlay.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.7);z-index:500;display:flex;align-items:center;justify-content:center;';
  let rows = '';
  let count = 0;
  Object.entries(all).forEach(([convId, list]) => {
    list.forEach(b => {
      count++;
      rows += `<div class="bm-row" data-conv="${convId}" style="padding:10px;border-bottom:1px solid var(--border);cursor:pointer;font-size:12.5px;color:var(--text);">${b.text.replace(/</g,'&lt;')}</div>`;
    });
  });
  if (!count) rows = '<div style="padding:20px;text-align:center;color:var(--muted);font-size:13px;">No bookmarked messages yet. Use the 🔖 icon on any message.</div>';
  overlay.innerHTML = `<div style="background:var(--panel);border:1px solid var(--border);border-radius:14px;padding:20px;width:90%;max-width:420px;max-height:70vh;overflow-y:auto;">
    <h3 style="margin:0 0 12px;font-size:16px;">🔖 Bookmarked Messages</h3>
    <div>${rows}</div>
    <button id="bm-close" style="margin-top:14px;width:100%;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:9px;cursor:pointer;font-family:inherit;">Close</button>
  </div>`;
  document.body.appendChild(overlay);
  overlay.querySelectorAll('.bm-row').forEach(row => {
    row.addEventListener('click', () => { openConversation(row.dataset.conv); overlay.remove(); });
  });
  overlay.querySelector('#bm-close').addEventListener('click', () => overlay.remove());
  overlay.addEventListener('click', e => { if (e.target === overlay) overlay.remove(); });
}
if (bookmarksBtn) bookmarksBtn.addEventListener('click', showBookmarksModal);

// ─── Chat statistics ────────────────────────────────────────────────────────
const statsBtn = document.getElementById('stats-btn');
async function showStatsModal() {
  const convs = await fetch('/api/conversations').then(r => r.json()).then(d => d.conversations || []).catch(() => []);
  const streak = await fetch('/api/streak').then(r => r.json()).then(d => d.streak || 0).catch(() => 0);
  let totalMsgsGuess = 0;
  if (activeConvId) {
    const d = await fetch('/api/conversations/' + activeConvId).then(r => r.json()).catch(() => null);
    if (d && d.messages) totalMsgsGuess = d.messages.length;
  }
  const overlay = document.createElement('div');
  overlay.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.7);z-index:500;display:flex;align-items:center;justify-content:center;';
  overlay.innerHTML = `<div style="background:var(--panel);border:1px solid var(--border);border-radius:14px;padding:20px;width:90%;max-width:380px;">
    <h3 style="margin:0 0 14px;font-size:16px;">📊 Chat Statistics</h3>
    <div style="display:grid;grid-template-columns:1fr 1fr;gap:10px;">
      <div style="background:var(--bg);border-radius:8px;padding:12px;text-align:center;">
        <div style="font-size:22px;font-weight:700;color:var(--accent);">${convs.length}</div>
        <div style="font-size:11px;color:var(--muted);margin-top:2px;">Total Chats</div>
      </div>
      <div style="background:var(--bg);border-radius:8px;padding:12px;text-align:center;">
        <div style="font-size:22px;font-weight:700;color:var(--accent);">🔥 ${streak}</div>
        <div style="font-size:11px;color:var(--muted);margin-top:2px;">Day Streak</div>
      </div>
      <div style="background:var(--bg);border-radius:8px;padding:12px;text-align:center;">
        <div style="font-size:22px;font-weight:700;color:var(--accent);">${convs.filter(c=>c.pinned).length}</div>
        <div style="font-size:11px;color:var(--muted);margin-top:2px;">Pinned</div>
      </div>
      <div style="background:var(--bg);border-radius:8px;padding:12px;text-align:center;">
        <div style="font-size:22px;font-weight:700;color:var(--accent);">${totalMsgsGuess}</div>
        <div style="font-size:11px;color:var(--muted);margin-top:2px;">Messages (this chat)</div>
      </div>
    </div>
    <button id="stats-close" style="margin-top:16px;width:100%;background:none;border:1px solid var(--border);color:var(--muted);border-radius:8px;padding:9px;cursor:pointer;font-family:inherit;">Close</button>
  </div>`;
  document.body.appendChild(overlay);
  overlay.querySelector('#stats-close').addEventListener('click', () => overlay.remove());
  overlay.addEventListener('click', e => { if (e.target === overlay) overlay.remove(); });
}
if (statsBtn) statsBtn.addEventListener('click', showStatsModal);

// ─── Bookmark button on AI/user messages ───────────────────────────────────
let _msgIndexCounter = 0;
const _origBuildActions2 = buildMsgActions;
buildMsgActions = function(row, textNode, role) {
  const actions = _origBuildActions2(row, textNode, role);
  const myIndex = _msgIndexCounter++;
  row.dataset.msgIndex = myIndex;
  const bmBtn = document.createElement('button');
  bmBtn.type = 'button'; bmBtn.title = 'Bookmark';
  bmBtn.textContent = isBookmarked(activeConvId, myIndex) ? '🔖' : '📑';
  bmBtn.addEventListener('click', () => {
    const on = toggleBookmark(activeConvId, myIndex, textNode.textContent || textNode.innerText || '');
    bmBtn.textContent = on ? '🔖' : '📑';
  });
  actions.appendChild(bmBtn);
  return actions;
};

// ─── Command palette (Ctrl+K) ──────────────────────────────────────────────
function showCommandPalette() {
  const existing = document.getElementById('cmd-palette-overlay');
  if (existing) { existing.remove(); }
  const overlay = document.createElement('div');
  overlay.id = 'cmd-palette-overlay';
  overlay.style.cssText = 'position:fixed;inset:0;background:rgba(0,0,0,.5);z-index:600;display:flex;align-items:flex-start;justify-content:center;padding-top:12vh;';
  overlay.innerHTML = `<div style="background:var(--panel);border:1px solid var(--border);border-radius:12px;width:92%;max-width:480px;box-shadow:0 10px 50px rgba(0,0,0,.4);overflow:hidden;">
    <input id="cmd-input" placeholder="Type a command or search chats..." autocomplete="off"
      style="width:100%;box-sizing:border-box;padding:14px 16px;background:transparent;border:none;border-bottom:1px solid var(--border);color:var(--text);font-size:14px;outline:none;font-family:inherit;">
    <div id="cmd-results" style="max-height:320px;overflow-y:auto;"></div>
  </div>`;
  document.body.appendChild(overlay);
  const input = overlay.querySelector('#cmd-input');
  const results = overlay.querySelector('#cmd-results');
  input.focus();

  const staticCommands = [
    { label: '+ New chat', action: () => { startNewChat(); } },
    { label: '⚙ Open Settings', action: () => { settingsModalOverlay.style.display = 'flex'; } },
    { label: '📊 Chat Statistics', action: showStatsModal },
    { label: '🔖 Bookmarked Messages', action: showBookmarksModal },
    { label: '⭐ Toggle Starred View', action: () => archivedToggleBtn && archivedToggleBtn.click() },
    { label: '⬇ Export current chat', action: () => exportBtn.click() },
    { label: '☰ Toggle sidebar', action: () => sidebarToggle.click() },
  ];

  async function renderResults(query) {
    results.innerHTML = '';
    const q = query.trim().toLowerCase();
    const cmdMatches = staticCommands.filter(c => !q || c.label.toLowerCase().includes(q));
    cmdMatches.forEach(c => {
      const row = document.createElement('div');
      row.textContent = c.label;
      row.style.cssText = 'padding:10px 16px;cursor:pointer;font-size:13.5px;color:var(--text);';
      row.addEventListener('mouseenter', () => row.style.background = 'var(--accent-dim)');
      row.addEventListener('mouseleave', () => row.style.background = '');
      row.addEventListener('click', () => { c.action(); overlay.remove(); });
      results.appendChild(row);
    });
    if (q) {
      const convs = await fetch('/api/conversations').then(r => r.json()).then(d => d.conversations || []).catch(() => []);
      const chatMatches = convs.filter(c => c.title.toLowerCase().includes(q)).slice(0, 8);
      chatMatches.forEach(c => {
        const row = document.createElement('div');
        row.textContent = '💬 ' + c.title;
        row.style.cssText = 'padding:10px 16px;cursor:pointer;font-size:13.5px;color:var(--muted);';
        row.addEventListener('mouseenter', () => row.style.background = 'var(--accent-dim)');
        row.addEventListener('mouseleave', () => row.style.background = '');
        row.addEventListener('click', () => { openConversation(c.id); overlay.remove(); });
        results.appendChild(row);
      });
    }
    if (!results.children.length) {
      results.innerHTML = '<div style="padding:16px;text-align:center;color:var(--muted);font-size:12.5px;">No matches.</div>';
    }
  }
  renderResults('');
  input.addEventListener('input', () => renderResults(input.value));
  overlay.addEventListener('click', e => { if (e.target === overlay) overlay.remove(); });
  input.addEventListener('keydown', e => { if (e.key === 'Escape') overlay.remove(); });
}

// ─── Extra keyboard shortcuts ───────────────────────────────────────────────
document.addEventListener('keydown', e => {
  const mod = e.ctrlKey || e.metaKey;
  if (mod && e.key.toLowerCase() === 'k') { e.preventDefault(); showCommandPalette(); return; }
  if (mod && e.shiftKey && e.key.toLowerCase() === 'o') { e.preventDefault(); startNewChat(); return; }
  if (mod && e.key.toLowerCase() === 'b') { e.preventDefault(); sidebarToggle.click(); return; }
  // Alt+M: jump straight to the message box from anywhere on the page
  if (e.altKey && !mod && e.key.toLowerCase() === 'm') { e.preventDefault(); input.focus(); return; }
  // "/" focuses the message input, unless already typing somewhere
  if (e.key === '/' && document.activeElement.tagName !== 'INPUT' && document.activeElement.tagName !== 'TEXTAREA') {
    e.preventDefault(); input.focus();
  }
});

// ─── Auto-generate a smart AI title after the first exchange in a new chat ──
const _origStreamReply = streamReply;
streamReply = async function(opts) {
  const wasNewChat = !activeConvId;
  await _origStreamReply(opts);
  if (wasNewChat && activeConvId && !(opts && opts.regenerate)) {
    fetch('/api/conversations/' + activeConvId + '/generate-title', {
      method: 'POST', headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify(getUserApiKeys())
    }).then(() => loadConversationList()).catch(() => {});
  }
};
</script>
</body>
</html>
"""

# ── Public, read-only "shared chat" page ──────────────────────────────────
# Deliberately a separate, minimal template: no login, no sidebar, no
# composer, no access to the viewer's own conversations — it only ever
# fetches /api/share/<id>, which itself only ever returns the one shared
# conversation's messages.
SHARE_PAGE = r"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<meta name="robots" content="noindex">
<link rel="icon" type="image/png" href="/icon.png">
<title>Shared chat · Mythic AI</title>
<style>
  :root { --bg:#1a1a1a; --panel:#2a2a2a; --border:#3a3a3a; --text:#ececec;
    --muted:#8e8ea0; --accent:#10a37f; --user-bubble:#2a2a2a; --ai-bubble:#1a1a1a; }
  * { box-sizing:border-box; margin:0; padding:0; }
  body { background:var(--bg); color:var(--text); font-family:-apple-system,BlinkMacSystemFont,
    "Segoe UI",Inter,sans-serif; min-height:100vh; }
  header { position:sticky; top:0; background:var(--bg); border-bottom:1px solid var(--border);
    padding:14px 20px; display:flex; align-items:center; justify-content:space-between; gap:10px;
    z-index:10; }
  header .brand { display:flex; align-items:center; gap:8px; font-weight:700; color:var(--accent);
    font-variant:small-caps; letter-spacing:.5px; }
  header .brand img { width:26px; height:26px; border-radius:7px; display:block; }
  header .badge { display:inline-flex; align-items:center; gap:6px; font-size:11px; color:var(--muted);
    border:1px solid var(--border); border-radius:20px; padding:5px 12px; }
  header .badge svg { width:13px; height:13px; flex-shrink:0; }
  header a.cta { background:var(--accent); color:#fff; text-decoration:none; font-size:12.5px;
    font-weight:700; padding:8px 14px; border-radius:8px; }
  #wrap { max-width:760px; margin:0 auto; padding:20px; }
  #title { font-size:19px; font-weight:700; margin-bottom:18px; }
  .msg-row { display:flex; flex-direction:column; margin-bottom:14px; max-width:82%; }
  .msg-row.user { margin-left:auto; align-items:flex-end; }
  .msg-row.ai { margin-right:auto; align-items:flex-start; }
  .msg { padding:11px 15px; border-radius:18px; line-height:1.6; font-size:14.5px;
    white-space:pre-wrap; word-wrap:break-word; }
  .msg-row.user .msg { background:var(--user-bubble); border-bottom-right-radius:4px; }
  .msg-row.ai .msg { background:var(--ai-bubble); border-bottom-left-radius:4px; }
  #state { text-align:center; color:var(--muted); padding:60px 20px; font-size:14px; }
  #footer-cta { text-align:center; padding:14px 20px 40px; color:var(--muted); font-size:13px; }
  #footer-cta a { color:var(--accent); text-decoration:none; font-weight:600; }
  code { background:var(--panel); padding:1px 5px; border-radius:4px; font-size:.92em; }
  pre { background:var(--panel); border:1px solid var(--border); border-radius:10px; padding:10px 12px;
    overflow-x:auto; margin:8px 0; }
  #continue-bar { max-width:760px; margin:0 auto; padding:0 20px 20px; }
  #continue-btn { display:flex; align-items:center; justify-content:center; gap:8px; width:100%;
    background:var(--accent); color:#fff; border:none; border-radius:12px; padding:13px;
    font-size:14px; font-weight:700; cursor:pointer; font-family:inherit; }
  #continue-btn:hover { opacity:.92; }
  #continue-btn:disabled { opacity:.6; cursor:default; }
  #continue-error { max-width:760px; margin:0 auto; padding:0 20px 12px; color:#ef4444; font-size:12.5px; display:none; }
</style>
</head>
<body>
<header>
  <span class="brand"><img src="/icon.png" alt="Mythic AI"> Mythic AI</span>
  <span class="badge">
    <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" stroke-linecap="round" stroke-linejoin="round">
      <path d="M10 13a5 5 0 0 0 7.54.54l3-3a5 5 0 0 0-7.07-7.07l-1.72 1.71"/>
      <path d="M14 11a5 5 0 0 0-7.54-.54l-3 3a5 5 0 0 0 7.07 7.07l1.71-1.71"/>
    </svg>
    Shared chat · read-only
  </span>
  <a class="cta" href="/">Open Mythic AI</a>
</header>
<div id="wrap">
  <div id="title"></div>
  <div id="messages"></div>
  <div id="state" style="display:none;"></div>
</div>
<div id="continue-error"></div>
<div id="continue-bar" style="display:none;">
  <button id="continue-btn" type="button">
    <svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
      <path d="M21 15a2 2 0 0 1-2 2H7l-4 4V5a2 2 0 0 1 2-2h14a2 2 0 0 1 2 2z"/>
    </svg>
    Continue this conversation
  </button>
</div>
<div id="footer-cta">Want a conversation like this? <a href="/">Try Mythic AI</a></div>
<script>
function esc(s) { return s.replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;'); }
function renderInline(text) {
  let html = esc(text);
  html = html.replace(/```(\w*)\n?([\s\S]*?)```/g, (_, lang, code) =>
    `<pre><code>${code.trim()}</code></pre>`);
  html = html.replace(/`([^`]+)`/g, '<code>$1</code>');
  html = html.replace(/\*\*(.+?)\*\*/g, '<strong>$1</strong>');
  html = html.replace(/\n/g, '<br>');
  return html;
}
const shareId = location.pathname.split('/').filter(Boolean).pop();
(async () => {
  const stateEl = document.getElementById('state');
  try {
    const r = await fetch('/api/share/' + encodeURIComponent(shareId));
    const d = await r.json();
    if (!r.ok || d.error) {
      stateEl.textContent = d.error || 'This shared chat could not be found.';
      stateEl.style.display = 'block';
      return;
    }
    document.title = (d.title || 'Shared chat') + ' · Mythic AI';
    document.getElementById('title').textContent = d.title || 'Shared chat';
    const wrap = document.getElementById('messages');
    (d.messages || []).forEach(m => {
      const row = document.createElement('div');
      row.className = 'msg-row ' + (m.role === 'user' ? 'user' : 'ai');
      const bubble = document.createElement('div');
      bubble.className = 'msg';
      bubble.innerHTML = renderInline(m.text || '');
      row.appendChild(bubble);
      wrap.appendChild(row);
    });
    if (!(d.messages || []).length) {
      stateEl.textContent = 'This chat has no messages yet.';
      stateEl.style.display = 'block';
    } else {
      document.getElementById('continue-bar').style.display = 'block';
    }
  } catch (e) {
    stateEl.textContent = 'Network error loading this shared chat.';
    stateEl.style.display = 'block';
  }
})();

const continueBtn = document.getElementById('continue-btn');
const continueErr = document.getElementById('continue-error');
continueBtn.addEventListener('click', async () => {
  continueErr.style.display = 'none';
  continueBtn.disabled = true;
  const origHTML = continueBtn.innerHTML;
  continueBtn.textContent = 'Setting up your copy…';
  try {
    const r = await fetch('/api/share/' + encodeURIComponent(shareId) + '/continue', { method: 'POST' });
    const d = await r.json();
    if (!r.ok || d.error || !d.conversation_id) {
      continueErr.textContent = d.error || 'Could not continue this chat right now.';
      continueErr.style.display = 'block';
      continueBtn.disabled = false;
      continueBtn.innerHTML = origHTML;
      return;
    }
    // This becomes the visitor's OWN private, editable copy — the
    // original owner's conversation is never touched.
    location.href = '/?open=' + encodeURIComponent(d.conversation_id);
  } catch (e) {
    continueErr.textContent = 'Network error: ' + e.message;
    continueErr.style.display = 'block';
    continueBtn.disabled = false;
    continueBtn.innerHTML = origHTML;
  }
});
</script>
</body>
</html>
"""

@app.route("/sw.js")
def service_worker_js():
    """Real service worker served from a proper URL so push events work.
    Blob-URL service workers cannot receive push events — this route is
    required for Web Push to function."""
    sw = r"""
const CACHE = 'mythic-ai-v3';

self.addEventListener('install', e => {
  e.waitUntil(caches.open(CACHE).then(c => c.addAll(['/'])));
  self.skipWaiting();
});

self.addEventListener('activate', e => {
  e.waitUntil(
    caches.keys().then(keys =>
      Promise.all(keys.filter(k => k !== CACHE).map(k => caches.delete(k)))
    )
  );
  self.clients.claim();
});

self.addEventListener('fetch', e => {
  if (e.request.method !== 'GET') return;
  if (e.request.url.includes('/api/')) return;
  e.respondWith(
    fetch(e.request)
      .then(resp => {
        const clone = resp.clone();
        caches.open(CACHE).then(c => c.put(e.request, clone));
        return resp;
      })
      .catch(() => caches.match(e.request))
  );
});

self.addEventListener('push', e => {
  let data = { title: 'Mythic AI', body: 'You have a new message', icon: '/icon.png', url: '/' };
  try {
    if (e.data) {
      const parsed = e.data.json();
      data = { ...data, ...parsed };
    }
  } catch {}

  e.waitUntil(
    self.registration.showNotification(data.title, {
      body:    data.body,
      icon:    data.icon,
      badge:   '/badge.png',
      tag:     'mythic-ai-reply',
      renotify: true,
      vibrate: [200, 100, 200],
      data:    { url: data.url || '/' },
      actions: [
        { action: 'open',    title: '💬 Open Chat' },
        { action: 'dismiss', title: '✕ Dismiss'   },
      ],
    })
  );
});

self.addEventListener('notificationclick', e => {
  e.notification.close();
  if (e.action === 'dismiss') return;
  const target = (e.notification.data && e.notification.data.url) || '/';
  e.waitUntil(
    clients.matchAll({ type: 'window', includeUncontrolled: true }).then(list => {
      for (const c of list) {
        if (c.url.includes(self.location.origin) && 'focus' in c) {
          c.navigate(target);
          return c.focus();
        }
      }
      if (clients.openWindow) return clients.openWindow(target);
    })
  );
});
"""
    return Response(sw, mimetype="application/javascript",
                    headers={"Service-Worker-Allowed": "/"})


@app.route("/health")
def health_check():
    """Lightweight endpoint for uptime pingers (UptimeRobot, cron-job.org,
    etc.) to hit every 10-14 minutes, keeping a Render free-tier instance
    from spinning down. Does no real work — just confirms the process is alive."""
    return jsonify({"status": "ok", "time": time.time()})


@app.route("/manifest.json")
def pwa_manifest():
    manifest = {
        "id": "/",
        "name": "Mythic AI",
        "short_name": "Mythic AI",
        "description": "Smart AI assistant by Aarav Singh",
        "start_url": "/",
        "display": "standalone",
        "background_color": "#1a1a1a",
        "theme_color": "#10a37f",
        "orientation": "any",
        "scope": "/",
        "lang": "en",
        "categories": ["productivity", "utilities"],
        "icons": [
            # "any" and "maskable" are declared as SEPARATE icon entries
            # (not combined "any maskable" on one entry) — combining them
            # is flagged by Chrome's Installability panel as likely to
            # render with wrong padding on some platforms.
            {"src": "/icon.png",     "sizes": "192x192", "type": "image/png", "purpose": "any"},
            {"src": "/icon-512.png", "sizes": "512x512", "type": "image/png", "purpose": "any"},
            {"src": "/icon.png",     "sizes": "192x192", "type": "image/png", "purpose": "maskable"},
            {"src": "/icon-512.png", "sizes": "512x512", "type": "image/png", "purpose": "maskable"},
        ],
        "screenshots": [
            {"src": "/screenshot-wide.png",   "sizes": "1280x800",  "type": "image/png", "form_factor": "wide"},
            {"src": "/screenshot-narrow.png", "sizes": "750x1334",  "type": "image/png", "form_factor": "narrow"},
        ],
        "shortcuts": [
            {
                "name": "New Chat",
                "url": "/",
                "description": "Start a new chat",
                "icons": [{"src": "/icon-96.png", "sizes": "96x96", "type": "image/png"}],
            },
        ],
        # Lets navigator.getInstalledRelatedApps() (called client-side) detect
        # an existing Mythic AI install even when the current tab is NOT
        # running in standalone mode — e.g. the user opened the plain URL in
        # a normal browser tab after already installing the app. Without
        # this, the page has no way to see that state and the Install button
        # stays visibly (but harmlessly) shown. "prefer_related_applications"
        # is deliberately omitted/false so this does NOT suppress
        # beforeinstallprompt for genuinely new installs.
        "related_applications": [
            {"platform": "webapp", "url": get_public_origin() + "/manifest.json"},
        ],
    }
    return Response(
        json.dumps(manifest),
        mimetype="application/manifest+json",
        headers={"Cache-Control": "public, max-age=86400"},
    )


@app.route("/robots.txt")
def robots_txt():
    """Tells search engine crawlers what to index. Everything under /api/,
    per-account invite links, and internal admin/dashboard pages are
    disallowed since they're either private (unique to one account) or not
    meaningful search results — only the public home page should be indexed.

    On Render specifically: disallow EVERYTHING. Vercel (PREFERRED_PUBLIC_ORIGIN)
    is the only domain meant to show up in Google — Render is treated purely
    as a backend deployment target, not a public search result."""
    if _is_deindexed_host():
        return Response("User-agent: *\nDisallow: /\n", mimetype="text/plain")

    origin = get_public_origin()
    lines = [
        "User-agent: *",
        "Allow: /$",
        "Disallow: /api/",
        "Disallow: /invite/",
        "Disallow: /legacy-invite/",
        "Disallow: /a/",
        "Disallow: /share/",
        "Disallow: /api-usage",
        "Disallow: /analytics",
        "Disallow: /claim-owner/",
        "",
        f"Sitemap: {origin}/sitemap.xml",
    ]
    return Response("\n".join(lines), mimetype="text/plain")


@app.route("/sitemap.xml")
def sitemap_xml():
    """Minimal sitemap — this is a single-page chat app behind an anonymous
    session, so there's really only one meaningful public URL (the home
    page) worth listing for crawlers. lastmod is set to "today" on every
    request since the app's content is dynamic/always current.

    Always lists PREFERRED_PUBLIC_ORIGIN (Vercel), even if fetched from
    Render, so a crawler that finds this file on Render is pointed straight
    back at the Vercel URL rather than sitemap-ing itself."""
    origin = PREFERRED_PUBLIC_ORIGIN
    lastmod = datetime.date.today().isoformat()
    xml = f"""<?xml version="1.0" encoding="UTF-8"?>
<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">
  <url>
    <loc>{origin}/</loc>
    <lastmod>{lastmod}</lastmod>
    <changefreq>weekly</changefreq>
    <priority>1.0</priority>
  </url>
</urlset>"""
    return Response(xml, mimetype="application/xml")


# ── Official Mythic AI logo (embedded, WEBP for size) ────────────────────
# The real 3D "M" mark, stored once as WEBP (much smaller than PNG for a
# gradient/glow image like this) and decoded + resized on demand below.
# This is what every icon/favicon/PWA-icon route now serves, instead of the
# old programmatically-drawn placeholder mark.
_LOGO_SOURCE_WEBP_B64 = """
UklGRqJyAABXRUJQVlA4IJZyAADQmwGdASoAAgACPj0ci0QiIaEmpdWLANAHiU3YS7IlHyNmlEB64KYpCdXKGtFPTP8B/e/8j/lv
73+1Xy6/6P25ev3ZP+m+5f3uPMv3X/bf4v/Qf+3/R//////dX/rf87/Me8D9Vf+P3BP41/Of9N/aP83/8f9N8Xv7gfA39rPzD+Dv
9j/0H/W/zn7//MV/vf+5/vfel/if9/+zHwCf0T/E/+L8//jm9kb0Dv2i/83rlf/b/ef73///Rt/UP9V/9v9J+//0K/0j+2/9P8+f
kA/6///9gD/df//3Nv4B+///s8Az9qfTD8b+4P5T+gP4/9w/rf8R+3P+I/9X+n4Wz/j9Dv5h+W/2X+J/a7/D/uN9If+Hx9/LfEL/
H/53/oPyz/wP7icETQb+qf3X/d/4v/Kf+j/HerT/leqXiB/r7/vPzp/xvtZeLX657AX9S/wv/C/zX5gfTR/Yf/D/Yflz7hvzT/N/
+z/W/AT/N/63/v/75/nf/b/pv/////vS///uR/ef//+6p+0f/9KKiTHmaXSzwQXPR1mI4sKNA7MsulnggtZNQmGMsaL/bfIc+tdK
Wx/7Atoly0uChD3m7qhklA3dsODd2s/bEavRjHeShaMOtQLXU210pBd7eob7lnEk8gYeQjEq85pcA4S7G752nddRx7pI3akdGPJU
7bn0Ntuhw03t+ztinmJAhKGEzbtEqcBuMWtjTyrEIjSqEAE8/6fnL9X/zelk7SZwzGarAV+1EfOHJpjb0SzzUO4Ph69ReRWpR059
ZwgbJ74VN2JgQpeS6Nv4tjUfPERc5/aU9PvWrOs9/qFNafZ85L8gZCz44fKyroDcATSUefJ2IcCh7mE+mPN++yf0x//+7ieSUqhi
zDTeQ1QIFQ9gWKRRRcRsce6M0LJhqYW6q3PVZTY1m/iDewHLGbp/NPoDemHYfRQK95QCToo0CRq2hPC8cvKbiUpA3vhFVTvigEqY
PPtj69l4W2gRnAsF5qgBkHlEvtR3xEZUv2MolQrN9ABkVEBZ4gLG9VRoGfTzgpEkRgW/rWo2GogPxQTVTCHOKpXM9ErD++iqjx6K
vOfP/be8Bxc7wCpJzIbSqTDQ/+EYxwFlP6bMKR2vGO+0PPqmTMwdZzQ7q42tn2jGFDN+9BahcjtlcJBdib/PbmrmJiV9t36+QLal
OUqPbVQe4AAzANrnVAEK0OomqDPcEUAw1f2t6YxfvgKt/ZTauaopqy6sK1Dla8T5w/Uz18HZUCgUsfrpKEUHtzcPqg7lXx3+OBt7
Bu0lFKVsGHUv+4ggMqu3XcsMWbG/V89Faup/Yo5uG3Cv7fC1EtgHZ2gZRdE+zX952arO7/6ev4IVyIM75s+QIRDb3UOqLaltjnTx
+1S1LfgnKo5fYrJ1UMGuigB5GTuBVRaeaVnb47OlAmY/O/3sJ3ue76mb+4jNYQJu/5OFcI1hyv6SblDJ1oQO/oMH7uOq6RBQOVcF
kWxR7X0zd704HqMQOUlPdBNqP38Qu6j4+amhfWarwsw9cNnhhYl8dP0iyuwPS/txtsMA1A3ChLb+KoBisp3vw3dWWxRdzJTmECgJ
xT2t7oOJrFyU4aFcxs21rDBVoWlZkMRXjzkLUPjgoPzXmBYOwdfWMTe4Tliehawphvueuo/tcIYbdqsq1tyKJ7OyUyc6GIFACar9
ZfqNFGMEmpxajFBP3c2llj6ms0mVXaiDImAEyBVRGWlkLIQ0yUEth900+o9RWq4OxC5SR3jSx4kFPkGMpMydcDIzJh8ffNjylitJ
zz5CbZt69UAg0hjDlTDORdcSs/DmlKu1sO1fGOmyv9RQkUOhwSctxk2W6bsgT1W3TRLNaYLTOeL7vrPDrd2axlB1y26uDdQuRcsH
oJcF6CUdYtCY5lp8R6LIN/oZTLptveXVk8Dg20gYue4noNSFjXIT6JBjB+/ELkQX9I/2NhCJW+/EGPzfctUJGTRG57ZIyn2iZyx6
JevTnaXBh1Jlqu94ovZRHBgCNgGJ0XUmRU9kxbpxW5tVQ51mPYkpgIVU3O9FZ093/rPzVXQOejWnbn5Owi1VlkKhRHKNehm7o6mo
UyCei9uHCkmya1uEqEn+n1PJ2Gwtpil93mbg7l7zGw/3BL6nwCg4qY3WP7gWLBZDFWI5UGIQOBnq7oXr0aY+Rv0rIlpZClaTDI8x
YuHJ2r6rn+sFq52F+YQd9owpQbAK9SNQjlbZ1voS5igNdOJq3dkVXlbtaPP2vYXin34G2m5YCXj/v5fPuWQGuSHEn+DH+7Ua7+Jx
1HyOa9nJXnrvPmC3ahyt0X1kWnLdET3DGJopJ6vte7237It3UQUZPXAwQqJStTvXRbC3Iy75ReMRQ3rhJb9ZvaW0uzhdiJtxvWT8
QQuc+Jts/eswWDU0261cO6Xlf5t7VE9VlxRu3YxFTDaEBe7L9tvGDBkqu/17sNbGHKBSI7802CDiQRPcUNst8T2FvTFaIOLOVESp
4dp/T4vzsb3OieIN40qbnIG3KbTbr+XZKXwJxgyiyLoHIisyz/roQQpeRrghA2xRgJzP41i6ewqLv5Sf/a8v+FfG7LM6s4UAHvPs
uDUNJh5SNVwdxfH3Ky/FC5Rs+Pi2q9XdTUj8EaRtDAA6qWRtU4r3IyUJpMPTXcTq9YZqCSg84mlas1s9hmrb7RIfLaD3oEPhykmW
Mm5JmaYVjNDRpXXgS0/AyMotO4Km1tWdrdYw+It8ThkO0S/GuVPzr5v/ms18Z55tOviTNNDg+sIbzr3nY675RFZvlpcKoooedHM2
WH9ObNzL0of4SNcFhHGhjPxIJMPzsaRPRhehV0mK9Zuq/e8Wcj0QQA3WynKmNJ6c6sKPbaoYGbZVSZs3EhOBuZV9Zhjl5X7bJEkv
GqCRTNBXhbqka6d+T7Bvv/CqVkB5/QJLpjBysutUT/9kXhk/F589l+sZjj80cMm6zxN2u8dfouMXIKKQrWOvSIoSg+dJZSWxd0vw
4CrsFANvF0Od9r4tSCLlawymHDYAxBHyZROXf4AJApxiiqFylcrBpTVjAvfT8bsZoumw14BMQa8G4r/JqyIu4Rx8I5b0/aP0nplA
giiV6NybC2GYbDdPb9P9vMGPXuNC8w/p/uDgIjFI+66SaL6r1IRYJob9MljpAG/VgBPzuK69YzzEO25pr1V99BMBric9em2KiAGX
l7MB26GgUCvZiykJxToOWbHcIO08exe3vXQowRDuTZ1qFd5aONwRax/NFUusyUQ1JLaDngrOMsejVvgXrFMeR/gAcmv5KrDlnobe
DWSpAWXBNewN97WEunYm3eaWLhIxOn3hdbb+6KikjIojwOVC7feLs0vQE7EwKwLWwZyYghdkOYd4lsGHSNq6Gr0ju/lDIJCWjcCt
tOUC/GtxejaYOI4By0a9YCQfBmJGy537/XpJWNT7k4MIKfadubzc2KhW2xdrMA06+rEqy8tgMt8WELPaMBChvYIrFKkW11F2krI3
MMhsBpMR75awBenR/wbec1ksnfWzlbJksiipzk+acouAablW+ZqxgK33oaj7RriRbjXDdiakQF5u/SfYjXElhrTefcdR3Ok/Vx3y
Lr8Hj1OJfoAE27tmPSnNhOcY4poWAdd5XCkprsX1CrjuFuxrpgU+gEV5ZwRNPj3xtN0+gNMYSh2dsAoZrpQnk9nVDedA36+g1lUf
OipnsVKj/xGvctDXwuyOHIe7KSQGt205I5UOFtAn5S+b9deEPuLx8TAQVZmRgvnaW4feIZqAtBz6FPk4oqAmNASKdC9tcdYJpEAz
87ftEBLXP4NCk0c/2MxxkXKBrB1pJucnnuDEqkP6hKHCx8hzMwTgjhUYYwav/pl8tzKg/4I+FDaW7XzNIenR0ZzAXDwmX1cncXT5
jWwM9yh/V+8Bnk9qKOGdPz/IfHXMj/pdGNT/Mcwy5aUzMT2VKkeFXxDBQvH5RmLqwfYhOfv7LPD5w0aO/FP3PkQo0ngOjRmK5bUF
kPOUwVgbvYQYtIpj8uwgbdqK5rqDerZj46q5kOjSxBmm1mBDsKOcHeiOoBdgMepN/ayqh4hn2Kf/euK/+Y0vf70LZYMVfLCFHzEr
bUoZ4uqpOBTaz/9MEvUO7FTBNCHCrUP0S5QQhF2eN9gw/cKjfPjQPXhocfanWojdk5UAXWQl00viDULwQF/xH9Cf/5Ld2KG3/99Q
7M1q5QDORgIqt3duQOFe4adh22tLSmBL+7xZpcN4Famcts+Ool3BWlpQ8GRfLkEH2cILzpJ2rpmJBq6a/w/Kt821hY9amWi1oO+d
MU4dZTQtWwbsqzESGsHp0Cme8EvQ8xs0ulnggl9nYskONYlo9mdH2CgzrOcOc9KsoBs81tAfhuJuQYB2ZZc1P7Hofnxs3P29UI3g
guejrMRxUoRR2ZZdLPBBc47hQeOaw0AA/v9J9P6r4ZSa+zakFm5grYAAAAAAKO4HQwghp8DQALHxBRAnT1TRSEZcIa2MkTNRuhwm
JM/tNY1MS4cOpg2klWUFDYACDTuMAEIQpSOuww6EXcdMZRfiDvU9EXsZ3XS2YYr2W6R+XxK+i8/Cx2lE6AjnilDSufutaVtA/zNH
Qh4SaOA+5SOaQWpwiqx76TsNmUiF5kOdAwPFs3cn0SFOQaudI/ZSrwnwFZUKrSNadCBYujX/TE+Jhzg7W+tA20M+u8/60Z/yer6g
d5kZeFrRoQUeeovzJ5R6Gj44xCKziFljtTfQNbhYYsL7cutjeYXwMuT9VZzPXEGLVTx6fAwr8FMztvkkxcalPo+lC5grr3O9r5Zl
UtmKZbQPDPa6SOtvwxhIt5xJ04gX5R9qxUTfuzTA+NuI0tkYvxvDyqD21hGYbO9qrVIGwieWQzwXYnqYIx55NI8a2ZuONUm0vXgk
Cuk79iRl7Oaqt1YnEpIAxWkjT7JE7DHQARFnachBef0ooZe8GxDAdE++Xh1XyTXgrglQKPAYb+E46mKs48XlkyGbL7Z9hk3NO/QH
Q84nIuqcHVL+wLKmJxGEeIx6/vdHakJ/HUPtUqaKSo885mzGDpMoDSRJmoa461LZ415ywjeUWby8OivzpuJGm/sjrCckfJTFm2L6
mvi1sOt+eAwEQnz48gAVTpc3ngSw1vNz7A9eXZGm7qKhc4vh+aR4VP/DDRG8sTH6LAW4Zxe2Lp5Dkl0TDKqpeNhAln3ZYPgunbvc
U1DW1JrNZmk+fGpl+YrcnFnApWNCpw8A2jyz96tKdvMzZ+hmujETZ31hclxU0QTaQvS3nPnDYlHkv44gpgeNc5QWKcVJNNBZZaB0
zGOHPGQsBXBMk/rJUMp9jqTKQBkwrH6lrzs9Zi+JSds6TkiahLe2Sk6bQXPeOQmkro+6UxXKVJRiAnoXWdlwFOX8+5zI2lLnZs60
L3oLQzl0T36m8S29/nXrGHbvPtDjBOrmDGySKkTSfihdVOlyk8nXZE2rvD35dF5zyj8a9FlLNGycpMokiNeS3LHNXYE+Htk47IfS
59AneSqDISSBTGlIDifrplRJ0AeISZ174ZTd60U1jL9etnkXGvOg13KO8ycOe/Ml5Ps5Jjfqk7QQw/Ob5cbRnKz/wE+bt9bcRSW3
SgCinfggOVSLLvt1M8au0itq/hRM3rty4wcKddhIXW1VxLV/XEzR0D0ZVnaX9A/fNldXeSmLFttCwcgkYgc39doUYAWGYJHKSueV
icGYaB7jlaCs4tu9JytohDOPcgprA3Qp4MXvhM4A7IWefY5I7Vli+D/LwX2eQ/tttffIzvagpHuB2/VICHadANKg/UvoQ0qHyMKW
Wvcs/68q1F0AHl6MtQnRQR/eoLe+/GOnhADlpUPicEahlLRceVZOnf6LEGI2FaXwF8NXpyLSYR6hQFqoQtWN3q6GPliqD61An1Ov
4FYsb94/ygJ8TeXg71JAyxgehdWbAsgAaec+e6G1Ro1j5fLUcV61Kgbp2zxOHnafb4/uxCnp7DLJTFbMvx3S26HXWcVcq3cw/IIe
vl28CESJ1eDpTGnyDlnfes+96qWauksqzIMAwqkBC9FyBzYIoV/Be4juwe8lIrJ6cZC0gUPApGnLvvdT81jC9DUIq/huXfcKZo3D
P7BlrxryHZ/9dber4Yxb7E27BtVYSTpr5ZP9sGkToPA5fpaijZCSH2M7Yx1nyycT/CbeoavUwF5SMn/WBtq2xbDjkWm9AKORUV9w
nTbbONFOthXfc06prP3Xcr3JXjvuBpzy5ftqfo3H1bAP8/OOGkZ4Bk+s7EArDuvBapvF2TCXcllQsmBi36Qolb0bxi6o2/+0Vt2l
bsDeYLsT9UHxlyVUFBEdcL7X7tIm0HqMYD3czkJbNzRbtM34FR6uorsRXXQ0CXhcaEswmoNaz1+F0djFz/z9oE+rw37WG+UtM3R7
2v+gUWsFBBLHdcWrggq8vpO/dAW4yh5f4SSwACLbQUbWZObDJZxaCNLeRewDJSHboNGkdMsTcmzrmV14JaH+qNZJstmFrcuh26Ne
gOueowbh+lIV6Avq+oPBzh9PmwpGqEo5a5fRi64AAE7NTf3B9SJ3KluZovAogOVfoEl1y7c0K+qVMsKiw7/GY7WhHixB1c2vFBmk
jqZDMDK4O0Ja9y+i87ZVeK8GyvfM7p9+Yl4emEVubLST/gdJjfudD7V1VaaNh3WnEFwrJIulgRTiHB1DMBjBE6Fu4T076fl57IQJ
rllpfFWeqxzagaAhFYcVAsBnwESsRmv7TesXiVRZ/Il7Gs7zov/J3KkGBwoz9OG93w1bfRVOX32P7c7008vH3Kak7TEU6ewv7Ksf
usCI2Vwimp4vXTCzjpMYW9XW4uMyXZtJfpTjCduLzNtOU7DLc+qndUZKfULBxH4gqCKRVAjD02y2yVda0YUj4Zs3WLlyjz2EmtRG
mpTEve7Cmq8tZcVU674scXJYcnBXuEKPBbrd6hJaxVDQZYbdwi4piQ3RxPqhd6c74lbr2nfrOk5Yt+XFsD8M7hQ/O0jEPQHOHxnG
UoOYZ8vg37pHk1gFR+ZqxHwV59XmXxHs6IkTpAWXrdOLpJNXgU0sChewQSNrXM2FCfzp34dmKtOTzTll38BZyCN6QFpg2J6kGTR+
iLiVlHwc/yRG1z42/8xwu4WWLmskvYpxobGM43AdPQrzzaF0MajjJYSbXmMs/5cX3Vvw1pOTQvJTkUVA9LMTl95f5tdx/Qak2OoH
2zuwxeA75qb8HrraP8bdjm4rcjE3EB0gbpZQC7p+sYPA9bISWH+TbzYKwQk/BEkcTjr0LjuC03L+z71IGyscDjx72GIf++LZQ7ac
xj+tNYPlqgCXWe/fdiFV4mRBpiIHsaTnj0N4wSvXEf7bMBsfq6KCiMH0PjODCNdvaYHHWZoO6efh4VqGI/Zs1OI8qwyBtLakHN0V
xED3rOdpimKt4oAbRRnmjVmVuPJZdIpK79DiZ68AlAFPTzFpWh7JyaRLL3CLL+cPzQE9d1BJx99s39dJmJLKKWsboFHN3vZWfefX
hWgPRi8+qF20nKrAADeK65SQlnLE/zlPamr1REUVagWltuFPq+lmwJUpUr3QIOzKBTJ77xAmMqR+SWtqv17aEe/0dNHZqd/D4c7E
1dXWjqAEZ9qw99FXYMGLtlFPmg46T2ez7A28v7Ka37+7ndUNsiHDzhiNt6foiYwYfwc3lYXTsZPyCYTE/kCcYLvDtAquj379+nQo
pGalMjm4ILkSwMoJMB0Pm+lupeY/GjCK7epC4a72oxPkAgerMg3Prnj2XG5PxWtf8SJptE7mHLG1WlWD/1Ua5lvipCTRIwovM8KJ
3t1Mj08uCqMSu8eSBuM3RZbU19gaMfDvHek3GxXXWOe+KvGpIL/DxRDj0xdOSHuFtVcYbWXEVnA0ON4mEOXiO9OcGdtEXLp+j7Th
DmGQFUwSR3WV0Mw1Cz4VveaLeML3Rsp660aslE4ulhrsfl3w/gY5WgpAjTe3qFIJA1NPkNj5iyOYxiZmoL2Me15qUfv/Zhu+EL9u
ATcbwbBX/RttWlnhm9/35Fo1R2Hvk/L23iYmQkw+gD1x+tsS0nfTBF3bD9wYE3XqswMAX9GufNcQ//gS3anmpVbYI6/bVyhzO1og
JCM6YSBFNXMbR+hvvesbtnF808iuk4s5kvOnhcjookoKNhPC7cHN0Fop/usA3agzntXZdKcYsdLwX+jhQ475TYgUL7xzTi2RtzBn
GnyCIhVZyP5x3IscFv9Ob+GClVci1n/s4T8tpJKqmyn0R8/chKYozLuI6VV3bJvONl3SwPb/wyppawq07ZpM3abDKJsp4m0tKpB/
MxaWXodVNdWep6rrCMkWi3zB1HDX7CYzM+bXVqR/Sb72Ke+k3OOgwV8HJQB9PDyJ8TmUFUK14DMUkGx/5SRe/2YyGgHnyHgPdXmb
LibLq/RkMkpNaSYs37Gg4wRWcfts2PbE7dnjoS8aElBEE03hVn/pQ9gDliTCplRFm8LiYja0dulCa5J0sOKJrBce6lV4Kh2iwqJy
5RkCh/JQYRjoJAo965WgQ1rMJph0zEYP0KHSjxoA36LegktT+htK9PNdIEO1tdIAtbUMZ5ftu6jHaTpZWJ+JupzuDBiWrCc9RNkr
5NN+z977xg6FruUGzwLrwq7+FWkDuIpaNIvcpcuwt7Q9wJyq0WvYFR+FG3+8lAnCmlRQ1hqaRgwVol6M6le/eGnmLVneD5YA5CIq
bdPcpXSezmH6fjp0Lz8kMMLW2BGq1B6/A63gFH6qSVFQTjiQqjkasZtH/CNdvxba/ZCO8JboCnzOydfYhvW3NbmXRbTMW1+pGUGZ
w5kQYyxdHiVAT0Y362/39U+42DLuaDVKKzqLhdWEjd9VBsDg/IyYSwsDLzFmPXjWJzyuDm24SWJuKXRrwK9mSAj/7Rgt+CH0jbfA
Bkk9IylJerJ81vMflMpc1nqz5unBWQwH5GibAAm0rpUtyf9npZQlBjMxL5w23aqUah1f9+fL/m1iGk/pIa4KwrySN++UdBIJaZWb
JqTkWSQ3O2n+RZh/kvDm2LsYKqewQC3Kd90M0SSz6/0ZgxOflF1tRDt4SKf1ZnerTs6NiIf4QUhdJHqwdr85pDNeWGXNUtGpyVOh
FgcUa0qreeXIb/LdJ0sBftcN97NWINn3NspDME5pZYUaB7LB6WMYtx/FBazJmK7ms1eNhgz8gYx7TLmkau0nDEBznBfFj8krOL8d
n8NMGGpf+Ov6hUdUwU6NedIb37f8ipYMgAWLd6ayZK78d2gIB6S56Rh1EHdnAhjl4oz1P1LCx+hrmNxm5dbjEXjliTuwt5IhAzVs
vy2yp+YzvExRGBPpMxM9KNAPforVlZteLoz8JASwnskbvGENDYwJqM6hdWWuYmVRoV8fky7wn0bTiI3mIUEmhVYhXuBp3SLOeU+X
nCxyERrz2w193WAmuiHIQEFu43YYpkxiIw6FkzNLlCaSmFyKrYYnMXWJpUdrQY2BaUCZosGgoBgLB35Kw4eQHerENtMinlF80FeO
YPe3mtQVyRPX5l+94Zh//piDt9XssLargZjVHhp2yACmz/gX1YytXgjbDPQkvW6s5lDLLuBagcenIoqhuPGbD0d+QHvUzuzxxZUO
2oUbi+/NisgbXwjApUhat2guzf0yPpgpBwYe4XLLJadr3+XwmFb+wuwIvfLsIpRrq+LAQIciQAE8bbOAWKiA/oi3wQku9hxD3KgB
t2gK7E750KzG0xZYI+E/dykEoVht4SQwLEmlcyk2CgZjKA1lT38nYKgV7PDx3JGpjn5L24AexRNVIA4BKUSF1aOoMhlqoCRwBjPK
V83MRW4qs2sGVq78HEUA28pUc3Zbh8EBCMq4Eb1+j1Nshkio+EDB+6SwUI6UoLkKw5e5+Z89gpmYgGuAs1xgBKpCvfBVZ8EhpTq+
M9h2nLKoX96DFBhGTmggQF4anS776BrCaHw/psi3gCRWqElUCnPfjyMwnJLfi2YAF6PmXK9MloRDuF3gqjcx6wk1rPm6ild9/7SY
azH9ebdyJe/lJtYRx+S/NKoXlrTMc38EAcRiZ3dOYUTwP10SPsasY1LQZvKIY4FRLmz8zOiOxITm1IZ/fXZ5cUnNo+nWVPVfPxzU
Pkb7DirBlm7O6YSldMHZK8OxuIZHm4PRhXrqrjujZs9/UBd65OIbLVg/ykiAS7FYmLo4tXtTlqu8XTeqx6K70By2vh+HAuvzC3M1
vcKUaeTF5ftcQawz5MOaX5briTKRNvP8q/iadFoh01Y5TbpVSmzpo8+1gFTCjHaEMSVZOo3Pe3YJqXIKuzVt8X5WIhrVY/CESoS2
0oXmAQgbM8yGykgB//NCXu01N3n5JV2/OCGntIsUhIOqhE3XB3/5hFDjmM0XI1zbp1mDop5it+zVmMpP7piI//FkecX/hXApZjuK
PBocqzoWdI/THAxgqbLO3TZe9CNsITBpRZCgK1dn2NqT4+Z3SXGhcWQXDmCCmnDuNzMQEp30j6MzQke5tPDOvH+mp4hnHndm1qvt
V7WkTTn2CBLDYhE53/+rteT4f2Nqhd52FTpn67Equ8aQifjmL77XPAh+QxQKXU40mfX0ZvgsUHRQWb6adyLJ1NnD1g8+JguJvsa3
T4na12VlqZSCFzRfcL1yOaMogjh+XN/gCrBkoCNBxEOgoKrTJnGFZXRdsLjOfBA3xqTxg16JzZRkFrxTl+hVM95Iq4uPs8Az5p9u
E2RoC5Oz1mb+54PZmWjNOxBQ+BxmudyUpfHpQKBstEs19mbyhDDfPWn/IFmRuM+ntCINoER/up7xlvXTXfA0oNJ3dDZO56p7jCHu
peZcbYVOOfiGannyKLwbgQwPcLEeR1EfpyLSRDXj56xuFFsiaZWL0j/P0T5pMoPx9SMYs4O1gI9+nze6BBeGj1CTNyGpEUCAltk9
0mzEjv+y6ub/By0uMeRSBq+3/amkp9e9y5m1hnpZ7qgH3qoVr8Q9kaIWqEX6Np158O4o75+DhuMldma3WBaDorvXXokyaRmI0fp9
VmObQJWYK+Reqdr4VgFp7rT/DHkV3Bwn4twG3c/8yr+43w0cejs0JOfe0Y+L9p3Ld/CtXZL8pZWCeEL2kM3jRND9XQ8/yoG3D1WV
DG1j53Jij31Y86UG9ro3dJtWldT3xwRZDHLjOwv2PYMfsiG7pDZ6MnGoLteJhcV8CQD9l5T/LdwDLwCoMGqeXMBOP/taAJ9lnA6u
3kanh6BedIfhJQ4y+LCftqSO6vzo+GQbz0ae0slxMvGBIrCuko+Q6LyoFybwS9+z4fQ3w8IxL++kKWk1CkRgfhzLnIL2mOhb6W3b
lc9fugLdi1iUIx1O5hFE/yNr+Eyc/i73SjNrQbVwG9Sd0rq/A5aZFmjGIxQQ8O7D30haAVetojlYbf8w6ocDg6LCSY0Jyc0wqG0F
5dU4bqudEwCNX4TYrU3xiyXnjP8/gu4A6JrpQ1ZYWD0Etv/ZZtH0z8ZYLXs28tXlzqzmimq9fqlWU4YjCfS0xQWZ9gRTdfGJERi8
39Uk/n07/4qXif/fH/a2NKPrwQpsXpXeuTwevv8LHHfbCYPbf1TqF2Cdy2S/QejLKsPXjWUKErwKbH34+P0uxlqucmX8Rm1s/fVn
qSrdU1STRLPpp+F1yDLGXzHrfZDJOnvuYeWNGDLenOK6y8BlIxVKXNXzYcffACnB7vhVM3efCP2IevepwuNPwRhCEx1sqs75dwrq
HMN+LmJh6ogecGqSxmlCeqGisBAofu/6msbXF/du7d7JAt4XGOzlo9GfVFG6CrKhu+RviCFjHVjg7VKMKSFElHxRzb1BQKPc2kHz
H/zUnkJ4+qJYLZ1w2Y+OKtNe8udFVlgxa9+ChjEZ7ThKVeh4HGat4RyLn7XjDougsH4YoS+lf9XXvhexS3wOwBR+Ntw4V5RxTZvm
K4Gwme+7cjLhwh2MiTf4LYjmgp+ZYPZ2omP+BopH+2c3aaw3ZZzA6OhHBdVSjjfyCqVwtm9Kggf9lgwpjQ89LiIBjx8eaK34L7H7
mjmfqAfziMIVq4TO9SCa4nXeX6sn1ik1Yv2tijWPdNG7yDhMbamGLlGo1dSdRIF427ASZZmhvw0iSnMPZIIwd1RfkRiy9NnWuQTM
E/G1sd/uOpG3QsmWh0aubiYJwRMp5STI5LRe3JTOtF5DPjbj3UhALug4zahYpNm8XJ1wabfDRtHNxY++ZIBlLNVlAJK74G+aIZJz
ZMAoS596Cdwcta16GI3W7y6Ts1VqwVtRC/NDLbriVQY61LP2tVNu6ExTAnrK5XtkLSt1vf99u/o6qKYrLBaHrPWnzoZsZkoQJPvv
0mKwjYMSv5ygH3LDUwcFcDRn7isjf6a8r/otLQ57geCW+a4Gpx1G308fz7XOnCqXOu+raG5a3PrFtSjtC5Opf2RwvcEd9LP9x4EK
CjsJ1wZWERGlc1FXZ1k/u3kDIn50fQZ4rc/D2mh4hqTO+Kf0+aAGVdL8tBeTiVrL98tRLl808TJXfM2jjAPKodsn7/lzeUDPrk/g
8cdyT/royh4OvwpMKgd2eTYDqM1gW/Tuj4vY2ljr9TEostA8l3KjMq0Kjc6rQ4pNJWt2KfDHG+HYgCXEvDazs6dVpJN6wCFHfJYM
pAfms9Av5Px2keVexhXe1t2M5NWhlVqN4qxd5f7EPCc0QhpkrrRKxXQKjLxI/cV65z6nENNLWiTIQo14wQONTzjSNciMt2GHk2KG
I6n9niL5dB6E3Ww3SI9wjRc5yB0rAR5d9IoIN2DmpefcNVlGNtX6e01wLPediuq6BXnQTXLdKnzFRk1SGi9dig7fIlQGrBzqEtgq
MEHqHrvIeaEUPvzxNf0niuCddThuRiYmBU086n7F2Pv/RqxlpkoiFYBwEblexlWAqUkCXtH98LwFVbNrKqMOwcqNhCTrV52M/iuV
CP0z2XaACn6Uz+0xRilr1DqGX8DkxPd8kOaiF6AIf8QDMUtYK9B0mAd+j5AE0qHiU+pDxsVBcti3dgRDTlseYXLaW/s4yBlr1HC5
+pLxMY1FdrSA0FYS3WomUyponcBrleok9kgQ0alw0RdaV+Bh0utVxY5nJ9UD76XFWVIuZ8NoSnXrsUU6hmNHOVTAfPlMkwNjoC7j
SO+TSGYeHNsSuPVpddpWU8SXDhxDAFJyUcpv4bO3SbXwsdfB6h5KHXnYFqHJq0NnES70lbPWT/CU8O7ZAn7vb/qHmzn1Bh7ROyMU
1APNrqHDvg7SzpJG/X3LF9iWFGZGVAsK4grAdE73SoJHhylkN0QpfQHo1WYedBKHhKutoACBKXH7TCllwhu23l2eJu94pos/rR1c
hyoEe+gkX1AwphvxZ9r03GYchX40pAmybffNeVBb1mgwL4rzzV1U3XVzqr/Uvq8A6Fev7+uHXsntctOPdTV9t2yPzPbDm6jXUoyV
XmR5wzwYszEEow81R2ieaPVfVDUrzf1iAqa6CbFPu38IC6rdaxCgv25iuVJBpJl4T1Y20+oBc4xAHQps+RGOvyl4yfOAzVf/6DUY
NT6venBGiYLljFV6XNr1CM8dcNZ4YJ0t8h9E5dCtnP87544iWh6++sQYmDyZyY4eNnbC6Xkw6cRQ+HxOv8xSnGDEol0GOtgS62ZJ
iF6J6Dpp4Sf9owpvoE2hQRSpHomOBXAfUQ4A4HAVvG4ocX1dfwTA11eN2y0BQeKPWKQ29RGowk0I0M07r2WzUkJFK5KY3ehcGdKZ
5D8gfdeICgcgZeeEfTO8Aq1eRfNUSB+PYzaqYGevtz35lNUfZ477YDxsnQpjmnuFvkuVIb9/1BB1LPD6Wtc+aWu4wqOUnD9ODUVv
p25oQC21pzgrn+rkw40XDk0pa+2+xeTFPT39ZKZPVJsm2pHXIIqpv1dluuhSIOBxObF/OP6GtSXpeGkWwymlZbWcSJT0gNaioAsm
vPc6Gkz0FiJRqRqLjKVVbehZ/5LNXNvMchWNeVkRiqILxLW5726VJYPaSFD7TptpogrHPIIaFEQqVxU/h6Ugv33Db8gMxlC0W9hc
GXisvR7gfJ5BdI8DvkczMHocprHU1LuUBP1yT0xyIGdbRokoIV/zgp1ElHYPCW//MiV2RNYnr42Fxtkt6e/FxtyU/8GbrC2XTvow
Kyhtg5dnzCVJAQM0sB06XGtq/anC3pKYKri3aBiBW8sjGc4kJHqsGwzJny07GACF32ohfB3Ja7aWZjVNbCgNknCkqXUfyFGJrHvD
kJs91lA19pqQG2Uuvn+u/yZbnWNRU4x3U2h9i3nlkN72IoVcMzTmGhjRKIqnMMKNVxqiJauUbsxakp8GWtdrg5ZnD8q1NCxffvMK
S9dhPO164JSDNCvyGMntpKnEzoV1P908LKS9whUEIItLLH8W4O9KE/qAl3Uy4emn+Tnpc5h5yRk0PzOZQAD9T2fuFoFAxDcOEBNq
PNFmXNeC9PAPPfYxxoOjFA4y+a9H24Mk/oidPSrzlenszTg3N/pkNSKyY47vMEvzBZ3cSk1IDh+dKuQs4dNdq4+UIvOMKw27B4m+
S19wmUK1sjtCNEj6I4mo9b6tFmyXjMPOIUWXgumtnEEPXk8NCWQjRdSJBiXO6w14tec2wznQwmmV4FDYpzLsd+ILM+RUr1GCTumU
IEJDW5yiukwg2xhBgaoBidcUePJhv32Mnp1vNDdhewJOpkB2v/Gl0S+GuhShfQS2CoOo92C/8T3mUZWC5aN4Tdv8QZSpqgCFdxLK
1E9DBIxNgaUZjqzEAmM4iCgWKwrU2c6m/pkCthiSsQUtBwNdXG/nc6mJ61tpn3Z2ok8x7obfG/G4+OgUWbphGg00oSIohrlx50at
2x3uiUfXbF3SzDcIgzmyKOWeSys5MVefsbv9aclvX8myVUazSrrqFByg9PQeTisKLDWprxiceLSSwsEu/wZja+1ANg4eqS9pfoW+
1JF9qqzE+VWHiCEvVAid1QJAXsPRW8NqaIBZ04dAYItzkw3df2DQQyUlpGtSfTNsZzWkJKlqi4lPuyHieXYZpeY8qY+/tr8jAUWj
pNelkECoDIi0akjWOth2j79Yh08yZKrwFvmgO/QdQy7Ezkn3KpMyQBsCVNPRkdidJXe8/DVfJ4XOYEYjy57hrBCN9qiPYZg7eJ+5
fvFnoz+4ItVfAxAG6a+THaeAz0uq08/oc5S/Du5BDZ1vRpA3i2kXPmTwaE8tqos/14oO+Gj6YwMgD8It9mEUVdnRW6o869y6taVz
rxUeuyTp/ScMyqR3YnYLeBp+FEUr5eggYEcCC5hWradseSw9Apte2cPUwlfHT2qlkmPzzKFctL8EUdSGuL2h1YJglT/GsZfE+jie
uZ+addp3/McZUI47eZCn1GfpbLo2hUUWh+tlxDcZWMaFaWWSp48qacy4uTmhye4Jb/dr/JTngtMn3GybwI+RFyxmg+XsLdV+10u8
yLpcrbZaueJgy35ON35cDVlpWuKOlG5GwtQFZxoCyDx6XOISRehWu2A9A5IAGGvvLw46XdqVXXScRu2RKdWEbBGglZdFCt92ep+Q
wzdwbuZSMfb3WPK9npI4aRRf/zUfirL8UBVItrq4OO3khktSlUT2EVaYllKA79by8FDo03804hPacFD9DL7x2nC9IB4BTTx9sabp
Z+yhjzA0Jg1DiS8fQxT8UlvsZpOUVmzTPCIBG05jwdgOgcMHw/2biSoZT1QH9fJ2188QY68YCDNW2pEaXdcXHNn7Y0/FsUJYr/je
BtVHvkQEL6POxQwkUY/veRwPOYLkZMWT1Wjz1UH1R02irqGO2IYjlZw8/OIfdhjcJ+kV2zNRu2IAG/Wi8TPCipXODe1OKVNFMUZY
GTV9+7qP9Gqeb1fYSEUovUzJ6js4UQjGJXvMajemcSVTvp9czIp2drp5UAGzI6Ku16/52spAe4Ws02gD8kA/Pi3C7LNj1DC2le5/
07mGsaXiT+5yiGJtVYYVSB82aaODv3/H+IAzR/YSYKvpJkV4v+h8sv/g4P7H+lb+exHCn8HYInhIjlmm3L/cajbUV4L7nZQzAnAs
kYN2keVvcSmKU2G/GmwQjrl1WStUNsJOYfAS5NSV7Sav0+loGu+Xl9PfOl44jngVu+i0tDfX7Q2lBRL3Y+gfv/yNtgnU66Tpi+Ca
YE87rgayRJ+RAcB39ifkOOqRsKSX1+QIkrM7ZR0pOIqZnMD7SjEjJN2T7ez2R/cOz3Ff2cs1MF9dYbJKcaO79SosiHhoB7cQ05MO
eLuJQNAnoZCj5qbP6eIHZO6ZTNa0zeqL4NlmiELl5Nxbge9OfWhFeYUAw6BtSzISFQ/G8jhFRYynDHimgvb0ofQxznPbELkrcB/7
RlrZ84DNqe6/gKG5tUP49655268nDu16H74Ilu24GAT4eE0tuKjR7AUb04LG/dZPWB1QUerbCaKzY4ARVnYAdEILXNQ/z8noByJS
yPwIMOcwqvjqCZjRn776NOUpqLV6PI6yEcf/L3K7zcRFwccsxrpM5OwFqV8SG+S8Lt8IhGBl2IqqHwUZCU705qjqKEEjJsBAQX2D
mlTuWF+ksZNmydu02VMLMQrYpXL2o69KfedXJ6eEnRruBWcl93kPZwVaINzQTkfG5FCA3AOqFZ195bKc5OSEpsCthQe5ngPBV3WA
VQnPWobrJqTKp/43hVO4Ez6NEf500/GfDmeqogF0ZPpd1YyikViU+vvwdKzbvQMKi2v76lyf+pG6CQgvRjVZphT7VKOrs9vID+0d
kKtTyyOTptLsyPH1t1BfYinbK/60+R+k/8H9iK+USjB7suWxiYKGgaIj+0ymElCCBqgE0GFwfenUhYoaPdR7+uCmKQdiXQrRD+xf
xsvyNm94FwLr1g4BwWeaqDyB9f0x8utAhv8HIv6tId8+dntLPQPaXQMRz/qrLU+U2vdtqQASXi7mfl1sfXb7FhC7AM1g7UQPZZL1
zATmPK+4b86OiJEMwnsTUKu82x7tUXQDyOedhVnOC9yD6xbu8Lp4SQLn/+oDZvDg6HV3gmb9IuQzeyACvH0LLMnEqRWlRW2tOfVZ
MDsRX1ZeYgDMWK9mEeQAsoraI4TAoEC/tDGE4YPenygq/Na6pzmOaIAHAWJzVvukcpSLtxByy2V141sFaxb7vWZr0ZB3bjiAMXsl
7cmxIWMrF/AFNhyT+u/2VzCjzVzlch/smzJoR8IdyqUAM3RzMqoX8CHItyRTXQrWi/d0XFPUneKnvS1tHIQSjxYImLzPpOidMjWM
iFnbvl1PY/Otz4G3/fNH6nr79MeNih3j6SHLTBHPNWRGDI5iwU6VUbfV7ZnwZg96dNmxIcgsTdLJIDEdRWp44WQseQbuO/4JvSqo
Wr7TS6G/gZEka8SXsvB5Dfn0M8M4c+cus7iRDhu32hXGTGgegBAWu6mhuu+X3laN+spuB2ospaW2DvfLQgXVnkIJb3++uOh8Aild
FjUAs5aT3C/yPUrrFPZaYLhf+7fFWzB1lyqNrLP4F3i4ii3gU7ueivcVMdAFjZnJ1MgOlJ2OlCunQzElD3+7/v/ADtFVT9Et6ECA
d0NHkOeMNihEanz1EcSyB5dIUdJDpMH0xJCo/RF8hIpE3aWFCrVcYiIiYGv+sEx9vAyfwa5KdGNUEqQINpgRIo0Jg+2qd2TZgn4D
GHVZRXIZp+2G5/Ui/gPPqbghfcP8ib+lUJp838tdYHgV/RM/L80dbqZZNpWFU611iwVMfaRt1pev5Dk4LsViJ9oJhgxRC4y+Te7Y
wt93sSdyHFCtT+RAYAAjxK9goCJFZ0arDNI+R0S67VtUdI7NQodhFPM+Q8FMs/SsQfUYxvFDn00fe3YUWDc2b23/AnPWuoKVh6Oy
/Z1qKRnLO1L5OoEv38NhD5Ph8oQbI4g6OnkyV1MCrqILn4stIERrMcwxZne6j84wz/OGd0fwg84erdy3YWP5MiDFIxvx3pWf8dEP
y9eEGm98NhptKxv2Giv8QzCiS80AH/N7JfvoNFnL/O3RDsazFm4isGI8/1L0yPmtzoGjZsa3RcFzb+enj56dB5dwltdQYFK+VJtD
C+SEYdVAEbEUmynPpmlnAlAI8iaFWzvGSd7LDTxcdTISKu9d7NgrGFJ8lfG8s9tkYirWexKumfshi5QJOIUJvhBy7ZY5jtXySx+L
AX3rMwoNoCeodv6z+SDMnGT8Q20m32M1z7kLzfSS0TJs4WE507lnX6R6xqno3C75grfQMPL2nyf63f3V5TUIi9cRkFC8BHrdE/IX
0DoR3FjKe0jUeVq2ITxNhA7jORJzCvT3wnqf2h8gQeIUVm1qCQ0y5+lcPm4Be7fFJuZAu77khjCJsIrYM8yqy8aLcUNFtO5XXFlp
ltt2U5p32fOBNwjHrqnBCnecj1soME3Fzhhp7JnRuGXInTxWnCRLDEnGByLWQWByCDX2AWHVujcevqN94c823Gv1v+yJdpPsBbht
LCJJask6ZxGfE0UOBX2UzVDLHibLPFJG0M9aXCkO2kODIyhAECL/U+gLqU9hnDwiKmuNuzv1ktxbwU6tO6rz5OpAYVuBLjnk9xK/
QJu1NitwIFbYhk9Px1zzj7bFLeRJP7YvcuSP9Rn8tKtU/58unAnm/L31zP8iS2T36EhTBTnWYZFFiPnkDl9p/7tTXQL6T+T2yKC2
GWnG0FoewLTcL4ySU3Adbz9fdRdYH8ugTbEPMeq0Or3vnL7uYg36uzBHwA3GDbdvYcPBHtiFfqAhh7mV+GCyTJy8n5s5ivAjaj0s
3nZB7FA1MbjOl4doQOSdPqtpRFBIbL3hps+oU0Y+jFDZBj48o06FjzGhs2PZPgSsB/rBwtS2vE093UWUeBhWbi13r/wGuakF36cH
N3Qb/+h/QYRhZBZJZtOaw16E+/GolcmJVfUlI2MRAebVH6JBzXnIMD/8Q3ijw2WSMkMTSMb9TNHPWRzgZEvfx3iW6NP9j86Y6NEH
fpSEGxwjd8vSQnyu83PE6V8RpDWbEE3fWCd3WS4MJn06zNU/NCs/EZWhQ+ra/fsY4hEjrQd/GYJRUXLyj5HyW6cE6c2JNNT0ijxN
BWQg3K+WHdvrSMRRTXhTtQ6t6NrLugJKCRu2bwA9T1jSY7ujkDBLo+7gq/0L8leM6FicUvjF+UT0ASlxu+bmSzTMpYqAByZ0V8ck
U02SLRgzeEaEjX/nT5LPO3PklqhZkcwF6Ki0zyOmOWVYgjoLFVa6q5uEKY49r3sdZrV0vgCLxMnpijRx/YVyhFCQ2DUx5mVVdxzD
0GwGUfs7UqZwXR68Rkb0cLnOgYG8UYbr8q7+zxr2wceBT3FQH0+rHuQs39trl0jRsP61pAwhtXngHwUeDdV0b8sTZGmJ/d1YLtxz
MZA61U+qcJbqvxGx5C0fLFOmgmJADMVw/aVQDxXYyjbynLb9wsbPC7+TK3wV49jGhu2syklFOu3W6UId7cKXEC4bDgpmeiX4ykwJ
doUTTk0byPaUEUozhSA1WB6zPnUhpLh5w5Nm3Y3Qqgi9TjwBZNL7SS6us3igGTdtO2Hagt+p/p0TQ4JK7FMbmXT+i+E3OYgSf4EC
XsP3ctU9xaN/RheqPdkmsUzAyJiAsBleUZ5uxf5s+Ufr/mXSYzwCVO7W6HomSZCrqyKEBy2M0DL6yNJyyiPgJPZjhf4pe5KXWiEX
aYf412u0+C9eGbCKkkAKZAfu2GnxY7UvEtGVZx2XGKzEl7tNB1y8qiGhbAnL2/iOo/lRCay+PaAdtMCGd6+KV56Op3W5ST/iQDpW
osWlNKmU4fqp29OqtVDwdtJSTNj8IOnaLF36ydHvsfFaj+j5YXj2GyEvLsgRvrenwDDMyX0gT5CpY6vPBLHQ/LJtjodTMNr5KO2r
ytjUCxkch/Bl9MLfzgA0Ux3MjfixdzPZmF48UdP5nKISpjISviocrlSvWcky+BmIWwMEsT7PwWURxJ1MFcN6a7/VxKmS/JcEn1DI
ob/MX/dkXfzUcy0ePThnHHgKwV2zAASuMTHXWMnopBOoycWhSSbr52YpilH1c2f7k+v6oawcNFYIwAlLfayzahuvpXx446eX+Pdc
+D4mF6kzN5EyY+2EHQ2w/msoX8xy6s1K1AgmgX0xX86Vlmx8xfii6ziXzhYIRJGOJEevnzh3l5b4uDPhFsZoYeNUfwurWKXaQdYr
30w4MjWnSsItq08JeUTjy/eBR/Z++5WVpfiHhcsZsniIdLxXgxI5+mJz6FM9quvDkJltGvjddPIZKZzi7P9phlBLftWmjGjK37Qj
RjwFcRzbJEENu/qJBCo0LxGzxwOx9M1LYZzdWSd16O97qBz5YI0KTEOyNtofmuXLGSxI5E2TgMtvlM1c0a15/V6rwCHrmhF7akjR
dbyKm1hOFZUgbLVpeMV3qiyanq4n4/agsMIian3TRO04IM4EY66ZjCVavR0UTO5kqyT8YKOcLxz8mjxWg8vriQpWXp2O4QKTuHWq
gsF2OPh/NMsbFxMP2Xa58rfKnK25T0T6cOartqgxROLbQZRuUZA+xUshL/OA7iOvl6qBpS5282lcc9S54zIc3wuaaYuuUhPfMjOB
ksYuzzSL+rMF9yXf3ZXfh+8eTfCkJTAKF3BKZOOXchfiDK4s2j/Wy3oHbrogk11nb4Jk0/g1TksC/zlZUQ0q+f9LxrhdmB3OJOBj
s4TKFJPtHsZNPCitdre/mHv+RTHaz6Gdt9Zbg/IDc3m0ZtN+jFJrnLI4mVuAoVFI87YEoQuJRX7bg+fcQK47jmt6/h+qIaZaZtuG
n3Kbu/JVnM92lgwgURHHM0KCwk1c4aaNVwFjFJLN2kR3RpUsy9uULMtSQhJbbNQoILKbCsdgGIpcDPkk8F5tp1/2M2Sz/uq2QCKY
JdH3cF2RI5SZRfGZ1RV7LXiAn5WoKh1WoeL0fSgP4GfX8Y8qs5dy9vk1b77RsOUBFGi170ZDCLFutXDLbLcbpY50YNSCIQ8g/xc3
oHvOaRSCYUd1+qudcgL+tTqi1aJYrjPaEdIN9m/FI62vOzjaVu4VFcAmeL4I2QzJQ4q7nZS8gG9/qUhvXh5tB6raq9rqvAibvGbc
JHolNwc4hqNQgQgXhLs+3JMdidKZtLPJLNWW4s2vRAuu4O+PdPqNEduEMdN8BL8TvLVzUIzQwCK0O7dZu20q5nHFVwssnDUnRR/D
1J4npot1ev3X9f99Su8NUUtPIvzF+v/94+7EyfhNSEIEF7wgzNJDsXwK6XZKjjdXBsxwE5gdL9EaHb1No8oN06qLnZLsxnzHMxPq
4YTxL9v0q6eLX5xpxWRa1Eb3iVDIMqUb0QumkutSlkK8z8+3SUJbo1v6WFURTfcHjdBfp2o49sOpdy9sC5/OgFzzbfozSdbDAS7r
ueGpskOT6IYu3jwxyu/QjC/wzhqXG8iNQ5gdJlhJuZFqGTniLb/c+sHa2KLKBDQW7AtyUHnsrsqsc7tml1W9IPZrAhiJRlhZXs/3
1Qu2xJHcNl9VNPD+GxWwwRxA5BVq7tEOY0qAsQv/nnPqAsVDsO/dBFBh4DgCRAFrXwYZVoXcQEKnkvaVIpRKc2AISBrZZBAANWEC
3DpZuEP5DZKebYRFQaQMYt9BIOFTbfC49nNFwV3QyMrVtA8tSfWjJVa+sdtrM+FSuvfEEkLkdjqD7pLI9PFSmkOYjB26PFWuMpuO
RvRPaz42GbIuJ5GBGxWDJgURA/qA20OYfclzvZr+RscA/K1qdRIo+8rlxQJPjtG6icSMJWBDn1klNFiDEVpxNgVSTvcXQkid/VuS
od/3KdduiULb4qclZCrh5pHKcm4+HKFxCVt9iSuefCAB7Fj1zeXv6CzxPckl0VsHBuiAcvTy5n+fIhpQuJdskOb0/0gCTvanX9Kc
WfPkfCHJ9yeZ8gIWLMDOuMyABdFNL+DBwocq+CFtU4K3FUWZUktldTvLLvWGzR0lAEp9Ksxqvi+EuriVT46uHqPOwCVfkMDwnAqg
I/dEO9JWvNBYgAJrm2vMJrniJnw9CsBn4oWBOLFwW7qNlgkFUGWX9ALRdVGv5XHLYs5LV7vB3sUARLwP9TWdrNyfhUG8Jc5H6kIG
r125Ytd7cRkvqeophlTgEA2WJG9eH7pBOSJBS1fBSeXEkQag7DqOYtuKReyQk/Dd8hpaQF6AY9blSS3az52adPnMsqqcjKyBpS5R
p99ifYwLL7pFF9eL/aDyaXkzIE+JWxYOftFejX31fqjMS1TzW5NgZ8bUXm3thXw1muQ864F2cZx6+Ao3GyOkMfweI0O5QGz/wWvY
gbPpfZMR4gsKXMwjJO4GPa8Hxg/iL3CkLItUg5Hbu/cYb0G4oPHKF5XvzYrjxlAuSKhSTjSKKHi7xVVSfTs9IULNuphuuqCkWbFN
gvyJmUx46pqL74GCz0Jj5f/7nOsoZUpRVhGU6dq9LJuGtLBsbnslK5nsq5Uw/JwlheXy3q3hpqPcjGOxpbpFVdzUvcNVx8tc1BKQ
qAdb8djdRPHi2yQm9WaZBv9NHrnAtB/7Um1x8nG//q8EGyOaNF+gLL2DheVVUEnNl8wB7Bey7BA9jTfS6T4KJd+7QUP/et5Pn0yf
kTPfL/y8ZBSi74T1Jp8tuRIfg9sFWLqWwy1qsgWViBS5HBLdf5wCm9NQdjI7Xv60krGBOLkEEAhV9Ju/KCotcoZKAc+IK0VSXLch
teelBfnPEsSjIof/7AC9mJn3EHfoQN0riHgZ7R9+o9dpeq9bN/kdtYGa8Bx/oGl3mrCYERzsg989tOw/5mRR3x5NGU7GExD+efeR
NrpAxuOEBYmLx5rOEG+kVwCT/Y87BE2uIpCvk7dAzoN/TccuKvxD0rHRVSGz1dXtmJuhUKS9U/geiRBe6P38eqJIlD40kbx3us6U
z+JFUCgZur22tNoOo7o9NUg9bh8lyzgsPMhCphw9yguIgTYgpCoSQRGBpWaAtWPjecqsjjdzcC84bTHqRc4YCZu11xwBfmFaEt1t
0GTAsu+xre12zMV73NY9DTeVt+Wz1gUMhSN8T4GFN/mQXhalWrNMv51h8cjaWbZhlSBr9vZLv73UwWrJpvB3D0YwwMCto0oc3H8s
5Ntw/zoTkhr3hP7be7Ho+IevaaCP6ibJhW/zx27nt8UMA06KwgAGnXNZnvYUUbLc34UHRwAyh/kCbD7m42sLdfj8SmAMbpA16ad3
uJ6n68m25o6C4Lvb4A0s4ZMGsR72n39tBbJL7gRA0B7jXCY43Y/Z24bY0BmQGh2d30+jjgtb89VOju/bm3XK/j6NjgABb52n6Rvt
z9qArpJLxEbmCsdJU4W8uqiFcoWpU2IQdpzRzxn3SKkmc/V5Y7fEvNj5sbTQ+CsAA8mg/RbOkAMcFdn3sdF364qw3mwwK1yHbRVW
HWge2fU8+MrLKs/uo+5qf6u/v0Q2fiy9NxsSwdBhVNuIqduQyQ5450vDfZsaUSenZzFF+bHd2NiwxsWWRkVJXO0ShEQHfRGGUtKm
cfI12Iait7wgV9Qe4rwFrThJW9Z/CwOskSveyNAhDcaGrKmrGsLS1wuysXsT9/EBjUFU3UxJoAmJWv5ddimjSuZK/sedNKpdAYBx
cMY3/4WViWvZ1x24UmwhrAgx7rLduk/HFrzCxQQf5UYDK8oy7+hmNivJ+Yh8aZv/gMvW2p+z8li8Uzxn4/wU+4zXPN870TIIrZ5w
7M2Ea0rrUcjmgATBYek2rMVWwAkhwPmsxavifnrHT7CxRsBu1ZFEvNpkOPYnJG0OO/phXjDgsth9zF/CgXfRplYkQrIyi56vzg5Z
rdIIPonUd4EJOdTIakPN6tbliG/ElAB+Gp/HvQcm+UpRuWJGazNgA1Ypl3GGf9CmfaNi5cnZlAOHv/Um+0n9vlW4RsyGu+/lpwz8
zh6Gkr9DNqJzm2dwejo/uPPBUep5HCWzdU+QqlVTfd1ViVjhmYSmvAaBubWFuuQ/bzZPBMPJNbt+kkpHCyMXGHiXcLmaFKzau8CK
vstdbHLn9pv4h6/5nHva3Qcn7d6cf5Q1gELF2g033eCXKBAaOy+cZKcIKwIH4SW3cyn6/ZvGNuku22K5Q54FUnK2ERXC251VYStB
2b7DjS4jhqHinChNUDmSPD2ugOCiho02LNNmhULVnCg4TdPB1npmhZ2ooRJKt8xlybJ1TCDMrTD+zeDcTLhBE4rpo1Inu7aghXLS
o7J4pQsFfI794HYA2MVONgLI8lijIGnvrN4HB8f8OFOkLeriwVzkSVUb7Ong7q/cXC213PVvcljnC6EEgWyH5e6dLICBjk3+/dnC
y6+tEROy3wrUCGqRzfd/fECpdJwn2Myj5q+dhyyHXip+bs+Xa/LMRZXsazF+WG7+qMwM9Im9dQQ7gfpL5ZkNIo0Aap0Q+o9HISg6
ZGzG258A8IQVuSPMP9xXXEl9cbszIJWgufX9EHUJ8mK+mDTs2iA8BNDQ5YQxc3LNIGtFo/688lf9Z+RikehrI2j7+q5LSp/H87EB
eqk43g9iyrhwHq7u1F/ZgWllBuPi+5vpBSUZDLXGmON39Ug3qbfuVBhBucamcP7Ra3Gx6nqcc9Di+TPo/E1sE1qwM8K5uaMsGXrW
yK7QMeFgrwh2FcxFsappwZMBYoHNsF/NX3UH2ui4ezdTW712xmAdokiiv44fcjuIN6SZ0xx898yAEZqzRa7NqjC0dT2P8/tpc3t4
B5VFVXdDrmCsF2jm827mtVxImdM/CV+JfMf1qHt6Jmkf7hh7jkqu+4OcKzuYKjRkaTXQbYFO6pzE8L+DpNWnJC2G/kDxn0KkDQp6
OdFHMCO0xj7i8vONgihVjOfLNI8/QGQduyD+Ls2H9acFHT3j8WmocNfGomtRO3W81Lkik3CC2GnD+tq7SwdTriFpbJQeRcGw5zQ3
Tc5aiBGcdR6/24mEJFNnlQilhyJCQktCa6SDaOq7+xNCQnGcRERRxGcEO1QGTKmjX2KBGlb7M5Uvaencb+llov7KdQ1VPpQrlLXi
koBY3AFaxjkCmdRae5DDoKyjfBXVpG7GWLgmTYX576Z7Bbp9ccrn5NBxfgs3xn8MyH1jVQp8Jq7hR4V9oN8x1UjDK9rLZco6cDzg
zL+sYOaLzNHTLS1cdoqgJdDvCftxgIQ0PH9AxikRpu5k5ztbH4ap/lTDrvoxjgFJD6Yp4qSx7njWtiOJ2+Lhf2fJqE3KCXtGV+nO
c7bzZqZwIDH2S0dCvvDYnfRDhvTsr/rR8Mpoxs5OPjbTYyUxkx9s9eCwQxxn/jm+snbntmc00wGo+tSmXy30NwifOXeKl67inzRP
9+D8zxqHJPAU1pjYQ6L0pK6e1gaT1VXDjwvJpM9KjXxYKATC58srRd1vZA5nUVcUjqn9vjT4zufmpQCPXTm9Yus8Pw8ltDWyxPmg
m3ic7e9BIqO7VqjlSZ5suVhv1QB+37QFUbDljPyMdsbaYEC4nEpe3gm6PrxQQUEWPxi8VhW368bZTLvlk3N9rdr6RJ8JL3f6+rgp
FNnd2eyvzXh74wry+EmusqKgxQTWBnbyGXuOorxEX6Rab3v8zq3Y/zC5O/EZLy9HqywFmv1ZItnJ7b0ucIHQm3NH5TN0YJGX4jVR
vEV/zYfextYbo0Ho+AIchQEbkYWzkEe/5Z1kBmjAUFRRE7KypRiY1g+oEj8NhidiLHGtxvimM5PejiL7t9rxyT+SaaCyLbl6rc0p
+vWVw++bCLlX5//AEpqMkpjIzGxWWk+sHfy7p3MGsXRXVmeYBbHpZRd80mkkhRqlqTbANRVB8yoq68cUj1ULwzyk9EJUXlNW6VHo
oS2tRYnRc0ovy0ZBsBwN9faTWynvkZyc7gBDmWmwTfzVWev+JtCdE+T8vl9RR7Y6OJwG1+mDvxrr85ZYSRa4BKkm8LkdWlzD5+B0
1W/6n1mO6Z5E3pbNSDncZj/VPTgH5UOQvvRJPNafXjFQpZ7qtt15MT7UZ1FluyeNd/ulOUNpy7/she7X+ZJLe+SISildC9EvAog9
lawl3zaZc9C9Arr0ff/aGotzrgFQ8w//qmFw0TCxkSyZKWMDbNcGYmgvM3RseUfrf+8x1jDfCNgRXLmW2pOZsFFiRz87ID5G101o
hdKlrcEEnrnaNlOhqxqL2RvXvYVLN4MVY1v8bW1javdRNRomv88HvdQmgblO5d3RWszlvbUYItLGqSRFA8iysvFPplxEqFs38E+o
YU+sB+02ufiqj4t1lJp3R3ESBRsMgF6hmbxUE77ogutR2eiNuhusybJzwc7b9+out3mhtrNwYwBoxXNXbmyK+3JS/5CcB5zkR3lr
UypjKHUvRSq3q/3nnlob0Ed6f1l+uhVoFDZGlPKTxvwD9tbbaZKfL2wRL8Qd0QOlyDroEvzdUt/mMwD+pfgXJQwpPNaRIxubDiHk
kl561HDDKhC+vXM6PlGtkOCWZwEVd1vrpRvJSvpkWgUK3uEkOrIlZfzvFIBhPQM9YryKP63ioDubxCA42w78XHiro6M6BA6Ri6mI
sPihEp1FdFlkJIiRpd1aauUWFFQxD9X3LTb0fBUgCHEt9hpFuEd7nBoAbJbSIIWyMXe4G9s/8zsoLTACQE0HQ7tel9f1zBYo+tCT
AnqptWf5y9rBL/CrGPqICq5O1NaqXc7kiSUHzcq4uZi+Cucv2vbHvazv7r6OfzazU7HeXNl0EeVT6kMFm57QzwlBzNL8dHf9dAW6
laszx8SU+uVJ7rej3hiSgRQ4qREmeVW8HY5+ua1DRy113Oi3oD4/Hqywz3FoPemiYepfabgooaQNJRAkqfQPTc+DKD27tNRVoCWW
CMRIecvqI2dvd/FhURtbt+En+I5B3DRqNISnyXPygNxqZdFTHDOltL3k/aCQWvVdEXvgVr2l1zZK5gANz/RwqjScoU+wUYcYMs2f
ARUT7W9ql1ZcxgjG78jFqvwR6uLWRZ/aph+BMHTL4sGMnPrnOra7yPNZxpFIB0UyGsZzRAbPiyIg5tmV3zVJXVHEj/zegO18gA8M
E5vs8BPFs8Cdo/8oGXwU8W7pNc8o5vaXAjNhkMJFMAArQFSz2E2ylHEEZzSp2wb7Ag7AZFnAq1bFSNwF0TBP5x7VfOpwGis4eAkT
NiWpqnenOH3tFi+eja7y909Dn/p4zbjPwIVD/uw+XvmrwUYKibWZAwhz62VpzeDWy/+HiPs3/Nt15NkY9ME+5/n/i+znc973izxV
2G6fJBssTls/GfrlvfCAhu9a4ndBTi11wDLfVcldAXxf3bnOrecwpudp8dCtnClbS4d4w8LWl/NX5nVSalGVGtMJCXAjnqsWKWS+
d4c+NaRhOqWeiveNjiwtB6J5fOgaWGdD8gezAgVsf48KqBz4razBBkcvfhmbwQhNIoduhvQ6XZneN0pOWa8fIlFK7nY0pCyd5wa8
jekZ/uYHt2Nd+yg5F68anJq1BbyzTmQ5BhQwLvCd2OKz0IjQjPDH0lHaTTx0Hwg4i4tDw477rde8rhoaNTqD2HAMiaE9+Bm3cWsw
9xsbyrL4o/Flj1FjtqCAHnDwpAGqNsFI0STWeWOSpkNoWIBCbtVNhv8pVpPf6sj9KevS5p5NaeImgJNVQhcHK9hXjYSZFMSLRJUy
vu4JWg0eWVoQ0cBxNPYyfLdVcouAoKduOhO3yimfoqF9Al7DxHnauKdFM2IAd0vUhuQtCpN+NCkTR7uw+c6hWXsStYSTRpifafrh
AQXpQmppHoVInupEin5kCh8wCAxFnilKggHssdYBrot1kciloLMNNdQEF7YBjGxPaC7wKt7LiRQ1CKlWWITKr1gHsxXxB4KCyWtS
NEu9A3pS5cejchiJEhZs16ITey9mSh+AxbydBn0JMTALB5tUvB/xveWU2TvJDHwHesXEtmhFKmm8kNPg720H4MkLEJC6z2tsxHhZ
cSIu4y2PAVnU/DZhnXmAAKJ4dci8Gel5s2R4mrWTsuNIOHYMuyHmJEpAEepnj/VkPBeC0HLqVLQRL4a7DlN0GnOy7oEw6XcvHWEc
iuxUuav2/n76/YvGAMKj+hpnBTAnvhLtuODAhVSx14ulPnvJOYjt3iRTFVyV4AD6s2ob7KgNHbksQwfBsSAyohk2prwN9QWvqQfg
T00zaoE6wnVjV9zexatjenRCT14t4ESE7KOQHMsJAIIHgaT7TRAXlojUusIRephdYIY8TDjLIjwcHov0PSMxhAOZWUKEF0IYFfAU
TRVVHSaa89yOL98paw2YUIHr+YgQ5Dw0xTV1WBahm1vyr58wJupSiy1lmjSnQBCTnCr4arXU4ZZAJPrHQ0VGkP0iWu575QAenbam
ZS8E1HAzF2KzUWzA6jfVcrh6fSTxrUY79EcLePdb0k3QFfPzBQuRlTI5iyxiXzAd2whveL8UFaDWOgA5SSmOFlMaEMOtteia1Aec
YPvW+Rowgtbpd/BZbu3vozJ8J1NviJ2p2N4Sgz3tiEuMc0dpuGs+Pmvjc1L7XtB1D9IrzcZagYlljVjvqTSw7IhG33K5sYvDq6ui
4IwY5nZu3ZMzhdoiIPCB1N0kAZL7NgHhJ2lvrPbchVlhb3k51WyGkBpMFJAcar4jrOcIY1XrhGzh4qYSlLybZcLClm4BYaRnMjM6
npek0pQJN40lE1QiRa9MJBuNRRT31osoJtNcE39IPqRbvVD6Ham/VOMvbJzdf4E3NUi7rPWr2lHOdWknIZmjNCZcepQ/ohUaEZ+d
L3ilQudXCB2q2kSgcCaA8FgGjjpGEE4KBXLjrYjYKeOvG6bo51JSGF7FSE8yS/vUeh88V6R7SycTdvW4nzuOS9Jnpl/oNtP2kHz9
nMi5BYV9cJkpJExDVg0iADkc4qJxQNT02S7lSA5aiewbZCQ4dt8Dq5ELlTM4PtS4Fz29ydEsTPrECn2qdsKkXpy6TPB7rMiihLlW
DHmEH02lRXZQbub+RN9/t1c2i2vgivffGNSDZwGNGkrO0yDXbELv0bqVykSQtKXSODI2T0SrvoEY+TlDqM/DKZGTeUgz4kutEmHa
3L7zBy/WNTjU+JJzVfhLnT8oJ+U9mPLrvEZZ9EebHwOIq0XgZVW8d4oUBjJCWitU7V5au20rcszvVb56sKYxDK5p2v+KcNI0t19I
c9HsvkZn9sk3NO7Pi+4aUw3KHHOaYzr1kJ1WZIvXf3r6ll4R2usru5lpmTSUgsRJN9FHeyznXd/bO2/UdYLKMc5oaHduLw6zSuIy
64xgmxbSqe5OIfCdbQlc20SNQuZmodIZLkRzqldh3CZfUvgyEnKOy+JDGDMAGTYsiCUldWMFOfvK3kUaF2dOuzFUKz8fZqwbsevs
wsfcVTAzZk01guCC05WVdwQAxTJwB2hEY0YvxpQ2ReRuPLJunvld4lQAoGSXmNBWK+Uu6moIFfrPQKvS6f4Yc8aO36E7FD2abc9x
mpvGbYIKsCq0GWw0Pt/xdPueNqQk5cx1hwm3uk/tBE7SDt1weGoYWJxR9w4pblUSB61oC2QWTT0DF4XVfb2Uvp1I2jVtrVsQ23uf
VeNGiK6TE8Q30olaJ7QE4z3Wsvbb2ePDSW8aNQDOA32HkNR0WNYvy7VD4fK9w7/yubMxRgQl7pkamPT/BdqCRenglbkCpPtROy8b
yqAzvgGCfLlICvolHXLCJ5+/ZHAhKaPaxd0+G7Zee3hpbfgBp9pPuTBWCUHBya1fbYVhkKilMwxURe/Na97zXWCHQQ0FNTwo+GEN
xnyeyCJqX1H5Kdyfp6oR5g+/ICNryniFDkvjDLUd67lPvTmg3pjux3voexnmYob+XJmYf8/pv+ivZfc0s7/SfsznknhhQXHwkeN0
ZAb1yLR7UZczf6Pa90wueZqaCCJP586Xnlu7v4hSy6Ixv0jezTDsNz56327EV3gQta9hpAMG5Ko513ptny8sWrBlWEl0eZwY9Onr
xTcxOyCNrZ/gUtuwDrZftnh9WkxLWXpcyv+CBqUGUx16nBg4P5DpkHhKbpUEtAq2WvGDPx+VU3f/38V1fimSn3DPMdaxEXv4Gl9P
Wluts6k9PyK2M9pdWNtro3MVykT3myB3efPsRCvi7bArp+Kj22HNjYhTS/YRUK9MrZPlxobZ6IAz98F2RwSbv6Vs0HGdebaCFKez
3cp6rgR9d6CRzPHJFYXeyZ5w5HhU4/uIuPcltD7X1/hzVXCsNcWtHuF77edyNyULpp/ESCVwkR+VRqaBVaevUmFZBn3KhtY9UwH1
tTZN/MU6/w23TmYcvaXZ/ekTIc3FI/th0fFbU34XTZX5Xbg/C/xegGDa5ro+X1gs1C6Wf/Tmczl+k+Jy1BSxWkaFw9+7NoQln3ey
sAaVXMOr+X1FOI/gdr5pEo6OKjUw6aKYDKNVZa5SyP4zrHqXHt1554sbFhT3QMYi+V2WGqhJUDqPovnquMAKCx6UZjAbWRbLcTzY
eyMf10l6lKMyVS7UMY25cXTpcjQ8W0pERO1+KJnDTVLS4ek26n89VwmdFfd13PhnTZRQ+rlDnlEkB5CmatnuHeU3ltunQ4dPfUhM
xy3RqEUY+y0l1ATGH7oZ8JA6FehwJ1qBi7A+/e6HlFvsROY+vg0TBVgZkuJNRiZevIm8I5uoko7WFqApbO/ng+Z4/oA/wB2RpMEt
CqQ1A9Z20RF3ZbmJjsz+hdyTtOwXuLk9lWLWlk0/S4GaWTwAUgrRbsYe9LhpUdYTDsZ4u4snOuTtsRCUybMTr+5+BeWU1+NmdHvq
KlP1wAQkZOb8lbgJrUdntxq/Zj5QhRalhMI5aRwjdFx/cuYaP3QkoeFeZ/3DSe5NUsE6kLxpSB4/8nIOlnR0okJAWa0kTtqfScZW
RNo/Rq80s8hKy2kSryTiE+DSU1wjpHsA0famFssCNzYYTYoLDujOJPab7cZ1UuQWLB/Gxxil4omO5hSs7j879JX8tpQLvJXy9HKY
3k4nM5r/CiJ1ZDlai0ulhCCH4uAQSo3sbdMs3promlO5oHGvJr0QO1qKihMB8SKnxjw4q6GUvR+dpPYSie9Il7wFjlhnWchpIQ8S
1o4cJMr7Wu5sGIogXH8W59kI01hyfrKaGJKrwWF2Joobb6+l5hM2Dynq9kpPJU4HzPyTkF8AWsBlbSsR4QqBhpsSdupeTuUfB8YX
GzzOl5QBaCUOvtK7AG2fy83FIT2wIcatlHRtNN+viGYHtX7aL+BNy7EPFGhDoEfCh4jGjdUFA26VImccy/yXSouum4Tkj58Mayrz
BhG+CiTJuc8ug00BbdIkLSuGwr5td1IdjnsngRVHPSTyTVqDIZMeD30YrWKT/a+wHt0Bww1N2TxcWNNKMiVcznU0dTLPUmfFtZLX
YQBtEIn2V6AK7BN0ikb0N74NiUl88XXN9mEL+gNDuht0t7cs9XxNEj2F2xryKOhtZ7ihTGs3nlV2qsSqY9htHliDP2VPTKiDzb8r
ano4VeZrd6LHRTpjHN4fVT6NeDqiwa7D3xLqlIdnnV0xGQBehinJBIK77yz8hltA2lsrSmtUX9kkr4vWe0de2rXnqp6ae7YRlF2J
i3AmY7y0PfTSbOEql5H5fE3gW99qxe7Y5Eal6lqqVMxAB0Fuw8m1r7iUAcndLW93oVIZT5qp7I6kfM5OenElK/x1gH+/tAhQn7t/
Msi6x8ulS349PxC+OO7x3PLtffoxjSBa7eNNFNZ5lchk38Syx3gB7i/jCMiR2RtbzDul4WpGa1KAa2ve1tYJrOyfbz00Ko++qYL/
O3uZCZLMQMe90C8I35GaY8NSZiS/1+2hxqQiOl1YwoK+uPKiH4dAoXQqO+dXLWcRqf2+7t/v6kIRCwvFYPNnvqb9wIAjhTpqzEox
UJ/xRmE8E8UqV7SNPTFnj0WL8EXXbrIM5UfDNruj0FK0wHeQJ2nZv++QOE4tFBmydOrrZMz3Jsps4borW5TxF+ORJ6LH+3SvmnUn
FijK7MZodAI288scaicZ4YokAP1SesDjDZdtETP99XdBRiMHsquTI7YYN2ghb3+oci8xzVd5HFNmTYZnO13t8T4q1IA4oGNHf+tc
iTprbM7N5fLHMo5jKCj7cS2Qr6zWKkwNM8pbVAhK8q1qOOExUH6dlDWSy0u0H4XlvHxf1hpT4bMCRfmfs0vrfVZtX87RlpH9IU/u
K7rMBXsYShLOGzsY9iA1uTIc7f33CClae8m9561LRCUbdkige0BoC/UlAd6pR2NB6gvDdr+YbKrcS9jdwb+llWxWt98xcQV5fu9U
BkoQWmm76mJYqlxyMJaHTZJVwkN05HYz7igmWP115YSGDIyj6oU71aZugcS0BXf0DYOwSHAla/yoArFIxwWKZn38u4EuxWO6b6Xd
xzwhQyldbpdGKEw7kFQ1ohT2as9xRWsS+JiKHUJ2rCM0ZOIWnLSZv0th+cz1BkPRS1guiYB+xIs4RHuEGhWy40GFsLYh1OQ08HD5
40vXSkFtH3loxqRBnFNEi+U1WotALf8oqqpq3EfPFETolCiOn6WLMxx3NAW3jJiv0pY+RQqnX7IeX0scjV9hk0nb9ZLm7XyqkkE7
oIr6BCwWdCkPYEoJnQLX7YPwvfftt+yQmaOB8W2vwKExbxRksKQnj6iNohvwDVbNsJimOZ3+Wf22Lhgttkz6FcVXkbDst51L/GcA
w7cDQYAwJH14xX0+dl8OITHPy5ToQn3VW88z3HF0jlX1nXDgJXk+Yk0vf6MU78SU6jqMxRew8u7WaNDH32GTVhCFQJ2DHrY2jZfy
zXD/eqBfTghCX0AR5g9h95iVZEm1jM/LxZ28aU2hqUUtK441wAAp/GAAbBJvWjghNy12pmYeva4r+9Urg+XBMlF3p4BGiLl7xwV4
AO8WN0OTnPwFS8uhRme6H+pzxMAR/z165JZM7CUTJELTdGahtnF2kQJGw4SAPvtd0Uk/wvksNtaRNczef7j0AX9mkmsxtpsKanUx
EITvg24kS9x6fF84h7qdqkH5X2U9xGQnRpXp3TSKr6e6iF2OL7h/VDh9jNcNW5E44o+00ReZDTXacEHvQg3v+jsla5pa5FC09s/b
ImsDVUicV3F4L4CM7VsUOsQ7E35l65xNKJYNMYglolE+uXEwvnkWzO6oaKwxTlhFjgCw+9LRek9EGlwjNhXfr5U6VTzSHdeEu0sV
hYDqy5xC3Q8Ggf4i2eYOr1iS9ERmMHf1LjZklTaZS8ob2bQv9skFAEFlWyHmFaM0QGe27AfIGBLTQOhKsjyxTrrc7Gn5YVNguwuY
OQtd8Sqbr+SHxBvrBzdwefPs5A8Sivx9WmgrX2uF4IEpn/cvasowx7azvaCLV7hoRN11pmAI5SwJn4qvEejGXfPMxLd2x8RtXgNv
i6LvsbPW8rYG4WeSPvMT9JvV6D4+atEqm1fh+q4wVVCZmVmXHLI+SOxj+3t370LkUrhXOFXK2bx5y3WaO7Zj+iRuX4YM8CmdNMDh
FZJq7K78mxKP2D6FBktJD7AJ5h74Tkd6URc3upvTX0DhEWtxH6Snzk1L8YJdX/YnMzyH6qVUEDVvU6Gw5A2FFbWj1R+bgcuYNwA/
QDScchAv6F5rLV0qjXyjw33lm6WlJnCwDtQgfrjkPBpnCBk2KB500B9guruzTesJ9PocLcRc8QYSfAIdVWDRu7+BR/CnYUNR26h/
eanoL9cPQpEV44pqdqSXEFw3m4qjLQVTary66+7xChKf1AZ8S0ENDmCh8xA8IUkKYAZ7si18ddzM4FyOzKt3pnDc1WVznR0vFe23
WqYP9qLL+TMnGac8pVhktvlqJ7cy8H4MxUNziQiJFksSNMMTy339n4oOr7ANXtm8JStlif/Zajf5q1xcp5gkvDRhCDXpSv8KMXle
4qa3dS4THv0EKuhFGCsFQD1udMqHOT+TYjosm51iLgBRGuFVuMqVE7Ib2IofWT89F8AIO+E6mfHaHmJNabu/ZKvJ1u5cCtcM7tTL
ysy00Iw1qWpWqjcRNNTLYk//VKepWee/P3j5wa226PHBNH0yHo93Wx3S+1koiMwqpc6hXRv2yqjxZKiClnp0uPvcjXvAkzLkRKwh
N5+Naf/Ru5ph28dtwa/UHvfjxohB+MaQwGAIkNYEKATzYG3bJCEYThAcTCLrOlUT+GX66I5a643VGk+m0covAcT4/tHZlMJYD6vD
6zJD4vd62K6vid+nxPHzyQgdjriOQ6HOweWeyDZ5d5C7gqnj89cfFTiNm92QPOPLmaAySNHI4h1nRv1V/8mEJ3K7gHp0rsc5F0D/
TVMO8aFh9lZGe65YPrw+aS0v+sz49VYuZoSKq5tcJrAtgpE0mj/Wzn5YGS3+2kKFOhGTyw86TDMDw9XbvwfZnAHlFbmVAOU9Om6p
nr3uIyBUSBkRreiHsqjtteE6w0Lt4h8n2yarXfHXp8a2LfWcna2Xii5cCp57OZBJiKmJedl777Ip+mI+UncK+9/ALNb45+Y/PFjY
8nX7sdzVSOe0wOFq3Pd+BsF+C0HSPQLzbmh0ikB44T1iN4u0lJVhKXWySvwolLe4u6w1mtN52eS8L5haWXtehMy0kIyLqURCMtky
F1lYbDhuO3/CabStq2+SzYhkko0SJbNJj5P2nflkIoIQ3yTMy0kX07OkbbErgxPIh880vYOT6iKGTxowIGynRkElr50V6Ka0/lDm
d2MO5ecttuHp5MSwN0gPh+Spw0E25j1ZdhqP/a9iV9nYAAAA2xgIlKvvZRuNZmQoIkHyawExJjwNnmT7kz5NbD5L6df3gVbYs04p
4svCuEC6YtVHbT+zmfqg5eamTQUJznOyPWNnlMJz+vqLGgWY2AA5i58AiugbdQq6X05MKpa2zCOg1dQfzVcpSlN5BSK3YpRMkork
qGj4pWeexd9TQ7XNt7kkqRRDCn2+uwXtq6qcN6pzmbDjK8WtS9pk4ENQ3y3VnYkkZvh5LaZmGWfiIZHuDENdBf/l1A5JjAuY1txY
1RC5OQEGiK4NODpgakSlR1l5xfWWZvKzQRqjnZGksde9bykkPxc12vOOYNueQ3HUkb7NyhJ17m//hrMvwoGIyPTyl/H53XHLZW4g
VKA1l63t2xO2vuLpCU0d3wXNvbhZ0vQNDKq5s2hcFAhx3rFPBIIli/H3fVGhjO/lzl6XNIVwauR9hN36UgoXYDYJp1o83wRQXvBj
5hM6/OQxXoWJns4rtd8JosFZOYhDWoYMUvL80gz8skRQX4n9Ui9dzz4tT3E6A12j89Xdv3h0jpVIKUD7uP3S6qCh4gznkojw8woC
bUEDXzkY7TiyS6PkWM/W4QK9yCk5HAWwaULjVMAYBUnoMr73n8dFqqDaNLxl64y+p1lo0KDTlRZVCCMePso0sDAzMTg/6ZAYl1/E
eAAEZt4GeQryGnwtTphVqBICv/QyxB83+QhXX/aUPta2ryje10qy4M8kuOdY8vNK4Jgzqla4Fhb+9+KW294eCFk21A6/Uth28694
oX5wIhJ2ne+FTZbV/FdY5bO/7VyCbv4LBcd9ixJE4pMOpCc3lK3OuQAYAsjBPB+Wq98p8+nLWsJxXovrGuiLcCDcvuTIgvrAm7yx
Je+TClCUtLfA7POjHBe3SJFgRCMIJkp38ktjKpwwbtXzyaDtMqs0R2miX+rnXcA666c7CYw80oG7+6RVdGTs7DTc6cbRRKAewIgs
3QRqnvkM3TQOd0bPhKt84+aKa6Rs4RTsaemblXHyAppuHxjS2juc9E8cabJWaTXCefr9oKQnG6mHBYqLQWrheqwttxLTqjT71i2s
t/z/KDbtG3e4XtfkZuWiRblsOHPkqshbIEB3ZLB0OaPF1ViE0pXQTgUcfdaJkrk9/FGBdtfcC0KO9wCZpm+MNZ5Wflrs/CiYG3cV
Os/3nGcPFfJcZlnj6TwR4Kn2PHE7Q47otCeppxWHA7ZcDlK/qD5ltdngJOsWLSZneVjOFRkELYfB55RSTKY2Au+tZhj35T8ROz4w
30SCuRl0htry1k3pwMuLOT/MF0koBJN/Ub6bpAVZ6Dh9Q4zsVPaQHUUvay4UTV0ONM3fVh3zq/RSY6AzlL5Z/lelM85/URRA/nMi
0Fn9itieOX2AFJCiMEpKiFdPDRyhaveveolrom4J5ETarhgVUOtUINaNAyo2b1eJNMkS4DbCr0RVYq6NcBAU4+842ZpYPPe1rMqD
oxJAAACDa2zsVmj4DEQGn6Fj2gy3yaoGT9vnlZ8iZf7nom35EWbF2Nj/NI4oX+j4LOZ/Kpk5so03m2kbhj+lLiRa4SmbXFj0TkhT
vRnB/zZx+qaEsgH4MmIGIL4VuvyPVjMu0JiFFCslyo7o+nDxXev24I1+rjlE9IdJoM/tc6TI/id6qKTOK15KbZ3RcJ+cJdqy8gqG
NNgQJDTle+LnCqgolsRTbJFMgfjlPZOk0VoAOvjL/My20Wyrpfnc+89+KeCqfXvqLSTex/zp1PesKBN0Y5m//7b5bRD+o6W8gjIA
GcvwpvRzeI8fyJElEp6Oey0FzPFBrV88ZdnnpB4IK++EYnUPkW7PF4kyPTwemE9eafFKjNok2islMC7rM0hjK6t1L2/hRBA8iOyH
rBAvIFEgXL0WX9gwmyikMhUztF4OhDbglIkxStM0LLkYt+NzXfH4z6xH9jLXCg0hvjQ5QBAcYPgzVx/cAbsDkjSDhPkfFI2GtDVr
v8Kgnf5JH4KOvBwEuKCgvhF34TcRIDAcsGLK3u6sqscSrUU0YEg8mGxvlZT9aKp1JHcKUnuQZVHWyot27bfZgKdZmFgNkjTNdv/p
i5Ew14yuTAv/pJaecedHOAzn2eAY9nYMijRJA4ce9lndQxjABpeRlCtjio7M+o0s3hL9a7oktSnpyCFplBpZVO6qGtXM3xeO8q2y
h+DO1thK3V3o9bID/QeAEjq5V3T6iYNf2ZS7HvdN4MvJKVhvoTS3Fkzbm8HTWxU59H0wowgyW1dNbQFlJmpz/oh8GiebUJ9EJdlW
aqJBAFwcHbPxPTGkh7ifa5E+971/9C3L7iOTa1QzelO3djc76j8TdDj4bFkgzFDFVIP96f5QTeeuMPJa7YKUWMpchPmgw6yl3FEL
DUy8446ezU82EA6rpLGUzooH4L+aObQBCYGyeLiNmCQ44Zs630xoh8AAADYAfsAPHt6fP2uYNwVk51ikeEEXxNbUOHEAWtUYDPtX
nclaRhuVS6OOA14Hr5ynQ+GgfymuAFYww1hty9j04YL7bwVjIHiOwcwHz1roAeoyXy5UjGHEGHFbfSYKCtuCficupmw+5cYQUQxK
ImWO/BUUuo8Vmt8uL9VnJIWh58BOM0Hm2xzTNc2C7YDvskuEdUdnOmg4fFJDh/GrN7ywVtbhrG3bZDmQfHCC9uKdRDvsEKX5lS9K
oqV9LgKZ7USXNOqR9A+4kZbTETBaO9d5mMgto0lQ+w5J4ktwhgvVoZBy5vHKR4YuTdHf2HwRs0EO35CoBVus1W2ecJURuy8mpR5e
uxM7drL/gt69Gvr0sxcF5PxQflV3vr/iS5rmjiD12t0M8EeZenj7f82XkZNIZciqWBc3c5z8CFZiia25CdLkx4bxQFqEqnPlaCpl
ZF5mC6ph0yMsMWgSeGKb233EYzEGf7zdyUB8faVVOhC7UoxHs1jl9bSxXV1BOhiSi1/46GB2mg3MQ7l9dYBhkh0fip4LUxm0zK53
2Cq4kjN/FElAFvV6jRrF+g5c/Ul95e1DuVb7DLPe/13KnoGY9qiT1PauK2jyv0zJtiPg3bprT/gpD7XtG+3QeY2D8hskJkboWwsi
txreoEbNACeM4/HR22W0DpCFDc4oIVUDKc4i8yihBxhpUj8Zvbz1lMUvP5bgcHXAxVIyE7XelYG+VJ56KfMpGP1tAIMnqEJ34/ud
NaLJ3K2xUh9qfN1Rm3/uCCAxw5n6quRNvnjqpLg5fwUEhOWDcMWniSOnf8rNFkunbNnNPPWINauTmfyNoX01NaLrTqharO+0zAZe
818k411RnzEMhJVIOmtzWlAUPtxxlq23Ky9dkxdqaHhzw6hrHWjT4q4ZC0dkIriaKp7WDyo603h7jHnf6YL+cWP4utvSomLPn1UT
NK/+AZ9FVZtj/Anru9tMoI+HMOuTjP10zf0oSZgghbPyO29WdbUKzebBxjlKGUoeoy3mgoF4ABwqOKlQnGU0uXrclE1VlnPz3jb3
TbTzDD7vHAJiIgVf77j+9BFE39EfxmO5k9erULDmLVRIFv8p6tYT9BFtoiYoBdXEdi7reAWBRcn3LdvBJGnJMDqLeP83ZuPWSxOE
QKWKn5LzzAVSeA/Z7exRr9M+7jOhu/PTI5l6voCE/O+2MkmE4vWLkWWS3KT4iFxXtBew1YcVjBR19+/4GIYNjpHoI3UEm/dlsNck
EDy0z/+SDJQGNFa8FiVDwXpyMp5ZxsNS1hpBMmQbUrR/CSX2xR/4wgPV8FyjpzB5luAEWDLayUNewHw9wAAAApIsFSXW8XyTy676
1QEixbAg3+frpnhVWGUri95FH//FYI5VJjbfholpU/vV3dU67xXARRy2Th7uSYzPQvYlOc7JcFudbkOLpibfDzIPZUtOZMM+E+Jr
uQHQDYNsGEHJ1NhteBZ9KdlPHXJCGDRzOfxsmNtlYNLVEeVRunZNGB1/ScDjHCdMeq9CTyaUkWM3vvnvNHp2EzIuC7AFruvUUsKU
8cR1EUhB7qwAh5kRyJgchWlrjxB84hNIfkSXy1k1a33A58Tn6KlYve3yeyGxF11waG8XI5puRWBj0yUTYasJorOKFiqW7T8TYZyV
rwjvMnJezDZH+faHK92hcMxKLpwjZt+JY3xzcRSFewR2c1ifyqSlwOMWcQsUFwoTNhXsAvKCIcyth+I33JiakETDUhloJgvoAcVz
ZFDRdWGMJlmSMQKiVxZPDon0EdcY00X8I6xL3SrQL6Lx8E+HskUCIS0cJU/1WOmGxdZZOXKrokHikOshM8zsVNgjtedfyaCavoKA
XdmGWVzxrTFub5KkvQnBBOcOjLxIDE34WKLisdSGoC+wzpz48Q4hjsB5LQm392gQZyVnQWbl7cp7J5YitnVXahIPqb7vxlhoqvfM
G8EkaJVVL92c6OE/aPh+VAeArNgAreDt663OZIqDyBhGk2acIDqN/IusNZQcDMTAAAAG+I2RmQAABmP22GC73lhAMUPlc6hy641B
GcYKbIq412YmFnxi3DnA+yjOweHIUt0SexfVSt6QgHrzb1RGIvdHQIk/CYZUCtsLfKABgyfHxS1YDkgAACencumUIxCSbEgW54gY
dytUqB9SgAAAAAAAAgQ1i7DvPCgiwjuDjY1CAAA=
""".replace("\n", "")

_logo_source_cache = None

def _get_logo_source_image():
    """Decodes the embedded official logo once, returns a PIL RGBA Image."""
    global _logo_source_cache
    if _logo_source_cache is None:
        import io
        raw = base64.b64decode(_LOGO_SOURCE_WEBP_B64)
        _logo_source_cache = Image.open(io.BytesIO(raw)).convert("RGBA")
    return _logo_source_cache


def _make_mythic_icon_png(size=192):
    """Returns the Mythic AI logo as PNG bytes, resized to size×size.
    Reverted to the original programmatically-drawn teal badge + white 'M'
    mark (see _make_mythic_icon_png_fallback) instead of the embedded
    "3D M mark" artwork."""
    return _make_mythic_icon_png_fallback(size)


def _make_mythic_icon_png_fallback(size=192):
    """Generate a real PNG icon for Mythic AI programmatically using only stdlib.
    Draws the teal rounded-rect background + white M-shape — no Pillow needed.
    Only used as a last-resort fallback if Pillow / the embedded logo aren't
    available — see _make_mythic_icon_png above for the normal path."""
    import struct, zlib

    W = H = size
    img = bytearray(W * H * 4)

    def set_pixel(x, y, r, g, b, a=255):
        if 0 <= x < W and 0 <= y < H:
            i = (y * W + x) * 4
            img[i], img[i+1], img[i+2], img[i+3] = r, g, b, a

    def fill_rect(x0, y0, x1, y1, r, g, b, a=255):
        for y in range(max(0,y0), min(H,y1)):
            for x in range(max(0,x0), min(W,x1)):
                set_pixel(x, y, r, g, b, a)

    def circle_aa(cx, cy, radius, r, g, b):
        for y in range(cy-radius-1, cy+radius+2):
            for x in range(cx-radius-1, cx+radius+2):
                d = ((x-cx)**2 + (y-cy)**2)**0.5
                alpha = max(0, min(255, int((radius+0.5-d)*255)))
                if alpha > 0 and 0 <= x < W and 0 <= y < H:
                    i = (y*W+x)*4
                    existing_a = img[i+3]
                    blend = alpha / 255
                    img[i]   = int(img[i]   * (1-blend) + r * blend)
                    img[i+1] = int(img[i+1] * (1-blend) + g * blend)
                    img[i+2] = int(img[i+2] * (1-blend) + b * blend)
                    img[i+3] = min(255, existing_a + alpha)

    cr = size // 4
    cx, cy = W // 2, H // 2
    fill_rect(cr, 0, W-cr, H, 16, 163, 127)
    fill_rect(0, cr, W, H-cr, 16, 163, 127)
    circle_aa(cr,   cr,   cr, 16, 163, 127)
    circle_aa(W-cr, cr,   cr, 16, 163, 127)
    circle_aa(cr,   H-cr, cr, 16, 163, 127)
    circle_aa(W-cr, H-cr, cr, 16, 163, 127)

    s = size / 40
    pts = [
        (int(10*s), int(28*s)),
        (int(10*s), int(12*s)),
        (int(20*s), int(22*s)),
        (int(30*s), int(12*s)),
        (int(30*s), int(28*s)),
    ]
    lw = max(2, size // 14)

    def draw_line(x0, y0, x1, y1):
        dx, dy = x1-x0, y1-y0
        steps = max(abs(dx), abs(dy), 1)
        for i in range(steps+1):
            x = int(x0 + dx*i/steps)
            y = int(y0 + dy*i/steps)
            for ox in range(-lw//2, lw//2+1):
                for oy in range(-lw//2, lw//2+1):
                    set_pixel(x+ox, y+oy, 255, 255, 255)

    for i in range(len(pts)-1):
        draw_line(pts[i][0], pts[i][1], pts[i+1][0], pts[i+1][1])

    def png_chunk(name, data):
        crc = zlib.crc32(name + data) & 0xffffffff
        return struct.pack('>I', len(data)) + name + data + struct.pack('>I', crc)

    raw_rows = b''
    for y in range(H):
        raw_rows += b'\x00'
        raw_rows += bytes(img[y*W*4:(y+1)*W*4])

    ihdr = struct.pack('>II', W, H) + bytes([8, 6, 0, 0, 0])
    compressed = zlib.compress(raw_rows, 9)

    png  = b'\x89PNG\r\n\x1a\n'
    png += png_chunk(b'IHDR', ihdr)
    png += png_chunk(b'IDAT', compressed)
    png += png_chunk(b'IEND', b'')
    return png



# ── PWA install screenshots (base64 PNG) ─────────────────────────────────
# Chrome's "richer" install UI requires at least one screenshot with
# form_factor 'wide' (desktop) and at least one without 'wide' (mobile) —
# see the 'screenshots' array in /manifest.json below. Static app-preview
# mockups, embedded here so no external asset files are needed.
_SCREENSHOT_WIDE_B64 = """
iVBORw0KGgoAAAANSUhEUgAABQAAAAMgCAIAAADz+lisAACy10lEQVR4nOzdeVwTd/4/8E8ChIQbuQTlkkMQxfuoIlIUiuCBVvFE
cYvYbiuiX60/qa3Xlq1HRdF6W7dd6wXaiuCBKBYRkIoiVQ6RUznkkCNAOELy+2O62SyEECAQIK/no4/dyXw+8/m85zMD5s18Zoam
q6tLAIAQa2vrvLw8WUcBAAAAAAA9hS7rAAAAAAAAAAB6AxJgAAAAAAAAkAtIgAEAAAAAAEAuIAEGAAAAAAAAuYAEGAAAAAAAAOQC
EmAAAAAAAACQC0iAAQAAAAAAQC4gAQYAAAAAAAC5gAQYAAAAAAAA5AISYAAAAAAAAJALSIABAAAAAABALiABBgAAAAAAALmABBgA
AAAAAADkgqKsAwAAAAAAgN5Wv3etrEPoK1S2nu65xs3MzHqu8YEtLy+vJ5ql6erq9kS7AP2OtbV1D/2YAQAAAPQdSH1FknoaTKW+
VVVV0m1WfmhpaZEeSIOlkwBzgz/vfiMA3aG48YdutoAEGAAAAAY8ZL9iSDEHNjMzQ+orFVpaWtL9it7dBBipL/Qp3UmDkQADAADA
wIbst0NSyYGR/UqXdHPgbiXAyH6hD+pyDtyjCbDIf2969IYTAAAAAGHIfiXUzW9oyH57ghRz4K4/BAvZL/RN3ODPuz8dWorE/GND
FSENBgAAAADoHV28AozsF/q4LuTAUr8C3Km/syINBgAAgJ6Dy7+d0uUvZrj823OkdRG4K+8BRvYLfZ/Mz9LO/jODf5YAAAAAAHpa
VxJgABCva9ms5FuZm5unCvn1119bVVBUVIyJiRGuo66u3oWQ2uPs7Cxo2dLSUio1O+XQoUPCezds2LC2dYRH6euvv5ZW1wAAAADQ
f3U6AZb5hTUACcnqXO3OtdyubWthYTFhwgThNa6urjo6Ol0OgyLIHj/77LNuNiVdGhoa06dPF14zZ84cWQUDAAAAAP0IrgADDARL
ly4V83GA+eijj5SUlITXuLu702g0WcUDAAAAAP0FEmAAaer+rbxda8HZ2VlPT49atrGxGTNmTDfDkJb79+/b/8fr16+l0qbgem9T
UxO1YGRkNHbsWKk0DgAAAH1K5qLNlSt3epnbt1oG6BokwAD924sXLwghioqKixYtotYsW7aMWvjzzz9bVd60aRM1qzkmJkZR8b9v
QVNWVk5ISKCKfH19T58+nZqaKij97LPPBNOhDQwM2sYwZcqUn376KSkpKT4+/siRIxYWFoIiMfcA29jY7NixIzw8/PHjx/Hx8Rcv
XvTx8VFTUxO/v0ZGRoL0/vz5842NjdQyZkEDAAD0jocen1au3Hlq2sIut3D5w+WVK3fun+guxagAJIQEGEBqpPUk5061c/PmTepp
+4sWLVJQUNDQ0HB3dyeEFBQUPHr0qFXlS5cu8Xg8QoiOjo6Tk5NgvYODg6qqKiGEx+PduHGjU9HOnz//5MmTY8eOZTKZampqM2bM
OHPmTId57CeffHLp0qWPP/7YzMyMxWKpqanZ2dlt2rRpzZo14jecM2eOYLZzZGSkYB9dXV1bzYsGAAAA6F98fb+k/kNfPQcJMED/
1tTU9NtvvxFC9PT0nJ2dFy5cqKysTAi5fPkyn89vVbmoqOjBgwfUsuCKMSHko48+ohYSEhLevXu3du1ae/v/Ti46fvy4YBrzu3fv
WrXp5eW1efPmqVOn7tu3j1qjo6Pj6ekpJmYXF5cNGzbQ6XRCSGpq6ooVKyZPnjxv3ry9e/e+f/9e/P56eHhQC2/evMnKyrp//z71
UUNDw9HRUfy2AAAAIHXUtOQdY2ddneldsGTbi4UbV1uOp4o+s5nyx7z1xcu+SvHccHDSnCEqmoSQhDl/dx1iTQjxHT6pcuXO7MVf
kv9cVa5cubN0+dd/zFv/he1UGe6RrAjnhz2dKw7UviTRdxPgY1M9qR+D6Nn/vRr2/+ydqJW2Wvoyieq3WasqV+5MXRDQtc0F8Zuo
aUkzLJBvly9fpq7rLl++3MvLixDS0NBAZcVtXbhwgVqYMmWKkZERIURZWVmQOrZ9o1KHfvnll7t379bW1p4/f76mpoZaKfK9RAJr
1/71Q11dXf3ZZ5/9+eefHA4nLy/vl19++eWXX8RsaGdnZ25uTi3fu3ePEPL777+3tLRQawS5MQAAAPQyD2Mb/4TwTY8jhqhofj/Z
Q5+pNnqQYdAEN05Ls83V792jzv1ZVeI7fCIh5IOIY1GFrwghZzKTtM/vtAjdRwiZHnlC+/zOQed3TQo/yuPz9ox3nWtiK+Nd6l1t
M8OeyxUHal8S6u0E+ODkOZUrd1pp6Eq+yXidIXOM5esHAKBTCgsL4+LiCCHjx48fOnQoIeTmzZtsNltk5aSkJOphVHQ6feHChYSQ
6dOnq6ioEEKqqqpiYmI62/vjx48Fy9XV1dSClpZWe/VVVVVtbGyo5YcPH7YXp0jCN/pSCXB1dXVycjK1xtHRUbqvOwYAAAAJ/Zz1
tLC+OrroNSFEgUY3VdfWUGISQoxVNb3M7c3VB13NfbHrWbT4RviEn1dbeb84mxDiPtSmF8IGOdR3rwAL2z7GmY53nAC07+LFi8If
L126JEllT09POp3u6upKfYyMjGxubu5s18KTlgUXY8XQ1NQULJeWlkreEZ1Od3Nzo5bLysoEj+miMmFCCIPBEOwLAAAA9KbShlpC
SDPvr28CCjRafGn+1bw/NRnMfRPdI1x8sr2+3DF2lshtaYQWYOeQOPfzt0sDK1fu/NRmCiHESEWj14IHuaLYcRVZe99YP1xTz8vc
/lLO87alH5uNWjt8op2WAY1Gy2W/v5yTevpVUmMLlxCywNTux+mLCSEf3jyV8r5olpFVqPMKQsik8KNZNeUfm4064/AxIcT51qln
FUVtW55nMsJ3+ER7bUMlOv1ZRdE/nt9PLC2gimiE5jt80irLcZbqOuWNdWdf/XH45V9P4jk1beHi/zyZncNtzmZXhOb+eSQtnk/4
Me5+YwYZUUXPPQMIIRlVpR9EHJPmYIG8io+PLygoMDExIYQ8e/YsIyNDTOUbN24EBASoq6vr6+u7uLgI5j+3N2tavLZ3GosnuEpM
CNHX78S9DFOnTtXR0aGW9fT0nj8X8Qthzpw5V69e7VQ8AAAA0BNa+DzfuKv/lxQ5QsvAdYhVgJ1DgJ1D8Iu4muaGVl8dPhpqvWPs
rDpu06zbZzKryg5Mcv+b9UQFObv6debMvlZzg8+c2Ye+ekI/uAIc/CKOELLN/kMGXaFV0fYxzmccPjZV054ddW7sb4fruE17xruG
Oq+gLhf/XpLL4/MJIQ4GZoL/JYQ4DP7vx8omzvP3xW073Wrv9JOj11R9053P7tpePfj/ntxeaDpSUDpUVdNWU2/e3Z+Opscbq2rt
HOviZPjXHY9+j65pn9+pfX6nwYU9H98/b6qmvWucy99tpxBCPrx5am/qA6ra6N8OaZ/fiewXpIXP51+5coVaFn/5lxDS0NAguNf3
q6++ouY/p6enZ2ZmtmqTWlBQaP2j1x11dXWC/Hz69OmST1qW5EVH48aNMzQ07HpwAAAAICVjdYxCpswzU9N+VlF0r+g1IaS2uZHT
0kwIKWuoI4SYqGkp0uiEEGW6AiGkhc9rbmmx0dKbZzJCpoHLjHBm2NNZ4kDtSxK9cQV4jfWEg5P+55tr0rwvqIUWPk/3l93iN/93
9tM11hOGqQ9abTVeeL2hikaAnQMh5GTG4xeVJYSQH9ITJuuZTDcwdza0jC7Ket9Yn1pZPGaQkcNgs6Pp8dMHm9U0N6grKU83MD/3
6sn0wWaEkAfFObw2168Gs9Q3j3IkhFzPT/tXVjIh5EVlyZd/3BRU4PJ5O55F1zY3huW92DJqBiFkzCCjB8U5wo008VoSSvOTyt7M
NLJcZD7qh/QEiQcM+iuVrael8iYkla2nu7DVzz///PPPP0tY+eLFiytXrqTT6YKbddte/n3//j11xXXs2LHq6uqdullXvNOnT3//
/feEEE1NzWPHju3du/f169f6+vozZswghPz73/9uu4mKioqzszO1fPv27S+//J8/JQ4bNoyKn0ajeXh4nDlzRlqhAgAAQNekvi+J
Lcn9bsLskdoGLXx+Qml+0PMYao70D+kJo7QHzxg8rGzFN/eKXnvF/HI8I3HpsNFRbp8klBbEFGcLJlTKm97MDwdqXx3qB1OguTze
t8/vn3VYtHmk46Xc/056HKdjpECjE0J2jJ3V6o4CKw2d6KIsQsiD4pwxg4w+0DfRZDBHDzKMeJNhraE7zcDUUEXDQl2HEPKgOLtt
j+N1h1B/jvqj/K3IkMo4tbXNjYQQaq41IYSp8NdIfqBvumWU4+hBhloMluC+ZVM17W4NAYC0FRYWPnz4kEo4CSFNTU2RkZGt6jx4
8ODjjz8mhEycOJF63W5aWtrSpUu73/vdu3cPHz68fv16Op0+evRowYOpCSGnT4tO/mfOnMlkMqllwauPBHJycgQzwJEAAwAA9Kjp
kSeEPw4POyBYruM2aZ/fKfgYlvdnWN6fbVvIqCp1unlSeE3gk9uBT24LPvo9uiayfeFlgK7pjQT43Ksn5149oZYPTp6zxmoCdReu
5C38mvcywM5hlPbgFRZjBStp5K/0ckNi+M+vn4rcMKY4O8DOQUOJuW74ZAUaPe5dXimndu3wSb7WE6kK90UlwAJ8IvrmRi6fJ7KC
kYrGVeeVLEWl3/JfbnwcUdXEuTrT29nQQt7uYZBn3b8I3LXLv11w4cIFQQJ8//59wRuMBA4cOMDlcp2cnPT09Kh39krR2bNnHz16
tHTp0vHjxxsYGLS0tOTn59++fTssLExkfcH85+bmZuqR163cv3/fx8eHEGJhYWFjYyP+LmgAAAAAkE/94AowIYRP+HtS7l35cIWO
sopgZXJFYQufp0CjT9E3aS8Bflz2hsNtZikqfWY7hRAS9y7vHYe9dvgkv+GTCCGvayre1lW33Sq5/K+WJ+oOPUket63QntGDDFmK
SoSQq3kvqpo4dBptuOb/vPCp7XRrgC7Izc21t+94atDx48ePHz/eXmlSUlJDQwN1WVXk46/q6uq+/fbbb7/9tm3R/fv3RQYwb948
CWsSQjIyMnbu3Nlu9P9r3bp14iscPHjw4MGDwmskHCUAAAAAkB/94CFYlLuFWYKHMFOK62uoZy8vMhu1yGyUiiLDTE17sbn9Lde/
Ua8dI4Q0tnDjS/MJIVoMVnlDXWZV2aN3+XzCV1NSJoTEtHP5t4TDPvjiISHE09RuleU4dSXlEVoGh6bM7TDI9KpS6sYG1yFW6krK
m0c6DlHRFK5QUFdFLYwehOf0DFjduYTba5d/mUzmmjVrqOw3MzMzIQH3qAMAAADAwNdvEmBCyK6U1u/O3pNyzzfu6h/lbw9OnpOz
+MsrzitmDDbf8exuTXODoI7g2VSPSvP5hP++sf5l5TtqTXsJMCEk6HnM3x6GJpQW/GP8R5kfb943cbbIlzC1kldbueZhaEZV6ZJh
o/+Y94Umg/l7yf88Geta3otLOc8rGut/cvSqXLnz4OSOn2oL/VHX8they35TU1OTkpL8/f0JIVwud+/evZ19lREAAAAAQH9E09XV
7biWEG7w5z0UCoDUKW78QfLK1tbWeXl5Uuy9UzcD91r2SwhJTU0lhNTW1mZlZR0/fjwxMbHXugYAAACZkMqLKuRHl7+YmZmZVVVV
STUW+IuWlpZUvqv3j3uAAfoj6ldnh//e9GbqS8GdsQAAAPJGWi9rlAfd+W6Wl5eHHLgnSCv7JUiAAXqamDS491NfAAAAAAB5hgQY
oDcg1wUAAADZwkVgSXT/OxsuAkudFC//kv71ECwAAAAAAOgy/EVePGmNT15enpaWllSaAulmvwQJMAAAAACA/EAO3B7pjgyVAyMN
7g5qAKWb/ZIuPAWa4EHQ0E906hHQpAeeAg0AAADQZ2E6tECP/lHAzMys5xof2HromznuAQYAAAAAkDu4FNw7cH2lr+nKFOjOXlgD
6H04SwEAAAAAoJUu3gOM7AL6MpyfAAAAAADQVtcfgoUcA/omnJkAAAAAACBSt54CjUwD+hqckwAAAAAA0J7uPgSLyjfwXGiQOaS+
AAAAAAAgnnSeAo3cAwAAAAAAAPq4bk2BBgAAAAAAAOgvkAADAAAAAACAXEACDAAAAAAAAHIBCTAAAAAAAADIBSTAAAAAAAAAIBeQ
AAMAAAAAAIBckM5rkAAGBjMzM1mHAAAAAAAAPQUJMMB/5eXlyToEAAAAAADoKZgCDQAAAAAAAHIBCTAAAAAAAADIBSTAAAAAAAAA
IBeQAAMAAAAAAIBcQAIMAAAAAAAAcgEJMAAAAAAAAMgFJMAAAAAAAAAgF5AAAwAAAAAAgFxAAgwAAAAAAAByAQkwAAAAAAAAyAUk
wAAAAAAAACAXkAADAAAAAACAXEACDAAAAAAAAHJBsWub6ejor1v3JSEkPf35r7/+W7B+27Z9NBqdEMLjtVRXVz17lpiYGCNYf+jQ
jvr6OuF2hNdPnjxj5sy5FRWlv/xyYv367VQ7lGfPEm7fvubt/fnQoWYPHtyMj79PCJkxw23atFn5+a9/+eWkqemwDz/00NMzpNFo
9fW1T57EJSY+6NquAQAAAAAAwIDUxQTY3n4iIYQQvrW1HZPJamjgCJceOrSDz+cvW7bO2dmjtrb6xYunHTb4wQfOH37oXlZWcuHC
ybo6tqAd4YQ5IuKyr+//OTi4ZmT8qaCg8MEHHzY3N0VGXlFUVFy0aE1dHfv48X/yeDwrqxHq6ppd2y8AAAAAAAAYqLoyBZpGo40c
Oa6lpSU5OV5BQXHEiDFt63A49fn5rwkhgwcbd9jgtGmzPvzQvbS0+JdfTgiy37bevy/7/ffbioqKHh5eHh5edLrC/fuRVVXvtbQG
KSsz+Xx+U1NjfX3t8+dJcXF3u7BfAAAAAAAAMIB15Qqwubm1urrmq1cvkpPjx4+fNmrUhKdPE1rVYbFUTE0tCSElJW87bHDixOnv
35f/8stxDqdeeH1AwC5qISLicmrqH4SQpKRYG5tRxsbmhJCCguzk5HhCSFVVRW1tjY6OfkDAruLiN7m5r548iWt1URoAAAAAAADk
XFcS4FGjJhBCXr58Vl7+rrS0eMgQUx0dvYqKMkGFgIBdPB6vpqbywYNbL14kd9ggn8/T0hpkZmaVnv5ceH3be4b5fN6tW2G+vv9H
CImMDCWETwjhcrk///zD1KnOFhY2xsbmxsbmNjb2Z85834VdAwAAAAAAgIGq0wmwsjJz+PCRhJAFC7wXLPCmVo4aNeHBg1uCOm0T
V/EiIq64uy/y9FxBo9HS0lLEV66qek8tVFe/F1pZcfNmKCHEyMhk9eov9PUN296ZDAAAAAAAAPKs0/cA29qOVlRUiou7GxS0OSho
8759/6+lhTty5HgajdblILKz00ND/9XSwps/f7md3djObq6lpbNw4SpjY3MGQ7m5uamlpaWmpgrZLwAAAAAAAAjr9BVg6vnP1AOu
CCFcLrewMN/ExMLMzDI3N6vLceTkZISG/rh48d/mzVsueAGS4B7g16/Tr1w52962zc2NTCZr4cJVTKYKl9tcWFhw796NLkcCAAAA
AAAAAxJNV1dX1jEA9AnW1tZ5eXmyjgIAAAAAAHpKV16DBAAAAAAAANDvIAEGAAAAAAAAuYAEGAAAAAAAAOQCEmAAAAAAAACQC0iA
AQAAAAAAQC4gAQYAAAAAAAC5gAQYAAAAAAAA5AISYAAAAAAAAJALSIABAAAAAABALiABBgAAAAAAALmABBgAAAAAAADkAhJgAAAA
AAAAkAtIgAEAAAAAAEAuIAEGAAAAAAAAuYAEGAAAAAAAAOQCEmAAAAAAAACQC4qyDqCP+u2333788cfw8HBZBwIAAAAggpaWlqxD
AAAQraqqStYhtGsAXgG+fPny4sWLZR0FAAAAgPRp/YesAwEAaFdf/k3V2wnw119/nZGRkZGR8fLlS2m1efXq1QULFgg+LlmyJDQ0
VFqNAwAAAPQFffbbJABAe/rgLy7pJMDc4M+p/zqsuWfPHhsbm08//VQq/QIAAADIg772DRIAQHJ96jeYFBJg4bxXkhy4QydOnMjI
yEhPT793795nn31Go9Go9Wpqal999dW9e/eSkpJ2797NYrEIIQcPHrSzs/vnP/+ZkZFx+fJl8p8p0CtWrPj1118FbQ4dOjQtLW3I
kCGEEBUVla+++iomJuaPP/44ceKEkZGRmGCYTOaxY8dOnjzJYrFEbrh06dIbN24I6puamr58+VJ8mwAAAAAS6oPXTwAAOqvv/Crr
bgLcNuPtfg786aef2tjYjBw5cv369QsWLPDw8KDW79u3b8SIEevWrZs1a9aLFy8mT55MCNm0adPLly+3bdtmY2OzZMkSQSM3btyw
sLCwtbWlPi5YsCApKamwsJAQsnfvXlNTU29v7w8//DA3N/fo0aN0uuhx0NXVPX/+fGlp6d///ncOhyNyw/DwcENDw3HjxlGbLF26
ND4+vqioqJuDAAAAANBHvi8CAEhFX/id1ncfgtXS0pKWlvbvf//b2dmZEGJoaOjs7Lxt27bXr1/X1NRcuXLlwYMHYjavqam5d+/e
woULCSE0Gm3BggVXr14lhAwePNjFxWX79u1v376tra3dv3+/iYmJpaVl2xYsLS0vXbp0586dnTt3trS0tLdhfX399evXly5dSghh
MBgLFiy4cuVKTwwIAAAAAAAAdEdffA2Sm5ubn5+fubk5Ncn56dOnhJChQ4e2tLS8efNG8nauXbu2f//+ffv2TZgwQV1d/e7du1Q7
hJDY2FjhmkOGDHn16lWrzT/++OPKyspffvmF+ihmw4sXL169ejUoKMjR0bGlpSUmJqazuwwAAADQSl+4VAIAIF1aWlqyfUlSdxNg
xY0/tJrzrLjxh+40qKOj8/3332/YsCEhIaG+vn7VqlXu7u6EkLdv3yooKBgbGxcUFLTahMfjiWzq0aNHjY2NTk5OLi4uN2/ebGho
IIQUFhbyeLypU6d2OO4HDhxwcHA4d+7c2rVra2pqxGz4+vXr1NTUBQsWuLi4XLt2jcvldmnXAQAAAP6C7BcABirZ5sBSmAItnPF2
M/slhDCZTDqdzmazm5ubR48e7ePjQ60vLi5+8OBBUFCQhYWFurq6l5eXk5MTVVRWVjZ8+HAFBYVWTfF4vOvXr69atcrV1ZWa/0y1
c//+/e+++87MzIzJZI4aNero0aMMBqNtJC0tLVu2bHn16tXPP/+so6MjfsMLFy74+vqOHTsWb2ACAAAAAADom6RzD7Dixh+o/zqs
uXjx4oyMjBMnTigoKFAvBNbW1hauUFhYuH///oMHDyYnJ2/bti0qKkpQRKWjZ86ciY6OtrOze/z4MbX+zJkzjo6Oqamp1FOghV29
enXixIlv375NTU0VrNy6dWtOTs7p06cTEhK2b9/+66+/NjU1iYyWx+N9/fXXiYmJ58+fHzx4sJgN7969y+fzHz9+3PYCNQAAAECn
4PIvAAxsMvwtR9PV1ZVV3wMJg8F4+PDh7t27IyMjZR0LdJG1tXVeXp6sowAAAEACDAADn6xmQffdp0D3IzQabcWKFQ0NDXfu3JF1
LAAAAAAAACBaX3wKdL+TkpLCZrMDAwPx+CsAAADoJlz+BQB5IKtHYeEKsBSMHj3awcGh1RuSAAAAAEC84ODgdevWdWqToKCgjRs3
9lA83dfHwxP29ddfb926tde660cj0zvEDEgXfi5AckiAAQAAAPo9TU3NL7744uLFi/fu3bt69ep33303duxYWQcFvURHR8ff3//K
lStRUVFnz551cXHp/Rj27t3r7+/fR9oMDAwMCQkRfBw/fvydO3fWr19Po9F6LYY+1QUIwxRoAAAAgP5NT0/v2LFjRUVF//jHP7Ky
stTV1S0tLVevXv38+XMejyfr6KDHeXp6ZmdnX7x4sa6ububMmV999VV1dXVSUpKs4/ofgYGBMul3xowZ33zzzU8//fTzzz/LJID2
yGpAAAkwAAAAQP+2YcMGQsiWLVuoFzRWVFRUVFQIXhi5efPmefPmEUJqamrS0tJCQkLevn1LCAkODi4uLmYwGBMnTqTT6dHR0UeO
HGmbMHt4eKxbt27hwoWCZ53s2rVLQUFh+/btSkpKa9eudXV1VVdXz87OPnHixNOnT6k6P/74Y1RU1KVLl6iPO3fu5HA4e/fubRu8
pqbm9u3b28ZAo9G8vLw8PT319fWLioouX74cERHRatvg4OCSkhIVFZWxY8fS6fSIiIhTp06JzPmDg4PLy8tZLNaoUaOUlJQiIyND
Q0MDAgImTJhQW1t7/vz5a9euSdip5P22N/KSD76YFoSdPXtWsHzjxo3Vq1fb2dm1TYAZDEZAQICzs3N9fX1CQgKLxRK80bO9jsQc
feGWAwMDP/jgA0LIokWLCCErV67cuHFjSUkJi8WaMmVKZmbmhg0bgoKCysrKgoODJRzAtm0WFBSIOeVEmjt37saNGw8dOhQeHt6F
M7ltDMXFxR0GMGXKlF27drm7u7e0tAwdOvTChQvXr1///vvvCSFr1661tbXdtGkTIYQaEBaL1XY3Sfs/F9B9mAINAAAA0I+xWKxp
06b99ttvgmSmlQMHDjg6Ojo6Onp7e5eWlgYFBSkoKFBF7u7uycnJy5Yt+/LLL93c3GbPnt1283v37jEYjKlTp1IfNTQ0HBwcbt68
SQjx9fV1dnb+6quvPD09ExIS9u/fb2ho2Nn424th9erVHh4eu3fv9vDw+P777/38/GbOnNl2cw8Pj8zMzOXLl2/fvt3Dw2Px4sXt
deTi4hIbG7t06dI9e/Z4eXmdPn06Ojp6/vz5ISEh/v7+pqamkncqYb9iRl7MjkveQltMJtPNzU1DQyM+Pr5t6bp168aOHbt+/Xof
H5/GxkYHB4cOOxJz9IUFBQUlJCSEhYVRjVApnLu7+9OnTxcsWED9gaazAyiyzU6dcitWrAgICNi9e3d4eLj4fWmv2bYxSBLA8+fP
GQyGjY0NIWTs2LHV1dWC+xHGjh2bkpLS4W4SyU4P6BokwAAAAAD9mJGRkYKCwps3bzqsWVlZefjwYRMTEzMzM2pNXFzcrVu36uvr
09PTExIS7O3t227V0NAQHR3t4eFBfXRxcampqUlMTGQwGIsWLTp9+vTLly/ZbPa5c+fy8vLE5J/tERkDg8FYvnx5SEhIenp6Q0ND
SkrKtWvXBDEIy8jIuHDhQm1tbUpKyi+//OLl5SWmo6ioKA6Hk5CQkJeXl5ycfP/+fQ6H8+DBg9LS0pEjR0reaaf6JaJGvr0d71QL
wszNzWNjY6OiojZv3rx///7MzMxWFZSVlefPn3/q1KmsrKyampoffvihrKysw47aO/piQhV48uRJeHg4h8MRWdqpAaR06pSzs7Nb
t27dt99+++DBA2pN989kCWtyOJxXr15RSe/YsWOvXbs2ePBgHR0dJpNpY2Pz7NmzDveUdPL0gE7BFGgAAACAfox6rg+fz2+vgrm5
uZ+fn52dnaamJlXZwMAgOzubECI8n5bNZuvp6YlsISIi4sSJEzo6OhUVFe7u7rdv3+bxeEZGRkpKSmlpaYJqaWlp7aVnYoiMwdTU
lMlkHjhwQLCPNBqtqKio7eYZGRnCy3p6eqqqqnV1deI7qq2tbfVRXV1d8k4l7FfMyLe3462Ib0FYbm6uo6Ojmpqas7Pztm3b6uvr
Hz16JFzByMiIwWAIEuOWlpasrCxJOhJ59EWOSduQxJRKfuCEd0HyU+7t27d0Ot3b2/vp06eCd+1080yWvOazZ8/Gjh17/vz50aNH
h4aGjhs3buzYsVVVVVwuNz09Xcw+CscvWBbzswldgAQYAAAAoB8rKipqaWkxMTERWUqj0fbt25eYmOjn51deXs7j8aKjowXTaMWk
zcIyMjJycnLc3NySkpKsrKx27txJ/pN4t+qrvQbFPH1X5CZ0Op0Q4uPjk5eXJz42CXehbc22G0reqST9ih95qbTQVm1tbXh4+JQp
U+bPn98qARbTqfiORB59STQ3N4splfzACcfZdk177VRXV+/YsePQoUOHDx8OCAiorKwk3T6TJa+ZkpKycOFCMzMzFRWVzMxMKh+u
qqp68eKF4A5k8bowPiAhTIEGAAAA6Mfq6+vj4+M9PT0ZDEbbUl1dXQMDg0uXLpWUlHC5XEtLS0XFrlz/iIiIcHd39/DwSE1NpaZb
FxYWNjc329raCurY2trm5+dTy2w2W11dXVA0dOjQTnWXn5/f2Ng4efLkDmsKB2BjY1NeXi7+KqJUOpWk3+6PfJdbEFmtqKioubmZ
ujGVEKKgoGBlZSVhR22PfltcLrdTLxmS5MC1alP8KddWZWWlv79/S0vL4cOHBw0a1N6+iG9WOAbJA3j+/LmSktLy5cv//PNPHo/3
7Nkz6iJwqxuARe4m9DQkwAAAAAD92+HDh+l0+v79+0eMGKGkpDRo0KDJkycHBwfT6fTKyko2m+3h4cFkMocNG7Z169audREVFWVg
YDB37lzBA5CamprCwsJ8fX1tbW3V1NRWr15tZmYWGhpKlaamprq4uJiYmKiqqnp7ewtyLQk1NDRcuHDBx8dn1qxZqqqqgwcP9vT0
XLVqVduaNjY2y5cvV1VVHTNmzIoVK65cudK1HexUp5L02/2Rl7yF3bt329vbq6ioaGtre3l5TZo06datW63qNDY2Xr9+3c/Pz8rK
SkND4/PPPxfMqu2wo7ZHv613795ZWVmxWCwJ906SA9eqTfGnnEjV1dUBAQFNTU0hISE6Ojoi90V8s8IxSB4AdRuwq6sr9Yzoly9f
6uvrjxgxQuQNwJ0dOugmTIEGAAAA6N9KS0t9fX29vb2/+eYbPT296urqrKysn3/+mcfj8Xi8HTt2+Pv7L1mypKKiIiwszNjYuAtd
1NbW/v777w4ODvfv3xesPHPmDJ1O/+6779TU1LKzs7ds2VJcXEwVXbhwwdDQ8MSJExwO5+HDhyKn44p37ty5qqqqVatWbdu2raKi
Ij4+/qeffmpbLTIycvjw4RcvXlRQUIiMjOxOAix5p5L0y+VyuznykrcQFha2evXqESNGcLnc3NzcwMBAkU+BPnnyJIvFOnr0KPUa
pLi4OAk7Enn0WwkNDd2+ffv169eZTObKlSs73DtJDlyrNgsKCsSccu2pqakJCAg4ePDgkSNHNmzYUFZW1qkzuVUMkgfw9OlTW1tb
6pJvU1PTy5cvbWxsRN4A3HY3Oxw96A6arq6urGMA6BOsra0luecHAACgR2lpack6BNEOHjz47t07ke/ylZXg4OCMjIyTJ0/KSb8y
JN2jL9sB7INnstwSPJ+sN2EKNAAAAAB0YNy4cePGjRM/3RQGqoF09AfSvkDXYAo0AAAAAIhz9epVFRWVU6dO5eTkyDoW6G0D6egP
pH2BLsMUaIC/YAo0AAD0BX12CjQAgHRhCjQAAAAAAABAT0ECDAAAAAAAAHIBCTAAAAAAAADIhU4/BMva2ron4gDoCa9evZJ1CAAA
AAAA0Fd0OgFGRgEAAAAAAAD9EaZAAwAAAAAAgFxAAgwAAAAAAAByAQkwAAAAgFyg0+lbt26NjIyMjY21sbGRdTh9WlBQ0MaNG2Ud
xV++/vrrrVu39p1OpX4iyWQHQW51+h5gAAAAAOiPpk2b5ujo6OPjU1ZWJutYumLv3r2FhYUhISGyDkTe9eiJNGXKFG9vb0tLy+rq
6tDQ0NDQUKl3AXIOCTAAAACAXBgyZEhxcXE/zX6h7+jRE2nhwoWHDx/Oy8tzcnLavn17Tk5OcnJyT3QEcgsJMAAAAED/tnnz5nnz
5hFCampq0tLSQkJC3r5926rOnj17ZsyYQQiJjY0tLCxctmyZkpLS2rVrXV1d1dXVs7OzT5w48fTpU6oyjUZbvHixp6envr5+Tk7O
0aNHU1NTf/zxx6ioqEuXLlF1du7cyeFw9u7dSwiZNGmSn5+fiYlJRUXFjRs3Ll26xOPxWgXQXnfBwcHFxcUMBmPixIl0Oj06OvrI
kSNtNw8MDPzggw8IIYsWLSKErFy58s2bN15eXlSQRUVFly9fjoiIoCoHBweXl5ezWKxRo0YpKSlFRkaGhoYGBARMmDChtrb2/Pnz
165do6qVlJSoqKiMHTuWTqdHREScOnWqbdc0Gq29jiTZ8faOjpgdZzAYAQEBzs7O9fX1CQkJLBarqamp7XFnMpkbN250cnKqq6tL
SEjQ1taurq6mjoiYg9tekYSdSn4iiYlBTF9ffvkltRATE7Nt2zZdXd22MQB0B+4BBgAAAOjfDhw44Ojo6Ojo6O3tXVpaGhQUpKCg
0KrO119/ffLkyczMTEdHx2XLlhFCfH19nZ2dv/rqK09Pz4SEhP379xsaGlKVP/nkE29v72PHjs2fP//gwYMuLi5iemexWN9+++3N
mzfnzZu3ceNGNTU1CwuLttXEdOfu7p6cnLxs2bIvv/zSzc1t9uzZbTcPCgpKSEgICwuj9rSgoGD16tUeHh67d+/28PD4/vvv/fz8
Zs6cKajv4uISGxu7dOnSPXv2eHl5nT59Ojo6ev78+SEhIf7+/qamplQ1Dw+PzMzM5cuXb9++3cPDY/HixW27bq8jCXdczNFpb8fX
rVs3duzY9evX+/j4NDY2Ojg4iBz5zz77zN7e/osvvvDx8amvrxeuJma02yuSsFPJTyQxMXTYF51O37RpU0lJyaNHj0SGAdBlSIAB
AAAABojKysrDhw+bmJiYmZmJr8lgMBYtWnT69OmXL1+y2exz587l5eVR6R+TyVyyZMnp06fj4uLq6uoyMjK+//57MU1pa2srKys/
evSooaGhpKTk1KlTWVlZkndHCImLi7t161Z9fX16enpCQoK9vX2He8pgMJYvXx4SEpKent7Q0JCSknLt2jUPDw9Bhbi4uKioKA6H
k5CQkJeXl5ycfP/+fQ6H8+DBg9LS0pEjR1LVMjIyLly4UFtbm5KS8ssvv3h5eUnekSQ7Lqzt0RG548rKyvPnz6daq6mp+eGHH0RO
NmYymXPmzBFUO378eGlpaYej3V6RhJ2KPBAiGxQTgyR9BQYG2tnZ+fv719bWShIGgOQwBRoAAACgfzM3N/fz87Ozs9PU1KTRaIQQ
AwOD7OxsMZsYGRkpKSmlpaUJ1qSlpVGJmYmJibKy8vPnzyXsvbi4ODEx8fjx4/fu3Xv27FlycnJjY6Pk3RFChCdss9lsPT29Djs1
NTVlMpkHDhygPtJoNBqNVlRUJKgg3GZtbW2rj+rq6tRyRkaGYH1GRoaenp6qqmpdXZ0kHUmy40Ts0RG540ZGRgwGIzMzk1rf0tIi
Mq82NDRUUlISVOPxeIJqYka7vSIJO21LTINiYhDf17Bhw1xdXb28vHC/OvQEJMAAAAAA/RiNRtu3b19iYqKfn195eTmPx4uOjm47
BbrtVm3X8Pl8QRG1LEkLfD5/69ato0ePnjBhwtq1azdv3rxp06a8vDwJu5Okr7bodDohxMfHp1VHAq3abK+LDrsW05GEOy7m6Ijp
XcIxEa4mGOQOD67IIsk7bbW5yAbFdyS+L+q+X2S/0EMwBRoAAACgH9PV1TUwMLh06VJJSQmXy7W0tFRU7PgKR2FhYXNzs62trWCN
ra1tfn4+ISQ/P7+xsXHMmDGtNmGz2YILp4SQoUOHCpb5fH5KSsqZM2fWrFlTXl4uPBW5w+4kx+VyBWkVFeTkyZM71UJbwiHZ2NiU
l5cLX/7tsKMOd7wLR6eoqKi5uVnwfl0FBQUrK6u21YqLi4Wr0el0S0tLalnMaLdXJGGnbbXXoJgYOuwrKSnJ0dGxpaVFkgAAOgsJ
MAAAAEA/VllZyWazPTw8mEzmsGHDtm7dKslWTU1NYWFhvr6+tra2ampqq1evNjMzo9652tDQcOXKFV9f32nTpqmoqNjY2Pzf//0f
ISQ1NdXFxcXExERVVdXb21uQtNjZ2W3evNnS0pLBYFhZWenr6xcWFkreneTevXtnZWXFYrGoIC9cuODj4zNr1ixVVdXBgwd7enqu
WrWqUw0SQmxsbJYvX66qqjpmzJgVK1ZcuXKlVQUxHUmy4104Oo2NjdevX/fz87OystLQ0Pj8889FzglvaGiIiIhYu3YtVe3TTz/V
19enisSMdntFEnbaVnsNiomhw76cnJxiY2PV1NQkCQCgszAFGgAAAKAf43K5O3bs8Pf3X7JkSUVFRVhYmLGxsSQbnjlzhk6nf/fd
d2pqatnZ2Vu2bCkuLhYU1dbWrl+/XldX9/Xr18eOHSOEXLhwwdDQ8MSJExwO5+HDh4LH86anp1tYWAQGBhobG1dWVoaHh4eHh3eq
OwmFhoZu3779+vXrTCZz5cqV586dq6qqWrVq1bZt2yoqKuLj43/66adONUgIiYyMHD58+MWLFxUUFCIjI9smwISQ9jqSZMe7dnRO
njzJYrGOHj1KvSUoLi5OZLXjx48LV0tKShK8TEj8wRVZJGGnbbXXoJgYutwXQPfR8HItAIq1tXV79xEBAAD0Gi0tLVmHIC+Cg4Mz
MjJOnjwp60CkgEajnT9/nnoXsaxjAZBUVVVV73eKK8AAAAAAAP3PuHHjTE1NY2Ji6HT6smXL9PX1Y2JiZB0UQF+HBBgAAAAAoP95
+fKlg4PD2bNnWSzW69ev/f393717J+ugAPo6TIEG+AumQAMAQF+AKdAAICdkMgUaT4EGAAAAAAAAuYAEGAAAAAAAAOQCEmAAAAAA
AACQC0iAAQAAAAAAQC505SnQ27bto9HolZXlJ07s5fP5CgoKn3/+lZqaRmNjw/ffbyeEKCgoOjjMGjFirKamVmNjY37+6wcPbr5/
Xy7Yls/nNTU1VVVVpKenPn78oKWlRVAk6OXZs4Rbt64SQkxNLT780ENPz5BGo9XX1z55EpeY+EAqOw8AAAAAAADyo+uvQdLW1h0+
fFRGRurIkePV1DSEixYuXGVlNSIxMSY+/r6RkcnHH682NbU8e/ZgTU0VVeHw4V2KikrTp7s6Oc02NbW4ePE0IXyq6NChHfX1df+N
T1Fp0aI1dXXs48f/yePxrKxGqKtrdjlmAAAAgD6uqqoKD4IGgAFPJo+AJl2eAt3Swi0pKfzggw8JoU2Z4vT2bZ6gyMjIxMpqRF0d
+8GDWw0NnJyczOfPk1gslYkTpwu3UFNTdfNmWGVlhbm5tYWFTXsdaWkNUlZm8vn8pqbG+vra58+T4uLudi1mAAAAAAAAkGddvwc4
MTHG0NDYxWWejo5+QkKMYL2h4VBCSFnZOx6PR60pKSkkhBgaGrdqgc/nFRbmE0KMjc0EKwMCdgUGHggMPGBvP5EQUlVVUVtbo6Oj
HxCwy9v7cwcHFyaT1eWYAQAAAAAAQG51fQp0enqqk9P7iROnl5YW5eZmtinnC5ZoNFp7jVBFfP5/K7eaAs3lcn/++YepU50tLGyM
jc2Njc1tbOzPnPm+y2EDAAAA9HGYBQ0AA5us5j+T7iTAfD7v8ePfP/pogfDlX0JIcfFbQoie3mA6nU5dBDYwMCKEFBe/adUCjUY3
MjIhhLx9my+mo6qqips3QwkhRkYmq1d/oa9vyGSyGho4XY4cAAAAAAAA5FDXE2BCSHLyo+TkR4QQRcX/tlNUVJCVlWZlNWLGDLeE
hBhDQ2N7+4kcTv0ffzwU3lZDQ2v6dFdtbZ3c3FfZ2RntdaGlpePs7PHHHw/fvStqbm5qaWmpr69B9gsAAAADGy4CA8BAJcPLv6Sb
CXB7rl372cFhlp3d2MmTZzQ2NubkZD54cFPwCGhCyIYNO5qbm6uqKn7//XZiYozwfOmAgF3UwuvX6VeunG1ubmQyWQsXrmIyVbjc
5sLCgnv3bvREzAAAAAB9CnJgABh4ZJv9EkJourq6so0AoI+wtrbOy8uTdRQAAAD/hQQYAAYYmSfAXX8KNAAAAAD0KJl/UwQAkKK+
8DutR6ZAAwAAAIBUUN8XcSkYAPq1vpD6UnAFGAAAAKCv6zvfHQEAOqtP/QbDFWAAAACAfgCXggGg3+lTqS8FCTAAAABAvyH4NolM
GAD6rD6Y9wogAQYAAADof/ry90sAgD4L9wADAAAAAACAXEACDAAAAAAAAHKh01OgGdZDeyIOgJ7Q9OqtrEMAAAAAAIC+ohMJMFJf
6HeokxZpMAAAAAAAEMmnQCP7hf4LZy8AAAAAABAJE2DkD9Df4RwGAAAAAICOE2BkDjAw4EwGAAAAAJBzeAo0AAAAAAAAyIUOEmBc
NIOBBOczAAAAAIA8wxVgAAAAAAAAkAu9mgA3eDlxrY17s0cAAAAAAAAACq4AAwAAAAAAgFxQ7IlGefrazeOteYMH8RUV6NV1Cllv
lf7MIdyWzrbTsGiG4otcxYyCnggSoH/R0tKSdQgAAH1LVVWVrEMAAIB+RvoJcIuJfqPbZKUXOYyEl7RaDk9DtcVqaIvpYIXsQqn3
BTDgIe8FAGiP4DckMmEAAJCQ9BPgJsfRihn5SvEvqY/09zX0x2mCUr6WWuO8aTx9LVLXwIj7U+FNKSGk0X1Ki6kBIYTGrldMz1dK
fkUIaXSZwNPTavpwbNOHY+nvKpnXYqUeKkBfhtQXAEBC1C9MpMEAANAhKSfAPG11vrqKYuab9ipwhxsr3/mD/r6mebRF08xxrH/d
JoQo30wkhBAajaej0eg6kVZTr5j1VvnukwZNVUyBBvmE7BcAoLO0tLSQAwMAgHjSfggWS5kQQqtraK9cKTWbXlpJuC2KL3L5LGW+
KvO/ZXw+vbxa6c+cFrPBUo4KoP/Q0tJC9gsA0DX4FQoAAOJJewo0p5EQwldl0mo57VRoov6fRj0TS1GBENJiYdQ8zpqnqUaUFAgh
9JL3Uo4KoJ/A9zYAgO7DpWAAAGiPlBNgeiWbxq7nWhsz3lVKuAmfpdzoMkH5zh/0t2W0Zi7X3oJrOeQ/ZdKNDgAAAAAAAOSX9N8D
zIh9zrU1bZ4ygq+pShQVeIPUmyfbtlgMaXcDRQVCo5GmZhqPxzMY1DzaQlBCq2/gDdIgdJrUgwTog3D5FwBAWvAbFQAARJL+U6AV
CkqZv8U1j7du+NiRr/DXe4AV8kvaq09j1yslvGxymcBXVqKXVSvkFPEMBlFFSs+ymj4cWz9qGL2sCk+BhoEN39UAAKQLE6EBAKAt
mq6urphihvXQXgsFoBc0vXrbXpG1tXVeXl4vxvI/kAADAEgdEmAAAGhF+lOgAaCzkP0CAPQE/HYFAIBWkAADAAAAAACAXEACDAAA
AAAAAHIBCTCAjGGGHgBAz8HvWAAAEIYEGGCgCQ4OXrduXffbCQoK2rhxY/fb6bkGBzDhsRIzbl9//fXWrVul2K/UG5QtMUOHsxEA
AEA+IQEG6K+srKxiY2OPHTsm2zD27t3r7+8vxQavXr26Zs0akUWBgYEhISHCa+zt7WNjY01MTKQYwM6dO319faXYYO/T0dHx9/e/
cuVKVFTU2bNnXVxcej+GVgdr/Pjxd+7cWb9+PY3W6Ve7S+Uck/qJCgAAAP2R9N8DDAC9Y+7cuRkZGSNGjDA1Nc3Pz5d1ONBTAgMD
O7uJp6dndnb2xYsX6+rqZs6c+dVXX1VXVyclJfVEeJKYMWPGN99889NPP/3888+92a+YoevCqAIAAMAAgAQYoF9SVlZ2cXHZtWuX
l5fXnDlzfvjhB+FSLS2tXbt2jR07lk6nR0REnDp1isfjTZo0yc/Pz8TEpKKi4saNG5cuXeLxeIQQJSWltWvXurq6qqurZ2dnnzhx
4unTp626+/HHH6Oioi5dukR93LlzJ4fD2bt3b2Bg4AcffEAIWbRoESFk5cqVBQUFNBrNy8vL09NTX1+/qKjo8uXLERERwq15eHis
W7du4cKFXC6XWrNr1y4FBYXt27d3c1jEdB0cHFxSUsJisaZMmZKZmRkQECA+SEJIeyMmbPPmzfPmzSOE1NTUpKWlhYSEvH0r4l3T
3Rn8oKCgsrKy4OBgQgiDwQgICHB2dq6vr09ISGCxWE1NTW27O3v2rGD5xo0bq1evtrOzE5kAi2mwvV3r7OGbO3fuxo0bDx06FB4e
Ln5zkaMh8hyTZNyEh65tI59++qlgVIODg4uLixkMxsSJE+l0enR09JEjR6gDJPkZ9ebNm/bOhFY1o6Kieuj8BwAAgA5hCjRAv+Tk
5FRXV5eUlBQeHv7RRx8pKv7PH7M8PDwyMzOXL1++fft2Dw+PxYsXs1isb7/99ubNm/Pmzdu4caOampqFhQVV2dfX19nZ+auvvvL0
9ExISNi/f7+hoaGEYQQFBSUkJISFhTk6Ojo6OhYUFBBCVq9e7eHhsXv3bg8Pj++//97Pz2/mzJnCW927d4/BYEydOpX6qKGh4eDg
cPPmze4OSkddu7u7P336dMGCBRs2bOgwSDEjJuzAgQPUvnt7e5eWlgYFBSkoKLSqI8XBX7du3dixY9evX+/j49PY2Ojg4CB+QJhM
ppubm4aGRnx8fGcbbG/XOnX4VqxYERAQsHv3bir7Fb+5yNEQeY51dtxENiLM3d09OTl52bJlX375pZub2+zZs6n1kp9R4s8E4Zo9
d/4DAABAhzpIgJteibiUAdBPDaTzec6cOZGRkTweLy4ujs/nT5s2Tbg0IyPjwoULtbW1KSkpv/zyi5eXl7a2trKy8qNHjxoaGkpK
Sk6dOpWVlUUIYTAYixYtOn369MuXL9ls9rlz5/Ly8hYvXtzlwBgMxvLly0NCQtLT0xsaGlJSUq5du+bh4SFcp6GhITo6WrDSxcWl
pqYmMTGR+tjY2NjQ0NBe+2PGjIkVcvToUcm7fvLkSXh4OIfDEV9z586dZ86caW/E2lNZWXn48GETExMzM7NWRdIafGVl5fnz51Ob
19TU/PDDD2VlZe1VNjc3j42NjYqK2rx58/79+zMzM7vcYKtdE3/4hNnZ2a1bt+7bb7998OCBYGV7m0s+GlI/aQkhcXFxt27dqq+v
T09PT0hIsLe3J505o8QMV9uakg8gAAAASB2mQAP0P0OGDBk1atSePXsIIVwu9+bNm3PmzPn9998FFTIyMoSX9fT0qqurExMTjx8/
fu/evWfPniUnJzc2NhJCjIyMlJSU0tLSBPXT0tLapnCSMzU1ZTKZBw4coD7SaDQajVZUVNSqWkRExIkTJ3R0dCoqKtzd3W/fvi2Y
Xcxms9lsdnvtp6SkCD/KyN7eXpADd9h1bm6u5EEWFxeLHLFWzM3N/fz87OzsNDU1qcc7GRgYZGdnS9JUZwffyMiIwWAIUtmWlhYx
OXlubq6jo6Oampqzs/O2bdvq6+sfPXrUqQbF7JqYwyfs7du3dDrd29v76dOnVVVVgvUiN5d8NKR+0lKhCpbZbLaenh7pzBlFOjoT
hGsSiQcQAAAApK7jBLjp1VuG9dBeCAWgRw2wy790Oj0sLEywhsfj6evrl5aWUh/5fH7brbZu3Tp69OgJEyasXbt28+bNmzZtysvL
a/tIXhqNJnLzVnXaK6LT6YQQHx+fvLw8MS1kZGTk5OS4ubklJSVZWVnt3LlTUNTldzh12HVzc7PkQfL5fJEjJlyHRqPt27cvMTHR
z8+vvLycx+NFR0e3nQLdXlNdG/wOKwirra0NDw+fMmXK/PnzHz16NH369G+//ZYq+uc//0n9oURkg+J3TczhE1ZdXb1jx45Dhw4d
Pnw4ICCgsrJSzOaSj0bXxk08kZtLfkZ1eCYIalIkHEAAAACQOonuAR5ImQPIp4F0DisoKMyePXvnzp2OQv788093d3dBHVtbW8Gy
jY1NeXl5XV0dn89PSUk5c+bMmjVrysvLqRmYhYWFzc3NwvVtbW3bPlOazWarq6sLPg4d+t8/inG5XOGEJD8/v7GxcfLkyR3uSERE
hLu7u4eHR2pq6ps3bzoxBO2QvGsJa4ocMWG6uroGBgaXLl0qKSnhcrmWlpatbsYW35SEgy9QVFTU3NxsY2NDfVRQULCysupwZwkh
gqgePnwoOGdu3bolpsEOd03Cw1dZWenv79/S0nL48OFBgwaJ2VzMaLQ6xzo7bpRWjUhC8jNK8jNBQOrnPwAAAEhC0odgDaT8AeTN
ADt7P/jgA3V19cePHwuvfPjwobu7O3XBihBiY2OzfPlyVVXVMWPGrFix4sqVK3Z2dps3b7a0tGQwGFZWVvr6+oWFhYSQpqamsLAw
X19fW1tbNTW11atXm5mZhYaGtuo0NTXVxcXFxMREVVXV29tbOO969+6dlZUVi8WiPjY0NFy4cMHHx2fWrFmqqqqDBw/29PRctWpV
2x2JiooyMDCYO3duq8f/iHkPsHiSdy1JzfZGTFhlZSWbzfbw8GAymcOGDdu6davIwLo5+AKNjY3Xr1/38/OzsrLS0ND4/PPPqZm6
be3evdve3l5FRUVbW9vLy2vSpEm3bt3qVIMd7lp7h6+t6urqgICApqamkJAQHR2d9jYXMxqtzrHOjhulVSOSkPyMkvBMENZqBNas
WSPyGAEAAIB0deIeYCqLwHRo6EcGWOpL8fDwePr0aV1dnfDK2NjYL774Yvz48X/88QchJDIycvjw4RcvXlRQUIiMjLxy5QohxMLC
IjAw0NjYuLKyMjw8XPBU3jNnztDp9O+++05NTS07O3vLli3FxcWtOr1w4YKhoeGJEyc4HM7Dhw+F7yYNDQ3dvn379evXmUwm9Yqa
c+fOVVVVrVq1atu2bRUVFfHx8T/99FPbHamtrf39998dHBzu378vrcGRsGtJaqanp7c3YgJcLnfHjh3+/v5LliypqKgICwszNjZu
25eYpiQZfGEnT55ksVhHjx6l3loUFxcnslpYWNjq1atHjBjB5XJzc3MDAwPbewp0ew12uGudOnw1NTUBAQEHDx48cuTIhg0bysrK
RG7e3mi0Pcc6O25tG+kwZoqEZ5SEZ4Kwnjj/AQAAoEM0XV1dWccA0CdYW1uLv221h2hpafV+p33EwYMH3717t3fvXlkHAl3RzcOH
o48R6DXCz2ADAAA5h/cAA4BsjBs3bty4cR3OXIW+qZuHD0cfIwAAACATeA0SAMjA1atXVVRUTp06lZOTI+tYoNO6efhw9DECAAAA
soIp0AB/wRRoAIABCVOgAQBAAFOgAQAAAAAAQC4gAQYAAAAAAAC5gAQYAAAAAAAA5MKASoB9fTeNHDle1lEAAAAAAABAXzSgEmAA
AAAAAACA9vTIa5CMjEymTZs5dKi5kpLS+/flL18+/eOPOC63uSf6AgAAAAAAAJCE9BPgYcNsFi/2efLk0b17EWx2lZaWjp3dOCur
Eenpz6XeFwAAAAAAAICEpJ4A02bPXvj8edK9ezeoz2VlJQ8e3KSWVVXVZs2ab25uxeeTnJyM6OhwDqfexsbeyWn2iRN7CSEzZrhN
mzbr2LF/VlVVGBmZLFvmFxz8DY/HE7khIcTXd9Pr1+lDh5oNHjz05s1QQRBKSkqenivpdPq1a/9ubm6S9j4CQA8KDg7OyMg4efJk
26KgoKCysrLg4GDJW+vCJl3Wo319/fXXTU1Ne/fu7VNRyRvpnpwAAADQ+6ScAOvq6mtqDvrzz2SRpQsWeDc3N505c5BOp8+du3Te
vGWXL5/Ny8vS1vbW0NCqqakyM7N6/77c3Nzq2bMKc3PrgoJsHo/X3oZUm6NHT7p69afCwgI+nzd1qjMhRFVV3cvrb8XFb+/c+ZXP
50l3BwH6gsDAQDc3N0JIS0sLm83Ozc29f/9+RERES0uLVNpnMpkuLi6enp5WVlb/+Mc/oqKiOhWVwBdffJGamiqVkKQuMDBw8ODB
/v7+sg5ECjQ0NFauXDlt2rTBgwfX1dUVFhZGRUXdunWroaGhR/vt6fPQ2tp6xYoVY8aMUVBQoNLOrKwsCbe1srI6e/bsixcv/v73
v0slGAAAABgYpJwAq6ioEULY7Oq2RVpaOiYmFj/8EFRbW0MIiYr6be3azWpqGrW1NSUlb83NrTMyUnV1DaKjwy0sbJ49SzQzs8rM
/FP8hoSQJ08evX2bJ+hFV9fA0fGjZ88SEhJipLtrAH1KSkqKv78/nU7X0tIaP378unXrXFxcNm3a1NQkhSkP7u7uVlZW+/btO336
dBei6n4A7QkMDOyFTbqsN/sS0NPTO3bsWFFR0d69ezMzM5WVlYcMGeLq6jpr1qyIiIiejqpHz8NPPvnk9u3bhw4d4vP5fn5+hw4d
8vHxKSsrk2TbuXPnZmRkjBgxwtTUND8/v/vBdEgmRx8AAAA6S8oJcH19LSFEXV2zpqaqVZGGhhaPx6uufk99rKwsp1bW1tbk5WWZ
mVlxOHWFhXmvX6c7O89hMJSHDjW9c+ea+A0JIdXVlcK9jB49sb6+Ljk5Xrr7BdA38Xi89+/f37179+XLl//+978XLVp04cIFQsjm
zZvnzZtHCKmpqUlLSwsJCXn79q2Hh8e6desWLlzI5XKpzXft2qWgoLB9+/ZWzV67dk1aEf74449RUVGXLl2iPu7cuZPD4VDzeCdN
muTn52diYlJRUXHjxo1Lly5RMz4IIZqamtu3b584cSKdTo+Ojj5y5AhVJJhlGhwcXF5ezmKxRo0apaSkFBkZGRoaGhAQMGHChNra
2vPnzwt2QcKJqcHBwcXFxQwGo1WnIkdSsElJSQmLxZoyZUpmZuaGDRva60t8IyL7JYQwGIyAgABnZ+f6+vqEhAQWiyUyqwwICCCE
bNmyhSptbGysqalJT08XVBCOSklJae3ata6ururq6tnZ2SdOnHj69Kn4MCTRQ+fh1q1bhQ/Q7NmzJ0yYcOvWrQ7jUVZWdnFx2bVr
l5eX15w5c3744QfhdoSP2ps3b9o7NO2dnx2enGJ2UPyZIOEpLaYRAAAA6JCUX4NUXl5aXf1+5MhxbYtqaqrodLqGhhb1UVtbl1pJ
CMnNzTIzszQ3t87NzaqrY9fW1kycOJ3DqS8vfyd+w7ZiYm6WlZUsW+bHZLKku2sAfVlRUVFiYuKMGTOojwcOHHB0dHR0dPT29i4t
LQ0KClJQULh37x6DwZg6dSpVR0NDw8HB4ebNm53qyNXVNTY2VlNTs5sBs1isb7/99ubNm/Pmzdu4caOampqFhYWg1N3dPTk5edmy
ZV9++aWbm9vs2bPbtuDi4hIbG7t06dI9e/Z4eXmdPn06Ojp6/vz5ISEh/v7+pqamnQ1JZKciR1J4k6dPny5YsGDDhg1iWu6wEZE7
u27durFjx65fv97Hx6exsdHBwaFtyyoqKlOnTv3tt98kvOLq6+vr7Oz81VdfeXp6JiQk7N+/39DQUHwYndJz56G6ujqdTq+rq6M+
ij8PnZyc6urqkpKSwsPDP/roI0XF//lTr/BRa+/QiDk/OxwoMTso/kyQ8JQW3wgAAACIJ/X3APNv3bo2evTkDz/0GDRIV0lJSU9v
sJPTbFvb0VVVFQUFOW5uC9XUNDQ0tFxdPbOzM6iruG/e5CorM0eOHJ+b+4oQkpubNWWKE7VMCBGzYVs8Hu/69V/KykpWrvxMVVVN
2nsH0HcVFBQIkhmBysrKw4cPm5iYmJmZNTQ0REdHe3h4UEUuLi41NTWJiYlSjGHMmDGx/yF++rS2traysvKjR48aGhpKSkpOnTol
fHtnXFzcrVu36uvr09PTExIS7O3t27YQFxcXFRXF4XASEhLy8vKSk5Pv37/P4XAePHhQWlo6cuTIzgYvvlPhkRSsfPLkSXh4OIfD
kbALkY2I7FdZWXn+/PnUsNTU1Pzwww8ip/4aGRkpKCi8efNGkt4ZDMaiRYtOnz798uVLNpt97ty5vLy8xYsXS7L7kuuh83DDhg2l
paVJSUmSxDBnzpzIyEgejxcXF8fn86dNmyZcKvKotTo0Ys7PDgdKkh1s70zo1CktshEAAAAQT/qvQcrJyTh//ti0abNWr/an3gP8
4sXTrKw0Qsivv/7bxWW+r+//EcLPycmMjg6nNmlp4b55k6uvb1RaWkIIyc19NWnS9Nzc/34bbm9Dkfh8/s2bobNmzVu58vMLF06I
vCEZYEDi8/nUgrm5uZ+fn52dnaamJo1GI4QYGBhkZ2dHREScOHFCR0enoqLC3d399u3bkk9zpURFRYl5Jpbk9wAXFxcnJiYeP378
3r17z549S05ObmxsFJQKT+lks9l6enptWxCuU1tb2+qjurq6JGG016Cg0/ZGkqqWm5srScviGxHZr5GREYPByMzMpNa3tLSIef6T
4LgTQiIiIjQ0NAghCQkJwlOIqTaVlJTS0tIEa9LS0gS5kyRjLiGpn4dr166dMmXKhg0bBI/1EnMeDhkyZNSoUXv27CGEcLncmzdv
zpkz5/fffxdUED5q7UUo5vyUZKDa20HJzwQxp7T4RgAAAEA86SfAhJCiooLQ0B/brq+rY//223mRm1y8eEqwnJ2dHhS0WZINz5w5
2N7H6Ohw8XkywABjYmJSXFxMCKHRaPv27UtMTPTz8ysvL+fxeNHR0dQkyYyMjJycHDc3t6SkJCsrq507d/ZykNT3dUIIn8/funXr
6NGjJ0yYsHbt2s2bN2/atCkvL09Q2mFTrepIskmnGiRiR5LS3NzcYbMdNiIm8g53qqioqKWlxdjYWLBmzpw5hJA9e/YwGIy2kbRd
I+ii+wNIkfp5uGbNmkWLFm3ZsiUjI0OSAObMmUOn08PCwgRreDyevr5+aWkp9VFw1MREKOb8lGSgRO5gZ88EkR112AgAAACIJ/Up
0AAgA4aGhlOmTImNjSWE6OrqGhgYXLp0qaSkhMvlWlpaCt8DGRER4e7u7uHhkZqaKuHU2e5gs9nCF2OHDh0qWObz+SkpKWfOnFmz
Zk15eblgymjfIX4ke7SRoqKi5uZmGxsb6qOCgoKVlVXbatTzsTw9PZWUlDpss7CwsLm52dbWVrDG1tZWuk9Ilvp5uGbNmqVLl27Z
skXC92kpKCjMnj17586djkL+/PNPd3f3tpXFR9jN87PtDsrwdAIAAAABJMAA/RidTtfW1p41a9bhw4czMzOvXr1KCKmsrGSz2R4e
Hkwmc9iwYa3mwUZFRRkYGMydO7ezj7+idPYhWKmpqS4uLiYmJqqqqt7e3oIszs7ObvPmzZaWlgwGw8rKSl9fv7CwsAvx9CjxI9mj
jTQ2Nl6/ft3Pz8/KykpDQ+Pzzz9vb07yoUOHFBQUjh49Om7cOA0NDSaTaWVlZWRk1Pb6YVNTU1hYmK+vr62trZqa2urVq83MzEJD
Q8VHsmbNmg4fvNxD5+GqVavay37bOw8/+OADdXX1x48fC698+PChu7s7nd763zsxEXb//Gy7gzI8nQAAAEAAfzkG6Jeox03xeDw2
m52Tk3Px4sUbN25Qr13hcrk7duzw9/dfsmRJRUVFWFiY8BTZ2tra33//3cHB4f79++01PmHChIMH/7qhYPv27du3b4+MjKTeXdRZ
Fy5cMDQ0PHHiBIfDefjw4aNHj6j16enpFhYWgYGBxsbGlZWV4eHh4eF97p4F8SPZ042cPHmSxWIdPXqUuswbFxcnslppaeknn3yy
YsWKzZs3GxgYcDickpKS+Pj4X3/9tW3lM2fO0On07777Tk1NLTs7e8uWLdR05S7r0fNw9erVSkpKR48eFaw5d+7cuXPnxMTj4eHx
9OlTwcOiKbGxsV988cX48eP/+OMP4fViIuz++dl2B2V7OgEAAACFpqurK+sYAPoEa2trwT2ovUlLS6uXezx48OC7d++6ltACSMuA
Pw8H/A72I1VVVbIOAQAA+gpMgQaQL+PGjRs3blyHE18BetSAPw8H/A4CAAD0U5gCDSBHrl69qqKicurUqZycHFnHAvJrwJ+HA34H
AQAA+i9MgQb4i/xMgQYAkCuYAg0AAAKYAg0AAAAAAAByAQkwAAAAAAAAyAUkwAAAAAAAACAXkAADAAAAAACAXEACDAAAAAAAAHIB
CTAAAAAAAADIBSTAAAAAAAAAIBeQAAP0P8HBwevWreuPjQ8MX3/99datWzu7VVBQ0MaNGzus1tnx7/7xEgQmYYT9lPDe9bU9lUk8
XTuNAQAA+jtFWQcAADBwBAYGDh482N/fn/o4fvz4oKCgiIiIo0ePyjYwKQoMDHRzcyOEtLS0sNns3Nzc+/fvR0REtLS0SKV9HR2d
FStWODg4aGlpvXnz5tKlS3fv3pUwqnPnzp07d45aY29vf/ToUU9Pz/fv30slME1NTW9v72nTpunr61dVVWVlZV2+fPnZs2edbWfv
3r2FhYUhISFSiQoAAAA6BQkwAECPmDFjxjfffPPTTz/9/PPPhJDAwEBZRySaIDDJI0xJSfH396fT6VpaWuPHj1+3bp2Li8umTZua
mpq6H4+np2d2dvbFixfr6upmzpz51VdfVVdXJyUldbhhU1PT0qVLf/vtt8rKyralwnvXhWOhp6d37NixoqKif/zjH1lZWerq6paW
lqtXr37+/DmPx+tsawAAACArSIAB+iUtLa1du3aNHTuWTqdHREScOnWK+ha+efPmefPmEUJqamrS0tJCQkLevn1LCAkODi4vL2ex
WKNGjVJSUoqMjAwNDQ0ICJgwYUJtbe358+evXbvWYeNKSkpr1651dXVVV1fPzs4+ceLE06dPqU3EFAljMpkbN250cnKqq6tLSEjQ
1taurq7eu3dvcHBwSUkJi8WaMmVKZmbmhg0bRDY4b968Tz75ZOHChYKLjd988w2Lxdq2bRshhEajeXl5eXp66uvrFxUVXb58OSIi
om0MYoaouLiYwWBMnDiRTqdHR0cfOXKE2nEGgxEQEODs7FxfX5+QkMBisTrM9ObOnbtx48ZDhw6Fh4dTa4KCgsrKyoKDg8X3JWzq
1KnffPNNcHDwqFGjRMbcXv07d+78+OOPUVFRly5doop27tzJ4XD27t3bditBYMIRSoLH471///7u3bsvX77897//vWjRogsXLogZ
YQ8Pj3Xr1i1cuJDL5VIt7Nq1S0FBYfv27cLNnj17VrB848aN1atX29nZSZIAP3/+XEtLa82aNQcPHhSzm62WJTx1N2zYQAjZsmUL
degrKioqKioeP35Mlba3y5MmTfLz8zMxMamoqLhx48alS5f+3//7fx988AEhZNGiRYSQlStXFhQUCHfUp05jAACAgQf3AAP0Sx4e
HpmZmcuXL9++fbuHh8fixYup9QcOHHB0dHR0dPT29i4tLQ0KClJQUKCKXFxcYmNjly5dumfPHi8vr9OnT0dHR8+fPz8kJMTf39/U
1LTDxn19fZ2dnb/66itPT8+EhIT9+/cbGhp2WCTss88+s7e3/+KLL3x8fOrr6x0cHARF7u7uT58+XbBgAZVpiGwwJiZGTU1twoQJ
1CZMJtPBwSEqKor6uHr1ag8Pj927d3t4eHz//fd+fn4zZ85sG4OYIXJ3d09OTl62bNmXX37p5uY2e/Zsav26devGjh27fv16Hx+f
xsZG4bBFWrFiRUBAwO7duwXZb1vt9SXw0UcfffPNN7t3775z546YmEXWFx+e1BUVFSUmJs6YMYP62F609+7dYzAYU6dOpappaGg4
ODjcvHmzvWaZTKabm5uGhkZ8fDy1xtXVNTY2VlNTU2R9Pp9/8uTJuXPnGhsbSx68JKcui8WaNm3ab7/91l7GKHKXWSzWt99+e/Pm
zXnz5m3cuFFNTc3CwiIoKCghISEsLIyq3yr7JX3pNAYAABiQkAAD9EsZGRkXLlyora1NSUn55ZdfvLy8WlWorKw8fPiwiYmJmZkZ
tSYuLi4qKorD4SQkJOTl5SUnJ9+/f5/D4Tx48KC0tHTkyJHiG2cwGIsWLTp9+vTLly/ZbPa5c+fy8vKo3FhMkTAmkzlnzpxTp05l
ZWXV1NQcP368tLRUUPrkyZPw8HAOhyOmQTab/fjxYxcXF2oTR0fHlpaWR48eUZssX748JCQkPT29oaEhJSXl2rVrHh4eYsZQ5BDd
unWrvr4+PT09ISHB3t6eEKKsrDx//nxB2D/88ENZWZmYZu3s7NatW/ftt98+ePBATDWRfQksXrx4w4YNW7duFeR+7cUsvn6vKSgo
aJs3toq2oaEhOjpacFBcXFxqamoSExPbtmZubh4bGxsVFbV58+b9+/dnZmZKGEZSUtLz58/9/PwkrC/hqWtkZKSgoPDmzZsOGxTe
ZW1tbWVl5UePHjU0NJSUlFCnUIfx9JHTGAAAYKDCFGiAfikjI0N4WU9PT1VVta6uztzc3M/Pz87OTlNTk0ajEUIMDAyys7MJIcKT
Zmtra1t9VFdXF9+4np6ekpJSWlqaoCgtLY36zm1kZNRekTBDQ0MlJSVBMsPj8YTzgdzcXMGymAajoqK2bdvGZDIbGhpcXFwePHhA
XZQzNTVlMpkHDhyg6tNoNBqNVlRU1HboJBwiNputp6dHBcNgMARht7S0iE9j3r59S6fTvb29nz59WlVVJaZa274orq6uWlpan376
qaAjMTGLrC8TfD6fWhATbURExIkTJ3R0dCoqKtzd3W/fvi3yBtrc3FxHR0c1NTVnZ+dt27bV19dTf+aIiooSXPBvz4kTJ06ePDli
xAhJYpbw1KX2QrCDbYnc5YSEhMTExOPHj9+7d+/Zs2fJycmNjY3i4+k7pzEAAMBAhSvAAP2SyO/iNBpt37595eXl1LRJJycnLpcr
mBjZahMx3+bba7ztGqqmmCLxjQtv2NzcLElf8fHxPB7PwcFBW1t7woQJgnSITqcTQnx8fJycnJycnGbMmOHo6Lh06dK27Ug+RO2F
LV51dTX1gKjDhw9ra2t3ocGMjAw2my2YuSo+5rb1RRIM6fTp02P/Q/wmnWJiYlJcXNxhtBkZGTk5OW5ublZWVlZWVmLmPxNCamtr
w8PDExMT58+fL3kkGRkZMTExn332mSSVJTx1i4qKWlpaTExM2mtE5C7z+fytW7fu3r27sbFx7dq1Fy9ebJtat9J3TmMAAICBqs8l
wKqq6oGBB1RUVGUdCECfZmtrK1i2sbEpLy+vq6vT1dU1MDC4dOlSSUkJl8u1tLRUVOzKLA+RjRcWFjY3NwsX2dra5ufnE0LEFAkr
Li5ubm62sbGhPtLpdEtLS5EBiGmwqanpwYMHLi4uzs7O79+/f/78OVUhPz+/sbFx8uTJ4netC0NUVFQkHLaCgoKVlZX4TSorK/39
/VtaWg4fPjxo0CDxldsqKCjYsGHDzJkzqdcpdRhzq/oUNpstfFV/6NCh1MLDhw8d/+PWrVudjU0kQ0PDKVOmxMbGShJtRESEu7u7
h4dHamqqJJOKu3AOnz592s7OjnrWlHgSnrr19fXx8fGenp4MBqNtI2J2mc/np6SknDlzZs2aNeXl5dRkZi6X2zbxpvSp0xgAAGBA
6nMJMABIwsbGZvny5aqqqmPGjFmxYsWVK1cIIZWVlWw228PDg8lkDhs2bOvWrVJsvKmpKSwszNfX19bWVk1NbfXq1WZmZqGhoeKL
hDU0NERERKxdu9bKykpDQ+PTTz/V19cXGYD4Bu/evTtp0qT58+ffvXtXMIG2oaHhwoULPj4+s2bNUlVVHTx4sKen56pVq1q13IUh
amxsvH79up+fHxX2559/LjxduT3V1dUBAQFNTU0hISE6Ojod1m8lPz/f39+fymkliVm4PrUmNTXVxcXFxMREVVXV29u7s9nOmjVr
OkyP6XS6trb2rFmzDh8+nJmZefXqVSLBCEdFRRkYGMydO7e9y7+7d++2t7dXUVHR1tb28vKaNGmSIBLxD8ESKCwsvHHjRttbeduS
8NQlhBw+fJhOp+/fv3/EiBFKSkqDBg2aPHlycHAwnU5vb5ft7Ow2b95saWnJYDCsrKz09fULCwsJIe/evbOysmKxWG176WunMQAA
wMDTI/cAGxmZTJs2c+hQcyUlpffvy1++fPrHH3FcbnPHWwKAZCIjI4cPH37x4kUFBYXIyEgqR+VyuTt27PD391+yZElFRUVYWFin
HocrvnFCyJkzZ+h0+nfffaemppadnb1lyxZq1qv4ImHHjx9nsVhHjx6lXsSSlJTU3mN1xTSYkpLy/v17MzOznTt3Cm9y7ty5qqqq
VatWbdu2raKiIj4+/qeffmrVbNeG6OTJk8Jhx8XFdbgJIaSmpiYgIODgwYNHjhyhHm3dKfn5+Rs2bDh8+DCfz5ckZuH6R44cuXDh
gqGh4YkTJzgczsOHD6l7aKVlzJgxsbGxPB6PzWbn5ORcvHjxxo0b1MuNOhzh2tra33//3cHB4f79+yIbDwsLW7169YgRI7hcbm5u
bmBgYBee7PWvf/3ro48+EnnBthUJT93S0lJfX19vb+9vvvlGT0+vuro6Kyvr559/5vF4PB5P5C6np6dbWFgEBgYaGxtXVlaGh4dT
TwUPDQ3dvn379evXmUxm29cg9bXTGAAAYICh6erqSrfFYcNsFi/2efLk0bNniWx2lZaWjp3duHfvCtPTn0uyuaqq+oYNOw4d2lFf
XyfdwADEs7a2zsvL6/1+tbS0er/TvoBGo50/f556OaqsY5Fr+/bty83NPX78eK/1ePDgwXfv3ol8KXHv6P1dBtkS8zg6AACQN1K/
AkybPXvh8+dJ9+7doD6XlZU8ePDXPDdVVbVZs+abm1vx+SQnJyM6OpzDqSeEqKlpuLsvNjEZVlNTlZz837/0MxjKM2bMHj58pLIy
882b3Dt3rlVXV0o7YADoPePGjTM1NY2JiaHT6cuWLdPX14+JiZF1UHJNW1vbwsJC/BubpGvcuHHjxo3729/+1ms9ttL7uwwAAAB9
h5QTYF1dfU3NQX/+mSyydMEC7+bmpjNnDtLp9Llzl86bt+zy5bPU+ro69vHj/2QwlD09Vwrqz527TFFR8fz5YxxOvYODy6JFPj/+
eAgPsQTov16+fOng4HD27FkWi/X69Wt/f/93797JOij5NXbs2L179yYmJkZHR/dOj1evXlVRUTl16lROTk7v9NhK7+8yAAAA9ClS
ToBVVNQIIWx2ddsiLS0dExOLH34Iqq2tIYRERf22du1mNTUNJSUlY2Pzw4d31tXV1tXVxsTcXL58HSFEXV1z+PCRISG7qfr370ds
2rRHT29waamIu7MAoF9obGwMCQkJCQmRdSBACCHPnj1zdXXtzR4//vjj3uyurd7fZQAAAOhTpJwA19fXEkLU1TVraqpaFWloaPF4
vOrq99THyspyaqWiolJzc1NdXe1/1ldQC1pagwgh/v7fCDeiqamNBBgAAAAAAAC6QMoJcHl5aXX1+5EjxxUWtn6PYk1NFZ1O19DQ
onJjbW1daqWSkpKSEkNVVY3KgbW1/3pnZnV1JZ/PP3RoB3WfMAAAAAAAAEB3SP09wPxbt66NHj35ww89Bg3SVVJS0tMb7OQ029Z2
dFVVRUFBjpvbQjU1DQ0NLVdXz+zsjNramsrKirdv81xdF6ioqGlp6Tg5uVMN1dRUZWW9nDt36aBBekpKSoaGxosW+Sgo9Mh7mwAA
AAAAAGDAk/5rkMhf7wGeNXSoGfUe4Bcvnj55EsflNquqqru4zDczsyKEn5OTGR0dTr3rSF1d0919sbGxOfUU6I8+WkC9BonBUJ4+
3WX48FGqquplZSXx8fdevXop9WgBKHgNEgDAgITXIAEAgECPJMAA/RESYACAAQkJMAAACEh9CjQAAAAAAABAX4QEGAAAAAAAAOQC
EmCAgSYoKGjjxo3i6wQHB69bt0785pLU6VFiAugCOp2+devWyMjI2NhYGxsbmbcj8PXXX2/durX77QAAAABAh5AAA/Q/gYGBsbGx
a9asEayxt7ePjY0dNGiQDKMi/wksMDBQeOVnn30WGxt74MCBXghg7969/v7+IoumTZvm6Ojo4+Pj6OiYkZHR5S6k1Q4AAAAA9D68
VQigX2pqalq6dOlvv/1WWVnZqqhV/tlZkmwupk5JSYmTk9Phw4fr6uoIIYqKim5ubiUlJd0JSSqGDBlSXFxcVlbWR9oBAAAAgN6H
BBigX3r+/LmWltaaNWsOHjzYqigoKKisrCw4OJgQMmnSJD8/PxMTk4qKihs3bly6dInH41HVNDU1t2/fPnHiRDqdHh0dfeTIEapI
eHOBqVOnfvPNN8HBwXfu3GmvDqWgoOD9+/czZ84MDw+nNqyrq0tPT9fU1KQqbN68ed68eYSQmpqatLS0kJCQt2/fUkXBwcElJSUs
FmvKlCmZmZmCUFsFQKPRvLy8PD099fX1i4qKLl++HBERQQgJDAz84IMPCCGLFi0ihKxcubKgoIDafM+ePTNmzCCExMbGFhYWLlu2
TElJae3ata6ururq6tnZ2SdOnHj69KnIMDZs2CAIo1PttFfEYDACAgKcnZ3r6+sTEhJYLFZTU5Pkhx4AAAAAugxToAH6JT6ff/Lk
yblz5xobG7dXh8Viffvttzdv3pw3b97GjRvV1NQsLCwEpe7u7snJycuWLfvyyy/d3Nxmz57dXjsfffTRN998s3v3bir77VBkZKS7
uzu1PGfOnMjISOHSAwcOODo6Ojo6ent7l5aWBgUFKSgoCEf19OnTBQsWCKedrQJYvXq1h4fH7t27PTw8vv/+ez8/v5kzZxJCgoKC
EhISwsLCqPYF2S8h5Ouvvz558mRmZqajo+OyZcsIIb6+vs7Ozl999ZWnp2dCQsL+/fsNDQ3Fh9HZdtorWrdu3dixY9evX+/j49PY
2Ojg4CDJqAIAAABA9yEBBuivkpKSnj9/7ufn114FbW1tZWXlR48eNTQ0lJSUnDp1KisrS1AaFxd369at+vr69PT0hIQEe3t7kY0s
Xrx4w4YNW7dujY+PlzCwe/fuDRs2zNzcXE9Pb/z48bdv3xZZrbKy8vDhwyYmJmZmZoKVT548CQ8P53A47QXAYDCWL18eEhKSnp7e
0NCQkpJy7do1Dw8PCWOjMBiMRYsWnT59+uXLl2w2+9y5c3l5eYsXLxYTRmfbaa9IWVl5/vz51LGoqan54YcfMJsaAAAAoNdgCjRA
P3bixImTJ0+OGDFCZGlxcXFiYuLx48fv3bv37Nmz5OTkxsZGQalg4jEhhM1m6+nptW3B1dVVS0vr008/Fc6cO1RfXx8TE+Pu7l5b
W5uUlFRRUSFcam5u7ufnZ2dnp6mpSaPRCCEGBgbZ2dlUaW5urvgATE1NmUym4JFaNBqNRqMVFRVJHh4hxMjISElJKS0tTbAmLS1N
OA9vFUYX2mmvyMjIiMFgZGZmUitbWlo6NbYAAAAA0B24AgzQj2VkZMTExHz22WciS/l8/tatW3fv3t3Y2Lh27dqLFy8K53h8Pl+S
9tlstpjZ0e2JjIz86KOPPDw8Ws1/ptFo+/btKy8vp+YtOzk5cblc4SnQzc3N4gOg0+mEEB8fHycnJycnpxkzZjg6Oi5durRT4VGJ
d6s1wgPSKowutCO+C0kGHwAAAACkDgkwQP92+vRpOzs76uFPbfH5/JSUlDNnzqxZs6a8vLyzU4ULCgo2bNgwc+bM9t4t1J7U1NSa
mhplZeWEhATh9bq6ugYGBpcuXSopKeFyuZaWloqK4uahtA0gPz+/sbFx8uTJIutzudy2mWdbhYWFzc3Ntra2gjW2trb5+fkd75jE
7bRXVFRU1NzcLHiBsIKCgpWVVWf7BQAAAICuQQIM0L8VFhbeuHFD+P5VATs7u82bN1taWjIYDCsrK319/cLCws62n5+f7+/v34Uc
eOXKlfPnz29paRFeWVlZyWazPTw8mEzmsGHDtm7d2tkAGhoaLly44OPjM2vWLFVV1cGDB3t6eq5atYqq/O7dOysrKxaLJb7Npqam
sLAwX19fW1tbNTW11atXm5mZhYaGdmoHxbfTXlFjY+P169f9/PysrKw0NDQ+//xz4cnna9asuXXrVmfDAAAAAAAJ4R5ggH7vX//6
10cffcRgMFqtT09Pt7CwCAwMNDY2rqysDA8Pp15N1Fn5+fkbNmw4fPgwn88/cuRId0Llcrk7duzw9/dfsmRJRUVFWFiYmKdYtxfA
uXPnqqqqVq1atW3btoqKivj4+J9++omqGRoaun379uvXrzOZTOHXILV15swZOp3+3XffqampZWdnb9mypbi4uAt7JKad9opOnjzJ
YrGOHj1KvQYpLi6uC/0CAAAAQBfQdHV1ZR0DQJ9gbW2dl5fX+/1qaWn1fqcAAPKjqqpK1iEAAEBfgSnQAAAAAAAAIBeQAAMAAAAA
AIBcQAIMAAAAAAAAcgEJMAAAAAAAAMgFJMAAAAAAAAAgF5AAAwAAAAAAgFxAAgwAAAAAAAByAQkwAAAAAAAAyAUkwAAyVlVVJesQ
AAAGLPyOBQAAYUiAAQAAAAAAQC4gAQYAAAAAAAC5gAQYQPYwQw8AoCfgtysAALSCBBgAAAAAAADkAhJggD4BlykAAKQLv1cBAKAt
JMAAfQW+qwEASAt+owIAgEhIgAEAAAAAAEAuIAEG6ENwyQIAoPvwuxQAANqjKOsAAOB/UN/btLS0ZBwHAEA/hNQXAADEwxVggL4I
3+EAADoLvzkBAKBDuAIM0EfhUjAAgISQ+gIAgISQAAP0aYJvdciEAQBaQd4LAACd1ekEeMuWICUlRquV0dHXk5IeKigoOjjMGjFi
rKamVmNjY37+6wcPbr5/X04IEVO0bds+Go1+6NCO+vo6QYOmphYffuihp2dIo9Hq62ufPIlLTHzQrR0F6OfwPQ8AAAAAoJs6nQDv
3x9ICBk2zGbpUt+iooJ//StEULRw4SorqxGJiTHx8feNjEw+/ni1qanl2bMHa2qqxBSJiElRadGiNXV17OPH/8nj8aysRqira3Zj
HwEAAAAAAACk9xAsIyMTK6sRdXXsBw9uNTRwcnIynz9PYrFUJk6cLqZIZFNaWoOUlZl8Pr+pqbG+vvb586S4uLvSihMAAAAAAADk
k9TuATY0HEoIKSt7x+PxqDUlJYWEEEND46qqivaKRDZVVVVRW1ujo6MfELCruPhNbu6rJ0/iGho40goVAAAAAAAA5JDUX4PEFyzR
aDSJi/4Hl8v9+ecfUlIeczh1xsbmjo4frVz5dymHCQAAAAAAAHJGaleAi4vfEkL09AbT6XTqSq+BgREhpLj4jZii9lqrqqq4eTOU
EGJkZLJ69Rf6+oZMJgsXgQEAAAAAAKDLpHYFuKioICsrTVVVfcYMNyaTZW5ubW8/kcOp/+OPh2KKRDalpaWzcOEqY2NzBkO5ubmp
paWlpqYK2S8AAAAAAAB0hzTfA3zt2s8ODrPs7MZOnjyjsbExJyfzwYOb1HOexRRRAgJ2UQuvX6dHRl5mMlkLF65iMlW43ObCwoJ7
925IMU4AAAAAAACQQzRdXV1ZxwDQJ1hbW+fl5ck6CgAAAAAA6ClSfwgWAAAAAAAAQF+EBBgAAAAAAADkAhJgAAAAAAAAkAtIgAEA
AAAAAEAuSPMp0AAAAAAAPUFLS0vWIQCApKqqqmQdQruQAAMAAABAH4W8F6A/Evzk9sFMGAkwAAAAAPQ5SH0BBgDqB7lPpcG4BxgA
AAAA+hZkvwADSZ/6icYVYAAAAADoK/rUF2UAkJa+cykYV4ABAAAAoE9A9gswsPWFn3EkwAAAAAAAACAXkAADAAAAgOz1hUtDANDT
ZP6TjgQYAAAAAGRM5t+JAaDXyPbnHQkwAAAAAAAAyAUkwAAAAAAgS7j8CyBvZPhTjwQYAAAAAAAA5AISYAAAAAAAAJALSIABAAAA
QGYw/xlAPsnqZ1/6CfDSpWtdXOZLvVkAAAAAAKmg0+lbt26NjIyMjY21sbGRdTjQdwUFBW3cuFHWUUgEZ7WEFGUdAAAAAABAJwQG
Brq5uRFCWlpaSktLY2Ji/vWvfzU0NEjewrRp0xwdHX18fMrKyroWw969ewsLC0NCQrq2OXQWBrxD3T+r5UTPJsC+vpvy8l4bGhoP
HjykuroyIuKysbH5pEmODIZyevrzW7fC+Hy+l9cnlpa2hPCrqytTUpIePbpHCJ8Qoq6u6e6+2NjYvKam6unTBFdXz0OHdtTX1xFC
GAzlGTNmDx8+UlmZ+eZN7p0716qrK3t0RwAAAACg70hJSfH391dUVLS3t//222+ZTOahQ4ck33zIkCHFxcXIE6BDgYGBsg5BUjir
JdTjV4CHDx959erP79+Xubp6Ll++Li0t5cyZ71ksFW/vz3Nz7dPTn1+5cpYQQqPR9fUNFy70rqqqePnyGSFkwQLvmpqqY8f+yWAo
e3quEG5z7txlioqK588f43DqHRxcFi3y+fHHQ3w+v6f3BQAAAAD6Di6X+/Tp01u3bk2fPv3QoUM0Gs3Ly8vT01NfX7+oqOjy5csR
ERFUzeDg4JKSEhaLNWXKlKamJurmw9jY2MLCwmXLlonZkEajLV68mCrKyck5evRoampqYGDgBx98QAhZtGgRIWTlypUFBQXtBSmy
BULI5s2b582bRwipqalJS0sLCQl5+/YtFWpxcTGDwZg4cSKdTo+Ojj5y5AiPx2vVrIeHx7p16xYuXMjlcqk1u3btUlBQ2L59u5KS
0tq1a11dXdXV1bOzs0+cOPH06VNCyI8//hgVFXXp0iWq/s6dOzkczt69e9vGzGQyN27c6OTkVFdXl5CQoK2tXV1dTdVsb6zEhC3h
ccnMzHzz5o3IMRE54GKaFdbNcQ4KCiorKwsODqY2KS8vZ7FYo0aNUlJSioyMDA0NDQgImDBhQm1t7fnz569duya+UzFjK+HuEEJE
Ht89e/bMmDGDCJ3VIrcF0gsPwfrjj7iSkrdNTY1//vmEwWDcvXudw6l//748Nzdr8OAhgmp8Pu/du8InT+KsrOwIIdraukOHmkVF
/VpfX1tVVfHgwU1BTXV1zeHDR0ZGXqmqet/Y2HD/foSWlo6e3uCe3hEAAAAA6MtWr17t4eGxe/duDw+P77//3s/Pb+bMmYJSd3f3
p0+fLliwYN68eSdPnszMzHR0dKTyBDEbfvLJJ97e3seOHZs/f/7BgwddXFwIIUFBQQkJCWFhYY6Ojo6OjmKy3/ZaIIQcOHCA2tzb
27u0tDQoKEhBQUEQanJy8rJly7788ks3N7fZs2e3bfbevXsMBmPq1KnURw0NDQcHh5s3bxJCfH19nZ2dv/rqK09Pz4SEhP379xsa
GnZqJD/77DN7e/svvvjCx8envr7ewcFBkkFuL2wJj8uGDRvaGxORAy6+WYFujnMrLi4usbGxS5cu3bNnj5eX1+nTp6Ojo+fPnx8S
EuLv729qatphp+2NrYS7Q9o5vl9//XWrsxra0+MJMDVpmRDC5XKbm5ubm5uojy0tXEVFJUKIre3oTz7ZuGVLUGDgARcXTw0NLUKI
urpmU1OjYNvKyveCBrW0BhFC/P2/CQw8EBh4YNu2/crKTE1N7Z7eEQAAAADoUxQUFMaMGePm5hYfH89gMJYvXx4SEpKent7Q0JCS
knLt2jUPDw9B5SdPnoSHh3M4nFaNiNmQyWQuWbLk9OnTcXFxdXV1GRkZ33//facilKSFysrKw4cPm5iYmJmZUWvi4uJu3bpVX1+f
np6ekJBgb2/ftuWGhobo6GjBDrq4uNTU1CQmJjIYjEWLFp0+ffrly5dsNvvcuXN5eXmLFy/uVMxz5sw5depUVlZWTU3N8ePHS0tL
Oxyr9sLu2nFpOyatdNhsW10b51bi4uKioqI4HE5CQkJeXl5ycvL9+/c5HM6DBw9KS0tHjhwpvtP2xlby3en+8QUZPwRLVVXd03PF
1as/5+VlNTU1TZrkMGLEWEIIm13NYCirqKhSObC29iDBJtXVlXw+/9ChHRxOvcziBgAAAADZGTNmTGxsbEtLS1lZWWRk5NmzZ01N
TZlM5oEDB6gKNBqNRqMVFRUJNsnNzRXZlJgNTUxMlJWVnz9/3uU4xbRgbm7u5+dnZ2enqalJo9EIIQYGBtnZ2YQQwXRZQgibzdbT
0xPZeERExIkTJ3R0dCoqKtzd3W/fvs3j8YyMjJSUlNLS0gTV0tLS2ksjRTI0NFRSUsrMzKQ+8ni8rKwsaln8IIsMu1PHRcyYtNJh
s5K0KeE4CxPepLa2ttVHdXV18Z22N7aS7073jy/IOAFWVFSi0WiNjQ0tLS1DhphMmjSjtraGEFJZWV5YmO/quiAq6jcGQ9nJyV2w
SU1NVVbWy7lzl0ZH32Czq3R1B0+bNvPXX8+3tHBltx8AAAAA0Huoh2AJr6HT6YQQHx+fvLw8kZs0NzeLXC9mQyp16c6DZtprgUaj
7du3LzEx0c/Pr7y8nMfjRUdHC2bJSthjRkZGTk6Om5tbUlKSlZXVzp07BT226ktkg21rChPeRFBT/CCL7EXy4yJ+TDrbrCRtduHI
ttqkvYEVvyNtx1bC3SGdOb7Qnh6fAi1edfX7e/ciPT1XbN787axZ8zIy/hQU/frrv5lM5t//vm3Jkk/S01MJIS0tLVRRePjFiorS
pUt9AwJ2ffTRgtTUP5D9AgAAAMiz/Pz8xsbGyZMnS3FDqmjMmDFti7hcrvgEUnwLurq6BgYGly5dKikp4XK5lpaWiopduS4VERHh
7u7u4eGRmpr65s0bQkhhYWFzc7Otra2gjq2tbX5+PiGEzWYLLlESQoYOHSqyzeLi4ubmZsGLZOl0uqWlpfDudGqQJd9E/Ji0GnAJ
m5XWOHeKmE7bG1vJR0nM8QUJSf8MuHTptGD5zJmDguXCwvz9+//7GPHIyCvUwuPHDx4/ftC2nZqaqkuXzlDLFha2DQ2cxsa/Xu/W
1NR4717EvXuiH4wGAAAAAPKmoaHhwoULPj4+FRUVCQkJ6urqU6ZM0dDQ+Pnnn7u8YUNDw5UrV3x9fSsqKp49e2ZiYkI9oIgQ8u7d
O2traxaLRd286urqun379rlz51ZXV7dqXGQLlZWVbDbbw8Pj559/NjIy2rp1a9f2Oioq6u9///vcuXMFtxY3NTWFhYX5+vq+efPm
zZs3H3/8sZmZ2fbt2wkhqampH3300Z07dyoqKhYuXGhlZfXq1SuRAxIREbF27dq3b9++e/du5cqV+vr6XR5kyTcRPyatBlzCZqU1
zp0iptP2xlbyURJzfEFCMp4CLcaIEWN4PF5OTqaWlo6zs0da2jNZRwQAAAAAfde5c+eqqqpWrVq1bdu2ioqK+Pj4n376qZsbnjlz
pra2dv369bq6uq9fvz527Bi1PjQ0dPv27devX2cymStXrjQ0NMzKymKz2W0bF9kCl8vdsWOHv7//kiVLKioqwsLCjI2Nu7DLtbW1
v//+u4ODw/3794V7pNPp3333nZqaWnZ29pYtW4qLiwkhFy5cMDQ0PHHiBIfDefjw4aNHj9pr9vjx4ywW6+jRo/X19QkJCUlJSU1N
TR2OVXsk3ET8mLQa8IKCAkmaldY4d4r4TtsbW8kHtr3jCxKi6erqyjoG0VgslY8+Wjhs2PDm5qaMjD9jYiK5XNF3bgBIhbW1dYf3
XQAAAIB0Ua/k7e+OHz9+6tSpZ89kcMHm4MGD7969E/k6X6mg0Wjnz5+/ceOG4AXCIC0Y26qqqt7vtO9eAeZw6n/77bysowAAAAAA
6MBnn30mk37HjRs3bty4v/3tb1Jv1tTUNCYmhk6nL1u2TF9fPyYmRrpdyC2Mrcz13QQYAAAAAADac/XqVRUVlVOnTuXk5Ei35Zcv
Xzo4OJw9e5bFYr1+/drf3//du3fS7UJuYWxlru9OgQboZZgCDQAA0PsGxhRoAOgCmUyBlvFrkAAAAAAAAAB6BxJgAAAAAAAAkAtI
gAEAAAAAAEAuIAEGAAAAAAAAuYAEGAAAAAAAAOQCEmAAAAAAAACQC0iAAQAAAAAAQC4gAQYAAAAAuRMcHLxu3TqRRUFBQRs3buxs
g13bqmt6sy+AAUZR1gEAAAAAAHRCYGCgm5ub8JovvvgiNTVVVvGIFxgYOHjwYH9//55ofMqUKd7e3paWltXV1aGhoaGhoT3RC8BA
ggQYAAAAAPqZlJSUHkopCSGBgYG9tlXXCPpauHDh4cOH8/LynJyctm/fnpOTk5yc3GthAPRHSIABAAAAoN/78ccfo6KiLl26RH3c
uXMnh8PZu3cvIWTSpEl+fn4mJiYVFRU3bty4dOkSj8cjhGhqam7fvn3ixIl0Oj06OvrIkSPU+qCgoLKysuDgYEJIcHBweXk5i8Ua
NWqUkpJSZGRkaGhoQEDAhAkTamtrz58/f+3aNapH4a3aExwcXFxczGAw2na6efPmefPmEUJqamrS0tJCQkLevn0r2KqkpITFYk2Z
MiUzM3PDhg2Cvr788kuqTkxMzLZt23R1dXtgaAEGFNwDDAAAAAADFovF+vbbb2/evDlv3ryNGzeqqalZWFhQRe7u7snJycuWLfvy
yy/d3Nxmz54tsgUXF5fY2NilS5fu2bPHy8vr9OnT0dHR8+fPDwkJ8ff3NzU17VQ87XV64MABR0dHR0dHb2/v0tLSoKAgBQUF4a2e
Pn26YMGCDRs2tG2TTqdv2rSppKTk0aNHnQoGQA4hAQYAAACAfmbMmDGx/3H69GkxNbW1tZWVlR89etTQ0FBSUnLq1KmsrCyqKC4u
7tatW/X19enp6QkJCfb29iJbiIuLi4qK4nA4CQkJeXl5ycnJ9+/f53A4Dx48KC0tHTlyZKci77DTysrKw4cPm5iYmJmZCVY+efIk
PDycw+GIbDMwMNDOzs7f37+2trZTwQDIIUyBBgAAAIB+RvJ7gIuLixMTE48fP37v3r1nz54lJyc3NjZSRYI5xoQQNputp6cnsgXh
arW1ta0+qqurdyry9jo1Nzf38/Ozs7PT1NSk0WiEEAMDg+zsbKo0Nze3vQaHDRvm6urq5eVVVlbWqUgA5BOuAAMAAADAQEPlkIQQ
Pp+/devW3bt3NzY2rl279uLFi4Irq3w+X5KmWlWTcCsJWxNEu2/fvvLycj8/v5kzZzo5OXG5XOEp0M3Nze01SN33i+wXQEK9nQD7
+m4aOXJ81+r4+PiPGTO5Z+ICAAAAgH6MzWYLX4wdOnSoYJnP56ekpJw5c2bNmjXl5eUeHh6yCFAcXV1dAwODS5culZSUcLlcS0tL
RUVJ52kmJSU5Ojq2tLT0aIQAA0ZfuQLc2eQWyTAAAAAACKSmprq4uJiYmKiqqnp7e1tZWVHr7ezsNm/ebGlpyWAwrKys9PX1CwsL
ZRtqW5WVlWw228PDg8lkDhs2bOvWrZJv6+TkFBsbq6am1nPhAQwkPXIPsKGh8dSpM42NzZSVmdXVlQUFOQkJMZWV5WI2+de/QjrV
RXv1ly5dO2zY8Js3Q1NSHlNr9PUNfX3/r7m5af/+3ns5GwAAAAD0pgsXLhgaGp44cYLD4Tx8+FDwPOT09HQLC4vAwEBjY+PKysrw
8PDw8HDZhtoWl8vdsWOHv7//kiVLKioqwsLCjI2NZR0UwMBEk/rrwiwsbBYt8klOjn/2LLG6ulJdXdPEZJixsXlExGVCiK/vpsTE
31+8EPeG7vbq+Pj4p6Q8FmS2Ii1dulZNTYPH4/34418vYXNzWzh0qLm2tg4SYBDP2to6Ly9P1lEAAADIFy0tLVmHAACyUVVV1fud
Sv0KMM3NbeHz50nR0X/9aa2ysryysvz58yRBDR0dvRUrPjU0NK6trblz59fc3FekTXKrq2vg7f25gYFRZWX5nTu/vn2b16obMcnw
69fpI0eOMzIyKSoqYDCUR4wY+/vvt5yd51ClDIbyjBmzhw8fqazMfPMm986da9XVldIeBAAAAAAAAOhzpHwPsK6uvqbmoD//FHeB
d9SoCTExNw8f3vnnn8nz5i0nhNa2zrhxH8TG3j5yZE9aWsqSJb4qKp24q4HP5z17ljh+/FRCyMiR49+8ya2urhKUzp27bNAg3fPn
jx09+o+KitJFi3wEDwkEAAAAAACAAUzKCTCVqbLZ1dTHkSPHBwYeoP4T1Pnjj4dFRQXNzc3JyY9UVdXU1TXatvP0aUJ+fnZjY0NC
QgybXT18eOfeMJ6SkmhtPZLFUhk37oPk5HjBenV1zeHDR0ZGXqmqet/Y2HD/foSWlo6e3uCu7CoAAAAAAAD0K1KeAl1fX0sIUVfX
rKmpIoS8eJH84kWyubnVsmXrBHXq6mqpBS63mRCiqKjUtp2qqgqh5ffq6pqdCqOurjY7O2P27EVKSoycnExLS1tqvZbWIEKIv/83
wpU1NbVLS4s71T4AAAAAAAD0O1JOgMvLS6ur348cOa6wML877VCZqmA5K+tlZ1tITn7k7f35vXsRhPz3bePV1ZV8Pv/QoR0cTn13
wgMAAAAAAIB+R+rvAebfvn1t9OjJM2fOGTRIT0FBUVVV3cTEorOtjB37gYmJBYOhPHmyk4aGVmbmi8628OZNblDQ5sePHwivrKmp
ysp6OXfu0kGD9JSUlAwNjRct8lFQ6JF3QQEAAAAAAECfIv3cLzs749///mHatJmrVn2hrMxks6uKi9929jW/KSmJTk6z9fUNKysr
Ll8+Q82slorw8IvTp7ssXeqrqqpeVlYSH3+vpYUrrcYBAAAAAACgz5L+e4AB+im8BxgAAKD34T3AAHJLJu8BlvoUaAAAAAAAScnk
GzAAyJysfvaRAAMAAAAAAIBcQAIMAAAAAAAAcgEJMAAAAADIEmZBA8gbGf7UIwEGAAAAAAAAuYAEGAAAAABkDBeBAeSHbH/ekQAD
AAAAgOwhBwaQBzL/SUcCDAAAAAAAAHIBCTAAAAAA9AkyvzQEAD2qL/yMK8o6AAAAAACAv1Dfj7W0tGQcBwBIVV9IfSm4AgwAAAAA
fUvf+a4MAN3Xp36icQUYAAAAAPocXAoGGAD6VOpLQQIMAAAAAH2U4NszMmGAfqQP5r0CSIABAAAAoK/ry9+nAaAfwT3AAAAAAAAA
IBeQAAMAAAAAAIBcQAIMAAAAAAAAcgEJMAAAAAAAAMgFJMAAAAAAAAAgF5AAAwAAAAAAgFzo7QTYx8d/zJjJvdwpAAAAAAAAQC8l
wMh7AQAAAAAAQLYUpd7i0qVrhw0bLviYk5N56dLpf/0rpDttfvrp/7tx42JhYX7bXpqaGsvKSmJibhYUZAuv53K5tbU1b9/m/fHH
w+LiN2Ji605gAAAAAAAA0F9IPwEmhPzxx8O7d69LqzUdHT0mk1lUVCCyF2Vl5owZsxcvXnPixHd1dbWC9QoKClpag8aMmbx69fpf
f/05M/NFT8QGAAAAAAAA/UWPJMBt+fj4p6Q8Tkl5LLySwVCeMWP28OEjlZWZb97k3rlzrbq6su22lpZ2r1+n8/l8kS03NjbExERM
mDB1yBDTV69eCta3tLRUVJTduxfBZKq4ui7IzHxJiOgWAAAAAAAAQB7I8inQc+cuGzRI9/z5Y0eP/qOionTRIh8ajda2mpXViNev
07vcy8uXT9XVNXV09LoRKQAAAAAAAPR7PXIFeOLE6RMnTqeWw8L+9erVi7Z11NU1hw8fGRKyu7a2hhBy/37Epk179PQGl5YWC1dj
MlmGhsY5OZnt9aWszHR0dGtqaiosbD1HmlJTU0UIYbFUJI8NoI8wMzOTdQgAA0deXp6sQwAAAAAZk9k9wFpagwgh/v7fCK/U1NRu
lQBbWNi8fZvb1NTYtgUqlW1ubiorKwkLO1dXxxbZkYaGFiGEw6mTPDYAmaNS37fs97IOBGDgoH6skAYDAADIs166B7it6upKPp9/
6NAODqdeTDUrK7usrJciiyRMZe3sxrHZ1RUV5V0MFKB3IfUF6CHUjxXSYAAAAHkms3uAa2qqsrJezp27dNAgPSUlJUND40WLfBQU
/ichp9Ppw4YNz8rqyg3AdLqCjo6es/OcUaMmREX9iidgQb9gZmb2lv0e2S9Az6F+xHB/AQAAgHyS2RVgQkh4+MXp012WLvVVVVUv
KyuJj7/X0sIVrmBsbM5mV1dXdy4ZoKZGt7Rwa2tr3rzJ++mnI4L3AAP0ZVT2K+soAOQClQPjOjAAAIC8oenq6so6hnbNnDm3paXl
wYObsg4E5IK1tbUMvw0j+wXofUPVByEHBgAAkCuyfA1Sh6ys7F6/TpN1FAAAAAAAADAQ9OkrwAC9SYZXgHH5F0BWcBEYAABArvTp
K8AAAAAAAAAA0oIEGEDGcPkXQIbwRGgAAAC5ggQYAAAAAAAA5AISYAAAAAAAAJALSIABAAAAAABALiABBgAAAAAAALmABBgAAAAA
AADkAhJgAAAAAAAAkAtIgAEAAAAAAEAuIAEGAAAAAAAAuYAEGAAAAAAAAOQCEmAAAAAAAACQC0iAAQAAAAAAQC4gAQYAAAAAAAC5
gAQYAAAAAAAA5AISYAAAAAAAAJALSIABAAAAAABALiABBgAAAAAAALmABPi/fH03jRw5XtZRiKaqqh4YeEBFRVXWgQAAAAAAAPRX
SIABAAAAAABALihKvcWlS9dWVJTevXtdsMbHxz87O+Phwyip9yUTFhY2M2fO1dDQys/Pjor6tbq6UmS1QYP0nJ09hg41o9MViory
Hz269+ZNLiHkb38LePLkUWrqH70bNQAAAAAAgLyTfgI8sNHpCgsWeMfG3nn6NMHExGLEiLEJCfdF1lyyxLewMO/s2eDGxoYhQ0yn
TnW+fPlsL0cLAAAAAAAAAr2dAKuqqs2aNd/c3IrPJzk5GdHR4RxOPSHE13dTXt5rQ0PjwYOHVFdXRkRcNjY2nzTJkcFQTk9/futW
GJ/PZzCUZ8yYPXz4SGVl5ps3uXfuXKuurrSxsXdymn3ixF5CyIwZbtOmzTp27J9VVRVGRibLlvkFB3/D4/HEdPr6dfrQoWaDBw+9
eTNUEKSSkpKn50o6nX7t2r+bm5uE46fT6QoKCu/eFXG5zTk5GTk5Ge3sprq2ts61az+x2dWEkNzcV7m5rwghnp4rBw8eOmfOkjlz
lhQVFfzrXyGEkPXrv46IuJSbm0UIUVfXXL/+64MHv25o4Kipabi7LzYxGVZTU5WcHE+1PG7cB+PHTzt9+gD1cdAg3XXrvjx27J/t
XYgGAAAAAAAASm/fA7xggTeTyTxz5uC5c4c0NLTmzVsmKBo+fOTdu9cPH95VVFSwfPk6HR39M2e+P3fukJXVCBsbe0LI3LnLBg3S
PX/+2NGj/6ioKF20yIdGo+XlZWlr62poaBFCzMys3r8vNze3IoSYm1sXFGTzeDzxnY4ePenBg1vff/91WloKtUZVVX3lyr+z2TVX
rpxrlf0SQrjc5ry81/PmLdPVNRCzm3V1te/fl0+f7jpkiKmi4n//yvDbb+dLSt5GRFwOCtpMZb/ix4rLbT5+/J+hoT/a20+gVr54
8VRDQ2voUDPq49ixH+TmvkL2CwAAAAAA0KEeSYAnTpweGHhA8J+RkQm1XktLx8TE4vbtX2tra2pqqqKifrOwsFVT06BK//gjrqTk
bVNT459/PmEwGHfvXudw6t+/L8/NzRo8eIi6uubw4SMjI69UVb1vbGy4fz9CS0tHT29wQwOnpOStubm1sjJTV9cgIeG+ubk1IcTM
zIq6piq+0ydPHr19m8fn86iPuroGq1evz8hIvX37qmClMGfnOU1Njffu3Vix4tMhQ0wJIXQ6fevW79rkw/xffjnOZlfPm7ds8+Zv
V69eP3z4qE6Noba2jrGx+Z071+rqaisrK2JiblLrm5oaX7xIHjfuA0KIgoKivf3EZ88ed6plAAAAAAAA+dQjU6D/+ONhq4dgUQsa
Glo8Hq+6+j31sbKynFpZW1tDCKmvr6PWc7nc5uZmwdXXlhauoqKSltYgQoi//zfCHWlqapeWFuflZZmZWXE4dYWFea9fpzs7z2Ew
lIcONb1z51qHnba6djp69MT6+jrBfONWVFRUJ092PH58b1VVRXNzk5fXJ+HhF1pauHV1teXl71pVZrOr79z5lRCirMwcOXL8woWr
Llw4kZ+fLeEYqqtrNTc31dXV/ifsCkFRcnL83/628e7d6xYWtjxeS1ZWmoRtAgAAAAAAyLNevQe4pqaKTqdraGjV1FQRQrS1damV
kmxbXV3J5/MPHdpB3b4rLDc3a/785Q0N9bm5WXV17NramokTp3M49VRS2qlOY2JuDhs2fNkyv8uXzzQ0cFqVKigo0mh0JSUlQkhW
Vtpvv51fuHAVh1P3+++3xUTe2NiQnPxo8mTHoUPN8/Oz+Xx+qwpcbrOiohK1rKqqRi2w2VVKSgxVVTUqB9bWHiSoX17+rqgo395+
4vDho1JTn/B4LWJ6BwAAAAAAAEqv3gNcVVVRUJDj5rZQTU1DQ0PL1dUzOzuDuhLboZqaqqysl3PnLh00SE9JScnQ0HjRIh8FBUVC
yJs3udRVVupBU7m5WVOmOFHLne2Ux+Ndv/5LWVnJypWfCXJRATa7uqAgZ86cJfr6RkpKSs3NTdXVlZqag9qmoOrqml5efzMzs2Iy
WcrKzLFjP9DUHFRUlE8Iqa2t0dc3pNP/O/Lv3hXZ209iMlmamtoffjiHWllZWfH2bZ6r6wIVFTUtLR0nJ3fh9pOT46dM+XDoUNOU
FMx/BgAAAAAAkEhvPwX611//7eIy39f3/wjh5+RkRkeHS75tePjF6dNdli71VVVVLysriY+/19LCJYS0tHDfvMnV1zcqLS0hhOTm
vpo0aTp1A3AXOuXz+Tdvhs6aNW/lys8vXDhBPcZZ4OrVfzk6ui1Z8omKilpFRWlycjybXb1w4aqGhobs7HRBNTa7+unThA8++NDQ
0JhGo1VWlkdEXKZCSkiI8fDwmjDBoaTkLfUcrPv3I+fOXbp+/TdVVRXJyfHUQ7yosN3dF//979uop0ALbqUmhGRmvnB19czLy6Zm
dAMAAAAAAECHaLq6urKOATpNQUHR3/+bO3euCR5eDd1nbW2dl5fX+/2amZm9Zb/v/X4BgDJUfZBMfvYBAACg9/X2a5BAGmgTJkzj
cpszMv6UdSQAAAAAAAD9Rm9PgYbu+/LLfzY2ciIiLuPxVwAAAAAAAJJDAtz/7Nv3/2QdAgAAAAAAQP+DKdAAAAAAAAAgF5AAAwAA
AAAAgFxAAgwAAAAAAAByAQkwAAAAAAAAyAUkwAAAAAAAACAXkAADAAAAAACAXEACDAAAAAAAAHIBCTAAAAAAAADIBSTAAAAAAAAA
IBeQAAMAAAAAAIBcQAIMAAAAAAAAcgEJMAAAAAAAAMgFJMAAAAAAAAAgF5AAAwAAAAAAgFxAAgwAAAAAAAByAQkwAMD/b+/eg6u6
DwOP/4QlIZCEhBCYhwARQALzim3ANhhDHL/ABj9qO3acZpKWnUl3dzKd7I7bZmab7G6n6XTaXSeTSdodp029btzEie0kNimOASc2
hvCwjQGDxeuCBMjIQm+BHujuH9fVUvOGK/T4fT7jyXDPPed3flcZNPfL79x7AACIggAGAAAgCv0ygFeu/NrMmTf29iwAAADoT/pl
AAMAAMClSnMAz5u36I/+6M9Oe3jr17/+N5MmlaUeZmZm/cmf/NXkydMuddg/+IM/nj17XtpmCQAAQHzSHMCJxJ7hw0cUFAxPPZw4
cWpNTXVp6dTUw/HjSzMyMiorD6T3pAAAAHBBmekdrqamurm5sbR06rZtmzIyBk2Y8KlVq56/5ZbPpJ4tLZ16+PCh9va27OzBixcv
LS+fOXhwTmXlgdWrX2hoqAshPProH06ZMj2EZEND3bvvblq/fk0IyQce+MLo0SX33fe5++773JEjh374w++EEEaMGPnEE18ZM2Z8
c3Pj6tUvHjhQEUI417ArV35t795dJSWlo0eXrFr1/Pvvv5veVw0AAEDfl+YADiEkEntTATxmTElLS1NFxc577300J2fIyZMnSkun
7t27K4SwfPnjmZmZzz77vRMnWm+99c6HH/7SP/zDU8lk8ic/+UEIISNj0KhRYx566Pfr62t37nznpZeeLSoq3rJl/Xvvbe4+y6xZ
c1944ZmamqPz5y9eseLz3/72fw8hea5hQwhz5sz/2c/+6fDhQ8lkV9pfMgAAAH1f+r8EK5HYU1o6JYRQWjolkdjb1XXq8OHExIlT
Bg/OGT163IEDFfn5BeXlM1955Sf19cfb2k6uXftyYeGIkSNHd4+QTHZ9+OHhLVvenDp1xrnOsnnzG0eOHOro6Ni6dX1ubl5+/rDz
D7tly/qqqoT6BQAAiFZPrADvyc3NLy6+duLEqe+8syGEkEjsKy2dmkx2dXR0HDlSOW7chBDCV7/656cfVVAw/Nixo9Onz1mw4Pai
opFZWdkhhKqqxLnO0tLSnPpDZ2dHCCEzMysvL/9cw4YQUtdCAwAAEK30B3BjY/3x4zVTpkwvKZn40kv/N4Rw8ODeFSs+n0x2HTq0
v6vrVENDXTKZfOqpb5w40Xr6gbm5+Q888MTPfvZMIrGnvb19/vxbr7vu+tRTqcuYz+9cwwIAAEDoofsAJxJ7brpp8fHjH6VatLq6
Kjc3b/r0OYnEnhBCY2P9nj07ly9/rKhoZFZW1pgx4x9++EvXXJOZmZmVkZHR1nby1KlT48ZNmD9/cfeAzc2No0aNGTTofLM917A9
8QIBAADod3qkDw8c2HPDDQt27Hg79TCZTFZW7p86dUYqgEMIv/jFc4sW3fnYYytzc/NraqrfemvNqVOdDQ3H16x55YEHnsjJGVpd
XbV79/aSkomp/TdsWHfvvY/OnXtrdXVV6lugz+qsw/bECwQAAKDfySguLu7tOUCfUFZWlkgkrv55S0tLq5qOX/3zAikl+UW98ncf
ALj6euQSaAAAAOhrBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQwAAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERB
AAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQwAAAAURDAAAAAREEAA/SU3CG5f/zlPx2SM6S3JwIAQAj9PYBX
rvzazJk39vYsAC7NY/d9cWbZnN6eBQBAdDJ7ewIA53Rt8Zh5s28Zd21JdtbgxpaGw9WHtmz/XX1jXa9M5vMrvvTurq3v79l+5UP9
y8vPXPkgAABcqvQHcFHRyNtvv7ekpHTQoGuOHDm4fv2aysoDaT8LMOCVlnxq+e2/t233229t/U1jc0Nebv64a8fPn73g1Tdf6e2p
AQDQL6U/gD/3uZWHDyd+8IP/3dZ2cty4iQsW3P7jH/8g7WcBBrzbb7l75573frtpTephfWNdfWPdzj3vpR5mZWUvvOG2yRPLsrMG
H/mwat3GVxubG0IIT9z/Bwcq944ZNe7a4jEtJ5pf3/jrg4cPXPiQqn1jR427dsTo19b/atrkmZPGTw4hNDY37KjYtmnbWyGEZUvu
HzVi9F233nvXrfdW1xz5l5efOddouUPz7liwdNzo8c0tTdt2bz3rS3vsvi/uqNi2o2Jb6uyVRw+OLh4zcsS1Tc2Nr775yrhrS66f
MS87K7viwK41b61OJpMhhPvveOTMWYUQ8obm37Fw6dhrS5pbmt7b/faSm+/8++e+feLkiXNNDwAgZmkO4Nzc/OHDR7zwwj81NTWE
EA4cqDhwoCL1VHb24MWLl5aXzxw8OKey8sDq1S80NNSFEB599A+nTJkeQrKhoe7ddzetX78mhGQIYeXKr+3du6ukpHT06JJVq57f
t2/3bbfdU1Z23eDBQ3bt2vbaa7/o6GgPIYwYMfKJJ74yZsz45ubG1atf7D4d0K8VFRYPyyvYtfec1xvfvei+zGsyn//Vj06ePHHT
pxfed/tDz/3yh6lWnD5l5ivrXvyo7qMbZsy7+7bl/+e571zwkBlTZ7+89oXqmiPJZPKDA7tCCBkZGSOLRi1b8kBjU/3u/e+vev3n
hcOGn34J9LlGW7bk/hMnWn/407/LyspetuSBi3mxUyaWvbz2xbrG40tuuuOhux+rOPD+sy/9IGfwkEeWfWHKxMSexO4Qws9fe/7M
WYUQln3m/qaWpn/86d9lZ2UvXXz/xfx8AACileYvwWppaT5+/KNFi+4aN25iZua/q+vlyx8vKip+9tnvffe7f1Fbe+zhh7+UkZER
QvjJT37wl3/5X7/1rT/56U//afbsuTNmfLr7kDlz5r/++q/+9m//2/vvv7tixeOjR4/78Y9/8L3v/WV1ddXEiVNS+8yaNXfdulXf
/vY3t2/fumLF50PISO8rAnrF0JyhIYTm1ubUw2mTZ/zxl/809V8IIS83f8rEsl+vX9XYVN/e0fbmlnWFwwpHFBandn5n55bqmqOd
nR3bdr09NGdo3tD8Cx6ybdfWo8cOn96HyWTyWO2H7+7a+qkJU8+c3rlGK8gvHHft+LUbVreebG1oqn9zy+sX82LffX/Lsdrqjo72
XXt3ZGdl/+Z3a062naxvrKs8khg1YvTpe35iVoXDho8dVfL6xldPnGxtaKpfv/X180/von/8AAADU9ovgU7+8z9/f8GC21eseLyg
YPjRo1UbN77+wQfb8/MLystnfuc7/6O5uTGEsHbty1/72v8cOXL0sWNHPz4s2fXhh4e3bHlz6tQZO3e+k9q4Zcv6qqpECGHYsMKp
U2d8//t/VVf3UQjhnXc2dp9v8+Y3jhw5FELYunX94sX35OcPSy0+A/1a68nWEELe0LymlsYQwu59O3fv2zlhbOlDdz8WQijIKwwh
/IfP/efTDxmWV/BRXU0IofVkS2rLqVOdIYTUP8ad/5Cm5sbujVNLp82bfcvwgqKszKwQwpFjVWdO71yjtXe0d3R2pCYfQmhouqjv
6+re/9Spzo7Ojo7Oju6H3f+SeNZZ5Q3N7+hoP3HyxL+drv7800u9WACAaKX/M8BNTQ2rV78YQhg8OGfmzBsfeuiLP/rR33V1dYUQ
vvrVPz99z4KC4ceOHZ0+fc6CBbcXFY3MysoOIaSKNyV1jXQIobCwKJnsqqurPfN0LS0fLxB1dnaEEDIzs9L+ioCr73j9R43NDdMm
zzxac+TMZxubG5LJ5N8/952TbScucsDzH5IMH6/9Dh2Su2zJ/S+vfbHyaKK9o/36GfPKJ03/eJ/T1ofPNVpBfmFWZtbQnKGppi3I
L7zI6Z3fuWbV3NqUlZU9JGdIqoG7T3cZPx8AgBj04G2Q2tpObt26/qabbispmbR9+5ZkMvnUU984caL19H1yc/MfeOCJn/3smURi
T3t7+/z5t1533fVnDlVffzwjY9Dw4SNSK8BADNZuWH3f7Q91nurY8cG2xpaGwdk5JaMnpJ5qamncX7nn7kX3/nbT2qbWphGFxfPn
3LLq9Z+fOnXqXKNd5CGZ12RmZGS0d7Sd6jo1ZtS4G2bMa/m3y7BbWpuLh48aNGhQV1fXuUZraKo/cqzqM7fctXbDq9lZ2QtvXJKW
H8W5ZlXfWHf02OElN9/1+sZfn366y/j5AADEIM0BnJ9fsHTp723a9EZ1dVUymbzuuusLCoqOHDnY2Fi/Z8/O5csfe+21XzY11RcX
j1648LMvvvhsZmZWRkZGW9vJU6dOjRs3Yf78xc2nXYjYrbGxfu/e9++773O/+tVPm5oapk+f09zctHfv++mdPNCnJKr2P7/qn+fP
ueXRe38/Ozu7uaXpw4+qu++gu/q3L9/06YUP3PXo0CG5tXU1m9/bcMG6u5hDGpsb3ti8buniFYMH5xyr/XBv4oMxo8alntqy/Xd3
3Lr009NvPFZb/S8vP3Ou0Vat+/kdC5d++eGvpL4FevTIMVf+ozjPrFa9/vPPLrznyw9/pbm1aWfFe6NHjjl1quvyfj4AAANeRnFx
mr8WZcqU6fPmLRozZnxGRkZd3UebN7+5ffuWEEJ29uBFi+4sL5+Vm5tfU1P91ltrKip2hhBuumnJzTcvzskZWl1ddfjwoZKSiT/8
4XdCCCtXfm3jxt/s2PHxTUQGD85ZsmTZ1KnXZWVl79697bXXftnR0X76PpmZmU8++VfdnxOGS1VWVpZIJK7+eUtLS6uajl/98zLw
TCqZfM/i5d//56d6eyL9TEl+Ua/83QcArr70BzD0UwKY/qhs0vSuZNehwwcK8gvvuW354Q+r1m5Y3duT6mcEMADEowc/AwxATzt0
JHH7LXfduXBpR0fH3oMfXOSNlwAA4iSAAfqxk20nVr3+896eBQBA/zCotycAAAAAV4MABgAAIAoCGAAAgCgIYAAAAKIggAEAAIiC
AAYAACAKAhgAAIAoCGAAAACiIIABAACIggAGAAAgCgIYAACAKAhgAAAAoiCAAQAAiIIABgAAIAoCGAAAgCgIYAAAAKIggAEAAIiC
AAYAACAKfS6Ac3Pzv/71vxk6NLe3JwIAAMCA0ucCGAAAAHpCZk8MOnbshIULP1tSMikrK+v48Y927nx78+Y3Ozs7euJcAAAAcDHS
H8Cf+tS0Rx750pYt69esebmpqb6wcMSMGTdMnXrdrl3b0n4uAAAAuEhpD+CMpUsf2rZt05o1v0w9rqmpfv31Vak/5+bm3XHH/ZMm
TU0mw/79u1977RcnTrSGEPLyhi1b9siECZ9qbKzfuvWt7rGyswcvXry0vHzm4ME5lZUHVq9+oaGhLt0TBgAAIApp/gxwcfGogoKi
7du3nvXZBx/8/ZycnKef/l//+I9PDRtWuGLF493bOzs7vv/9bz3//D/Mnj23e//lyx8vKip+9tnvffe7f1Fbe+zhh7+UkZGR3gkD
AAAQiTQH8NCheSGEpqaGM58qLBwxYcLkf/3XF5ubGxsb61999aXJk6fn5Q0bPnzE+PGTVq9+oaWlua6udt26j5eL8/MLystnvvLK
T+rrj7e1nVy79uXCwhEjR45O74QBAACIRJovgW5tbQ4h5OcXNDbWf+KpYcMKu7q6GhqOpx7W1X2U2piZmdXR0d7S0vxv22tTfygs
LAohfPWrf376IAUFw48dO5reOQMAABCDNAfwRx8da2g4PnPmDYcPH/zEU42N9YMGDRo2rDDVxsOHF6c2ZmVlZWVl5+bmpRp4+PCi
1P4NDXXJZPKpp76R+pwwAAAAXIm03wc4+atfvTBnzk2f+cy9RUXFWVlZI0eOXrJk6fTpc+rraw8d2n/PPQ/l5Q0bNqzwrrse2Ldv
d3NzY11dbVVV4q67Hhw6NK+wcMSSJctSAzU21u/Zs3P58seKikZmZWWNGTP+4Ye/dM01PXLfJgAAAAa8jOLi4rQPOnbshIUL7ygp
KU3dB3jHjre3bHmzs7MjNzf/zjvvLy2dGkJy//4PXnvtF62tLSGE/PyCZcseGT9+UupboO+++8GnnvpGa2tLdvbgRYvuLC+flZub
X1NT/dZbayoqdqZ9tpBSVlaWSCSu/nlLS0urmo5f/fMCKSX5Rb3ydx8AuPp6JIChPxLAECcBDADxSPsl0AAAANAXCWAAAACiIIAB
AACIggAGAAAgCgIYAACAKAhgAAAAoiCAAQAAiIIABgAAIAoCGAAAgCgIYAAAAKIggAEAAIiCAAYAACAKAhgAAIAoCGAAAACiIIAB
AACIggAGAAAgCgIYAACAKAhgAAAAoiCAAQAAiIIABgAAIAoCGAAAgCgIYAAAAKIggAEAAIiCAAYAACAKAhgAAIAoCGAAAACiIIAB
AACIggAGAAAgCgIYAACAKAhgAAAAoiCAAQAAiIIABgAAIAoCGAAAgCgIYAAAAKIggAEAAIiCAAYAACAKAhgAAIAoCGAAAACiIIAB
AACIggAGAAAgCgIYAACAKAhgAAAAoiCAAQAAiIIABgAAIAoCGAAAgCgIYAAAAKIggAEAAIiCAAYAACAKAhh6WSKRKMkv6u1ZQKRK
8osSiURvzwIAuEoEMAAAAFEQwND7LAJDr7D8CwCxEcAAAABEQQBDn2ARGK4yy78AECEBDH2FBoarRv0CQJwye3sCwP+XSCRKS0tD
CFVNx3t7LjAwpf6ZSf0CQJwEMPQtqfflMhjSTvoCAAIY+qLTMxhIC+kLAAhg6Lu8XwcAgDTyJVgAAABEQQADAAAQBQEMAABAFAQw
AAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQw
AAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQw
AAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQw
AAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQw
AAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQw
AAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQw
AAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQw
AAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQBQEMAABAFAQw
AAAAURDAAAAAREEAAwAAEAUBDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABRyFiwYEFvzwEAAAB6XEZxcXFvzwH6
hLKyskQi0duzAAAAeopLoAEAAIiCAAYAACAKAhgAAIAoCGAAAACiIIABAACIggAGAAAgCgIYAACAKAhgAAAAoiCAAQAAiIIABgAA
IAoCGAAAgCgIYAAAAKIggAEAAIiCAAYAACAKmb09AeBilZaW9vYUAADgfBKJRG9P4XwEMPR13d1bX1/fm/MAAIAL6X7v2jdLWABD
35X69aF7AQDoL7rfu6beyva1DPYZYOijSktL6+vr1S8AAP1R6q1sX/sQnwCGvihVv709CwAAuCJ9rYEFMPQ56hcAgAGjTzWwAIa+
Rf0CADDA9J0GFsDQh6hfAAAGpD7SwAIY+gr1CwDAANYXGlgAQ5+gfgEAGPB6vYEFMAAAAFEQwND7LP8CABCJ3l0EFsAAAABEQQAD
AAAQBQEMvcz1zwAARKUXr4IWwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMAAAABEQQADAAAQhczengAAAABX25E/e+RcT4391vNX
cyZXkwAGAACIxXm698x9Bl4JC2AYaFaufPLMjU8//ddXfyYAAPQdF5O+Zz1kIGWwAIYB4qzde+azShgAIEKXUb+nHztgGlgAQ7/3
ifQ9a+J275P6gwwGAIjHldRv9wgDo4EziouLe3sO0CeUlZUlEomrf97S0tL6+vrLPvxSl3YtBQMAROXK67dbGhu4sLCwV957C2D4
WL8L4CtJWRkMABCDNNZvSroauLcC2H2AoV+6woLtPur8nxwGAKD/Snv99tCYV5MAhn7sStZvrf0CAAxgPVeq/bqBBTD0P+n6IqvU
CBaBAQAGmJ5u1P7bwAIY+pmLr9+VK5+8YNxqYACAqNR94Zu9PYXeJIAB0m/Zss/Nm7f4rE/deeeDCxbccZXnAwDE44LLs3Vf+OaV
Z3A/XQQWwNCf9MRdfK98EXjx4mUrVz65ePGy0zfOn79k5con77mn3/xmvJJXcffdv3fzzZ/tydmFESOuXbnyyeXLnzh94+LFy+69
9/EePS8AMFClJYP7HQEMpEFzc+OkSeVZWdmph4MGDZo6dUZzc2PvzupSXZ1X8etfv/jWW69d6lHTps2uqTk6atTYwsIR6Z0PABCz
2DI4s7cnAAwE9fW12dk5kydP3717WwhhwoQp7e3tNTVHc3KGdO8za9a86dOvz83Nb2qq37598wcfvJfaXlIyae7c2woLi1pbm3fv
fm/79k3JZPKsG2+99a5p0z4dQmhrO3ns2JGNG9c0NNSlBsnMzFqw4M5Jk8o6OtoPHdo3ZMjQkydPvPHGv57/1Jf0KsrKZs2fv/hH
P/peV1dXav/bb18xaNA1HR1t48dPHj8+zJx5Ywjhpz/9QX19bQghJ2fIkiX3jhtXmpExaN++XRs3rkkmkyGEO+98sKWlKdXAM2fO
nT79+ry8/OPHa373u3XV1VVnnVhmZubkydetW/fLmTPnlpfP/t3v1l3J/1kAwAB2eVcm133hm8Of/eZlnCtdtwW+agQw9Bs9cf1z
ytNP/3XqG7OuZPCKivfKyman0rG8fHZFxfbhw4u7n73++gWTJ09ft+6XdXUfjRw5+rOfvb+zs2Pfvl1ZWdl33PHgpk2vV1RsHzJk
6LRpc4qKRjY21p+5sbb22Jtvvvrmm6+GEIYMyb3xxlvvvPOhF174x1SOzp+/ePTokpdffq65ufHTn7552rQ53ZV7rlNf6qvYv3/3
Lbd8dsKEKYlERQhh8OCciROnrlnz0qFD+3JyhjQ01G/cuOb0ocrKZr3xxq/Wr/91YeGIZcs+V1v7YUXF9tN3uPHGRdOnf/qNN/71
yJGDhYUjyspmnSuAS0vLOzraq6oOZGZmLVx45+bNv+mOcACAtEitA19GBvcvLoEG0mPfvl1FRcXDhxfn5uaPGzdxz54d3U9dc03m
nDk3bdjwWk3N0c7OjqNHK3fufLusbHYIISdnaGZm5qFDezs7O5qaGjZv/m1t7bGzbjz9XCdOtGzY8FpBQVHqeuDMzKxp0+Zs3vyb
2toP29pObNr0ektL0wVPfamvIpXN5eWzUg+nTLmure1EZeX+cw118ODeioodHR3tNTVHDx3aO3r0+NOfzczMmjVr3pYtbxw8uCe1
z/r1r55rqGnTZn/wwXvJZPLgwT3JZHLixCnn2hMA4EoM+CuirQAD6dHR0b5//wdlZbPa29uqqg60tjZ3P1VYOCIzM+ueex5NPczI
yAghNDXVp/63snL/ihVf2Ldv19Gjh44cOdjZ2XnWjSGE4cOL5827bdSosTk5Q1ND5eUNO368Jj+/YNCgaz766MPUxmQyWVv74QVP
famvIoTwwQfvrVjxhaFD81pbm8vKZu/ZszN1VfNZNTYe7/5zW9vJ3Nz8058tLCzKzMysrq4890/0Y8OGDb/22pJ1614OIXR1dVVU
bC8vn3PgQMUFDwQAuDwDeDVYAEO/d/4vcD7rsz1xHXUIoaLivTvueKCjo2PjxrWnb09l589+9g+pD8d+wurVPx0zZvy4caVz596W
kzNk1aof19fXnnXj3Xc/XFm5/6WXnmltbU4mk1/+8n/JyDj9MpbTWzTjYk59Sa8ihFBTc7Surmbq1JlVVftHjBi1du0vzjPOudP4
383wgsrLZ2dkZDz++B+dNnIyNze/e5UbAKAnDMgMFsBA2lRXV7W1nczOHlxZue/07fX1tZ2dnePHf+pcFXr0aOXRo5VbtrzxwANf
7P6Sp09s3LFjS17esO3bN6e+lrm4+NpBgz6u36amhq6uU8XFo5uaGkIIGRkZI0aMqqo6cDGnvvhXkbJ793szZtyQm5tXXV3V0PDx
Gm9XV1eqtC9eamKjR5ecf2KDBg0qK5u5du0v9u/f3b3xvvseLyub9c47b13SGQEALskAS98UAQz93rmWc3vuS7PO4/nnnz5zY2dn
x3vv/e6GGxa2tjYfOrRv8OCc8eMn5+TkvPPOhlGjxpaVzXr//XcaGo4XFhbl5uY3NtafdeOJEy1tbSfLy2e9886GYcMKFy1aevr4
u3dvmzv3tsbGutSXYHVfb3yeU1/qq0jZu3fnTTctmTZtTurruFKamhqLi6/Nysru6Gi/yB9UZ2fHjh1b5s5d1NracvTooYKCovLy
2Wd+DHj8+MmDB+d84pPGicSemTPnvvvuv3sJN9ywcObMuc888+0zH37iKQCA8xuQ6ZsigIGr4e2315882Xr99QsWL17W0tJcWbkv
tYBZU3O0qGjkkiXLhg0rOnGiZffubbt3vxtCOHNjMplcu/bnt9xyx6xZ81pbW3bs2FJQsKh7/E2bfrNwYfZ9932+o6O9snJ/VdWB
U6dOnf/Ul6e9ve3AgYqJE6ecviS7c+eWJUvue+KJ/5SZmdV9G6QL2rLlt+3tJ2+55fahQ/Nqa2s2bTrLzY3Ky2cdOXLoE1198OCe
m2++fezYiZf9KgAAzmoAp29KRnFx8YX3ggiUlZUlEomrf97S0tL6+vqL3PmSFnUvfudeWSvuUY88sjJ19+CeGHzp0kebmxu7bzIM
ANB3XPA+wOf6kufLS9/Lvg9wYWFhr7z3dhskYCAYO3biddddn5MzdMiQ3Jtu+kxu7rADB3Zf+LDLOtHYsRN37tzaE4MDAFyhyyvS
q1y/vcgl0MBAcOzY4YkTpzz44JeysrJqa4+98spzqe/KSq/HH/+P2dnZmzf/9vjxmrQPDgBw9Q34a54/QQBDf/L003+9cuWTK1c+
mcbLlQfG9c+dnZ0bNqzZsGFNj57luee+16PjAwBcNbGlb4oABgAAGDjGfuv5838SOC3p2x+vfw4+Awz9TmqpNrVse8E9L7iuOzCW
fwEAuEhxLvx2E8DQ/1x8A5+f+gUAGJB6enm2ny7/BgEM/dqVNPCV9zMAAH1WzzVq/63fIIChn+petr28ju0+yvIvAMBA1ROl2q/r
N4SQUVxc3NtzgD6hrKysV27GXVpaWl9ff9mHX2rKSl8AgKic/wuxLkka67ewsLBX3nsLYPhYPw3gcMYi8FnL9mL2AQBgQEpLA6d3
7VcAQy/rvwGccpHXQktfAIAIXWEDp/3K594KYPcBhgHi/J8K1r0AADG74M2Bz39seifTiwQwDDRaFwCAM6U69pIyeCClb4oABgAA
iEV3056nhAde93YTwAAAANEZwJV7Hu4DDAAAQBQEMAAAAFEQwAAAAERBAAMAABAFAQwAAEAUBDAAAABREMDQyxKJRGFhYW/PAgAA
rpLCwsJEItErpxbAAAAAREEAAwAAEAUBDL3PVdAAAESiF69/DgIYAACASAhg6BMsAgMAMOD17vJvEMDQd2hgAAAGsF6v3yCAoU/R
wAAADEh9oX6DAIa+RgMDADDA9JH6DQIY+iANDADAgNF36jcIYOibNDAAAANAn6rfIIChz0o1sAwGAKA/Sr2V7VP1G0LI7O0JAOeU
+n1RWlqaelhfX997cwEAgAvrXr/pa+mbIoChr+v+3dFdwgAA0Df1ze7tJoCh3+jjv00AAKCP8xlgAAAAoiCAAQAAiIIABgAAIAoC
GAAAgCgIYAAAAKIggAEAAIiCAAYAACAKAhgAAIAoCGAAAACiIIABAACIggAGAAAgCgIYAACAKAhgAAAAoiCAAQAAiML/A7PInSrn
TKETAAAAAElFTkSuQmCC
""".replace("\n", "")
_SCREENSHOT_NARROW_B64 = """
iVBORw0KGgoAAAANSUhEUgAAAu4AAAU2CAIAAABFtaRRAABrWElEQVR4nO3deXxU5aH44ZOwBZKwySJBERBc2BFFcUFxRYwKoni9
VStXUSvWVkvRWw3qdSst163ys6LVqlQLgguKuIEQBWxdgCCissUNomwhbGFLfn/MvXNTICGEbC88z6d/TM6c5Z2ZlPn6njOThLS0
tAgAIEyJVT0AAICykzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCk
DAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABq1nVAwCAcrBp5JCqHsKBpd6tT1b1EP5H
QlpaWlWPAQDKTsRUoeoQNE4wARAwHVO1qsPzL2UACFV1eB+lyl8FKQNAkKr8HZS4qn0tpAwA4dEx1U0VviJSBgAImJQBIDCmZKqn
qnpdpAwAEDApAwAEzLf9VpTdzrNVh68SAoD9iZQpfyWcLIzdJWgAoLxImfJUyiueBA0AlJdqlzLDht1f0YcYNep3FbHbvb1ye9PI
IXtVM3//+987dOhQdMkFF1yQnZ2902onnHDCmDFjii557733brnllr0a216ZO3duYuL/XHR1+eWXZ2VlVcQmZXPHHXcMGjSo6JL+
/fsvXbp0tysXfYZnzZp1/fXXV9CogND9vxP7X9a2WxRFp08ZM2f18n3Z1W1dTru1y2lRFHV99eFvN+SWZpNXz7zy1IPbfrcxt8sr
D+/64wHIZb/lo2yfQNvHz61dfPHFuy685JJL9mWfu8rIyMj6XxkZGeW78wpVs2bNs88+e6eF6enpVTIYACqIlCkH+1Ik+7LtBRdc
ULt27aJLDjrooD59+pR5h/uZU045pWHDhjst7NevX0JCQlUMB4AKUe1OMFXQ2Z/9UsOGDc8888w333wzvmTAgAE1a1a717Sq7HYC
Ji0trXv37p999lnljwfYX8VO8Xy/cd0jX8y8st0x7VIPWrVl41++/viRBTNjKxzb5JDbupzWpXGL1Fq1l65f897yxU98+Y/lm/Le
73dtt8ZpsXXm9f91FEVf5v7U643/N+akiy5p0yW2fPP2bUvWr35p2fw/fTGrMCqsisdX3Xnb21f7/uWGe3vRTFGXXHJJPGUSEhIG
Dhy4j4OpHN26davoQ6SkpPTu3Tv+46pVq5o0aRK7nZ6eLmWAcndIcoOjGzS94N1nbzj6hN92PvWu7mfNW7Ni+oqlSTVqTTj98ga1
kwZ/8NKU779qldzw3EOOvKnDSbd9MqXPm2N2e63MtTNfvnbmy1EU1U6s0aPJIeP6/Pvdx5y1o7Bg9MLZVff4qi8nmIK0YMGC2I0e
PXq0adMmdrtXr14tW7aM3f7888933eq6666LX/Xy0EMP7XRv69at4/f+85//TElJeeSRR7KysopefHPJJZfE18nMzNzt2A4++OA7
7rjjrbfe+vTTT2fMmPHHP/6xbdu2O60zd+7c+H66dOmy071HHnnkrbfe+ve///2DDz747LPP3n333b/85S//9m//1qBBg1I9O1EU
RdHZZ59dp06d2O0NGzaMHj266F21atUq/a4ASmN7YcGdc97L3bp5Qvb//Ascm3Fpk9qoQe2kKIpaJTdsUid5Ud6qR7+YedsnU0qz
z60FO2b/9M0/V34XRdHFbTpX2NjDJmX2SXn9vYm93U9mZuZPP/0Uux2/+DfeHAsWLFi4cOGuW40fP37Lli2x26eddlrTpk2L3tu3
b9/47bfeemvDhg17NaSYk0466dVXXx00aFBaWlqtWrUaNWp0zjnnjB079vDDDy/N5rVr1x4xYsRLL730s5/9rEOHDg0aNKhZs2bz
5s2PO+643/3ud7/97W9LP5LzzjsvfnvGjBnvvffejh07Yj/Wr1+/6IQNQLlYuXnDhm1boijasmN7bElSjZpRFGVvWLtmy6Yoiu4+
5qzPL7r5u0v/c8Lplx/ftFUJu+rV7LCXz7hiySXDV//szrWX33VGWrsoig5LaVThjyFMUiZIO3bsePnll2O3L7jggjp16jRt2vS0
006LLXnppZd2u9XatWvjZ6Nq1KjRv3//oveec8458dsTJkwo28B+8Ytf1KtXb6eFKSkppayQBx54YLcfy9pbzZs379GjR/zHqVOn
rlu37tNPP40vKRo6AOVie2FB7MZOV7Rs3r5t4NSxb3y3cFX+xiiKUmrVOSOt3fjT/z25Zu3d7CWK0urVn3j65X1aHJ6Zs+zwl/7Q
aOxd01YsiaKoho8sFEPKhOrll18uKCiIoqhBgwZnnXXWgAEDatSoEUXRhg0bpkwpdt5y7Nix8dsDBw6Mf7lLu3bt4hMnixYtin3R
y69+9asuXboUDaOXXnqpy/8qbmLjueeeO/vss88+++zp06fHF55wwgl169Yt+RH16dPnrLPOiv+4ZMmSm2666fTTTz/xxBMHDRr0
8MMPxyei9ui8886LP7QtW7Z8+OGHURRNnTo1vkLv3r1TU1NLuTeAfTR3zfIrZoxrP+GPh7/0hz99MSuKovq1kg5NaRhFUUHhzlfy
dm3com7NWlEUTcz+PHfr5sSEhCMbNKn0IYdEyoQqJyfngw8+iN2+5JJLLrrootjtN998c/PmzcVttWjRon/+85+x22lpab169Yrd
Lnp2qcxTMlEUzZo1a9SoUTk5OTk5OaNGjYovT0xMjF/HU5wLLrggfnvt2rWDBw+ePn36qlWrNmzY8OWXXz799NOPPvpoKYdRdNJl
5syZ+fn5URRNmzat8H//yahdu/auXzkDUBGaJaVMOP3ys1secXDd1E3bt20v2BFF0cr8jUvzVkdR9O3G3NhqXRu3iN1YmPvTtoId
URSd3bJ9aq06wzr1bllvL64UPABVu08whfttv5VvwoQJp556ahRF3bt3jy8s7uxS3PPPP9+zZ8/Y7YsvvnjmzJlRkZTJz89/4403
yjykyZMnx2//+OOPRe/a46xM0e8ynjJlSm5ubtnGcMQRR7Rv3z7+Y3wy5scff1ywYEGnTp1iP6anp0+cOLFshwAovZ/yNzy84MPB
Rxz7x+P6HVwvde2WTa9+s2Dk/BlbC3ZEUfRy9uenHtz2rJbtn+09KCFKeGbRJ7f8443BH7x0R9fTL23b9eyW7Sdmfz4jZ+mpB+/8
+Qniql3KUHoffPBBTk7OwQcfHF8yf/78r776quStMjMzv/3221atWkX/e/HvQQcdFPsxiqK33357/fr1ZR7SihUr4re3bt1a9K49
fjFd/fr147eXLy/7F4EX/TqZHTt2zJgxI/7jtGnT4ilzzDHHtGjRouiAAUrphlmv3jDr1fiP/d97rui9327IbTT2rqJLPvwx+8Mf
s3e7q60FO34x65WdFk7+7svJ331Z3NF3OtxOPx6AnGAKWEFBQfzi35g9TslEUVRYWPjCCy/Ebscu/i2vs0vRv+ZL4S4ngEuWl5cX
v52Wlla2ASQmJvbr1y/+Y40aNT788MP4B79vuumm+F0JCQku/gXYD1S7WZmwzv7Uu/XJcvk8dpm/Iu/ll1++/vrrY5e4btiw4a23
3irNVq+++uqNN96YkpISRdHAgQPjzbF48eJ58+bttHLRIoldWVxBFi5c2KLF/5wqPvfcc//85z+X4RzTcccd16xZs1KufN555z31
1FN7ewgAqhWzMmH76aef4idQXn/99dj1rXu0adOm+HROWlpa/ILc3U7JFJ0s6d69e4sWLSrobxi99tpr8duNGjV65plnTjvttIMO
Oig5OfnII48cMmTIzTffvMed7NVfizz88MOPOuqosowVgGqj2s3KBGffJ2bKPCUT86tf/aoMW73wwguXX355/BPLURRt2bLl9ddf
33XNL7/8v/O1bdq0efvtt2O3//a3v40cObIMhy7O+++//95775155pmxHw8//PCdPrI0adKkkvdQp06dM844I/7j3//+9/vv3/kq
8qSkpOnTp8e//CY9Pb3oAwQgOGZlDlDLly+fNm1a0SXFXfA7Y8aMb7/9tnJGddttt+3Lp4r69OkTO2sWs9v0yc/Pf++99+I/nnvu
uUV7DoDg+Ee8HOzLtMo+Tsnsi6JflxcVf8Hvli1bBg8ePGHChOXLl2/btq1Ch7R169a77777kksu+dvf/rZw4cL169dv3779p59+
+vjjj3//+9//8Y9/LHnzopfxLlu2bLd/iCqKoqKTT02bNj3++OPLZfAAVImEMn9UhJ2U4TRTFXZMFEWJiYkzZ85MTk6OomjJkiUD
BgyowsEAlF55/f07yl2VvK+ZlSk3e/v6VW3HNGzY8Fe/+lWsY6IoGj9+fBUOBgDKrNpd9hv0t/3G6mSP/7lQtRGTkZER/xvaMatW
rXr11VeraDgAsE/MypS/erc+WVyslHBXVSkoKLjzzjtL+LNNANVNdfuHlJiqel2q3azMfqP6/z9t9erVX3zxxeOPP17c5bEAUP25
7BeAILn4t1qpwv+Ad4IJgCBV/8nvA0cVfyC3Co8NAPtCzVQHVf4qSBkAAlbl76MHuOrw/LtWBoD9gUtnKll1iJgYKQMABMwJJgAg
YFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkA
IGBSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICA1WzdunVVjwEAoIwS0tLSqnoMAABl
5AQTABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIA
QMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMy
AEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABKxmVQ8AqkDDhg2reggA+4/c3NwqPLqU4QCiYAAqQvxf1yppGinDAUHEAFSC
2D+2lRw0rpVh/6djACpTJf+ra1aG/ZmIAagSlTk9Y1aG/ZaOAahalfPvsJQBAAImZdg/mZIBqA4q4V9jKcN+SMcAVB8V/W+ylAEA
AiZl2N+YkgGobir0X2YpAwAETMoAAAGTMuxXnF0CqJ4q7t9nKQMABMwfLoAoMTGxX79+/fr1a9u2bUJCwvLly2fPnj1hwoQ1a9ZU
9dBKKyMjo0GDBsOGDdvbDYcNG9amTZuhQ4dWxKjKVxUO9brrrvvZz34WRdH27dt//PHHqVOnjh07Nj8/f48btm/f/i9/+cvll1/+
7bffVvww4QAlZTjQJSYm3nPPPZ06dXr88cdnz569ZcuWjh073nDDDWlpaXfddVdVj45q4YknnnjiiSeiKKpdu3bnzp3vvPPOFi1a
3HPPPVU9LiCKnGCCSy655MQTTxw2bNhbb721bt26/Pz8Tz/99Be/+MXcuXOremhUO1u3bv30008nTpx42mmn1apVq6qHA0SRWRkY
MGDARx99tGjRoqILt27d+uqrr0ZRNHjw4MGDB8eW/PDDDxMnTpw0aVJ8tYyMjCZNmmRlZZ1xxhlNmjRZtmzZww8/vHDhwp0Ocfrp
p99xxx0XX3xx/IzVUUcdNWbMmBtuuOHzzz9PSEi45JJLBgwY0Lx589zc3KlTpz755JNbt26NomjUqFFr1qy5//77i2512WWX/fDD
D7s+kFq1al199dXFjaRPnz7//u//3rZt282bN0+bNm306NFbtmzZaQ+xh7Nw4cJTTz21adOmX3/99UMPPbTTMxNfs2nTposXL+7V
q1fTpk2XLFkyatSoww477Morr2zZsuWKFStGjx49e/bsvTr6Xg2ghNellC9KyTsp2Y4dOxISEnZ7V0JCwlVXXXXhhRempqYuWbJk
ypQppTliyb8hpRkSHMjMynBAq1+/flpa2hdffFHcCs8880zv3r179+6dnp4ee2vp06dP0RW6d++empp64403XnzxxTk5Offee2/N
mjv/F0JmZua6devS09PjSy688MIlS5bE3qUGDhw4ZMiQp556Kj09/Z577jnzzDNvvvnmMjyWEkZy3nnn3X333TNmzBgwYMCNN97Y
uXPnW265pbidtGjR4pe//OW//du//fTTTw8++GBycvJu1+zWrVtBQcH1119/2WWXJSQkPPzwwwMHDhwxYkR6evqsWbPuuuuu+vXr
7+3RSz+Akl+X0rwoe9zJbtWqVatLly4XXXTRuHHjtm3btusKgwYNGjRo0B//+Mfzzz//mWeeufbaa0tzxJJ/Q4CSSRkOaKmpqVEU
bdiwYY9r5ufnz5w58+233z7rrLOKLs/JyXn00UfXrFmTl5f3/PPPN23atGXLljttu3379kmTJl1wwQWJiYlRFKWkpJxxxhmxWZ8o
iq666qpXXnll6tSpmzZtmjNnzpgxY/r169esWbO9fSzFjSQhIWHIkCGZmZljx47Ny8vLzs4eNWrUOeec06JFi113smnTppEjR65a
tWrVqlUjR46sWbPm+eefv9vDxaZe1q1bt3Llytdffz01NfXBBx/Mzs7evHnzCy+8ULdu3SOPPHJvj75XA4jZ7etSmhdljzvZ1dSp
U6dOnfrYY499+eWXzz777G7XueKKK8aPHz9z5sxNmzbNnj37hRdeKM0RS/4NAUrmBBMHtPXr10dRlJKSUtwKzZo1u+aaa3r06NG4
ceMaNWpEUbTT+Y7vv/++oKAgdnvdunVRFDVo0GDX/UyaNOmKK67o1avXzJkz+/btW1BQ8O6770ZRdNBBB9WvX3/+/PnxNbOyshIS
Etq0afPTTz/t1WMpbiTNmzdv3LjxrFmz4msuXLiwoKCgffv2K1as2Gkny5Yt27RpU+z25s2bly5d2rZt2+IOV1hYGLsdexrjH9KJ
/Ribldmro5d+ACW/LqV8Ufb44u7kjDPOqFWrVrt27X7zm9+MHj36uuuu2759e9EVYq/mggUL4kt2mlYp4YjF/YYAeyRlOKDl5eWt
WLGiQ4cOxa1w//33r1u37re//e3y5cvz8/OHDh3as2fPoivE3zJLtnr16g8++KB///4zZ87s37//O++8E3/DjqIo3gRRFBV3EUbJ
d5Uwkth/6N9222233XZb0eVNmzYtzciLU3TMxS2puKOX/LqU8kXZ44u7q23bti1cuPChhx56/PHHu3Xr9sknnxS9t+QXqOQjlvwb
ApRAynCge/nll6+//vp27dotXrw4vrB27dr9+vV76623jjjiiP/8z/9cunRpbHmnTp3KfKCJEyf+6U9/Sk9Pb9Wq1Z133hlbuHr1
6ry8vM6dO3/44YfxQxQWFi5btiyKory8vNgpsJjDDjusDMfNycnJzc194oknJk+evMeV27RpU69evdibaN26ddu2bRsfWNns1dFL
OYCkpKR9f132ZSfFfXZp1apVeXl5HTt2/Oc//7nrPvd4xN3+hgB75FoZDnQvvfTSRx999N///d/nnHNO/fr169Sp06NHj9h/c+fn
5+fk5PTt27dhw4aNGjW64YYbOnbsWOYDZWVlLV269Oabb54/f/6SJUviy//6178OGDCgT58+9erV69q165AhQ958883Y2aU5c+b0
7NmzR48eSUlJXbt2vfrqq8tw3IKCgieeeOL6668/9dRTU1NTk5OTu3btet999+12XqRevXq33npr06ZNmzRpMnz48B07drz++utl
fsh7e/RSDqBcXpe92smIESN69erVsGHDunXrduvW7eabb87Ozp43b96ua44dO3bQoEEnnXRSvXr1evXqddlll5X+iMX9hgAlMyvD
ga6goOD2229PT0/v37//LbfcEvu231mzZk2cODGKooyMjF//+tcTJkzYsmXLJ598Mnny5BLORu3RK6+8MmzYsJ0u55w4cWJiYuK1
116bkZER/zB27K7Jkycfcsghd9xxR1JS0tdff/23v/2thM/+lGDy5Mnr1q277LLLbr/99m3bti1evPill15auXLlrmvOmTNnxYoV
f/rTn2Kfhf7Nb35Tmmuiy+vopR9Aubwupd/JCy+8cMUVV9x6661169ZduXJlZmbm3//+991+gmn8+PHJycnDhw9PSUlZunTpk08+
WfTzaHs84m5/Q4CSJaSlpVX1GKDcVPM/J3nVVVcNHDjwoosu2u27YJUr818/2G8GUOWq+W8I7KPc3NyK2K0TTFBJWrZsOXDgwOK+
jwT8hkDZOMEEleHpp59u1arVzJkzx40bV9VjoTryGwJl5gQT+5VqfoIJ4EDmBBMAwM6kDAAQMCkDAARMygAAAZMyAEDApAwAEDAp
AwAETMoAAAGTMrBnw4YNGz16dGUe8brrrnvnnXfiP9aoUWPEiBHvvPPOcccdVwlH79Sp0x/+8Ic33njj9ddf/93vfpeamloJBwUo
GynDge7WW2999NFHq/POa9Wqde+99/bq1WvYsGEff/xxJYzkZz/72fjx4y+66KLf/OY3xx13XNn+HDdA5fA3mGDPRo0aVVWHTkpK
uv/++9u3b//rX//6q6++qpyD/ud//mfsxtdff/3ZZ5916dKlco4LUAZShgPasGHDzjvvvCiKMjMzoyh67rnnnnrqqYyMjCZNmixY
sODcc89t1KjR4MGDBw4c2KZNm6FDh0ZRFLt34cKFp556atOmTb/++uuHHnpo0aJFpdn5X/7yl0suuWTAgAHNmzfPzc2dOnXqk08+
uXXr1uKGl5yc/Ic//KFFixa//OUvs7Ozoyg6/fTT77jjjosvvnjNmjWxdY466qgxY8bccMMNn3/+eUJCwm73v9uHWdzKRQdw8MEH
H3/88R988EG5PNsAFUHKcEAbNWpUjRo1WrZsedNNNxVd3r179x9++GHIkCGrVq3adavu3bvn5ub+8pe/jKLoxhtvfPDBB//t3/5t
48aNe9z5xRdfPGTIkN///vezZ88+8sgjR4wYkZKSMnLkyN2OrUaNGo888khqauqNN964fPny2MLMzMx169alp6c/99xzsSUXXnjh
kiVLPv/88yiKBg4cuNv97/ZhFrdyfIUmTZo89NBD2dnZjzzyyN48qQCVyrUysBtr1qx58MEHd9sxURRt2rRp5MiRq1atWrVq1ciR
I2vWrHn++eeXZrdXXXXVK6+8MnXq1E2bNs2ZM2fMmDH9+vVr1qzZbleuUaNGmzZtvvnmm59++im+cPv27ZMmTbrgggsSExOjKEpJ
STnjjDNeffXVMux/jyvfdddd69evHzZsWH5+fmkeHUCVkDKwG99888327duLu3fZsmWbNm2K3d68efPSpUvbtm27x30edNBB9evX
nz9/fnxJVlZWQkJCmzZtdrv+1q1bb7/99h49etx11101atSIL580aVLjxo179eoVRVHfvn0LCgrefffdvd3/HleuU6dOp06dxo8f
r2OAak7KwG5s27atgvZcWFgYv52QkFDyyh999NEdd9zRq1evojWzevXqDz74oH///lEU9e/f/5133ol31d7uv4SVk5OTExMTt2zZ
UvIeAKqclOFAt3379tjJmtJr06ZNvXr1Yrfr1q3btm3bZcuW7XHnq1evzsvL69y5c/zeTp06FRYWFrdtzOzZszMyMk488cQRI0bE
dzVx4sSePXump6e3atXqtddeK83+d3qYexzMmjVrevfu7YJfoPqTMhzoli9ffuihhzZp0qT0m9SrV+/WW29t2rRpkyZNhg8fvmPH
jtdff700O//rX/86YMCAPn361KtXr2vXrkOGDHnzzTeLXgqzW7NmzRoxYsQpp5wSr5msrKylS5fefPPN8+fPX7JkSXzNEva/68Ms
eTAtW7bMzMw844wzSv+0AFQJn2DiQPf666937dr1ueeeS0lJiX1KeY+bzJkzZ8WKFX/6059iH8b+zW9+s2HDhtLs/C9/+UtiYuK1
116bkZER//xzaQY5c+bMO++88+67777jjjvuvffegoKCV155ZdiwYfELfmMmTpxY3P53fZglrAwQkIS0tLSqHgOUm4YNG1b0ITIy
Mho0aDBs2LCKPlDJrrrqqoEDB1500UUVd1kPQPnKzc2tiN06wQThadmy5cCBA8eNG6djAJxggsA8/fTTrVq1mjlz5rhx46p6LABV
zwkm9iuVcIIJgLJxggkAYGdSBgAImJQBAAImZQCAgEkZACBgUob9SgVdHg/APqq4f5+lDAAQMCkDAARMyrC/cY4JoLqp0H+ZpQwA
EDApw37IxAxA9VHR/yZLGfZPagagOqiEf42lDAAQMCnDfsvEDEDVqpx/h2tWwjGgqsT+X9SwYcMqHgfAAaYy/2PSrAz7P9MzAJWp
kv/VNSvDAcH0DEAlqJL/dJQyHEDi/x/TNADlqGonv6UMByKnnAD2G66VAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkAIGBS
BgAImJQBAAImZQCAgEkZACBgUgYACFjV/zlJf6MY4ADnL7yyL6osZRQMADHxdwRNQxlUQcqIGAB2K/YGIWjYK5V9rYyOAaBk3inY
K5U3K+NXE4BSMj1D6VXSrIyOAWBvee+gNHwYGwAIWGWkjKwGoGy8g7BHFZ4yfgsB2BfeRyiZE0wAQMAqNmWkNAD7zrsJJTArAwAE
TMoAAAGrwJQxHwhAefGeQnHMygAAAauyv4y9W7Vr17700kvPOuusgw8+eP369V988cVLL72UlZVV7gd66KGHatSocdNNN+323uuu
u+5nP/tZFEWFhYWbNm3KycmZO3fuyy+//N1335XjGGrVqnXiiSeed955PXv2zMrKKm4wJY8wbu7cuXu1h7iMjIwGDRoMGzZs17uG
DRvWpk2boUOH7u0+S9jwuuuuGzhw4Nlnn12GoZbtiOWlffv2f/nLXy6//PJvv/224o4CwN6qXilz2223HXvssffff/+8efOSkpI6
dOhwzTXX/OY3v9m2bVsljyQ/Pz/2dpucnHz44Ydfeumlf/3rX++5557p06eX1yFOPvnkM84449VXX61Zs2bNmnv9QsRHCAAHsuqV
Mqeeeur48eM/+uijKIo2b948c+bMmTNnVu2QNm7cmJWVlZWVNWLEiN/97ndz584trz9v9v7777///vtRFF100UXlssNyN2rUqEre
sMwq/4gAVBPVK2VWrVrVuXPnlJSUDRs27HrvqFGj1qxZc//998d+POqoo8aMGXPZZZf98MMPGRkZTZo0WbBgwbnnntuoUaPBgwcv
W7asT58+//7v/962bdvNmzdPmzZt9OjRW7ZsKfPYnn/++TPPPLNPnz6vvPLK4MGDBw8eHEXR1q1bf/jhh4kTJ06aNCm22umnn37H
HXdcfPHFa9asKTrOG2644fPPPy/94bp06fLYY4/ddttts2bNKuUmJTw/URR17Njxuuuua9++/datW+fOnfv444/n5OTE1qxVq9bV
V199xhlnNGnSZNmyZQ8//PDChQujfz1rk5GR0bRp08WLF/fq1atp06ZLliwZNWrUYYcdduWVV7Zs2XLFihWjR4+ePXt2bIelPN0T
e9WysrJ2PXQURSU8ybu+3AMHDizuiHvcT3EDSEhIuOqqqy688MLU1NQlS5ZMmTKllC8EAJWpel32O2rUqEMPPfTVV1/905/+NHTo
0G7dupV+2+7duzds2HDIkCGnnXbasmXLzjvvvLvvvnvGjBkDBgy48cYbO3fufMstt8RXLiwsLCws3KuxZWdn5+fnH3744VEUPfPM
M7179+7du3d6enosU/r06RNbLTMzc926denp6fENL7zwwiVLluxVx5S7xMTEkSNHLlq06NJLL7388sunTZtWdITdu3dPTU298cYb
L7744pycnHvvvXe3J7y6detWUFBw/fXXX3bZZQkJCQ8//PDAgQNHjBiRnp4+a9asu+66q379+ns7sBIOXcKTHO3ycpdwiD3up7gB
DBo0aNCgQX/84x/PP//8Z5555tprr93bRwdAJaheKfPxxx9feumld9xxx7x58zp06PDoo4/+/ve/L+V1JGvWrHnwwQdXrVoVRVFC
QsKQIUMyMzPHjh2bl5eXnZ09atSoc845p0WLFrGVN23atGnTpr0aW2Fh4ebNm5OTk4suzM/Pnzlz5ttvv33WWWfFlmzfvn3SpEkX
XHBBYmJiFEUpKSmxC2L26lhRFGVlZfXu3buEKZmkpKTMIo4//vgS9paSklK/fv3Zs2fn5eWtX79+xowZTz31VPzenJycRx99dM2a
NXl5ec8//3zTpk1btmy5605iUy/r1q1buXLl66+/npqa+uCDD2ZnZ2/evPmFF16oW7fukUceubcPszSH3vVJjv715S6l3e6nhAFc
ccUV48ePnzlz5qZNm2bPnv3CCy/s7aMDoBJUrxNMURTl5+d/9NFHsctlzjrrrIyMjAsvvHDixIl73PCbb77Zvn177Hbz5s0bN25c
tAMWLlxYUFDQvn37FStWRFG0fv36WrVq7dXAEhIS6tWrt3HjxiiKmjVrds011/To0aNx48Y1atSIomjRokXxNSdNmnTFFVf06tVr
5syZffv2LSgoePfdd/fqWKWx62W/l1xySXEr5+XlTZky5d57750+ffq8efM++eSTlStXxu/9/vvvCwoKYrfXrVsXRVGDBg123cn3
338fn8pav359FEXxz/LEfizDrEwJhy75SS76cpes5P0UN4CDDjqofv36CxYsiK9ZtfNqABSn2qVMUe++++6NN9549NFH7/behISE
oj8W/ZRTbEbktttuu+2224qu07Rp09iNkSNH7u1g2rRpU6dOndi74P33379u3brf/va3y5cvz8/PHzp0aM+ePeNrrl69+oMPPujf
v//MmTP79+//zjvv7O0MULnY6fl54IEHJkyYcOyxx/bp02fYsGFjx4595plnYnfF38tLtuspub09SberEg5d8pNc+g+1lbyf4gaw
07MHQLVVvU4wXX311UV/rFOnTkpKSrwD8vLyUlNT4/cedthhxe0nJycnNzd35MiRvf9VaWZ3inPFFVds3rz5/fffT0pKOuKII15+
+eWlS5fm5+dHUdSpU6edVp44cWLPnj3T09NbtWr12muvlfmge2WPz8+iRYtefPHF4cOHP/bYYz//+c/3dl6qMpXmSa7Q/axatSov
L69jx47xJWUbAAAVrXqlzL//+7//4Q9/OOqoo2rXrt2qVauMjIyEhIQ333wzdu+cOXN69uzZo0ePpKSkrl277tQ9RRUUFDzxxBPX
X3/9qaeempqampyc3LVr1/vuuy8+K/PQQw89+uijpRlSvXr1unTpcu+9955yyin33XdfXl5efn5+Tk5O3759GzZs2KhRoxtuuKHo
G15MVlbW0qVLb7755vnz5y9ZsqQMT0WXLl0yMzNPPPHE0m9SwvPTtm3bESNGdOzYMSkpqVGjRp07d87Jyan8b+spvdI8yRW9n7Fj
xw4aNOikk06qV69er169LrvssjIMAICKVr1OMP3iF7/o37//iBEjmjVrtm7duq+++mro0KFffvll7N7Jkycfcsghd9xxR1JS0tdf
f/23v/2t6IeSdjJ58uR169Zddtllt99++7Zt2xYvXvzSSy8VvUCkZLGLamOX+ubk5MyZM2fw4MHxb/vNyMj49a9/PWHChC1btnzy
ySeTJ0/u0KHDTnt45ZVXhg0bVsIFvy1btnzxxRfjP2ZmZkZRtC9fJlvC85OdnT1r1qyhQ4e2a9cuPz9//vz5w4cPL9tRKk1pnuQK
3c/48eOTk5OHDx+ekpKydOnSJ5988uabby7DAACoUAlpaWkVtOsD/E9/XXXVVQMHDrzooouq8+QHQEDK6xtK2c9UrxNM+42WLVsO
HDhw3LhxOgYAKlT1OsG0f3j66adbtWo1c+bMcePGVfVYAGA/5wQTAGFwgondcoIJAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAIm
ZQCAgEkZACBg1e7bfhMTE88999xzzz338MMPr1mz5k8//bR48eJXX311zpw5sRWGDRvWpk2boUOHVtAArrvuup/97GdRFBUWFm7a
tCknJ2fu3Lkvv/xy/G9JlotWrVpdcsklxx13XOPGjb///vtXXnnljTfeKCwsLP0eEhMTx48f36xZs335C5TFycjIaNCgwbBhw3Z7
b0W/BFWrffv2f/nLXyriWQWgIlSvlElMTLzvvvs6duz4xBNPzJ49e/369bG3/EceeWTgwIGl/7vW+yg/P//ss8+Ooig5Ofnwww+/
9NJL//rXv95zzz3Tp08vr0Pcc889a9euzcjI+P7770866aT//M//bNy48bPPPlv6PfTo0aNp06Zr167t16/fn//85/IaGACEpXql
zKBBg0444YRrr7120aJFsSVLliz5/e9/P3/+/PiMxahRoyptPBs3bszKysrKyhoxYsTvfve7uXPnltfXZr/77rsvvPBCQUFBFEXv
vfdejx49Lrzwwr1KmX79+mVlZc2ZM+f8888fM2ZMbFeVozJfAgAoWfVKmQEDBnz00UfxjombPHly/HbRsxsJCQmXXHLJgAEDmjdv
npubO3Xq1CeffHLr1q1RFGVkZDRp0iQrK+uMM85o0qTJsmXLHn744YULF5ZtYM8///yZZ57Zp0+fV155JYqiwYMHDx48OIqirVu3
/vDDDxMnTpw0aVIURaeffvodd9xx8cUXr1mzJrbhUUcdNWbMmBtuuOHzzz8vusOxY8cW/XHLli01atSI/9ilS5fHHnvstttumzVr
1m7Hk5KScsoppzz88MNz5sy56qqrjj/++NmzZ8fuij3wBQsWnHvuuY0aNRo8ePBpp52229FGUdSxY8frrruuffv2W7dunTt37uOP
P56TkxO7q1atWldfffVun734S1Dy4y3uWYoNsmnTposXL+7Vq1fTpk2XLFkyatSoww477Morr2zZsuWKFStGjx4df0Ql7Keoklcr
4fchISHhqquuuvDCC1NTU5csWTJlypTdPucAVE/V6LLf+vXrt2jR4osvvij9JgMHDhwyZMhTTz2Vnp5+zz33nHnmmTfffHP83u7d
u6empt54440XX3xxTk7OvffeW7NmGdMtOzs7Pz//8MMPj/34zDPP9O7du3fv3unp6bF37j59+kRRlJmZuW7duvT09PiGF1544ZIl
S3bqmJ2kpKT06dMnKyur9OM588wzExIS3n///R9++GHBggX9+vUrem/37t0bNmw4ZMiQ0047bdmyZcWNNjExceTIkYsWLbr00ksv
v/zyadOmFR15aZ69kh9vcceN6datW0FBwfXXX3/ZZZclJCQ8/PDDAwcOHDFiRHp6+qxZs+6666769euX/GzvZI+rFfeIBg0aNGjQ
oD/+8Y/nn3/+M888c+2115b+hQCgylWjlElNTY2iaMOGDfElV155Zeb/Kvp+GXfVVVe98sorU6dO3bRp05w5c8aMGdOvX79mzZrF
7s3JyXn00UfXrFmTl5f3/PPPN23atGXLlmUbW2Fh4ebNm5OTk3danp+fP3PmzLfffvuss86Komj79u2TJk264IILEhMToyhKSUk5
44wzXn311RL2nJiYmJGRkZSU9MQTT8QXZmVl9e7du7gpmSiK+vXr9+GHH27cuDGKorfffvvEE0+Mv/FHUbRmzZoHH3xw1apVJY82
JSWlfv36s2fPzsvLW79+/YwZM5566qn4yqV59kr5eHc6bkxs6mXdunUrV658/fXXU1NTH3zwwezs7M2bN7/wwgt169Y98sgjS7Of
XRW3WnGP6Iorrhg/fvzMmTM3bdo0e/bsF154oYSdA1DdVKMTTOvXr4+iKCUlJb7kueeee+655xo0aPD666/vuv5BBx1Uv379+fPn
x5dkZWUlJCS0adPmp59+iqLo+++/j19Bsm7duiiKGjRoULaxJSQk1KtXL5YOURQ1a9bsmmuu6dGjR+PGjWMnhuInxSZNmnTFFVf0
6tVr5syZffv2LSgoePfdd0vY87Bhw4499thbb731+++/L+Vg2rRpc9RRR8UvrJk6deovf/nLs88+e8KECbEl33zzzfbt2+PrFzfa
vLy8KVOm3HvvvdOnT583b94nn3xS9MLqUj57JTzeEp6l2P7j1z/FXvr4J4ZiP8bjrOT97PFhlvyIYr9FCxYsiK9W8hQaANVNNUqZ
vLy8FStWdOjQYa+2KvoB5oSEhKJ3leOVsG3atKlTp078rfH+++9ft27db3/72+XLl+fn5w8dOrRnz56xu1avXv3BBx/0799/5syZ
/fv3f+eddzZt2lTcbocOHXruuedmZGR88sknpR9M7HTSAw88UHThueeeG0+Zbdu2Fb2rhNE+8MADEyZMOPbYY/v06TNs2LCxY8c+
88wzsbtK+eyV8HhLOG70ry9ccUtKs5/Sr7bbR7TT7wwAwalGJ5iiKHr11VePP/74tm3blmbl1atX5+Xlde7cOb6kU6dOhYWFy5Yt
K/eBXXHFFZs3b37//fejKEpKSjriiCNefvnlpUuX5ufnx45bdOWJEyf27NkzPT29VatWr732WnH7HDx48CWXXPJf//VfH374YelH
UqNGjbPPPjt+XUjM0KFD27dv3759+13X3+NoFy1a9OKLLw4fPvyxxx77+c9/XqtWrdIPJma3j3ePxy2lUu6nzIdbtWpVXl5ex44d
40vKNk4Aqkr1Splx48b94x//eOihh2KfvqlVq9bBBx984YUXRsX8J/tf//rXAQMG9OnTp169el27dh0yZMibb74ZO7tUgvT09MzM
zFatWu1xPPXq1evSpcu99957yimn3HfffXl5eVEU5efn5+Tk9O3bt2HDho0aNbrhhhuKvhFGUZSVlbV06dKbb755/vz5S5Ys2e2e
L7744p///OcPPPBALI920qVLl8zMzBNPPHHXu0444YRGjRrt9A03n3/++cqVK88999xd1y9htG3bth0xYkTHjh2TkpIaNWrUuXPn
nJycnWZ0SmO3j3ePz1IplXI/+3K4sWPHDho06KSTTqpXr16vXr0uu+yyMowTgKpSjU4wRVFUUFBw++23n3vuueedd95NN91Uu3bt
lStXfvfdd//1X/81Y8aMXdefOHFiYmLitddem5GREf8w9r4PIykpKTMzM3apb05Ozpw5cwYPHlz0234zMjJ+/etfT5gwYcuWLZ98
8snkyZN3Oi/2yiuvDBs2rIQLfq+44orExMTbb7/99ttvjy8855xzNm/eXPLY+vXr98033+w081RYWJiZmXn22Wc//vjju25S3Giz
s7NnzZo1dOjQdu3a5efnz58/f/jw4SUfvTi7fbx7fJZKqZT7KfPhxo8fn5ycPHz48JSUlKVLlz755JNFPwcHQDWXkJaWVkG7btiw
YQXtufq76qqrBg4ceNFFF5VhkiNEB9rjBapEeX1JKfuZ6nWCaf/QsmXLgQMHjhs37gB5Xz/QHi8A1Ur1OsG0H3j66adbtWo1c+bM
cePGVfVYKsOB9ngBqG6cYAIgDE4wsVtOMAEAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABKx6pcx11133zjvvxH+s
UaPGiBEj3nnnneOOOy6KomHDho0ePbo0+8nIyBg1atReHboMm+yq6AhLP9rrrrsuMzMzMzNzxowZU6ZMeeaZZ371q18deuih+ziY
Eg40bdq0F1988ZprrklKSir9hg8++GDRhffff/+jjz4au122B76TxMTE9PT0//f//t9bb7319ttvP/PMM9dee23jxo3LsKvilHls
5aJ9+/al/JPsAJRe9f3DBbVq1fqv//qvbt26DRs2LCsrq6qHU7Hy8/PPPvvsKIqSk5MPP/zwSy+99K9//es999wzffr0cjzKE088
8cQTT0RRVLt27c6dO995550tWrS45557SrNtQUHBscce27179zlz5pTjkOISExPvueeeTp06Pf7447Nnz96yZUvHjh1vuOGGtLS0
u+66q2z7vPXWW1u2bHnTTTeV60gBqF6qacokJSXdf//97du3//Wvf/3VV1/FFu77rElFKzrCso1248aNWVlZWVlZI0aM+N3vfjd3
7tyK+KLurVu3fvrppxMnTrzyyit///vfl+bPQK5bt27x4sXXXXfd9ddfv+u9+/7AL7nkkhNPPPHaa69dtGhRbMmnn376i1/8ol+/
fmXYGwAHjup1gikmOTn5v//7v1u3bv3LX/4y3jHRv54dyMjIeOSRR66++uoXXnjhnXfeeeKJJ44++ujd7q1BgwajR49+4oknGjVq
NHjw4NgZlvfee+/ZZ5+94IIL9rjJqFGjfve738XvOuqoozIzM1u2bLnbDcvlPEvM888/n5SU1KdPn9iPxY389NNPnzZtWtGzMLER
durUaY+H2LFjR0JCQux2ly5dMjMzTzzxxBLWHzNmTIcOHU4++eRd7yrugSckJAwaNOjFF1+cNm3ayy+/PHTo0Nq1a+925wMGDPjo
o4/iHROzdevWV199NXa7hNeuY8eOjz766JQpU1577bW777774IMPjg3jvPPO69atW2yra665ZteD9unT58knn5w6deobb7xxyy23
1KlTZ7djK+HQJf8eJiQkDB48+NVXX506deqYMWNK86IAsLeqXcrUqFHjkUceadKkyY033pidnV3Cmt27d09NTb3xxhsvvvjinJyc
e++9t2bNnSeZDj300D//+c9r16696aab1q5d+8wzz/Tu3bt3797p6eljxoy54YYb4q1Q3Cbl++hKLzs7Oz8///DDD4/9WNzIMzMz
161bl56eHt/wwgsvXLJkyeeff17CzmvVqtWlS5eLLrpo3LhxpZmSifnqq69mzJgxZMiQxMTS/toMHDhwyJAhTz31VHp6+j333HPm
mWfefPPNu65Wv379tLS0L774ooRdFfcMJCYmjhw5ctGiRZdeeunll18+bdq02LMxatSoyZMnz507N7bVU089tdMOzzvvvLvvvnvG
jBkDBgy48cYbO3fufMstt+zVoWNK+D0cNGjQoEGD/vjHP55//vmxS39K+bwBUHrVMWXatGnzzTff/PTTTyWvmZOT8+ijj65ZsyYv
L+/5559v2rTpTpMlXbt2ffzxxz/88MMRI0Zs2bKl6F35+fkzZ858++23zzrrrFJuUskKCws3b96cnJy80/KdRr59+/ZJkyZdcMEF
sbxISUk544wz4jMZuzV16tSpU6c+9thjX3755bPPPhtbmJWV1bt371mzZpU8qieffPKwww4788wzS/korrrqqldeeWXq1KmbNm2a
M2fOmDFj+vXr16xZs51WS01NjaJow4YNpdnnTs9ASkpK/fr1Z8+enZeXt379+hkzZuxaLbtKSEgYMmRIZmbm2LFj8/LysrOzR40a
dc4557Ro0aL0h44p4ffwiiuuGD9+/MyZMzdt2jR79uwXXnihNA8QgL1S7a6V2bp164gRI+6777677rrrzjvv3LFjR3Frfv/99wUF
BbHb69ati6KoQYMG8XvbtWv34IMPPv3003/729/iC5s1a3bNNdf06NGjcePGNWrUiKKo6BmN3W5SgmOPPTb+oZ4PP/yw6HmofZeQ
kFCvXr2NGzfuceSTJk264oorevXqNXPmzL59+xYUFLz77rsl7PmMM86oVatWu3btfvOb34wePfq6667bvn17KUf17bffvvXWW//x
H/8xbdq0Pa580EEH1a9ff/78+fElWVlZCQkJbdq02alT169fH0VRSkpKCXsr7hnIy8ubMmXKvffeO3369Hnz5n3yyScrV67c49ia
N2/euHHjoum2cOHCgoKC9u3br1ixopSHjinu9zD28BcsWBBfs+SpMgDKptrNykRR9NFHH91xxx29evW66667Yu8cuxV//9itnJyc
r7/++swzz2zUqFF84f3333/QQQf99re/Pffcc3v37j1u3Lii56R2u8lO4heXRFH0ySef9P5f5dsxURS1adOmTp068bfMEka+evXq
Dz74oH///lEU9e/f/5133tm0aVPJO9+2bdvChQsfeuih9u3bd+vWba8G9swzzzRt2vT8888v5fqFhYXx20WfvaLy8vJWrFjRoUOH
EvZTwjPwwAMP/PKXv/zmm2/69Onz97//ffDgwXscVWwS67bbbsv8X9OnT69Zs2bTpk336tBR8b+HxT1YAMpXdUyZKIpmz56dkZFx
4oknjhgxovRXZhS1YcOGW265JTc3d/To0c2bN4+iKCkp6Ygjjnj55ZeXLl2an58fRdFOl2HuukkURXl5ebHTHzGHHXZY2R/V3rji
iis2b978/vvvl2bkEydO7NmzZ3p6eqtWrV577bVSHqJWrVplGNiPP/742muvXXnllXv8TprVq1fn5eV17tw5vqRTp06FhYXLli3b
deWXX375hBNOaNeuXdGFtWvXjiXaHp+BRYsWvfjii8OHD3/sscd+/vOfxx7a9u3bi/vlycnJyc3NHTlyZO9/NXHixJ3W3OOhi7Nq
1aq8vLyOHTsWffil2RCAvVJNUyaKolmzZo0YMeKUU04pc81s3rx5+PDhy5YtGz16dKtWrfLz83Nycvr27duwYcNGjRrdcMMNRd9m
drtJFEVz5szp2bNnjx49kpKSunbtevXVV5ft4aSnp5fm69Hq1avXpUuXe++995RTTrnvvvvy8vKiKNrjyLOyspYuXXrzzTfPnz9/
yZIlxe18xIgRvXr1atiwYd26dbt163bzzTdnZ2fPmzcvKt0nmOKef/75evXqHXPMMXtc869//euAAQP69OlTr169rl27Dhky5M03
39ztVVAvvfTSRx999N///d/nnHNO/fr169Sp06NHj8cffzw2aVTCM9C2bdsRI0Z07NgxKSmpUaNGnTt3zsnJiV3LvHz58kMPPbRJ
kya7Hq6goOCJJ564/vrrTz311NTU1OTk5K5du9533327zsqU5temOGPHjh00aNBJJ51Ur169Xr16XXbZZaXcEIDSq3bXyhQ1c+bM
O++88+67777jjjvuvffeMuxh27ZtGRkZt91222OPPfab3/wmIyPj17/+9YQJE7Zs2fLJJ59Mnjx515MaO20yefLkQw455I477khK
Svr666//9re/Ffc5l32RlJSUmZkZu9Q3Jydnzpw5gwcP/u677+Ir7HHkr7zyyrBhw0q+4PeFF1644oorbr311rp1665cuTIzM/Pv
f/976T/BFLd27dqXXnrpyiuv3OOaEydOTExMvPbaazMyMnJzc6dOnfrkk0/uds2CgoLbb789PT29f//+t9xyS0JCwvLly2fNmhWf
JinuGcjOzp41a9bQoUPbtWuXn58/f/784cOHxzZ5/fXXu3bt+txzz6WkpDz33HM7XQ48efLkdevWXXbZZbfffvu2bdsWL1780ksv
7fY6m9L82uzW+PHjk5OThw8fnpKSsnTp0ieffHK3H+ACYF8kpKWlVdCuGzZsWEF7rrZuu+22Zs2aVUTr7NFVV101cODAiy66qAxp
Ur6q8EkA9m8V8ZWh7Aeq7wmm4NSvX//oo48u4fxOxWnZsuXAgQP36ktiKkgVPgkAHJikTPk49thjJ0yYsGrVqsr/7pCnn3762Wef
/eyzz8aNG1fJh95JFT4JABywnGACIAxOMLFbZmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAJWvVLmuuuuy8zM
fPDBB4suvP/++x999NHY7WHDho0ePbrknWRkZIwaNaq4e4vuoYQ1SzhQbJA7faFtYmLiq6++mpmZWZo/sri3YwYAilO9UiaKooKC
gmOPPbZ79+5VPZCSbN68uXnz5h07dowv6dmzZ506dXbs2FFpY7j11lvjhQcAB6yaVT2Ana1bt27x4sXXXXfd9ddfv+u9+z51Uco9
lLxafn5+VlZW3759FyxYEFvSt2/f6dOnn3322fs4PABgr1S7WZkoisaMGdOhQ4eTTz5517uKnvfp2LHjo48+OmXKlNdee+3uu+8+
+OCD46vVqlXr6quvfuGFF955550nnnji6KOP3u0eimrQoMHo0aOfeOKJRo0albBa3FtvvXX66afXqlUriqLk5OSTTz75rbfeit87
ePDgzMzMzMzM995779lnn73ggguKbpuRkfHII49ce+21r7zyyvTp09u0aVPCSPr06fPkk09OnTr1jTfeuOWWW+rUqRMb3nnnndet
W7fYUa655poShgoA+7HqmDJfffXVjBkzhgwZkphY7PASExNHjhy5aNGiSy+99PLLL582bVp6enr83u7du6empt54440XX3xxTk7O
vffeW7NmSfNPhx566J///Oe1a9fedNNNa9euLc0gP/roox07dpx00klRFJ1++umrV6/OysqK3/vMM8/07t27d+/e6enpY8aMueGG
G/r06VN08+7duzds2HDIkCGnnXbasmXLihvJeeedd/fdd8+YMWPAgAE33nhj586db7nlliiKRo0aNXny5Llz58aO8tRTT5VmzACw
/6mOKRNF0ZNPPnnYYYedeeaZxa2QkpJSv3792bNn5+XlrV+/fsaMGUXfznNych599NE1a9bk5eU9//zzTZs2bdmyZXG76tq16+OP
P/7hhx+OGDFiy5YtpRzhjh073nvvvb59+0ZR1Ldv37fffruwsHDX1fLz82fOnPn222+fddZZRZevWbPmwQcfXLVqVQkjSUhIGDJk
SGZm5tixY/Py8rKzs0eNGnXOOee0aNGilIMEgP1etbtWJubbb7996623/uM//mPatGm7XSEvL2/KlCn33nvv9OnT582b98knn6xc
uTJ+7/fff19QUBC7vW7duiiKGjRosNv9tGvX7sEHH3z66af/9re/7e0gp0yZMmbMmI4dO3bq1On+++8velezZs2uueaaHj16NG7c
uEaNGlEULVq0qOgK33zzzfbt20seSfPmzRs3bjxr1qz4OgsXLiwoKGjfvv2KFSv2drQAsF+qprMyURQ988wzTZs2Pf/884tb4YEH
HvjlL3/5zTff9OnT5+9///vgwYPjd8U7Zo9ycnK+/vrrM888M3Zhyl5ZtGjRN998c+eddy5YsOCHH34oetf9999/0EEH/fa3vz33
3HN79+49bty4nc5wbdu2bY8jiZ1fu+222zL/1/Tp02vWrNm0adO9HSoA7K+qb8r8+OOPr7322pVXXpmUlFTcOosWLXrxxReHDx/+
2GOP/fznP49dhLtXNmzYcMstt+Tm5o4ePbp58+Z7u/mUKVMOPvjgohf8RlGUlJR0xBFHvPzyy0uXLs3Pz4+iqFOnTmUYSU5OTm5u
7siRI3v/q4kTJ0ZRtH379hKuJQKAA0S1fi98/vnn69Wrd8wxx+x6V9u2bUeMGNGxY8ekpKRGjRp17tw5Jydn16mO0ti8efPw4cOX
LVs2evToVq1a7dW248aN692796RJk4ouzM/Pz8nJ6du3b8OGDRs1anTDDTcU/Qaa0o+koKDgiSeeuP7660899dTU1NTk5OSuXbve
d999sVmZ5cuXH3rooU2aNIlvnp6enpmZubcPAQCCVq1TZu3atS+99NJu5x6ys7NnzZo1dOjQSZMmPfvss3Xq1Bk+fHiZD7Rt27aM
jIzPPvvssccea9++/T4M+X9kZGQ0adJkwoQJY8eObd68+eTJk8s2ksmTJ48cOXLQoEETJ04cP378f/zHf0yZMiV2VdDrr7++cOHC
5557zoexATiQJaSlpVXQrhs2bFhBewbgAJSbm1vVQ6A6qtazMgAAJZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABKwC
U8Z3GQFQXrynUByzMgBAwKQMABCwik0Z84EA7DvvJpTArAwAELAKTxkpDcC+8D5CySpjVsZvIQBl4x2EPXKCCQAIWCWljKwGYG95
76A0albakWK/kQ0bNqy0IwIQKBFD6VX2CSa/nQCUzDsFe6XyZmXiTM8AsFsihjKogpSJif++ahqAA5yCYV9UWcrE+Q0GAMrMh7EB
gIBJGQAgYFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACVpV/TrJ169ZVeHTggJKdnV3VQwAq
RBWkTLxgvl+/pvKPDhyY4v/yaBrYzySkpaVV5vFat26tYIAqdEhqYzUD+5PKS5nYfxLpGKDKHZLaODI9A/uLSkoZkzFAdWN6BvYP
lfEJJh0DVEPfr1/jwwewH/BhbAAgYBWeMqZkgGrLxAzsByo2ZXQMUM2pGQidE0wAQMCkDAAQsApMGWeXgCA4xwRBMysDAARMygAA
AZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoA
AAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzK
AAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARM
ygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAE
TMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMA
BEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkD
AARMygAAAatZ1QOgsiUkJHRo17lD+85NGjWrUaPG+o15K1f/NP+rOd+t+Kaqh1Y+mjZu/rMLBz/38pNr1q3e6a4zTux7UKMm4yeP
rZKBAVARpMyBJSEh4fzTL2rRrOWHn87I/m5x/tb8RvUP6t7x2IF9L3tq/OgNG9dX9QCjKIrOOuncBvUbTZjyQlUPBIAASJkDyzEd
e7Y+5PAXX3925ZofY0tWrf3p3Q/fXP7j94WFhVU7tkowddZbVT0EAMqZlDmwdD2qe/b3S+IdE7dgUVbRH49oc9SxnU84qGHTbdu3
fr3sy8x/Tt2+Y3sURX17n59cL2X5j98f2bZDSr2UVWtXTf/HOzkrV5Ryw5yVyzu061yvbvLY155u3/rIE7qdHEXR9h3b163PnfvF
J/O/mhtF0Rkn9u14RNcoin49+LYoiv45b9aszzJL2HPMCd1P7nxk96Q6SavWrPziXx9LUUVPMPXtfX5KcuqqNT+1PuTwlOTUVWtW
Tpv9VuMGB/XsemKD1EZ5G9Zlfjx12XdL4vvf7WijKEpISOjV/ZROR3St879H79Pr7PjprZKHDcC+kzIHkKQ6deunNvz863klr9bx
iK5nnXTuzE9nvPzVuHp1k/uddsHpvc5558PJsXsPbXHYmtxVL705dkdBwem9zknvc9HTEx4vKCgozYbr1ue+8PpfN27aEEXR6rUr
P5rzYRRFtWrWOjStdd/e52/Zmv/1si+nznorMSFhpxNMJe+5R6eex3Q87q3MN75f8U3Lgw/t2/uCUj4hhxzc6qfVP46b/FyNGjXP
P/2igX0vW5O7+o33X12/Me+Ebif1O/XCv7z0eP6WzVEUfTTnw92ONoqiHp2O73p0j7cyX/8h59u05oece+r/Hb3kYQNQLnyC6QBS
p3adKIq2bNsSX9Kz64m/Hnxb7H+djugaW3jiMb0Xf/PVx1mz87dsXpO7auqst49u16l+asPYvXkb1k3/x3sbN2/M37L546xZKcmp
Des3Ks2GGzdvnDb77VjHFLVt+7al3y5auOTzI9t2LG7kJe/5uC69Pvv846XfLtq6beuy75Z8Ov+jUj4heetzM/85dXP+5g0b18//
al6d2knTZr+zJnfVtm1bP5n/j1q1ajdvcvAeR3ts5xM+W/DPZd8t3rpta/b3Sz/O+r+jlzxsAMqFWZkDyJatW6IoqlOrTnzJP+fN
+ue8WXWT6l532a9iS+qnNEiumxw/sRJFUc7K5YWFhc0aN8tbnxtFUW7e2vhVNZu3bI6iKKlO3dJsuHbd6tjkTUxqcv0Tj+l9aIvD
6tVNTkxMjKJo19NepRlScr2UpDp1V6z8IX7v8p9+2N1udmNt3toiT05+bJBFf0yqnVTyaJPrpSTVScpZubzo2Eoz7FKOEIA9kjIH
kPwtm/PW5x7cNK2EdRISEqIoOuvkfmed3K/o8pR6qbEbxV0dvMcNd+zYUXT5+WdctDl/86vvjl+3Pnfb9m29e55+WMu2ZdtzJSj9
aOOqw7ABDgRS5sAy76s5Jx1zapNGTVetXbnbFfI2rNucv+nDT2cs2NMlNfuyYa2atZoddPCk9ybEh9Gi2SHxe3cUFCQm/N+pz5L3
vHHThvwtm1s0bfnND8tiS9Katdyrke/LaDdu2rBla37zJi3iR4+XYpmfSQD2imtlDiyfff7P7O+XXNT3sg7tO9dLqlcjsUb9lAad
j+weRVFssqWwsPDDT2ec3OO0dq2PrFM7qU7tOi0PPvT8My7a41zCXm24bfu2vA3rOrTrXDepXr2keqcc16dFkbmidRtyGzZonFwv
pZR7/jjro2M6Hdf20Ha1a9Vuc+jhPTofXw7PVKlH+3HWRz069WxzyOG1atU+rGWbYzufUIYnBIAyMytzYCksLHx92ssd2nfp1L7r
acefWSOx5oZN69fmrZkyY9Li7K9i6yz4el5+/qYenY/ve0r6joIdK9f8NGfBxxs27fnb8/ZqwzemvdKn19lXD7phx47t3y7PXrAo
Kz6f8flX8w5pfuiVA66pUzsp9mHskvf82YJ/1q5V+8yT+9WpXWf12pWzPs3s0+vs8ni2SjXaTz//R+3adc465bzY0T/OmnXKcadv
L9ixt08IAGWTkJZW0pUT+6J169bfr19TQTuH6qnLUd179zzj8bEP7SjYsee1qTYOSW2cnZ1d1aMAysIJJtgnBzVqeuIxveunNoyd
3jqh28lfLJqvYwAqjRNMsE/W5K7atn3bgLMHpSbXz9uwbu7CTz+d/4+qHhTAAUTKwD4pLCz8OGv2x1mzq3ogAAcoJ5gAgIBJGQAg
YFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkA
IGBSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAImZQCAgEkZ
ACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACJmUAgIBJ
GQAgYFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICA
SRkAIGBSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAImZQCA
gEkZACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAImZQCAgEkZACBgUgYACFgFpkx2dvYh
qY0rbv8A5eKQ1MbZ2dlVPQqgjMzKAAABkzIAQMAqNmWcYwKqOWeXIHQVPiujZoBqS8fAfsAJJgAgYJWRMiZmgGrIlAzsHxLS0tIq
50itW7eOouj79Wsq53AAxYn9x5WOgf1D5aVMTOvWrdUMUIVMxsB+prJTJvrf6ZnIDA1QieKnuXUM7GeqIGXi4k0DUNEUDOyvqjJl
AAD2kQ9jAwABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQ
MCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwA
EDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQM
ABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCk
DAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDA
pAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBA
wKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIA
QMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMy
AEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGT
MgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAAB
kzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAA
AZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoA
AAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzK
AAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARM
ygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAE
TMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMA
BEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkD
AARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDAp
AwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAw
KQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQ
MCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwA
EDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQM
ABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCk
DAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDA
pAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBA
wKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIA
QMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMy
AEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGT
MgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAAB
kzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAA
AZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoA
AAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzK
AAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARM
ygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAE
TMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMA
BEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkD
AARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDAp
AwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAw
KQMABEzKAAABkzIAQMCkDAAQMCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwKQMABAwKQMABEzKAAABkzIAQMCkDAAQ
MCkDAARMygAAAZMyAEDApAwAEDApAwAETMoAAAGTMgBAwGq2bt26qscAAFBGCWlpaVU9BgCAMnKCCQAImJQBAAImZQCAgEkZACBg
UgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAJWs/IP2bp168o/KABQObKzsyvzcJWXMvGCyc3N
rbSDAgCVLP6OXzlNk5CWllbRx4g9JAUDAAeUhg0bRhUfNBWeMq1btxYxAHDAatiwYYXWTMVe9qtjAOAAl5ubW6GXyVZgyugYACCq
4JqpqJTRMQBAXMXVTIWkjI4BAHZSQTVT/imjYwCA3aqIminnlNExAEAJyr1m/OECACBg5ZkypmQAgD0q34kZszIAQMCkDAAQsHJL
GWeXAIBSKsdzTGZlAICASRkAIGBSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkA
IGBSBgAImJQBAAJWs6oHsHvXXDO8uLueeuoPlTkSANhfLf/PS4q7K+2BlypzJPsiIS0trVx21Lp169zc3H3cSQkFsytNAwBlUELB
7KrimqZhw4bZ2dn7vp/qkjI7RUwJmVL6NQGAovYqYoqqiKDZr1KmaJ2UMk3KsAkAHMjK3DEx5V4z+0/KxKOkDEWyL9sCwIFjHzsm
pnxrZj9JmViL7GOIlMtOAGB/VS4dE1OONVNeKVOVH8YurwSJ7WGvLhkGgANEOXZMue+tXFRZypTvVIqaAYBdVUR5VLeaqZqUqYhT
QmoGAIqquOaoVjVTBSmzt7VxzTXD92oTNQMAFV0b1admquwEU0VcpevKXwDYo7WX31XVQyhPlZ0ylTNlYmIGgAPZHqdM1l5+174H
TTWZmKmaWZmKmz7Zlz0fd9yp11wz/JJLhhRdmJCQ8LOfDb3mmuHdu/fa59FVgXJ/UKedlt63b7G/uyeffPb55/+sDONMSEi47LJf
XHPN8IYNDyq6/LjjTr3qqpvLsEMA9qhcgqbK+cvY/2Lbtq2pqfWbNfu/79o55JA2NWvWKigoqMJR7aN9eVCnnNL3vPMuq8jR/Y+0
tMOSk1M3b954xBGdKuFwAMSFHjSVmjKV8112+/JRpu3bt33zzZL27f/v3bR9+05Ll35ZWFhYbuOrdJX5oD788J3XX/9bGTY84ojO
OTnfffnlvHbtOiYkJJT7wAAOHGU771O2mqkO55hqVvUAqp1Fiz4/9dR+H300dceOHbVr1znssHZTprzUrl3Houu0bXtUly49GzVq
un371qVLv/zHP97fvn17FEXNmqUdd9ypTZo03759+4oV33788Yz169eVsPyYY0465piToijasWN7Xl7uggWfffnl3NghEhISevQ4
+cgju9Spk7Rmzcqvv55/4olnTZjwl9zc1SWPYW8fVNu2R512WvqLLz6+efPG2MpNmhzcv/+Vr7/+t/btOx55ZJfof7tw7tzZn3zy
QRRFNWrU6NHj5LZtj05OTlm7dtWsWe+tXLkitu3JJ5/dqFHTeM20bXtU587HNWrUdMuWzYsWLZgzZ+aOHTt2HV7t2nVat24/a9Z7
K1Z81737iYcc0va775aU6dUDoOxiNdNo7F1VPI69JGV29v33SwsLC1q1ards2Vdt2x61adPGnJzviq5w5JFdTjml78cfZ06Z8lLd
uvVOP/2CE088OzPzzYSEhHPOGbho0YL33nsliqK0tFZHHNHl008/KG55FEWffTbzs89mRlFUs2atli0PO+209K1b85cu/TKKoi5d
enbocMz06W+sWPHdwQcfctpp6aUZQxkeVHb21/n5m446qsucObNjSzp06L5mzcoff/zhxx9/SEhIrF+/0eTJLxbdW4sWrdauXfXG
Gy8UFBScdNJZZ57Zf9y4J3Y9XXX00d169Trzn/+cvmTJy4mJCe3bd2rRotX33y/bdXiHH94hihKWLftq69YtP/20/IgjOksZgKoS
XNC4VmZnBQUFixcvjJ2Oad++0+LFC3Za4dhjT8nO/nrevI+2bNmcm7v6ww/fbt++Y2pqg9q169SpU/e775Zu2ZK/ZUv+smVfx3ql
uOVFbd++7ZtvFi9atODww4+OLenS5fj58z/+9tsl27Zt/e67pVlZ/yjNGMrwoAoKCr78ct5RR3WLndapXbtO27ZHLVw4p4SnaP36
dbNnT928eeOWLZvnzv0oOTm1fv1GO60Tm1X68su5n3/+yebNGzdu3DB37ke77Zgoio44otO33y7eunVLFEWLFi047LDD69SpW8IA
AKhoAV1AU11mZfZ4aUtxK1TElTeLFn1+4YVXNGuW1rx5yxkz/mWqIzW1Qd26yd9++39zBitXrigsLDzooGbZ2Yu+/vrzM8/sv2zZ
Vzk53/3wwzcbN66PomjLlvzdLo+iKDk59dhje6eltapbNzkxMTGKotWrf4qiqF69lDp1kuJnbaIo+umn5aUZQ+y81d4+qC+/nNet
W69WrQ7/5pvFRxzRqbCwcPHiL0p4fvLy1savs9myZXMURUlJO5dHSkr9pKR6y5d/W8J+Yho1atK0aYu5c/9nTmjp0oW9ep3erl2H
BQs+3eO2AFSoIGZoqkvKVCurV/+Ym7v69NMv+PHHH/Ly1v7rnQlRFPXufW7v3ucWXZqcnBpFUWbmmwsWfNKyZes2bY466aRz5s37
KHb+qLjlZ5110ZYtm99+e0JeXu727duOP77PIYe0ie/zXy/LTdjpdnFjKMOD2rRpQ3b210cf3f2bbxYfdVS3xYsXbNu2tYRdle56
4diA97zmEUd0jqLorLMu2mmhlAGoctU8YmKqS8qUMLlSOZ972smiRZ8ff3yfuXM/2mn5hg3r8vM3ffxx5ldfZe12w9Wrf1q9+qes
rH926NC9V68z5879qKBgx26XJyYmNmnS/J13Xl6zZmVs2+bNW8ZubNq0YcuW/GbNWvzwQ3ZsSbNmLUo/hr19UFEUffHFZ+np/37k
kV0aNjxo2rRJ8eUFBQVl+zxRbJBpaYdlZy8qYbXExMR27TrErxmKad685fnn/+ygg5rF5qgAqHxBREyMa2V2b/78j5966g/xzxPF
FRYWfvxx5nHHndqmzRF16iTVrl3n4IMPPeusAcnJKY0bN+3TJ71Zs7SaNWvVrZvcvPkhGzasKyjYUdzy7du3bdiQ1759p6SkenXr
1jv++NOKfvVLVtY/OnU67tBD29aqVfuQQ9p06XJ8acZQtgcVRVFOzvdr1qw86aSzfvzxh3haRVG0fn1ugwaN69Xbw553VVhY+Mkn
Hx51VLeOHXskJdVLTk7p1u2EonNOMYce2rZu3eRly74quvDHH3/YuHFDbLZmJ0ce2SX+NXpFb+/6IwBl02jsXQF1TFR9ZmUC8tVX
Wfn5m7t06Xnqqeft2LFjzZqfPv/8040bN2zatPHbb5ccf3yfgw5qtn37th9//OGttyZEUbR27ardLo+i6L33Xj3xxDMvu+z67du3
L1/+zddfz2/a9H9mX7Ky/lmrVp3evfvVqVNnzZqVc+fOPv74Pjt2bC95DPvyuL74Ys7JJ5+90wW/X3457+CDD7344qtr164T/zB2
KX355dxt27Z07tyzZ8/TNm/euHjxghUrvttpnSOO6Jybu3rt2lU7Lc/O/rpduw7/+Mf0Mj0UAMoirIKJS0hLS9vzWqXQunXr3Nzc
Pa5WhrNFe7tJlZyQqmhHH939hBP6PPvsI7HTVRWhe/cTO3bs8eKL/2+3X/0CQEBK/ua63X40qWwdk/bAS2XYKqZhw4bZ2dll3jzO
rEw11ahRk7Ztj/7666z8/M0HH3zoMcec+PXX8yuuY+rXb9SxY4/58z/WMQAHmkAnY+KqJmWuuWZ4Bc2a7Dd/Ezs3d/X27dv69r0k
Obn+hg3rvvjis3nz/llBx7rooqsaNGj8zTeL58//uIIOAUA1FHrExFR2yjz11B8qoTb2g7NLhYWF8+Z9NG/ebj5tVO5efvmvlXAU
ACpN2gMvlXyOqVwiZl/OLpWjKvsEU0UEzX4zJQMAFWf/mIyJq4KU2dspk6ee+sNebbIfTMkAwD6q6CmTajIlE1XVrEysNsp3EmW/
/OASAJRZxdVG9emYqApPMJVvzegYANhVRTRHteqYqGq/7be8akbHAEBxyrc8qlvHRJX/FXm7iqdMGVpkX7YFgANHyR9oKqXy7Zjy
+oq8qk+Z6F8nZkoZJWXYBAAOZPtYM+U+H7NfpUy0y2mmPf6h7NKsCQAUVeaaqYjzSvtbysTs1XUzIgYAymCvgqbiLo7ZP1MmroSm
UTAAUC5KaJpKuLx3P08ZAGD/Vl4pU5UfxgYA2EdSBgAImJQBAAImZQCAgEkZACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEA
AiZlAICASRkAIGBSBgAImJQBAAImZQCAgJVbymRnZzds2LC89gYA7McaNmyYnZ1dLrsyKwMABEzKAAABK8+UcY4JANijcjy7FJmV
AQCCVs4pY2IGAChB+U7JRBUxK6NmAIDdKveOiSroBJOaAQB2UhEdE1XctTJqBgCIq6COiSr0sl81AwBEFdkxUUV/gknNAMABrkI7
JoqihLS0tIrbe0zr1q2jKMrNza3oAwEA1UdsOqNCOyaqnJSJiQVNpGkAYL8WPyFT0RETU3kpExdvGgBg/1M5BRNXBSkDAFBe/OEC
ACBgUgYACJiUAQACJmUAgIBJGQAgYFIGAAiYlAEAAiZlAICASRkAIGBSBgAImJQBAAImZQCAgEkZACBgUgYACNj/B9YDA4uEBdfL
AAAAAElFTkSuQmCC
""".replace("\n", "")

_ICON_CACHE = {}

def _get_icon(size):
    if size not in _ICON_CACHE:
        _ICON_CACHE[size] = _make_mythic_icon_png(size)
    return _ICON_CACHE[size]


def _make_mythic_badge_png(size=96):
    """Generate the small monochrome 'status bar' badge used by Web Push
    notifications. This MUST be a white silhouette on a fully transparent
    background — no colored/opaque backdrop.

    Android reads only the ALPHA channel of the 'badge' image: every pixel
    with any opacity is treated as "part of the icon" and gets flattened
    into a solid tinted shape. The old code pointed /badge.png at the same
    PNG used for the app icon (/icon.png), which has an opaque teal
    rounded-square background — so Android saw one big solid opaque
    rectangle and rendered a plain filled box in the status bar instead of
    the "M" mark. Drawing just the white line-art on a transparent canvas
    (no fill_rect/circle_aa background pass) fixes that."""
    import struct, zlib

    W = H = size
    img = bytearray(W * H * 4)  # fully transparent (all zeros) to start

    def set_pixel(x, y, a=255):
        if 0 <= x < W and 0 <= y < H:
            i = (y * W + x) * 4
            img[i], img[i+1], img[i+2], img[i+3] = 255, 255, 255, a

    s = size / 40
    pts = [
        (int(10*s), int(28*s)),
        (int(10*s), int(12*s)),
        (int(20*s), int(22*s)),
        (int(30*s), int(12*s)),
        (int(30*s), int(28*s)),
    ]
    lw = max(2, size // 12)  # slightly thicker than the app icon's mark, reads better at tiny status-bar sizes

    def draw_line(x0, y0, x1, y1):
        dx, dy = x1-x0, y1-y0
        steps = max(abs(dx), abs(dy), 1)
        for i in range(steps+1):
            x = int(x0 + dx*i/steps)
            y = int(y0 + dy*i/steps)
            for ox in range(-lw//2, lw//2+1):
                for oy in range(-lw//2, lw//2+1):
                    set_pixel(x+ox, y+oy)

    for i in range(len(pts)-1):
        draw_line(pts[i][0], pts[i][1], pts[i+1][0], pts[i+1][1])

    def png_chunk(name, data):
        crc = zlib.crc32(name + data) & 0xffffffff
        return struct.pack('>I', len(data)) + name + data + struct.pack('>I', crc)

    raw_rows = b''
    for y in range(H):
        raw_rows += b'\x00'
        raw_rows += bytes(img[y*W*4:(y+1)*W*4])

    ihdr = struct.pack('>II', W, H) + bytes([8, 6, 0, 0, 0])
    compressed = zlib.compress(raw_rows, 9)

    png  = b'\x89PNG\r\n\x1a\n'
    png += png_chunk(b'IHDR', ihdr)
    png += png_chunk(b'IDAT', compressed)
    png += png_chunk(b'IEND', b'')
    return png


_BADGE_CACHE = {}

def _get_badge(size):
    if size not in _BADGE_CACHE:
        _BADGE_CACHE[size] = _make_mythic_badge_png(size)
    return _BADGE_CACHE[size]


@app.route("/icon.png")
def pwa_icon_192():
    return Response(_get_icon(192), mimetype="image/png",
                    headers={"Cache-Control": "public, max-age=604800"})

@app.route("/icon-96.png")
def pwa_icon_96():
    return Response(_get_icon(96), mimetype="image/png",
                    headers={"Cache-Control": "public, max-age=604800"})

@app.route("/icon-512.png")
def pwa_icon_512():
    return Response(_get_icon(512), mimetype="image/png",
                    headers={"Cache-Control": "public, max-age=604800"})


def _make_ico_from_png(png_bytes: bytes, size: int) -> bytes:
    """Wraps a PNG in a minimal valid .ico container. Since Windows Vista,
    ICO entries are allowed to hold raw PNG data directly (instead of the
    old uncompressed BMP format) — this is exactly what real .ico files
    made by icon editors do for modern sizes, so Explorer/shortcuts render
    it correctly. Needed because Windows .url shortcuts silently fail to
    show a custom icon when IconFile points at a plain .png — it wants a
    real .ico container."""
    import struct
    count = 1
    # ICONDIR: reserved(2)=0, type(2)=1 (icon), count(2)
    header = struct.pack("<HHH", 0, 1, count)
    # ICONDIRENTRY: width, height (0 means 256), color count, reserved,
    # planes, bitcount, bytes in resource, offset of resource data
    w = size if size < 256 else 0
    h = size if size < 256 else 0
    entry = struct.pack("<BBBBHHII", w, h, 0, 0, 1, 32, len(png_bytes), 6 + 16 * count)
    return header + entry + png_bytes


_ICO_CACHE = None

def _get_ico():
    global _ICO_CACHE
    if _ICO_CACHE is None:
        _ICO_CACHE = _make_ico_from_png(_get_icon(256), 256)
    return _ICO_CACHE


@app.route("/icon.ico")
def pwa_icon_ico():
    """A real .ico file (not a renamed PNG) for use as a Windows shortcut/
    .url icon — see _make_ico_from_png for why this is necessary."""
    return Response(_get_ico(), mimetype="image/x-icon",
                    headers={"Cache-Control": "public, max-age=604800"})

@app.route("/badge.png")
def pwa_badge():
    """Badge icon for push notifications (status bar). Must be a white-on-
    transparent silhouette, NOT the colored app icon — see
    _make_mythic_badge_png() for why reusing /icon.png here caused a plain
    solid box to show up in notifications instead of the logo."""
    return Response(_get_badge(96), mimetype="image/png",
                    headers={"Cache-Control": "public, max-age=604800"})

@app.route("/favicon.ico")
def favicon():
    return Response(_get_icon(192), mimetype="image/png",
                    headers={"Cache-Control": "public, max-age=604800"})


_SCREENSHOT_WIDE_BYTES   = base64.b64decode(_SCREENSHOT_WIDE_B64)
_SCREENSHOT_NARROW_BYTES = base64.b64decode(_SCREENSHOT_NARROW_B64)

@app.route("/screenshot-wide.png")
def pwa_screenshot_wide():
    """Desktop-form-factor preview used by manifest.json's 'screenshots'
    array, so Chrome can show its richer install UI on desktop."""
    return Response(_SCREENSHOT_WIDE_BYTES, mimetype="image/png",
                    headers={"Cache-Control": "public, max-age=604800"})

@app.route("/screenshot-narrow.png")
def pwa_screenshot_narrow():
    """Mobile-form-factor preview used by manifest.json's 'screenshots'
    array, so Chrome can show its richer install UI on mobile."""
    return Response(_SCREENSHOT_NARROW_BYTES, mimetype="image/png",
                    headers={"Cache-Control": "public, max-age=604800"})


def render_page():
    """Serves PAGE with the canonical-URL placeholders filled in — shared by
    every route that returns the main app shell (/, /invite/<token>,
    /legacy-invite/<code>) so SEO tags never leak the literal
    __CANONICAL_URL__ placeholder text.

    Canonical/OG/JSON-LD URLs always point at PREFERRED_PUBLIC_ORIGIN
    (Vercel), even when the page is actually being served from Render —
    that's what tells Google "this Render copy and the Vercel copy are the
    same page, treat the Vercel one as canonical." On top of that, Render
    itself gets a hard noindex,nofollow so it's excluded from search
    entirely rather than just de-prioritized."""
    origin = PREFERRED_PUBLIC_ORIGIN
    html = PAGE.replace("__CANONICAL_URL__", origin + "/").replace("__CANONICAL_ORIGIN__", origin)
    if _is_deindexed_host():
        html = html.replace(
            '<meta name="robots" content="index, follow">',
            '<meta name="robots" content="noindex, nofollow">',
        )
    return Response(html, mimetype="text/html")


@app.route("/")
@login_required
def index():
    return render_page()


@app.route("/api/invite-link", methods=["GET"])
@login_required
def api_invite_link():
    # Each account gets its OWN unique, permanent link (…/invite/<token>) that
    # always logs back into THIS SPECIFIC account's private chats — not a
    # single link shared by everyone. Same token every time for the same
    # account (persisted — see get_or_create_account_token), so it never
    # changes across reloads/redeploys.
    token = get_or_create_account_token(current_username())
    return jsonify({"invite_url": get_public_origin() + "/invite/" + token})


@app.route("/invite/<token>")
def account_link_landing(token):
    # Opens the ONE specific account this token belongs to — created the
    # first time that account visited /api/invite-link. Unknown/invalid
    # tokens fall through to the login page rather than erroring.
    user_id = resolve_account_token(token)
    if user_id:
        session["user_id"] = user_id
        session["authenticated"] = True
        session.permanent = True
        return render_page()
    return redirect("/login")


@app.route("/legacy-invite/<code>")
def invite_landing(code):
    # Legacy shared-account link (everyone who opens it lands in the SAME
    # owner account) — kept working for any old links already handed out,
    # but no longer surfaced anywhere in the UI. The 🔗 button now issues
    # the per-account /invite/<token> links above instead.
    session["user_id"] = get_or_create_owner_id()
    session["authenticated"] = True
    session.permanent = True
    return render_page()


# --- Claim owner status (needed on serverless hosts like Vercel) -------------
# On Vercel, each request can land on a different instance with its own
# ephemeral /tmp, so a locally-written owner_user_id.txt does NOT reliably
# make you "the owner" on every subsequent request — see get_or_create_owner_id
# above. If OWNER_USER_ID is set (a fixed id you choose, e.g. a UUID), that
# value is always returned as the owner id, no file/DB needed. This route
# lets YOUR browser's session get pinned to that exact id, once, so every
# future request from you compares equal to it. Protect it with OWNER_SECRET
# (also an env var) so nobody else can claim it.
#
# Setup on Vercel:
#   1. Generate two random values locally:
#        python -c "import secrets,uuid; print('OWNER_USER_ID=', uuid.uuid4()); print('OWNER_SECRET=', secrets.token_hex(16))"
#   2. Set OWNER_USER_ID, OWNER_SECRET, and FLASK_SECRET_KEY as environment
#      variables in your Vercel project settings, then redeploy.
#   3. Visit https://<your-app>.vercel.app/claim-owner/<OWNER_SECRET> once,
#      in the browser you want to use as the owner. You only need to do
#      this again if you ever clear cookies or switch browsers/devices.
@app.route("/claim-owner/<secret>")
def claim_owner(secret):
    owner_secret = _os.environ.get("OWNER_SECRET", "").strip()
    if not owner_secret:
        return Response(
            "OWNER_SECRET is not set on the server, so no one can claim "
            "ownership this way. Set OWNER_USER_ID, OWNER_SECRET, and "
            "FLASK_SECRET_KEY as environment variables first, then reload "
            "this page.", mimetype="text/plain"), 400
    if not secrets.compare_digest(secret, owner_secret):
        return Response("Incorrect secret.", mimetype="text/plain"), 403
    if not _OWNER_USER_ID_ENV:
        return Response(
            "OWNER_SECRET is set but OWNER_USER_ID is not. Set OWNER_USER_ID "
            "too (any fixed string, e.g. a UUID), then reload this page.",
            mimetype="text/plain"), 400
    session["user_id"] = _OWNER_USER_ID_ENV
    session["authenticated"] = True
    session.permanent = True
    return Response(
        "You're now recognized as the account owner on this browser. "
        "<a href='/api-usage'>Go to API keys →</a>",
        mimetype="text/html")


# --- API key management (open to everyone — no owner gate) --------------------
@app.route("/api/keys", methods=["GET"])
@login_required
def api_keys_list():
    return jsonify({"keys": list_api_keys(current_username())})

def _fmt_dt(iso_str):
    if not iso_str:
        return "Never"
    try:
        dt = datetime.datetime.fromisoformat(iso_str.replace("Z", "+00:00"))
        return dt.strftime("%b %-d, %Y, %-I:%M %p")
    except Exception:
        return iso_str

@app.route("/api-usage")
@login_required
def api_usage_page():
    html = """<!DOCTYPE html>
<html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1">
<title>API Keys · Mythic AI</title>
<style>
  * { box-sizing:border-box; }
  body { font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;
          background:#0f1115; color:#f2f2f2; margin:0; padding:32px 24px 60px; }
  .wrap { max-width:1000px; margin:0 auto; }
  a.back { color:#9a9ea6; text-decoration:none; font-size:14px; display:inline-block; margin-bottom:20px; }
  a.back:hover { color:#fff; }
  .headrow { display:flex; justify-content:space-between; align-items:flex-start; flex-wrap:wrap; gap:14px; margin-bottom:8px; }
  h1 { font-size:30px; margin:0 0 6px; }
  .sub { color:#9a9ea6; font-size:14px; margin-bottom:30px; max-width:560px; line-height:1.5; }
  .gen-btn { background:#e8532a; color:#fff; border:none; border-radius:8px; padding:11px 18px;
             font-size:13px; font-weight:700; letter-spacing:.3px; cursor:pointer; white-space:nowrap; }
  .gen-btn:hover { background:#d1471f; }
  .totals { display:flex; gap:16px; margin-bottom:30px; flex-wrap:wrap; }
  .totals .box { flex:1; min-width:150px; background:#1a1d24; border:1px solid #2a2e37;
                   border-radius:14px; padding:18px; text-align:center; }
  .totals .num { font-size:40px; font-weight:800; line-height:1.1; }
  .totals .label { font-size:12px; color:#9a9ea6; margin-top:6px; letter-spacing:.3px; text-transform:uppercase; }
  table { width:100%; border-collapse:collapse; background:#1a1d24; border:1px solid #2a2e37; border-radius:14px; overflow:hidden; }
  thead th { text-align:left; font-size:11px; letter-spacing:.5px; text-transform:uppercase; color:#9a9ea6;
             padding:14px 16px; border-bottom:1px solid #2a2e37; }
  tbody td { padding:16px; border-bottom:1px solid #22252c; font-size:14px; vertical-align:middle; }
  tbody tr:last-child td { border-bottom:none; }
  tbody tr:hover { background:#20232b; }
  .name-cell { font-weight:700; font-size:15px; }
  .key-cell { font-family:monospace; color:#c7cad1; display:flex; align-items:center; gap:8px; }
  .copy-btn { background:none; border:1px solid #3a3e47; color:#c7cad1; cursor:pointer; font-size:11.5px;
              padding:4px 9px; border-radius:6px; font-weight:600; white-space:nowrap; }
  .copy-btn:hover { border-color:#7be3ab; color:#7be3ab; }
  .calls-cell { font-weight:700; font-size:18px; }
  .state-pill { font-size:11px; font-weight:700; letter-spacing:.4px; border:1px solid; border-radius:20px; padding:3px 10px; display:inline-block; }
  .state-active { color:#1a9e5c; border-color:#1a9e5c; }
  .state-revoked { color:#c0392b; border-color:#c0392b; }
  .revoke-btn { background:none; border:1px solid #3a3e47; color:#c0392b; border-radius:6px; padding:6px 10px;
                font-size:12px; cursor:pointer; }
  .revoke-btn:hover { background:#c0392b; color:#fff; border-color:#c0392b; }
  .rename-btn { background:none; border:1px solid #3a3e47; color:#9a9ea6; border-radius:6px; padding:6px 10px;
                font-size:12px; cursor:pointer; margin-right:6px; }
  .rename-btn:hover { border-color:#e8532a; color:#e8532a; }
  .options-cell { display:flex; gap:6px; flex-wrap:wrap; }
  .empty { color:#9a9ea6; font-size:15px; padding:50px 0; text-align:center; }
  .modal-overlay { display:none; position:fixed; inset:0; background:#000a; align-items:center; justify-content:center; z-index:50; }
  .modal { background:#1a1d24; border:1px solid #2a2e37; border-radius:14px; padding:26px; width:min(90vw,420px); }
  .modal h3 { margin:0 0 14px; font-size:18px; }
  .modal input { width:100%; padding:10px 12px; border-radius:8px; border:1px solid #3a3e47; background:#0f1115;
                 color:#fff; font-size:14px; margin-bottom:14px; }
  .modal-actions { display:flex; gap:10px; justify-content:flex-end; }
  .modal-actions button { border-radius:8px; padding:9px 16px; font-size:13px; cursor:pointer; border:none; }
  .btn-cancel { background:#2a2e37; color:#f2f2f2; }
  .btn-confirm { background:#e8532a; color:#fff; font-weight:700; }
  .new-key-box { background:#0f1115; border:1px solid #1a9e5c; border-radius:8px; padding:12px; margin-bottom:14px;
                 font-family:monospace; font-size:13px; word-break:break-all; color:#7be3ab; }
</style></head>
<body>
  <div class="wrap">
    <a class="back" href="/">← Back to chat</a>
    <div class="headrow">
      <div>
        <h1>API keys</h1>
        <div class="sub">Create and manage API keys for authenticating requests to Mythic AI. These keys allow programmatic access to your app.</div>
      </div>
      <button class="gen-btn" onclick="openCreateModal()">GENERATE API KEY</button>
    </div>
    <div class="totals" id="totals-row"></div>
    <table>
      <thead><tr>
        <th>Name</th><th>API Key</th><th>Created At</th><th>Calls</th><th>State</th><th>Options</th>
      </tr></thead>
      <tbody id="keys-tbody"></tbody>
    </table>
    <div class="empty" id="empty-msg" style="display:none;">No API keys yet. Click "Generate API Key" to create one.</div>
  </div>

  <div class="modal-overlay" id="create-overlay">
    <div class="modal">
      <h3>Generate API key</h3>
      <input type="text" id="create-label" placeholder="Key name (optional)" maxlength="100">
      <div id="new-key-result"></div>
      <div class="modal-actions">
        <button class="btn-cancel" onclick="closeCreateModal()">Close</button>
        <button class="btn-confirm" id="create-confirm-btn" onclick="doCreateKey()">Generate</button>
      </div>
    </div>
  </div>

<script>
function fmtDate(iso) {
  if (!iso) return 'Never';
  const d = new Date(iso);
  if (isNaN(d)) return iso;
  return d.toLocaleString(undefined, { month:'short', day:'numeric', year:'numeric', hour:'numeric', minute:'2-digit' });
}

async function loadKeys() {
  const res = await fetch('/api/keys');
  if (res.status === 403) {
    document.querySelector('.wrap').innerHTML = '<a class="back" href="/">← Back to chat</a><div class="empty">Only the account owner can view API keys.</div>';
    return;
  }
  const data = await res.json();
  const keys = data.keys || [];
  const totalsRow = document.getElementById('totals-row');
  const tbody = document.getElementById('keys-tbody');
  const emptyMsg = document.getElementById('empty-msg');

  const activeCount = keys.filter(k => k.active).length;
  const totalCalls = keys.reduce((s,k) => s + (k.request_count || 0), 0);
  totalsRow.innerHTML = `
    <div class="box"><div class="num">${activeCount}</div><div class="label">Active Keys</div></div>
    <div class="box"><div class="num">${totalCalls}</div><div class="label">Total Calls</div></div>
    <div class="box"><div class="num">${keys.length}</div><div class="label">Total Keys</div></div>`;

  if (!keys.length) {
    tbody.innerHTML = '';
    emptyMsg.style.display = 'block';
    return;
  }
  emptyMsg.style.display = 'none';

  tbody.innerHTML = keys.map(k => `
    <tr>
      <td class="name-cell" id="name-cell-${k.id}">${(k.label || '(unnamed key)').replace(/</g,'&lt;')}</td>
      <td><div class="key-cell"><span>${k.key_prefix || ''}</span>
        <button class="copy-btn" title="Copy key prefix" onclick="copyKeyPrefix('${(k.key_prefix||'').replace(/'/g,"")}', this)">📋 Copy</button></div></td>
      <td>${fmtDate(k.created_at)}</td>
      <td class="calls-cell">${k.request_count || 0}</td>
      <td><span class="state-pill ${k.active ? 'state-active' : 'state-revoked'}">${k.active ? 'ACTIVE' : 'REVOKED'}</span></td>
      <td><div class="options-cell">
        <button class="rename-btn" onclick="renameKey('${k.id}', ${JSON.stringify(k.label || '')})">✎ Rename</button>
        ${k.active ? `<button class="revoke-btn" onclick="revokeKey('${k.id}')">Revoke</button>` : ''}
      </div></td>
    </tr>`).join('');
}

function copyKeyPrefix(prefix, btn) {
  navigator.clipboard.writeText(prefix).then(() => {
    const orig = btn.textContent;
    btn.textContent = '✓ Copied';
    setTimeout(() => { btn.textContent = orig; }, 1200);
  });
}

async function renameKey(id, currentLabel) {
  const newLabel = prompt('New name for this key:', currentLabel || '');
  if (newLabel === null) return;
  const trimmed = newLabel.trim();
  if (!trimmed) { alert('Name cannot be empty.'); return; }
  const res = await fetch('/api/keys/' + id, {
    method: 'PATCH',
    headers: { 'Content-Type': 'application/json' },
    body: JSON.stringify({ label: trimmed }),
  });
  const data = await res.json();
  if (data.error) { alert(data.error); return; }
  loadKeys();
}

async function revokeKey(id) {
  if (!confirm('Revoke this key? Apps using it will stop working immediately.')) return;
  await fetch('/api/keys/' + id, { method: 'DELETE' });
  loadKeys();
}

function openCreateModal() {
  document.getElementById('create-label').value = '';
  document.getElementById('new-key-result').innerHTML = '';
  document.getElementById('create-overlay').style.display = 'flex';
}
function closeCreateModal() {
  document.getElementById('create-overlay').style.display = 'none';
  loadKeys();
}
async function doCreateKey() {
  const label = document.getElementById('create-label').value.trim();
  const btn = document.getElementById('create-confirm-btn');
  btn.disabled = true;
  try {
    const res = await fetch('/api/keys', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ label }),
    });
    const data = await res.json();
    if (data.api_key) {
      document.getElementById('new-key-result').innerHTML =
        '<div class="new-key-box" style="display:flex;gap:8px;align-items:center;">' +
          '<input type="text" id="standalone-new-key-input" readonly value="' + data.api_key.replace(/"/g, '&quot;') + '" ' +
            'style="flex:1;background:transparent;border:none;color:#7be3ab;font-family:monospace;font-size:13px;min-width:0;">' +
          '<button type="button" id="standalone-new-key-copy-btn" ' +
            'style="background:#1a9e5c;color:#04140b;border:none;padding:8px 14px;border-radius:6px;cursor:pointer;font-size:12px;font-weight:700;white-space:nowrap;">Copy</button>' +
        '</div>' +
        '<div style="font-size:12px;color:#9a9ea6;margin-bottom:10px;">Copy this now — it will not be shown again. (Use the Copy button — double-clicking to select may only grab part of the key due to the hyphens.)</div>';
      const stCopyBtn = document.getElementById('standalone-new-key-copy-btn');
      if (stCopyBtn) {
        stCopyBtn.addEventListener('click', async () => {
          const keyInput = document.getElementById('standalone-new-key-input');
          if (!keyInput) return;
          try {
            keyInput.select();
            document.execCommand('copy');
            stCopyBtn.textContent = '✓ Copied!';
            setTimeout(() => { stCopyBtn.textContent = 'Copy'; }, 2500);
          } catch (e) {
            try {
              await navigator.clipboard.writeText(keyInput.value);
              stCopyBtn.textContent = '✓ Copied!';
              setTimeout(() => { stCopyBtn.textContent = 'Copy'; }, 2500);
            } catch (e2) {
              alert('Copy failed. Click once inside the box, then Ctrl+A, Ctrl+C.');
            }
          }
        });
      }
      loadKeys();
    } else if (data.error) {
      alert(data.error);
    }
  } catch (e) {
    alert('Could not create key.');
  } finally {
    btn.disabled = false;
  }
}

loadKeys();
</script>
</body></html>"""
    return Response(html, mimetype="text/html")


@app.route("/analytics")
@login_required
def analytics_dashboard():
    """Comprehensive analytics dashboard for viewing usage stats, trends, and exporting conversations."""
    html = """<!DOCTYPE html>
<html>
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Analytics Dashboard · Mythic AI</title>
  <style>
    * { box-sizing: border-box; }
    html, body { margin: 0; padding: 0; }
    body { 
      font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
      background: #0f1115;
      color: #f2f2f2;
      padding: 32px 20px;
    }
    .wrap { max-width: 1200px; margin: 0 auto; }
    h1 { font-size: 28px; margin: 0 0 8px; font-weight: 700; }
    .subtitle { color: #9a9ea6; font-size: 14px; margin-bottom: 30px; }
    
    .metrics {
      display: grid;
      grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
      gap: 16px;
      margin-bottom: 40px;
    }
    .metric {
      background: #1a1d24;
      border: 1px solid #2a2e37;
      border-radius: 12px;
      padding: 20px;
    }
    .metric-value { font-size: 32px; font-weight: 800; margin-bottom: 6px; }
    .metric-label { font-size: 11px; color: #9a9ea6; text-transform: uppercase; }
    
    .tabs {
      display: flex;
      gap: 0;
      border-bottom: 1px solid #2a2e37;
      margin-bottom: 20px;
    }
    .tab-button {
      padding: 12px 18px;
      background: none;
      border: none;
      color: #9a9ea6;
      font-weight: 600;
      font-size: 14px;
      cursor: pointer;
      border-bottom: 2px solid transparent;
      transition: all 0.2s;
    }
    .tab-button:hover { color: #fff; }
    .tab-button.active { color: #fff; border-bottom-color: #e8532a; }
    
    .tab-content { display: none; }
    .tab-content.active { display: block; }
    
    .card {
      background: #1a1d24;
      border: 1px solid #2a2e37;
      border-radius: 12px;
      padding: 24px;
      margin-bottom: 16px;
    }
    
    .search-row {
      display: flex;
      gap: 10px;
      flex-wrap: wrap;
      margin-bottom: 20px;
    }
    .search-row input,
    .search-row select {
      padding: 10px 14px;
      border: 1px solid #3a3e47;
      background: #0f1115;
      color: #fff;
      border-radius: 8px;
      font-size: 13px;
      font-family: inherit;
    }
    .search-row button {
      padding: 10px 24px;
      background: #e8532a;
      color: #fff;
      border: none;
      border-radius: 8px;
      cursor: pointer;
      font-weight: 700;
      font-size: 13px;
      transition: all 0.2s;
    }
    .search-row button:hover { background: #d1471f; }
    
    .results {
      display: flex;
      flex-direction: column;
      gap: 10px;
      max-height: 500px;
      overflow-y: auto;
    }
    .result {
      background: #0f1115;
      border: 1px solid #2a2e37;
      border-radius: 8px;
      padding: 12px 14px;
      cursor: pointer;
      transition: all 0.2s;
    }
    .result:hover { border-color: #e8532a; background: #1a1d24; }
    .result-title { font-weight: 700; margin-bottom: 4px; }
    .result-meta { font-size: 12px; color: #9a9ea6; }
    
    .export-grid {
      display: grid;
      grid-template-columns: repeat(auto-fill, minmax(130px, 1fr));
      gap: 12px;
    }
    .export-item {
      background: #1a1d24;
      border: 1px solid #2a2e37;
      border-radius: 8px;
      padding: 12px;
      text-align: center;
      cursor: pointer;
      font-weight: 700;
      font-size: 13px;
      transition: all 0.2s;
    }
    .export-item:hover { border-color: #e8532a; color: #e8532a; }
    
    table {
      width: 100%;
      border-collapse: collapse;
      font-size: 13px;
    }
    thead { background: #0f1115; }
    th {
      text-align: left;
      padding: 12px;
      font-size: 11px;
      color: #9a9ea6;
      text-transform: uppercase;
      font-weight: 600;
      border-bottom: 1px solid #2a2e37;
    }
    td { padding: 12px; border-bottom: 1px solid #2a2e37; }
    tr:hover { background: #1a1d24; }
    
    .loading { text-align: center; color: #9a9ea6; padding: 40px 20px; }
    .error { color: #ef4444; padding: 12px; background: #2a1515; border-radius: 8px; }
    
    .format-select {
      padding: 10px 14px;
      border: 1px solid #3a3e47;
      background: #0f1115;
      color: #fff;
      border-radius: 8px;
      font-size: 13px;
      font-family: inherit;
      width: 200px;
      margin-bottom: 20px;
    }
  </style>
</head>
<body>

<div class="wrap">
  <h1>📊 Analytics Dashboard</h1>
  <p class="subtitle">View usage statistics, search conversations, and export your data</p>

  <div class="metrics">
    <div class="metric">
      <div class="metric-value" id="metric-chats">-</div>
      <div class="metric-label">Total Chats</div>
    </div>
    <div class="metric">
      <div class="metric-value" id="metric-messages">-</div>
      <div class="metric-label">Total Messages</div>
    </div>
    <div class="metric">
      <div class="metric-value" id="metric-folders">-</div>
      <div class="metric-label">Folders</div>
    </div>
    <div class="metric">
      <div class="metric-value" id="metric-week">-</div>
      <div class="metric-label">This Week</div>
    </div>
  </div>

  <div class="tabs">
    <button class="tab-button active" onclick="showTab('search')">🔎 Search</button>
    <button class="tab-button" onclick="showTab('export')">📥 Export</button>
    <button class="tab-button" onclick="showTab('trends')">📈 Trends</button>
  </div>

  <!-- SEARCH TAB -->
  <div id="search" class="tab-content active">
    <div class="card">
      <h2 style="margin-top: 0;">Search Conversations</h2>
      <div class="search-row">
        <input type="text" id="search-query" placeholder="Search..." />
        <input type="date" id="search-start" />
        <input type="date" id="search-end" />
        <select id="search-folder">
          <option value="">All Folders</option>
        </select>
        <button onclick="performSearch()">Search</button>
      </div>
      <div id="search-results" class="results"></div>
    </div>
  </div>

  <!-- EXPORT TAB -->
  <div id="export" class="tab-content">
    <div class="card">
      <h2 style="margin-top: 0;">Export Conversations</h2>
      <label style="display: block; margin-bottom: 8px; font-weight: 700;">Format:</label>
      <select id="export-format" class="format-select">
        <option value="json">JSON</option>
        <option value="html">HTML</option>
        <option value="csv">CSV</option>
      </select>
      <div id="export-list" class="export-grid"></div>
    </div>
  </div>

  <!-- TRENDS TAB -->
  <div id="trends" class="tab-content">
    <div class="card">
      <h2 style="margin-top: 0;">Usage Trends (Last 30 Days)</h2>
      <table>
        <thead>
          <tr>
            <th>Date</th>
            <th>Requests</th>
            <th>Tokens</th>
            <th>Users</th>
          </tr>
        </thead>
        <tbody id="trends-table"></tbody>
      </table>
    </div>
  </div>

</div>

<script>
  // Show/hide tabs
  function showTab(tabName) {
    // Hide all tabs
    var tabs = document.querySelectorAll('.tab-content');
    tabs.forEach(function(tab) {
      tab.classList.remove('active');
    });
    
    // Deactivate all buttons
    var btns = document.querySelectorAll('.tab-button');
    btns.forEach(function(btn) {
      btn.classList.remove('active');
    });
    
    // Show selected tab
    document.getElementById(tabName).classList.add('active');
    
    // Activate clicked button
    event.target.classList.add('active');
    
    // Load data if needed
    if (tabName === 'trends') {
      loadTrends();
    }
    if (tabName === 'export') {
      loadExportList();
    }
  }

  // Load initial dashboard
  function loadDashboard() {
    fetch('/api/analytics/dashboard')
      .then(r => r.json())
      .then(data => {
        document.getElementById('metric-chats').textContent = data.dashboard.total_conversations;
        document.getElementById('metric-messages').textContent = data.dashboard.total_messages_sent;
        document.getElementById('metric-folders').textContent = Object.keys(data.dashboard.folders_breakdown).length;
        document.getElementById('metric-week').textContent = data.dashboard.usage_this_week.total_requests;
        
        // Fill folder dropdown
        var folderSelect = document.getElementById('search-folder');
        Object.keys(data.dashboard.folders_breakdown).forEach(function(folder) {
          var opt = document.createElement('option');
          opt.value = folder;
          opt.textContent = folder + ' (' + data.dashboard.folders_breakdown[folder] + ')';
          folderSelect.appendChild(opt);
        });
      })
      .catch(function(err) {
        console.error('Failed to load dashboard:', err);
      });
  }

  // Search conversations
  function performSearch() {
    var query = document.getElementById('search-query').value;
    var startDate = document.getElementById('search-start').value;
    var endDate = document.getElementById('search-end').value;
    var folder = document.getElementById('search-folder').value;
    
    var resultsDiv = document.getElementById('search-results');
    resultsDiv.innerHTML = '<div class="loading">Searching...</div>';
    
    var filters = {};
    if (startDate) filters.start_date = startDate;
    if (endDate) filters.end_date = endDate;
    if (folder) filters.folder = folder;
    
    fetch('/api/analytics/search', {
      method: 'POST',
      headers: { 'Content-Type': 'application/json' },
      body: JSON.stringify({ query: query, filters: filters })
    })
      .then(r => r.json())
      .then(data => {
        if (data.found === 0) {
          resultsDiv.innerHTML = '<div class="loading">No results found</div>';
          return;
        }
        
        var html = '';
        data.conversations.forEach(function(conv) {
          html += '<div class="result" onclick="exportConv(\'' + conv.id + '\')">';
          html += '<div class="result-title">' + (conv.title || 'Untitled').substring(0, 50) + '</div>';
          html += '<div class="result-meta">' + (conv.messages ? conv.messages.length : 0) + ' messages</div>';
          html += '</div>';
        });
        resultsDiv.innerHTML = html;
      })
      .catch(function(err) {
        resultsDiv.innerHTML = '<div class="error">Error: ' + err.message + '</div>';
      });
  }

  // Load export list
  function loadExportList() {
    var listDiv = document.getElementById('export-list');
    listDiv.innerHTML = '<div class="loading">Loading conversations...</div>';
    
    fetch('/api/analytics/dashboard')
      .then(r => r.json())
      .then(data => {
        var convs = data.dashboard.recent_conversations || [];
        if (convs.length === 0) {
          listDiv.innerHTML = '<div class="loading">No conversations</div>';
          return;
        }
        
        var html = '';
        convs.forEach(function(conv) {
          html += '<div class="export-item" onclick="exportConv(\'' + conv.id + '\')">';
          html += (conv.title || 'Untitled').substring(0, 20);
          html += '</div>';
        });
        listDiv.innerHTML = html;
      })
      .catch(function(err) {
        listDiv.innerHTML = '<div class="error">Error: ' + err.message + '</div>';
      });
  }

  // Export a conversation
  function exportConv(convId) {
    var format = document.getElementById('export-format').value;
    var url = '/api/conversations/' + encodeURIComponent(convId) + '/export?format=' + format;
    window.location.href = url;
  }

  // Load trends
  function loadTrends() {
    var tbody = document.getElementById('trends-table');
    tbody.innerHTML = '<tr><td colspan="4" class="loading">Loading trends...</td></tr>';
    
    fetch('/api/analytics/trend?days=30')
      .then(r => r.json())
      .then(data => {
        if (!data.trend || data.trend.length === 0) {
          tbody.innerHTML = '<tr><td colspan="4" class="loading">No trend data</td></tr>';
          return;
        }
        
        var html = '';
        data.trend.forEach(function(row) {
          html += '<tr>';
          html += '<td>' + row.date + '</td>';
          html += '<td>' + row.requests + '</td>';
          html += '<td>' + row.tokens + '</td>';
          html += '<td>' + row.unique_users + '</td>';
          html += '</tr>';
        });
        tbody.innerHTML = html;
      })
      .catch(function(err) {
        tbody.innerHTML = '<tr><td colspan="4" class="error">Error: ' + err.message + '</td></tr>';
      });
  }

  // Allow Enter in search
  document.addEventListener('DOMContentLoaded', function() {
    var searchInput = document.getElementById('search-query');
    if (searchInput) {
      searchInput.addEventListener('keypress', function(e) {
        if (e.key === 'Enter') {
          performSearch();
        }
      });
    }
  });

  // Load on startup
  loadDashboard();
</script>

</body>
</html>"""
    return Response(html, mimetype="text/html")






@app.route("/api/keys", methods=["POST"])
@login_required
def api_keys_create():
    data = request.get_json(silent=True) or {}
    label = (data.get("label") or "").strip()
    raw_key, record = create_api_key(label, current_username())
    return jsonify({
        "api_key": raw_key,  # shown ONCE — the frontend must display and never re-fetch this
        "id": record["id"],
        "key_prefix": record["key_prefix"],
        "label": record["label"],
        "created_at": record["created_at"],
    })

@app.route("/api/keys/<key_id>", methods=["DELETE"])
@login_required
def api_keys_revoke(key_id):
    ok = revoke_api_key(key_id, current_username())
    return jsonify({"revoked": ok})

@app.route("/api/keys/<key_id>", methods=["PATCH"])
@login_required
def api_keys_rename(key_id):
    data = request.get_json(silent=True) or {}
    new_label = (data.get("label") or "").strip()
    if not new_label:
        return jsonify({"error": "label is required"}), 400
    ok = rename_api_key(key_id, new_label, current_username())
    if not ok:
        return jsonify({"error": "key not found"}), 404
    return jsonify({"renamed": ok, "label": new_label})


# ══════════════════════════════════════════════════════════════════════════════
# ─── ANALYTICS & USAGE TRACKING API ENDPOINTS ────────────────────────────────
# ══════════════════════════════════════════════════════════════════════════════

@app.route("/api/analytics/usage", methods=["GET"])
@login_required
def api_get_usage_report():
    """Get detailed usage analytics for authenticated user's API keys and activity.
    Query params: days_back (default 30), start_date, end_date, key_id (filter by specific key)"""
    username = current_username()
    days_back = request.args.get("days_back", 30, type=int)
    start_date = request.args.get("start_date")
    end_date = request.args.get("end_date")
    key_id = request.args.get("key_id")
    
    report = get_usage_report(username=username, key_id=key_id, start_date=start_date, 
                              end_date=end_date, days_back=days_back)
    return jsonify(report)


@app.route("/api/analytics/search", methods=["POST"])
@login_required
def api_search_conversations():
    """Full-text search conversations with advanced filtering.
    POST body: { 'query': 'string', 'filters': { 'start_date': 'YYYY-MM-DD', 
    'end_date': 'YYYY-MM-DD', 'folder': 'name', 'archived': bool, 
    'pinned': bool, 'min_messages': int } }"""
    username = current_username()
    data = request.get_json(silent=True) or {}
    query = data.get("query", "").strip()
    filters = data.get("filters", {})
    
    results = search_conversations(username, query=query, filters=filters)
    return jsonify({"found": len(results), "conversations": results})


@app.route("/api/conversations/<conv_id>/export", methods=["GET"])
@login_required
def api_export_conversation(conv_id):
    """Export a conversation in requested format (json, csv, html).
    Query param: format (default 'json')"""
    username = current_username()
    format_type = request.args.get("format", "json").lower()
    
    if format_type not in ("json", "csv", "html"):
        return jsonify({"error": "format must be json, csv, or html"}), 400
    
    content, mimetype, filename = export_conversation(conv_id, username, format_type)
    if not content:
        return jsonify({"error": "conversation not found"}), 404
    
    return Response(content, mimetype=mimetype, 
                    headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@app.route("/api/analytics/dashboard", methods=["GET"])
@login_required
def api_analytics_dashboard():
    """Get a quick dashboard summary of user's activity.
    Includes: recent activity, top models, conversation trends."""
    username = current_username()
    report = get_usage_report(username=username, days_back=7)
    
    # Get recent conversations
    convs = list_conversations(username)
    recent = sorted(convs, key=lambda c: c.get("updated_at", ""), reverse=True)[:10]
    
    # Get folder breakdown
    folders = {}
    for conv in convs:
        folder = conv.get("folder") or "Uncategorized"
        if folder not in folders:
            folders[folder] = 0
        folders[folder] += 1
    
    return jsonify({
        "username": username,
        "dashboard": {
            "total_conversations": len(convs),
            "total_messages_sent": sum(len(c.get("messages", [])) for c in convs),
            "recent_conversations": recent,
            "folders_breakdown": folders,
            "usage_this_week": report["summary"],
            "generated_at": _dt.utcnow().isoformat()
        }
    })


@app.route("/api/admin/stats", methods=["GET"])
@login_required
def api_admin_stats():
    """Get system-wide admin statistics (requires admin role).
    Includes: total users, total requests, top models, daily trends."""
    username = current_username()
    
    # Check if user is admin (you can add a proper admin role check here)
    # For now, allow super-admin access
    is_admin = username in getattr(_app, 'SUPER_ADMINS', [])
    if not is_admin and os.environ.get("SUPER_ADMIN_USERNAME") != username:
        return jsonify({"error": "admin access required"}), 403
    
    stats = get_admin_stats(include_daily=True)
    return jsonify(stats)


@app.route("/api/analytics/export-bulk", methods=["POST"])
@login_required
def api_export_bulk():
    """Export multiple conversations at once in ZIP format.
    POST body: { 'conversation_ids': ['id1', 'id2', ...], 'format': 'json|csv|html' }"""
    username = current_username()
    data = request.get_json(silent=True) or {}
    conv_ids = data.get("conversation_ids", [])
    format_type = data.get("format", "json").lower()
    
    if not conv_ids:
        return jsonify({"error": "conversation_ids required"}), 400
    if format_type not in ("json", "csv", "html"):
        return jsonify({"error": "format must be json, csv, or html"}), 400
    
    # Create ZIP in memory
    import zipfile as _zipfile
    from io import BytesIO as _BytesIO
    
    zip_buffer = _BytesIO()
    with _zipfile.ZipFile(zip_buffer, "w", _zipfile.ZIP_DEFLATED) as zip_file:
        for conv_id in conv_ids:
            content, mimetype, filename = export_conversation(conv_id, username, format_type)
            if content:
                zip_file.writestr(filename, content)
    
    zip_buffer.seek(0)
    return Response(zip_buffer.getvalue(), mimetype="application/zip",
                    headers={"Content-Disposition": 'attachment; filename="conversations.zip"'})


@app.route("/api/analytics/filters-options", methods=["GET"])
@login_required
def api_search_filter_options():
    """Get available filter options for conversation search (folders, date range, etc)."""
    username = current_username()
    convs = list_conversations(username)
    
    folders = set(c.get("folder") for c in convs if c.get("folder"))
    folders.add("Uncategorized")
    
    dates = [c.get("updated_at", "")[:10] for c in convs if c.get("updated_at")]
    min_date = min(dates) if dates else None
    max_date = max(dates) if dates else None
    
    return jsonify({
        "available_folders": list(folders),
        "date_range": {"min": min_date, "max": max_date},
        "total_conversations": len(convs),
        "filters_available": ["start_date", "end_date", "folder", "archived", "pinned", "min_messages"]
    })


@app.route("/api/analytics/trend", methods=["GET"])
@login_required
def api_usage_trend():
    """Get usage trends over time (daily breakdown).
    Query params: days (default 30)"""
    username = current_username()
    days = request.args.get("days", 30, type=int)
    
    report = get_usage_report(username=username, days_back=days)
    trend_data = []
    
    for day_entry in report["summary"]["daily_breakdown"]:
        date = day_entry["date"]
        stats = day_entry["stats"]
        trend_data.append({
            "date": date,
            "requests": stats.get("total_requests", 0),
            "tokens": stats.get("total_tokens", 0),
            "unique_users": len(stats.get("unique_users", []))
        })
    
    return jsonify({
        "period_days": days,
        "trend": trend_data,
        "generated_at": _dt.utcnow().isoformat()
    })


@app.route("/v1/chat/completions", methods=["POST"])
def v1_chat_completions():
    auth = request.headers.get("Authorization", "")
    raw_key = auth[7:].strip() if auth.lower().startswith("bearer ") else ""
    if not verify_api_key(raw_key):
        return jsonify({"error": {"message": "Invalid or missing API key. Include "
                                              "'Authorization: Bearer aarav-...'.",
                                   "type": "invalid_request_error"}}), 401

    data = request.get_json(silent=True) or {}
    incoming = data.get("messages") or []
    if not incoming:
        return jsonify({"error": {"message": "'messages' is required.",
                                   "type": "invalid_request_error"}}), 400

    system_prompt = SYSTEM_PROMPT
    gemini_messages = []
    for m in incoming:
        role = m.get("role", "user")
        content = m.get("content", "")
        if role == "system":
            system_prompt = content or system_prompt
            continue
        gemini_messages.append({
            "role": "user" if role == "user" else "model",
            "parts": [{"text": content}],
        })

    stream_requested = bool(data.get("stream"))
    model_name = data.get("model") or "mythic-2"

    if stream_requested:
        def sse():
            for chunk in auto_stream_chunks(None, gemini_messages, system_prompt):
                payload = {
                    "id": "chatcmpl-" + uuid.uuid4().hex[:24],
                    "object": "chat.completion.chunk",
                    "model": model_name,
                    "choices": [{"index": 0, "delta": {"content": chunk}, "finish_reason": None}],
                }
                yield f"data: {json.dumps(payload)}\n\n"
            yield "data: [DONE]\n\n"
        return Response(stream_with_context(sse()), mimetype="text/event-stream")

    full_reply = "".join(auto_stream_chunks(None, gemini_messages, system_prompt))
    return jsonify({
        "id": "chatcmpl-" + uuid.uuid4().hex[:24],
        "object": "chat.completion",
        "created": int(time.time()),
        "model": model_name,
        "choices": [{
            "index": 0,
            "message": {"role": "assistant", "content": full_reply},
            "finish_reason": "stop",
        }],
        "usage": {  # not tracked precisely — placeholder for client compatibility
            "prompt_tokens": None, "completion_tokens": None, "total_tokens": None,
        },
    })


# --- Public API: /v1/images/generations (OpenAI-compatible, requires API key) -----
@app.route("/v1/images/generations", methods=["POST"])
def v1_images_generations():
    """OpenAI-compatible endpoint: POST with prompt + style, get back image URLs."""
    auth = request.headers.get("Authorization", "")
    raw_key = auth[7:].strip() if auth.lower().startswith("bearer ") else ""
    if not verify_api_key(raw_key):
        return jsonify({"error": {"message": "Invalid or missing API key.",
                                   "type": "invalid_request_error"}}), 401
    
    # Reuse the existing /api/generate-image logic by calling it internally
    # We'll just set the session and call the endpoint function
    session["user_id"] = get_or_create_owner_id(preferred_id=raw_key)
    return generate_image()


# --- Public API: /v1/code/execute (OpenAI-compatible, requires API key) ----------
@app.route("/v1/code/execute", methods=["POST"])
def v1_code_execute():
    """OpenAI-compatible endpoint: POST with code + language, get back output/errors."""
    auth = request.headers.get("Authorization", "")
    raw_key = auth[7:].strip() if auth.lower().startswith("bearer ") else ""
    if not verify_api_key(raw_key):
        return jsonify({"error": {"message": "Invalid or missing API key.",
                                   "type": "invalid_request_error"}}), 401
    
    # Reuse the existing /api/execute-code logic
    session["user_id"] = get_or_create_owner_id(preferred_id=raw_key)
    return api_execute_code()


@app.route("/api/conversations", methods=["GET"])
@login_required
def api_list_conversations():
    show_archived = request.args.get("archived", "0") == "1"
    convs = list_conversations(current_username())
    if show_archived:
        convs = [c for c in convs if c.get("archived")]
    else:
        convs = [c for c in convs if not c.get("archived")]
    return jsonify({"conversations": convs})


@app.route("/api/conversations", methods=["POST"])
@login_required
def api_create_conversation():
    """Creates an empty 'New chat' conversation immediately, so it shows up
    in the sidebar right away instead of only appearing after the first
    message is sent. The conversation carries no messages yet — sending the
    first message into it works exactly the same as any other conversation."""
    username = current_username()
    conv_id = str(uuid.uuid4())
    save_conversation(username, conv_id, {"title": "New chat", "messages": []})
    return jsonify({"id": conv_id, "title": "New chat"})


# Patterns that identify a conversation as internal-tooling leakage rather
# than a real chat — e.g. old follow-up-suggestion or tone/length-prefix
# requests saved before the ephemeral-request fix existed. Matched against
# either the saved title or the first user message.
_JUNK_CONV_PATTERNS = (
    "based on this ai reply, suggest",
    "[instructions:",
)


def _conv_is_junk(conv_summary, username):
    title = (conv_summary.get("title") or "").strip().lower()
    if any(title.startswith(p) for p in _JUNK_CONV_PATTERNS):
        return True
    # Title alone might be a truncated/renamed version — check the first
    # real user message too, for conversations saved before titles were
    # cleaned up server-side.
    full = load_conversation(username, conv_summary["id"])
    if not full:
        return False
    for m in full.get("messages", []):
        if m.get("role") != "user":
            continue
        text = "".join(p.get("text", "") for p in m.get("parts", []) if "text" in p).strip().lower()
        if text:
            return any(text.startswith(p) for p in _JUNK_CONV_PATTERNS)
    return False


@app.route("/api/conversations/cleanup-junk", methods=["POST"])
@login_required
def api_cleanup_junk_conversations():
    """Finds and deletes stray conversations created by internal tooling
    (old follow-up-suggestion / tone-prefix requests) rather than real user
    chats — see _JUNK_CONV_PATTERNS. Safe to call any time; only ever
    deletes conversations matching those specific patterns."""
    username = current_username()
    convs = list_conversations(username)
    removed = []
    for c in convs:
        if _conv_is_junk(c, username):
            delete_conversation(username, c["id"])
            removed.append(c.get("title", c["id"]))
    return jsonify({"status": "done", "removed_count": len(removed), "removed_titles": removed})


@app.route("/api/folders", methods=["GET"])
@login_required
def api_list_folders():
    return jsonify({"folders": list_folders(current_username())})


@app.route("/api/conversations/<conv_id>", methods=["GET"])
@login_required
def api_get_conversation(conv_id):
    data = load_conversation(current_username(), conv_id)
    if data is None:
        return jsonify({"error": "not found"}), 404
    simplified = []
    for m in data.get("messages", []):
        role = "user" if m["role"] == "user" else "ai"
        text_parts = [p.get("text", "") for p in m["parts"] if "text" in p]
        entry = {"role": role, "text": "".join(text_parts)}
        if m.get("attachment_meta"):
            entry["attachment"] = m["attachment_meta"]
        simplified.append(entry)
    return jsonify({"messages": simplified, "title": data.get("title", "New chat")})


@app.route("/api/conversations/<conv_id>", methods=["DELETE"])
@login_required
def api_delete_conversation(conv_id):
    delete_conversation(current_username(), conv_id)
    return jsonify({"status": "deleted"})


@app.route("/api/conversations/<conv_id>", methods=["PATCH"])
@login_required
def api_rename_conversation(conv_id):
    """Updates one or more of: title, folder, pinned, archived, persona_id,
    custom_instructions. At least one field must be present; unspecified
    fields are left unchanged."""
    data = request.get_json(force=True) or {}
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404

    changed = {}
    if "title" in data:
        new_title = (data.get("title") or "").strip()[:120]
        if not new_title:
            return jsonify({"error": "title cannot be empty"}), 400
        conv["title"] = new_title
        conv["title_is_custom"] = True
        changed["title"] = new_title
    if "folder" in data:
        folder = (data.get("folder") or "").strip()[:60] or None
        conv["folder"] = folder
        changed["folder"] = folder
    if "pinned" in data:
        conv["pinned"] = bool(data.get("pinned"))
        changed["pinned"] = conv["pinned"]
    if "archived" in data:
        conv["archived"] = bool(data.get("archived"))
        changed["archived"] = conv["archived"]
    if "persona_id" in data:
        # null/empty clears the persona back to the default assistant.
        pid = data.get("persona_id") or None
        if pid and not get_persona_for(username, pid):
            return jsonify({"error": "persona not found"}), 404
        conv["persona_id"] = pid
        changed["persona_id"] = pid
    if "custom_instructions" in data:
        instructions = (data.get("custom_instructions") or "").strip()[:2000]
        conv["custom_instructions"] = instructions
        changed["custom_instructions"] = instructions

    if not changed:
        return jsonify({"error": "no recognized fields to update "
                                  "(expected title/folder/pinned/archived/persona_id/custom_instructions)"}), 400

    save_conversation(username, conv_id, conv)
    return jsonify({"status": "updated", **changed})


@app.route("/api/conversations/<conv_id>/edit-message", methods=["POST"])
@login_required
def api_edit_message(conv_id):
    """Edits a previously-sent USER message and branches the conversation
    from that point: every message at or after message_index (the edited
    message itself, its old AI reply, and anything after) is dropped, then
    a fresh user entry with the new text is appended in its place.

    Deliberately does NOT touch the streaming/model logic at all — the
    frontend follows this call with the existing regenerate=True request to
    /api/chat, which already knows how to stream a reply for "last message
    is a user message with no reply yet". This endpoint only rewrites
    history; /api/chat (unchanged) does the actual generation."""
    data = request.get_json(force=True) or {}
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404

    try:
        message_index = int(data.get("message_index"))
    except (TypeError, ValueError):
        return jsonify({"error": "message_index must be an integer"}), 400
    new_text = (data.get("new_text") or "").strip()
    if not new_text:
        return jsonify({"error": "new_text cannot be empty"}), 400

    messages = conv.get("messages", [])
    if not (0 <= message_index < len(messages)):
        return jsonify({"error": "message_index out of range"}), 400
    original = messages[message_index]
    if original.get("role") != "user":
        return jsonify({"error": "message_index does not point to a user message"}), 400

    # Preserve any non-text parts (e.g. an attached image's inline_data) from
    # the original message so editing the caption doesn't silently drop the
    # attachment — only the text itself changes.
    new_parts = [{"text": new_text}]
    for part in original.get("parts", []):
        if "inline_data" in part:
            new_parts.append(part)
    new_entry = {"role": "user", "parts": new_parts}
    if "attachment_meta" in original:
        new_entry["attachment_meta"] = original["attachment_meta"]

    conv["messages"] = messages[:message_index] + [new_entry]
    save_conversation(username, conv_id, conv)
    return jsonify({"status": "ok", "message_index": message_index})


@app.route("/api/conversations/<conv_id>/duplicate", methods=["POST"])
@login_required
def api_duplicate_conversation(conv_id):
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    new_id = str(uuid.uuid4())
    new_conv = {
        "title": (conv.get("title") or "New chat") + " (copy)",
        "messages": json.loads(json.dumps(conv.get("messages", []))),  # deep copy
        "folder": conv.get("folder"),
        "pinned": False,
        "archived": False,
    }
    save_conversation(username, new_id, new_conv)
    return jsonify({"status": "duplicated", "id": new_id, "title": new_conv["title"]})


# ---------------------------------------------------------------------------
# Phase 1 additions: reactions, message pins, conversation starring, drafts,
# split/merge. All additive to the existing conversation JSON shape — no
# existing field is renamed or removed, so old saved conversations keep
# working unchanged (missing keys just default to empty/False on read).
# ---------------------------------------------------------------------------

_VALID_REACTIONS = {"like", "dislike", "heart"}
_REACTION_EMOJI = {"like": "👍", "dislike": "👎", "heart": "❤️"}


@app.route("/api/conversations/<conv_id>/messages/<int:message_index>/reaction", methods=["POST"])
@login_required
def api_set_message_reaction(conv_id, message_index):
    """Sets or clears the current user's reaction on a message.
    Body: {"reaction": "like"|"dislike"|"heart"|null}. One reaction per
    message per user; sending the same reaction again clears it (toggle)."""
    data = request.get_json(force=True) or {}
    reaction = data.get("reaction")
    if reaction is not None and reaction not in _VALID_REACTIONS:
        return jsonify({"error": f"reaction must be one of {sorted(_VALID_REACTIONS)} or null"}), 400

    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    messages = conv.get("messages", [])
    if not (0 <= message_index < len(messages)):
        return jsonify({"error": "message_index out of range"}), 400

    msg = messages[message_index]
    if reaction is None or msg.get("reaction") == reaction:
        msg.pop("reaction", None)
        new_value = None
    else:
        msg["reaction"] = reaction
        new_value = reaction

    save_conversation(username, conv_id, conv)
    return jsonify({"status": "ok", "message_index": message_index, "reaction": new_value,
                     "emoji": _REACTION_EMOJI.get(new_value)})


@app.route("/api/conversations/<conv_id>/messages/<int:message_index>/pin", methods=["POST"])
@login_required
def api_toggle_message_pin(conv_id, message_index):
    """Pins/unpins a single message within a conversation (distinct from
    conversation-level 'pinned', which is about sidebar ordering).
    Body: {"pinned": true|false}."""
    data = request.get_json(force=True) or {}
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    messages = conv.get("messages", [])
    if not (0 <= message_index < len(messages)):
        return jsonify({"error": "message_index out of range"}), 400

    pinned = bool(data.get("pinned", True))
    if pinned:
        messages[message_index]["pinned"] = True
    else:
        messages[message_index].pop("pinned", None)

    save_conversation(username, conv_id, conv)
    return jsonify({"status": "ok", "message_index": message_index, "pinned": pinned})


@app.route("/api/conversations/<conv_id>/pinned-messages", methods=["GET"])
@login_required
def api_list_pinned_messages(conv_id):
    """Returns every pinned message in the conversation, in order, for the
    'Pinned Messages' panel. Each entry includes its index so the frontend
    can jump back to it in the full transcript."""
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    pinned = []
    for i, m in enumerate(conv.get("messages", [])):
        if m.get("pinned"):
            text = "".join(p.get("text", "") for p in m.get("parts", []) if "text" in p)
            pinned.append({"message_index": i, "role": m.get("role"), "text": text})
    return jsonify({"pinned": pinned})


@app.route("/api/conversations/<conv_id>/star", methods=["POST"])
@login_required
def api_toggle_star_conversation(conv_id):
    """Stars/unstars a conversation as a favorite. Kept separate from the
    existing 'pinned' field (which controls sidebar top-of-list ordering) so
    the two concepts — 'keep at top' vs 'favorite' — don't collide.
    Body: {"starred": true|false}."""
    data = request.get_json(force=True) or {}
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    conv["starred"] = bool(data.get("starred", True))
    save_conversation(username, conv_id, conv)
    return jsonify({"status": "ok", "starred": conv["starred"]})


@app.route("/api/conversations/favorites", methods=["GET"])
@login_required
def api_list_favorite_conversations():
    username = current_username()
    convs = [c for c in list_conversations(username) if c.get("starred")]
    return jsonify({"conversations": convs})


@app.route("/api/conversations/<conv_id>/draft", methods=["GET"])
@login_required
def api_get_draft(conv_id):
    """Returns the auto-saved unsent composer text for this conversation,
    if any. Drafts are stored on the conversation record itself, so they
    survive navigation, refresh, and reconnects for free."""
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    return jsonify({"draft": conv.get("draft", "")})


@app.route("/api/conversations/<conv_id>/draft", methods=["PUT"])
@login_required
def api_save_draft(conv_id):
    """Upserts the draft text. The frontend should debounce calls to this
    (e.g. every 1-2s while typing, and on blur/navigation) rather than
    calling on every keystroke. Body: {"text": "..."}."""
    data = request.get_json(force=True) or {}
    text = (data.get("text") or "")[:20000]
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    if text:
        conv["draft"] = text
    else:
        conv.pop("draft", None)
    save_conversation(username, conv_id, conv)
    return jsonify({"status": "ok"})


@app.route("/api/conversations/<conv_id>/split", methods=["POST"])
@login_required
def api_split_conversation(conv_id):
    """Creates a new conversation containing messages from message_index
    onward, leaving the original conversation completely untouched.
    Body: {"message_index": int}."""
    data = request.get_json(force=True) or {}
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    try:
        message_index = int(data.get("message_index"))
    except (TypeError, ValueError):
        return jsonify({"error": "message_index must be an integer"}), 400
    messages = conv.get("messages", [])
    if not (0 <= message_index < len(messages)):
        return jsonify({"error": "message_index out of range"}), 400

    new_id = str(uuid.uuid4())
    new_conv = {
        "title": (conv.get("title") or "New chat") + " (split)",
        "messages": json.loads(json.dumps(messages[message_index:])),
        "folder": conv.get("folder"),
        "split_from": conv_id,
    }
    save_conversation(username, new_id, new_conv)
    return jsonify({"status": "ok", "id": new_id, "title": new_conv["title"]})


@app.route("/api/conversations/merge", methods=["POST"])
@login_required
def api_merge_conversations():
    """Merges two conversations into a NEW conversation, ordered by each
    message's original position (both source conversations are left
    intact). Exact duplicate consecutive user+reply pairs from the second
    conversation that also appear in the first are skipped on a best-effort
    text-match basis. Body: {"conv_id_a": "...", "conv_id_b": "..."}."""
    data = request.get_json(force=True) or {}
    username = current_username()
    id_a, id_b = data.get("conv_id_a"), data.get("conv_id_b")
    conv_a = load_conversation(username, id_a) if id_a else None
    conv_b = load_conversation(username, id_b) if id_b else None
    if conv_a is None or conv_b is None:
        return jsonify({"error": "one or both conversations not found"}), 404

    def _msg_text(m):
        return "".join(p.get("text", "") for p in m.get("parts", []) if "text" in p).strip()

    msgs_a = conv_a.get("messages", [])
    seen_texts = {_msg_text(m) for m in msgs_a if _msg_text(m)}
    merged = list(msgs_a)
    skipped = 0
    for m in conv_b.get("messages", []):
        t = _msg_text(m)
        if t and t in seen_texts:
            skipped += 1
            continue
        merged.append(m)

    new_id = str(uuid.uuid4())
    new_title = f'{conv_a.get("title", "Chat")} + {conv_b.get("title", "Chat")}'[:120]
    new_conv = {"title": new_title, "messages": merged, "merged_from": [id_a, id_b]}
    save_conversation(username, new_id, new_conv)
    return jsonify({"status": "ok", "id": new_id, "title": new_title,
                     "message_count": len(merged), "duplicates_skipped": skipped})


@app.route("/api/conversations/<conv_id>/messages/<int:message_index>/undo-send", methods=["POST"])
@login_required
def api_undo_send(conv_id, message_index):
    """Deletes a just-sent message (and any reply after it) within the
    undo grace period. The frontend enforces the grace-period countdown
    (e.g. 5s) client-side and only calls this if the user taps Undo in
    time; server just does the deletion of everything from that index on,
    same mechanics as edit-message's history trim."""
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    messages = conv.get("messages", [])
    if not (0 <= message_index < len(messages)):
        return jsonify({"error": "message_index out of range"}), 400
    conv["messages"] = messages[:message_index]
    save_conversation(username, conv_id, conv)
    return jsonify({"status": "ok", "removed_from": message_index})


def _share_url_for(share_id):
    return get_public_origin() + "/share/" + share_id


@app.route("/api/conversations/<conv_id>/share", methods=["GET"])
@login_required
def api_get_share_status(conv_id):
    """Returns whether this conversation currently has an active public
    share link, and its URL if so."""
    username = current_username()
    if load_conversation(username, conv_id) is None:
        return jsonify({"error": "not found"}), 404
    share_id = get_active_share_id(username, conv_id)
    if not share_id:
        return jsonify({"shared": False})
    return jsonify({"shared": True, "share_id": share_id, "share_url": _share_url_for(share_id)})


@app.route("/api/conversations/<conv_id>/share", methods=["POST"])
@login_required
def api_create_share(conv_id):
    """Creates (or reuses an existing) public share link for this
    conversation — anyone with the link can view a read-only copy, no
    login required."""
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    if not conv.get("messages"):
        return jsonify({"error": "This chat has no messages yet — nothing to share."}), 400
    share_id = create_share_link(username, conv_id)
    return jsonify({"status": "shared", "share_id": share_id, "share_url": _share_url_for(share_id)})


@app.route("/api/conversations/<conv_id>/share", methods=["DELETE"])
@login_required
def api_revoke_share(conv_id):
    """Revokes any active public share link for this conversation — the old
    link stops working immediately."""
    username = current_username()
    revoke_share_link(username, conv_id)
    return jsonify({"status": "revoked"})


@app.route("/api/share/<share_id>", methods=["GET"])
def api_public_share(share_id):
    """Public, unauthenticated endpoint — returns a read-only view of a
    shared conversation. No user identity is exposed, only the messages."""
    ref = resolve_share_link(share_id)
    if not ref:
        return jsonify({"error": "This share link is invalid or has been revoked."}), 404
    conv = load_conversation(ref["username"], ref["conv_id"])
    if conv is None:
        return jsonify({"error": "This shared chat is no longer available."}), 404
    simplified = []
    for m in conv.get("messages", []):
        role = "user" if m["role"] == "user" else "ai"
        text_parts = [p.get("text", "") for p in m.get("parts", []) if "text" in p]
        entry = {"role": role, "text": "".join(text_parts)}
        if m.get("attachment_meta"):
            entry["attachment"] = m["attachment_meta"]
        simplified.append(entry)
    return jsonify({"title": conv.get("title", "Shared chat"), "messages": simplified})


@app.route("/api/share/<share_id>/continue", methods=["POST"])
@login_required
def api_continue_shared_chat(share_id):
    """Lets a visitor fork their own private, editable copy of a shared
    conversation so they can keep chatting. This NEVER touches the original
    owner's conversation — it copies the messages into a brand-new
    conversation under the visitor's own (anonymous, cookie-based) session,
    same as the existing "Duplicate" feature."""
    ref = resolve_share_link(share_id)
    if not ref:
        return jsonify({"error": "This share link is invalid or has been revoked."}), 404
    source_conv = load_conversation(ref["username"], ref["conv_id"])
    if source_conv is None:
        return jsonify({"error": "This shared chat is no longer available."}), 404

    visitor_username = current_username()
    new_id = str(uuid.uuid4())
    new_conv = {
        "title": (source_conv.get("title") or "Shared chat").strip(),
        "title_is_custom": True,
        "messages": json.loads(json.dumps(source_conv.get("messages", []))),  # deep copy
        "folder": None,
        "pinned": False,
        "archived": False,
    }
    save_conversation(visitor_username, new_id, new_conv)
    return jsonify({"status": "forked", "conversation_id": new_id})


@app.route("/share/<share_id>")
def public_share_page(share_id):
    """Public, unauthenticated read-only page rendering a shared chat —
    intentionally a separate, minimal template with no sidebar and no
    access to the viewer's own conversations. It does offer a "Continue
    this conversation" action, which forks a private copy for the visitor
    (see api_continue_shared_chat) rather than editing the original."""
    return Response(SHARE_PAGE, mimetype="text/html")


def _collect_full_reply(chunks):
    return "".join(chunks)


@app.route("/api/conversations/<conv_id>/generate-title", methods=["POST"])
@login_required
def api_generate_title(conv_id):
    """Asks the AI to write a short, punchy title from the first exchange in
    the conversation, replacing the naive first-40-characters title. Safe to
    call any time; falls back to leaving the title unchanged if the AI can't
    be reached."""
    username = current_username()
    conv = load_conversation(username, conv_id)
    if conv is None:
        return jsonify({"error": "not found"}), 404
    messages = conv.get("messages", [])
    if not messages:
        return jsonify({"error": "conversation has no messages yet"}), 400

    convo_excerpt = []
    for m in messages[:4]:
        text = "".join(p.get("text", "") for p in m.get("parts", []) if "text" in p)
        if text:
            speaker = "User" if m["role"] == "user" else "Assistant"
            convo_excerpt.append(f"{speaker}: {text[:300]}")
    excerpt = "\n".join(convo_excerpt)
    if not excerpt.strip():
        return jsonify({"error": "no text content to summarize"}), 400

    title_prompt = [{"role": "user", "parts": [{"text":
        "Below is a snippet of a chat conversation. Write ONE short, specific "
        "title (1-3 words) describing what the conversation is actually "
        "ABOUT — not a description of your task, not the words 'chat', "
        "'title', or 'conversation' themselves. No quotes, no punctuation "
        "at the end, plain text only. Prefer a simple category-style label "
        "when that fits.\n\n"
        "Examples:\n"
        "- If the user says 'hi' and gets a greeting back → 'Greeting'\n"
        "- If they ask about Python loops → 'Python Loops'\n"
        "- If they discuss weekend plans → 'Weekend Plans'\n\n"
        f"Conversation:\n{excerpt}\n\nTitle:"
    }]}]
    data = request.get_json(silent=True) or {}
    user_groq_key = (data.get("groq_api_key") or "").strip()
    user_cerebras_key = (data.get("cerebras_api_key") or "").strip()

    try:
        raw_title = _collect_full_reply(
            auto_stream_chunks(None, title_prompt, SYSTEM_PROMPT, user_groq_key, user_cerebras_key)
        ).strip()
    except Exception:
        raw_title = ""

    raw_title = raw_title.strip().strip('"').strip("'")
    raw_title = re.sub(r'^\[Instructions:.*?\]\s*', '', raw_title, flags=re.DOTALL)
    # Take just the first line/sentence in case the model added extra
    # commentary despite instructions — truncate rather than reject outright,
    # so a slightly-verbose reply still produces a usable title.
    raw_title = raw_title.split("\n")[0].strip()
    raw_title = re.sub(r'^Title:\s*', '', raw_title, flags=re.IGNORECASE).strip()

    # Safety filter: reject titles that describe the TASK instead of the
    # conversation's actual content (the model sometimes echoes back
    # phrases like "Generating Chat Titles" or "Chat Title" for very short/
    # generic exchanges like a bare "hi"). Fall back to a sensible default
    # derived from the first user message instead of saving junk.
    _meta_phrases = ("generating chat title", "chat title", "conversation title",
                      "title generation", "generate title", "untitled")
    if not raw_title or any(p in raw_title.lower() for p in _meta_phrases):
        first_user_text = ""
        for m in messages:
            if m.get("role") == "user":
                first_user_text = "".join(p.get("text", "") for p in m.get("parts", []) if "text" in p)
                break
        raw_title = (first_user_text[:40].strip() or "New chat")

    if not raw_title:
        return jsonify({"status": "unchanged", "title": conv.get("title", "New chat")})

    new_title = raw_title[:60]
    conv["title"] = new_title
    save_conversation(username, conv_id, conv)
    return jsonify({"status": "generated", "title": new_title})


# ── Lightweight in-memory response cache ─────────────────────────────────────
# Used for deterministic, repeat-prone calls (follow-up suggestions, chat-title
# generation) where an identical prompt is likely to recur and re-hitting the
# API adds latency/quota use for no benefit. NOT used for normal chat turns,
# since conversation context differs on every message. TTL + size-capped so it
# never grows unbounded; resets on process restart (no external cache needed).
_response_cache = {}
_RESPONSE_CACHE_TTL_SECONDS = 600
_RESPONSE_CACHE_MAX_ENTRIES = 300


def _cache_key(label, model, messages):
    raw = json.dumps(messages, sort_keys=True) + "|" + label + "|" + (model or "")
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def _cache_get(key):
    entry = _response_cache.get(key)
    if not entry:
        return None
    text, ts = entry
    if time.time() - ts > _RESPONSE_CACHE_TTL_SECONDS:
        _response_cache.pop(key, None)
        return None
    return text


def _cache_set(key, text):
    if len(_response_cache) > _RESPONSE_CACHE_MAX_ENTRIES:
        _response_cache.pop(next(iter(_response_cache)), None)
    _response_cache[key] = (text, time.time())


def to_openai_messages(gemini_messages, system_prompt):
    """Convert stored Gemini-format messages to OpenAI-compatible chat format.
    Used by both Groq and Cerebras (both use the same OpenAI-style chat API)."""
    msgs = [{"role": "system", "content": system_prompt}]
    for m in gemini_messages:
        role = "user" if m["role"] == "user" else "assistant"
        text = "".join(p.get("text", "") for p in m["parts"] if "text" in p)
        msgs.append({"role": role, "content": text})
    return msgs


def to_openai_vision_messages(gemini_messages, system_prompt):
    """Like to_openai_messages, but if the LAST user turn carries an image
    (inline_data with an image/* mime type — a Video Call frame, Screen
    Share frame, or a regular image attachment), that turn is sent as a
    real multimodal content block so a vision-capable model can actually
    see it. Earlier turns stay text-only — we don't re-send old frames,
    which keeps requests small and avoids re-uploading stale video/screen
    snapshots from earlier in the conversation."""
    msgs = [{"role": "system", "content": system_prompt}]
    last_idx = len(gemini_messages) - 1
    for i, m in enumerate(gemini_messages):
        role = "user" if m["role"] == "user" else "assistant"
        text = "".join(p.get("text", "") for p in m["parts"] if "text" in p)
        image_part = None
        if i == last_idx and role == "user":
            for p in m["parts"]:
                inline = p.get("inline_data")
                if inline and str(inline.get("mime_type", "")).startswith("image/"):
                    image_part = inline
                    break
        if image_part:
            content = []
            if text:
                content.append({"type": "text", "text": text})
            content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{image_part['mime_type']};base64,{image_part['data']}"},
            })
            msgs.append({"role": role, "content": content})
        else:
            msgs.append({"role": role, "content": text})
    return msgs


def groq_vision_stream_chunks(messages, api_key=None, model=None):
    """Vision-capable completion via Groq's multimodal model. Used for Video
    Call, Screen Share, and regular image attachments. No Cerebras fallback
    here — Cerebras isn't used as a vision provider in this app — if this
    yields nothing, the caller falls back to the normal text-only path."""
    yield from _openai_style_stream(
        "https://api.groq.com/openai/v1/chat/completions",
        api_key or GROQ_API_KEY, model or GROQ_VISION_MODEL, messages, "Groq-Vision",
    )


# TEMP DEBUG: collects a human-readable reason for each provider failure
# within a single request, so the actual cause can be shown to the user
# instead of only being printed to server logs. Reset at the start of every
# auto_stream_chunks() call. Remove this once providers are confirmed working.
_LAST_PROVIDER_ERRORS = []


def _openai_style_stream(url, api_key, model, messages, provider_label):
    """Shared streaming logic for Groq/Cerebras (both are OpenAI-compatible).
    Yields nothing at all on ANY failure (auth, rate limit, timeout, invalid
    model, network error, 4xx/5xx) so the caller can silently fall through to
    the next provider without ever exposing a provider error to the user.

    On Vercel/serverless, streaming responses are buffered and can be cut off,
    so we fall back to a single non-streaming request that returns the full
    reply in one go — the frontend still displays it, just not word-by-word."""
    if not api_key:
        msg = f"[{provider_label}] skipped: no API key configured"
        print(msg)
        _LAST_PROVIDER_ERRORS.append(msg)
        return

    # Non-streaming path on serverless — avoids Vercel's edge buffer cutting
    # the response short mid-stream.
    if IS_SERVERLESS:
        try:
            resp = requests.post(
                url,
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={"model": model, "messages": messages, "stream": False, "max_tokens": 2048},
                timeout=45,
            )
        except requests.RequestException as e:
            msg = f"[{provider_label}] network error: {e}"
            print(msg)
            _LAST_PROVIDER_ERRORS.append(msg)
            return
        if resp.status_code != 200:
            try: body_preview = resp.text[:500]
            except Exception: body_preview = "<unreadable>"
            msg = f"[{provider_label}] HTTP {resp.status_code}: {body_preview}"
            print(msg)
            _LAST_PROVIDER_ERRORS.append(msg)
            return
        try:
            obj = resp.json()
            content = obj["choices"][0]["message"]["content"] or ""
        except (json.JSONDecodeError, KeyError, IndexError, TypeError) as e:
            msg = f"[{provider_label}] bad JSON: {e}"
            print(msg)
            _LAST_PROVIDER_ERRORS.append(msg)
            return
        # Yield the full reply in reasonable chunks so the frontend still
        # renders progressively even though the network delivered it all at once
        step = 80
        for i in range(0, len(content), step):
            yield content[i:i + step]
        return

    # Streaming path — the normal, preferred flow on always-on hosts
    max_retries = 2
    resp = None
    for attempt in range(max_retries + 1):
        try:
            resp = requests.post(
                url,
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={"model": model, "messages": messages, "stream": True, "max_tokens": 2048},
                stream=True, timeout=60,
            )
        except requests.RequestException as e:
            msg = f"[{provider_label}] network error: {e}"
            print(msg)
            _LAST_PROVIDER_ERRORS.append(msg)
            return
        if resp.status_code == 429 and attempt < max_retries:
            # Rate-limited — brief backoff, then retry the same model/provider
            # before giving up on it entirely.
            print(f"[{provider_label}] rate-limited (429), retrying in "
                  f"{1.5 * (attempt + 1):.1f}s (attempt {attempt + 1}/{max_retries})")
            time.sleep(1.5 * (attempt + 1))
            continue
        break

    if resp.status_code != 200:
        try:
            body_preview = resp.text[:500]
        except Exception:
            body_preview = "<unreadable>"
        msg = f"[{provider_label}] HTTP {resp.status_code}: {body_preview}"
        print(msg)
        _LAST_PROVIDER_ERRORS.append(msg)
        return

    for raw_line in resp.iter_lines(decode_unicode=False):
        if not raw_line:
            continue
        try:
            line = raw_line.decode("utf-8")
        except UnicodeDecodeError:
            continue
        if not line.startswith("data:"):
            continue
        data_str = line[5:].strip()
        if data_str == "[DONE]":
            break
        try:
            obj = json.loads(data_str)
            content = obj["choices"][0]["delta"].get("content", "")
            if content:
                yield content
        except (json.JSONDecodeError, KeyError, IndexError):
            continue


def _stream_with_model_fallback(url, api_key, primary_model, fallback_model, messages, label):
    """Tries `primary_model` first (e.g. a model the person picked in the
    Model Manager). If it yields nothing at all — invalid/decommissioned
    model, 404, etc. — automatically retries once on `fallback_model` (the
    server's configured default) before giving up on this provider, so a
    stale or unavailable model choice doesn't just silently fail."""
    models_to_try = [primary_model]
    if fallback_model and fallback_model != primary_model:
        models_to_try.append(fallback_model)
    for m in models_to_try:
        got_any = False
        for chunk in _openai_style_stream(url, api_key, m, messages, f"{label}:{m}"):
            got_any = True
            yield chunk
        if got_any:
            return


def groq_stream_chunks(messages, api_key=None, model=None):
    """Stream from Groq (primary chat provider). Uses a person's own key
    (from Settings) when provided, otherwise the server's GROQ_API_KEY. An
    optional `model` overrides GROQ_MODEL for this request (falls back to
    GROQ_MODEL automatically if the override is unavailable)."""
    yield from _stream_with_model_fallback(
        "https://api.groq.com/openai/v1/chat/completions",
        api_key or GROQ_API_KEY, model or GROQ_MODEL, GROQ_MODEL, messages, "Groq",
    )


def cerebras_stream_chunks(messages, api_key=None, model=None):
    """Stream from Cerebras (automatic fallback if Groq is unavailable). Uses
    a person's own key (from Settings) when provided, otherwise the
    server's CEREBRAS_API_KEY. An optional `model` overrides CEREBRAS_MODEL
    for this request (falls back to CEREBRAS_MODEL automatically if the
    override is unavailable)."""
    yield from _stream_with_model_fallback(
        "https://api.cerebras.ai/v1/chat/completions",
        api_key or CEREBRAS_API_KEY, model or CEREBRAS_MODEL, CEREBRAS_MODEL, messages, "Cerebras",
    )


def _quick_completion(messages, api_key_groq=None, api_key_cerebras=None, max_tokens=20):
    """Non-streaming, short completion used for auxiliary tasks like AI title
    generation — tries Groq then Cerebras (same silent-fallback pattern as
    chat), returns plain text or None if both fail. Kept deliberately small
    (max_tokens) since this is just for a 3-6 word chat title, not a real
    reply, so it stays fast and cheap."""
    groq_key = (api_key_groq or "").strip() or GROQ_API_KEY
    cerebras_key = (api_key_cerebras or "").strip() or CEREBRAS_API_KEY

    for key, model, url, label in (
        (groq_key, GROQ_MODEL, "https://api.groq.com/openai/v1/chat/completions", "Groq"),
        (cerebras_key, CEREBRAS_MODEL, "https://api.cerebras.ai/v1/chat/completions", "Cerebras"),
    ):
        if not key:
            continue
        try:
            resp = requests.post(
                url,
                headers={"Authorization": f"Bearer {key}", "Content-Type": "application/json"},
                json={"model": model, "messages": messages, "stream": False, "max_tokens": max_tokens, "temperature": 0.4},
                timeout=15,
            )
            if resp.status_code != 200:
                continue
            obj = resp.json()
            content = (obj["choices"][0]["message"]["content"] or "").strip()
            if content:
                return content
        except Exception as e:
            print(f"[TitleGen/{label}] failed: {e}")
            continue
    return None


def generate_smart_title(first_user_message, first_ai_reply, api_key_groq=None, api_key_cerebras=None):
    """Asks the AI for a short, natural chat title based on the first
    exchange — the same pattern ChatGPT/Claude use, instead of just
    truncating the raw first message. Falls back to make_title() if the
    AI call fails for any reason, so a title is always produced."""
    user_msg = (first_user_message or "").strip()
    user_msg = re.sub(r'^\[Instructions:.*?\]\s*', '', user_msg, flags=re.DOTALL)
    ai_reply = (first_ai_reply or "").strip()

    if not user_msg and not ai_reply:
        return "New chat"

    # Deterministic check for simple greetings — guarantees the title is
    # always exactly "Greeting" for these, instead of trusting the AI to
    # follow that instruction (which can occasionally hallucinate an
    # unrelated title instead).
    greeting_only = re.sub(r'[^a-z]', '', user_msg.lower())
    _GREETING_WORDS = {
        "hi", "hii", "hiii", "hello", "hellooo", "hey", "heyy", "heyyy",
        "yo", "sup", "hola", "namaste", "hii there", "helloo",
    }
    if greeting_only in _GREETING_WORDS:
        return "Greeting"

    prompt = (
        "Generate a short, natural chat title (1-4 words, no quotes, no punctuation "
        "at the end, no emoji, title case) that summarizes what this SPECIFIC "
        "conversation is about. Do NOT describe the task of generating a title — "
        "write about the actual topic discussed. Prefer a simple category-style "
        "label when that fits — for a simple greeting like 'hi'/'hello', use "
        "'Greeting'. Reply with ONLY the title text, nothing else.\n\n"
        f"User: {user_msg[:400]}\n"
        f"Assistant: {ai_reply[:400]}"
    )
    messages = [
        {"role": "system", "content": "You generate concise chat titles. Reply with only the title, no extra text."},
        {"role": "user", "content": prompt},
    ]
    result = _quick_completion(messages, api_key_groq, api_key_cerebras, max_tokens=16)
    if result:
        title = result.strip()
        title = title.strip('"\'“”‘’').strip()
        title = title.split("\n")[0].strip()
        title = re.sub(r'^\[.*?\]\s*', '', title).strip()  # strip any leaked [Instructions: ...] / bracketed prefix
        title = re.sub(r'[.!?]+$', '', title).strip()
        _meta_phrases = ("generating chat title", "chat title", "conversation title",
                          "title generation", "generate title", "untitled")
        looks_bad = (
            len(title) < 3 or
            any(ch in title for ch in '[]"“”') or
            title.lower().startswith(('based on', 'here is', 'here\'s', 'sure,', 'title:')) or
            any(p in title.lower() for p in _meta_phrases)
        )
        if title and not looks_bad:
            return title[:60]
    return make_title(first_user_message)


def auto_stream_chunks(gemini_payload, gemini_messages, system_prompt=None,
                        user_groq_key=None, user_cerebras_key=None,
                        groq_model=None, cerebras_model=None):
    """Groq first, Cerebras as a silent automatic fallback.
    Never asks the user to pick a provider and never exposes provider errors —
    if Groq yields nothing (rate limit, timeout, invalid model, network error,
    4xx/5xx), we just move on to Cerebras with no visible interruption.
    If the person supplied their own API key(s) in Settings, those are tried
    first (and exclusively, in that provider's slot) before the server's key.
    `groq_model`/`cerebras_model` optionally override the chosen model within
    each provider (see Model Manager) — invalid choices auto-fall-back to the
    server's default model for that provider before moving to the next provider."""
    _LAST_PROVIDER_ERRORS.clear()
    sp = system_prompt or SYSTEM_PROMPT
    groq_key = (user_groq_key or "").strip() or GROQ_API_KEY
    cerebras_key = (user_cerebras_key or "").strip() or CEREBRAS_API_KEY

    # If the most recent user turn carries an image (a Video Call frame, a
    # Screen Share frame, or a regular image attachment), try the vision
    # model FIRST so Mythic AI can actually see it. Falls through to the
    # normal text-only path below if the vision call fails or yields nothing
    # (the image is simply dropped from context at that point).
    has_image = False
    if gemini_messages:
        last = gemini_messages[-1]
        if last.get("role") == "user":
            has_image = any(
                str((p.get("inline_data") or {}).get("mime_type", "")).startswith("image/")
                for p in last.get("parts", [])
            )
    if has_image and PROVIDER in ("auto", "groq") and groq_key:
        vision_msgs = to_openai_vision_messages(gemini_messages, sp)
        collected = False
        try:
            for chunk in groq_vision_stream_chunks(vision_msgs, groq_key):
                collected = True
                yield chunk
        except Exception as e:
            print(f"[Groq-Vision] unexpected error: {e}")
        if collected:
            return

    openai_msgs = to_openai_messages(gemini_messages, sp)

    order = []
    if PROVIDER in ("auto", "groq") and groq_key:
        order.append(("Groq", lambda: groq_stream_chunks(openai_msgs, groq_key, groq_model)))
    if PROVIDER in ("auto", "cerebras") and cerebras_key:
        order.append(("Cerebras", lambda: cerebras_stream_chunks(openai_msgs, cerebras_key, cerebras_model)))

    if not order:
        # Kept short and free of internal setup instructions — see chat
        # request for context.
        yield "I'm not able to reply right now — please try again shortly."
        return

    for _name, fn in order:
        collected = False
        try:
            for chunk in fn():
                collected = True
                yield chunk
            if collected:
                return
        except Exception as e:
            print(f"[{_name}] unexpected error: {e}")

    # All configured providers failed. TEMP DEBUG: surface the actual reason
    # directly in the chat reply so it's visible without checking server
    # logs. Remove this block (and go back to the generic message) once
    # providers are confirmed working.
    debug_detail = " | ".join(_LAST_PROVIDER_ERRORS) if _LAST_PROVIDER_ERRORS else "no providers attempted"
    yield f"I couldn't get a reply just now.\n\n[DEBUG] {debug_detail}"


# --- Model selector (cosmetic tiers over the same underlying providers) -----
VIP_PASSWORD = os.environ.get("VIP_PASSWORD", "1254")

MODEL_CATALOG = [
    {"id": "mythic-1", "name": "Mythic 1", "vip": False},
    {"id": "mythic-2", "name": "Mythic 2", "vip": False},
    {"id": "mythic-3", "name": "Mythic 3", "vip": False},
    {"id": "mythic-vip", "name": "Mythic VIP 🔒", "vip": True},
]
DEFAULT_MODEL_ID = "mythic-2"


@app.route("/api/models", methods=["GET"])
@login_required
def api_models():
    return jsonify({"models": MODEL_CATALOG, "default": DEFAULT_MODEL_ID})


@app.route("/api/vip-status", methods=["GET"])
@login_required
def api_vip_status():
    return jsonify({"vip": bool(session.get("vip_unlocked"))})


@app.route("/api/vip-unlock", methods=["POST"])
@login_required
def api_vip_unlock():
    data = request.get_json(force=True) or {}
    password = data.get("password") or ""
    if VIP_PASSWORD and password == VIP_PASSWORD:
        session["vip_unlocked"] = True
        session.permanent = True
        return jsonify({"success": True})
    return jsonify({"success": False})


@app.route("/api/code/run", methods=["POST"])
@login_required
def api_code_run():
    """Runs a snippet in Code Workspace for any non-web language (Python,
    C++, C, Java, Node.js, TypeScript, Go, Ruby) via the Judge0 execution
    engine and returns stdout/stderr/compile output. See JUDGE0_API_KEY
    above for how to point this at a more reliable host."""
    data = request.get_json(silent=True) or {}
    lang = (data.get("language") or "").strip().lower()
    code = data.get("code") or ""
    stdin = data.get("stdin") or ""

    lang_id = JUDGE0_LANGUAGE_IDS.get(lang)
    if lang_id is None:
        return jsonify({"error": f"Unsupported language: {lang or '(none)'}"}), 400
    if not code.strip():
        return jsonify({"error": "There's no code to run yet."}), 400
    if len(code) > 60000:
        return jsonify({"error": "That's too much code for the online runner — trim it down a bit."}), 400

    headers = {"Content-Type": "application/json"}
    if JUDGE0_API_KEY:
        headers["X-RapidAPI-Key"] = JUDGE0_API_KEY
        headers["X-RapidAPI-Host"] = JUDGE0_API_HOST

    payload = {
        "language_id": lang_id,
        "source_code": base64.b64encode(code.encode("utf-8")).decode("ascii"),
        "stdin": base64.b64encode(stdin.encode("utf-8")).decode("ascii") if stdin else "",
    }

    try:
        resp = requests.post(
            f"{JUDGE0_BASE_URL}/submissions",
            params={"base64_encoded": "true", "wait": "true"},
            headers=headers, json=payload, timeout=25,
        )
    except requests.exceptions.RequestException:
        return jsonify({"error": "The code runner is unreachable right now. Please try again in a moment."}), 502

    if resp.status_code == 401 or resp.status_code == 403:
        return jsonify({"error": "Code runner authorization failed. If you're using RapidAPI, double-check JUDGE0_API_KEY."}), 502
    if resp.status_code == 429:
        return jsonify({"error": "The code runner is rate-limited right now — wait a few seconds and try again."}), 429
    if not resp.ok:
        return jsonify({"error": f"Code runner returned an error (HTTP {resp.status_code})."}), 502

    try:
        result = resp.json()
    except ValueError:
        return jsonify({"error": "Code runner returned an invalid response."}), 502

    def _decode(s):
        if not s:
            return ""
        try:
            return base64.b64decode(s).decode("utf-8", errors="replace")
        except Exception:
            return ""

    status = (result.get("status") or {}).get("description", "Unknown")
    return jsonify({
        "stdout": _decode(result.get("stdout")),
        "stderr": _decode(result.get("stderr")),
        "compile_output": _decode(result.get("compile_output")),
        "message": _decode(result.get("message")),
        "status": status,
        "time": result.get("time"),
        "memory": result.get("memory"),
    })


@app.route("/api/streak", methods=["GET"])
@login_required
def api_streak():
    """Returns the current user's daily chat streak length (0 if they've
    never chatted or their streak has already lapsed)."""
    username = current_username()
    return jsonify({"streak": _get_user_streak(username)})


@app.route("/api/push/vapid-public-key", methods=["GET"])
def push_vapid_key():
    if not VAPID_PUBLIC_KEY:
        return jsonify({"error": "push not configured"}), 503
    # length_ok: a valid P-256 uncompressed-point key, base64url-encoded
    # with no padding, is always exactly 87 characters. If this is false,
    # the env var was pasted with extra/missing characters (whitespace,
    # truncation, wrong key type) and the browser will reject it with
    # "applicationServerKey is not valid".
    return jsonify({
        "publicKey": VAPID_PUBLIC_KEY,
        "length": len(VAPID_PUBLIC_KEY),
        "length_ok": len(VAPID_PUBLIC_KEY) == 87,
    })


@app.route("/api/push/subscribe", methods=["POST"])
@login_required
def push_subscribe():
    """Save (or update) a browser's push subscription, tagged with the
    current (anonymous) username so re-engagement notifications can be
    targeted at this specific person later."""
    data = request.get_json(force=True) or {}
    sub = data.get("subscription")
    if not sub or not sub.get("endpoint"):
        return jsonify({"error": "invalid subscription"}), 400
    sub_id = str(uuid.uuid5(uuid.NAMESPACE_URL, sub["endpoint"]))
    sub = dict(sub)
    sub["_username"] = current_username()
    _save_push_subscription(sub_id, sub)
    return jsonify({"status": "subscribed", "id": sub_id})


@app.route("/api/push/unsubscribe", methods=["POST"])
@login_required
def push_unsubscribe():
    data = request.get_json(force=True) or {}
    endpoint = (data.get("endpoint") or "").strip()
    if endpoint:
        sub_id = str(uuid.uuid5(uuid.NAMESPACE_URL, endpoint))
        _delete_push_subscription(sub_id)
    return jsonify({"status": "unsubscribed"})


@app.route("/api/push/test-reengagement", methods=["POST"])
@login_required
def push_test_reengagement():
    """Manually fire one reengagement pass right now (bypassing the hourly
    wait) so you can verify VAPID/subscriptions/logic are actually working
    without sitting around for an hour. Remove or protect this in production."""
    if not _PUSH_AVAILABLE:
        return jsonify({"error": "pywebpush not installed"}), 500
    if not VAPID_PRIVATE_KEY or not VAPID_PUBLIC_KEY:
        return jsonify({"error": "VAPID_PRIVATE_KEY / VAPID_PUBLIC_KEY not set on the server"}), 400
    _load_push_subscriptions()
    usernames = _subscribed_usernames()
    if not usernames:
        return jsonify({"error": "no push subscriptions registered yet — enable notifications in Settings first"}), 400
    _run_reengagement_pass()
    return jsonify({"status": "ran", "subscribed_users": len(usernames)})


@app.route("/api/push/test", methods=["POST"])
@login_required
def push_test():
    send_push_notification(
        title="Mythic AI",
        body="🎉 Push notifications are working!",
        url="/",
    )
    return jsonify({"status": "sent"})


# ── Cron endpoint for scheduled notifications ─────────────────────────────────
# SCHEDULE SUMMARY:
#   - Render (or any always-on host): nothing to configure — the in-process
#     background thread (_reengagement_loop) fires automatically every
#     _REENGAGEMENT_CHECK_INTERVAL_SECONDS (1 hour).
#   - Vercel (serverless): the in-process thread never runs there (functions
#     freeze between requests), so this endpoint must be triggered by an
#     EXTERNAL cron, once a day at 12:00 (noon). Configure it with:
#
#   vercel.json:
#     {
#       "crons": [
#         { "path": "/api/cron/reengagement", "schedule": "0 12 * * *" }
#       ]
#     }
#
#   "0 12 * * *" = once a day at 12:00 in the project's configured cron
#   timezone (UTC by default on Vercel — adjust the hour if you need a
#   specific local noon).
#
# Protection: set CRON_SECRET as an environment variable and pass it in a
# header:  Authorization: Bearer <CRON_SECRET>
#          or in a query string: /api/cron/reengagement?secret=<CRON_SECRET>
# If CRON_SECRET is unset, the endpoint is open (fine for dev only).
CRON_SECRET = os.environ.get("CRON_SECRET", "").strip()

@app.route("/api/cron/reengagement", methods=["GET", "POST"])
def cron_reengagement():
    if CRON_SECRET:
        provided = ""
        auth = request.headers.get("Authorization", "")
        if auth.startswith("Bearer "):
            provided = auth[7:].strip()
        elif "x-vercel-cron-secret" in request.headers:
            provided = request.headers.get("x-vercel-cron-secret", "").strip()
        else:
            provided = request.args.get("secret", "").strip()
        if provided != CRON_SECRET:
            return jsonify({"error": "unauthorized"}), 401
    try:
        _run_reengagement_pass()
        return jsonify({"status": "ok", "subscribers": len(_push_subscriptions)})
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── /api/health: quick diagnostic so admins can see what's configured ────────
# Gated behind ADMIN_SECRET — without it, anyone who guesses this URL could
# see which underlying AI providers (Groq/Cerebras) and model names power
# the app. Always set ADMIN_SECRET in production (Render env vars).
@app.route("/api/health", methods=["GET"])
def api_health():
    if not _require_admin():
        return jsonify({"status": "ok"}), 200
    supabase_status = {"configured": bool(SUPABASE_URL and SUPABASE_KEY)}
    if supabase_status["configured"]:
        try:
            r = requests.get(sb("conversations?limit=1"), headers=sb_headers(), timeout=8)
            supabase_status["reachable"] = r.status_code == 200
            supabase_status["status_code"] = r.status_code
            if r.status_code != 200:
                supabase_status["error"] = r.text[:300]
        except Exception as e:
            supabase_status["reachable"] = False
            supabase_status["error"] = str(e)

        # A GET succeeding only proves the API key + table exist for reads.
        # The actual bug people hit is on INSERT (missing column, wrong
        # type, RLS blocking writes) — so actually try writing and deleting
        # a throwaway row and report the exact PostgREST error if it fails.
        if supabase_status.get("reachable"):
            test_id = "healthcheck-" + uuid.uuid4().hex[:8]
            try:
                w = requests.post(
                    sb("conversations"),
                    headers={**sb_headers(), "Prefer": "return=minimal"},
                    json={
                        "id": test_id, "username": "healthcheck", "title": "healthcheck",
                        "updated_at": time.time(), "messages": [],
                        "folder": None, "pinned": False, "archived": False,
                    },
                    timeout=10,
                )
                supabase_status["write_test"] = {
                    "status_code": w.status_code,
                    "success": w.status_code in (200, 201, 204),
                }
                if w.status_code not in (200, 201, 204):
                    supabase_status["write_test"]["error"] = w.text[:500]
                else:
                    requests.delete(
                        sb(f"conversations?id=eq.{test_id}"),
                        headers=sb_headers(), timeout=10,
                    )
            except Exception as e:
                supabase_status["write_test"] = {"success": False, "error": str(e)}

    return jsonify({
        "app": "Mythic AI",
        "serverless": IS_SERVERLESS,
        "providers": {
            "groq":     {"configured": bool(GROQ_API_KEY),     "model": GROQ_MODEL},
            "cerebras": {"configured": bool(CEREBRAS_API_KEY), "model": CEREBRAS_MODEL},
        },
        "storage": {
            "supabase": supabase_status,
            "using": "supabase" if supabase_status["configured"] else "local_json_files",
            "note": ("If 'configured' is true but 'reachable' is false or status_code isn't 200, "
                     "check that the 'conversations' table exists with the right columns, and that "
                     "SUPABASE_KEY is your service-role/secret key with insert/select access — see "
                     "the 'error' field above for the exact Supabase response.") if supabase_status["configured"] else
                    ("No SUPABASE_URL set — conversations are saved to local disk on the server. "
                     "On Render this works fine while the instance stays up, but is wiped on redeploy "
                     "and isn't shared across instances."),
            "shares_note": ("Share links (/share/<id>) use a separate `shares` table in Supabase when "
                             "configured: columns id (text, primary key), username (text), conv_id (text), "
                             "revoked (bool, default false), created_at (float8). If that table doesn't "
                             "exist yet, share links automatically fall back to local storage instead of "
                             "breaking, but they'll be lost on redeploy — create the table for persistence."),
        },
        "image_generation": {
            "nano_banana":  bool(NANO_BANANA_API_KEY),
            "huggingface":  bool(HF_API_KEY),
            "pollinations": True,  # always available, no key needed
        },
        "push_notifications": {
            "configured": _PUSH_AVAILABLE and bool(VAPID_PRIVATE_KEY) and bool(VAPID_PUBLIC_KEY),
            "subscribers": len(_push_subscriptions),
            "cron_endpoint": "/api/cron/reengagement",
            "cron_schedule": "Render: every 1 hour (built-in thread). Vercel: once daily at 12:00 via external cron.",
            "cron_secret_set": bool(CRON_SECRET),
        },
        "hint": ("Add GROQ_API_KEY or CEREBRAS_API_KEY as environment variables "
                 "if 'configured' is false. On Vercel, also set up a cron job "
                 "hitting /api/cron/reengagement once a day at 12:00 for notifications.")
    })


# --- Weather (Open-Meteo — free, no API key needed, works for any country/city) ---
_WMO_ICON = {
    0: ("☀️", "Clear sky"), 1: ("🌤", "Mainly clear"), 2: ("⛅", "Partly cloudy"), 3: ("☁️", "Overcast"),
    45: ("🌫", "Fog"), 48: ("🌫", "Freezing fog"),
    51: ("🌦", "Light drizzle"), 53: ("🌦", "Drizzle"), 55: ("🌧", "Dense drizzle"),
    56: ("🌧", "Freezing drizzle"), 57: ("🌧", "Freezing drizzle"),
    61: ("🌧", "Light rain"), 63: ("🌧", "Rain"), 65: ("🌧", "Heavy rain"),
    66: ("🌧", "Freezing rain"), 67: ("🌧", "Freezing rain"),
    71: ("🌨", "Light snow"), 73: ("🌨", "Snow"), 75: ("❄️", "Heavy snow"), 77: ("❄️", "Snow grains"),
    80: ("🌦", "Rain showers"), 81: ("🌧", "Rain showers"), 82: ("⛈", "Violent rain showers"),
    85: ("🌨", "Snow showers"), 86: ("❄️", "Snow showers"),
    95: ("⛈", "Thunderstorm"), 96: ("⛈", "Thunderstorm with hail"), 99: ("⛈", "Thunderstorm with hail"),
}


def _wmo(code):
    return _WMO_ICON.get(code, ("🌡", "Unknown"))


def _aqi_label(us_aqi):
    if us_aqi is None:
        return None
    if us_aqi <= 50: return f"{us_aqi} (Good)"
    if us_aqi <= 100: return f"{us_aqi} (Moderate)"
    if us_aqi <= 150: return f"{us_aqi} (Unhealthy for sensitive groups)"
    if us_aqi <= 200: return f"{us_aqi} (Unhealthy)"
    if us_aqi <= 300: return f"{us_aqi} (Very unhealthy)"
    return f"{us_aqi} (Hazardous)"


# ── "Paste URL" helpers: ordinary webpages + flipbook viewers ────────────────

# Known flipbook/page-turn viewer platforms — if the URL's host matches one
# of these, it's almost certainly a JS canvas/image viewer with no real text
# in the page source, so we go straight to the embedded-PDF search / OCR path
# instead of trying (and failing) to scrape "readable text" from the shell.
_FLIPBOOK_DOMAINS = (
    "flippingbook.com", "issuu.com", "yumpu.com", "calameo.com",
    "anyflip.com", "joomag.com", "publitas.com", "heyzine.com",
    "flipsnack.com", "fliphtml5.com", "mmdigital.co.in",
)
# Telltale strings that show up in a flipbook viewer's HTML/JS even on
# self-hosted / white-labelled platforms we don't know by domain.
_FLIPBOOK_HTML_SIGNATURES = (
    "flipbook", "flip-book", "df-page", "page-flip", "pageflip",
    "turn.js", "df-parent", "flippingbook",
)


def _looks_like_flipbook(url: str, html: str) -> bool:
    host = urllib.parse.urlparse(url).netloc.lower()
    if any(d in host for d in _FLIPBOOK_DOMAINS):
        return True
    sample = (html or "")[:20000].lower()
    return any(sig in sample for sig in _FLIPBOOK_HTML_SIGNATURES)


def _find_embedded_document_link(html: str, base_url: str):
    """Many flipbook widgets (and some webpages) are generated from an
    uploaded PDF/DOCX that's still linked somewhere in the HTML or JS config
    — e.g. a hidden 'download original' link, or a `"pdfUrl": "...pdf"` entry
    in an embedded script. Grabbing that directly is far more reliable than
    OCR, so we check for it first. Returns an absolute URL or None."""
    if not html:
        return None
    candidates = []
    # href="....pdf" / .docx / .txt (typical <a> download links)
    for m in re.finditer(r'''href=["']([^"']+\.(?:pdf|docx|txt))(?:[?"'][^"']*)?["']''', html, re.IGNORECASE):
        candidates.append(m.group(1))
    # "file": "....pdf" / "pdfUrl": "....pdf" style JS config values
    for m in re.finditer(r'''["'](?:pdf|pdfUrl|file|source|bookFile|fileUrl)["']\s*:\s*["']([^"']+\.(?:pdf|docx|txt))["']''', html, re.IGNORECASE):
        candidates.append(m.group(1))
    for c in candidates:
        absolute = urllib.parse.urljoin(base_url, c)
        if absolute.lower().split("?")[0].endswith((".pdf", ".docx", ".txt")):
            return absolute
    return None


def _extract_readable_webpage_text(html: str) -> str:
    """Pulls readable body text out of an ordinary webpage, stripping nav,
    scripts, styles, and other chrome. Uses BeautifulSoup when available;
    falls back to a crude regex tag-strip so this still degrades gracefully
    if bs4 isn't installed."""
    if _BS4_AVAILABLE:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "header", "footer", "noscript", "svg", "form"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
    else:
        text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", html)
        text = re.sub(r"(?s)<[^>]+>", " ", text)
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]
    return "\n".join(lines)


def _webpage_title(html: str, fallback: str) -> str:
    m = re.search(r"(?is)<title[^>]*>(.*?)</title>", html or "")
    if m:
        title = re.sub(r"\s+", " ", m.group(1)).strip()
        if title:
            return title
    return fallback


def _resolve_basic_html_index(url: str, html: str):
    """FlipBuilder-family flipbook software ('Flip PDF', and similar clones)
    publishes an accessibility/SEO fallback under a 'basic-html' path
    alongside its JS page-turn viewer. Per FlipBuilder's own documentation,
    that fallback is generated by extracting the *real text* of the source
    PDF for search engines — so it's real readable text, not just images.
    Returns the absolute URL of that fallback's index page, or None."""
    if "basic-html" in url.lower():
        return url  # the user may have pasted this fallback URL directly
    m = re.search(r'''href=["']([^"']*basic-html/index\.html[^"']*)["']''', html, re.IGNORECASE)
    if m:
        return urllib.parse.urljoin(url, m.group(1))
    return None


def _extract_flipbuilder_basic_html_text(index_url: str, max_pages: int = 300):
    """Reads a FlipBuilder-style 'basic-html' fallback page by page. Returns
    (text, note) on success, or (None, reason) if it turns out to be
    image-only after all (some publishers disable the text-extraction option)."""
    try:
        r = requests.get(index_url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
        r.raise_for_status()
    except requests.RequestException as e:
        return None, f"Couldn't reach the basic-HTML version of this flipbook: {e}"

    html = r.text
    found = re.findall(r'''href=["']([^"']*/page(\d+)\.html)["']''', html, re.IGNORECASE)
    if not found:
        return None, "This looks like a Flip PDF/FlipBuilder-style flipbook, but no basic-HTML pages were found."
    by_num = {}
    for frag, num in found:
        by_num.setdefault(int(num), frag)
    ordered_frags = [by_num[n] for n in sorted(by_num)][:max_pages]

    texts = []
    for frag in ordered_frags:
        page_url = urllib.parse.urljoin(index_url, frag)
        try:
            pr = requests.get(page_url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
            pr.raise_for_status()
        except requests.RequestException:
            continue
        page_text = _extract_readable_webpage_text(pr.text).strip()
        if page_text:
            texts.append(page_text)

    full_text = "\n\n".join(texts).strip()
    if not full_text:
        return None, ("Found this flipbook's basic-HTML pages, but they contained no extractable "
                       "text (this publisher likely disabled text-extraction, leaving image-only pages).")
    note = f"Read {len(texts)} page(s) from this flipbook's built-in accessible text version (no OCR needed)."
    return full_text, note


def _extract_flipbook_via_ocr(url: str, max_pages: int = 40):
    """Best-effort flipbook reader: opens the page in a headless browser,
    clicks through it page by page, screenshots the viewer, and OCRs each
    screenshot. This is inherently fragile — every flipbook platform uses a
    different DOM/selector for its "next page" control — so it tries a list
    of common selectors and stops early if none of them seem to advance the
    book. Returns (text, note) or (None, error_message).

    Requires: playwright (+ `playwright install chromium`), pytesseract,
    and the tesseract-ocr binary installed on the host. See the import block
    near the top of this file for install instructions."""
    if not _FLIPBOOK_OCR_AVAILABLE:
        missing = []
        if not _PLAYWRIGHT_AVAILABLE:
            missing.append("playwright (`pip install playwright` + `playwright install chromium`)")
        if not _OCR_AVAILABLE:
            missing.append("pytesseract (`pip install pytesseract`) + the tesseract-ocr system package")
        if not _WATERMARK_AVAILABLE:
            missing.append("Pillow (`pip install Pillow`)")
        return None, ("This looks like a flipbook viewer. Reading it requires OCR support that "
                       "isn't installed on this server yet. Missing: " + "; ".join(missing) + ".")

    # Common "next page" selectors across popular flipbook platforms —
    # tried in order; the first one that's clickable AND changes the page
    # content wins. Add more as you encounter new platforms.
    NEXT_SELECTORS = [
        ".df-next-page", ".flipbook-next", ".next-page", "[aria-label='Next page']",
        "[aria-label='next']", ".turn-page-next", "button.next", ".df-arrow-right",
    ]
    texts = []
    try:
        with sync_playwright() as p:
            # --no-sandbox: Chromium refuses to run as root without this, and
            # most container platforms (including Render's Docker runtime) run
            # as root by default. --disable-dev-shm-usage avoids crashes from
            # Docker's small default /dev/shm size.
            browser = p.chromium.launch(args=["--no-sandbox", "--disable-dev-shm-usage"])
            page = browser.new_page(viewport={"width": 1400, "height": 1000})
            page.goto(url, timeout=30000, wait_until="networkidle")
            page.wait_for_timeout(2000)  # let the viewer finish initial render

            next_selector = None
            for sel in NEXT_SELECTORS:
                if page.locator(sel).count() > 0:
                    next_selector = sel
                    break

            seen_hashes = set()
            for i in range(max_pages):
                shot = page.screenshot()
                img_hash = hashlib.md5(shot).hexdigest()
                if img_hash in seen_hashes:
                    break  # page didn't change — we've reached the end or got stuck
                seen_hashes.add(img_hash)
                try:
                    from io import BytesIO
                    page_text = pytesseract.image_to_string(Image.open(BytesIO(shot)))
                except Exception:
                    page_text = ""
                if page_text.strip():
                    texts.append(page_text.strip())
                if not next_selector:
                    break  # can't advance — only OCR the first visible page/spread
                try:
                    page.click(next_selector, timeout=3000)
                    page.wait_for_timeout(900)
                except Exception:
                    break
            browser.close()
    except Exception as e:
        return None, f"Flipbook OCR failed: {e}"

    full_text = "\n\n".join(texts).strip()
    if not full_text:
        return None, ("Couldn't read any text from this flipbook — OCR ran but found nothing "
                       "recognizable. The page images may be too low-resolution or stylized to OCR.")
    note = f"Read via OCR from a flipbook viewer ({len(texts)} page(s) scanned — may be incomplete)."
    return full_text, note


@app.route("/api/fetch-url-document", methods=["POST"])
@login_required
def api_fetch_url_document():
    """Pulls readable text from whatever kind of 'book URL' the user pastes:
      1. A direct PDF/DOCX/TXT download link       -> parsed like a file upload
      2. An ordinary webpage (article, story, etc.) -> readable text scraped from the HTML
      3. A flipbook-style viewer (FlippingBook, Issuu, Yumpu, mmdigital, ...)
           a. first tries to find an embedded/original PDF link on the page
           b. falls back to headless-browser + OCR page-by-page (if configured)
    Caps document downloads to DOCUMENT_UPLOAD_BYTES, streaming so we abort
    early instead of reading an arbitrarily large body first."""
    DOC_EXTENSIONS = (".pdf", ".docx", ".doc", ".txt")
    DOC_CONTENT_TYPES = (
        "application/pdf",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
    )

    data = request.get_json(force=True) or {}
    url = (data.get("url") or "").strip()
    if not url or not (url.startswith("http://") or url.startswith("https://")):
        return jsonify({"error": "Please enter a valid http:// or https:// URL"}), 400

    def _download_and_extract_doc(doc_url, timeout=20):
        """Shared path for downloading + parsing a direct PDF/DOCX/TXT link."""
        r = requests.get(doc_url, timeout=timeout, stream=True, headers={"User-Agent": "Mozilla/5.0"})
        r.raise_for_status()
        ct = r.headers.get("Content-Type", "").split(";")[0].strip()
        raw = bytearray()
        for chunk in r.iter_content(chunk_size=65536):
            raw.extend(chunk)
            if len(raw) > DOCUMENT_UPLOAD_BYTES:
                limit_mb = DOCUMENT_UPLOAD_BYTES // (1024 * 1024)
                raise ValueError(f"That file is larger than {limit_mb}MB — please download it and upload it directly instead.")
        fname = (doc_url.split("/")[-1].split("?")[0] or "document").strip() or "document"
        txt, note = extract_text_from_attachment(fname, ct, bytes(raw))
        return txt, note, fname, ct

    path = urllib.parse.urlparse(url).path.lower()

    # ── Case 1: direct document link ─────────────────────────────────────
    if path.endswith(DOC_EXTENSIONS):
        try:
            text, note, filename, content_type = _download_and_extract_doc(url)
        except ValueError as e:
            return jsonify({"error": str(e)}), 400
        except requests.RequestException:
            return jsonify({"error": "Could not reach that link — please double-check it's a direct, "
                                      "publicly accessible file link and try again."}), 502
        if content_type and content_type not in DOC_CONTENT_TYPES:
            return jsonify({"error": "That link didn't actually return a PDF/DOCX/TXT file "
                                      "(the server sent something else back)."}), 400
        if text is None:
            return jsonify({"error": "Couldn't read that as a PDF, DOCX, or TXT file."}), 400
        if not text.strip():
            return jsonify({"error": note or "No readable text was found in that document."}), 400
        return jsonify({"text": text, "note": note, "filename": filename})

    # ── Fetch the page as HTML so we can figure out what we're looking at ──
    try:
        resp = requests.get(url, timeout=20, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
    except requests.RequestException:
        return jsonify({"error": "Could not reach that link — please double-check the URL and try again."}), 502

    content_type = resp.headers.get("Content-Type", "").split(";")[0].strip()

    # A no-extension URL might still redirect straight to a real document.
    if content_type in DOC_CONTENT_TYPES:
        fname = (url.split("/")[-1].split("?")[0] or "document").strip() or "document"
        text, note = extract_text_from_attachment(fname, content_type, resp.content)
        if text and text.strip():
            return jsonify({"text": text, "note": note, "filename": fname})

    html = resp.text

    # ── Case 2/3: figure out if this is a flipbook, and try the embedded-PDF shortcut either way ──
    embedded_doc_url = _find_embedded_document_link(html, url)
    if embedded_doc_url:
        try:
            text, note, filename, doc_ct = _download_and_extract_doc(embedded_doc_url)
            if text and text.strip() and (not doc_ct or doc_ct in DOC_CONTENT_TYPES):
                note = (note + " " if note else "") + "(Found the original document embedded in that page.)"
                return jsonify({"text": text, "note": note.strip(), "filename": filename})
        except Exception:
            pass  # fall through to OCR / webpage scraping below

    if _looks_like_flipbook(url, html):
        text, note_or_error = _extract_flipbook_via_ocr(url)
        if text:
            title = _webpage_title(html, "flipbook")
            return jsonify({"text": text, "note": note_or_error, "filename": f"{title}.txt"})
        return jsonify({"error": note_or_error}), 501 if not _FLIPBOOK_OCR_AVAILABLE else 502

    # ── Case 2: ordinary webpage ─────────────────────────────────────────
    text = _extract_readable_webpage_text(html)
    if not text.strip():
        return jsonify({"error": "That page didn't have any readable text to extract."}), 400
    title = _webpage_title(html, url.split("/")[-1] or "page")
    return jsonify({"text": text, "note": "Extracted readable text from a webpage.", "filename": f"{title}.txt"})


# ── Live news (Google News RSS — free, no API key needed) ───────────────────
# Groq/Cerebras models have no live web access at all (see SYSTEM_PROMPT),
# so "tell me today's news" could never be answered with anything real
# without this. Google News publishes a public RSS feed per search query/
# topic with no authentication required, which is enough to give the model
# actual, current headlines to summarize instead of guessing.
import xml.etree.ElementTree as _ET

_NEWS_TOPIC_QUERIES = {
    "top": "when:1d",
    "world": "world news when:1d",
    "india": "India news when:1d",
    "business": "business news when:1d",
    "technology": "technology news when:1d",
    "sports": "sports news when:1d",
    "entertainment": "entertainment news when:1d",
    "science": "science news when:1d",
    "health": "health news when:1d",
}


def fetch_google_news(query, max_items=8):
    """Returns (list_of_items, error). Each item: {title, source, link,
    published}. Uses Google News' public RSS search endpoint — no API key,
    no rate-limit key required for light use."""
    try:
        params = {"q": query, "hl": "en-US", "gl": "US", "ceid": "US:en"}
        url = "https://news.google.com/rss/search?" + urllib.parse.urlencode(params)
        resp = requests.get(url, timeout=12, headers={"User-Agent": "Mozilla/5.0"})
        resp.raise_for_status()
        root = _ET.fromstring(resp.content)
        items = []
        for item in root.findall(".//item")[:max_items]:
            title = (item.findtext("title") or "").strip()
            link = (item.findtext("link") or "").strip()
            pub_date = (item.findtext("pubDate") or "").strip()
            source_el = item.find("source")
            source = (source_el.text or "").strip() if source_el is not None else ""
            # Google News RSS titles are usually "Headline - Source" — strip
            # the trailing " - Source" if we already have it separately, to
            # avoid showing the source name twice.
            if source and title.endswith(" - " + source):
                title = title[: -(len(source) + 3)].strip()
            if title:
                items.append({"title": title, "source": source, "link": link, "published": pub_date})
        return items, None
    except requests.RequestException as e:
        return [], f"Could not reach the news service: {e}"
    except _ET.ParseError as e:
        return [], f"Could not read the news feed: {e}"


@app.route("/api/news", methods=["POST"])
@login_required
def api_news():
    data = request.get_json(force=True) or {}
    topic = (data.get("topic") or "").strip().lower()
    query = (data.get("query") or "").strip()

    if query:
        search_q = query
        heading = f'News about "{query}"'
    elif topic and topic in _NEWS_TOPIC_QUERIES:
        search_q = _NEWS_TOPIC_QUERIES[topic]
        heading = topic.capitalize() + " News"
    else:
        search_q = _NEWS_TOPIC_QUERIES["top"]
        heading = "Top Stories"

    items, error = fetch_google_news(search_q, max_items=10)
    if error and not items:
        return jsonify({"error": error}), 502
    return jsonify({"heading": heading, "items": items})


# Detects a message that's actually asking for live news, so /api/chat can
# transparently fetch real current headlines and hand them to the model as
# grounding context — instead of the model just guessing or refusing since
# it has no live web access on its own (see SYSTEM_PROMPT).
_NEWS_INTENT_RE = re.compile(
    r"\b(latest|today'?s?|current|breaking|recent)\b.{0,20}\bnews\b"
    r"|\bnews\b.{0,20}\b(today|now|update|headlines?)\b"
    r"|\bwhat'?s\s+(happening|going on)\b"
    r"|\btop\s+(headlines?|stories)\b"
    r"|\bnews\s+about\b",
    re.IGNORECASE,
)


def _maybe_fetch_news_context(user_message):
    """If the message looks like a news request, fetches real current
    headlines and returns a context block to prepend to the model's view of
    the conversation. Returns None if the message isn't news-related, or if
    the fetch failed (in which case the model just answers normally without
    pretending to have live data — see SYSTEM_PROMPT's honesty rule)."""
    if not user_message or not _NEWS_INTENT_RE.search(user_message):
        return None
    # Pull out a specific subject after "news about X" / "X news" if present,
    # otherwise just fetch general top headlines.
    topic_match = re.search(r"news\s+(?:about|on|regarding)\s+(.+)", user_message, re.IGNORECASE)
    query = topic_match.group(1).strip(" ?.!") if topic_match else ""
    search_q = query if query else _NEWS_TOPIC_QUERIES["top"]
    items, error = fetch_google_news(search_q, max_items=6)
    if error or not items:
        return None
    lines = [f"[Live news results for: {query or 'top stories'} — use these to answer, "
              f"citing headlines/sources naturally, don't just dump a list:]"]
    for it in items:
        src = f" ({it['source']})" if it["source"] else ""
        lines.append(f"- {it['title']}{src}")
    return "\n".join(lines)


@app.route("/api/weather", methods=["POST"])
@login_required
def api_weather():
    data = request.get_json(force=True) or {}
    location_name = (data.get("location") or "").strip()
    lat = data.get("lat")
    lon = data.get("lon")
    display_name = location_name

    try:
        if lat is None or lon is None:
            if not location_name:
                return jsonify({"error": "Enter a city or use your location."}), 400
            geo = requests.get(
                "https://geocoding-api.open-meteo.com/v1/search",
                params={"name": location_name, "count": 1, "language": "en", "format": "json"},
                timeout=10,
            ).json()
            results = geo.get("results") or []
            if not results:
                return jsonify({"error": f'Could not find "{location_name}".'}), 404
            top = results[0]
            lat, lon = top["latitude"], top["longitude"]
            parts = [top.get("name")]
            if top.get("admin1") and top.get("admin1") != top.get("name"):
                parts.append(top["admin1"])
            if top.get("country"):
                parts.append(top["country"])
            display_name = ", ".join(p for p in parts if p)

        fr = requests.get(
            "https://api.open-meteo.com/v1/forecast",
            params={
                "latitude": lat, "longitude": lon,
                "current": "temperature_2m,relative_humidity_2m,apparent_temperature,weather_code,"
                           "wind_speed_10m,pressure_msl,visibility",
                "hourly": "temperature_2m,weather_code",
                "daily": "weather_code,temperature_2m_max,temperature_2m_min,uv_index_max,"
                         "sunrise,sunset",
                "timezone": "auto",
                "forecast_days": 7,
            },
            timeout=10,
        ).json()

        current = fr.get("current", {})
        hourly_raw = fr.get("hourly", {})
        daily_raw = fr.get("daily", {})

        if not display_name:
            try:
                rev = requests.get(
                    "https://geocoding-api.open-meteo.com/v1/reverse",
                    params={"latitude": lat, "longitude": lon, "language": "en", "format": "json"},
                    timeout=8,
                ).json()
                rr = (rev.get("results") or [None])[0]
                if rr:
                    display_name = ", ".join(p for p in [rr.get("name"), rr.get("country")] if p)
            except Exception:
                pass
            display_name = display_name or f"{lat:.2f}, {lon:.2f}"

        icon, condition = _wmo(current.get("weather_code"))

        aqi = None
        try:
            aq = requests.get(
                "https://air-quality-api.open-meteo.com/v1/air-quality",
                params={"latitude": lat, "longitude": lon, "current": "us_aqi"},
                timeout=8,
            ).json()
            aqi = _aqi_label((aq.get("current") or {}).get("us_aqi"))
        except Exception:
            pass

        hourly = []
        times = hourly_raw.get("time", [])
        temps = hourly_raw.get("temperature_2m", [])
        codes = hourly_raw.get("weather_code", [])
        now_iso = current.get("time", "")
        start_idx = 0
        for i, t in enumerate(times):
            if t >= now_iso:
                start_idx = i
                break
        for i in range(start_idx, min(start_idx + 8, len(times))):
            hi, _ = _wmo(codes[i] if i < len(codes) else None)
            hour_label = times[i][11:16] if len(times[i]) >= 16 else times[i]
            hourly.append({"time": hour_label, "icon": hi, "temp": round(temps[i]) if i < len(temps) else None})

        daily = []
        d_times = daily_raw.get("time", [])
        d_max = daily_raw.get("temperature_2m_max", [])
        d_min = daily_raw.get("temperature_2m_min", [])
        d_codes = daily_raw.get("weather_code", [])
        weekday_names = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
        for i, t in enumerate(d_times):
            try:
                y, m, dd = [int(x) for x in t.split("-")]
                import datetime as _dt
                wd = weekday_names[_dt.date(y, m, dd).weekday()]
            except Exception:
                wd = t
            di, _ = _wmo(d_codes[i] if i < len(d_codes) else None)
            daily.append({
                "day": "Today" if i == 0 else wd,
                "icon": di,
                "max": round(d_max[i]) if i < len(d_max) else None,
                "min": round(d_min[i]) if i < len(d_min) else None,
            })

        weather = {
            "location": display_name,
            "icon": icon,
            "condition": condition,
            "temp": round(current.get("temperature_2m", 0)),
            "feels_like": round(current.get("apparent_temperature", 0)),
            "humidity": current.get("relative_humidity_2m"),
            "wind_speed": round(current.get("wind_speed_10m", 0)),
            "pressure": round(current.get("pressure_msl")) if current.get("pressure_msl") is not None else None,
            "visibility": round((current.get("visibility") or 0) / 1000, 1) if current.get("visibility") is not None else None,
            "uv": (daily_raw.get("uv_index_max") or [None])[0],
            "sunrise": (daily_raw.get("sunrise") or [None])[0][11:16] if (daily_raw.get("sunrise") or [None])[0] else None,
            "sunset": (daily_raw.get("sunset") or [None])[0][11:16] if (daily_raw.get("sunset") or [None])[0] else None,
            "aqi": aqi,
            "hourly": hourly,
            "daily": daily,
        }
        return jsonify({"weather": weather})
    except requests.RequestException as e:
        return jsonify({"error": f"Weather service unavailable: {e}"}), 502
    except Exception as e:
        return jsonify({"error": f"Could not process weather data: {e}"}), 500


@app.route("/api/chat", methods=["POST"])
@login_required
def chat():
    data = request.get_json(force=True) or {}
    user_message = (data.get("message") or "").strip()
    conv_id = data.get("conversation_id")
    attachment = data.get("attachment")  # {name, mimeType, dataBase64} or None
    user_name = (data.get("user_name") or "").strip()[:60]  # what Mythic AI should call the user
    requested_model = (data.get("model") or DEFAULT_MODEL_ID).strip()
    regenerate = bool(data.get("regenerate"))
    ephemeral = bool(data.get("ephemeral"))
    # Optional per-person "bring your own API key" override, set in Settings.
    user_groq_key = (data.get("groq_api_key") or "").strip()
    user_cerebras_key = (data.get("cerebras_api_key") or "").strip()
    # Optional model overrides from the Model Manager (falls back to the
    # server's default model automatically if unavailable — see
    # _stream_with_model_fallback).
    groq_model_override = (data.get("groq_model") or "").strip() or None
    cerebras_model_override = (data.get("cerebras_model") or "").strip() or None
    continue_reply = bool(data.get("continue_reply"))

    if ephemeral:
        if not user_message:
            return jsonify({"error": "message is required"}), 400
        temp_messages = [{"role": "user", "parts": [{"text": user_message}]}]

        def generate_ephemeral():
            for chunk in auto_stream_chunks(None, temp_messages, SYSTEM_PROMPT,
                                             user_groq_key, user_cerebras_key):
                yield chunk.encode("utf-8")

        return Response(stream_with_context(generate_ephemeral()),
                         mimetype="text/plain")

    if regenerate:
        if not conv_id:
            return jsonify({"error": "conversation_id is required to regenerate"}), 400
    elif not user_message and not attachment:
        return jsonify({"error": "message or attachment is required"}), 400

    if attachment:
        try:
            raw = base64.b64decode(attachment.get("dataBase64", ""), validate=True)
        except Exception:
            return jsonify({"error": "invalid attachment data"}), 400
        att_mime = (attachment.get("mimeType") or "")
        att_name = (attachment.get("name") or "").lower()
        is_document = (att_mime == "application/pdf" or att_name.endswith((".pdf", ".docx", ".txt", ".md")))
        effective_limit = DOCUMENT_UPLOAD_BYTES if is_document else MAX_UPLOAD_BYTES
        if len(raw) > effective_limit:
            limit_mb = effective_limit // (1024 * 1024)
            return jsonify({"error": f"attachment too large (max {limit_mb}MB for this file type)"}), 400

    username = current_username()
    conv = load_conversation(username, conv_id) if conv_id else None
    if conv is None:
        if regenerate:
            return jsonify({"error": "conversation not found"}), 404
        conv_id = str(uuid.uuid4())
        conv = {"title": make_title(user_message), "messages": [],
                "folder": None, "pinned": False, "archived": False}

    messages = conv.setdefault("messages", [])

    if regenerate:
        if messages and messages[-1]["role"] == "model":
            messages.pop()
        if not messages or messages[-1]["role"] != "user":
            return jsonify({"error": "nothing to regenerate"}), 400
    else:
        user_parts = []
        if user_message:
            user_parts.append({"text": user_message})
        attachment_meta = None
        if attachment:
            mime_type = attachment.get("mimeType", "application/octet-stream")
            filename = attachment.get("name", "file")
            extracted_text, extract_note = extract_text_from_attachment(filename, mime_type, raw)
            if extracted_text is not None:
                doc_block = f"\n\n[Attached file: {filename}]\n{extracted_text}"
                if extract_note:
                    doc_block += f"\n[Note: {extract_note}]"
                user_parts.append({"text": doc_block})
            elif mime_type.startswith("image/"):
                # No caption needed — the raw image bytes are attached as
                # inline_data below and get routed to the vision model by
                # auto_stream_chunks(), so Mythic AI actually sees this frame.
                pass
            elif extract_note:
                user_parts.append({"text": f"\n\n[Attached file: {filename} — {extract_note}]"})
            user_parts.append({
                "inline_data": {"mime_type": mime_type, "data": attachment["dataBase64"]}
            })
            attachment_meta = {"name": filename, "mimeType": mime_type}

        user_entry = {"role": "user", "parts": user_parts}
        if attachment_meta:
            user_entry["attachment_meta"] = attachment_meta
        messages.append(user_entry)

        _update_user_activity(username)

    effective_system_prompt = SYSTEM_PROMPT
    requested_mode = (data.get("mode") or "default").strip()
    mode_addition = MODE_PROMPTS.get(requested_mode, "")
    if mode_addition:
        effective_system_prompt += " " + mode_addition
    if user_name:
        effective_system_prompt += (
            f" The user has told you their preferred name is \"{user_name}\". "
            f"Address them as {user_name} naturally where it fits (e.g. greetings, "
            f"acknowledgements) — don't force it into every single reply."
        )
    if requested_model == "mythic-vip" and session.get("vip_unlocked"):
        effective_system_prompt += (
            " The user is on the VIP tier — feel free to go deeper and be more "
            "thorough than usual when it's helpful, without padding for its own sake."
        )

    # Feature 15 — per-conversation custom system prompt/instructions, kept
    # on the conversation record itself (separate from global settings) so
    # it only affects this one chat.
    conv_custom_instructions = (conv.get("custom_instructions") or "").strip()
    if conv_custom_instructions:
        effective_system_prompt += "\n\nCustom instructions for this conversation:\n" + conv_custom_instructions

    # Features 49-50 — active persona for this conversation, if any.
    persona = get_persona_for(username, conv.get("persona_id"))
    if persona and persona.get("instructions"):
        effective_system_prompt += (
            f"\n\nYou are currently acting as the persona \"{persona['name']}\" — "
            f"follow these persona instructions: {persona['instructions']}"
        )

    # Feature 14 — long-term memory (no-op string if disabled/empty).
    effective_system_prompt += get_memory_context_block(username)


    # Live news grounding — only kicks in when the message actually looks
    # like a news request (see _NEWS_INTENT_RE). Fetches real current
    # headlines via Google News RSS and hands them to the model as context,
    # so "what's today's news" gets a real answer instead of the model
    # either refusing or guessing (it has no live web access on its own).
    if not regenerate:
        news_context = _maybe_fetch_news_context(user_message)
        if news_context:
            effective_system_prompt += (
                "\n\nThe user is asking about current news. Below are REAL, "
                "just-fetched headlines — use them to answer naturally (mention "
                "a few relevant headlines and their sources), don't just paste "
                "the raw list, and don't claim these are your own knowledge:\n\n"
                + news_context
            )

    is_first_exchange = (not regenerate) and len(messages) == 1  # just the user message so far

    def generate():
        full_reply = []
        chunk_source = auto_stream_chunks(None, messages, effective_system_prompt,
                                           user_groq_key, user_cerebras_key)

        for chunk in chunk_source:
            full_reply.append(chunk)
            yield chunk.encode("utf-8")
        reply_text = "".join(full_reply)
        messages.append({"role": "model", "parts": [{"text": reply_text}]})

        # AI-generated smart title — only for the very first exchange in a
        # conversation, and only if the person hasn't already renamed the
        # chat themselves (title_is_custom guards against overwriting that).
        if is_first_exchange and not conv.get("title_is_custom"):
            try:
                conv["title"] = generate_smart_title(
                    user_message, reply_text, user_groq_key, user_cerebras_key
                )
            except Exception as e:
                print(f"[SmartTitle] failed, keeping fallback title: {e}")

        save_conversation(username, conv_id, conv)

    resp = Response(stream_with_context(generate()), mimetype="text/plain")
    resp.headers["Content-Type"] = "text/plain; charset=utf-8"
    resp.headers["X-Conversation-Id"] = conv_id
    # Tell any intermediate proxy (nginx, Render's edge, Cloudflare, etc.)
    # NOT to buffer this response. Without this, some proxies collect the
    # entire streamed reply server-side before sending anything to the
    # browser at all — on a mobile connection that delay can exceed the
    # proxy's own idle timeout, so the connection gets killed and the
    # browser sees nothing, ever. This header is the standard fix.
    resp.headers["X-Accel-Buffering"] = "no"
    resp.headers["Cache-Control"] = "no-cache, no-transform"
    return resp


@app.route("/api/temp-image/<img_id>", methods=["GET"])
def serve_temp_image(img_id):
    entry = _TEMP_IMAGES.get(img_id)
    if not entry:
        return jsonify({"error": "not found or expired"}), 404
    return Response(entry["data"], mimetype=entry["mime_type"])


@app.route("/api/generate-file", methods=["POST"])
@login_required
def generate_file():
    """Generates a downloadable file (PDF, Word doc, or plain text) from
    text content — used for "generate a PDF / document / downloadable
    file" requests. Returns base64 file bytes + filename + mime type."""
    try:
        data = request.get_json(force=True) or {}
        content = (data.get("content") or "").strip()
        fmt = (data.get("format") or "pdf").strip().lower()
        title = (data.get("title") or "Mythic AI Document").strip()[:100]
        filename_base = "".join(
            c for c in title if c.isalnum() or c in " -_"
        ).strip().replace(" ", "-") or "Mythic-AI-Document"

        if not content:
            return jsonify({"error": "content is required"}), 400

        if fmt in ("docx", "word", "doc"):
            docx_bytes = generate_docx_bytes(title, content)
            if docx_bytes is not None:
                return jsonify({
                    "file": base64.b64encode(docx_bytes).decode("utf-8"),
                    "filename": f"{filename_base}.docx",
                    "mimeType": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                })
            return jsonify({
                "file": base64.b64encode(content.encode("utf-8")).decode("utf-8"),
                "filename": f"{filename_base}.txt",
                "mimeType": "text/plain",
                "note": "Word (.docx) generation isn't set up on this server "
                        "(needs `pip install python-docx`) — sent as a plain "
                        "text file instead.",
            })

        if fmt in ("txt", "text"):
            return jsonify({
                "file": base64.b64encode(content.encode("utf-8")).decode("utf-8"),
                "filename": f"{filename_base}.txt",
                "mimeType": "text/plain",
            })

        pdf_bytes = generate_pdf_bytes(title, content)
        return jsonify({
            "file": base64.b64encode(pdf_bytes).decode("utf-8"),
            "filename": f"{filename_base}.pdf",
            "mimeType": "application/pdf",
        })
    except Exception as e:
        return jsonify({"error": f"File generation failed: {e}"}), 500


@app.route("/api/generate-image", methods=["POST"])
@login_required
def generate_image():
    try:
        data = request.get_json(force=True) or {}
        prompt = data.get("prompt", "").strip()
        image_b64 = data.get("imageBase64")
        mime_type = data.get("mimeType", "image/jpeg")
        style = data.get("style", "").strip()
        # Per-request watermark settings (from Settings → Image Watermark in
        # the frontend, or sane defaults if not sent). Every branch below
        # routes its result through _finalize_generated_image() so the mark
        # can never be skipped for any current or future image provider.
        watermark_opts = data.get("watermark") or {}
        if not prompt:
            return jsonify({"error": "prompt required"}), 400

        full_prompt = f"{prompt}, {style}" if style else prompt

        prompt_lower = full_prompt.lower()
        is_book      = any(w in prompt_lower for w in ["book cover", "book", "novel cover", "textbook"])
        is_portrait  = any(w in prompt_lower for w in ["portrait", "person", "face", "selfie", "photo of"])
        is_logo      = any(w in prompt_lower for w in ["logo", "icon", "emblem", "badge"])
        is_anime     = any(w in prompt_lower for w in ["anime", "ghibli", "manga", "cartoon", "illustration"])
        is_landscape = any(w in prompt_lower for w in ["landscape", "scenery", "nature", "city", "skyline", "aerial"])

        quality_tail = ", masterpiece, best quality, highly detailed, sharp focus, professional"

        if is_book:
            enhanced = (
                f"{full_prompt}, professional book cover design, elegant typography layout, "
                f"dramatic lighting, rich colors, visually striking, award-winning cover art, "
                f"publishing industry standard{quality_tail}"
            )
            width, height = 512, 768
        elif is_portrait:
            enhanced = (
                f"{full_prompt}, cinematic portrait photography, soft studio lighting, "
                f"8K ultra-detailed, DSLR quality, photorealistic{quality_tail}"
            )
            width, height = 512, 768
        elif is_logo:
            enhanced = (
                f"{full_prompt}, clean vector style, minimalist, professional brand design, "
                f"flat design, scalable, high contrast{quality_tail}"
            )
            width, height = 768, 768
        elif is_anime:
            enhanced = (
                f"{full_prompt}, vibrant anime art, clean linework, beautiful coloring, "
                f"Studio Ghibli quality, cel shading{quality_tail}"
            )
            width, height = 768, 768
        elif is_landscape:
            enhanced = (
                f"{full_prompt}, epic wide shot, golden hour lighting, ultra-wide, "
                f"breathtaking scenery, National Geographic quality{quality_tail}"
            )
            width, height = 896, 512
        else:
            enhanced = f"{full_prompt}{quality_tail}"
            width, height = 768, 768

        negative = (
            "blurry, low quality, pixelated, distorted, deformed, ugly, bad anatomy, "
            "watermark, signature, text errors, garbled text, poorly drawn, disfigured, "
            "oversaturated, washed out, extra limbs, duplicate, clone, artifact, noise"
        )

        if NANO_BANANA_API_KEY:
            image_urls = None
            if image_b64:
                try:
                    raw = base64.b64decode(image_b64, validate=True)
                    if len(raw) > MAX_UPLOAD_BYTES:
                        return jsonify({"error": "image too large (max 8MB)"}), 400
                    img_id = _store_temp_image(raw, mime_type)
                    base_url = get_public_origin()
                    image_urls = [f"{base_url}/api/temp-image/{img_id}"]
                except Exception:
                    pass
            try:
                task_id, err = nano_banana_submit(enhanced, image_urls=image_urls)
                if not err:
                    result_url, err = nano_banana_poll(task_id)
                    if not err:
                        img_resp = requests.get(result_url, timeout=30)
                        img_resp.raise_for_status()
                        return jsonify({"image": _finalize_generated_image(img_resp.content, watermark_opts)})
            except Exception:
                pass

        if HF_API_KEY:
            try:
                resp = requests.post(
                    "https://api-inference.huggingface.co/models/black-forest-labs/FLUX.1-schnell",
                    headers={"Authorization": f"Bearer {HF_API_KEY}"},
                    json={"inputs": enhanced},
                    timeout=90,
                )
                if resp.status_code == 200 and resp.headers.get("content-type", "").startswith("image/"):
                    return jsonify({"image": _finalize_generated_image(resp.content, watermark_opts)})
            except Exception:
                pass

        seed = int(time.time()) % 99999
        encoded_prompt   = urllib.parse.quote(enhanced)
        encoded_negative = urllib.parse.quote(negative)
        poll_url = (
            f"https://image.pollinations.ai/prompt/{encoded_prompt}"
            f"?width={width}&height={height}&seed={seed}"
            f"&negative={encoded_negative}"
            f"&model=flux&enhance=true&nologo=true"
        )
        img_resp = requests.get(poll_url, timeout=120)
        if img_resp.status_code == 200 and img_resp.headers.get("content-type", "").startswith("image/"):
            return jsonify({"image": _finalize_generated_image(img_resp.content, watermark_opts)})
        return jsonify({"error": f"Image generation failed (status {img_resp.status_code}). Please try again."}), 502

    except Exception as e:
        # Log the real detail server-side only — raw exception text can
        # contain the underlying provider's URL/hostname, which we don't
        # want to expose to the browser.
        print(f"[image-generation] error: {e}")
        return jsonify({"error": "Image generation failed. Please try again."}), 500


# ══════════════════════════════════════════════════════════════════════════
# NEW FEATURE: real, sandboxed code execution (Python/JS/etc.) for the Code
# tab and for "Run" on non-HTML Artifacts.
#
# Execution happens on Piston (emkc.org's free public code-execution API),
# NOT on this server. That matters: this Render instance holds real secrets
# (GROQ/CEREBRAS/SUPABASE keys) in its environment, so running arbitrary,
# untrusted code IN THIS PROCESS would let anyone with the invite link read
# those secrets or pivot into this box. Piston runs the code in its own
# disposable sandbox instead, so a malicious snippet can at worst mess with
# a throwaway container that has none of this app's secrets.
# Still gated behind the existing VIP session flag (server-side, not just a
# client-side lock icon) — it's real third-party compute, not something to
# leave wide open to every anonymous visitor of the public invite link.
# ══════════════════════════════════════════════════════════════════════════
PISTON_API = "https://emkc.org/api/v2/piston/execute"
PISTON_RUNTIMES_API = "https://emkc.org/api/v2/piston/runtimes"
_piston_runtimes_cache = {"ts": 0.0, "data": []}
_PISTON_LANGUAGE_ALIASES = {
    "py": "python", "python3": "python", "js": "javascript", "node": "javascript",
    "nodejs": "javascript", "ts": "typescript", "sh": "bash", "shell": "bash",
    "c++": "cpp", "golang": "go",
}
_PISTON_FILENAMES = {
    "python": "main.py", "javascript": "main.js", "typescript": "main.ts",
    "bash": "main.sh", "c": "main.c", "cpp": "main.cpp", "java": "Main.java",
    "go": "main.go", "rust": "main.rs", "ruby": "main.rb", "php": "main.php",
}

def _piston_runtimes():
    now = time.time()
    if _piston_runtimes_cache["data"] and (now - _piston_runtimes_cache["ts"]) < 3600:
        return _piston_runtimes_cache["data"]
    try:
        r = requests.get(PISTON_RUNTIMES_API, timeout=10)
        if r.status_code == 200:
            _piston_runtimes_cache["data"] = r.json()
            _piston_runtimes_cache["ts"] = now
    except Exception as e:
        print(f"[Piston] failed to fetch runtimes: {e}")
    return _piston_runtimes_cache["data"]


@app.route("/api/execute-code", methods=["POST"])
@login_required
def api_execute_code():
    if not session.get("vip_unlocked"):
        return jsonify({"error": "Unlock VIP mode first to run code."}), 403
    data = request.get_json(force=True) or {}
    code = (data.get("code") or "")
    if not code.strip():
        return jsonify({"error": "No code provided."}), 400
    if len(code) > 20000:
        return jsonify({"error": "Code is too long (max 20,000 characters)."}), 400
    lang_in = (data.get("language") or "python").strip().lower()
    lang = _PISTON_LANGUAGE_ALIASES.get(lang_in, lang_in)

    runtimes = _piston_runtimes()
    version = "*"
    if runtimes:
        match = next((r for r in runtimes if r.get("language") == lang
                       or lang in (r.get("aliases") or [])), None)
        if not match:
            supported = sorted({r.get("language") for r in runtimes})
            return jsonify({"error": f"'{lang_in}' isn't a supported language.",
                             "supported": supported[:30]}), 400
        version = match.get("version", "*")

    filename = _PISTON_FILENAMES.get(lang, "main.txt")
    try:
        resp = requests.post(PISTON_API, json={
            "language": lang,
            "version": version,
            "files": [{"name": filename, "content": code}],
            "stdin": (data.get("stdin") or "")[:5000],
            "run_timeout": 8000,
            "compile_timeout": 10000,
        }, timeout=25)
        if resp.status_code != 200:
            return jsonify({"error": f"Execution service returned HTTP {resp.status_code}."}), 502
        result = resp.json()
        run = result.get("run") or {}
        compiled = result.get("compile") or {}
        return jsonify({
            "stdout": run.get("stdout", ""),
            "stderr": run.get("stderr", "") or compiled.get("stderr", ""),
            "exit_code": run.get("code"),
            "language": lang,
            "version": version,
        })
    except requests.exceptions.Timeout:
        return jsonify({"error": "Execution timed out."}), 504
    except Exception as e:
        return jsonify({"error": f"Execution failed: {e}"}), 500


# ══════════════════════════════════════════════════════════════════════════
# NEW FEATURE: Cowork mode — real multi-step task execution, not just a
# different system prompt. Breaks the task into steps, works through each
# one (carrying forward what earlier steps produced), then synthesizes a
# single final answer. Uses the existing _quick_completion() Groq→Cerebras
# fallback helper, so it inherits the same silent-provider-fallback behavior
# as the rest of the app.
# ══════════════════════════════════════════════════════════════════════════
@app.route("/api/cowork/run", methods=["POST"])
@login_required
def api_cowork_run():
    if not session.get("vip_unlocked"):
        return jsonify({"error": "Unlock VIP mode first to use Cowork."}), 403
    data = request.get_json(force=True) or {}
    task = (data.get("task") or "").strip()
    if not task:
        return jsonify({"error": "No task provided."}), 400
    task = task[:4000]
    user_groq = (data.get("groqKey") or "").strip() or None
    user_cerebras = (data.get("cerebrasKey") or "").strip() or None

    plan_text = _quick_completion(
        [
            {"role": "system", "content": (
                "You are a meticulous project planner. Break the user's task into 3 to 6 "
                "short, concrete, sequential steps a capable assistant could execute one "
                "at a time. Reply with ONLY a numbered list, one short step per line, "
                "nothing else — no preamble, no explanation."
            )},
            {"role": "user", "content": task},
        ],
        user_groq, user_cerebras, max_tokens=300,
    ) or ""
    steps = []
    for line in plan_text.splitlines():
        line = re.sub(r"^\s*[\d]+[\.\)]\s*", "", line).strip("-• \t")
        if line:
            steps.append(line)
    steps = steps[:6] or [task]

    step_results = []
    running_context = ""
    for i, step in enumerate(steps, 1):
        step_prompt = (
            f"Overall task: {task}\n\n"
            f"Progress so far:\n{running_context[-3000:] or '(nothing yet)'}\n\n"
            f"Do ONLY this step ({i} of {len(steps)}): {step}\n"
            "Give a direct, concrete result for this step alone — don't restate the "
            "whole task or repeat earlier steps."
        )
        result = _quick_completion(
            [
                {"role": "system", "content": "You are a focused task executor. Do exactly "
                                               "the step given, concisely and concretely."},
                {"role": "user", "content": step_prompt},
            ],
            user_groq, user_cerebras, max_tokens=600,
        ) or "(no result — provider unavailable)"
        step_results.append({"step": step, "result": result})
        running_context += f"\nStep {i} — {step}:\n{result}\n"

    final_answer = _quick_completion(
        [
            {"role": "system", "content": "You synthesize completed work into one final, "
                                           "coherent answer or deliverable."},
            {"role": "user", "content": (
                f"Task: {task}\n\nStep-by-step work completed:\n{running_context[-6000:]}\n\n"
                "Write the final, complete answer for the task, weaving the step results "
                "together into one coherent response. Don't just list the steps again."
            )},
        ],
        user_groq, user_cerebras, max_tokens=1400,
    ) or running_context

    return jsonify({"task": task, "steps": step_results, "final_answer": final_answer})


# ══════════════════════════════════════════════════════════════════════════
# NEW FEATURE: search across every one of this visitor's conversations
# (not just the currently open one).
# ══════════════════════════════════════════════════════════════════════════
@app.route("/api/search", methods=["GET"])
@login_required
def api_search_messages():
    q = (request.args.get("q") or "").strip()
    if len(q) < 2:
        return jsonify({"results": [], "query": q})
    q_lower = q.lower()
    username = current_username()
    results = []
    for c in list_conversations(username):
        conv = load_conversation(username, c["id"])
        if not conv:
            continue
        for idx, m in enumerate(conv.get("messages", []) or []):
            text = m.get("text") or ""
            pos = text.lower().find(q_lower)
            if pos == -1:
                continue
            start = max(0, pos - 40)
            snippet = text[start:start + 160].strip()
            results.append({
                "conv_id": c["id"],
                "title": c.get("title") or "Untitled",
                "role": m.get("role"),
                "msg_index": idx,
                "snippet": ("…" if start > 0 else "") + snippet + ("…" if start + 160 < len(text) else ""),
            })
            if len(results) >= 50:
                break
        if len(results) >= 50:
            break
    return jsonify({"results": results, "query": q})


# ══════════════════════════════════════════════════════════════════════════
# NEW FEATURE: reminders / scheduled tasks. Reuses the existing Web Push
# infrastructure (send_push_notification_to_user, built for re-engagement
# notifications) to actually deliver the reminder at the right time.
# ══════════════════════════════════════════════════════════════════════════
_REMINDERS_FILE = _os.path.join(_DATA_DIR, "reminders.json")
_reminders_lock = threading.Lock()

def _load_reminders():
    with _reminders_lock:
        if _os.path.exists(_REMINDERS_FILE):
            try:
                with open(_REMINDERS_FILE, encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {}

def _save_reminders(data):
    with _reminders_lock:
        try:
            with open(_REMINDERS_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f)
        except Exception as e:
            print(f"[reminders] failed to save: {e}")


# ---------------------------------------------------------------------------
# Phase 2: Long-term memory (feature 14). Same file-backed pattern as
# reminders above — one JSON file, keyed by a random id, each row tagged
# with the owning username. A user's "memory enabled" toggle lives in the
# same file under a special key so it survives without a separate store.
# ---------------------------------------------------------------------------
_MEMORY_FILE = _os.path.join(_DATA_DIR, "memories.json")
_memory_lock = threading.Lock()

def _load_memory_store_file():
    with _memory_lock:
        if _os.path.exists(_MEMORY_FILE):
            try:
                with open(_MEMORY_FILE, encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {"entries": {}, "settings": {}}

def _save_memory_store_file(store):
    with _memory_lock:
        try:
            with open(_MEMORY_FILE, "w", encoding="utf-8") as f:
                json.dump(store, f)
        except Exception as e:
            print(f"[memory] failed to save: {e}")


# --- Memory: Supabase-backed when configured, local JSON file otherwise -----
# Tables needed on Supabase:
#   create table memories (
#     id text primary key, username text not null, text text not null,
#     created_at double precision, updated_at double precision
#   );
#   create table memory_settings (
#     username text primary key, enabled boolean default true
#   );

def memory_enabled_for(username):
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"memory_settings?username=eq.{username}"), headers=sb_headers(), timeout=10)
            if r.status_code == 200 and r.json():
                return bool(r.json()[0].get("enabled", True))
        except Exception as e:
            print(f"[Supabase] memory_enabled_for failed: {e}")
        return True  # default on if no row / lookup failed
    store = _load_memory_store_file()
    return store.get("settings", {}).get(username, {}).get("enabled", True)


def _list_memories(username):
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"memories?username=eq.{username}&order=created_at.asc"),
                              headers=sb_headers(), timeout=10)
            if r.status_code == 200:
                return r.json()
        except Exception as e:
            print(f"[Supabase] list memories failed: {e}")
        return []
    store = _load_memory_store_file()
    return [m for m in store.get("entries", {}).values() if m.get("username") == username]


def get_memory_context_block(username, max_chars=1500):
    """Builds the block of remembered facts to splice into the system
    prompt for this user's next chat request. Returns "" if memory is
    disabled or there's nothing saved yet, so callers can unconditionally
    append the result without an extra branch."""
    if not memory_enabled_for(username):
        return ""
    mine = _list_memories(username)
    if not mine:
        return ""
    mine.sort(key=lambda m: m.get("created_at", 0))
    lines = [f"- {m['text']}" for m in mine]
    block = "\n".join(lines)
    if len(block) > max_chars:
        block = block[:max_chars] + "…"
    return (
        "\n\nThe user has asked you to remember these facts/preferences about "
        "them from past conversations — use them naturally where relevant, "
        "don't just recite the list back:\n" + block
    )


@app.route("/api/memory", methods=["GET"])
@login_required
def api_list_memory():
    """Optional ?q=... does a simple case-insensitive substring search."""
    username = current_username()
    q = (request.args.get("q") or "").strip().lower()
    mine = _list_memories(username)
    if q:
        mine = [m for m in mine if q in m["text"].lower()]
    mine.sort(key=lambda m: m.get("created_at", 0), reverse=True)
    return jsonify({"memories": mine, "enabled": memory_enabled_for(username)})


@app.route("/api/memory", methods=["POST"])
@login_required
def api_add_memory():
    data = request.get_json(force=True) or {}
    text = (data.get("text") or "").strip()[:500]
    if not text:
        return jsonify({"error": "'text' is required"}), 400
    # Never silently save anything that looks like a password, card number,
    # or other secret — memory is meant for preferences, not credentials.
    lowered = text.lower()
    if any(w in lowered for w in ("password", "credit card", "card number", "cvv", "ssn", "social security")):
        return jsonify({"error": "This looks like sensitive/credential information — "
                                  "Mythic AI won't store that in long-term memory."}), 400
    mid = uuid.uuid4().hex[:12]
    username = current_username()
    entry = {"id": mid, "username": username, "text": text, "created_at": time.time()}
    if SUPABASE_URL:
        try:
            r = requests.post(sb("memories"), headers=sb_headers(), json=entry, timeout=10)
            if r.status_code not in (200, 201):
                print(f"[Supabase] add_memory failed: HTTP {r.status_code} — {r.text[:300]}")
        except Exception as e:
            print(f"[Supabase] add_memory exception: {e}")
    else:
        store = _load_memory_store_file()
        store.setdefault("entries", {})[mid] = entry
        _save_memory_store_file(store)
    return jsonify({"memory": entry})


@app.route("/api/memory/<mid>", methods=["PUT"])
@login_required
def api_edit_memory(mid):
    data = request.get_json(force=True) or {}
    text = (data.get("text") or "").strip()[:500]
    if not text:
        return jsonify({"error": "'text' is required"}), 400
    username = current_username()
    if SUPABASE_URL:
        try:
            r = requests.patch(sb(f"memories?id=eq.{mid}&username=eq.{username}"), headers=sb_headers(),
                                json={"text": text, "updated_at": time.time()}, timeout=10)
            if r.status_code not in (200, 204) or (r.status_code == 200 and not r.json()):
                return jsonify({"error": "not found"}), 404
        except Exception as e:
            print(f"[Supabase] edit_memory exception: {e}")
            return jsonify({"error": "storage error"}), 500
        return jsonify({"memory": {"id": mid, "username": username, "text": text}})
    store = _load_memory_store_file()
    entry = store.get("entries", {}).get(mid)
    if not entry or entry.get("username") != username:
        return jsonify({"error": "not found"}), 404
    entry["text"] = text
    entry["updated_at"] = time.time()
    _save_memory_store_file(store)
    return jsonify({"memory": entry})


@app.route("/api/memory/<mid>", methods=["DELETE"])
@login_required
def api_delete_memory(mid):
    username = current_username()
    if SUPABASE_URL:
        try:
            requests.delete(sb(f"memories?id=eq.{mid}&username=eq.{username}"), headers=sb_headers(), timeout=10)
        except Exception as e:
            print(f"[Supabase] delete_memory exception: {e}")
        return jsonify({"status": "deleted"})
    store = _load_memory_store_file()
    entry = store.get("entries", {}).get(mid)
    if not entry or entry.get("username") != username:
        return jsonify({"error": "not found"}), 404
    del store["entries"][mid]
    _save_memory_store_file(store)
    return jsonify({"status": "deleted"})


@app.route("/api/memory/settings", methods=["POST"])
@login_required
def api_set_memory_enabled():
    """Body: {"enabled": true|false} — global on/off switch. When off,
    get_memory_context_block() returns "" so /api/chat stops injecting any
    saved memories, without deleting them (an easy re-enable later)."""
    data = request.get_json(force=True) or {}
    username = current_username()
    enabled = bool(data.get("enabled", True))
    if SUPABASE_URL:
        try:
            headers = {**sb_headers(), "Prefer": "resolution=merge-duplicates,return=minimal"}
            requests.post(sb("memory_settings"), headers=headers,
                          json={"username": username, "enabled": enabled}, timeout=10)
        except Exception as e:
            print(f"[Supabase] set_memory_enabled exception: {e}")
    else:
        store = _load_memory_store_file()
        store.setdefault("settings", {})[username] = {"enabled": enabled}
        _save_memory_store_file(store)
    return jsonify({"enabled": enabled})


# ---------------------------------------------------------------------------
# Phase 2: Personas (features 49-50). A saved persona is a name + avatar
# emoji + instructions the user can attach to any conversation; selecting
# one just appends its instructions to that conversation's system prompt.
# Supabase-backed when configured, local JSON file otherwise. Table:
#   create table personas (
#     id text primary key, username text not null, name text not null,
#     avatar text, description text, instructions text, created_at double precision
#   );
# ---------------------------------------------------------------------------
_PERSONAS_FILE = _os.path.join(_DATA_DIR, "personas.json")
_personas_lock = threading.Lock()

def _load_personas_store_file():
    with _personas_lock:
        if _os.path.exists(_PERSONAS_FILE):
            try:
                with open(_PERSONAS_FILE, encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {}

def _save_personas_store_file(data):
    with _personas_lock:
        try:
            with open(_PERSONAS_FILE, "w", encoding="utf-8") as f:
                json.dump(data, f)
        except Exception as e:
            print(f"[personas] failed to save: {e}")


def _list_personas(username):
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"personas?username=eq.{username}&order=created_at.asc"),
                              headers=sb_headers(), timeout=10)
            if r.status_code == 200:
                return r.json()
        except Exception as e:
            print(f"[Supabase] list personas failed: {e}")
        return []
    all_p = _load_personas_store_file()
    return [p for p in all_p.values() if p.get("username") == username]


def get_persona_for(username, persona_id):
    if not persona_id:
        return None
    if SUPABASE_URL:
        try:
            r = requests.get(sb(f"personas?id=eq.{persona_id}&username=eq.{username}"),
                              headers=sb_headers(), timeout=10)
            if r.status_code == 200 and r.json():
                return r.json()[0]
        except Exception as e:
            print(f"[Supabase] get_persona_for failed: {e}")
        return None
    all_p = _load_personas_store_file()
    p = all_p.get(persona_id)
    return p if p and p.get("username") == username else None


@app.route("/api/personas", methods=["GET"])
@login_required
def api_list_personas():
    username = current_username()
    mine = _list_personas(username)
    mine.sort(key=lambda p: p.get("created_at", 0))
    return jsonify({"personas": mine})


@app.route("/api/personas", methods=["POST"])
@login_required
def api_create_persona():
    data = request.get_json(force=True) or {}
    name = (data.get("name") or "").strip()[:60]
    if not name:
        return jsonify({"error": "'name' is required"}), 400
    pid = uuid.uuid4().hex[:12]
    username = current_username()
    persona = {
        "id": pid, "username": username, "name": name,
        "avatar": (data.get("avatar") or "🤖")[:8],
        "description": (data.get("description") or "").strip()[:200],
        "instructions": (data.get("instructions") or "").strip()[:2000],
        "created_at": time.time(),
    }
    if SUPABASE_URL:
        try:
            r = requests.post(sb("personas"), headers=sb_headers(), json=persona, timeout=10)
            if r.status_code not in (200, 201):
                print(f"[Supabase] create_persona failed: HTTP {r.status_code} — {r.text[:300]}")
        except Exception as e:
            print(f"[Supabase] create_persona exception: {e}")
    else:
        all_p = _load_personas_store_file()
        all_p[pid] = persona
        _save_personas_store_file(all_p)
    return jsonify({"persona": persona})


@app.route("/api/personas/<pid>", methods=["PUT"])
@login_required
def api_update_persona(pid):
    data = request.get_json(force=True) or {}
    username = current_username()
    fields = {}
    for field, cap in (("name", 60), ("avatar", 8), ("description", 200), ("instructions", 2000)):
        if field in data:
            fields[field] = (data.get(field) or "").strip()[:cap]
    if not fields:
        return jsonify({"error": "no recognized fields"}), 400

    if SUPABASE_URL:
        try:
            r = requests.patch(sb(f"personas?id=eq.{pid}&username=eq.{username}"),
                                headers=sb_headers(), json=fields, timeout=10)
            if r.status_code not in (200, 204):
                return jsonify({"error": "not found"}), 404
        except Exception as e:
            print(f"[Supabase] update_persona exception: {e}")
            return jsonify({"error": "storage error"}), 500
        return jsonify({"persona": {"id": pid, "username": username, **fields}})

    all_p = _load_personas_store_file()
    p = all_p.get(pid)
    if not p or p.get("username") != username:
        return jsonify({"error": "not found"}), 404
    p.update(fields)
    _save_personas_store_file(all_p)
    return jsonify({"persona": p})


@app.route("/api/personas/<pid>", methods=["DELETE"])
@login_required
def api_delete_persona(pid):
    username = current_username()
    if SUPABASE_URL:
        try:
            requests.delete(sb(f"personas?id=eq.{pid}&username=eq.{username}"), headers=sb_headers(), timeout=10)
        except Exception as e:
            print(f"[Supabase] delete_persona exception: {e}")
        return jsonify({"status": "deleted"})
    all_p = _load_personas_store_file()
    p = all_p.get(pid)
    if not p or p.get("username") != username:
        return jsonify({"error": "not found"}), 404
    del all_p[pid]
    _save_personas_store_file(all_p)
    return jsonify({"status": "deleted"})


@app.route("/api/reminders", methods=["GET"])
@login_required
def api_list_reminders():
    username = current_username()
    all_r = _load_reminders()
    mine = [r for r in all_r.values() if r.get("username") == username and not r.get("fired")]
    mine.sort(key=lambda r: r.get("fire_at", 0))
    return jsonify({"reminders": mine})


@app.route("/api/reminders", methods=["POST"])
@login_required
def api_create_reminder():
    data = request.get_json(force=True) or {}
    text = (data.get("text") or "").strip()[:200]
    fire_at = data.get("fire_at")
    if not text or fire_at is None:
        return jsonify({"error": "'text' and 'fire_at' (unix timestamp, seconds) are required."}), 400
    try:
        fire_at = float(fire_at)
    except (TypeError, ValueError):
        return jsonify({"error": "'fire_at' must be a numeric unix timestamp."}), 400
    if fire_at < time.time() - 60:
        return jsonify({"error": "That time is in the past."}), 400
    rid = uuid.uuid4().hex[:12]
    all_r = _load_reminders()
    all_r[rid] = {
        "id": rid, "username": current_username(), "text": text,
        "fire_at": fire_at, "created_at": time.time(), "fired": False,
    }
    _save_reminders(all_r)
    return jsonify({"reminder": all_r[rid]})


@app.route("/api/reminders/<rid>", methods=["DELETE"])
@login_required
def api_delete_reminder(rid):
    all_r = _load_reminders()
    if rid in all_r and all_r[rid].get("username") == current_username():
        del all_r[rid]
        _save_reminders(all_r)
    return jsonify({"ok": True})


def _reminders_check_loop():
    """Every 30s, fires any reminder whose time has come, via Web Push.
    Render/always-on only, same limitation as the re-engagement loop below —
    Vercel's serverless functions freeze between requests so a background
    thread there would never actually get to run on schedule."""
    while True:
        try:
            all_r = _load_reminders()
            now = time.time()
            changed = False
            for rid, r in list(all_r.items()):
                if not r.get("fired") and r.get("fire_at", 0) <= now:
                    send_push_notification_to_user(
                        r.get("username", ""), "⏰ Reminder", r.get("text", ""), url="/",
                    )
                    r["fired"] = True
                    changed = True
            if changed:
                _save_reminders(all_r)
        except Exception as e:
            print(f"[reminders] loop error: {e}")
        time.sleep(30)


_reminders_thread_started = False

def _start_reminders_thread_once():
    global _reminders_thread_started
    if _reminders_thread_started or IS_SERVERLESS:
        return
    threading.Thread(target=_reminders_check_loop, daemon=True).start()
    _reminders_thread_started = True

_start_reminders_thread_once()


# ══════════════════════════════════════════════════════════════════════════
# NEW FEATURE: export ALL conversations as one downloadable .zip backup, and
# re-import them later (e.g. after clearing cookies, or on another device/
# browser — since accounts are anonymous per-browser-cookie, this is also
# the only way to move your chats to a different browser at all).
# ══════════════════════════════════════════════════════════════════════════
@app.route("/api/backup/export", methods=["GET"])
@login_required
def api_backup_export():
    import io as _io, zipfile as _zipfile
    username = current_username()
    buf = _io.BytesIO()
    manifest = []
    with _zipfile.ZipFile(buf, "w", _zipfile.ZIP_DEFLATED) as zf:
        for c in list_conversations(username):
            conv = load_conversation(username, c["id"])
            if not conv:
                continue
            payload = dict(conv)
            payload["id"] = c["id"]
            zf.writestr(f"conversations/{c['id']}.json",
                        json.dumps(payload, ensure_ascii=False))
            manifest.append({"id": c["id"], "title": c.get("title")})
        zf.writestr("manifest.json", json.dumps({
            "exported_at": time.time(), "count": len(manifest), "conversations": manifest,
        }, ensure_ascii=False))
    buf.seek(0)
    return Response(buf.read(), mimetype="application/zip", headers={
        "Content-Disposition": f'attachment; filename="mythic-ai-backup-{int(time.time())}.zip"',
    })


@app.route("/api/backup/import", methods=["POST"])
@login_required
def api_backup_import():
    import io as _io, zipfile as _zipfile
    username = current_username()
    f = request.files.get("file")
    if not f:
        return jsonify({"error": "No file uploaded."}), 400
    raw = f.read()
    if len(raw) > 25 * 1024 * 1024:
        return jsonify({"error": "Backup file too large (max 25MB)."}), 400
    try:
        zf = _zipfile.ZipFile(_io.BytesIO(raw))
    except Exception as e:
        return jsonify({"error": f"Not a valid backup file: {e}"}), 400
    imported = 0
    for name in zf.namelist():
        if not name.startswith("conversations/") or not name.endswith(".json"):
            continue
        try:
            payload = json.loads(zf.read(name).decode("utf-8"))
        except Exception:
            continue
        # Always import under a brand-new id, even if the backup file still
        # has its original id — prevents an import from silently clobbering
        # an existing chat that happens to share the same id.
        payload.pop("id", None)
        new_id = uuid.uuid4().hex[:12]
        save_conversation(username, new_id, payload)
        imported += 1
    return jsonify({"imported": imported})


_reengagement_thread_started = False

def _start_reengagement_thread_once():
    """Starts the hourly notification background thread exactly once.
    Render/always-on only — never starts on Vercel (use the cron route)."""
    global _reengagement_thread_started
    if _reengagement_thread_started or IS_SERVERLESS:
        return
    threading.Thread(target=_reengagement_loop, daemon=True).start()
    _reengagement_thread_started = True


_start_reengagement_thread_once()

if __name__ == "__main__":
    active = []
    if PROVIDER in ("auto", "groq") and GROQ_API_KEY:
        active.append(f"Groq({GROQ_MODEL})")
    if PROVIDER in ("auto", "cerebras") and CEREBRAS_API_KEY:
        active.append(f"Cerebras({CEREBRAS_MODEL})")
    providers_str = " → ".join(active) if active else "none configured! (users can still supply their own key in Settings)"
    image_provider = "NanoBanana (image-to-image supported)" if NANO_BANANA_API_KEY else (
        "HuggingFace FLUX (text-to-image only)" if HF_API_KEY else "Pollinations (text-to-image, no key needed)"
    )
    print(f"Starting Mythic AI at http://localhost:5000")
    print(f"Providers (Groq primary, Cerebras fallback): {providers_str}")
    print(f"Image generation: {image_provider}")
    print(f"Re-engagement notifications: Render/always-on -> hourly background "
          f"thread. Vercel -> daily at 12:00 via /api/cron/reengagement (external cron).")
    if IS_SERVERLESS:
        print("NOTE: detected a serverless environment (Vercel) — see the "
              "IS_SERVERLESS comment near the top of this file for the "
              "limitations that come with running this app there.")

    app.run(host="0.0.0.0", port=5000, debug=False, threaded=True)
